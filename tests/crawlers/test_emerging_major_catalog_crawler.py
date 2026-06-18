import json
import sys
from io import BytesIO
from pathlib import Path

from docx import Document
from openpyxl import Workbook

ROOT = Path(__file__).resolve().parents[2]
SRC_ROOT = ROOT / "src"
if str(SRC_ROOT) not in sys.path:
    sys.path.insert(0, str(SRC_ROOT))

from major_intel.crawlers.emerging_major_catalog_crawler import (
    _normalize_ocr_major_code,
    _parse_catalog_ocr_words,
    _parse_filing_ocr_words,
    crawl_emerging_major_catalog,
    discover_attachments_from_html,
    load_source_rows,
    parse_docx_major_candidates,
    parse_pdf_major_candidates_from_text,
    parse_xlsx_major_candidates,
)


def _xlsx_bytes() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "新增专业"
    sheet.append(["序号", "学校名称", "专业代码", "专业名称", "门类", "专业类", "学位授予门类", "修业年限"])
    sheet.append([1, "示例大学", "080717T", "人工智能", "工学", "计算机类", "工学", "四年"])
    sheet.append([2, "示例大学", "082012T", "低空技术与工程", "工学", "航空航天类", "工学", "四年"])
    buffer = BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def _source_row() -> dict[str, str]:
    return {
        "source_id": "moe_2024",
        "title": "教育部关于公布2024年度普通高等学校本科专业备案和审批结果的通知",
        "url": "http://www.moe.gov.cn/source.html",
        "source_domain": "www.moe.gov.cn",
        "source_level": "A",
        "source_type": "annual_filing_approval",
        "issuing_org": "教育部",
        "published_date": "2025-04-22",
        "source_year": "2024",
        "notes": "seed",
    }


def _word(text: str, x: float, y: float, w: float = 20, h: float = 20) -> dict[str, float | str]:
    return {"text": text, "x": x, "y": y, "w": w, "h": h}


def test_load_source_rows_requires_seed_columns(tmp_path):
    seed = tmp_path / "sources.csv"
    seed.write_text(
        "source_id,title,url,source_domain,source_level,source_type,issuing_org,published_date,source_year,notes\n"
        "moe_2024,教育部通知,http://www.moe.gov.cn/source.html,www.moe.gov.cn,A,annual_filing_approval,教育部,2025-04-22,2024,seed\n",
        encoding="utf-8",
    )

    rows = load_source_rows(seed)

    assert rows[0]["source_id"] == "moe_2024"
    assert rows[0]["source_year"] == "2024"


def test_discover_attachments_from_html_resolves_relative_links():
    html = """
    <html><body>
      <p>附件：1.<a href="./W020250422312780758186.pdf">2024年度普通高等学校本科专业备案和审批结果</a></p>
      <p>2.<a href="/files/catalog.xlsx">普通高等学校本科专业目录（2025年）</a></p>
      <a href="https://example.com/news.html">普通网页</a>
    </body></html>
    """

    attachments = discover_attachments_from_html(
        html.encode("utf-8"),
        "http://www.moe.gov.cn/srcsite/A08/moe_1034/s4930/202504/t20250422_1188239.html",
    )

    assert [item["file_type"] for item in attachments] == ["pdf", "xlsx"]
    assert attachments[0]["attachment_url"].endswith("/202504/W020250422312780758186.pdf")
    assert attachments[0]["attachment_title"] == "2024年度普通高等学校本科专业备案和审批结果"


def test_parse_xlsx_major_candidates_extracts_major_rows(tmp_path):
    xlsx = tmp_path / "majors.xlsx"
    xlsx.write_bytes(_xlsx_bytes())
    source_row = _source_row()
    attachment = {
        "attachment_url": "http://www.moe.gov.cn/majors.xlsx",
        "attachment_title": "2024年度普通高等学校本科专业备案和审批结果",
        "file_type": "xlsx",
    }

    candidates = parse_xlsx_major_candidates(
        xlsx,
        source_row=source_row,
        attachment=attachment,
        captured_at="2026-06-12T00:00:00+08:00",
    )

    assert [row["major_name"] for row in candidates] == ["人工智能", "低空技术与工程"]
    assert candidates[0]["major_code"] == "080717T"
    assert candidates[0]["candidate_status"] == "catalog_confirmed"
    assert candidates[0]["event_type"] == "filing_added"
    assert "人工智能" in candidates[0]["evidence_text"]


def test_parse_docx_major_candidates_extracts_table_rows(tmp_path):
    docx_path = tmp_path / "majors.docx"
    document = Document()
    table = document.add_table(rows=1, cols=8)
    headers = ["序号", "学校名称", "专业代码", "专业名称", "门类", "专业类", "学位授予门类", "修业年限"]
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = value
    row = table.add_row().cells
    for index, value in enumerate([1, "示例大学", "080717T", "人工智能", "工学", "计算机类", "工学", "四年"]):
        row[index].text = str(value)
    document.save(docx_path)

    candidates = parse_docx_major_candidates(
        docx_path,
        source_row=_source_row(),
        attachment={
            "attachment_url": "http://www.moe.gov.cn/majors.docx",
            "attachment_title": "2024年度普通高等学校本科专业备案和审批结果",
            "file_type": "docx",
        },
        captured_at="2026-06-12T00:00:00+08:00",
    )

    assert len(candidates) == 1
    assert candidates[0]["major_code"] == "080717T"
    assert candidates[0]["major_name"] == "人工智能"
    assert candidates[0]["degree"] == "工学"
    assert candidates[0]["study_years"] == "四年"
    assert candidates[0]["parsed_from"] == "docx"


def test_parse_docx_major_candidates_handles_two_column_catalog(tmp_path):
    docx_path = tmp_path / "catalog.docx"
    document = Document()
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "01"
    table.rows[0].cells[1].text = "学科门类：哲学"
    row = table.add_row().cells
    row[0].text = "0101"
    row[1].text = "哲学类"
    row = table.add_row().cells
    row[0].text = "010101"
    row[1].text = "哲学"
    document.save(docx_path)

    candidates = parse_docx_major_candidates(
        docx_path,
        source_row={**_source_row(), "source_type": "undergraduate_catalog"},
        attachment={
            "attachment_url": "http://www.moe.gov.cn/catalog.docx",
            "attachment_title": "普通高等学校本科专业目录（2012年）",
            "file_type": "docx",
        },
        captured_at="2026-06-12T00:00:00+08:00",
    )

    assert len(candidates) == 1
    assert candidates[0]["major_code"] == "010101"
    assert candidates[0]["major_name"] == "哲学"
    assert candidates[0]["discipline_category"] == "哲学"
    assert candidates[0]["major_class"] == "哲学类"
    assert candidates[0]["event_type"] == "catalog_added"


def test_parse_pdf_major_candidates_from_text_handles_filing_and_catalog_rows(tmp_path):
    source_row = _source_row()
    filing_attachment = {
        "attachment_url": "http://www.moe.gov.cn/filing.pdf",
        "attachment_title": "2024年度普通高等学校本科专业备案和审批结果",
        "file_type": "pdf",
    }
    catalog_attachment = {
        "attachment_url": "http://www.moe.gov.cn/catalog.pdf",
        "attachment_title": "普通高等学校本科专业目录（2025年）",
        "file_type": "pdf",
    }
    filing_text = """
    序号       学校名称       专业名称       专业代码       学位授予门类 修业年限 备注
    24   北京邮电大学     低空技术与工程       083203TK   工学  四年    新专业
    28   北京林业大学     人工智能          080717T    工学  四年
    """
    catalog_text = """
    08    学科门类：工学
    0807      电子信息类
    080717T    人工智能
    083203TK    低空技术与工程
    """

    filing_candidates = parse_pdf_major_candidates_from_text(
        filing_text,
        source_row=source_row,
        attachment=filing_attachment,
        raw_path=tmp_path / "filing.pdf",
        captured_at="2026-06-12T00:00:00+08:00",
    )
    catalog_candidates = parse_pdf_major_candidates_from_text(
        catalog_text,
        source_row={**source_row, "source_type": "undergraduate_catalog"},
        attachment=catalog_attachment,
        raw_path=tmp_path / "catalog.pdf",
        captured_at="2026-06-12T00:00:00+08:00",
    )

    assert [row["major_name"] for row in filing_candidates] == ["低空技术与工程", "人工智能"]
    assert filing_candidates[0]["event_type"] == "filing_added"
    assert filing_candidates[0]["degree"] == "工学"
    assert [row["major_code"] for row in catalog_candidates] == ["080717T", "083203TK"]
    assert catalog_candidates[0]["major_class"] == "电子信息类"


def test_parse_filing_ocr_words_extracts_rows_and_repairs_codes(tmp_path):
    words_by_page = {
        1: [
            _word("53", 135, 100),
            _word("东北师范大学", 190, 100, 120),
            _word("数据科学", 450, 100, 80),
            _word("与大数据技术", 540, 100, 120),
            _word("08091OT", 760, 100, 70),
            _word("理学", 870, 100, 40),
            _word("四年", 980, 100, 40),
            _word("67", 135, 135),
            _word("东南大学", 190, 135, 90),
            _word("智能医学", 450, 135, 80),
            _word("工程", 540, 135, 50),
            _word("1010HT", 760, 135, 70),
            _word("工学", 870, 135, 40),
            _word("四年", 980, 135, 40),
        ]
    }

    candidates = _parse_filing_ocr_words(
        words_by_page,
        source_row=_source_row(),
        attachment={
            "attachment_url": "http://www.moe.gov.cn/filing.pdf",
            "attachment_title": "2019年度普通高等学校本科专业备案和审批结果",
            "file_type": "pdf",
        },
        raw_path=tmp_path / "filing.pdf",
        captured_at="2026-06-12T00:00:00+08:00",
    )

    assert [row["major_code"] for row in candidates] == ["080910T", "101011T"]
    assert [row["major_name"] for row in candidates] == ["数据科学与大数据技术", "智能医学工程"]
    assert candidates[0]["parsed_from"] == "pdf_ocr"
    assert candidates[0]["degree"] == "理学"
    assert candidates[0]["study_years"] == "四年"


def test_parse_catalog_ocr_words_handles_split_codes_and_multiline_names(tmp_path):
    words_by_page = {
        1: [
            _word("25经济学", 160, 100, 90),
            _word("金融学类", 300, 100, 90),
            _word("02031", 450, 100, 65),
            _word("OT金融科技", 580, 100, 90),
            _word("经济学", 805, 100, 60),
            _word("四年", 930, 100, 40),
            _word("信用风险管理与", 580, 130, 130),
            _word("31法学", 160, 160, 70),
            _word("法学类", 300, 160, 70),
            _word("030104T", 450, 160, 80),
            _word("法学", 805, 160, 40),
            _word("四年", 930, 160, 40),
            _word("法律防控", 580, 190, 80),
        ]
    }

    candidates = _parse_catalog_ocr_words(
        words_by_page,
        source_row={**_source_row(), "source_type": "undergraduate_catalog"},
        attachment={
            "attachment_url": "http://www.moe.gov.cn/catalog.pdf",
            "attachment_title": "普通高等学校本科专业目录（2020年版）",
            "file_type": "pdf",
        },
        raw_path=tmp_path / "catalog.pdf",
        captured_at="2026-06-12T00:00:00+08:00",
    )

    assert [row["major_code"] for row in candidates] == ["020310T", "030104T"]
    assert [row["major_name"] for row in candidates] == ["金融科技", "信用风险管理与法律防控"]
    assert candidates[0]["discipline_category"] == "经济学"
    assert candidates[0]["major_class"] == "金融学类"
    assert candidates[0]["parsed_from"] == "pdf_ocr"


def test_normalize_ocr_major_code_repairs_common_misreads():
    assert _normalize_ocr_major_code("08091OT") == "080910T"
    assert _normalize_ocr_major_code("1010HT") == "101011T"
    assert _normalize_ocr_major_code("0203mK") == "020301K"


def test_crawl_emerging_major_catalog_writes_candidates_attachments_and_manifest(tmp_path):
    source_csv = tmp_path / "sources.csv"
    source_csv.write_text(
        "source_id,title,url,source_domain,source_level,source_type,issuing_org,published_date,source_year,notes\n"
        "moe_2024,教育部通知,http://www.moe.gov.cn/source.html,www.moe.gov.cn,A,annual_filing_approval,教育部,2025-04-22,2024,seed\n",
        encoding="utf-8",
    )
    page_html = """
    <html><head><title>教育部通知</title></head><body>
      <a href="./majors.xlsx">2024年度普通高等学校本科专业备案和审批结果</a>
      <a href="./catalog.pdf">普通高等学校本科专业目录（2025年）</a>
    </body></html>
    """

    def fetcher(url, timeout_seconds):
        if url == "http://www.moe.gov.cn/source.html":
            return page_html.encode("utf-8")
        if url == "http://www.moe.gov.cn/majors.xlsx":
            return _xlsx_bytes()
        if url == "http://www.moe.gov.cn/catalog.pdf":
            return b"%PDF-1.4 sample"
        raise AssertionError(f"unexpected url: {url}")

    summary = crawl_emerging_major_catalog(
        source_csv=source_csv,
        raw_dir=tmp_path / "raw",
        processed_dir=tmp_path / "processed",
        logs_dir=tmp_path / "logs",
        run_id="test_run",
        fetcher=fetcher,
    )

    assert summary["success_count"] == 1
    assert summary["attachment_count"] == 2
    assert summary["candidate_count"] == 2
    assert Path(summary["attachments_jsonl"]).exists()
    assert Path(summary["candidates_jsonl"]).exists()
    attachment_rows = [
        json.loads(line)
        for line in Path(summary["attachments_jsonl"]).read_text(encoding="utf-8").splitlines()
    ]
    assert {row["parse_status"] for row in attachment_rows} == {"ok", "needs_review"}
    candidate_rows = [
        json.loads(line)
        for line in Path(summary["candidates_jsonl"]).read_text(encoding="utf-8").splitlines()
    ]
    assert {row["major_name"] for row in candidate_rows} == {"人工智能", "低空技术与工程"}
