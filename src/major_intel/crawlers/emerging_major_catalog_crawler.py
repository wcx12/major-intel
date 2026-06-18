"""Crawl official emerging-major catalog sources and attachment evidence.

The first version is intentionally source-preserving: HTML pages and
attachments are saved before parsing, and parse failures become reviewable
attachment records instead of silent skips.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import os
import re
import subprocess
import tempfile
import time
from datetime import datetime, timezone
from http.client import IncompleteRead
from pathlib import Path
from typing import Any
from urllib.parse import unquote, urljoin, urlparse
from urllib.request import Request, urlopen


SOURCE_SCHEMA_VERSION = "emerging_major_source_document/v1"
ATTACHMENT_SCHEMA_VERSION = "policy_attachment/v1"
CANDIDATE_SCHEMA_VERSION = "emerging_major_candidate/v1"

REQUIRED_SOURCE_FIELDS = {
    "source_id",
    "title",
    "url",
    "source_domain",
    "source_level",
    "source_type",
    "issuing_org",
    "published_date",
    "source_year",
    "notes",
}

DEFAULT_USER_AGENT = "major-intel-emerging-major-crawler/1.0"

WINDOWS_OCR_SCRIPT = r"""
[Console]::OutputEncoding = [System.Text.UTF8Encoding]::new()
Add-Type -AssemblyName System.Runtime.WindowsRuntime
[Windows.Storage.StorageFile, Windows.Storage, ContentType=WindowsRuntime] | Out-Null
[Windows.Storage.Streams.IRandomAccessStream, Windows.Storage.Streams, ContentType=WindowsRuntime] | Out-Null
[Windows.Graphics.Imaging.BitmapDecoder, Windows.Graphics.Imaging, ContentType=WindowsRuntime] | Out-Null
[Windows.Graphics.Imaging.SoftwareBitmap, Windows.Graphics.Imaging, ContentType=WindowsRuntime] | Out-Null
[Windows.Media.Ocr.OcrEngine, Windows.Foundation, ContentType=WindowsRuntime] | Out-Null
[Windows.Media.Ocr.OcrResult, Windows.Foundation, ContentType=WindowsRuntime] | Out-Null
[Windows.Globalization.Language, Windows.Globalization, ContentType=WindowsRuntime] | Out-Null
function Await($op, [type]$resultType) {
    $method = ([System.WindowsRuntimeSystemExtensions].GetMethods() | Where-Object {
        $_.Name -eq 'AsTask' -and $_.GetParameters().Count -eq 1 -and $_.ToString().Contains('IAsyncOperation')
    })[0]
    $task = $method.MakeGenericMethod($resultType).Invoke($null, @($op))
    $task.Wait()
    $task.Result
}
$path = (Resolve-Path $env:OCR_IMAGE_PATH).Path
$file = Await ([Windows.Storage.StorageFile]::GetFileFromPathAsync($path)) ([Windows.Storage.StorageFile])
$stream = Await ($file.OpenAsync([Windows.Storage.FileAccessMode]::Read)) ([Windows.Storage.Streams.IRandomAccessStream])
$decoder = Await ([Windows.Graphics.Imaging.BitmapDecoder]::CreateAsync($stream)) ([Windows.Graphics.Imaging.BitmapDecoder])
$bitmap = Await ($decoder.GetSoftwareBitmapAsync()) ([Windows.Graphics.Imaging.SoftwareBitmap])
$engine = [Windows.Media.Ocr.OcrEngine]::TryCreateFromLanguage((New-Object Windows.Globalization.Language 'zh-Hans-CN'))
if ($null -eq $engine) { throw 'Windows zh-Hans-CN OCR engine is not available.' }
$result = Await ($engine.RecognizeAsync($bitmap)) ([Windows.Media.Ocr.OcrResult])
$words = New-Object System.Collections.Generic.List[object]
foreach ($line in $result.Lines) {
    foreach ($word in $line.Words) {
        $words.Add([PSCustomObject]@{
            text = $word.Text
            x = [int]$word.BoundingRect.X
            y = [int]$word.BoundingRect.Y
            w = [int]$word.BoundingRect.Width
            h = [int]$word.BoundingRect.Height
        })
    }
}
$words | ConvertTo-Json -Depth 4 -Compress
"""

DISCIPLINE_CATEGORIES = {
    "哲学",
    "经济学",
    "法学",
    "教育学",
    "文学",
    "历史学",
    "理学",
    "工学",
    "农学",
    "医学",
    "管理学",
    "艺术学",
    "交叉学科",
}

ATTACHMENT_EXTENSIONS = {
    ".pdf": "pdf",
    ".xls": "xls",
    ".xlsx": "xlsx",
    ".doc": "doc",
    ".docx": "docx",
    ".zip": "zip",
    ".rar": "rar",
}

CSV_FIELDS = [
    "candidate_id",
    "major_code",
    "major_name",
    "major_level",
    "discipline_category",
    "major_class",
    "degree",
    "study_years",
    "event_type",
    "event_year",
    "candidate_status",
    "source_title",
    "source_url",
    "attachment_url",
    "source_level",
    "evidence_text",
    "raw_path",
    "parsed_from",
    "captured_at",
    "warnings_json",
]


def load_source_rows(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = [dict(row) for row in csv.DictReader(handle)]
    if not rows:
        raise ValueError(f"{path} has no source rows")
    missing = REQUIRED_SOURCE_FIELDS - set(rows[0])
    if missing:
        raise ValueError(f"{path} missing required fields: {sorted(missing)}")
    return rows


def discover_attachments_from_html(raw_html: bytes | str, page_url: str) -> list[dict[str, str]]:
    html = _decode_bytes(raw_html) if isinstance(raw_html, bytes) else raw_html
    links = _extract_links(html)
    attachments: list[dict[str, str]] = []
    seen: set[str] = set()
    for href, text in links:
        if not href or href.lower().startswith(("javascript:", "mailto:")):
            continue
        attachment_url = urljoin(page_url, href)
        file_type = infer_file_type(attachment_url, title=text)
        if file_type == "html":
            continue
        if attachment_url in seen:
            continue
        seen.add(attachment_url)
        attachments.append(
            {
                "attachment_url": attachment_url,
                "attachment_title": _clean_space(text) or Path(urlparse(attachment_url).path).name,
                "file_type": file_type,
            }
        )
    return attachments


def parse_xlsx_major_candidates(
    path: Path,
    *,
    source_row: dict[str, str],
    attachment: dict[str, str],
    captured_at: str,
) -> list[dict[str, Any]]:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:  # pragma: no cover - dependency exists in test env.
        raise RuntimeError("openpyxl is required to parse xlsx attachments") from exc

    workbook = load_workbook(path, read_only=True, data_only=True)
    candidates: list[dict[str, Any]] = []
    for sheet in workbook.worksheets:
        rows = [[_cell_text(value) for value in row] for row in sheet.iter_rows(values_only=True)]
        candidates.extend(
            _parse_tabular_major_rows(
                rows,
                source_row=source_row,
                attachment=attachment,
                raw_path=path,
                captured_at=captured_at,
                parsed_from="xlsx",
            )
        )
    workbook.close()
    return candidates


def parse_xls_major_candidates(
    path: Path,
    *,
    source_row: dict[str, str],
    attachment: dict[str, str],
    captured_at: str,
) -> list[dict[str, Any]]:
    candidates, _ = _parse_xls_with_conversion(
        path,
        source_row=source_row,
        attachment=attachment,
        captured_at=captured_at,
    )
    return candidates


def parse_docx_major_candidates(
    path: Path,
    *,
    source_row: dict[str, str],
    attachment: dict[str, str],
    captured_at: str,
) -> list[dict[str, Any]]:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency exists in repo env.
        raise RuntimeError("python-docx is required to parse docx attachments") from exc

    document = Document(str(path))
    candidates: list[dict[str, Any]] = []
    for table in document.tables:
        rows = [[_clean_space(cell.text) for cell in row.cells] for row in table.rows]
        candidates.extend(
            _parse_tabular_major_rows(
                rows,
                source_row=source_row,
                attachment=attachment,
                raw_path=path,
                captured_at=captured_at,
                parsed_from="docx",
            )
        )
    if not candidates and _is_catalog_attachment(source_row, attachment):
        for table in document.tables:
            rows = [[_clean_space(cell.text) for cell in row.cells] for row in table.rows]
            candidates.extend(
                _parse_two_column_catalog_rows(
                    rows,
                    source_row=source_row,
                    attachment=attachment,
                    raw_path=path,
                    captured_at=captured_at,
                    parsed_from="docx",
                )
            )
    if candidates:
        return candidates
    return _parse_word_paragraph_major_rows(
        [paragraph.text for paragraph in document.paragraphs],
        source_row=source_row,
        attachment=attachment,
        raw_path=path,
        captured_at=captured_at,
        parsed_from="docx",
    )


def parse_doc_major_candidates(
    path: Path,
    *,
    source_row: dict[str, str],
    attachment: dict[str, str],
    captured_at: str,
) -> list[dict[str, Any]]:
    candidates, _ = _parse_doc_with_conversion(
        path,
        source_row=source_row,
        attachment=attachment,
        captured_at=captured_at,
    )
    return candidates


def parse_pdf_major_candidates(
    path: Path,
    *,
    source_row: dict[str, str],
    attachment: dict[str, str],
    captured_at: str,
) -> list[dict[str, Any]]:
    extracted_text = extract_pdf_text(path)
    candidates = parse_pdf_major_candidates_from_text(
        extracted_text,
        source_row=source_row,
        attachment=attachment,
        raw_path=path,
        captured_at=captured_at,
    )
    if candidates:
        return candidates
    if extracted_text.strip():
        return []
    return parse_pdf_major_candidates_with_windows_ocr(
        path,
        source_row=source_row,
        attachment=attachment,
        captured_at=captured_at,
    )


def parse_pdf_major_candidates_with_windows_ocr(
    path: Path,
    *,
    source_row: dict[str, str],
    attachment: dict[str, str],
    captured_at: str,
) -> list[dict[str, Any]]:
    words_by_page = extract_pdf_ocr_words(path)
    if not words_by_page:
        return []
    if _is_catalog_attachment(source_row, attachment):
        return _parse_catalog_ocr_words(
            words_by_page,
            source_row=source_row,
            attachment=attachment,
            raw_path=path,
            captured_at=captured_at,
        )
    return _parse_filing_ocr_words(
        words_by_page,
        source_row=source_row,
        attachment=attachment,
        raw_path=path,
        captured_at=captured_at,
    )


def parse_pdf_major_candidates_from_text(
    text: str,
    *,
    source_row: dict[str, str],
    attachment: dict[str, str],
    raw_path: Path,
    captured_at: str,
) -> list[dict[str, Any]]:
    if _is_catalog_attachment(source_row, attachment):
        return _parse_catalog_pdf_text(
            text,
            source_row=source_row,
            attachment=attachment,
            raw_path=raw_path,
            captured_at=captured_at,
        )
    return _parse_filing_pdf_text(
        text,
        source_row=source_row,
        attachment=attachment,
        raw_path=raw_path,
        captured_at=captured_at,
    )


def crawl_emerging_major_catalog(
    *,
    source_csv: Path,
    raw_dir: Path,
    processed_dir: Path,
    logs_dir: Path,
    run_id: str,
    fetcher=None,
    timeout_seconds: float = 20,
    sleep_seconds: float = 0.0,
) -> dict[str, Any]:
    sources = load_source_rows(source_csv)
    processed_dir.mkdir(parents=True, exist_ok=True)
    logs_dir.mkdir(parents=True, exist_ok=True)

    documents_jsonl = processed_dir / f"documents_{run_id}.jsonl"
    attachments_jsonl = processed_dir / f"attachments_{run_id}.jsonl"
    candidates_jsonl = processed_dir / f"emerging_major_candidates_{run_id}.jsonl"
    failure_log = logs_dir / f"{run_id}_failures.jsonl"
    manifest_path = logs_dir / f"{run_id}_manifest.json"

    actual_fetcher = fetcher or fetch_url_bytes
    success_count = 0
    failure_count = 0
    attachment_count = 0
    candidate_count = 0

    with (
        documents_jsonl.open("w", encoding="utf-8") as documents_handle,
        attachments_jsonl.open("w", encoding="utf-8") as attachments_handle,
        candidates_jsonl.open("w", encoding="utf-8") as candidates_handle,
        failure_log.open("w", encoding="utf-8") as failure_handle,
    ):
        for source in sources:
            captured_at = _now_iso()
            try:
                page_bytes = actual_fetcher(source["url"], timeout_seconds)
                page_raw_path = _write_raw_bytes(
                    page_bytes,
                    raw_dir=raw_dir,
                    run_id=run_id,
                    url=source["url"],
                    preferred_file_type="html",
                )
                attachments = discover_attachments_from_html(page_bytes, source["url"])
                document = _build_source_document(
                    source,
                    page_bytes=page_bytes,
                    raw_path=page_raw_path,
                    captured_at=captured_at,
                    attachment_count=len(attachments),
                )
                _write_jsonl(documents_handle, document)
                success_count += 1
            except Exception as exc:
                failure_count += 1
                _write_jsonl(
                    failure_handle,
                    _failure_row(source, source.get("url", ""), exc, failure_scope="source_page"),
                )
                continue

            for attachment in attachments:
                try:
                    attachment_bytes = actual_fetcher(attachment["attachment_url"], timeout_seconds)
                    detected_file_type = infer_file_type(
                        attachment["attachment_url"],
                        title=attachment.get("attachment_title", ""),
                        raw_bytes=attachment_bytes,
                    )
                    attachment["file_type"] = detected_file_type
                    attachment_raw_path = _write_raw_bytes(
                        attachment_bytes,
                        raw_dir=raw_dir,
                        run_id=run_id,
                        url=attachment["attachment_url"],
                        preferred_file_type=detected_file_type,
                    )
                    parsed_candidates, attachment_record = _parse_attachment(
                        source,
                        attachment=attachment,
                        raw_path=attachment_raw_path,
                        captured_at=captured_at,
                        parent_doc_id=document["doc_id"],
                    )
                    attachment_count += 1
                    candidate_count += len(parsed_candidates)
                    _write_jsonl(attachments_handle, attachment_record)
                    for candidate in parsed_candidates:
                        _write_jsonl(candidates_handle, candidate)
                except Exception as exc:
                    failure_count += 1
                    attachment_count += 1
                    _write_jsonl(
                        attachments_handle,
                        _attachment_record(
                            source,
                            attachment=attachment,
                            raw_path=None,
                            parent_doc_id=document["doc_id"],
                            parse_status="error",
                            row_count=0,
                            candidate_major_count=0,
                            warnings=[f"{type(exc).__name__}: {exc}"],
                        ),
                    )
                    _write_jsonl(
                        failure_handle,
                        _failure_row(
                            source,
                            attachment.get("attachment_url", ""),
                            exc,
                            failure_scope="attachment",
                        ),
                    )
            if sleep_seconds:
                time.sleep(sleep_seconds)

    manifest = {
        "run_id": run_id,
        "source_csv": str(source_csv),
        "raw_dir": str(raw_dir),
        "processed_dir": str(processed_dir),
        "logs_dir": str(logs_dir),
        "documents_jsonl": str(documents_jsonl),
        "attachments_jsonl": str(attachments_jsonl),
        "candidates_jsonl": str(candidates_jsonl),
        "failure_log": str(failure_log),
        "manifest_path": str(manifest_path),
        "success_count": success_count,
        "failure_count": failure_count,
        "attachment_count": attachment_count,
        "candidate_count": candidate_count,
        "finished_at": _now_iso(),
    }
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    return manifest


def infer_file_type(url: str, *, title: str = "", raw_bytes: bytes | None = None) -> str:
    path = unquote(urlparse(url).path).lower()
    title_lower = title.lower()
    for suffix, file_type in ATTACHMENT_EXTENSIONS.items():
        if path.endswith(suffix) or suffix in title_lower:
            return file_type
    if raw_bytes:
        head = raw_bytes[:4096]
        if head.startswith(b"%PDF"):
            return "pdf"
        if head.startswith(b"PK"):
            return _zip_family_file_type(raw_bytes)
    return "html"


def fetch_url_bytes(url: str, timeout_seconds: float = 20, max_attempts: int = 3) -> bytes:
    last_exc: Exception | None = None
    for attempt in range(1, max_attempts + 1):
        request = Request(url, headers={"User-Agent": DEFAULT_USER_AGENT})
        try:
            with urlopen(request, timeout=timeout_seconds) as response:
                return response.read()
        except IncompleteRead as exc:
            last_exc = exc
        except Exception as exc:
            status = getattr(exc, "code", None)
            if status is None or status < 500 or attempt == max_attempts:
                raise
            last_exc = exc
        if attempt < max_attempts:
            time.sleep(min(2 ** (attempt - 1), 4))
    if last_exc is not None:
        raise last_exc
    raise RuntimeError(f"failed to fetch {url}")


def write_candidates_csv(candidates: list[dict[str, Any]], path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=CSV_FIELDS)
        writer.writeheader()
        for row in candidates:
            writer.writerow(_candidate_csv_row(row))


def read_jsonl(path: Path) -> list[dict[str, Any]]:
    if not path.exists():
        return []
    rows: list[dict[str, Any]] = []
    with path.open("r", encoding="utf-8") as handle:
        for line in handle:
            if line.strip():
                rows.append(json.loads(line))
    return rows


def _parse_attachment(
    source_row: dict[str, str],
    *,
    attachment: dict[str, str],
    raw_path: Path,
    captured_at: str,
    parent_doc_id: str,
) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    file_type = attachment.get("file_type") or infer_file_type(attachment["attachment_url"])
    if file_type == "xlsx":
        candidates = parse_xlsx_major_candidates(
            raw_path,
            source_row=source_row,
            attachment=attachment,
            captured_at=captured_at,
        )
        return candidates, _attachment_record(
            source_row,
            attachment=attachment,
            raw_path=raw_path,
            parent_doc_id=parent_doc_id,
            parse_status="ok",
            row_count=_xlsx_row_count(raw_path),
            candidate_major_count=len(candidates),
            warnings=[],
        )
    if file_type == "xls":
        try:
            candidates, row_count = _parse_xls_with_conversion(
                raw_path,
                source_row=source_row,
                attachment=attachment,
                captured_at=captured_at,
            )
        except Exception as exc:
            return [], _attachment_record(
                source_row,
                attachment=attachment,
                raw_path=raw_path,
                parent_doc_id=parent_doc_id,
                parse_status="needs_review",
                row_count=0,
                candidate_major_count=0,
                warnings=[f"xls_parse_error:{type(exc).__name__}: {exc}"],
            )
        return candidates, _attachment_record(
            source_row,
            attachment=attachment,
            raw_path=raw_path,
            parent_doc_id=parent_doc_id,
            parse_status="ok" if candidates else "needs_review",
            row_count=row_count,
            candidate_major_count=len(candidates),
            warnings=[] if candidates else ["xls_converted_but_no_major_rows_detected"],
        )
    if file_type == "docx":
        try:
            candidates = parse_docx_major_candidates(
                raw_path,
                source_row=source_row,
                attachment=attachment,
                captured_at=captured_at,
            )
        except Exception as exc:
            return [], _attachment_record(
                source_row,
                attachment=attachment,
                raw_path=raw_path,
                parent_doc_id=parent_doc_id,
                parse_status="needs_review",
                row_count=0,
                candidate_major_count=0,
                warnings=[f"docx_parse_error:{type(exc).__name__}: {exc}"],
            )
        return candidates, _attachment_record(
            source_row,
            attachment=attachment,
            raw_path=raw_path,
            parent_doc_id=parent_doc_id,
            parse_status="ok" if candidates else "needs_review",
            row_count=_docx_row_count(raw_path),
            candidate_major_count=len(candidates),
            warnings=[] if candidates else ["docx_parsed_but_no_major_rows_detected"],
        )
    if file_type == "doc":
        try:
            candidates, row_count = _parse_doc_with_conversion(
                raw_path,
                source_row=source_row,
                attachment=attachment,
                captured_at=captured_at,
            )
        except Exception as exc:
            return [], _attachment_record(
                source_row,
                attachment=attachment,
                raw_path=raw_path,
                parent_doc_id=parent_doc_id,
                parse_status="needs_review",
                row_count=0,
                candidate_major_count=0,
                warnings=[f"doc_parse_error:{type(exc).__name__}: {exc}"],
            )
        return candidates, _attachment_record(
            source_row,
            attachment=attachment,
            raw_path=raw_path,
            parent_doc_id=parent_doc_id,
            parse_status="ok" if candidates else "needs_review",
            row_count=row_count,
            candidate_major_count=len(candidates),
            warnings=[] if candidates else ["doc_converted_but_no_major_rows_detected"],
        )
    if file_type == "pdf":
        try:
            candidates = parse_pdf_major_candidates(
                raw_path,
                source_row=source_row,
                attachment=attachment,
                captured_at=captured_at,
            )
        except Exception as exc:
            return [], _attachment_record(
                source_row,
                attachment=attachment,
                raw_path=raw_path,
                parent_doc_id=parent_doc_id,
                parse_status="needs_review",
                row_count=0,
                candidate_major_count=0,
                warnings=[f"pdf_parse_error:{type(exc).__name__}: {exc}"],
            )
        return candidates, _attachment_record(
            source_row,
            attachment=attachment,
            raw_path=raw_path,
            parent_doc_id=parent_doc_id,
            parse_status="ok" if candidates else "needs_review",
            row_count=_pdf_text_line_count(raw_path),
            candidate_major_count=len(candidates),
            warnings=[] if candidates else ["pdf_text_extracted_but_no_major_rows_detected"],
        )
    return [], _attachment_record(
        source_row,
        attachment=attachment,
        raw_path=raw_path,
        parent_doc_id=parent_doc_id,
        parse_status="needs_review",
        row_count=0,
        candidate_major_count=0,
        warnings=[f"{file_type}_parser_not_implemented"],
    )


def _attachment_record(
    source_row: dict[str, str],
    *,
    attachment: dict[str, str],
    raw_path: Path | None,
    parent_doc_id: str,
    parse_status: str,
    row_count: int,
    candidate_major_count: int,
    warnings: list[str],
) -> dict[str, Any]:
    return {
        "schema_version": ATTACHMENT_SCHEMA_VERSION,
        "parent_doc_id": parent_doc_id,
        "parent_source_id": source_row["source_id"],
        "source_title": source_row["title"],
        "source_url": source_row["url"],
        "source_year": source_row["source_year"],
        "source_level": source_row["source_level"],
        "attachment_url": attachment.get("attachment_url", ""),
        "attachment_title": attachment.get("attachment_title", ""),
        "file_type": attachment.get("file_type", ""),
        "raw_path": str(raw_path) if raw_path else "",
        "parse_status": parse_status,
        "row_count": row_count,
        "candidate_major_count": candidate_major_count,
        "warnings": warnings,
    }


def _build_candidate(
    *,
    source_row: dict[str, str],
    attachment: dict[str, str],
    raw_path: Path,
    captured_at: str,
    major_code: str,
    major_name: str,
    major_level: str,
    discipline_category: str,
    major_class: str,
    degree: str,
    study_years: str,
    event_type: str,
    evidence_text: str,
    parsed_from: str,
    row_index: int,
) -> dict[str, Any]:
    status = "catalog_confirmed" if major_code and major_name else "catalog_candidate"
    candidate_id = "emerging_major:" + _sha256_text(
        "|".join(
            [
                source_row["url"],
                attachment.get("attachment_url", ""),
                major_code,
                major_name,
                str(row_index),
            ]
        )
    )
    return {
        "schema_version": CANDIDATE_SCHEMA_VERSION,
        "candidate_id": candidate_id,
        "major_code": major_code,
        "major_name": major_name,
        "major_level": major_level,
        "discipline_category": discipline_category,
        "major_class": major_class,
        "degree": degree,
        "study_years": study_years,
        "event_type": event_type,
        "event_year": _to_int(source_row.get("source_year")),
        "candidate_status": status,
        "source_title": source_row["title"],
        "source_url": source_row["url"],
        "attachment_url": attachment.get("attachment_url", ""),
        "source_level": source_row["source_level"],
        "evidence_text": evidence_text,
        "raw_path": str(raw_path),
        "parsed_from": parsed_from,
        "captured_at": captured_at,
        "warnings": [],
    }


def extract_pdf_text(path: Path) -> str:
    pdftotext_text = _extract_pdf_text_with_pdftotext(path)
    if pdftotext_text:
        return pdftotext_text
    try:
        from pypdf import PdfReader
    except ImportError:
        return ""
    try:
        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages).strip()
    except Exception:
        return ""


def extract_pdf_ocr_words(path: Path) -> dict[int, list[dict[str, Any]]]:
    cache_path = path.with_name(path.name + ".ocr.jsonl")
    cached = _read_ocr_words_jsonl(cache_path)
    if cached:
        return cached
    with tempfile.TemporaryDirectory(prefix="major_intel_pdf_ocr_") as temp_dir:
        image_paths = _render_pdf_page_images(path, Path(temp_dir))
        rows: list[dict[str, Any]] = []
        for image_path in image_paths:
            page_match = re.search(r"-(\d+)\.png$", image_path.name)
            page = int(page_match.group(1)) if page_match else len(rows) + 1
            for word in _run_windows_ocr_image(image_path):
                word["page"] = page
                rows.append(word)
            _write_ocr_words_jsonl(cache_path, _group_words_by_page(rows))
        return _group_words_by_page(rows)


def _render_pdf_page_images(path: Path, target_dir: Path) -> list[Path]:
    target_dir.mkdir(parents=True, exist_ok=True)
    try:
        subprocess.run(
            ["pdftoppm", "-png", "-r", "150", str(path), str(target_dir / "page")],
            check=True,
            capture_output=True,
            text=True,
        )
    except (FileNotFoundError, subprocess.CalledProcessError):
        return []
    return sorted(target_dir.glob("page-*.png"))


def _run_windows_ocr_image(image_path: Path) -> list[dict[str, Any]]:
    try:
        completed = subprocess.run(
            [
                "powershell",
                "-NoProfile",
                "-ExecutionPolicy",
                "Bypass",
                "-Command",
                WINDOWS_OCR_SCRIPT,
            ],
            check=True,
            capture_output=True,
            encoding="utf-8",
            env={**os.environ, "OCR_IMAGE_PATH": str(image_path)},
        )
    except (FileNotFoundError, subprocess.CalledProcessError):
        return []
    try:
        data = json.loads(completed.stdout or "[]")
    except json.JSONDecodeError:
        return []
    if isinstance(data, dict):
        data = [data]
    return [
        {
            "text": _clean_space(item.get("text")),
            "x": float(item.get("x") or 0),
            "y": float(item.get("y") or 0),
            "w": float(item.get("w") or 0),
            "h": float(item.get("h") or 0),
        }
        for item in data
        if _clean_space(item.get("text"))
    ]


def _read_ocr_words_jsonl(path: Path) -> dict[int, list[dict[str, Any]]]:
    if not path.exists():
        return {}
    rows: dict[int, list[dict[str, Any]]] = {}
    with path.open("r", encoding="utf-8") as handle:
        for line in handle:
            if not line.strip():
                continue
            item = json.loads(line)
            rows.setdefault(int(item.get("page") or 0), []).append(item)
    return rows


def _write_ocr_words_jsonl(path: Path, words_by_page: dict[int, list[dict[str, Any]]]) -> None:
    with path.open("w", encoding="utf-8") as handle:
        for page, words in sorted(words_by_page.items()):
            for word in words:
                row = dict(word)
                row["page"] = page
                handle.write(json.dumps(row, ensure_ascii=False, sort_keys=True) + "\n")


def _group_words_by_page(rows: list[dict[str, Any]]) -> dict[int, list[dict[str, Any]]]:
    grouped: dict[int, list[dict[str, Any]]] = {}
    for row in rows:
        grouped.setdefault(int(row.get("page") or 0), []).append(row)
    return grouped


def _extract_pdf_text_with_pdftotext(path: Path) -> str:
    try:
        completed = subprocess.run(
            ["pdftotext", "-layout", str(path), "-"],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            check=False,
        )
    except FileNotFoundError:
        return ""
    if not completed.stdout:
        return ""
    return _decode_bytes(completed.stdout).strip()


def _parse_filing_pdf_text(
    text: str,
    *,
    source_row: dict[str, str],
    attachment: dict[str, str],
    raw_path: Path,
    captured_at: str,
) -> list[dict[str, Any]]:
    candidates: list[dict[str, Any]] = []
    for line_number, line in enumerate(text.splitlines(), start=1):
        parsed = _parse_filing_pdf_line(line)
        if not parsed:
            continue
        major_code, major_name, degree, study_years, evidence_text = parsed
        candidates.append(
            _build_candidate(
                source_row=source_row,
                attachment=attachment,
                raw_path=raw_path,
                captured_at=captured_at,
                major_code=major_code,
                major_name=major_name,
                major_level="本科",
                discipline_category="",
                major_class="",
                degree=degree,
                study_years=study_years,
                event_type=_infer_event_type(source_row, attachment),
                evidence_text=evidence_text,
                parsed_from="pdf",
                row_index=line_number,
            )
        )
    return candidates


def _parse_catalog_pdf_text(
    text: str,
    *,
    source_row: dict[str, str],
    attachment: dict[str, str],
    raw_path: Path,
    captured_at: str,
) -> list[dict[str, Any]]:
    candidates: list[dict[str, Any]] = []
    current_discipline = ""
    current_major_class = ""
    for line_number, line in enumerate(text.splitlines(), start=1):
        clean = _clean_space(line)
        if not clean:
            continue
        if clean in DISCIPLINE_CATEGORIES:
            current_discipline = clean
            current_major_class = ""
            continue
        discipline_match = re.match(r"^\d{2}\s+学科门类[:：]\s*(\S+)", clean)
        if discipline_match:
            current_discipline = discipline_match.group(1)
            current_major_class = ""
            continue
        class_match = re.match(r"^(\d{4})\s+(.+类)$", clean)
        if class_match:
            current_major_class = class_match.group(2)
            continue
        major_match = re.match(r"^(\d{6}[A-ZKTY]*)\s+(.+)$", clean)
        if major_match:
            major_code = major_match.group(1)
            major_name = _clean_space(re.sub(r"（注[:：]?.*?）", "", major_match.group(2)))
            if not _looks_like_major_name(major_name):
                continue
            candidates.append(
                _build_candidate(
                    source_row=source_row,
                    attachment=attachment,
                    raw_path=raw_path,
                    captured_at=captured_at,
                    major_code=major_code,
                    major_name=major_name,
                    major_level="本科",
                    discipline_category=current_discipline,
                    major_class=current_major_class,
                    degree="",
                    study_years="",
                    event_type="catalog_added",
                    evidence_text=clean,
                    parsed_from="pdf",
                    row_index=line_number,
                )
            )
            continue
        table_row = _parse_catalog_table_pdf_line(line)
        if not table_row:
            continue
        major_code, major_name, discipline_category, major_class, degree, study_years, evidence_text = table_row
        candidates.append(
            _build_candidate(
                source_row=source_row,
                attachment=attachment,
                raw_path=raw_path,
                captured_at=captured_at,
                major_code=major_code,
                major_name=major_name,
                major_level="本科",
                discipline_category=discipline_category or current_discipline,
                major_class=major_class or current_major_class,
                degree=degree,
                study_years=study_years,
                event_type="catalog_added",
                evidence_text=evidence_text,
                parsed_from="pdf",
                row_index=line_number,
            )
        )
    return candidates


def _parse_tabular_major_rows(
    rows: list[list[str]],
    *,
    source_row: dict[str, str],
    attachment: dict[str, str],
    raw_path: Path,
    captured_at: str,
    parsed_from: str,
) -> list[dict[str, Any]]:
    candidates: list[dict[str, Any]] = []
    header_index, header_map = _find_header(rows)
    if header_index is None:
        return candidates
    for row_index, row in enumerate(rows[header_index + 1 :], start=header_index + 2):
        if not any(row):
            continue
        major_name = _value_by_header(row, header_map, "major_name")
        major_code = _normalize_major_code(_value_by_header(row, header_map, "major_code"))
        if not _looks_like_major_name(major_name):
            continue
        evidence_text = " | ".join(value for value in row if value)
        candidates.append(
            _build_candidate(
                source_row=source_row,
                attachment=attachment,
                raw_path=raw_path,
                captured_at=captured_at,
                major_code=major_code,
                major_name=major_name,
                major_level="本科",
                discipline_category=_value_by_header(row, header_map, "discipline_category"),
                major_class=_value_by_header(row, header_map, "major_class"),
                degree=_value_by_header(row, header_map, "degree"),
                study_years=_value_by_header(row, header_map, "study_years"),
                event_type=_infer_event_type(source_row, attachment),
                evidence_text=evidence_text,
                parsed_from=parsed_from,
                row_index=row_index,
            )
        )
    return candidates


def _parse_two_column_catalog_rows(
    rows: list[list[str]],
    *,
    source_row: dict[str, str],
    attachment: dict[str, str],
    raw_path: Path,
    captured_at: str,
    parsed_from: str,
) -> list[dict[str, Any]]:
    candidates: list[dict[str, Any]] = []
    current_discipline = ""
    current_major_class = ""
    for row_index, row in enumerate(rows, start=1):
        values = [value for value in row if value]
        if len(values) < 2:
            continue
        code = _normalize_major_code(values[0])
        text = _clean_space(values[1])
        if re.fullmatch(r"\d{2}", code) and "学科门类" in text:
            current_discipline = text.split("：")[-1].split(":")[-1].strip()
            current_major_class = ""
            continue
        if re.fullmatch(r"\d{4}", code) and text.endswith("类"):
            current_major_class = text
            continue
        if not re.fullmatch(r"\d{6}[A-ZKTY]*", code):
            continue
        major_name = _clean_space(re.sub(r"（注[:：]?.*?）", "", text))
        if not _looks_like_major_name(major_name):
            continue
        candidates.append(
            _build_candidate(
                source_row=source_row,
                attachment=attachment,
                raw_path=raw_path,
                captured_at=captured_at,
                major_code=code,
                major_name=major_name,
                major_level="本科",
                discipline_category=current_discipline,
                major_class=current_major_class,
                degree="",
                study_years="",
                event_type="catalog_added",
                evidence_text=" | ".join(values),
                parsed_from=parsed_from,
                row_index=row_index,
            )
        )
    return candidates


def _parse_word_paragraph_major_rows(
    lines: list[str],
    *,
    source_row: dict[str, str],
    attachment: dict[str, str],
    raw_path: Path,
    captured_at: str,
    parsed_from: str,
) -> list[dict[str, Any]]:
    candidates: list[dict[str, Any]] = []
    for line_number, line in enumerate(lines, start=1):
        if _is_catalog_attachment(source_row, attachment):
            parsed_catalog = _parse_catalog_table_pdf_line(line)
            if parsed_catalog:
                major_code, major_name, discipline_category, major_class, degree, study_years, evidence_text = parsed_catalog
                candidates.append(
                    _build_candidate(
                        source_row=source_row,
                        attachment=attachment,
                        raw_path=raw_path,
                        captured_at=captured_at,
                        major_code=major_code,
                        major_name=major_name,
                        major_level="本科",
                        discipline_category=discipline_category,
                        major_class=major_class,
                        degree=degree,
                        study_years=study_years,
                        event_type="catalog_added",
                        evidence_text=evidence_text,
                        parsed_from=parsed_from,
                        row_index=line_number,
                    )
                )
                continue
        parsed_filing = _parse_filing_pdf_line(line)
        if not parsed_filing:
            continue
        major_code, major_name, degree, study_years, evidence_text = parsed_filing
        candidates.append(
            _build_candidate(
                source_row=source_row,
                attachment=attachment,
                raw_path=raw_path,
                captured_at=captured_at,
                major_code=major_code,
                major_name=major_name,
                major_level="本科",
                discipline_category="",
                major_class="",
                degree=degree,
                study_years=study_years,
                event_type=_infer_event_type(source_row, attachment),
                evidence_text=evidence_text,
                parsed_from=parsed_from,
                row_index=line_number,
            )
        )
    return candidates


def _parse_filing_ocr_words(
    words_by_page: dict[int, list[dict[str, Any]]],
    *,
    source_row: dict[str, str],
    attachment: dict[str, str],
    raw_path: Path,
    captured_at: str,
) -> list[dict[str, Any]]:
    candidates: list[dict[str, Any]] = []
    for page, words in sorted(words_by_page.items()):
        rows = _group_ocr_rows(words)
        for row_index, row in enumerate(rows, start=1):
            major_code = _normalize_ocr_major_code(_ocr_band_text(row, 735, 850))
            major_name = _clean_space(_ocr_band_text(row, 430, 735))
            if not major_code or not _looks_like_major_name(major_name):
                continue
            degree, study_years = _split_ocr_degree_years(_ocr_band_text(row, 850, 1125))
            evidence_text = " | ".join(
                value
                for value in [
                    _ocr_band_text(row, 90, 180),
                    _ocr_band_text(row, 180, 430),
                    major_name,
                    major_code,
                    degree,
                    study_years,
                    _ocr_band_text(row, 1125, 1300),
                ]
                if value
            )
            candidates.append(
                _build_candidate(
                    source_row=source_row,
                    attachment=attachment,
                    raw_path=raw_path,
                    captured_at=captured_at,
                    major_code=major_code,
                    major_name=major_name,
                    major_level="本科",
                    discipline_category="",
                    major_class="",
                    degree=degree,
                    study_years=study_years,
                    event_type=_infer_event_type(source_row, attachment),
                    evidence_text=evidence_text,
                    parsed_from="pdf_ocr",
                    row_index=page * 10000 + row_index,
                )
            )
    return candidates


def _parse_catalog_ocr_words(
    words_by_page: dict[int, list[dict[str, Any]]],
    *,
    source_row: dict[str, str],
    attachment: dict[str, str],
    raw_path: Path,
    captured_at: str,
) -> list[dict[str, Any]]:
    candidates: list[dict[str, Any]] = []
    for page, words in sorted(words_by_page.items()):
        rows = _group_ocr_rows(words)
        for row_index, row in enumerate(rows, start=1):
            code_band = _ocr_band_text(row, 430, 570)
            name_band = _ocr_band_text(row, 570, 790)
            major_code, major_name = _split_ocr_code_and_name(code_band, name_band)
            if not major_code:
                continue
            if not major_name:
                major_name = _ocr_name_continuation(rows, row_index - 1)
            if not _looks_like_major_name(major_name):
                continue
            degree, study_years = _split_ocr_degree_years(_ocr_band_text(row, 790, 1035))
            major_class = _clean_space(_ocr_band_text(row, 265, 430) + _ocr_code_band_prefix(code_band))
            evidence_text = " | ".join(
                value
                for value in [
                    _ocr_band_text(row, 150, 265),
                    major_class,
                    major_code,
                    major_name,
                    degree,
                    study_years,
                    _ocr_band_text(row, 1035, 1135),
                ]
                if value
            )
            candidates.append(
                _build_candidate(
                    source_row=source_row,
                    attachment=attachment,
                    raw_path=raw_path,
                    captured_at=captured_at,
                    major_code=major_code,
                    major_name=major_name,
                    major_level="本科",
                    discipline_category=_extract_discipline_category(_ocr_band_text(row, 150, 265)),
                    major_class=major_class,
                    degree=degree,
                    study_years=study_years,
                    event_type="catalog_added",
                    evidence_text=evidence_text,
                    parsed_from="pdf_ocr",
                    row_index=page * 10000 + row_index,
                )
            )
    return candidates


def _ocr_name_continuation(rows: list[list[dict[str, Any]]], anchor_index: int) -> str:
    parts: list[str] = []
    for index in [anchor_index - 1, anchor_index, anchor_index + 1]:
        if not 0 <= index < len(rows):
            continue
        row = rows[index]
        code_band = _ocr_band_text(row, 430, 570)
        name_band = _ocr_band_text(row, 570, 790)
        code, name = _split_ocr_code_and_name(code_band, name_band)
        if index != anchor_index and code:
            continue
        parts.append(name or name_band)
    return _clean_space("".join(parts))


def _split_ocr_code_and_name(code_band: str, name_band: str) -> tuple[str, str]:
    name_text = _clean_space(name_band)
    match = re.match(r"([0-9A-Za-z|!]+)(.*)", name_text)
    leading_code = match.group(1) if match else ""
    remaining_name = match.group(2) if match else name_text
    major_code = _normalize_ocr_major_code(_clean_space(code_band) + leading_code)
    if major_code:
        return major_code, _clean_space(remaining_name)
    major_code = _normalize_ocr_major_code(code_band)
    return major_code, name_text


def _ocr_code_band_prefix(code_band: str) -> str:
    match = re.match(r"([^0-9A-Za-z|!]+)", _clean_space(code_band))
    return match.group(1) if match else ""


def _extract_discipline_category(value: str) -> str:
    clean = _clean_space(re.sub(r"^[^一-龥]+", "", value))
    for category in sorted(DISCIPLINE_CATEGORIES, key=len, reverse=True):
        if category in clean:
            return category
    return clean if clean in DISCIPLINE_CATEGORIES else ""


def _split_ocr_degree_years(value: str) -> tuple[str, str]:
    clean = _clean_space(value).replace(" ", "")
    match = re.search(r"(二年|三年|四年|五年)", clean)
    if not match:
        return clean, ""
    degree = clean[: match.start()].strip("，,、")
    study_years = match.group(1)
    return degree, study_years


def _normalize_ocr_major_code(value: str) -> str:
    text = _clean_space(value).upper()
    digits = ""
    suffix = ""
    for char in text:
        additions = ""
        if char.isdigit():
            additions = char
        elif char in {"O", "Q", "D"}:
            additions = "0"
        elif char in {"I", "L", "|", "!"}:
            additions = "1"
        elif char == "H":
            additions = "11"
        elif char == "M":
            additions = "01"
        elif char in {"T", "K", "Y"} and len(digits) >= 4:
            suffix += char
            continue
        else:
            continue
        for digit in additions:
            if len(digits) < 6:
                digits += digit
    if len(digits) != 6:
        return ""
    suffix = "".join(char for char in suffix if char in {"T", "K", "Y"})[:2]
    return digits + suffix


def _group_ocr_rows(words: list[dict[str, Any]], tolerance: float = 8.0) -> list[list[dict[str, Any]]]:
    groups: list[list[dict[str, Any]]] = []
    centers: list[float] = []
    for word in sorted(words, key=lambda item: (_ocr_cy(item), float(item.get("x", 0)))):
        center = _ocr_cy(word)
        if groups and abs(center - centers[-1]) <= tolerance:
            groups[-1].append(word)
            centers[-1] = sum(_ocr_cy(item) for item in groups[-1]) / len(groups[-1])
        else:
            groups.append([word])
            centers.append(center)
    return [sorted(group, key=lambda item: float(item.get("x", 0))) for group in groups]


def _ocr_band_text(row: list[dict[str, Any]], x_min: float, x_max: float) -> str:
    return _clean_space(
        "".join(
            str(item.get("text") or "")
            for item in sorted(row, key=lambda value: float(value.get("x", 0)))
            if x_min <= _ocr_cx(item) < x_max
        )
    )


def _ocr_cx(item: dict[str, Any]) -> float:
    return float(item.get("x", 0)) + float(item.get("w", 0)) / 2


def _ocr_cy(item: dict[str, Any]) -> float:
    return float(item.get("y", 0)) + float(item.get("h", 0)) / 2


def _parse_catalog_table_pdf_line(line: str) -> tuple[str, str, str, str, str, str, str] | None:
    clean = _clean_space(line)
    if "专业代码" in clean or "序号" in clean:
        return None
    code_match = re.search(r"\b(\d{6}[A-ZKTY]*)\b", line)
    if not code_match:
        return None
    prefix = line[: code_match.start()].strip()
    suffix = line[code_match.end() :].strip()
    if not re.match(r"^\d+\b", prefix):
        return None

    prefix_parts = _split_pdf_table_cells(prefix)
    if prefix_parts:
        prefix_parts[0] = re.sub(r"^\d+\s*", "", prefix_parts[0]).strip()
        if not prefix_parts[0]:
            prefix_parts = prefix_parts[1:]
    prefix_tokens = _flatten_space_separated_tokens(prefix_parts)

    discipline_category = ""
    major_class = ""
    for token in prefix_tokens:
        if token in DISCIPLINE_CATEGORIES:
            discipline_category = token
        elif token.endswith("类"):
            major_class = token

    suffix_parts = _split_pdf_table_cells(suffix)
    if not suffix_parts:
        return None
    major_name = suffix_parts[0]
    major_name = _clean_space(re.sub(r"（注[:：]?.*?）", "", major_name))
    if not _looks_like_major_name(major_name):
        return None

    degree = ""
    study_years = ""
    for value in suffix_parts[1:]:
        if value.endswith("年"):
            study_years = study_years or value
        elif not re.fullmatch(r"\d{4}", value):
            degree = degree or value
    return code_match.group(1), major_name, discipline_category, major_class, degree, study_years, clean


def _split_pdf_table_cells(value: str) -> list[str]:
    return [_clean_space(part) for part in re.split(r"\s{2,}", value.strip()) if _clean_space(part)]


def _flatten_space_separated_tokens(values: list[str]) -> list[str]:
    tokens: list[str] = []
    for value in values:
        tokens.extend(part for part in value.split() if part)
    return tokens


def _parse_filing_pdf_line(line: str) -> tuple[str, str, str, str, str] | None:
    clean = _clean_space(line)
    if "专业代码" in clean or "序号" in clean:
        return None
    code_match = re.search(r"\b(\d{6}[A-ZKTY]*)\b", line)
    if not code_match:
        return None
    prefix = line[: code_match.start()].strip()
    suffix = line[code_match.end() :].strip()
    if not re.match(r"^\d+\b", prefix):
        return None
    prefix_parts = [part for part in re.split(r"\s{2,}", prefix) if part]
    if not prefix_parts:
        return None
    first_without_index = re.sub(r"^\d+\s*", "", prefix_parts[0]).strip()
    if first_without_index:
        prefix_parts[0] = first_without_index
    else:
        prefix_parts = prefix_parts[1:]
    if not prefix_parts:
        return None
    major_name = prefix_parts[-1]
    if not _looks_like_major_name(major_name):
        return None
    suffix_parts = suffix.split()
    degree = suffix_parts[0] if suffix_parts else ""
    study_years = suffix_parts[1] if len(suffix_parts) > 1 and suffix_parts[1].endswith("年") else ""
    return code_match.group(1), major_name, degree, study_years, clean


def _is_catalog_attachment(source_row: dict[str, str], attachment: dict[str, str]) -> bool:
    text = f"{source_row.get('source_type', '')} {attachment.get('attachment_title', '')}"
    return "目录" in text or "catalog" in text


def _build_source_document(
    source_row: dict[str, str],
    *,
    page_bytes: bytes,
    raw_path: Path,
    captured_at: str,
    attachment_count: int,
) -> dict[str, Any]:
    page_text = _decode_bytes(page_bytes)
    title = _extract_title(page_text) or source_row["title"]
    doc_id = "policy_doc:" + _sha256_text(source_row["url"])
    return {
        "schema_version": SOURCE_SCHEMA_VERSION,
        "doc_id": doc_id,
        "source_id": source_row["source_id"],
        "title": title,
        "url": source_row["url"],
        "source_domain": source_row["source_domain"],
        "source_level": source_row["source_level"],
        "source_type": source_row["source_type"],
        "issuing_org": source_row["issuing_org"],
        "published_date": source_row["published_date"],
        "source_year": source_row["source_year"],
        "captured_at": captured_at,
        "raw_path": str(raw_path),
        "content_sha256": hashlib.sha256(page_bytes).hexdigest(),
        "text_length": len(_clean_space(_strip_html(page_text))),
        "attachment_count": attachment_count,
    }


def _extract_links(html: str) -> list[tuple[str, str]]:
    try:
        from bs4 import BeautifulSoup
    except ImportError:  # pragma: no cover - BeautifulSoup exists in repo env.
        return re.findall(r"<a[^>]+href=[\"']([^\"']+)[\"'][^>]*>(.*?)</a>", html, flags=re.I | re.S)

    soup = BeautifulSoup(html, "html.parser")
    links: list[tuple[str, str]] = []
    for anchor in soup.find_all("a"):
        href = anchor.get("href") or ""
        text = anchor.get_text(" ", strip=True)
        links.append((href, text))
    return links


def _find_header(rows: list[list[str]]) -> tuple[int | None, dict[str, int]]:
    for index, row in enumerate(rows[:30]):
        header_map: dict[str, int] = {}
        for col_index, value in enumerate(row):
            normalized = value.replace(" ", "")
            if "专业代码" in normalized or normalized in {"代码", "专业代号"}:
                header_map.setdefault("major_code", col_index)
            if "专业名称" in normalized or normalized in {"专业名", "专业"}:
                header_map.setdefault("major_name", col_index)
            if normalized in {"门类", "学科门类"}:
                header_map.setdefault("discipline_category", col_index)
            if normalized in {"专业类", "所属专业类"}:
                header_map.setdefault("major_class", col_index)
            if "学位" in normalized:
                header_map.setdefault("degree", col_index)
            if "修业年限" in normalized or "学制" in normalized:
                header_map.setdefault("study_years", col_index)
        if "major_name" in header_map or "major_code" in header_map:
            return index, header_map
    return None, {}


def _value_by_header(row: list[str], header_map: dict[str, int], key: str) -> str:
    index = header_map.get(key)
    if index is None or index >= len(row):
        return ""
    return row[index]


def _looks_like_major_name(value: str) -> bool:
    value = _clean_space(value)
    if not value or len(value) < 2:
        return False
    if value in {"专业名称", "专业", "小计", "合计", "无"}:
        return False
    if re.fullmatch(r"\d+(\.\d+)?", value):
        return False
    return True


def _infer_event_type(source_row: dict[str, str], attachment: dict[str, str]) -> str:
    text = " ".join([source_row.get("source_type", ""), source_row.get("title", ""), attachment.get("attachment_title", "")])
    if "撤销" in text:
        return "revoked"
    if "annual_filing_approval" in text or "备案" in text:
        return "filing_added"
    if "审批" in text:
        return "approval_added"
    if "目录" in text:
        return "catalog_added"
    return "policy_candidate"


def _write_raw_bytes(
    raw_bytes: bytes,
    *,
    raw_dir: Path,
    run_id: str,
    url: str,
    preferred_file_type: str,
) -> Path:
    domain = (urlparse(url).hostname or "unknown").replace(":", "_")
    target_dir = raw_dir / run_id / domain
    target_dir.mkdir(parents=True, exist_ok=True)
    file_type = infer_file_type(url, raw_bytes=raw_bytes)
    if file_type == "html" and preferred_file_type:
        file_type = preferred_file_type
    suffix = f".{file_type}"
    target = target_dir / f"{_sha256_text(url)}{suffix}"
    target.write_bytes(raw_bytes)
    return target


def _parse_xls_with_conversion(
    path: Path,
    *,
    source_row: dict[str, str],
    attachment: dict[str, str],
    captured_at: str,
) -> tuple[list[dict[str, Any]], int]:
    with tempfile.TemporaryDirectory(prefix="major_intel_xls_") as temp_dir:
        converted_path = Path(temp_dir) / f"{path.stem}.xlsx"
        _convert_xls_to_xlsx(path, converted_path)
        row_count = _xlsx_row_count(converted_path)
        candidates = parse_xlsx_major_candidates(
            converted_path,
            source_row=source_row,
            attachment=attachment,
            captured_at=captured_at,
        )
    for candidate in candidates:
        candidate["raw_path"] = str(path)
        candidate["parsed_from"] = "xls"
    return candidates, row_count


def _parse_doc_with_conversion(
    path: Path,
    *,
    source_row: dict[str, str],
    attachment: dict[str, str],
    captured_at: str,
) -> tuple[list[dict[str, Any]], int]:
    with tempfile.TemporaryDirectory(prefix="major_intel_doc_") as temp_dir:
        converted_path = Path(temp_dir) / f"{path.stem}.docx"
        _convert_doc_to_docx(path, converted_path)
        row_count = _docx_row_count(converted_path)
        candidates = parse_docx_major_candidates(
            converted_path,
            source_row=source_row,
            attachment=attachment,
            captured_at=captured_at,
        )
    for candidate in candidates:
        candidate["raw_path"] = str(path)
        candidate["parsed_from"] = "doc"
    return candidates, row_count


def _convert_xls_to_xlsx(source_path: Path, target_path: Path) -> None:
    try:
        import pythoncom
        import win32com.client
    except ImportError as exc:  # pragma: no cover - platform dependent.
        raise RuntimeError("win32com is required to convert legacy xls attachments") from exc

    initialized = False
    excel = None
    workbook = None
    try:
        pythoncom.CoInitialize()
        initialized = True
        excel = win32com.client.DispatchEx("Excel.Application")
        excel.Visible = False
        excel.DisplayAlerts = False
        workbook = excel.Workbooks.Open(
            str(source_path.resolve()),
            UpdateLinks=0,
            ReadOnly=True,
            AddToMru=False,
        )
        workbook.SaveAs(str(target_path.resolve()), FileFormat=51)
    finally:
        if workbook is not None:
            workbook.Close(SaveChanges=False)
        if excel is not None:
            excel.Quit()
        if initialized:
            pythoncom.CoUninitialize()
    if not target_path.exists() or target_path.stat().st_size == 0:
        raise RuntimeError("xls conversion produced no xlsx output")


def _convert_doc_to_docx(source_path: Path, target_path: Path) -> None:
    try:
        import pythoncom
        import win32com.client
    except ImportError as exc:  # pragma: no cover - platform dependent.
        raise RuntimeError("win32com is required to convert legacy doc attachments") from exc

    initialized = False
    word = None
    document = None
    try:
        pythoncom.CoInitialize()
        initialized = True
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(
            str(source_path.resolve()),
            ConfirmConversions=False,
            ReadOnly=True,
            AddToRecentFiles=False,
            Visible=False,
        )
        document.SaveAs2(str(target_path.resolve()), FileFormat=16)
    finally:
        if document is not None:
            document.Close(SaveChanges=False)
        if word is not None:
            word.Quit()
        if initialized:
            pythoncom.CoUninitialize()
    if not target_path.exists() or target_path.stat().st_size == 0:
        raise RuntimeError("doc conversion produced no docx output")


def _xlsx_row_count(path: Path) -> int:
    try:
        from openpyxl import load_workbook

        workbook = load_workbook(path, read_only=True, data_only=True)
        count = sum(sheet.max_row or 0 for sheet in workbook.worksheets)
        workbook.close()
        return count
    except Exception:
        return 0


def _docx_row_count(path: Path) -> int:
    try:
        from docx import Document

        document = Document(str(path))
        table_rows = sum(len(table.rows) for table in document.tables)
        paragraph_rows = sum(1 for paragraph in document.paragraphs if _clean_space(paragraph.text))
        return table_rows + paragraph_rows
    except Exception:
        return 0


def _pdf_text_line_count(path: Path) -> int:
    text = extract_pdf_text(path)
    return sum(1 for line in text.splitlines() if line.strip())


def _candidate_csv_row(row: dict[str, Any]) -> dict[str, Any]:
    csv_row = {field: row.get(field, "") for field in CSV_FIELDS}
    csv_row["warnings_json"] = json.dumps(row.get("warnings", []), ensure_ascii=False)
    return csv_row


def _failure_row(source_row: dict[str, str], url: str, exc: Exception, *, failure_scope: str) -> dict[str, Any]:
    return {
        "source_id": source_row.get("source_id"),
        "url": url,
        "failure_scope": failure_scope,
        "error_type": type(exc).__name__,
        "error": str(exc),
    }


def _write_jsonl(handle, row: dict[str, Any]) -> None:
    handle.write(json.dumps(row, ensure_ascii=False) + "\n")


def _zip_family_file_type(raw_bytes: bytes) -> str:
    import zipfile

    try:
        with zipfile.ZipFile(__import__("io").BytesIO(raw_bytes)) as archive:
            names = set(archive.namelist())
    except zipfile.BadZipFile:
        return "zip"
    if any(name.startswith("xl/") for name in names):
        return "xlsx"
    if any(name.startswith("word/") for name in names):
        return "docx"
    return "zip"


def _extract_title(html: str) -> str:
    match = re.search(r"<title[^>]*>(.*?)</title>", html, flags=re.I | re.S)
    return _clean_space(_strip_html(match.group(1))) if match else ""


def _strip_html(value: str) -> str:
    return re.sub(r"<[^>]+>", " ", value or "")


def _decode_bytes(raw: bytes) -> str:
    for encoding in ("utf-8", "gb18030", "gbk"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _cell_text(value: Any) -> str:
    if value is None:
        return ""
    return _clean_space(str(value))


def _clean_space(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip()


def _normalize_major_code(value: str) -> str:
    return _clean_space(value).upper()


def _sha256_text(value: str) -> str:
    return hashlib.sha256(value.encode("utf-8")).hexdigest()


def _now_iso() -> str:
    return datetime.now(timezone.utc).astimezone().isoformat(timespec="seconds")


def _to_int(value: Any) -> int | None:
    try:
        return int(str(value))
    except (TypeError, ValueError):
        return None


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Crawl official emerging-major catalog sources.")
    subparsers = parser.add_subparsers(dest="command", required=True)

    crawl = subparsers.add_parser("crawl")
    crawl.add_argument("--source-csv", type=Path, default=Path("data/seeds/policy_document_sources.csv"))
    crawl.add_argument("--raw-dir", type=Path, default=Path("data/raw/policy_documents"))
    crawl.add_argument("--processed-dir", type=Path, default=Path("data/processed/policy_documents"))
    crawl.add_argument("--logs-dir", type=Path, default=Path("data/logs/policy_documents"))
    crawl.add_argument("--run-id", default="emerging_major_seed")
    crawl.add_argument("--timeout-seconds", type=float, default=20)
    crawl.add_argument("--sleep-seconds", type=float, default=0.0)

    args = parser.parse_args(argv)
    if args.command == "crawl":
        summary = crawl_emerging_major_catalog(
            source_csv=args.source_csv,
            raw_dir=args.raw_dir,
            processed_dir=args.processed_dir,
            logs_dir=args.logs_dir,
            run_id=args.run_id,
            timeout_seconds=args.timeout_seconds,
            sleep_seconds=args.sleep_seconds,
        )
        print(json.dumps(summary, ensure_ascii=False, indent=2))
        return 0
    return 2


if __name__ == "__main__":
    raise SystemExit(main())
