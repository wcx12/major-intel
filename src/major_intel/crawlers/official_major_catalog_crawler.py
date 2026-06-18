"""Crawl and parse official MOE major catalog documents."""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import re
import subprocess
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import requests


SCHEMA_VERSION = "official_major_catalog/v1"
REGISTER_PAGE = "https://zwfw.moe.gov.cn/zyyxzy/result.html"
DEFAULT_USER_AGENT = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
    "(KHTML, like Gecko) Chrome/126.0 Safari/537.36 major-intel/1.0"
)
RECORD_FIELDS = [
    "catalog_record_id",
    "catalog_year",
    "education_level",
    "major_level",
    "major_code",
    "major_name",
    "major_category_code",
    "major_category_name",
    "major_class_code",
    "major_class_name",
    "degree_type",
    "master_only",
    "source_level",
    "source_name",
    "source_url",
    "attachment_url",
    "captured_at",
]
DEFAULT_SOURCES = [
    {
        "source_id": "vocational_education_catalog_2021",
        "catalog_year": "2021",
        "source_name": "教育部关于印发《职业教育专业目录（2021年）》的通知",
        "source_url": "http://www.moe.gov.cn/srcsite/A07/moe_953/202103/t20210319_521135.html",
        "attachment_url": "http://www.moe.gov.cn/srcsite/A07/moe_953/202103/W020210319595911145604.docx",
        "attachment_name": "职业教育专业目录（2021年）.docx",
        "parser": "vocational_docx",
    },
    {
        "source_id": "graduate_catalog_2022",
        "catalog_year": "2022",
        "source_name": "国务院学位委员会 教育部关于印发《研究生教育学科专业目录（2022年）》的通知",
        "source_url": "http://www.moe.gov.cn/srcsite/A22/moe_833/202209/t20220914_660828.html",
        "attachment_url": "http://www.moe.gov.cn/srcsite/A22/moe_833/202209/W020220914572994461110.pdf",
        "attachment_name": "研究生教育学科专业目录（2022年）.pdf",
        "parser": "graduate_pdf",
    },
]


@dataclass(frozen=True)
class CatalogContext:
    catalog_year: str
    source_name: str
    source_url: str
    attachment_url: str
    captured_at: str


def crawl_official_major_catalogs(
    *,
    raw_dir: Path,
    processed_dir: Path,
    logs_dir: Path,
    reports_dir: Path,
    run_id: str,
    timeout_seconds: float = 60,
) -> dict[str, Any]:
    raw_run_dir = raw_dir / run_id
    raw_run_dir.mkdir(parents=True, exist_ok=True)
    processed_dir.mkdir(parents=True, exist_ok=True)
    logs_dir.mkdir(parents=True, exist_ok=True)
    reports_dir.mkdir(parents=True, exist_ok=True)

    captured_at = datetime.now(timezone.utc).astimezone().isoformat()
    records: list[dict[str, Any]] = []
    attachments: list[dict[str, Any]] = []
    failures: list[dict[str, Any]] = []

    for source in DEFAULT_SOURCES:
        source_id = source["source_id"]
        try:
            attachment_path = raw_run_dir / source["attachment_name"]
            download_file(source["attachment_url"], attachment_path, timeout_seconds=timeout_seconds)
            context = CatalogContext(
                catalog_year=source["catalog_year"],
                source_name=source["source_name"],
                source_url=source["source_url"],
                attachment_url=source["attachment_url"],
                captured_at=captured_at,
            )
            if source["parser"] == "vocational_docx":
                source_records = parse_vocational_catalog_docx(attachment_path, context)
            elif source["parser"] == "graduate_pdf":
                source_records = parse_graduate_catalog_pdf(attachment_path, context)
            else:
                raise ValueError(f"unsupported parser: {source['parser']}")
            records.extend(source_records)
            attachments.append(
                {
                    "source_id": source_id,
                    "attachment_url": source["attachment_url"],
                    "raw_path": str(attachment_path),
                    "record_count": len(source_records),
                    "status": "ok",
                }
            )
        except Exception as exc:  # pragma: no cover - covered by integration runs.
            failures.append({"source_id": source_id, "error": f"{type(exc).__name__}: {exc}"})

    records.sort(
        key=lambda row: (
            row["education_level"],
            row["catalog_year"],
            row["major_category_code"],
            row["major_class_code"],
            row["major_code"],
            row["major_name"],
        )
    )
    records_jsonl = processed_dir / f"official_major_catalog_records_{run_id}.jsonl"
    records_csv = processed_dir / f"official_major_catalog_records_{run_id}.csv"
    manifest_path = logs_dir / f"{run_id}_manifest.json"
    failures_path = logs_dir / f"{run_id}_failures.jsonl"
    attachments_path = processed_dir / f"official_major_catalog_attachments_{run_id}.jsonl"
    coverage_report = reports_dir / f"official_major_catalog_coverage_{run_id}.md"

    _write_jsonl(records_jsonl, records)
    _write_csv(records_csv, records, RECORD_FIELDS)
    _write_jsonl(attachments_path, attachments)
    _write_jsonl(failures_path, failures)
    _write_coverage_report(coverage_report, records, attachments, failures)

    summary = _coverage_summary(records)
    manifest = {
        "run_id": run_id,
        "source_count": len(DEFAULT_SOURCES),
        "attachment_count": len(attachments),
        "record_count": len(records),
        "failure_count": len(failures),
        "records_jsonl": str(records_jsonl),
        "records_csv": str(records_csv),
        "attachments_jsonl": str(attachments_path),
        "coverage_report": str(coverage_report),
        "coverage_summary": summary,
        "finished_at": datetime.now(timezone.utc).astimezone().isoformat(),
    }
    _write_json(manifest_path, manifest)
    return manifest


def download_file(url: str, target_path: Path, *, timeout_seconds: float = 60) -> None:
    response = requests.get(url, headers={"User-Agent": DEFAULT_USER_AGENT}, timeout=timeout_seconds)
    response.raise_for_status()
    target_path.parent.mkdir(parents=True, exist_ok=True)
    target_path.write_bytes(response.content)


def parse_vocational_catalog_docx(path: Path, context: CatalogContext) -> list[dict[str, Any]]:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency is present in the repo env.
        raise RuntimeError("python-docx is required to parse vocational catalog docx") from exc

    document = Document(path)
    table_specs = [
        ("secondary_vocational", "中等职业教育", "中职", 0),
        ("higher_vocational_associate", "高等职业教育专科", "高职专科", 1),
        ("vocational_undergraduate", "高等职业教育本科", "高职本科", 2),
    ]
    records: list[dict[str, Any]] = []
    for education_level, major_level, display_level, table_index in table_specs:
        if table_index >= len(document.tables):
            continue
        rows = [[_clean_cell(cell.text) for cell in row.cells] for row in document.tables[table_index].rows]
        records.extend(
            parse_vocational_catalog_rows(
                rows,
                education_level=education_level,
                major_level=major_level,
                display_level=display_level,
                context=context,
            )
        )
    return records


def parse_vocational_catalog_rows(
    rows: list[list[str]],
    *,
    education_level: str,
    major_level: str,
    display_level: str,
    context: CatalogContext,
) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    major_category_code = ""
    major_category_name = ""
    major_class_code = ""
    major_class_name = ""

    for row in rows:
        if len(row) < 3:
            continue
        first, code, name = row[0], row[1], row[2]
        if first == "序号" and code == "专业代码":
            continue
        category = _parse_vocational_category(first)
        if category and first == code == name:
            category_code, category_name = category
            if len(category_code) == 2:
                major_category_code = category_code
                major_category_name = category_name
            elif len(category_code) == 4:
                major_class_code = category_code
                major_class_name = category_name
            continue
        if not re.fullmatch(r"\d{6}K?", code or ""):
            continue
        record = _catalog_record(
            catalog_year=context.catalog_year,
            education_level=education_level,
            major_level=major_level,
            major_code=code,
            major_name=name,
            major_category_code=major_category_code,
            major_category_name=major_category_name,
            major_class_code=major_class_code,
            major_class_name=major_class_name,
            degree_type="",
            master_only=False,
            source_name=context.source_name,
            source_url=context.source_url,
            attachment_url=context.attachment_url,
            captured_at=context.captured_at,
            evidence_text=f"{display_level}|{major_category_code}{major_category_name}|{major_class_code}{major_class_name}|{code}{name}",
        )
        records.append(record)
    return records


def parse_graduate_catalog_pdf(path: Path, context: CatalogContext) -> list[dict[str, Any]]:
    return parse_graduate_catalog_text(extract_pdf_text(path), context)


def parse_graduate_catalog_text(text: str, context: CatalogContext) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    discipline_category_code = ""
    discipline_category_name = ""
    for raw_line in text.splitlines():
        line = _clean_cell(raw_line)
        if not line:
            continue
        category = re.fullmatch(r"(\d{2})\s+([\u4e00-\u9fff]+)", line)
        if category:
            discipline_category_code = category.group(1)
            discipline_category_name = category.group(2)
            continue
        major = re.fullmatch(r"(\d{4})\s+(.+)", line)
        if not major:
            continue
        code = major.group(1)
        official_name = major.group(2).strip()
        master_only = official_name.endswith("*")
        major_name = official_name.rstrip("*").strip()
        degree_type = "professional_degree_category" if code[2] >= "5" else "academic_first_level_discipline"
        records.append(
            _catalog_record(
                catalog_year=context.catalog_year,
                education_level="graduate",
                major_level="研究生教育",
                major_code=code,
                major_name=major_name,
                major_category_code=discipline_category_code,
                major_category_name=discipline_category_name,
                major_class_code="",
                major_class_name="",
                degree_type=degree_type,
                master_only=master_only,
                source_name=context.source_name,
                source_url=context.source_url,
                attachment_url=context.attachment_url,
                captured_at=context.captured_at,
                evidence_text=line,
            )
        )
    return records


def extract_pdf_text(path: Path) -> str:
    try:
        completed = subprocess.run(
            ["pdftotext", "-layout", str(path.resolve()), "-"],
            check=True,
            capture_output=True,
        )
        return completed.stdout.decode("utf-8", errors="replace")
    except (subprocess.CalledProcessError, FileNotFoundError):
        try:
            from pypdf import PdfReader
        except ImportError as exc:  # pragma: no cover - dependency is present in the repo env.
            raise RuntimeError("pdftotext or pypdf is required to parse graduate catalog pdf") from exc
        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)


def _catalog_record(
    *,
    catalog_year: str,
    education_level: str,
    major_level: str,
    major_code: str,
    major_name: str,
    major_category_code: str,
    major_category_name: str,
    major_class_code: str,
    major_class_name: str,
    degree_type: str,
    master_only: bool,
    source_name: str,
    source_url: str,
    attachment_url: str,
    captured_at: str,
    evidence_text: str,
) -> dict[str, Any]:
    identity = "|".join([catalog_year, education_level, major_code, major_name, major_category_code, major_class_code])
    return {
        "schema_version": SCHEMA_VERSION,
        "catalog_record_id": "official_major_catalog:" + hashlib.sha256(identity.encode("utf-8")).hexdigest()[:24],
        "catalog_year": catalog_year,
        "education_level": education_level,
        "major_level": major_level,
        "major_code": major_code,
        "major_name": major_name,
        "major_category_code": major_category_code,
        "major_category_name": major_category_name,
        "major_class_code": major_class_code,
        "major_class_name": major_class_name,
        "degree_type": degree_type,
        "master_only": bool(master_only),
        "source_level": "A",
        "source_type": "official_major_catalog",
        "source_name": source_name,
        "source_url": source_url,
        "attachment_url": attachment_url,
        "captured_at": captured_at,
        "evidence_text": evidence_text,
    }


def _parse_vocational_category(value: str) -> tuple[str, str] | None:
    match = re.fullmatch(r"(\d{4}|\d{2})(.+)", value.strip())
    if not match:
        return None
    return match.group(1), match.group(2).strip()


def _coverage_summary(records: list[dict[str, Any]]) -> list[dict[str, Any]]:
    grouped: dict[tuple[str, str], int] = {}
    for record in records:
        key = (str(record.get("education_level") or ""), str(record.get("major_level") or ""))
        grouped[key] = grouped.get(key, 0) + 1
    return [
        {"education_level": key[0], "major_level": key[1], "record_count": count}
        for key, count in sorted(grouped.items())
    ]


def _write_coverage_report(
    path: Path,
    records: list[dict[str, Any]],
    attachments: list[dict[str, Any]],
    failures: list[dict[str, Any]],
) -> None:
    lines = [
        "# 官方专业目录覆盖报告",
        "",
        f"- 目录专业记录数：{len(records)}",
        f"- 附件数：{len(attachments)}",
        f"- 失败数：{len(failures)}",
        "",
        "## 分层覆盖",
        "",
    ]
    for row in _coverage_summary(records):
        lines.append(f"- {row['major_level']}（{row['education_level']}）：{row['record_count']} 条")
    if attachments:
        lines.extend(["", "## 官方附件", ""])
        for attachment in attachments:
            lines.append(f"- {attachment['source_id']}：{attachment['record_count']} 条，{attachment['attachment_url']}")
    if failures:
        lines.extend(["", "## 失败项", ""])
        for failure in failures:
            lines.append(f"- {failure['source_id']}：{failure['error']}")
    lines.append("")
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text("\n".join(lines), encoding="utf-8")


def _write_json(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, ensure_ascii=False, indent=2), encoding="utf-8")


def _write_jsonl(path: Path, rows: list[dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8") as handle:
        for row in rows:
            handle.write(json.dumps(row, ensure_ascii=False, sort_keys=True) + "\n")


def _write_csv(path: Path, rows: list[dict[str, Any]], fields: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        for row in rows:
            writer.writerow(row)


def _clean_cell(value: Any) -> str:
    return re.sub(r"\s+", " ", "" if value is None else str(value)).strip()


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Crawl official MOE major catalog documents.")
    parser.add_argument("--raw-dir", type=Path, default=Path("data/raw/official_major_catalog"))
    parser.add_argument("--processed-dir", type=Path, default=Path("data/processed/official_major_catalog"))
    parser.add_argument("--logs-dir", type=Path, default=Path("data/logs/official_major_catalog"))
    parser.add_argument("--reports-dir", type=Path, default=Path("reports/official_major_catalog"))
    parser.add_argument("--run-id", default="official_major_catalog")
    parser.add_argument("--timeout-seconds", type=float, default=60)
    args = parser.parse_args(argv)
    manifest = crawl_official_major_catalogs(
        raw_dir=args.raw_dir,
        processed_dir=args.processed_dir,
        logs_dir=args.logs_dir,
        reports_dir=args.reports_dir,
        run_id=args.run_id,
        timeout_seconds=args.timeout_seconds,
    )
    print(json.dumps(manifest, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
