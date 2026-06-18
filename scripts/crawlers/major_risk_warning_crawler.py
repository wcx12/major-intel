"""Build a public major employment-warning dataset.

The dataset is intentionally evidence-first: raw source documents are fetched
and checksummed, while structured records keep their source id and confidence
level. The curated records below are limited to public pages/PDF text that can
be fetched without authentication.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import re
import shutil
import subprocess
import tempfile
import zipfile
from collections import Counter, defaultdict
from dataclasses import dataclass
from datetime import datetime, timezone
from io import StringIO
from pathlib import Path
from typing import Any
from urllib.parse import urlparse

import pandas as pd
import requests
from bs4 import BeautifulSoup
from requests.packages.urllib3.exceptions import InsecureRequestWarning


requests.packages.urllib3.disable_warnings(InsecureRequestWarning)


ROOT = Path(__file__).resolve().parents[2]
DEFAULT_RAW_DIR = ROOT / "data/raw/major_risk_warnings"
DEFAULT_OUTPUT_DIR = ROOT / "data/processed/major_risk_warnings"
DEFAULT_REPORT_DIR = ROOT / "reports/major_risk_warnings"
CATALOG_PATH = ROOT / "data/seeds/rysxai_professions.full.csv"

SCHEMA_VERSION = "major_employment_warning/v1"
CAPTURED_AT = datetime.now(timezone.utc).astimezone().isoformat(timespec="seconds")

SOURCE_FIELDS = [
    "source_id",
    "title",
    "url",
    "publisher",
    "published_date",
    "source_type",
    "coverage_note",
    "status",
    "http_status",
    "content_type",
    "content_length",
    "sha256",
    "raw_path",
    "text_path",
    "fetched_at",
    "error",
]

RECORD_FIELDS = [
    "record_id",
    "schema_version",
    "report_year",
    "graduate_cohort",
    "education_level",
    "risk_level",
    "reported_major_name",
    "standard_major_name",
    "major_code",
    "discipline",
    "major_category",
    "source_ids",
    "evidence_type",
    "evidence_text",
    "confidence",
    "notes",
    "captured_at",
]

METRIC_FIELDS = [
    "metric_id",
    "schema_version",
    "report_year",
    "graduate_cohort",
    "education_level",
    "reported_major_name",
    "standard_major_name",
    "major_code",
    "metric_name",
    "metric_value",
    "metric_unit",
    "metric_rank",
    "rank_scope",
    "source_ids",
    "evidence_text",
    "confidence",
    "captured_at",
]

OFFICIAL_WARNING_FIELDS = [
    "warning_id",
    "schema_version",
    "policy_year",
    "region",
    "education_level",
    "record_type",
    "warning_label",
    "reported_major_name",
    "standard_major_name",
    "major_code",
    "discipline",
    "major_category",
    "policy_action",
    "criterion_text",
    "source_row_no",
    "study_duration",
    "source_ids",
    "evidence_text",
    "confidence",
    "captured_at",
]


@dataclass(frozen=True)
class SourceSeed:
    source_id: str
    title: str
    url: str
    publisher: str
    published_date: str
    source_type: str
    coverage_note: str


SOURCES: list[SourceSeed] = [
    SourceSeed(
        "edu_cn_2011_red",
        "本科和高职高专就业红绿牌专业名单出炉",
        "https://www.edu.cn/edu/jiuye/jiu_ye_xin_wen/201106/t20110609_631368.shtml",
        "中国教育和科研计算机网/新华网",
        "2011-06-09",
        "news_html",
        "2011本科、高职高专红牌；说明这些专业就业率不好具有持续性。",
    ),
    SourceSeed(
        "sina_2012_red_green",
        "动画、法学、英语等9个本科专业因高失业风险被红牌警告",
        "https://news.sina.cn/sa/2012-06-12/detail-ikmxzfmk1054765.d.html",
        "新浪新闻",
        "2012-06-12",
        "news_html",
        "2012本科、高职高专红牌；部分本科绿牌。",
    ),
    SourceSeed(
        "edu_cn_2012_red_green",
        "2012年本科就业“红牌警告”专业",
        "https://www.edu.cn/edu/jiuye/jiu_ye_xin_wen/201206/t20120612_789760_2.shtml",
        "中国教育和科研计算机网",
        "2012-06-12",
        "news_html",
        "2012本科、高职高专红牌和绿牌公开文本。",
    ),
    SourceSeed(
        "gkzxw_2012_undergrad_yellow",
        "2012年高失业风险专业遭遇红牌警告",
        "https://gkzxw.com/gxzs/201206/1127423.html",
        "高考资讯网转载",
        "2012-06",
        "news_html",
        "2012本科黄牌专业公开文本；原文带“等”，仅结构化明确列名专业。",
    ),
    SourceSeed(
        "edu_cn_2013_undergrad_yellow_green",
        "2013年就业“红黄绿牌”专业",
        "https://www.edu.cn/jiu_ye_xin_wen_11362/20130610/t20130610_960286.shtml",
        "中国教育和科研计算机网/新华网-北京晚报",
        "2013-06-10",
        "news_html",
        "2013本科红牌、黄牌、绿牌公开文本。",
    ),
    SourceSeed(
        "haedu_2014_red_yellow_green",
        "高考哪些专业就业前景看好?",
        "https://gaokao.haedu.cn/zyjd/2015/0313/107038.html",
        "河南高考信息网",
        "2015-03-13",
        "news_html",
        "2014本科、高职高专红牌、黄牌、绿牌公开文本。",
    ),
    SourceSeed(
        "edu_cn_2015_red_yellow_green",
        "本科和高职就业红牌专业：法学连续亮红",
        "https://www.edu.cn/edu/jiuye/jiu_ye_xin_wen/201507/t20150721_1290946.shtml",
        "中国教育和科研计算机网/北京青年报",
        "2015-07-21",
        "news_html",
        "2015本科、高职高专红牌、黄牌、绿牌公开文本。",
    ),
    SourceSeed(
        "sina_2016_red_yellow_green",
        "彻底惊呆：大学本科“最没用”六大专业",
        "https://edu.sina.com.cn/zl/edu/2016-06-16/10233705.shtml",
        "新浪教育",
        "2016-06-16",
        "news_html",
        "2016本科、高职高专红牌、黄牌、绿牌公开文本。",
    ),
    SourceSeed(
        "whysw_2018_yellow",
        "高考志愿填报四大注意事项",
        "https://www.whysw.org/m/view.php?aid=12457",
        "文化艺术网/文化艺术报",
        "2018",
        "news_html",
        "2018本科、高职高专黄牌专业公开文本；仅结构化正文明确列名专业。",
    ),
    SourceSeed(
        "sczk_2017_archive_red_yellow_green",
        "哪些大学专业被亮了红牌？",
        "https://sczk.com.cn/News/15c583d093ea4c1ab18135326cf4a52a.html",
        "四川招考网转载中国青年报经济部",
        "2017-07-03",
        "news_html",
        "HTML表格汇总2009-2017大学红牌、黄牌、绿牌专业；作为二级归档源补齐早年黄牌和部分高职绿牌。",
    ),
    SourceSeed(
        "sina_2017_undergrad_yellow",
        "高考志愿填报：2017年本科就业的五大黄牌专业",
        "https://edu.sina.cn/gaokao/bkzn/2017-06-17/detail-ifyhfnqa4362128.d.html?vt=4",
        "新浪教育",
        "2017-06-17",
        "news_html",
        "2017本科就业黄牌专业五个专业的可抓取文本。",
    ),
    SourceSeed(
        "eol_2020_undergrad_red_2010_2019",
        "注意！这些曾经的热门专业竟惨遭教育部撤销",
        "https://www.eol.cn/shuju/uni/202003/t20200325_1718176.shtml",
        "中国教育在线",
        "2020-03-25",
        "news_html",
        "附录列出2010-2019本科红牌专业。",
    ),
    SourceSeed(
        "sohu_2017_archive_highvoc_red_2010_2016",
        "注意啦！高校绿牌专业和红牌专业出炉",
        "https://www.sohu.com/a/148680659_200190",
        "搜狐转载",
        "2017-06-14",
        "news_html",
        "整理2010-2016高职高专红牌、2015-2016黄牌等历史资料。",
    ),
    SourceSeed(
        "thepaper_2017_red_green",
        "2017大学生就业报告：音乐表演等连续三届成本科红牌专业",
        "https://www.thepaper.cn/newsDetail_forward_1707751",
        "澎湃新闻",
        "2017-06-12",
        "news_html",
        "2017本科、高职高专红牌和绿牌。",
    ),
    SourceSeed(
        "sina_2018_red",
        "2018大学生就业报告：2017本科毕业生月薪4774",
        "https://edu.sina.cn/gaokao/gkrx/2018-06-12/detail-ihcufqih0591640.d.html",
        "新浪教育",
        "2018-06-12",
        "news_html",
        "2018本科、高职高专红牌。",
    ),
    SourceSeed(
        "people_2019_red_green",
        "2019中国大学生就业报告：新一线城市吸引力不断增强",
        "https://edu.people.com.cn/n1/2019/0613/c1053-31134703.html",
        "人民网教育",
        "2019-06-13",
        "news_html",
        "2019本科、高职高专红牌和绿牌。",
    ),
    SourceSeed(
        "ms315_2019_undergrad_yellow",
        "2019本科就业红牌专业，多个艺术类专业上榜！",
        "https://www.ms315.com/html/20190803/201908031927421.htm",
        "中国美术高考网/人民日报转载",
        "2019-08-03",
        "news_html",
        "2019本科红牌、黄牌、绿牌公开文本；用于补齐本科黄牌专业。",
    ),
    SourceSeed(
        "sdwm_2020_red_green",
        "2020年中国大学生就业报告发布",
        "https://ckjrx.sdwm.edu.cn/info/1147/2466.htm",
        "山东外贸职业学院财会金融系转载",
        "2020-07-09",
        "news_html",
        "2020本科、高职红牌和绿牌。",
    ),
    SourceSeed(
        "sohu_myc_2022_red_green",
        "重磅！近5年红绿牌专业，全揭晓！",
        "https://www.sohu.com/a/557522813_121294",
        "麦可思研究/搜狐",
        "2022-06-15",
        "news_html",
        "2022本科、高职红牌和绿牌；近五年连续上榜提示。",
    ),
    SourceSeed(
        "zjnu_2023_red_green",
        "2023年红绿牌专业，公布！",
        "https://fgc.zjnu.edu.cn/2023/1208/c16734a454744/page.htm",
        "浙江师范大学发展规划处转载",
        "2023-12-08",
        "news_html",
        "2023本科、高职红牌和绿牌；红黄绿牌定义。",
    ),
    SourceSeed(
        "bluebook_pdf_2023_undergrad",
        "就业蓝皮书：2023年中国本科生就业报告",
        "https://xiuzhenorgweb.oss-cn-zhangjiakou.aliyuncs.com/uploads/files/20231127/cc79a43039780361b3bfe6540f797713.pdf",
        "社会科学文献出版社/麦可思",
        "2023",
        "pdf",
        "2023本科红黄绿牌；近五年本科绿牌表；可用pdftotext抽取。",
    ),
    SourceSeed(
        "scribd_2024_undergrad_text",
        "2024年中国本科生就业报告文本摘录",
        "https://www.scribd.com/document/811172311/2024%E5%B9%B4%E4%B8%AD%E5%9B%BD%E6%9C%AC%E7%A7%91%E7%94%9F%E5%B0%B1%E4%B8%9A%E6%8A%A5%E5%91%8A-%E9%BA%A6%E5%8F%AF%E6%80%9D%E7%A0%94%E7%A9%B6%E9%99%A2-Z-Library",
        "Scribd文本页",
        "2024",
        "web_text_excerpt",
        "2024本科红黄绿牌表的公开文本摘录；版权页不作为全文源。",
    ),
    SourceSeed(
        "china_2025_undergrad_yellow",
        "本科专业就业分“红黄绿牌”，美术学绿变黄，法学为啥年年红牌？",
        "https://3g.china.com/act/edu/13004042/20250110/47851550_1.html",
        "中华网教育",
        "2025-01-10",
        "news_html",
        "2024本科黄牌专业及2010-2024累计黄牌次数；用于替代Scribd受阻文本页。",
    ),
    SourceSeed(
        "labor_2024_red_green",
        "报告：大学生薪资增速明显放缓，工科春天到来？",
        "https://search.laborinfocn.com/articles/52253",
        "工劳网转载/检索",
        "2024-07",
        "news_html",
        "2024本科和高职红牌、绿牌摘要。",
    ),
    SourceSeed(
        "people_2025_green",
        "2025年绿牌专业榜，发布！",
        "https://www.peopleapp.com/rmharticle/30049323357",
        "人民日报客户端/教育在线",
        "2025-06-12",
        "news_html",
        "2025本科、高职绿牌；数据样本说明和部分就业质量指标。",
    ),
    SourceSeed(
        "eol_2025_undergrad_red",
        "2025年本科红牌专业公布！",
        "https://www.eol.cn/news/yaowen/202506/t20250630_2678129.shtml",
        "中国教育在线/麦可思研究",
        "2025-06-30",
        "news_html",
        "2025本科红牌；近五年本科红牌连续上榜说明。",
    ),
    SourceSeed(
        "sohu_2025_highvoc_red",
        "2025年高职红牌专业，发布！",
        "https://www.sohu.com/a/907025131_121294",
        "麦可思研究/搜狐",
        "2025-06-23",
        "news_html",
        "2025高职红牌；近五年高职红牌连续上榜说明和样本量。",
    ),
    SourceSeed(
        "gkztc_2025_undergrad_yellow",
        "2025年本科红黄绿牌专业揭晓",
        "https://app.gaokaozhitongche.com/news/h/QzeBkMPO",
        "高考直通车",
        "2025-11-24",
        "news_html",
        "2025本科黄牌、红牌、绿牌公开整理。",
    ),
    SourceSeed(
        "sina_2026_undergrad_green",
        "2026年本科绿牌专业都是工科，高学历人才涌入制造业",
        "https://finance.sina.cn/2026-06-11/detail-iniazihs2356736.d.html",
        "新浪财经转载",
        "2026-06-11",
        "news_html",
        "2026本科绿牌；2025届本科样本量和就业流向摘要。",
    ),
    SourceSeed(
        "gkztc_2026_green",
        "2026年绿牌专业榜，发布！",
        "https://app.gaokaozhitongche.com/news/h/jOv3kQLw",
        "高考直通车/麦可思研究",
        "2026-06-12",
        "news_html",
        "2026本科、高职绿牌；包含电气工程及其自动化就业质量指标。",
    ),
    SourceSeed(
        "peopleapp_2026_undergrad_stop_summary",
        "70所本科，官宣停招专业！",
        "https://www.peopleapp.com/rmharticle/30052250494",
        "人民日报客户端人民号 / 教育在线",
        "2026-05-29",
        "news_html",
        "二级统计源，根据高校信息公开网、最新招生专业目录等统计70所本科高校最新公布停招专业名单，共涉及525个本科专业，并列出停招数量较多的专业频次。",
    ),
    SourceSeed(
        "peopleapp_2026_undergrad_stop_frequency_image",
        "停招数量较多的本科专业（部分）",
        "https://cdnjdphoto.aikan.pdnews.cn/zhbj-20260528/image/content/bc7ac717cce4469fba24eeecd608e364.webp",
        "人民日报客户端人民号 / 教育在线",
        "2026-05-29",
        "secondary_news_image",
        "人民号文章内嵌长图，列出40个停招数量较多本科专业及停招高校数量、专业类。",
    ),
    SourceSeed(
        "qq_2026_undergrad_stop_summary",
        "70所本科，官宣停招专业！（附榜单）",
        "https://news.qq.com/rain/a/20260525A08HJW00",
        "腾讯新闻 / 掌上高考",
        "2026-05-25",
        "secondary_news_html",
        "掌上高考在腾讯新闻发布的原始传播页，统计70所本科高校最新公布停招专业名单、525个本科专业实例，并附停招数量较多本科专业榜单图片。",
    ),
    SourceSeed(
        "qq_2026_undergrad_stop_frequency_image_part1",
        "停招数量较多的本科专业图片（上）",
        "https://inews.gtimg.com/om_bt/OtYHKVP6p4nlol5-WildyDJBGtkkqYleLGKAuC-Bw0IXMAA/641",
        "腾讯新闻 / 掌上高考",
        "2026-05-25",
        "secondary_news_image",
        "掌上高考腾讯新闻页内嵌榜单图片上半部分，列出停招数量从16所至6所的本科专业。",
    ),
    SourceSeed(
        "qq_2026_undergrad_stop_frequency_image_part2",
        "停招数量较多的本科专业图片（下）",
        "https://inews.gtimg.com/om_bt/OJrPKjj9qFC2_Vmdy3o27nP5eEVB9AH0mvRDjE9xJ87hAAA/641",
        "腾讯新闻 / 掌上高考",
        "2026-05-25",
        "secondary_news_image",
        "掌上高考腾讯新闻页内嵌榜单图片下半部分，列出停招数量从6所至4所的本科专业。",
    ),
    SourceSeed(
        "ycwb_2021_highvoc_green",
        "2021年高职高专就业“绿牌专业”话你知",
        "https://ep.ycwb.com/epaper/xkb/h5/html5/2022-03/10/content_2580_477753.htm",
        "新快报",
        "2022-03-10",
        "news_html",
        "2021高职高专绿牌；2020届高职高薪专业和部分绿牌专业就业质量指标。",
    ),
    SourceSeed(
        "people_2025_green_income",
        "哪些专业就业好、收入高？",
        "https://www.peopleapp.com/rmharticle/30049394366",
        "人民日报客户端/教育在线",
        "2025-06-22",
        "news_html",
        "2025本科、高职绿牌；2024届本科高薪前十和部分绿牌就业质量指标。",
    ),
    SourceSeed(
        "edu_guizhou_2012_undergrad_warning",
        "贵州开始实施普通高校本科专业预警及退出机制",
        "https://www.edu.cn/ke_yan_yu_fa_zhan/gao_xiao_cheng_guo/gao_xiao_zi_xun/201207/t20120727_817328.shtml",
        "中国教育和科研计算机网/中国新闻网",
        "2012-07-27",
        "news_html",
        "贵州省本科专业预警名单与退出机制；原省教育厅页面已迁移/失效，使用可抓取转载页。",
    ),
    SourceSeed(
        "eol_guizhou_2021_undergrad_warning",
        "5个专业“亮红灯”！2021年贵州普通高校本科预警专业发布",
        "https://news.eol.cn/yaowen/202111/t20211103_2171351.shtml",
        "中国教育在线转载贵州教育发布",
        "2021-11-03",
        "news_html",
        "2021年贵州省普通本科高校专业预警名单，列出法学、数字媒体艺术、劳动与社会保障、公共事业管理、汉语言文学。",
    ),
    SourceSeed(
        "cctv_2016_shanghai_low_employment_warning",
        "预警 | 多年来这些本科专业屡被亮红灯 看看你的专业是否上榜",
        "https://m.news.cctv.com/2016/06/12/ARTIOXb7Dpc8vMXnZqAMuCAq160612.shtml",
        "央视网/央视新闻/新华社",
        "2016-06-12",
        "news_html",
        "上海2012-2016本科预警专业、2014全国就业率较低本科专业和各省市低就业率本科专业名单。",
    ),
    SourceSeed(
        "edu_cn_2014_low_employment_warning",
        "教育部公布低就业率本科专业名单 15专业被亮红牌",
        "https://www.edu.cn/edu/jiao_yu_bu/xin_wen_dong_tai/201410/t20141014_1189131.shtml",
        "中国教育和科研计算机网/新华网-大河报",
        "2014-10-14",
        "news_html",
        "教育部2014年公布近两年全国15个就业率较低本科专业，并列出北京至江苏10个地区的低就业率本科专业名单；作为央视/阳光高考名单的可抓取补充证据。",
    ),
    SourceSeed(
        "edu_cn_2014_low_employment_warning_page2",
        "教育部公布低就业率本科专业名单 15专业被亮红牌（分页2）",
        "https://www.edu.cn/edu/jiao_yu_bu/xin_wen_dong_tai/201410/t20141014_1189131_1.shtml",
        "中国教育和科研计算机网/新华网-大河报",
        "2014-10-14",
        "news_html",
        "教育部2014年公布近两年各地就业率较低本科专业名单分页，列出浙江至新疆生产建设兵团22个地区的低就业率本科专业名单。",
    ),
    SourceSeed(
        "shanghai_2012_undergrad_warning_doc",
        "上海市教育委员会关于2012年度对部分本科专业实施预警的意见",
        "https://edu.sh.gov.cn/cmsres/84/8443f41e939c45219faf0469814de858/8408eb135f049a169d22ef485f2d78e5.doc",
        "上海市教育委员会",
        "2012-02-02",
        "official_policy_doc",
        "上海市教委沪教委高〔2012〕4号附件，正文列出2012年度18个本科预警专业。",
    ),
    SourceSeed(
        "shanghai_2013_undergrad_warning_doc",
        "上海市教育委员会关于2013年度本科预警专业名单及相关事项的通知",
        "https://edu.sh.gov.cn/cmsres/6b/6b2821bcac014201836d6f3d77763f01/c61edd234625d9c1cedb4c86381896fc.doc",
        "上海市教育委员会",
        "2013-04-22",
        "official_policy_doc",
        "上海市教委沪教委高〔2013〕9号附件，正文列出2013年度15个本科预警专业。",
    ),
    SourceSeed(
        "shanghai_2014_undergrad_warning_doc",
        "上海市教育委员会关于公布2014年度本科预警专业名单的通知",
        "https://edu.sh.gov.cn/cmsres/97/9709d027fff64731b9d9e556f4993f55/07c1a477c7a0cf51f6a93aa56947e459.doc",
        "上海市教育委员会",
        "2014-04-01",
        "official_policy_doc",
        "上海市教委沪教委高〔2014〕11号附件，正文列出2014年度7个本科预警专业。",
    ),
    SourceSeed(
        "moe_2014_shanghai_warning_news",
        "上海公布本科预警专业名单 今年7个专业被亮黄牌",
        "https://www.moe.gov.cn/jyb_xwfb/s5147/201404/t20140403_166648.html",
        "中华人民共和国教育部",
        "2014-04-03",
        "official_news_html",
        "教育部政府门户网站转载中国教育报报道，确认上海市教委公布2014年度7个本科预警专业及预警依据。",
    ),
    SourceSeed(
        "shanghai_2016_undergrad_warning_doc",
        "上海市教育委员会关于公布2016年度本科预警专业名单的通知",
        "https://edu.sh.gov.cn/cmsres/46/469b46db6c384ebe942fa27b0361a4d3/940b741e26df2b28746271c95ab97a0b.doc",
        "上海市教育委员会",
        "2016-06-08",
        "official_policy_doc",
        "上海市教委沪教委高〔2016〕40号附件，正文列出2016年度10个本科预警专业。",
    ),
    SourceSeed(
        "caijing_2016_provincial_warning",
        "填报志愿要谨慎！粤沪等地发布本科专业预警名单",
        "https://m.21jingji.com/article/20160624/herald/753aa8fecd2931196a4bcdf81741939e.html",
        "21世纪经济报道",
        "2016-06-24",
        "news_html",
        "2016上海、广东、辽宁、河北等地本科预警专业和山东低就业率专业报道；广东、河北、山东为完整列名，辽宁仅结构化正文明确列出的部分专业。",
    ),
    SourceSeed(
        "liaoning_kjt_2016_undergrad_negative_list",
        "省内高校将按东北振兴需要增减专业",
        "https://kjt.ln.gov.cn/kjt/dfkj/5C2499A394DF4111805D90D4A67948A8/index.shtml",
        "辽宁省科学技术厅转载沈阳日报/沈阳网",
        "2016-05-19",
        "government_news_html",
        "辽宁省科技厅官网转载报道称辽宁省教育厅印发通知，建议高校2016年度暂缓申请增设66种本科专业，并列出完整名单。",
    ),
    SourceSeed(
        "hebei_2019_undergrad_application_policy_pdf_mirror",
        "河北省教育厅关于2019年度普通高等学校本科专业申报工作的通知附件",
        "https://jwc.hebeinu.edu.cn/upload/file/20190712/1562924199107790.pdf",
        "河北省教育厅（河北北方学院教务处镜像）",
        "2019-07-12",
        "official_policy_pdf_mirror",
        "河北省2019年度本科专业申报通知附件，含2018届本科毕业生分专业初次就业率统计表、急需本科专业名单和专业布点情况统计表。",
    ),
    SourceSeed(
        "hebtu_2019_performance_target_pdf",
        "河北师范大学2019年度职能绩效管理目标一览表",
        "https://fzghc.hebtu.edu.cn/dynamic/download.jsp?id=4c6a285e5ee84b789f71dc63bb9143d5",
        "河北师范大学发展规划处",
        "2020-01-02",
        "official_university_pdf",
        "河北师范大学2019年度职能绩效管理目标一览表说明撤销秘书学等10个本科专业并已上报教育部备案审批。",
    ),
    SourceSeed(
        "aufe_2019_undergrad_stop_list_html",
        "安徽财经大学本科专业设置、当年新增专业、停招专业名单信息一览表",
        "https://xxgk2019.aufe.edu.cn/2019/1015/c10245a138075/page.htm",
        "安徽财经大学信息公开网",
        "2019-10-15",
        "official_university_html",
        "安徽财经大学2019年本科专业设置表，逐项列出专业代码、学院、设置年份、学位门类、招生状态；含14个招生状态为“当年停招”的本科专业或方向。",
    ),
    SourceSeed(
        "aufe_2019_cancel_notice_page",
        "安徽财经大学关于撤销相关本科专业（方向）的通知",
        "https://xxgk2019.aufe.edu.cn/2023/1007/c10245a203984/page.htm",
        "安徽财经大学信息公开网",
        "2019-03-29",
        "official_university_html",
        "安徽财经大学信息公开网发布的校政字〔2019〕39号撤销本科专业（方向）通知页面，含官方PDF附件。",
    ),
    SourceSeed(
        "aufe_2019_cancel_notice_pdf",
        "安徽财经大学关于撤销相关本科专业（方向）的通知PDF",
        "https://xxgk2019.aufe.edu.cn/_upload/article/files/75/ee/9354818d42fa940c8f67cf62d132/63e236bd-0865-4fae-a59e-84d9fd694392.pdf",
        "安徽财经大学",
        "2019-03-26",
        "official_university_pdf",
        "安徽财经大学校政字〔2019〕39号扫描PDF，说明对公共事业管理等13个本科专业（方向）实施撤销。",
    ),
    SourceSeed(
        "henan_2020_undergrad_warning_policy",
        "河南省本科高校学科专业结构优化调整指导意见",
        "https://m.jyt.henan.gov.cn/2020/11-30/1912063.html",
        "河南省教育厅",
        "2020-11-30",
        "official_policy_html",
        "河南本科专业供需预警、负面清单、黄牌红牌预警和退出机制政策。",
    ),
    SourceSeed(
        "edu_moe_2025_undergrad_setup_notice",
        "教育部：加快布局急需紧缺专业",
        "https://www.edu.cn/rd/gao_xiao_cheng_guo/gao_xiao_zi_xun/202507/t20250702_2678614.shtml",
        "中国教育和科研计算机网/教育部高等教育司",
        "2025-07-02",
        "official_policy_html",
        "2025年度本科专业设置通知，要求省级发布急需本科专业清单和过剩专业预警清单。",
    ),
    SourceSeed(
        "moe_2025_2024_undergrad_cancel_notice_pdf",
        "教育部关于备案2024年度普通高等学校申请撤销的本科专业的通知",
        "https://jwc.hebeinu.edu.cn/upload/file/20250429/1745915356983921.pdf",
        "教育部文件（河北北方学院教务处镜像PDF）",
        "2025-04-01",
        "official_policy_pdf_mirror",
        "教育部教高函〔2025〕5号附件，列出2024年度普通高等学校撤销本科专业名单；PDF同时含教高函〔2025〕3号备案审批结果和2025年专业目录。",
    ),
    SourceSeed(
        "moe_2025_2024_undergrad_setting_result_news",
        "教育部公布2024年度普通高等学校本科专业备案和审批结果",
        "https://qspfw.moe.gov.cn/html/hotnews/20250427/22933.html",
        "教育部政务服务平台",
        "2025-04-27",
        "official_policy_html",
        "教育部政务服务平台新闻页，汇总2024年度普通高等学校本科专业新增、调整、停招、撤销专业点数，以及2025版本科专业目录规模。",
    ),
    SourceSeed(
        "moe_2025_2024_undergrad_setting_qna",
        "教育部高等教育司负责人就2024年度普通高等学校本科专业设置工作答记者问",
        "https://qspfw.moe.gov.cn/html/edudoc/20250526/23028.html",
        "教育部政务服务平台",
        "2025-05-26",
        "official_policy_html",
        "教育部政务服务平台答记者问，说明2024年度新增、调整、撤销、停招专业点数量，以及区域匹配度试点、省域专业调整比例和重点产业支撑指标。",
    ),
    SourceSeed(
        "moe_2026_undergrad_catalog_notice",
        "教育部关于公布《普通高等学校本科专业目录（2026年）》的通知",
        "http://www.moe.gov.cn/srcsite/A08/moe_1034/s3882/202604/t20260427_1434931.html",
        "中华人民共和国教育部",
        "2026-04-28",
        "official_policy_html",
        "教育部官网通知，说明组织开展2025年度普通高等学校本科专业设置和调整工作，并公布2026年本科专业目录PDF附件。",
    ),
    SourceSeed(
        "moe_2026_undergrad_catalog_pdf",
        "普通高等学校本科专业目录（2026年）",
        "http://www.moe.gov.cn/srcsite/A08/moe_1034/s3882/202604/W020260427440749576927.pdf",
        "中华人民共和国教育部",
        "2026-04-28",
        "official_policy_pdf",
        "教育部官网PDF附件，列出2026年普通高等学校本科专业目录，包括交叉学科门类下未来机器人、具身智能、脑机科学与技术等专业。",
    ),
    SourceSeed(
        "moe_2026_undergrad_catalog_release_news",
        "《普通高等学校本科专业目录（2026年）》发布",
        "http://www.moe.gov.cn/jyb_xwfb/gzdt_gzdt/s5987/202604/t20260428_1435016.html",
        "中华人民共和国教育部",
        "2026-04-28",
        "official_news_html",
        "教育部官网新闻，汇总2026年本科专业目录规模、交叉学科首批专业、十四五期间新增与撤销/停招专业布点数量以及年度专业调整比例。",
    ),
    SourceSeed(
        "moe_2026_vocational_specialty_setup_results_notice",
        "教育部关于公布2026年高等职业教育专科专业设置备案和审批结果的通知",
        "http://www.moe.gov.cn/srcsite/A07/moe_953/202604/t20260422_1434454.html",
        "中华人民共和国教育部",
        "2026-04-22",
        "official_policy_html",
        "教育部官网通知，公布2026年高等职业教育专科专业设置备案和审批结果，说明拟招生专业点69414个、国家控制布点专业申请12个、同意8个、不同意4个。",
    ),
    SourceSeed(
        "moe_2026_vocational_controlled_specialty_approval_pdf",
        "2026年新设高职专科国家控制布点专业审批结果",
        "http://www.moe.gov.cn/srcsite/A07/moe_953/202604/W020260422371180700581.pdf",
        "中华人民共和国教育部",
        "2026-04-22",
        "official_policy_pdf",
        "教育部官网PDF附件，逐项列出2026年新设高职专科国家控制布点专业审批结果，包括8个同意设置和4个不同意设置的专业点。",
    ),
    SourceSeed(
        "moe_2024_2023_undergrad_cancel_notice_pdf",
        "教育部关于备案2023年度普通高等学校申请撤销的本科专业的通知",
        "https://jwc.hebeinu.edu.cn/upload/file/20250318/1742285397351870.pdf",
        "教育部文件（河北北方学院教务处镜像PDF）",
        "2024-02-04",
        "official_policy_pdf_mirror",
        "教育部教高函〔2024〕7号附件，列出2023年度普通高等学校撤销本科专业名单；PDF同时含教高函〔2024〕6号备案审批结果和2024年专业目录。",
    ),
    SourceSeed(
        "moe_2023_2022_undergrad_setup_results_pdf",
        "教育部关于公布2022年度普通高等学校本科专业备案和审批结果的通知",
        "http://www.moe.gov.cn/srcsite/A08/moe_1034/s4930/202304/W020230419336779647503.pdf",
        "中华人民共和国教育部",
        "2023-04-04",
        "official_policy_pdf",
        "教育部教高函〔2023〕3号附件1，内含2022年度普通高等学校撤销专业名单。",
    ),
    SourceSeed(
        "moe_2022_2021_undergrad_setup_results_pdf",
        "教育部关于公布2021年度普通高等学校本科专业备案和审批结果的通知",
        "http://www.moe.gov.cn/srcsite/A08/moe_1034/s4930/202202/W020220228600206028276.pdf",
        "中华人民共和国教育部",
        "2021-12-10",
        "official_policy_pdf",
        "教育部教高函〔2021〕14号附件1，内含2021年度普通高等学校撤销专业名单。",
    ),
    SourceSeed(
        "moe_2021_2020_undergrad_setup_results_xls",
        "教育部关于公布2020年度普通高等学校本科专业备案和审批结果的通知",
        "http://www.moe.gov.cn/srcsite/A08/moe_1034/s4930/202103/W020210302545152199812.xls",
        "中华人民共和国教育部",
        "2021-03-01",
        "official_policy_xls",
        "教育部2020年度普通高等学校本科专业备案和审批结果附件1，内含撤销本科专业名单。",
    ),
    SourceSeed(
        "moe_2020_2019_undergrad_setup_results_pdf",
        "教育部关于公布2019年度普通高等学校本科专业备案和审批结果的通知",
        "http://www.moe.gov.cn/srcsite/A08/moe_1034/s4930/202003/W020200303365402032446.pdf",
        "中华人民共和国教育部",
        "2020-02-21",
        "official_policy_pdf",
        "教育部教高函〔2020〕2号附件1，扫描版PDF；作为2019年度撤销本科专业名单的官方原始依据，结构化数据使用另一个机器可读XLS镜像交叉补足。",
    ),
    SourceSeed(
        "moe_zwfw_2020_2019_undergrad_setup_notice_json",
        "教育部关于公布2019年度普通高等学校本科专业备案和审批结果的通知（政务服务平台JSON）",
        "https://zwfw.moe.gov.cn/portal/dynamic/get/08a45c1fd1994ce1a939f071fda0da16",
        "教育部政务服务平台",
        "2020-03-03",
        "official_policy_json",
        "教育部政务服务平台动态详情接口返回教高函〔2020〕2号正文，明确2019年度撤销专业367个，并链接官方扫描PDF附件。",
    ),
    SourceSeed(
        "beijing_2019_municipal_undergrad_setup_doc",
        "2019年度市属普通高等学校本科专业备案和审批结果",
        "https://jw.beijing.gov.cn/gjc/tzgg_15688/202003/P020200313667515717463.doc",
        "北京市教育委员会",
        "2020-03-13",
        "official_policy_doc",
        "北京市教委2019年度市属普通高等学校本科专业备案和审批结果Word附件，含新增备案、新增审批和2条撤销本科专业名单。",
    ),
    SourceSeed(
        "hunan_2019_undergrad_catalog_notice",
        "关于印发《湖南省省属普通本科高等学校2019年招生专业目录》的通知",
        "http://jyt.hunan.gov.cn/sjyt/xxgk/tzgg/201904/t20190423_5320082.html",
        "湖南省教育厅",
        "2019-04-10",
        "official_policy_html",
        "湖南省教育厅湘教发[2019]8号通知正文说明，招生专业目录备注栏标注“拟撤销”的专业从2019年起停止招生，相关学校须在2019年7月正式提出撤销申请。",
    ),
    SourceSeed(
        "hunan_2019_undergrad_catalog_xls",
        "湖南省省属普通本科高等学校2019年招生专业目录",
        "http://govnew.hnedu.cn:8090/zcms/contentcore/resource/download?ID=58942",
        "湖南省教育厅",
        "2019-04-10",
        "official_policy_xls",
        "湖南省省属普通本科高等学校2019年招生专业目录附件，备注列标注14个“拟撤销”本科专业。",
    ),
    SourceSeed(
        "scnu_2019_teaching_quality_report_pdf",
        "华南师范大学2018-2019学年本科教学质量报告",
        "https://statics.scnu.edu.cn/pics/xxgk/2019/1225/1577289427663507.pdf",
        "华南师范大学信息公开网",
        "2019-12-25",
        "official_university_pdf",
        "华南师范大学2018-2019学年本科教学质量报告在专业建设部分列出学校申请撤销的7个本科专业。",
    ),
    SourceSeed(
        "ynnu_2019_teaching_quality_report_pdf",
        "云南师范大学2018-2019学年本科教学质量报告",
        "https://xxgk.ynnu.edu.cn/__local/1/53/18/6350568E90EF467BC89A08A7285_98718786_12B322.pdf",
        "云南师范大学信息公开网络平台",
        "2019-12-23",
        "official_university_pdf",
        "云南师范大学2018-2019学年本科教学质量报告表3列出2019年申请撤销的7个本科专业，含专业代码和所属学院。",
    ),
    SourceSeed(
        "dzu_2019_teaching_quality_report_pdf",
        "德州学院2018-2019学年本科教学质量报告",
        "https://zpc.dzu.edu.cn/__local/9/2C/CE/D6A4442F1793CE3478291065190_5D6FECEF_2932E3.pdf",
        "德州学院教学质量督导与评估中心",
        "2022-05-19",
        "official_university_pdf",
        "德州学院2018-2019学年本科教学质量报告说明2018-2019学年停招产品设计等5个本科专业。",
    ),
    SourceSeed(
        "glut_2019_teaching_quality_report_pdf",
        "桂林理工大学2018-2019学年本科教学质量报告",
        "https://jwc.glut.edu.cn/__local/1/36/41/4FA6AD19A0F3121A40D9354FDD2_B89D001F_188E34.pdf?e=.pdf",
        "桂林理工大学",
        "2019-12",
        "official_university_pdf",
        "桂林理工大学2018-2019学年本科教学质量报告在专业建设部分列出当年停招的8个校内专业。",
    ),
    SourceSeed(
        "hrbu_2020_teaching_quality_report_pdf",
        "哈尔滨学院2019-2020学年本科教学质量报告",
        "https://www.hrbu.edu.cn/system/_content/download.jsp?urltype=news.DownloadAttachUrl&owner=1285693926&wbfileid=C6F2FFE661AEFD917CA21C6A45B7FBD5",
        "哈尔滨学院教学质量评估中心",
        "2020-12-29",
        "official_university_pdf",
        "哈尔滨学院2019-2020学年本科教学质量报告正文列出撤销6个本科专业，附表4列出当年停招的7个本科专业。",
    ),
    SourceSeed(
        "cust_2020_applied_statistics_application_pdf",
        "长春理工大学应用统计学专业备案申请表",
        "https://www.cust.edu.cn/docs/2020-07/20200722103725618520.pdf",
        "长春理工大学",
        "2020-07",
        "official_university_pdf",
        "长春理工大学应用统计学专业备案申请表在学校近五年专业增设、停招、撤并情况中列出停招3个专业和2019年撤销的5个本科专业。",
    ),
    SourceSeed(
        "suse_2019_teaching_quality_report_pdf",
        "四川轻化工大学2018-2019学年本科教学质量报告",
        "https://webadm.suse.edu.cn/_upload/article/files/4d/13/a538b58c4dce9410ef510b08465f/56a9ea3d-a2d4-43c7-b091-1ece5822f3aa.pdf",
        "四川轻化工大学",
        "2019",
        "official_university_pdf",
        "四川轻化工大学2018-2019学年本科教学质量报告在专业建设部分列出学校申请撤销的8个本科专业。",
    ),
    SourceSeed(
        "ourjiangsu_2020_2019_jiangsu_cancel_summary",
        "教育部最新发布！江苏高校新增96个本科专业 人工智能受青睐",
        "https://www.ourjiangsu.com/a/20200303/1583232487126.shtml",
        "我苏网",
        "2020-03-03",
        "secondary_news_html",
        "二级转载来源，按校列出江苏省2019年度普通高校撤销专业13个；用于补充2019年江苏学校撤销专业的可抓取文本证据，置信度保持medium。",
    ),
    SourceSeed(
        "cyol_2018_hdu_major_stop_adjustment",
        "高校专业并非上了保险 杭电“砍掉”11个专业",
        "https://zqb.cyol.com/html/2018-04/16/nw.D110000zgqnb_20180416_3-04.htm",
        "中国青年报",
        "2018-04-16",
        "secondary_news_html",
        "二级媒体来源，报道杭州电子科技大学未来两年内停招11个本科专业，并列出过去5年已撤销或停招的6个专业；按medium置信度结构化。",
    ),
    SourceSeed(
        "xpu_2019_major_dynamic_adjustment",
        "关于2019年申报新专业及现有专业动态调整公示",
        "https://jw.xpu.edu.cn/info/1097/1805.htm",
        "西安工程大学教务处",
        "2018-07-08",
        "official_university_html",
        "西安工程大学教务处公示列出2019年拟撤销5个专业、预警8个专业，并列出2020年拟停招3个专业。",
    ),
    SourceSeed(
        "zjnu_2020_2019_undergrad_setup_news",
        "省内领先！浙师大智能制造工程专业获教育部批准设置",
        "https://gxy.zjnu.edu.cn/_t952/p18657c3798/list.psp",
        "浙江师范大学工学院",
        "2020",
        "official_university_html",
        "浙江师范大学工学院页面引用教高函〔2020〕2号，说明学校撤销教育学、艺术教育、环境科学3个本科专业。",
    ),
    SourceSeed(
        "jxau_smart_agriculture_application_pdf",
        "江西农业大学智慧农业专业备案申请表",
        "https://www.jxau.edu.cn/__local/5/01/F6/AFB432C4524707BDC5E7CDD9B86_7FA60C84_F9465.pdf?e=.pdf",
        "江西农业大学",
        "2021",
        "official_university_pdf",
        "江西农业大学智慧农业专业备案申请表列出近五年专业增设、调整、停招、撤并情况，包括2018、2019、2020年撤销专业和2020、2021年停招专业。",
    ),
    SourceSeed(
        "hbu_jijianjiancha_application_pdf",
        "河北大学纪检监察专业备案申请表",
        "https://gz.moe.gov.cn/api/gjs/file/download/0ec3cb48ec484564a9b1bbb6fc390ff9",
        "教育部高校本科专业设置申请材料公示平台 / 河北大学",
        "2022-07-27",
        "official_university_pdf",
        "河北大学纪检监察专业备案申请表列出学校近五年专业增设、停招、撤并情况；用于补充可与教育部2019年度结果精确对应的撤销专业证据。",
    ),
    SourceSeed(
        "kmust_2025_stomatology_application_pdf",
        "昆明理工大学口腔医学专业申请表",
        "https://gz.moe.gov.cn/api/gjs/file/download/e945f4f10a9b45089ae9b9e3ef216e9d",
        "教育部高校本科专业设置申请材料公示平台 / 昆明理工大学",
        "2025-08-14",
        "official_university_pdf",
        "昆明理工大学口腔医学专业申请表列出学校近五年专业增设、停招、撤并情况；用于补充可与教育部2019年度结果精确对应的撤销专业证据。",
    ),
    SourceSeed(
        "nhky_2020_cybersecurity_application_pdf",
        "南昌航空大学科技学院网络空间安全专业申请表",
        "https://gz.moe.gov.cn/api/gjs/file/download/c93e8deb50c14b5a8d87d4b0c9fbb4f5",
        "教育部高校本科专业设置申请材料公示平台 / 南昌航空大学科技学院",
        "2020",
        "official_university_pdf",
        "南昌航空大学科技学院网络空间安全专业申请表列出2015-2020学年停招专业和2020年现报撤销待批专业。",
    ),
    SourceSeed(
        "qjnu_2025_area_studies_application_pdf",
        "曲靖师范学院区域国别学专业申请表",
        "https://www.qjnu.edu.cn/upload/files/2025/7/1b3bbc87d3eda4a5.pdf",
        "曲靖师范学院",
        "2025-07-23",
        "official_university_pdf",
        "曲靖师范学院区域国别学专业申请表列出近五年停招14个本科专业和撤销3个本科专业。",
    ),
    SourceSeed(
        "aqnu_2025_sports_training_application_pdf",
        "安庆师范大学运动训练专业申请表",
        "https://gz.moe.gov.cn/api/gjs/file/download/439da5f43d6d424d85c0c8e394b19077",
        "教育部高校本科专业设置申请材料公示平台 / 安庆师范大学",
        "2025-08-14",
        "official_university_pdf",
        "安庆师范大学运动训练专业申请表列出近五年停招4个本科专业，并逐年列出2022-2024年撤销9个本科专业。",
    ),
    SourceSeed(
        "yxu_2025_smart_energy_application_pdf",
        "玉溪师范学院智慧能源工程专业申请表",
        "https://gz.moe.gov.cn/api/gjs/file/download/e0767f9c05374f479975d3afc0f642ef",
        "教育部高校本科专业设置申请材料公示平台 / 玉溪师范学院",
        "2025-08-14",
        "official_university_pdf",
        "玉溪师范学院智慧能源工程专业申请表列出近五年撤销13个本科专业、停招7个本科专业。",
    ),
    SourceSeed(
        "yxu_2025_biological_breeding_application_pdf",
        "玉溪师范学院生物育种科学专业申请表",
        "https://gz.moe.gov.cn/api/gjs/file/download/06b6f05e8155465087b5a2cf6387386a",
        "教育部高校本科专业设置申请材料公示平台 / 玉溪师范学院",
        "2025-08-14",
        "official_university_pdf",
        "玉溪师范学院生物育种科学专业申请表列出近五年撤销13个本科专业、停招7个本科专业。",
    ),
    SourceSeed(
        "zcmu_2025_nursing_application_pdf",
        "浙江中医药大学护理学专业申请表",
        "https://gz.moe.gov.cn/api/gjs/file/download/054125142c9a4107b0f8abca54f01dc2",
        "教育部高校本科专业设置申请材料公示平台 / 浙江中医药大学",
        "2025-08-14",
        "official_university_pdf",
        "浙江中医药大学护理学专业申请表列出2023、2025年停招专业和2025年撤销药物制剂专业。",
    ),
    SourceSeed(
        "cxtc_2025_uav_smart_network_application_pdf",
        "楚雄师范学院无人机与智能网络专业申请表",
        "https://gz.moe.gov.cn/api/gjs/file/download/4392c44d2f2449aaaf89ed1ebdfff5ed",
        "教育部高校本科专业设置申请材料公示平台 / 楚雄师范学院",
        "2025-08-14",
        "official_university_pdf",
        "楚雄师范学院无人机与智能网络专业申请表列出近五年撤销应用物理学、运动康复等7个本科专业。",
    ),
    SourceSeed(
        "wust_2025_fintech_application_pdf",
        "武汉科技大学金融科技专业备案申请表",
        "https://jwc.wust.edu.cn/__local/6/BC/EB/E9C800060CF1C10BB542538B5B4_10D28370_7CDD3B.pdf?e=.pdf",
        "武汉科技大学教务处",
        "2025-07-18",
        "official_university_pdf",
        "武汉科技大学金融科技专业备案申请表列出近五年停招人文地理与城乡规划、汽车服务工程、电子商务、交通运输、马克思主义理论、光电信息科学与工程等专业，并列出近五年撤销风景园林。",
    ),
    SourceSeed(
        "bigc_2025_digital_economy_application_pdf",
        "北京印刷学院数字经济专业备案申请表",
        "https://www.bigc.edu.cn/docs//2025-07/ea95aa5ae5ff488fb68f98f56b9f4528.pdf",
        "北京印刷学院",
        "2025-07-24",
        "official_university_pdf",
        "北京印刷学院数字经济专业备案申请表列出近五年停招摄影、电子信息工程、物流管理、物流工程、工业设计、市场营销等本科专业。",
    ),
    SourceSeed(
        "hbkjxy_2025_cybersecurity_application_pdf",
        "河北科技学院网络空间安全专业备案申请表",
        "https://www.hbkjxy.edu.cn/UploadFiles/2025/102/F134000713362457.pdf",
        "河北科技学院",
        "2025-07-24",
        "official_university_pdf",
        "河北科技学院网络空间安全专业备案申请表逐年列出2020-2023年停招专业，并列出2021、2023、2024年撤销专业。",
    ),
    SourceSeed(
        "whwl_2025_ai_education_application_pdf",
        "武汉文理学院人工智能教育专业申请表",
        "https://gz.moe.gov.cn/api/gjs/file/download/fe8b1bb22ba34c31b447049f30f1b66a",
        "教育部高校本科专业设置申请材料公示平台 / 武汉文理学院",
        "2025-07-23",
        "official_university_pdf",
        "武汉文理学院人工智能教育专业申请表列出近五年停招旅游管理、互联网金融等6个专业，并列出近五年撤销环境工程、园艺等9个专业。",
    ),
    SourceSeed(
        "whwl_2025_medical_device_application_pdf",
        "武汉文理学院医疗器械与装备工程专业申请表",
        "https://gz.moe.gov.cn/api/gjs/file/download/5de28769d5094432acddd78880209b9a",
        "教育部高校本科专业设置申请材料公示平台 / 武汉文理学院",
        "2025-08-14",
        "official_university_pdf",
        "武汉文理学院医疗器械与装备工程专业申请表列出与人工智能教育申请表一致的近五年停招6个专业和撤销9个专业。",
    ),
    SourceSeed(
        "muc_2025_sports_training_application_pdf",
        "中央民族大学运动训练专业申请表",
        "https://gz.moe.gov.cn/api/gjs/file/download/0b8fca8fe8ca4fe3a0254bd8121ed21c",
        "教育部高校本科专业设置申请材料公示平台 / 中央民族大学",
        "2025-08-14",
        "official_university_pdf",
        "中央民族大学运动训练专业申请表说明近五年停招朝鲜语等10个专业；本数据集结构化其中直接点名的朝鲜语。",
    ),
    SourceSeed(
        "sus_2025_football_application_pdf",
        "上海体育大学足球运动专业申请表",
        "https://gz.moe.gov.cn/api/gjs/file/download/3318da651b8043c3a4c2b730983e013b",
        "教育部高校本科专业设置申请材料公示平台 / 上海体育大学",
        "2025-08-14",
        "official_university_pdf",
        "上海体育大学足球运动专业申请表列出特殊教育未招生，并说明艺术教育、公共事业管理、信息管理与信息系统停招并撤销。",
    ),
    SourceSeed(
        "qtnu_2025_sports_training_application_pdf",
        "琼台师范学院运动训练专业申请表",
        "https://gz.moe.gov.cn/api/gjs/file/download/999cd0122b0f443c8cf1b48cc698d1a0",
        "教育部高校本科专业设置申请材料公示平台 / 琼台师范学院",
        "2025-08-14",
        "official_university_pdf",
        "琼台师范学院运动训练专业申请表说明2023、2024、2025年连续三年停招艺术设计学专业。",
    ),
    SourceSeed(
        "gmc_2025_geriatric_medicine_application_pdf",
        "贵州医科大学老年医学与健康专业申请表",
        "https://gz.moe.gov.cn/api/gjs/file/download/6ff8097692584306a2e4bea38b22eb4d",
        "教育部高校本科专业设置申请材料公示平台 / 贵州医科大学",
        "2025-08-14",
        "official_university_pdf",
        "贵州医科大学老年医学与健康专业申请表说明2023年暂停翻译，2024年暂停英语和翻译招生。",
    ),
    SourceSeed(
        "gnnu_2025_ai_education_application_pdf",
        "赣南师范大学人工智能教育专业申请表",
        "https://gz.moe.gov.cn/api/gjs/file/download/77e046ab15714a52be86513ec95a0577",
        "教育部高校本科专业设置申请材料公示平台 / 赣南师范大学",
        "2025-08-14",
        "official_university_pdf",
        "赣南师范大学人工智能教育专业申请表逐年说明2022、2023、2025年停招和撤销专业，结构化原文直接点名项。",
    ),
    SourceSeed(
        "hbesxy_2025_stomatology_application_pdf",
        "湖北恩施学院口腔医学专业申请表",
        "https://gz.moe.gov.cn/api/gjs/file/download/bd92001eed1c41f296697dd4bb9865a5",
        "教育部高校本科专业设置申请材料公示平台 / 湖北恩施学院",
        "2025-08-14",
        "official_university_pdf",
        "湖北恩施学院口腔医学专业申请表说明2018-2022年停招专业数量，结构化其中直接点名的通信工程、社会学。",
    ),
    SourceSeed(
        "nwupl_2025_forensic_science_application_pdf",
        "西北政法大学司法鉴定学专业申请表",
        "https://gz.moe.gov.cn/api/gjs/file/download/199289b972fe46a4888014ec621e33c4",
        "教育部高校本科专业设置申请材料公示平台 / 西北政法大学",
        "2025-08-14",
        "official_university_pdf",
        "西北政法大学司法鉴定学专业申请表逐年列出2022-2025年未招生和撤销专业。",
    ),
    SourceSeed(
        "sicnu_2025_sports_tourism_application_pdf",
        "四川师范大学体育旅游专业申请表",
        "https://gz.moe.gov.cn/api/gjs/file/download/0c037719638f4795b5dd1eadd256b7e1",
        "教育部高校本科专业设置申请材料公示平台 / 四川师范大学",
        "2025-08-14",
        "official_university_pdf",
        "四川师范大学体育旅游专业申请表说明2020年停招酒店管理、服装设计与工程，2021年停招国际经济与贸易、工业工程、工业设计。",
    ),
    SourceSeed(
        "tjufezj_2025_fiscal_application_pdf",
        "天津财经大学珠江学院财政学专业申请表",
        "https://gz.moe.gov.cn/api/gjs/file/download/5823b6776c07436caeb311f03d2851db",
        "教育部高校本科专业设置申请材料公示平台 / 天津财经大学珠江学院",
        "2025-08-14",
        "official_university_pdf",
        "天津财经大学珠江学院财政学专业申请表说明2020年停招产品设计、服装与服饰设计，2024年撤销信息管理与信息系统、网络工程、俄语。",
    ),
    SourceSeed(
        "whtyxykj_2025_smart_sports_engineering_application_pdf",
        "武汉体育学院体育科技学院智能体育工程专业申请表",
        "https://gz.moe.gov.cn/api/gjs/file/download/267ad4cf9b5b45d8988ce2c532af6d0c",
        "教育部高校本科专业设置申请材料公示平台 / 武汉体育学院体育科技学院",
        "2025-08-14",
        "official_university_pdf",
        "武汉体育学院体育科技学院智能体育工程专业申请表说明2019年停招文化产业管理、2024年撤销表演。",
    ),
    SourceSeed(
        "ntu_2025_smart_energy_application_pdf",
        "南通大学智慧能源工程专业申请表",
        "https://gz.moe.gov.cn/api/gjs/file/download/558e75203edd4a179fe8ea43f60fdee5",
        "教育部高校本科专业设置申请材料公示平台 / 南通大学",
        "2025-08-14",
        "official_university_pdf",
        "南通大学智慧能源工程专业申请表列出近五年停招20个本科专业和撤销7个本科专业；本数据集先结构化停招项。",
    ),
    SourceSeed(
        "tjcu_2025_low_altitude_application_pdf",
        "天津商业大学低空技术与工程专业申请表",
        "https://gz.moe.gov.cn/api/gjs/file/download/6863270fb75e40079cb0c712db9186cd",
        "教育部高校本科专业设置申请材料公示平台 / 天津商业大学",
        "2025-08-14",
        "official_university_pdf",
        "天津商业大学低空技术与工程专业申请表列出近五年停招5个本科专业和撤销4个本科专业；本数据集先结构化停招项。",
    ),
    SourceSeed(
        "csuft_2025_low_altitude_application_pdf",
        "中南林业科技大学低空技术与工程专业申请表",
        "https://gz.moe.gov.cn/api/gjs/file/download/02a1295afb4e4d32bad97012ca5982d3",
        "教育部高校本科专业设置申请材料公示平台 / 中南林业科技大学",
        "2025-08-14",
        "official_university_pdf",
        "中南林业科技大学低空技术与工程专业申请表逐年列出2021、2023、2024、2025年停招19个本科专业。",
    ),
    SourceSeed(
        "guizhou_2019_undergrad_setup_results_pdf_mirror",
        "2019年度贵州省普通高等学校本科专业备案和审批结果",
        "https://www.gyiist.edu.cn/ewebeditor/uploadfile/2020/04/20200408160012335.pdf",
        "贵州省教育厅附件镜像（贵阳信息科技学院）",
        "2020-04-08",
        "secondary_pdf_mirror",
        "贵州省2019年度普通高等学校本科专业备案和审批结果PDF镜像，含11条撤销本科专业名单及专业代码、学位门类、修业年限。",
    ),
    SourceSeed(
        "yingkou_2019_undergrad_setup_results_xls_mirror",
        "2019年度普通高等学校本科专业备案和审批结果（机器可读XLS镜像）",
        "https://jyj.yingkou.gov.cn/jyj/0a6b5da3-3287-4e4a-8f4f-0600af8bc92e/4bf00567-0c06-494c-b5c4-6dc415d2b045/%E9%99%84%E4%BB%B63.%E8%BE%BD%E5%AE%81%E7%9C%81%E8%90%A5%E5%8F%A3%E5%B8%822025%E5%B9%B4%E6%A0%A1%E5%9B%AD%E6%8B%9B%E8%81%98%E6%95%99%E5%B8%88%E8%AF%95%E8%AE%B2%E6%95%99%E6%9D%90%E4%B8%80%E8%A7%88%E8%A1%A8.xls",
        "营口市教育局附件镜像",
        "2025",
        "secondary_xls_mirror",
        "文件名为营口市2025年招聘附件，但工作表Sheet1实际包含2019年度普通高等学校本科专业备案和审批结果及367条撤销本科专业名单；用于补足教育部2019扫描PDF的机器可读结构化数据。",
    ),
    SourceSeed(
        "eol_2020_2019_moe_direct_cancel_table",
        "教育部2019年度撤销本科专业",
        "https://gaokao.eol.cn/news/202007/t20200701_1736172.shtml",
        "中国教育在线",
        "2020-07-01",
        "secondary_news_html",
        "中国教育在线转列表格，覆盖2019年度教育部直属高校撤销本科专业11条，提供专业代码、学位授予门类和修业年限；按medium置信度补充字段。",
    ),
    SourceSeed(
        "moe_2019_2018_undergrad_setup_results_docx",
        "教育部关于公布2018年度普通高等学校本科专业备案和审批结果的通知",
        "http://www.moe.gov.cn/srcsite/A08/moe_1034/s4930/201903/W020190329560198491315.docx",
        "中华人民共和国教育部",
        "2019-03-21",
        "official_policy_docx",
        "教育部教高函〔2019〕7号附件，内含2018年度撤销本科专业名单。",
    ),
    SourceSeed(
        "moe_2018_2017_undergrad_setup_results_docx",
        "教育部关于公布2017年度普通高等学校本科专业备案和审批结果的通知",
        "http://www.moe.gov.cn/srcsite/A08/moe_1034/s4930/201803/W020180322531880180253.docx",
        "中华人民共和国教育部",
        "2018-03-15",
        "official_policy_docx",
        "教育部教高函〔2018〕4号附件，内含2017年度撤销本科专业名单。",
    ),
    SourceSeed(
        "moe_2017_2016_undergrad_setup_results_docx",
        "教育部关于公布2016年度普通高等学校本科专业备案和审批结果的通知",
        "http://www.moe.gov.cn/srcsite/A08/moe_1034/s4930/201703/W020170322299503870738.docx",
        "中华人民共和国教育部",
        "2017-03-13",
        "official_policy_docx",
        "教育部教高函〔2017〕2号附件，内含2016年度撤销本科专业名单。",
    ),
    SourceSeed(
        "moe_2016_2015_undergrad_setup_results_docx",
        "教育部关于公布2015年度普通高等学校本科专业备案和审批结果的通知",
        "http://www.moe.gov.cn/srcsite/A08/moe_1034/s4930/201603/W020190627383907890233.docx",
        "中华人民共和国教育部",
        "2016-03-03",
        "official_policy_docx",
        "教育部教高函〔2016〕2号附件，内含2015年度撤销本科专业名单。",
    ),
    SourceSeed(
        "moe_2015_2014_undergrad_setup_results_docx_mirror",
        "教育部：2014年度普通高等学校本科专业备案或审批结果",
        "https://www.htu.edu.cn/_upload/article/files/23/9b/269ec5b74ca5bdda94d26c548303/fd9a70c3-33be-43f9-88a4-69c5ac8500c1.docx",
        "河南师范大学发展规划处转载教育部附件",
        "2015-03-27",
        "official_policy_docx_mirror",
        "教育部2014年度普通高等学校本科专业备案或审批结果附件镜像，内含撤销本科专业名单。",
    ),
    SourceSeed(
        "sdipct_2013_undergrad_setup_results_doc_mirror",
        "2013年度普通高等学校本科专业备案或审批结果（教育部）",
        "https://jwc.sdipct.edu.cn/__local/B/74/5F/F5A97D8634E688F38DD34D55EE9_206F7A96_2A3200.doc?e=.doc",
        "山东石油化工学院教务处镜像",
        "2014-03",
        "official_policy_doc_mirror",
        "教育部2013年度普通高等学校本科专业备案或审批结果附件镜像，内含撤销专业名单。",
    ),
    SourceSeed(
        "haedu_2011_undergrad_cancel_doc_mirror",
        "教育部公布2011年度高等学校本科专业设置备案或审批结果：同意撤销的高等学校本科专业名单",
        "https://www.gaokao.haedu.cn/UserFiles/File/1330939153276.doc",
        "河南省阳光高考信息平台转载教育部附件",
        "2012-03-05",
        "official_policy_doc_mirror",
        "教育部2011年度高等学校本科专业设置备案或审批结果附件4，列出32条同意撤销的高等学校本科专业名单。",
    ),
    SourceSeed(
        "sina_2022_jiangxi_undergrad_structure_policy",
        "一图读懂：江西省高校本科专业结构优化调整指导办法",
        "https://edu.sina.cn/2022-06-02/detail-imizmscu4725808.d.html?vt=4",
        "新浪教育转载江西省教育厅微信公众号",
        "2022-06-01",
        "news_html",
        "转载江西省教育厅通知全文，明确低于50%黄牌、连续两年低于50%红牌并停止招生。",
    ),
    SourceSeed(
        "sichuan_2023_employment_enrollment_policy",
        "四川深化高等学校“就业—招生—培养”联动机制改革",
        "https://cbgc.scol.com.cn/news/4275479",
        "四川在线/川观新闻",
        "2023-07-03",
        "news_html",
        "四川就业—招生—培养联动机制，明确连续两年低于50%黄牌、连续三年低于50%红牌。",
    ),
    SourceSeed(
        "scol_2026_sichuan_undergrad_adjustment_stats",
        "四川高校新增本科专业120个，将列入相关高校2026年高考招生计划",
        "https://sichuan.scol.com.cn/ggxw/202604/83247506.html",
        "四川在线",
        "2026-04-30",
        "news_html",
        "四川在线报道称记者从省教育厅了解到，四川全省2025年共调整优化专业296个，其中新设120个、调整修业年限及学位授予门类专业点11个、停招132个、撤销33个。",
    ),
    SourceSeed(
        "fujian_2025_full_employment_policy",
        "关于全方位促进高质量充分就业的实施意见",
        "https://rst.fujian.gov.cn/zw/ztzl/zxzt/cjgzlcfjy/jyzcsd/zcwj_68203/202509/t20250910_7003576.htm",
        "福建省人力资源和社会保障厅",
        "2025-09-10",
        "official_policy_html",
        "福建就业红黄牌提示制度，对就业质量不高专业调减或停止招生。",
    ),
    SourceSeed(
        "e23_2024_shandong_teacher_warning_policy",
        "山东推进师范教育高质量发展：建立师范类专业红黄牌制度",
        "https://news.e23.cn/shandong/2024-09-07/2024090700025.html",
        "舜网/济南时报",
        "2024-09-07",
        "news_html",
        "报道山东省教育厅等5部门措施，明确师范类专业红黄牌制度和退出机制。",
    ),
    SourceSeed(
        "anhui_2022_structure_policy",
        "安徽：连续3年就业去向落实率低于60%的专业暂停招生",
        "https://news.eol.cn/yaowen/202208/t20220807_2240891.shtml",
        "中国教育在线",
        "2022-08-07",
        "news_html",
        "安徽深化高校学科专业结构改革方案公开报道，明确连续3年就业去向落实率低于60%的专业暂停招生。",
    ),
    SourceSeed(
        "acabridge_2026_wuhan_business_warning",
        "2所高校公示：拟撤销、停招专业名单",
        "https://www.acabridge.cn/zxhz/202605/t20260519_2735785.shtml",
        "学术桥转载西安文理学院、武汉商学院网站",
        "2026-05-20",
        "news_html",
        "武汉商学院2026年度拟预警和停招专业公示，列出3个拟预警专业和1个拟停招专业。",
    ),
    SourceSeed(
        "xpu_2019_dynamic_adjustment",
        "关于2019年拟增设新专业和专业动态调整的公示",
        "https://jw.xpu.edu.cn/info/1097/1805.htm",
        "西安工程大学教务处",
        "2018-07-08",
        "official_university_html",
        "西安工程大学2019年度拟撤销、拟停招、拟预警专业公示，列出专业名称、代码和学位门类。",
    ),
    SourceSeed(
        "wbu_2025_warning_stop",
        "武汉商学院2025年度拟申报本科专业、预备案专业、预警和停招专业公示",
        "https://www.wbu.edu.cn/2025/0718/c2191a82226/page.htm",
        "武汉商学院",
        "2025-07-18",
        "official_university_html",
        "武汉商学院2025年度拟预警专业4个、拟停招专业1个。",
    ),
    SourceSeed(
        "syuct_2025_major_cancellation",
        "沈阳化工大学2025年拟撤销专业公示",
        "https://jiaowu.syuct.edu.cn/info/1087/3030.htm",
        "沈阳化工大学教务处",
        "2025-07-22",
        "official_university_html",
        "沈阳化工大学2025年拟撤销停招满五年的5个本科专业。",
    ),
    SourceSeed(
        "ncu_2025_major_cancellation",
        "关于南昌大学2025年度本科专业调整方案的公示",
        "https://jwc.ncu.edu.cn/content.jsp?urltype=news.NewsContentUrl&wbnewsid=42011&wbtreeid=1541",
        "南昌大学教务处",
        "2025-07-23",
        "official_university_html",
        "南昌大学2025年度拟撤销戏剧影视文学等8个本科专业。",
    ),
    SourceSeed(
        "gznu_2025_major_cancellation",
        "关于2025年度本科专业设置情况的公示",
        "https://jwc.gznu.edu.cn/info/2002/89425.htm",
        "贵州师范大学教务处",
        "2025-07-25",
        "official_university_html",
        "贵州师范大学2025年度拟撤销冶金工程、工程管理、电子信息科学与技术、数字媒体艺术4个本科专业，并列出专业代码。",
    ),
    SourceSeed(
        "synu_2025_major_cancellation_docx",
        "关于我校2025年度拟增设本科专业、预申报本科专业、拟撤销本科专业的公示",
        "https://www.synu.edu.cn/_upload/article/files/05/d1/a1a5610e44c9babc3b1366afa57a/e9296253-a3cc-4056-b75b-1972ceb63e15.docx",
        "沈阳师范大学教务处",
        "2025-07-22",
        "official_university_docx",
        "沈阳师范大学2025年度拟撤销食品质量与安全、人力资源管理、会展经济与管理、动画4个本科专业，并列出专业代码。",
    ),
    SourceSeed(
        "ncpu_2025_major_cancellation_notice",
        "关于对2025年度拟新增设、撤销本科专业的公示",
        "https://jwc.ncpu.edu.cn/jiaoxuegaige/rencaipeiyang/html.php?c-4138.html",
        "南昌工学院教务处",
        "2025-07-14",
        "official_university_html",
        "南昌工学院2025年度拟新增设、撤销本科专业公示；正文以图片形式展示拟撤销专业表。",
    ),
    SourceSeed(
        "ncpu_2025_major_cancellation_table_image",
        "南昌工学院2025年度拟撤销本科专业表图片",
        "https://jwc.ncpu.edu.cn/uploadfile/2025/07/15/1752542800855385.png",
        "南昌工学院教务处",
        "2025-07-14",
        "official_university_image",
        "南昌工学院公示页内嵌拟撤销专业表，列出采购管理、交通运输、工业工程、测绘工程及专业代码、停招年份。",
    ),
    SourceSeed(
        "whxy_2025_major_cancellation",
        "武汉学院2025年拟撤销本科专业公示",
        "https://www.whxy.edu.cn/info/1012/17922.htm",
        "武汉学院教务处",
        "2025-07-17",
        "official_university_html",
        "武汉学院2025年度拟撤销投资学1个本科专业，并列出专业代码、学科门类和所属学院。",
    ),
    SourceSeed(
        "gufe_2025_major_cancellation_pdf",
        "关于贵州财经大学2025年度拟撤销专业名单的公示",
        "https://jwc.gufe.edu.cn/__local/2/7F/FA/8E9E8A643FCB035ADBB59EC3579_33678027_11793.pdf",
        "贵州财经大学教务处",
        "2025-07-25",
        "official_university_pdf",
        "贵州财经大学2025年度拟撤销投资学、公共事业管理、应用心理学、工程造价、房地产开发与管理、教育技术学、日语7个本科专业，并列出专业代码、学院和授予学位门类。",
    ),
    SourceSeed(
        "hcnu_2025_major_cancellation",
        "河池学院2025年拟撤销本科专业公示",
        "https://news.hcnu.edu.cn/info/1153/38819.htm",
        "河池学院教务处",
        "2025-07-25",
        "official_university_html",
        "河池学院2025年度拟撤销微电子科学与工程、网络工程2个本科专业，并列出专业代码、修业年限和学位授予门类。",
    ),
    SourceSeed(
        "wjut_2025_cancel_stop",
        "关于学校2025年申请增设专业和撤销停招专业的公示",
        "https://www.wjut.edu.cn/jiao-wu-bu/jiao-wu-notice/pageinfo21401.html",
        "皖江工学院教务部",
        "2025-05-15",
        "official_university_html",
        "皖江工学院2025年度拟撤销公共艺术、应用化学、材料成型及控制工程3个本科专业，并自2025年起停招汽车服务工程；正文列出专业代码、授予学位和停招起始年份。",
    ),
    SourceSeed(
        "hqu_2025_major_cancellation",
        "关于华侨大学2025年本科专业设置调整的公示",
        "https://www.hqu.edu.cn/info/1069/710563.htm",
        "华侨大学",
        "2025-07-22",
        "official_university_html",
        "华侨大学2025年度拟撤销自动化、会计学2个本科专业。",
    ),
    SourceSeed(
        "xauat_2022_cancel_stop",
        "关于2022年度专业调整的公示",
        "https://jwc.xauat.edu.cn/xxwz_nry.jsp?urltype=news.NewsContentUrl&wbnewsid=20671&wbtreeid=1072",
        "西安建筑科技大学教务处",
        "2022-07-13",
        "official_university_html",
        "西安建筑科技大学2022年度拟撤销产品设计、摄影2个本科专业，并继续停招软件工程1个本科专业。",
    ),
    SourceSeed(
        "nyist_2025_major_cancellation",
        "关于2025年度本科专业设置与调整的公示",
        "https://jwc.nyist.edu.cn/info/1032/12286.htm",
        "南阳理工学院教务处",
        "2025-07-11",
        "official_university_html",
        "南阳理工学院2025年度拟申请撤销停招五年及以上且无在籍学生的5个本科专业。",
    ),
    SourceSeed(
        "jxnu_2025_major_cancellation",
        "关于我校2025年度本科专业撤销情况的公示",
        "https://jwc.jxnu.edu.cn/Portal/ArticlesView.aspx?id=13229",
        "江西师范大学教务处",
        "2025-07-23",
        "official_university_html",
        "江西师范大学2025年度拟对网络工程、物联网工程、经济统计学、信息与计算科学、汉语言、翻译、劳动与社会保障7个本科专业实施撤销处理。",
    ),
    SourceSeed(
        "jgsu_2025_major_cancellation",
        "关于2025年拟新增和撤销本科专业的公示",
        "https://jwc.jgsu.edu.cn/info/1460/20377.htm",
        "井冈山大学教务处",
        "2025-07-07",
        "official_university_html",
        "井冈山大学2025年拟撤销商务英语、应用物理学、动画3个本科专业。",
    ),
    SourceSeed(
        "nbu_2025_major_cancellation",
        "关于撤销与新增设若干专业的公示",
        "https://jwc.nbu.edu.cn/info/1125/13344.htm",
        "宁波大学教务处",
        "2025-07-04",
        "official_university_html",
        "宁波大学2025年度申请撤销财务会计教育、新闻学、广告学、旅游管理与服务教育、产品设计、物联网6个本科专业。",
    ),
    SourceSeed(
        "cqnu_2025_teaching_committee_major_adjustment",
        "学校召开第四届教学指导委员会第二次工作会",
        "https://www.cqnu.edu.cn/info/1202/23051.htm",
        "重庆师范大学",
        "2025-06-12",
        "official_university_html",
        "重庆师范大学官方新闻披露教学指导委员会审议2025年拟新设专业、第二学士学位专业、撤销专业，教务处汇报拟撤销信息管理与信息系统专业。",
    ),
    SourceSeed(
        "qzuie_2025_major_cancellation_notice",
        "关于公布2025年度撤销部分停招本科专业的通知",
        "http://jwc.qzuie.edu.cn/index.php?s=news&c=show&id=1215",
        "泉州信息工程学院教务处",
        "2025-07-15",
        "official_university_html",
        "泉州信息工程学院教务处官网发布2025年度撤销部分停招本科专业通知；当前公开页面保留标题、来源和时间，未暴露正文专业名单。",
    ),
    SourceSeed(
        "qdu_2025_major_cancellation",
        "2025年青岛大学本科专业设置调整情况公示",
        "https://jwc.qdu.edu.cn/info/1009/6006.htm",
        "青岛大学教务处",
        "2025-07-22",
        "official_university_html",
        "青岛大学2025年度拟撤销近五年停招的数字媒体技术本科专业。",
    ),
    SourceSeed(
        "jzmu_2025_major_cancellation_notice",
        "关于2025年撤销本科专业的公示",
        "https://jwc.jzmu.edu.cn/info/1207/2257.htm",
        "锦州医科大学教务处",
        "2025-07-24",
        "official_university_html",
        "锦州医科大学2025年撤销本科专业公示说明保险学专业停招已超过五年且目前已无在校生，决定予以撤销。",
    ),
    SourceSeed(
        "pxc_2025_smart_construction_application_pdf",
        "萍乡学院智能建造专业备案申请表",
        "https://www.pxc.jx.cn/zhinengjianzao.pdf",
        "萍乡学院",
        "2025-07-25",
        "official_university_pdf",
        "萍乡学院智能建造专业备案申请表列出近五年专业增设、停招、撤并情况，包括2021年思想政治教育停招、2022年数字媒体技术停招、2023年工程造价停招、2025年拟撤销设施农业科学与工程。",
    ),
    SourceSeed(
        "hubu_2025_major_cancellation",
        "关于2025年度本科专业设置的公示",
        "https://jwc.hubu.edu.cn/info/1061/9376.htm",
        "湖北大学本科生院",
        "2025-07-18",
        "official_university_html",
        "湖北大学2025年度拟申请撤销连续停招五年及以上的工程管理、公共事业管理2个本科专业。",
    ),
    SourceSeed(
        "wbu_2024_warning_stop",
        "武汉商学院2024年度拟申报本科专业、预备案专业、预警和停招专业公示",
        "https://www.wbu.edu.cn/2024/0729/c2191a75781/page.htm",
        "武汉商学院",
        "2024-07-29",
        "official_university_html",
        "武汉商学院2024年度拟预警专业1个、拟停招专业2个。",
    ),
    SourceSeed(
        "hxu_2025_major_cancellation",
        "关于2025年度拟申请增设本科专业、预备案本科专业和撤销本科专业的公示",
        "https://www10.hxu.edu.cn/info/1212/292802.htm",
        "河西学院",
        "2025-07-21",
        "official_university_html",
        "河西学院2025年度拟申请撤销绘画本科专业。",
    ),
    SourceSeed(
        "tsinghua_2024_not_enrolling",
        "2023—2024学年度专业设置、当年新增专业、停招专业名单",
        "https://www.tsinghua.edu.cn/info/1098/108143.htm",
        "清华大学",
        "2024",
        "official_university_html",
        "清华大学2023—2024学年度专业设置页说明微机电系统工程、广告学不招生，并披露若干专业合并调整招生。",
    ),
    SourceSeed(
        "fudan_2024_stop_enrollment",
        "2024年学科专业设置、当年新增或停招专业名单",
        "https://xxgk.fudan.edu.cn/a7/62/c13422a698210/page.htm",
        "复旦大学信息公开网",
        "2024-04-19",
        "official_university_html",
        "复旦大学2024年信息公开页正文列出2024年停招应用化学、保密管理、电气工程及其自动化3个本科专业，并停招第二学士学位项目电子信息科学与技术（智能科学与技术方向）。",
    ),
    SourceSeed(
        "fudan_2023_stop_enrollment",
        "招生专业调整情况",
        "https://jwc.fudan.edu.cn/0a/ab/c25336a461483/page.htm",
        "复旦大学教务处",
        "2022-10-09",
        "official_university_html",
        "复旦大学教务处招生专业调整情况页正文列出2023年停招应用化学、电气工程及其自动化、保密管理3个专业，并链接2023专业设置PDF。",
    ),
    SourceSeed(
        "hust_2025_stop_enrollment",
        "2025年本科专业设置、新增专业、停招专业名单",
        "https://xxgk.hust.edu.cn/info/1054/3718.htm",
        "华中科技大学信息公开网",
        "2025-10-10",
        "official_university_html",
        "华中科技大学2025年本科专业一览表备注列标记17个暂停招生专业。",
    ),
    SourceSeed(
        "hust_2024_stop_enrollment_empty",
        "2024年本科专业设置、新增专业、停招专业名单",
        "https://xxgk.hust.edu.cn/info/1054/3306.htm",
        "华中科技大学信息公开网",
        "2024-10-28",
        "official_university_html",
        "华中科技大学2024年本科专业设置页已归档；解析表格未发现备注列标记停招/暂停招生的专业。",
    ),
    SourceSeed(
        "hust_2023_stop_enrollment_empty",
        "2023年本科专业设置、新增专业、停招专业名单",
        "https://xxgk.hust.edu.cn/info/1054/2976.htm",
        "华中科技大学信息公开网",
        "2023-10-27",
        "official_university_html",
        "华中科技大学2023年本科专业设置页已归档；解析表格未发现备注列标记停招/暂停招生的专业。",
    ),
    SourceSeed(
        "hust_2022_stop_enrollment_empty",
        "2022年本科专业设置、新增专业、停招专业名单",
        "https://xxgk.hust.edu.cn/info/1054/2731.htm",
        "华中科技大学信息公开网",
        "2022-10-27",
        "official_university_html",
        "华中科技大学2022年本科专业设置页已归档；解析正文/表格未发现停招/暂停招生专业。",
    ),
    SourceSeed(
        "hust_2021_stop_enrollment",
        "2021年本科专业设置、新增专业、停招专业名单",
        "https://xxgk.hust.edu.cn/info/1054/2246.htm",
        "华中科技大学信息公开网",
        "2021-10-23",
        "official_university_html",
        "华中科技大学2021年本科专业一览表备注列标记13个暂停招生专业。",
    ),
    SourceSeed(
        "hust_2020_stop_enrollment",
        "2020年本科专业设置、新增专业、停招专业名单",
        "https://xxgk.hust.edu.cn/info/1054/1802.htm",
        "华中科技大学信息公开网",
        "2020-10-22",
        "official_university_html",
        "华中科技大学2020年本科专业一览表备注列标记12个暂停招生专业和1个暂未招生专业。",
    ),
    SourceSeed(
        "hust_2019_stop_enrollment",
        "2019年本科专业设置、新增专业、停招专业名单",
        "https://xxgk.hust.edu.cn/info/1054/1803.htm",
        "华中科技大学信息公开网",
        "2019-10-23",
        "official_university_html",
        "华中科技大学2019年本科专业一览表备注列标记8个停止招生专业和1个暂未招生专业。",
    ),
    SourceSeed(
        "hust_2018_stop_enrollment",
        "2018年本科专业设置、新增专业、停招专业名单",
        "https://xxgk.hust.edu.cn/info/1054/1804.htm",
        "华中科技大学信息公开网",
        "2018-10-22",
        "official_university_html",
        "华中科技大学2018年本科专业一览表备注列标记5个已停止招生专业。",
    ),
    SourceSeed(
        "hust_2016_stop_enrollment_empty",
        "专业设置、当年新增专业、停招专业名单",
        "https://xxgk.hust.edu.cn/info/1054/1805.htm",
        "华中科技大学信息公开网",
        "2016-10-25",
        "official_university_html",
        "华中科技大学2016年专业设置页已归档；解析表格未发现备注列标记停招/暂停招生的专业。",
    ),
    SourceSeed(
        "sjtu_2024_stop_enrollment",
        "2024年专业设置及调整情况（新增、停招）",
        "https://gk.sjtu.edu.cn/Phone/View/4169",
        "上海交通大学信息公开网",
        "2024",
        "official_university_html",
        "上海交通大学信息公开网年度表格列出本科专业总数、在招专业数、新专业名单和当年停招建筑学、风景园林2个本科专业。",
    ),
    SourceSeed(
        "sjtu_2023_stop_enrollment",
        "2023年专业设置及调整情况（新增、停招）",
        "https://gk.sjtu.edu.cn/Phone/View/3369",
        "上海交通大学信息公开网",
        "2023",
        "official_university_html",
        "上海交通大学信息公开网年度表格列出本科专业总数、在招专业数、新专业名单和当年停招交通运输1个本科专业。",
    ),
    SourceSeed(
        "sjtu_2022_stop_enrollment",
        "2022年专业设置及调整情况（新增、停招）",
        "https://gk.sjtu.edu.cn/Phone/View/3272",
        "上海交通大学信息公开网",
        "2022",
        "official_university_html",
        "上海交通大学信息公开网年度表格列出本科专业总数、在招专业数、新专业名单和当年停招测控技术与仪器1个本科专业。",
    ),
    SourceSeed(
        "sjtu_2021_stop_enrollment_empty",
        "2021年专业设置及调整情况（新增、停招）",
        "https://gk.sjtu.edu.cn/Phone/View/3054",
        "上海交通大学信息公开网",
        "2021",
        "official_university_html",
        "上海交通大学信息公开网年度表格列出本科专业总数、招生专业总数和新专业名单；当年停招专业名单为“无”。",
    ),
    SourceSeed(
        "sjtu_2020_stop_enrollment",
        "2020年度专业设置及调整情况（新增、停招）",
        "https://gk.sjtu.edu.cn/Phone/View/2598",
        "上海交通大学信息公开网",
        "2020",
        "official_university_html",
        "上海交通大学信息公开网年度表格列出本科专业总数、招生专业总数、新专业名单和当年停招工程力学(海洋科学与技术)、生物科学(海洋科学)2个本科专业方向。",
    ),
    SourceSeed(
        "zufe_2026_training_system_news",
        "“新培养体系”锻造中流砥柱型人才，“一生一方案”赋能学生个性化成长",
        "https://www.zufe.edu.cn/info/1056/22597.htm",
        "浙江财经大学",
        "2026-05-18",
        "official_university_html",
        "浙江财经大学官网新闻说明学校以“增停撤改”优化专业结构，撤销2个本科专业、停招8个本科专业、新增5个本科专业、申报和预申报3个本科专业。",
    ),
    SourceSeed(
        "zufe_2024_teaching_quality_report_pdf",
        "浙江财经大学2023-2024学年本科教学质量报告",
        "https://jwc.zufe.edu.cn/__local/C/B7/63/7EBEF3A2E5AD2AD4A1F2DE8C099_C9B69026_F6B64.pdf",
        "浙江财经大学教务处",
        "2024-12-12",
        "official_university_pdf",
        "浙江财经大学2023-2024学年本科教学质量报告PDF，列出在招本科专业数、经管类专业占比、招生报到情况，并说明资产评估专业列入预警名单。",
    ),
    SourceSeed(
        "bnu_2025_stop_enrollment_index",
        "专业设置、当年新增专业、停招专业名单",
        "https://xxgk.bnu.edu.cn/xxxgkml/jxzl/zyszdnxzzytzzyms/index.htm",
        "北京师范大学信息公开网",
        "2025-10",
        "official_university_html",
        "北京师范大学信息公开目录页链接2024-2025学年专业设置、新增专业、停招专业名单PDF。",
    ),
    SourceSeed(
        "bnu_2023_stop_enrollment_pdf",
        "2022-2023学年专业设置、新增专业、停招专业名单",
        "https://xxgk.bnu.edu.cn/docs/2023-10/0b9f41bde9214f38a467cb1ab28b4800.pdf",
        "北京师范大学信息公开网",
        "2023-10",
        "official_university_pdf",
        "北京师范大学2022-2023学年专业设置、新增专业、停招专业名单PDF在备注列标记12个暂停招生专业。",
    ),
    SourceSeed(
        "bnu_2025_stop_enrollment_pdf",
        "2024-2025学年专业设置、新增专业、停招专业名单",
        "https://xxgk.bnu.edu.cn/docs/2025-10/b008f6b6195145ad9ad36764e8409a3d.pdf",
        "北京师范大学信息公开网",
        "2025-10",
        "official_university_pdf",
        "北京师范大学2024-2025学年专业设置、新增专业、停招专业名单PDF在备注列标记9个暂停招生专业。",
    ),
    SourceSeed(
        "uestc_2017_stop_enrollment",
        "电子科技大学专业设置、当年新增专业及停招专业情况",
        "https://xxgkw.uestc.edu.cn/info/1073/3842.htm",
        "电子科技大学信息公开网",
        "2019-03-20",
        "official_university_html",
        "电子科技大学信息公开页说明截至2017年8月学校共有本科专业66个，2017年新增5个专业、停招环境工程1个专业。",
    ),
    SourceSeed(
        "ecnu_2023_stop_enrollment",
        "2023年本科专业设置及新增专业、停招专业情况",
        "https://xxgk.ecnu.edu.cn/63/ec/c11846a549868/page.htm",
        "华东师范大学信息公开网",
        "2023-10-13",
        "official_university_html",
        "华东师范大学2023年本科专业一览表在备注列标记16个暂停招生专业。",
    ),
    SourceSeed(
        "ecnu_2022_stop_enrollment_empty",
        "2022年本科专业设置及新增专业、停招专业情况",
        "https://xxgk.ecnu.edu.cn/dd/5b/c11846a384347/page.htm",
        "华东师范大学信息公开网",
        "2022-10-01",
        "official_university_html",
        "华东师范大学2022年本科专业设置页已归档；解析表格未发现备注列标记暂停招生的专业。",
    ),
    SourceSeed(
        "ecnu_2021_stop_enrollment_empty",
        "2021年本科专业设置及新增专业、停招专业情况",
        "https://xxgk.ecnu.edu.cn/9c/8c/c11846a433292/page.htm",
        "华东师范大学信息公开网",
        "2021-10-21",
        "official_university_html",
        "华东师范大学2021年本科专业设置页已归档；解析表格未发现备注列标记暂停招生的专业。",
    ),
    SourceSeed(
        "ecnu_2020_stop_enrollment_empty",
        "2020年本科专业设置及新增专业、停招专业情况",
        "https://xxgk.ecnu.edu.cn/2b/63/c11846a338787/page.htm",
        "华东师范大学信息公开网",
        "2020-10-22",
        "official_university_html",
        "华东师范大学2020年本科专业设置页已归档；解析表格未发现备注列标记暂停招生的专业。",
    ),
    SourceSeed(
        "ecnu_2019_stop_enrollment_empty",
        "2019年本科专业设置及新增专业、停招专业情况",
        "https://xxgk.ecnu.edu.cn/d9/4d/c11846a317773/page.htm",
        "华东师范大学信息公开网",
        "2019-08-31",
        "official_university_html",
        "华东师范大学2019年本科专业设置页已归档；解析表格未发现备注列标记暂停招生的专业。",
    ),
    SourceSeed(
        "ecnu_2018_stop_enrollment_empty",
        "2018年本科专业设置及新增专业、停招专业情况",
        "https://xxgk.ecnu.edu.cn/d9/4b/c11846a317771/page.htm",
        "华东师范大学信息公开网",
        "2018-10-23",
        "official_university_html",
        "华东师范大学2018年本科专业设置页已归档；解析表格未发现备注列标记暂停招生的专业。",
    ),
    SourceSeed(
        "ecnu_2017_stop_enrollment_empty",
        "2017年本科专业设置及新增专业、停招专业情况",
        "https://xxgk.ecnu.edu.cn/3f/dc/c11846a147420/page.htm",
        "华东师范大学信息公开网",
        "2017-10-20",
        "official_university_html",
        "华东师范大学2017年本科专业设置页已归档；解析表格未发现备注列标记暂停招生的专业。",
    ),
    SourceSeed(
        "ecnu_2016_stop_enrollment_empty",
        "2016年本科专业设置及新增专业、停招专业情况",
        "https://xxgk.ecnu.edu.cn/3f/db/c11846a147419/page.htm",
        "华东师范大学信息公开网",
        "2016-10-19",
        "official_university_html",
        "华东师范大学2016年本科专业设置页已归档；解析表格未发现备注列标记暂停招生的专业。",
    ),
    SourceSeed(
        "ecnu_2015_stop_enrollment_empty",
        "2015年本科专业设置及新增专业、停招专业情况",
        "https://xxgk.ecnu.edu.cn/3f/da/c11846a147418/page.htm",
        "华东师范大学信息公开网",
        "2015-10-28",
        "official_university_html",
        "华东师范大学2015年本科专业设置页已归档；解析表格未发现备注列标记暂停招生的专业。",
    ),
    SourceSeed(
        "ecnu_2014_stop_enrollment_empty",
        "2014年本科专业设置及新增专业、停招专业情况",
        "https://xxgk.ecnu.edu.cn/3f/d9/c11846a147417/page.htm",
        "华东师范大学信息公开网",
        "2014-09-24",
        "official_university_html",
        "华东师范大学2014年本科专业设置页已归档；解析表格未发现备注列标记暂停招生的专业。",
    ),
    SourceSeed(
        "ecnu_2024_stop_enrollment",
        "2024年本科专业设置及新增专业、停招专业情况",
        "https://xxgk.ecnu.edu.cn/cb/f5/c29049a642037/page.htm",
        "华东师范大学信息公开网",
        "2024-10-15",
        "official_university_html",
        "华东师范大学2024年本科专业一览表在备注列标记22个暂停招生专业。",
    ),
    SourceSeed(
        "ecnu_2025_stop_enrollment",
        "2025年本科专业设置及新增专业、停招专业情况",
        "https://xxgk.ecnu.edu.cn/05/bb/c29049a722363/page.htm",
        "华东师范大学信息公开网",
        "2025-10-16",
        "official_university_html",
        "华东师范大学2025年信息公开页称停招专业24个；正文和meta描述实际列出22个专业名称，本数据集仅结构化这些明示名称。",
    ),
    SourceSeed(
        "sbs_2025_stop_enrollment",
        "2025年本科专业新增、停招情况",
        "https://xxgk.sbs.edu.cn/xxgklm/jxgl/zysz/28fc700b64fe4920802336a75edf32de.htm",
        "上海商学院信息公开网",
        "2025-09-27",
        "official_university_html",
        "上海商学院信息公开页列出2025年停招本科专业2个：园林、广告学。",
    ),
    SourceSeed(
        "shou_2025_stop_enrollment",
        "2025年新增和停招专业",
        "https://xxgk.shou.edu.cn/2025/1028/c7960a347413/page.htm",
        "上海海洋大学信息公开网",
        "2025-06-10",
        "official_university_html",
        "上海海洋大学信息公开页列出2025年停招行政管理、物流管理、工业工程、软件工程、朝鲜语5个本科专业。",
    ),
    SourceSeed(
        "shou_2024_teaching_quality_report_pdf",
        "上海海洋大学2023-2024学年本科教学质量报告",
        "https://xxgk.shou.edu.cn/_upload/article/files/b0/fd/c8b45f9a4463b270a7cdb8a3010e/4cce9f87-c6fa-4186-81eb-58ebdc7e1ec7.pdf",
        "上海海洋大学信息公开网",
        "2024",
        "official_university_pdf",
        "上海海洋大学2023-2024学年本科教学质量报告说明2023学年新增人工智能、环境科学与工程2个专业，并停招环境科学、环境工程、文化产业管理3个专业。",
    ),
    SourceSeed(
        "shu_2024_teaching_quality_report_pdf",
        "上海大学2023-2024学年本科教育教学质量报告",
        "https://oiqa.shu.edu.cn/2023-2024bkzlbgzz.pdf",
        "上海大学教育质量考察与评估办公室",
        "2024",
        "official_university_pdf",
        "上海大学2023-2024学年本科教育教学质量报告列出本科专业总数101个、在招专业数98个，已停招材料物理、工业设计、包装工程3个专业，并申请备案撤销包装工程专业。",
    ),
    SourceSeed(
        "imufe_2024_teaching_quality_report_pdf",
        "内蒙古财经大学2023-2024学年本科教学质量报告",
        "https://www.imufe.edu.cn/__local/4/F2/1B/87C0F11AE202E6B1C81E48CF3F7_6EEB1DCB_97E82.pdf",
        "内蒙古财经大学",
        "2024",
        "official_university_pdf",
        "内蒙古财经大学2023-2024学年本科教学质量报告列出本科专业55个、实际招生本科专业48个、暂停招生专业7个，并说明2023年撤销信息与计算科学等5个专业、物业管理等6个专业暂缓招生，2024年保险学和公共事业管理暂缓招生。",
    ),
    SourceSeed(
        "cqwu_2024_teaching_quality_report_pdf",
        "重庆文理学院2023-2024学年本科教学质量报告",
        "https://www.cqwu.edu.cn/SITE_ATTACHE/web2025_cqwu_edu_cn/2025-12-22/upload/file/20251222/dPca_1766409101570024524.pdf",
        "重庆文理学院",
        "2024",
        "official_university_pdf",
        "重庆文理学院2023-2024学年本科教学质量报告说明2024年主动撤销经济统计学专业，停招广播电视学、美术学、旅游管理与服务教育、金融数学等4个专业，并新增智慧农业和智能车辆工程2个专业。",
    ),
    SourceSeed(
        "nepu_2024_teaching_quality_report_pdf",
        "东北石油大学2023-2024学年本科教学质量报告",
        "https://xxgk.nepu.edu.cn/__local/1/A8/AF/4B35D620C85D4B952995AC45122_C48B4B19_1F6310.pdf",
        "东北石油大学信息公开网",
        "2024",
        "official_university_pdf",
        "东北石油大学2023-2024学年本科教学质量报告列出本科专业总数68个、当年招生专业53个，新增地球信息科学与技术、储能科学与工程2个专业，停招光电信息科学与工程、智能电网信息工程、财务管理3个专业，并撤销公共事业管理专业。",
    ),
    SourceSeed(
        "cqust_2024_teaching_quality_report_pdf",
        "重庆科技大学2023-2024学年本科教学质量报告",
        "https://xxgk.cqust.edu.cn/__local/F/9D/4C/89B8F4D48FA2427C9006E75B03B_D35F7597_F2269.pdf",
        "重庆科技大学信息公开网",
        "2024",
        "official_university_pdf",
        "重庆科技大学2023-2024学年本科教学质量报告列出本科专业70个、当年本科招生专业61个、9个专业暂停招生、当年停招7个；正文逐名提及矿物加工工程、过程装备与控制工程、地球物理学、材料物理等停招或暂停招生专业。",
    ),
    SourceSeed(
        "jxycu_2024_teaching_quality_report_pdf",
        "宜春学院2023-2024学年本科教学质量报告",
        "https://www.jxycu.edu.cn/_upload/article/files/99/1c/49e7b85f45a09687db7cb419045b/414725be-0fef-4ce0-8950-61478bd5b86a.pdf",
        "宜春学院",
        "2024",
        "official_university_pdf",
        "宜春学院2023-2024学年本科教学质量报告说明停招广播电视编导等14个专业，交叉融合生物工程和制药工程2个专业，并称2024年招生专业总数减少至53个。",
    ),
    SourceSeed(
        "sdua_2024_teaching_quality_report_pdf",
        "山东航空学院2023-2024学年本科教学质量报告",
        "https://xxgk.sdua.edu.cn/_upload/article/files/89/85/555b9b7042efbdbcfd4d53feef9a/956bed86-8afd-4147-86e8-e29d39b301d0.pdf",
        "山东航空学院信息公开网",
        "2024",
        "official_university_pdf",
        "山东航空学院2023-2024学年本科教学质量报告列出全校本科专业59个、当年招生专业52个、新增空间信息与数字技术、停招法语和飞行器适航技术，并说明撤销能源化学工程专业。",
    ),
    SourceSeed(
        "ustb_2024_teaching_quality_report_pdf",
        "北京科技大学2023-2024学年本科教学质量报告",
        "https://xxgk.ustb.edu.cn/attach/file/xinxigongkaimulu/jiaoxuezhiliang/benkejiaoxue/2024-12-05/1c217c649caaf1920249462747783058.pdf",
        "北京科技大学信息公开网",
        "2024",
        "official_university_pdf",
        "北京科技大学2023-2024学年本科教学质量报告说明学校设有61个本科专业，本科专业中共有思想政治教育、生态学、电子信息工程、矿物资源工程、工业工程、智能科学与技术6个专业停招。",
    ),
    SourceSeed(
        "gufe_2024_teaching_quality_report_pdf",
        "贵州财经大学2023-2024学年本科教学质量报告",
        "https://xxgk.gufe.edu.cn/__local/7/79/72/3C5D937F9F37FF057F4123005E4_962F6ADF_8AAEF.pdf",
        "贵州财经大学信息公开网",
        "2024",
        "official_university_pdf",
        "贵州财经大学2023-2024学年本科教学质量报告列出本科专业64个、当年招生专业54个、停招专业10个，并说明申请撤销英语、会展经济与管理、金融科技3个专业。",
    ),
    SourceSeed(
        "sues_2024_teaching_quality_report_pdf",
        "上海工程技术大学2023-2024学年本科教学质量报告",
        "https://infopub.sues.edu.cn/_upload/article/files/eb/77/2cf61d504f38b0e67bfb5ade0c90/789febc0-a022-4087-b7c1-6609a5a362fd.pdf",
        "上海工程技术大学信息公开网",
        "2024",
        "official_university_pdf",
        "上海工程技术大学2023-2024学年本科教学质量报告说明2023年发布本科专业综合得分并对5个专业预警，持续推进专业动态调整，撤销及停招专业4个；报告未逐名列出这4个专业。",
    ),
    SourceSeed(
        "sta_2025_stop_enrollment",
        "2025-2026学年学科专业建设、当年新增或停招专业名单",
        "https://xxgk.sta.edu.cn/d6/80/c3356a120448/page.htm",
        "上海戏剧学院信息公开网",
        "2025-10-27",
        "official_university_html",
        "上海戏剧学院信息公开页列出2025-2026学年当年停招动画、数字演艺设计2个本科专业。",
    ),
    SourceSeed(
        "lixin_2025_stop_enrollment",
        "2025年上海立信会计金融学院当年新增或停招专业名单",
        "https://www.lixin.edu.cn/gk/xxgklm/jhgl/hkzysz/dnxchtqzymc/150259.htm",
        "上海立信会计金融学院信息公开网",
        "2025-10-27",
        "official_university_html",
        "上海立信会计金融学院信息公开页说明2025年停招专业，并链接本科专业设置情况。",
    ),
    SourceSeed(
        "lixin_2025_undergrad_setting",
        "2025年上海立信会计金融学院本科专业设置情况",
        "https://www.lixin.edu.cn/gk/xxgklm/jhgl/hkzysz/hkzysz/150261.htm",
        "上海立信会计金融学院信息公开网",
        "2025-10-27",
        "official_university_html",
        "上海立信会计金融学院本科专业设置表标注金融工程（中美合作）、信用管理、房地产开发与管理、劳动与社会保障、日语等专业2025年停招。",
    ),
    SourceSeed(
        "tongji_2025_stop_enrollment",
        "2025级停招专业",
        "https://xxgk.tongji.edu.cn/index.php?classid=4368&newsid=20275&t=show",
        "同济大学信息公开网",
        "2025-09-16",
        "official_university_html",
        "同济大学信息公开页列出2025级停招海洋科学、汽车服务工程、物流管理、视觉传达设计、环境设计、产品设计6个本科专业。",
    ),
    SourceSeed(
        "tongji_2024_stop_enrollment",
        "2024级停招专业",
        "https://xxgk.tongji.edu.cn/index.php?classid=4368&newsid=18856&t=show",
        "同济大学信息公开网",
        "2024-09-23",
        "official_university_html",
        "同济大学信息公开页列出2024级停招广播电视学1个本科专业。",
    ),
    SourceSeed(
        "tongji_2023_stop_enrollment",
        "2023级停招专业名单",
        "https://xxgk.tongji.edu.cn/index.php?classid=4368&newsid=17241&t=show",
        "同济大学信息公开网",
        "2023-09-26",
        "official_university_html",
        "同济大学信息公开页列出2023级停招市场营销、行政管理、国际经济与贸易、化学工程与工艺、电子信息工程、海洋技术、港口航道与海岸工程、广告学、工业工程9个本科专业。",
    ),
    SourceSeed(
        "tongji_2022_stop_enrollment",
        "2022级停招专业名单",
        "https://xxgk.tongji.edu.cn/index.php?classid=4368&newsid=14481&t=show",
        "同济大学信息公开网",
        "2023-05-22",
        "official_university_html",
        "同济大学信息公开页列出2022级停招临床医学（拔尖卓越培养）、社会学2个本科专业。",
    ),
    SourceSeed(
        "tongji_2021_stop_enrollment",
        "2021级停招专业名单",
        "https://xxgk.tongji.edu.cn/index.php?classid=4368&newsid=13571&t=show",
        "同济大学信息公开网",
        "2021-10-26",
        "official_university_html",
        "同济大学信息公开页列出2021级停招康复治疗学1个本科专业，并提供校内专业编号。",
    ),
    SourceSeed(
        "tongji_2020_stop_enrollment",
        "2020级停招专业名单",
        "https://xxgk.tongji.edu.cn/index.php?classid=4368&newsid=12415&t=show",
        "同济大学信息公开网",
        "2020-10-29",
        "official_university_html",
        "同济大学信息公开页列出2020级停招物流工程、电子科学与技术、地质学3个本科专业，并提供校内专业编号。",
    ),
    SourceSeed(
        "ahtcm_2024_stop_enrollment",
        "2024年停招专业情况说明",
        "https://jwc.ahtcm.edu.cn/info/1271/20571.htm",
        "安徽中医药大学教务处（招生办）",
        "2024-10-17",
        "official_university_html",
        "安徽中医药大学教务处信息公开页说明2024年药物分析、中药资源与开发、中医儿科学、保险学专业停招。",
    ),
    SourceSeed(
        "ahtcm_2024_major_setting",
        "2024年专业设置情况",
        "https://jwc.ahtcm.edu.cn/info/1241/19951.htm",
        "安徽中医药大学教务处（招生办）",
        "2024-10-17",
        "official_university_html",
        "安徽中医药大学2024年专业设置情况表列出32个本科专业，并在招生状态列标注药物分析、保险学、中药资源与开发、中医儿科学停招。",
    ),
    SourceSeed(
        "ahtcm_2025_stop_enrollment",
        "2025年停招专业情况说明",
        "https://jwc.ahtcm.edu.cn/info/1271/24251.htm",
        "安徽中医药大学教务处（招生办）",
        "2025-10-17",
        "official_university_html",
        "安徽中医药大学教务处信息公开页说明2025年中医儿科学、保险学、人力资源管理专业停招。",
    ),
    SourceSeed(
        "ahtcm_2025_major_setting",
        "2025年本科专业设置情况",
        "https://jwc.ahtcm.edu.cn/info/1241/24141.htm",
        "安徽中医药大学教务处（招生办）",
        "2025-10-16",
        "official_university_html",
        "安徽中医药大学2025年本科专业设置情况表列出32个本科专业，并在招生状态列标注人力资源管理、保险学、中医儿科学停招。",
    ),
    SourceSeed(
        "shiep_2019_stop_enrollment",
        "2019年上海电力大学本科新增或停招专业名单",
        "https://xxgk.shiep.edu.cn/e2/23/c6157a254499/page.htm",
        "上海电力大学信息公开网",
        "2019-03-27",
        "official_university_html",
        "上海电力大学信息公开页列出2019年停招电力工程与管理专业。",
    ),
    SourceSeed(
        "shiep_2021_stop_enrollment",
        "2021年新增或停招专业名单",
        "https://xxgk.shiep.edu.cn/e2/20/c6157a254496/page.htm",
        "上海电力大学信息公开网",
        "2021-05-18",
        "official_university_html",
        "上海电力大学信息公开页列出2021年停招公共事业管理专业。",
    ),
    SourceSeed(
        "shiep_2023_stop_enrollment",
        "2023年新增、停招本科专业情况",
        "https://xxgk.shiep.edu.cn/e2/1c/c6157a254492/page.htm",
        "上海电力大学信息公开网",
        "2023-09-26",
        "official_university_html",
        "上海电力大学信息公开页列出2023年停招公共事业管理、材料科学与工程、网络工程，并给出专业代码和停招年份。",
    ),
    SourceSeed(
        "shiep_2024_teaching_quality_report_pdf",
        "上海电力大学2023-2024学年本科教学质量报告",
        "https://jwc.shiep.edu.cn/_upload/article/files/10/cd/e2308a7c4d35985b7aafdd205898/50b33fab-b7a4-412b-b05c-50d0ce829425.pdf",
        "上海电力大学本科生院",
        "2024",
        "official_university_pdf",
        "上海电力大学2023-2024学年本科教学质量报告正文和附表列出已停招专业，含机械电子工程、材料化学、材料科学与工程、网络工程、公共事业管理、物流管理、日语。",
    ),
    SourceSeed(
        "shiep_2025_stop_enrollment",
        "2025年新增、停招本科专业情况",
        "https://jwc.shiep.edu.cn/13/62/c238a267106/page.htm",
        "上海电力大学本科生院",
        "2025-09-25",
        "official_university_html",
        "上海电力大学本科生院保留2025年新增、停招本科专业情况信息公开页；公开HTML正文未暴露完整名单，需结合可抓取转载文本核验。",
    ),
    SourceSeed(
        "shiep_2026_professional_adjustment_meeting",
        "上海电力大学教学指导委员会召开2026年第一次会议",
        "https://www.shiep.edu.cn/30/24/c6407a274468/page.htm",
        "上海电力大学",
        "2026-03-25",
        "official_university_html",
        "上海电力大学官网新闻列出2026年专业优化调整方案，拟停招环境工程、光电信息科学与工程、信息与计算科学，拟撤销公共事业管理。",
    ),
    SourceSeed(
        "sohu_2025_shanghai_university_stop_summary",
        "专业大洗牌！停招多个本科专业，就业堪忧？",
        "https://www.sohu.com/a/952545365_121124337",
        "搜狐转载",
        "2025-11-09",
        "secondary_news_html",
        "可抓取二级汇总来源，列出浙江财经大学、上海海洋大学、上海戏剧学院、上海立信会计金融学院、安徽中医药大学、华东师范大学、上海电力大学、同济大学、上海理工大学等2025年停招/撤销专业信息。",
    ),
    SourceSeed(
        "usst_2025_adjustment_article",
        "中国科技网：培养学生“真本领”上海理工大学启动综合教改",
        "https://www.usst.edu.cn/2025/1105/c965a65807/page.htm",
        "上海理工大学官网转载中国科技网",
        "2025-11-03",
        "official_university_html",
        "上海理工大学官网转载稿明确2025年撤销假肢矫形工程、网络工程2个专业，并称停招制药工程等4个专业。",
    ),
    SourceSeed(
        "usst_undergrad_major_setting",
        "上海理工大学本科专业汇总表",
        "https://jwc.usst.edu.cn/zysz/list.htm",
        "上海理工大学教务处",
        "2025",
        "official_university_html",
        "上海理工大学教务处本科专业汇总表列出网络工程2019年停招、假肢矫形工程2020年停招，并提供专业代码和所属学院。",
    ),
    SourceSeed(
        "chu_2019_major_setting",
        "巢湖学院2019年本科专业设置及停招专业情况一览表",
        "https://www.chu.edu.cn/xxgk/2019/1029/c5128a92232/page.htm",
        "巢湖学院信息公开网",
        "2019-10-30",
        "official_university_html",
        "巢湖学院信息公开页表格列出2019年停招专业及专业代码、修业年限、学位门类和所属学院。",
    ),
    SourceSeed(
        "chu_2020_major_setting",
        "巢湖学院2020年本科专业设置及新增专业、停招专业情况一览表",
        "https://www.chu.edu.cn/xxgk/2020/1030/c5128a106088/page.htm",
        "巢湖学院信息公开网",
        "2020-10-30",
        "official_university_html",
        "巢湖学院信息公开页表格列出2020年停招专业及专业代码、修业年限、学位门类和所属学院。",
    ),
    SourceSeed(
        "chu_2021_major_setting",
        "巢湖学院2021年本科专业设置及新增专业、停招专业情况一览表",
        "https://www.chu.edu.cn/xxgk/2021/1020/c5128a124448/page.htm",
        "巢湖学院信息公开网",
        "2021-10-20",
        "official_university_html",
        "巢湖学院信息公开页表格列出2021年停招专业及专业代码、修业年限、学位门类和所属学院。",
    ),
    SourceSeed(
        "chu_2022_major_setting",
        "巢湖学院2022年本科专业设置及新增专业、停招专业情况一览表",
        "https://www.chu.edu.cn/xxgk/2022/1009/c5128a144314/page.htm",
        "巢湖学院信息公开网",
        "2022-10-09",
        "official_university_html",
        "巢湖学院信息公开页表格列出2022年停招专业及专业代码、修业年限、学位门类和所属学院。",
    ),
    SourceSeed(
        "chu_2023_major_setting_notice",
        "巢湖学院2023年本科专业设置及新增专业、停招专业情况一览表",
        "https://www.chu.edu.cn/xxgk/2023/1025/c5128a168196/page.htm",
        "巢湖学院信息公开网",
        "2023-07-12",
        "official_university_html",
        "巢湖学院信息公开页链接2023年本科专业设置及新增专业、停招专业情况PDF附件。",
    ),
    SourceSeed(
        "chu_2023_major_setting_pdf",
        "巢湖学院2023年本科专业设置及新增专业、停招专业情况一览表PDF",
        "https://www.chu.edu.cn/_upload/article/files/45/02/80246e934c6c97eef2630938aece/7ddad0e6-7b1c-40a8-94e5-37535131c7d7.pdf",
        "巢湖学院信息公开网",
        "2023-07-12",
        "official_university_pdf",
        "巢湖学院2023年本科专业设置PDF逐项列出专业代码、修业年限、学位门类、所属学院和2023年停招标记。",
    ),
    SourceSeed(
        "chu_2024_major_setting_notice",
        "巢湖学院2024年本科专业设置及新增专业、停招专业情况一览表",
        "https://www.chu.edu.cn/xxgk/2024/1016/c5128a186150/page.htm",
        "巢湖学院信息公开网",
        "2024-10-16",
        "official_university_html",
        "巢湖学院信息公开页链接2024年本科专业设置及新增专业、停招专业情况PDF附件。",
    ),
    SourceSeed(
        "chu_2024_major_setting_pdf",
        "巢湖学院2024年本科专业设置及新增专业、停招专业情况一览表PDF",
        "https://www.chu.edu.cn/_upload/article/files/b0/f1/aa5cdc8f440997bdbeb3942ec932/146007f0-7e0c-4f84-8702-7e4b45346ad0.pdf",
        "巢湖学院信息公开网",
        "2024-10-16",
        "official_university_pdf",
        "巢湖学院2024年本科专业设置PDF逐项列出专业代码、修业年限、学位门类、所属学院和2024年停招标记。",
    ),
    SourceSeed(
        "chu_2025_major_setting_notice",
        "巢湖学院2025年本科专业设置及新增专业、停招专业情况一览表",
        "https://www.chu.edu.cn/xxgk/2025/0603/c5128a194985/page.htm",
        "巢湖学院信息公开网",
        "2025-06-03",
        "official_university_html",
        "巢湖学院信息公开页链接2025年本科专业设置及新增专业、停招专业情况PDF附件。",
    ),
    SourceSeed(
        "chu_2025_major_setting_pdf",
        "巢湖学院2025年本科专业设置及新增专业、停招专业情况一览表PDF",
        "https://www.chu.edu.cn/_upload/article/files/03/8b/a428ae954bd883c1b5dde071074c/e4b3e251-0474-462c-9561-50ca8f2603a2.pdf",
        "巢湖学院信息公开网",
        "2025-06-03",
        "official_university_pdf",
        "巢湖学院2025年本科专业设置PDF逐项列出专业代码、修业年限、学位门类、所属学院和2025年停招标记。",
    ),
    SourceSeed(
        "szu_2025_major_setting_notice",
        "专业设置、当年新增专业、停招专业名单",
        "https://xxgk.szu.edu.cn/jxzl1/zysz_dnxzzy_tzzymd.htm",
        "深圳大学信息公开网",
        "2025-11-03",
        "official_university_html",
        "深圳大学信息公开页链接专业设置、当年新增专业、停招专业名单Word附件。",
    ),
    SourceSeed(
        "szu_2025_major_setting_docx",
        "专业设置、当年新增专业、停招专业名单25-11-3",
        "https://xxgk.szu.edu.cn/system/_content/download.jsp?urltype=news.DownloadAttachUrl&owner=1637113499&wbfileid=4800713",
        "深圳大学信息公开网",
        "2025-11-03",
        "official_university_docx",
        "深圳大学2025年专业设置Word附件逐项列出专业代码、专业名称和备注，其中标注当年停招或往年停招。",
    ),
    SourceSeed(
        "jiangnan_2025_major_setting_notice",
        "专业设置、当年新增专业、停招专业名单",
        "https://xxgk.jiangnan.edu.cn/info/1007/1108.htm",
        "江南大学信息公开网",
        "2025-10",
        "official_university_html",
        "江南大学信息公开页链接2025年本科专业设置、当年新增专业、停招专业名单PDF附件。",
    ),
    SourceSeed(
        "jiangnan_2025_major_setting_pdf",
        "江南大学2025年本科专业设置、当年新增专业、停招专业名单",
        "https://xxgk.jiangnan.edu.cn/system/_content/download.jsp?urltype=news.DownloadAttachUrl&owner=1246629957&wbfileid=14958817",
        "江南大学信息公开网",
        "2025-10",
        "official_university_pdf",
        "江南大学2025年专业设置PDF标注国际经济与贸易2025年停招，并标注汉语言文学2025年停招师范班级、保留非师范招生。",
    ),
    SourceSeed(
        "upc_2025_major_setting_pdf",
        "2025年本科专业设置情况（含新增专业、停招专业）",
        "https://jwc.upc.edu.cn/_upload/article/files/a4/c8/62f4fa1849e6b7a1bbe19b7cc241/c0aca6e6-57ec-42e2-8d7d-0b452e940a6a.pdf",
        "中国石油大学（华东）教务处",
        "2025",
        "official_university_pdf",
        "中国石油大学（华东）2025年本科专业设置PDF在备注列标注机械工程、土木工程、材料物理、材料化学、地理信息科学、物联网工程、市场营销停招且有在校生。",
    ),
    SourceSeed(
        "jci_2024_major_setting_notice",
        "专业设置、近三年新增和停招专业名单",
        "https://jwc.jci.edu.cn/info/1132/4429.htm",
        "景德镇陶瓷大学教务处",
        "2024-05-31",
        "official_university_html",
        "景德镇陶瓷大学教务处页面链接近三年新增和停招专业表、本科专业设置一览表附件。",
    ),
    SourceSeed(
        "jci_undergrad_major_setting_docx",
        "景德镇陶瓷大学本科专业设置一览表",
        "https://jwc.jci.edu.cn/system/_content/download.jsp?urltype=news.DownloadAttachUrl&owner=1412906336&wbfileid=3317203",
        "景德镇陶瓷大学教务处",
        "2024-05-31",
        "official_university_docx",
        "景德镇陶瓷大学本科专业设置一览表Word附件在备注列标注物流管理、公共事业管理为21停招，金融工程、翻译为22停招。",
    ),
    SourceSeed(
        "jci_2024_academic_year_major_status_notice",
        "2023-2024学年本科专业设置信息",
        "https://jwc.jci.edu.cn/info/1132/5089.htm",
        "景德镇陶瓷大学教务处",
        "2025-03-10",
        "official_university_html",
        "景德镇陶瓷大学教务处信息公开页说明2023-2024学年学校共有60个本科专业、53个招生专业、7个停招专业，并逐名列出停招名单。",
    ),
    SourceSeed(
        "jci_2024_academic_year_major_setting_docx",
        "景德镇陶瓷大学本科专业设置一览表（2023-2024学年）",
        "https://jwc.jci.edu.cn/system/_content/download.jsp?urltype=news.DownloadAttachUrl&owner=1412906336&wbfileid=3354102",
        "景德镇陶瓷大学教务处",
        "2025-03-10",
        "official_university_docx",
        "景德镇陶瓷大学2023-2024学年本科专业设置一览表Word附件提供停招名单中7个专业的专业代码、学位门类和表内序号。",
    ),
    SourceSeed(
        "cqnu_2023_major_setting_notice",
        "专业设置、当年新增专业、停招专业名单（2022-2023学年度）",
        "https://xxgk.cqnu.edu.cn/info/1229/4123.htm",
        "重庆师范大学信息公开网",
        "2023-10-20",
        "official_university_html",
        "重庆师范大学信息公开页链接2022-2023学年度专业设置、当年新增专业、停招专业名单Word附件。",
    ),
    SourceSeed(
        "cqnu_2023_major_setting_docx",
        "专业设置、当年新增专业、停招专业名单（2022-2023学年度）DOCX",
        "https://xxgk.cqnu.edu.cn/system/_content/download.jsp?urltype=news.DownloadAttachUrl&owner=1130054770&wbfileid=5047615",
        "重庆师范大学信息公开网",
        "2023-10-20",
        "official_university_docx",
        "重庆师范大学2022-2023学年度专业设置Word附件按专业列出专业代码、所在学院、授予学位和2022年招生情况，其中7个专业标注停招。",
    ),
    SourceSeed(
        "cqnu_2024_major_setting_html",
        "专业设置、当年新增专业、停招专业名单（2023-2024学年度）",
        "https://xxgk.cqnu.edu.cn/info/1229/4369_1.htm",
        "重庆师范大学信息公开网",
        "2024-11-13",
        "official_university_html",
        "重庆师范大学2023-2024学年度专业设置HTML分页表按专业列出专业代码、所在学院和招生情况，其中10个专业标注停招。",
    ),
    SourceSeed(
        "scnu_2025_stop_enrollment_pdf",
        "华南师范大学2025本科专业设置、新增专业、停招专业一览表",
        "https://statics.scnu.edu.cn/pics/xxgk/2025/1027/1761552190573638.pdf",
        "华南师范大学信息公开网",
        "2025-10-27",
        "official_university_pdf",
        "华南师范大学2025本科专业设置、新增专业、停招专业一览表逐项标注招生、停招、未招状态；结构化其中停招专业。",
    ),
    SourceSeed(
        "sufe_2025_stop_enrollment_notice",
        "2025上海财经大学本科专业设置情况（含停招专业情况）",
        "https://gongkai.sufe.edu.cn/c0/b7/c12259a245943/page.htm",
        "上海财经大学信息公开网",
        "2025-09-03",
        "official_university_html",
        "上海财经大学2025本科专业设置情况信息公开页；页面链接本科专业设置PDF，结构化记录以PDF附件为准。",
    ),
    SourceSeed(
        "sufe_2025_stop_enrollment_pdf",
        "2025上海财经大学本科专业设置情况（含停招专业情况）",
        "https://gongkai.sufe.edu.cn/_upload/article/files/b9/ec/ce4886d5414db59f3aa08772fe13/f8e82408-7b6d-4cfa-8fd0-4b061e26182a.pdf",
        "上海财经大学信息公开网",
        "2025-09-03",
        "official_university_pdf",
        "上海财经大学2025本科专业设置PDF用◆标注已停招专业；结构化其中12个已停招专业。",
    ),
    SourceSeed(
        "sufe_2024_stop_enrollment",
        "2024上海财经大学本科专业设置情况（含新增、停招专业情况）",
        "https://gongkai.sufe.edu.cn/62/6f/c12259a221807/page.htm",
        "上海财经大学信息公开网",
        "2024-09-04",
        "official_university_html",
        "上海财经大学2024本科专业设置HTML表以◆标注已停招专业；结构化其中11个已停招专业。",
    ),
    SourceSeed(
        "sufe_2023_stop_enrollment",
        "2023上海财经大学本科专业设置情况（含停招专业情况）",
        "https://gongkai.sufe.edu.cn/21/2e/c12259a205102/page.htm",
        "上海财经大学信息公开网",
        "2023-09-25",
        "official_university_html",
        "上海财经大学2023本科专业设置HTML表以◆标注已停招专业；结构化其中11个已停招专业。",
    ),
    SourceSeed(
        "sufe_2022_stop_enrollment",
        "2022上海财经大学本科专业设置情况（含停招专业情况）",
        "https://gongkai.sufe.edu.cn/d9/e5/c12259a186853/page.htm",
        "上海财经大学信息公开网",
        "2022-09-22",
        "official_university_html",
        "上海财经大学2022本科专业设置HTML表以◆标注已停招专业；结构化其中12个已停招专业。",
    ),
    SourceSeed(
        "sufe_2021_stop_enrollment",
        "2020-2021学年学科专业设置、当年新增或停招专业名单",
        "https://gongkai.sufe.edu.cn/5e/e8/c12259a155368/page.htm",
        "上海财经大学信息公开网",
        "2021-10-25",
        "official_university_html",
        "上海财经大学2020-2021学年本科专业设置HTML表以◆标注已停招专业；结构化其中12个已停招专业。",
    ),
    SourceSeed(
        "sufe_2020_stop_enrollment",
        "上海财经大学本科专业设置情况（2020）",
        "https://gongkai.sufe.edu.cn/13/02/c12259a135938/page.htm",
        "上海财经大学信息公开网",
        "2020-10-29",
        "official_university_html",
        "上海财经大学2020本科专业设置HTML表以◆标注已停招专业；结构化其中9个已停招专业。",
    ),
    SourceSeed(
        "sufe_2019_stop_enrollment",
        "上海财经大学本科专业设置情况（2019）",
        "https://gongkai.sufe.edu.cn/f4/3c/c12259a128060/page.htm",
        "上海财经大学信息公开网",
        "2019-10-23",
        "official_university_html",
        "上海财经大学2019本科专业设置HTML表以■标注停招专业；结构化其中11个停招专业。",
    ),
    SourceSeed(
        "sufe_2018_stop_enrollment",
        "上海财经大学本科专业设置情况（2018）",
        "https://gongkai.sufe.edu.cn/85/87/c12259a99719/page.htm",
        "上海财经大学信息公开网",
        "2018-10-18",
        "official_university_html",
        "上海财经大学2018本科专业设置HTML表以***标注已停招、以****标注2018年停止招生；结构化其中12个停招/停止招生专业。",
    ),
    SourceSeed(
        "sufe_2017_stop_enrollment",
        "本科专业设置、当年新增专业、停招专业名单",
        "https://gongkai.sufe.edu.cn/2c/7f/c12259a76927/page.htm",
        "上海财经大学信息公开网",
        "2017-10-24",
        "official_university_html",
        "上海财经大学2017本科专业设置HTML表以***标注2017年停止招生专业；结构化其中10个停止招生专业。",
    ),
    SourceSeed(
        "sufe_2017_2018_stop_enrollment_reference",
        "2017-2018学年本科专业设置一览表(含变动情况)",
        "https://gongkai.sufe.edu.cn/91/75/c12259a102773/page.htm",
        "上海财经大学信息公开网",
        "2017-10-11",
        "official_university_html",
        "上海财经大学2017-2018学年本科专业设置一览表归档页，与2018本科专业设置页停招标记相同；仅归档原始来源，不重复结构化。",
    ),
    SourceSeed(
        "sufe_2014_2016_stop_enrollment",
        "2014、2015、2016级本科专业设置（含当年新增专业、停招专业名单）",
        "https://gongkai.sufe.edu.cn/a2/6f/c12259a41583/page.htm",
        "上海财经大学信息公开网",
        "2016-09-22",
        "official_university_html",
        "上海财经大学2014、2015、2016级本科专业设置HTML表在备注列标注停招或未招生；结构化其中5个停招/未招生专业。",
    ),
    SourceSeed(
        "bjtu_2025_stop_enrollment",
        "专业设置、当年新增专业、停招专业名单",
        "https://www.bjtu.edu.cn/xxgkw/xxgk_xxgkml/jxzl/jxzl_bksjx/bksjx_zysz/e0b0788aec404090b131b2cd77ed15fe.htm",
        "北京交通大学信息公开网",
        "2025-09-30",
        "official_university_html",
        "北京交通大学信息公开页本科专业表备注列标记6个停招专业。",
    ),
    SourceSeed(
        "bjtu_2025_major_adjustment_notice",
        "关于2025年本科专业设置与调整的公示",
        "https://bksy.bjtu.edu.cn/informations/638884534783666188.html",
        "北京交通大学本科生院",
        "2025-07-18",
        "official_university_html",
        "北京交通大学本科生院2025年本科专业设置与调整公示页公开2025年度专业设置与调整申报材料附件。",
    ),
    SourceSeed(
        "bjtu_2025_major_adjustment_materials_zip",
        "2025年度专业设置与调整申报材料",
        "https://bksy.bjtu.edu.cn/Admin/SysManage/FileUploadHandler.ashx?action=upload&url=/UploadFiles/20250718/2025%E5%B9%B4%E5%BA%A6%E4%B8%93%E4%B8%9A%E8%AE%BE%E7%BD%AE%E4%B8%8E%E8%B0%83%E6%95%B4%E7%94%B3%E6%8A%A5%E6%9D%90%E6%96%99_638884533227163948.zip&name=2025%E5%B9%B4%E5%BA%A6%E4%B8%93%E4%B8%9A%E8%AE%BE%E7%BD%AE%E4%B8%8E%E8%B0%83%E6%95%B4%E7%94%B3%E6%8A%A5%E6%9D%90%E6%96%99.zip",
        "北京交通大学本科生院",
        "2025-07-18",
        "official_university_zip",
        "北京交通大学2025年度专业设置与调整申报材料zip附件含申请撤销专业材料目录，列出材料化学、生物信息学、给排水科学与工程、电子信息工程、汉语言、思想政治教育6个撤销材料。",
    ),
    SourceSeed(
        "jnu_2024_stop_enrollment",
        "2024年专业设置一览表（含当年新增专业、停招专业）",
        "https://xxgk.jnu.edu.cn/2025/0626/c7601a839427/page.htm",
        "暨南大学信息公开网",
        "2025-06-26",
        "official_university_html",
        "暨南大学2024年专业设置一览表备注列标记20个停招专业。",
    ),
    SourceSeed(
        "jnu_2023_stop_enrollment",
        "2023年专业设置一览表（含当年新增专业、停招专业）",
        "https://xxgk.jnu.edu.cn/2023/1031/c7601a770889/page.htm",
        "暨南大学信息公开网",
        "2023-10-31",
        "official_university_html",
        "暨南大学2023年专业设置一览表备注列标记20个停招专业。",
    ),
    SourceSeed(
        "jnu_2022_stop_enrollment",
        "2022年专业设置一览表（含当年新增专业、停招专业）",
        "https://xxgk.jnu.edu.cn/2022/0623/c7601a707049/page.htm",
        "暨南大学信息公开网",
        "2022-06-23",
        "official_university_html",
        "暨南大学2022年专业设置一览表备注列标记19个停招专业。",
    ),
    SourceSeed(
        "jnu_2021_stop_enrollment",
        "2021年专业设置一览表（含当年新增专业、停招专业）",
        "https://xxgk.jnu.edu.cn/2022/0623/c7601a707047/page.htm",
        "暨南大学信息公开网",
        "2022-06-23",
        "official_university_html",
        "暨南大学2021年专业设置一览表备注列标记14个停招专业。",
    ),
    SourceSeed(
        "jnu_2020_stop_enrollment",
        "2020年专业设置一览表（含当年新增专业、停招专业）",
        "https://xxgk.jnu.edu.cn/2022/0623/c7601a707039/page.htm",
        "暨南大学信息公开网",
        "2022-06-23",
        "official_university_html",
        "暨南大学2020年专业设置一览表备注列标记12个停招专业。",
    ),
    SourceSeed(
        "jnu_2019_stop_enrollment",
        "2019年专业设置一览表（含当年新增专业、停招专业）",
        "https://xxgk.jnu.edu.cn/2022/0623/c7601a707037/page.htm",
        "暨南大学信息公开网",
        "2022-06-23",
        "official_university_html",
        "暨南大学2019年专业设置一览表备注列标记2个停招专业。",
    ),
    SourceSeed(
        "jnu_2018_stop_enrollment",
        "2018年专业设置一览表（含当年新增专业、停招专业）",
        "https://xxgk.jnu.edu.cn/2022/0623/c7601a707031/page.htm",
        "暨南大学信息公开网",
        "2022-06-23",
        "official_university_html",
        "暨南大学2018年专业设置一览表备注列标记1个停招专业。",
    ),
    SourceSeed(
        "jnu_2025_stop_enrollment",
        "2025年专业设置一览表（含当年新增专业、停招专业）",
        "https://xxgk.jnu.edu.cn/2025/0626/c7601a839436/page.htm",
        "暨南大学信息公开网",
        "2025-06-26",
        "official_university_html",
        "暨南大学2025年专业设置一览表备注列标记23个停招专业。",
    ),
    SourceSeed(
        "njtech_2025_stop_enrollment",
        "2025年本科专业设置及新增、停招情况",
        "https://jwc.njtech.edu.cn/info/1017/6053.htm",
        "南京工业大学教务处",
        "2025-05-28",
        "official_university_html",
        "南京工业大学2025年本科专业设置及新增、停招情况表备注列标记17个2025年停招专业。",
    ),
    SourceSeed(
        "nuist_2025_stop_enrollment_image",
        "2025年停招本科专业列表",
        "https://webs.nuist.edu.cn/_upload/article/images/e0/07/0ea1bc0e4375bb2059ceb3a71db8/c3eac7b4-702e-4f6d-aa12-4f3422996432.jpg",
        "南京信息工程大学信息公开网",
        "2025-11-04",
        "official_university_image",
        "南京信息工程大学2025年停招本科专业列表图片列出水利科学与工程、市场营销、公共事业管理、保险学、翻译5个停招专业及代码。",
    ),
    SourceSeed(
        "nuist_2024_stop_enrollment_image",
        "2024年停招本科专业列表",
        "https://webs.nuist.edu.cn/_upload/article/images/d1/f2/eabedd664723bbca4467c2fec2f3/c706bbfc-13f9-4f4f-83c1-d7819302a0b7.jpg",
        "南京信息工程大学信息公开网",
        "2024-11-13",
        "official_university_image",
        "南京信息工程大学2024年停招本科专业列表图片列出水利科学与工程、市场营销、公共事业管理3个停招专业及代码。",
    ),
    SourceSeed(
        "nuist_2024_major_cancellation_image",
        "2024年撤销本科专业列表",
        "https://webs.nuist.edu.cn/_upload/article/images/c0/ef/3d8bf65841cba9c85041a078a9ba/2f5e8e10-0c40-41b6-bdd0-ed2f010503cc.jpg",
        "南京信息工程大学信息公开网",
        "2024-11-13",
        "official_university_image",
        "南京信息工程大学2024年撤销本科专业列表图片列出统计学、轨道交通信号与控制2个撤销专业及代码。",
    ),
    SourceSeed(
        "njtech_2024_stop_enrollment",
        "2024年本科专业设置及新增、停招情况",
        "https://jwc.njtech.edu.cn/info/1017/5420.htm",
        "南京工业大学教务处",
        "2024",
        "official_university_html",
        "南京工业大学2024年本科专业设置及新增、停招情况表备注列标记11个停招专业。",
    ),
    SourceSeed(
        "njtech_2023_stop_enrollment",
        "2023年本科专业设置及新增、停招情况",
        "https://jwc.njtech.edu.cn/info/1017/2829.htm",
        "南京工业大学教务处",
        "2023",
        "official_university_html",
        "南京工业大学2023年本科专业设置及新增、停招情况表备注列标记5个2022年停招专业。",
    ),
    SourceSeed(
        "njtech_2022_stop_enrollment",
        "2022年本科专业设置及新增、停招情况",
        "https://jwc.njtech.edu.cn/info/1017/2830.htm",
        "南京工业大学教务处",
        "2022",
        "official_university_html",
        "南京工业大学2022年本科专业设置及新增、停招情况表备注列标记4个当年停招专业。",
    ),
    SourceSeed(
        "njtech_2021_stop_enrollment",
        "2021年本科专业设置及新增、停招情况",
        "https://jwc.njtech.edu.cn/info/1017/2831.htm",
        "南京工业大学教务处",
        "2021",
        "official_university_html",
        "南京工业大学2021年本科专业设置及新增、停招情况表备注列标记4个当年停招专业。",
    ),
    SourceSeed(
        "njtech_2020_stop_enrollment",
        "2020年本科专业设置及新增、停招情况",
        "https://jwc.njtech.edu.cn/info/1017/2832.htm",
        "南京工业大学教务处",
        "2020",
        "official_university_html",
        "南京工业大学2020年本科专业设置及新增、停招情况表备注列标记7个2019年停招专业。",
    ),
    SourceSeed(
        "njtech_2019_stop_enrollment",
        "2019年本科专业设置及新增、停招情况",
        "https://jwc.njtech.edu.cn/info/1017/2833.htm",
        "南京工业大学教务处",
        "2019",
        "official_university_html",
        "南京工业大学2019年本科专业设置及新增、停招情况表备注列标记9个当年停招专业。",
    ),
    SourceSeed(
        "wit_2025_stop_enrollment",
        "武汉工程大学本科专业目录（2025年）",
        "https://xxgk.wit.edu.cn/info/1205/8677.htm",
        "武汉工程大学信息公开网",
        "2025-11-14",
        "official_university_html",
        "武汉工程大学本科专业目录（2025年）在专业名称中标注停招或当年停招专业。",
    ),
    SourceSeed(
        "whut_2025_undergrad_major_catalog",
        "武汉理工大学2025年本科专业目录",
        "http://xxgk.whut.edu.cn/jxgl/zyqk/zyqk/202511/t20251111_620592.shtml",
        "武汉理工大学信息公开网",
        "2025-09-01",
        "official_university_image_html",
        "武汉理工大学2025年本科专业目录以两张图片表格发布，备注列标注2020、2021、2024、2025年停招专业。",
    ),
    SourceSeed(
        "zsit_2025_major_setting_status",
        "学院专业设置（含当年新增专业、停招专业）一览表（2025年）",
        "https://jwc.zsit.edu.cn/info/1324/12700.htm",
        "绍兴理工学院教务处",
        "2025-11-02",
        "official_university_html",
        "绍兴理工学院教务处2025年专业设置表逐项列出专业代码、学科门类、学位授予、设置年份和2022-2025历年招生情况，其中“停”表示该年停招。",
    ),
    SourceSeed(
        "ujs_2025_stop_enrollment",
        "专业设置、当年新增专业、停招专业名单（2025）",
        "https://xxgk.ujs.edu.cn/info/1566/7685.htm",
        "江苏大学信息公开网",
        "2025-11-20",
        "official_university_html",
        "江苏大学2025年本科专业一览表备注列标记2024/2025年停招专业和暂缓招生专业。",
    ),
    SourceSeed(
        "ujs_2024_stop_enrollment",
        "专业设置、当年新增专业、停招专业名单（2024）",
        "https://xxgk.ujs.edu.cn/info/1566/7455.htm",
        "江苏大学信息公开网",
        "2024-05-30",
        "official_university_html",
        "江苏大学2024年本科专业一览表备注列标记停招专业。",
    ),
    SourceSeed(
        "ujs_2023_stop_enrollment",
        "专业设置、当年新增专业、停招专业名单（2023）",
        "https://xxgk.ujs.edu.cn/info/1566/5645.htm",
        "江苏大学信息公开网",
        "2023-11-09",
        "official_university_html",
        "江苏大学2023年本科专业一览表备注列标记停招专业。",
    ),
    SourceSeed(
        "ujs_2022_stop_enrollment",
        "专业设置、当年新增专业、停招专业名单（2022）",
        "https://xxgk.ujs.edu.cn/info/1566/5655.htm",
        "江苏大学信息公开网",
        "2023-11-09",
        "official_university_html",
        "江苏大学2022年本科专业一览表备注列标记停招专业。",
    ),
    SourceSeed(
        "ujs_2021_new_stop_empty",
        "2021年新增专业、停招专业名单",
        "https://xxgk.ujs.edu.cn/info/1566/5665.htm",
        "江苏大学信息公开网",
        "2021-10-26",
        "official_university_html",
        "江苏大学2021年新增专业、停招专业名单页面说明2021年停招专业名单为无，本轮作为年度覆盖页归档。",
    ),
    SourceSeed(
        "ujs_2021_stop_enrollment",
        "2021年本科专业设置",
        "https://xxgk.ujs.edu.cn/info/1566/5675.htm",
        "江苏大学信息公开网",
        "2021-10-26",
        "official_university_html",
        "江苏大学2021年本科专业设置表是否停招列标记停招专业。",
    ),
    SourceSeed(
        "ujs_2020_new_stop_empty",
        "2020年新增专业、停招专业名单",
        "https://xxgk.ujs.edu.cn/info/1566/5685.htm",
        "江苏大学信息公开网",
        "2020-10-15",
        "official_university_html",
        "江苏大学2020年新增专业、停招专业名单页面说明2020年新增专业、停招专业名单为无，本轮作为年度覆盖页归档。",
    ),
    SourceSeed(
        "ujs_2020_stop_enrollment",
        "2020年本科专业设置",
        "https://xxgk.ujs.edu.cn/info/1566/5695.htm",
        "江苏大学信息公开网",
        "2020-10-15",
        "official_university_html",
        "江苏大学2020年本科专业设置表是否停招列标记停招专业。",
    ),
    SourceSeed(
        "gdufe_2025_stop_cancel",
        "2025年专业设置、当年新增专业、停招专业名单",
        "https://www.gdufe.edu.cn/2025/1112/c3473a230185/page.htm",
        "广东财经大学",
        "2025-08-31",
        "official_university_html",
        "广东财经大学2025年专业设置页列出暂停招生本科专业4个、撤销本科专业2个。",
    ),
    SourceSeed(
        "nefu_2025_stop_enrollment",
        "2025年专业设置，当年新增专业、停招专业名单",
        "https://xxgk.nefu.edu.cn/info/1134/5881.htm",
        "东北林业大学信息公开网",
        "2025-10-27",
        "official_university_html",
        "东北林业大学2025年信息公开页正文列出停招专业7个。",
    ),
    SourceSeed(
        "whpu_2025_stop_enrollment",
        "2025年学校停招专业",
        "https://xxgkw.whpu.edu.cn/info/1184/3456.htm",
        "武汉轻工大学信息公开网",
        "2025-06-25",
        "official_university_html",
        "武汉轻工大学2025年学校停招专业页列出停招专业1个：给排水科学与工程。",
    ),
    SourceSeed(
        "hbue_2024_teaching_quality_report_pdf",
        "湖北经济学院2023-2024学年本科教学质量报告",
        "https://xxgk.hbue.edu.cn/_upload/article/files/41/5c/7a85fa1b4fc5ac2785ac9545faac/9f91e64f-484b-4cb0-bc62-1f0162a5bd89.pdf",
        "湖北经济学院信息公开网",
        "2024",
        "official_university_pdf",
        "湖北经济学院2023-2024学年本科教学质量报告列出15个暂停招生专业，其中3个为本年度停招，并说明申请撤销信用管理、物流工程、资产评估、社会工作、艺术设计学、广告学、应用统计学7个专业。",
    ),
    SourceSeed(
        "hgnu_2025_stop_enrollment_notice",
        "黄冈师范学院停招专业（2025）",
        "https://xxgk.hgnu.edu.cn/?Article57/Article50/Article91/Article167/677.html",
        "黄冈师范学院信息公开网",
        "2025-06-18",
        "official_university_html",
        "黄冈师范学院信息公开详情页链接2025年停招/撤销专业名单XLS附件。",
    ),
    SourceSeed(
        "hgnu_2025_stop_enrollment_xls",
        "黄冈师范学院停招专业（2025）XLS",
        "https://xxgk.hgnu.edu.cn/up/file/20250618/20250618170616801680.xls",
        "黄冈师范学院信息公开网",
        "2025-06-18",
        "official_university_xls",
        "黄冈师范学院2025年停招专业XLS列出6个停招专业和3个撤销专业，并提供所属院系、专业名称、专业代码、授予学位、调整类型及备注。",
    ),
    SourceSeed(
        "hgnu_2022_stop_enrollment",
        "黄冈师范学院停招专业（2022）",
        "https://xxgk.hgnu.edu.cn/?Article57/Article50/Article91/Article167/582.html",
        "黄冈师范学院信息公开网",
        "2022-07-11",
        "official_university_html",
        "黄冈师范学院2022年停招专业HTML表列出6个停招专业，并提供教学学院、专业名称、专业代码、授予学位、开始招生时间和批准文号。",
    ),
    SourceSeed(
        "huat_2025_stop_enrollment_xlsx",
        "2025年新增专业、停招专业名单",
        "https://xxgk.huat.edu.cn/system/_content/download.jsp?urltype=news.DownloadAttachUrl&owner=1905084890&wbfileid=15596502",
        "湖北汽车工业学院信息公开网",
        "2025-06-10",
        "official_university_xlsx",
        "湖北汽车工业学院2025年新增专业、停招专业名单附件列出2025年停招专业2个。",
    ),
    SourceSeed(
        "eurasia_2025_stop_enrollment",
        "2025年专业设置、当年新增专业、停招专业名单",
        "https://xxgk.eurasia.edu/info/1621/4639.htm",
        "西安欧亚学院信息公开网",
        "2025-09-16",
        "official_university_html",
        "西安欧亚学院2025年专业设置页明确软件工程（合作办学）停招。",
    ),
    SourceSeed(
        "csu_2023_stop_enrollment",
        "2023年停招专业名单",
        "https://xxgk.csu.edu.cn/info/1102/4986.htm",
        "中南大学信息公开网",
        "2023-10-26",
        "official_university_html",
        "中南大学2023年停招专业名单列出视觉传达设计、环境设计2个本科专业停止招生。",
    ),
    SourceSeed(
        "csu_2024_stop_enrollment",
        "2024年停招专业名单",
        "https://xxgk.csu.edu.cn/info/1102/5144.htm",
        "中南大学信息公开网",
        "2024-09-30",
        "official_university_html",
        "中南大学2024年停招专业名单列出产品设计、生物技术2个本科专业停止招生。",
    ),
    SourceSeed(
        "csu_2025_undergrad_setting",
        "2025年本科专业设置表",
        "https://xxgk.csu.edu.cn/info/1102/5315.htm",
        "中南大学信息公开网",
        "2025-09-30",
        "official_university_html",
        "中南大学2025年本科专业设置表信息公开页嵌入PDF，PDF停止招生列标记15个专业。",
    ),
    SourceSeed(
        "csu_2025_undergrad_setting_pdf",
        "中南大学2025年本科专业设置表PDF",
        "https://xxgk.csu.edu.cn/__local/A/B0/43/053D4EF2F3346F7CF6166DAD9EA_9A5CC778_3156F.pdf",
        "中南大学信息公开网",
        "2025-09-30",
        "official_university_pdf",
        "中南大学2025年本科专业设置表PDF逐项列出学院、专业代码、专业名称、授予学位、修业年限和停止招生列，其中15个专业停止招生列为“是”。",
    ),
    SourceSeed(
        "imnc_2025_teaching_quality_report_pdf",
        "呼和浩特民族学院2024-2025学年本科教学质量报告",
        "https://www.imnc.edu.cn/__local/5/2D/1B/B961C3AEC040EBD2B65C9BBC3C1_1B456114_D9463.pdf",
        "呼和浩特民族学院",
        "2025",
        "official_university_pdf",
        "呼和浩特民族学院2024-2025学年本科教学质量报告在本科专业设置情况中说明2025年停招信息管理与信息系统、视觉传达设计、行政管理、市场营销4个专业，并申请撤销这4个本科专业。",
    ),
    SourceSeed(
        "tjfsu_bhws_2025_teaching_quality_report",
        "天津外国语大学滨海外事学院2024-2025学年本科教学质量报告",
        "https://bhws.tjfsu.edu.cn/xxgk/p.jsp?id=115&pagenum=1&type=1",
        "天津外国语大学滨海外事学院信息公开网",
        "2025-10-27",
        "official_university_html",
        "天津外国语大学滨海外事学院2024-2025学年本科教学质量报告说明学院设有22个本科专业，并列出本学年停招国际事务与国际关系、新闻学、电子商务、财务管理、广告学5个专业。",
    ),
    SourceSeed(
        "sxufe_2025_stop_enrollment",
        "山西财经大学2025年停招本科专业名单",
        "https://jwc.sxufe.edu.cn/info/1880/8694.htm",
        "山西财经大学教务部",
        "2025-11-03",
        "official_university_html",
        "山西财经大学2025年停招本科专业名单表格列出汉语国际教育、软件工程、房地产开发与管理3个本科专业，并给出专业代码、专业类、修业年限、学位门类和开设年份。",
    ),
    SourceSeed(
        "hbmzu_2025_stop_enrollment",
        "停招专业（2025年）",
        "https://www.hbmzu.edu.cn/xxgkw/info/1139/5276.htm",
        "湖北民族大学信息公开网",
        "2025-05-29",
        "official_university_html",
        "湖北民族大学信息公开网2025年停招专业页面列出医学影像技术、广播电视编导2个本科专业。",
    ),
    SourceSeed(
        "hbmzu_2024_stop_enrollment",
        "停招专业（2024年）",
        "https://www.hbmzu.edu.cn/xxgkw/info/1139/4869.htm",
        "湖北民族大学信息公开网",
        "2024-05-15",
        "official_university_html",
        "湖北民族大学信息公开网2024年停招专业页面列出编辑出版学1个本科专业。",
    ),
    SourceSeed(
        "hbmzu_2023_stop_enrollment_none",
        "停招专业（2023年）",
        "https://www.hbmzu.edu.cn/xxgkw/info/1139/4868.htm",
        "湖北民族大学信息公开网",
        "2023-05-15",
        "official_university_html",
        "湖北民族大学信息公开网2023年停招专业页面说明2023年无停止招生专业。",
    ),
    SourceSeed(
        "hbmzu_2022_stop_enrollment",
        "停招专业（2022年）",
        "https://www.hbmzu.edu.cn/xxgkw/info/1139/4194.htm",
        "湖北民族大学信息公开网",
        "2022-10-27",
        "official_university_html",
        "湖北民族大学信息公开网2022年停招专业页面列出应用化学、翻译、人文地理与城乡规划3个本科专业。",
    ),
    SourceSeed(
        "hbmzu_2021_stop_enrollment_none",
        "停招专业名单（2021年）",
        "https://www.hbmzu.edu.cn/xxgkw/info/1139/3268.htm",
        "湖北民族大学信息公开网",
        "2021-10-12",
        "official_university_html",
        "湖北民族大学信息公开网2021年停招专业名单页面说明2021年无停止招生专业。",
    ),
    SourceSeed(
        "hbxytc_2025_vocational_stop_pdf",
        "襄阳职业技术学院2025年停招专业PDF",
        "https://xxgk.hbxytc.cn/__local/9/83/9E/6E12483AA01DAA1A79773AB4325_4770B32A_9227.pdf",
        "襄阳职业技术学院信息公开网",
        "2025",
        "official_college_pdf",
        "襄阳职业技术学院2025年停招专业PDF表格列出供应链运营、助产、计算机应用技术、物联网应用技术、酒店管理与数字化运营、大数据与会计、电子商务、服装与服饰设计8个高职专业。",
    ),
    SourceSeed(
        "hbxytc_2024_vocational_stop_pdf",
        "襄阳职业技术学院2024年停招专业名单PDF",
        "https://xxgk.hbxytc.cn/__local/C/A3/54/3EBAFC013B14E7007A73609ED80_256ACF5F_8092.pdf",
        "襄阳职业技术学院信息公开网",
        "2024",
        "official_college_pdf",
        "襄阳职业技术学院2024年停招专业名单PDF表格列出云计算技术应用、大数据与财务管理、数控技术、装配式建筑构件智能制造技术、助产、现代物流管理6个高职专业。",
    ),
    SourceSeed(
        "hbxytc_2023_vocational_stop_pdf",
        "襄阳职业技术学院2023年停招专业PDF",
        "https://xxgk.hbxytc.cn/__local/6/C8/0B/A9B41AE7790660FF5A060FE4628_61668F80_6EC2.pdf",
        "襄阳职业技术学院信息公开网",
        "2023",
        "official_college_pdf",
        "襄阳职业技术学院2023年停招专业PDF表格列出药学、数控技术2个高职专业。",
    ),
    SourceSeed(
        "hbxytc_2022_vocational_stop_pdf",
        "襄阳职业技术学院2022年停招专业PDF",
        "https://xxgk.hbxytc.cn/__local/A/39/35/B59D72F3FE89349CB77E48838D8_C7873F7A_545C.pdf",
        "襄阳职业技术学院信息公开网",
        "2022",
        "official_college_pdf",
        "襄阳职业技术学院2022年停招专业PDF表格列出动物医学、装配式建筑构件智能制造技术、轨道交通工程机械制造与维护、云计算技术应用、助产、大数据与财务管理6个高职专业。",
    ),
    SourceSeed(
        "hbxytc_2021_vocational_stop_pdf",
        "襄阳职业技术学院2021年停招专业PDF",
        "https://xxgk.hbxytc.cn/__local/7/D1/80/8F47A810123449A57EB3D783E4E_9143491A_20FA1.pdf",
        "襄阳职业技术学院信息公开网",
        "2021",
        "official_college_pdf",
        "襄阳职业技术学院2021年停招专业PDF表格列出20个高职专业，含口腔医学技术、眼视光技术、应用电子技术、计算机网络技术等。",
    ),
    SourceSeed(
        "hbxytc_2020_vocational_stop_pdf",
        "襄阳职业技术学院2020年停招专业PDF",
        "https://xxgk.hbxytc.cn/__local/A/EB/51/66A95030B94AC7AB7D18E4E30CB_BCC6661C_26245.pdf",
        "襄阳职业技术学院信息公开网",
        "2020",
        "official_college_pdf",
        "襄阳职业技术学院2020年停招专业PDF为Print2Flash生成文件，本轮归档原始PDF；当前文本抽取仅得到水印，未结构化专业行。",
    ),
    SourceSeed(
        "hzec_2025_major_setting_page",
        "2024-2025学年专业设置、当年新增专业、停招专业名单",
        "https://www.hzec.edu.cn/xxgk/jxzlxx/2024-11-18/22092.html",
        "惠州经济职业技术学院高校信息公开专栏",
        "2025-08-31",
        "official_college_html",
        "惠州经济职业技术学院2024-2025学年专业设置、当年新增专业、停招专业名单信息公开页嵌入PDF附件。",
    ),
    SourceSeed(
        "hzec_2025_major_setting_pdf",
        "惠州经济职业技术学院2024-2025学年专业设置、当年新增专业、停招专业名单PDF",
        "https://www.hzec.edu.cn/d/file/xxgk/jxzlxx/2025-10-30/b88e6fbd64c9665ed244df63bc4c0731.pdf",
        "惠州经济职业技术学院高校信息公开专栏",
        "2025-08-31",
        "official_college_pdf",
        "惠州经济职业技术学院2024-2025学年专业设置PDF表格列出2024年停招9个高职专业、2025年停招5个高职专业，并给出学院、专业代码、学制和新招/停招状态。",
    ),
    SourceSeed(
        "gdjmcmc_2025_major_setting_pdf",
        "2024-2025学年专业设置、当年新增专业、停招专业名单",
        "https://www.gdjmcmc.edu.cn/xxgk/jxzlxx/202511/P020251105370134462118.pdf",
        "广东江门中医药职业学院",
        "2025-11-05",
        "official_college_pdf",
        "广东江门中医药职业学院2024-2025学年专业设置、当年新增专业、停招专业名单PDF列出32个专业，其中智能医疗装备技术、食品智能加工技术、药品经营与管理、智慧健康养老服务与管理为2023年停招、2024年招生，老年保健与管理、生殖健康管理、呼吸治疗技术为2024年停招。",
    ),
    SourceSeed(
        "xzcit_2025_stop_page",
        "徐州工业职业技术学院2025年招生、停招、新增专业",
        "https://xxgk.xzcit.cn/2025/1021/c1931a63737/page.htm",
        "徐州工业职业技术学院信息公开网",
        "2025-10-22",
        "official_college_html",
        "徐州工业职业技术学院2025年招生、停招、新增专业信息公开页嵌入PDF附件。",
    ),
    SourceSeed(
        "xzcit_2025_stop_pdf",
        "徐州工业职业技术学院2025年招生、停招、新增专业PDF",
        "https://xxgk.xzcit.cn/_upload/article/files/8f/c0/a56c283c4a7ea7e37c50dca5ee90/533c2dd2-a968-4f95-bd26-f53321619635.pdf",
        "徐州工业职业技术学院信息公开网",
        "2025-10-22",
        "official_college_pdf",
        "徐州工业职业技术学院2025年招生、停招、新增专业PDF表格列出药品生物技术、食品检验检测技术、建设工程管理、大数据技术4个停招高职专业，并给出学院和专业代码。",
    ),
    SourceSeed(
        "xzcit_2024_stop_page",
        "徐州工业职业技术学院2024年停招专业",
        "https://xxgk.xzcit.cn/2024/1115/c1931a56015/page.htm",
        "徐州工业职业技术学院信息公开网",
        "2024-11-15",
        "official_college_html",
        "徐州工业职业技术学院2024年停招专业信息公开页嵌入PDF附件。",
    ),
    SourceSeed(
        "xzcit_2024_stop_pdf",
        "徐州工业职业技术学院2024年停招专业PDF",
        "https://xxgk.xzcit.cn/_upload/article/files/61/fe/28fdc3bc43fc9203e61cf4919383/54c8bae9-f365-4c84-a7c5-ff50936ea69b.pdf",
        "徐州工业职业技术学院信息公开网",
        "2024-11-15",
        "official_college_pdf",
        "徐州工业职业技术学院2024年停招专业PDF表格列出旅游管理、酒店管理与数字化运营、商务英语、城市轨道交通机电技术4个停招高职专业，并给出专业代码。",
    ),
    SourceSeed(
        "xzcit_2022_stop_html",
        "2022年停招专业",
        "https://xxgk.xzcit.cn/2023/1115/c1931a53162/page.htm",
        "徐州工业职业技术学院信息公开网",
        "2023-11-15",
        "official_college_html",
        "徐州工业职业技术学院2022年停招专业HTML表格列出皮具制作与工艺、煤化工技术、云计算技术应用3个停招高职专业，并给出专业代码。",
    ),
    SourceSeed(
        "xzcit_2021_stop_page",
        "2021停招专业",
        "https://xxgk.xzcit.cn/2022/1128/c1931a49492/page.htm",
        "徐州工业职业技术学院信息公开网",
        "2022-11-28",
        "official_college_html",
        "徐州工业职业技术学院2021停招专业信息公开页嵌入PDF附件。",
    ),
    SourceSeed(
        "xzcit_2021_stop_pdf",
        "徐州工业职业技术学院2021停招专业PDF",
        "https://xxgk.xzcit.cn/_upload/article/files/16/f6/eb4f86754cbdb4f7584b6f75132a/129a1ac3-e09e-4855-8c91-00c8b1434ece.pdf",
        "徐州工业职业技术学院信息公开网",
        "2022-11-28",
        "official_college_pdf",
        "徐州工业职业技术学院2021停招专业PDF表格列出光伏发电技术与应用、数字媒体艺术设计2个停招高职专业，并给出专业代码。",
    ),
    SourceSeed(
        "wust_2026_stop_enrollment_empty",
        "2026年停招本科专业",
        "https://xxgk.wust.edu.cn/info/2351/70202.htm",
        "武汉科技大学信息公开网",
        "2026-04-28",
        "official_university_html",
        "武汉科技大学信息公开网2026年停招本科专业页面公开存在，但抓取正文未列出停招专业明细，本轮归档为年度覆盖页。",
    ),
    SourceSeed(
        "wust_2025_stop_enrollment_empty",
        "2025年停招本科专业",
        "https://xxgk.wust.edu.cn/info/2351/69592.htm",
        "武汉科技大学信息公开网",
        "2025-10-15",
        "official_university_html",
        "武汉科技大学信息公开网2025年停招本科专业页面公开存在，但抓取正文未列出停招专业明细，本轮归档为年度覆盖页。",
    ),
    SourceSeed(
        "wust_2024_stop_enrollment_empty",
        "2024年停招专业",
        "https://xxgk.wust.edu.cn/info/2351/62542.htm",
        "武汉科技大学信息公开网",
        "2024-06-25",
        "official_university_html",
        "武汉科技大学信息公开网2024年停招专业页面公开存在，但抓取正文未列出停招专业明细，本轮归档为年度覆盖页。",
    ),
    SourceSeed(
        "wust_2024_teaching_quality_report_pdf",
        "武汉科技大学2023-2024学年本科教学质量报告",
        "https://xxgk.wust.edu.cn/__local/F/5F/FE/AF6758057A92FD641149E1FBA31_D7D352E0_12868E.pdf",
        "武汉科技大学信息公开网",
        "2024",
        "official_university_pdf",
        "武汉科技大学2023-2024学年本科教学质量报告列出2024年招生专业72个、新增5个、停招绘画等4个专业，并说明申请撤销人文地理与城乡规划和交通运输专业。",
    ),
    SourceSeed(
        "cuit_2024_teaching_quality_report_pdf",
        "成都信息工程大学2023-2024学年本科教学质量报告",
        "https://jwc.cuit.edu.cn/system/_content/download.jsp?urltype=news.DownloadAttachUrl&owner=1581786981&wbfileid=D480CA679E07860A7380FDCD3DD05C6C",
        "成都信息工程大学教务处",
        "2024-12-11",
        "official_university_pdf",
        "成都信息工程大学2023-2024学年本科教学质量报告列出本科专业总数60个、招生专业51个，实施专业红橙黄牌管理，并说明撤销材料物理和信息对抗技术2个专业。",
    ),
    SourceSeed(
        "wust_2023_stop_enrollment",
        "2023年停招专业",
        "https://xxgk.wust.edu.cn/info/2351/46831.htm",
        "武汉科技大学信息公开网",
        "2023-06-21",
        "official_university_html",
        "武汉科技大学信息公开网2023年停招专业HTML表格列出电子商务、物业管理、汽车服务工程、人文地理与城乡规划、交通运输、马克思主义理论、劳动与社会保障7个本科停招专业，并给出学院、专业代码、修业年限和学科。",
    ),
    SourceSeed(
        "ldxy_2024_teaching_quality_report_pdf",
        "陇东学院2023-2024学年本科教学质量报告",
        "https://www.ldxy.edu.cn/_upload/article/files/c0/04/fdad2c5c465e9b44b183458d85e1/172adfb3-3838-4484-a387-a5c86804ef6f.pdf",
        "陇东学院",
        "2024",
        "official_university_pdf",
        "陇东学院2023-2024学年本科教学质量报告列出2022年15个未招生专业、2023年24个未招生专业，并说明2023-2024学年主动申请撤销8个停招五年以上且无在校生的本科专业。",
    ),
    SourceSeed(
        "dlut_2024_teaching_quality_report_pdf",
        "大连理工大学2023-2024学年本科教学质量报告",
        "https://info.dlut.edu.cn/__local/F/F7/B8/84529792FB9CC0B998FE4DA3DE8_645ED224_7EB43.pdf",
        "大连理工大学信息公开网",
        "2024",
        "official_university_pdf",
        "大连理工大学2023-2024学年本科教学质量报告脚注列明历年停止招生专业、管理科学暂未招生，以及2024年申请撤销物流工程等6个本科专业。",
    ),
    SourceSeed(
        "sie_2024_teaching_quality_report_pdf",
        "沈阳工程学院2023-2024学年本科教学质量报告",
        "https://www.sie.edu.cn/__local/6/EA/55/C2B3B8751355FF39A5FAAE56ADC_47E9B74D_9CA59.pdf",
        "沈阳工程学院",
        "2024",
        "official_university_pdf",
        "沈阳工程学院2023-2024学年本科教学质量报告专业设置表标注机械电子工程、机械工艺技术、商务英语停招，并说明2023年撤销测控技术与仪器、2024年申请撤销机械电子工程和商务英语。",
    ),
    SourceSeed(
        "zjhzu_2024_teaching_quality_report_pdf",
        "湖州学院2023-2024学年本科教学质量报告",
        "https://jwc.zjhzu.edu.cn/_upload/article/files/79/10/f7e56d2a44589b0fe28b17617f56/d0b9e5ee-2c8d-491b-a213-0c0193d449be.pdf",
        "湖州学院教务处",
        "2024",
        "official_university_pdf",
        "湖州学院2023-2024学年本科教学质量报告列出36个本科专业、27个招生专业、9个停招专业，并说明2023年撤销历史学和美术学、2024年申请撤销物联网工程。",
    ),
    SourceSeed(
        "synu_2024_teaching_quality_report_pdf",
        "沈阳师范大学2023-2024学年本科教学质量报告",
        "https://www.synu.edu.cn/_upload/article/files/7e/2c/bc0dd65548ab818a636f9e54c06e/541c9bea-158c-437a-9a1e-3f53cdb6a58b.pdf",
        "沈阳师范大学",
        "2024",
        "official_university_pdf",
        "沈阳师范大学2023-2024学年本科教学质量报告披露学校现有68个本科专业、招生专业56个、停招专业12个，并称当年撤销9个专业。",
    ),
    SourceSeed(
        "nnutc_2024_teaching_quality_report_pdf",
        "南京师范大学泰州学院2023-2024学年本科教学质量报告",
        "https://news.nnutc.edu.cn/nanjingshifandaxuetaizhouxueyuan2023-2024xuenianbenkejiaoxuezhiliangbaogao.pdf",
        "南京师范大学泰州学院",
        "2024",
        "official_university_pdf",
        "南京师范大学泰州学院2023-2024学年本科教学质量报告说明申请撤销广告学、戏剧影视文学、园艺等3个长期未招生专业。",
    ),
    SourceSeed(
        "sccm_2024_teaching_quality_report_pdf",
        "四川音乐学院2023-2024学年本科教学质量报告",
        "https://www.sccm.edu.cn/upload/202412/02/202412021531358992.pdf",
        "四川音乐学院",
        "2024",
        "official_university_pdf",
        "四川音乐学院2023-2024学年本科教学质量报告说明2023年和2024年申请撤销工业设计、公共事业管理，并一并撤销戏剧学、戏剧影视导演、服装与服饰设计等未招生专业；支撑数据列出新媒体艺术为当年停招专业。",
    ),
    SourceSeed(
        "hunau_2024_teaching_quality_report_pdf",
        "湖南农业大学2023-2024学年本科教学质量报告",
        "https://jwc.hunau.edu.cn/shpg_2023/zlbg/202412/P020241210649033929365.pdf",
        "湖南农业大学教务处",
        "2024",
        "official_university_pdf",
        "湖南农业大学2023-2024学年本科教学质量报告说明近5年撤销植物科学与技术、表演2个专业，停招社会工作、信息工程、汽车服务工程、机械电子工程、水族科学与技术5个专业，并于2024年申请撤销信息工程、汽车服务工程、社会工作3个本科专业。",
    ),
    SourceSeed(
        "zust_2024_teaching_quality_report_pdf",
        "浙江科技大学2023-2024学年本科教学质量报告",
        "https://xxgk.zust.edu.cn/20251017160849780e0a.pdf",
        "浙江科技大学信息公开网",
        "2024",
        "official_university_pdf",
        "浙江科技大学2023-2024学年本科教学质量报告脚注列出2023-2024学年停招测控技术与仪器、物联网工程、包装工程、电子商务、汽车服务工程5个专业，并说明拟撤销2个专业。",
    ),
    SourceSeed(
        "chnu_2024_teaching_quality_report_pdf",
        "淮北师范大学2023-2024学年本科教学质量报告",
        "https://xxgk.chnu.edu.cn/upload/xxgk/contentmanage/article/file/2024/12/04/%E6%B7%AE%E5%8C%97%E5%B8%88%E8%8C%83%E5%A4%A7%E5%AD%A62023-2024%E5%AD%A6%E5%B9%B4%E6%9C%AC%E7%A7%91%E6%95%99%E5%AD%A6%E8%B4%A8%E9%87%8F%E6%8A%A5%E5%91%8A.pdf",
        "淮北师范大学信息公开网",
        "2024",
        "official_university_pdf",
        "淮北师范大学2023-2024学年本科教学质量报告说明完成7个专业撤销工作、申请撤销信息管理与信息系统，并停招广告学、戏剧影视文学、社会学、国际经济与贸易、审计学5个专业。",
    ),
    SourceSeed(
        "wzut_2024_teaching_quality_report_pdf",
        "温州理工学院2023-2024学年本科教学质量报告",
        "https://www.wzut.edu.cn/resources/202412/%E6%B8%A9%E5%B7%9E%E7%90%86%E5%B7%A5%E5%AD%A6%E9%99%A22023-2024%E5%AD%A6%E5%B9%B4%E6%9C%AC%E7%A7%91%E6%95%99%E5%AD%A6%E8%B4%A8%E9%87%8F%E6%8A%A5%E5%91%8A20241202031502540.pdf",
        "温州理工学院",
        "2024",
        "official_university_pdf",
        "温州理工学院2023-2024学年本科教学质量报告列出34个本科专业，其中停招专业6个，并在专业设置表逐项列出市场营销、音乐表演、人力资源管理、电子商务、车辆工程、广告学的停招年份。",
    ),
    SourceSeed(
        "czu_2024_teaching_quality_report_page",
        "池州学院2023-2024学年本科教学质量报告",
        "https://fz.czu.edu.cn/info/1030/1870.htm",
        "池州学院发展规划处",
        "2024-12-02",
        "official_university_html",
        "池州学院官网正文页公开2023-2024学年本科教学质量报告全文，列出校内停招专业材料化学、市场营销（专升本）、知识产权（安警院联培），支撑数据列出本科专业总数50、在招49、当年停招材料化学。",
    ),
    SourceSeed(
        "bwu_2024_teaching_quality_report_pdf",
        "北京物资学院2023-2024学年本科教学质量报告",
        "https://xxgk.bwu.edu.cn/system/_content/download.jsp?urltype=news.DownloadAttachUrl&owner=1172872658&wbfileid=3108850",
        "北京物资学院信息公开网",
        "2024-12-05",
        "official_university_pdf",
        "北京物资学院2023-2024学年本科教学质量报告称2023年撤销经济统计学、英语、劳动关系3个专业，并在支撑数据中列出本科专业总数26、在招专业数24。",
    ),
    SourceSeed(
        "xcu_2024_teaching_quality_report_pdf",
        "许昌学院2023-2024学年本科教学质量报告",
        "https://www.xcu.edu.cn/system/resource/storage/download.jsp?mark=NUFBRkQxNUM0RkRDQTJEMzEzODY5MkMzMzUzQzFCMTYvOEIzMDdDQTAvOTc4QzM%3D",
        "许昌学院",
        "2024",
        "official_university_pdf",
        "许昌学院2023-2024学年本科教学质量报告说明本学年停招通信工程、网络工程、交通设备与控制工程、酒店管理、财务管理、绘画、汉语国际教育、商务英语、工商管理等9个本科专业，并申请撤销城乡规划、市场营销、音乐表演、社会体育指导与管理4个专业。",
    ),
    SourceSeed(
        "jlenu_2024_teaching_quality_report_pdf",
        "吉林工程技术师范学院2023-2024学年本科教学质量报告",
        "https://xxgk.jlenu.edu.cn/__local/5/A4/48/709198BE7B5448A668655A635D8_7C5DE637_171A87.pdf",
        "吉林工程技术师范学院信息公开网",
        "2024-11",
        "official_university_pdf",
        "吉林工程技术师范学院2023-2024学年本科教学质量报告列出本科专业54个、有在校生专业53个、在招专业41个，并在本科专业设置表注释中说明44-53项为停招专业。",
    ),
    SourceSeed(
        "jlenu_2023_major_cancellation",
        "2023年拟撤销专业公示",
        "https://www.jlenu.edu.cn/info/1112/7593.htm",
        "吉林工程技术师范学院",
        "2023-08-12",
        "official_university_html",
        "吉林工程技术师范学院2023年拟撤销行政管理本科专业。",
    ),
    SourceSeed(
        "jlenu_2025_teaching_quality_report_pdf",
        "吉林工程技术师范学院2024-2025学年本科教学质量报告",
        "https://xxgk.jlenu.edu.cn/__local/0/76/FE/D9C45819DEF23882A78CBC6703D_FA6291F5_191101.pdf",
        "吉林工程技术师范学院信息公开网",
        "2025-12",
        "official_university_pdf",
        "吉林工程技术师范学院2024-2025学年本科教学质量报告称现有42个本科专业招生，2025年撤销汽车服务工程、人工智能、环境设计、表演4个专业，新增新能源汽车工程、新能源汽车工程技术、网络空间安全3个专业。",
    ),
    SourceSeed(
        "wsyu_2025_stop_enrollment",
        "5-28-63 武昌首义学院2025年停招专业名单",
        "https://www.wsyu.edu.cn/xxgkw/2025/0630/c944a95442/pagem.htm",
        "武昌首义学院信息公开网",
        "2025-06-30",
        "official_university_html",
        "武昌首义学院2025年停招专业名单HTML表列出11个本科专业条目，含校内专业代码、校内专业名称、标准专业名称、专业代码、所属单位、设置年份、学制、学位门类和招生状态。",
    ),
    SourceSeed(
        "wsyu_2024_stop_enrollment",
        "5-28-63 武昌首义学院2024年停招专业名单",
        "https://www.wsyu.edu.cn/xxgkw/2024/0530/c944a95443/pagem.htm",
        "武昌首义学院信息公开网",
        "2024-05-30",
        "official_university_html",
        "武昌首义学院2024年停招专业名单HTML表列出11个本科专业条目，含校内专业代码、校内专业名称、标准专业名称、专业代码、所属单位、设置年份、学制、学位门类和招生状态。",
    ),
    SourceSeed(
        "wsyu_2023_stop_enrollment",
        "5-28-63 武昌首义学院2023年停招专业名单",
        "https://www.wsyu.edu.cn/xxgkw/2023/0324/c944a95423/pagem.htm",
        "武昌首义学院信息公开网",
        "2023-03-24",
        "official_university_html",
        "武昌首义学院2023年停招专业名单HTML表列出11个已停招本科专业条目，含标准专业名称、专业代码、校内专业名称、所属单位、设置年份、学制、学位门类和招生状态。",
    ),
    SourceSeed(
        "wsyu_2022_stop_enrollment_empty",
        "5-28-63 武昌首义学院2022年停招专业名单",
        "https://www.wsyu.edu.cn/xxgkw/2022/0909/c944a95410/pagem.htm",
        "武昌首义学院信息公开网",
        "2022-09-09",
        "official_university_html",
        "武昌首义学院2022年停招专业名单页面说明2021-2022学年学校无停招专业，本轮作为年度覆盖页归档。",
    ),
    SourceSeed(
        "wsyu_2021_stop_enrollment",
        "5-28-63 武昌首义学院2021年停招专业名单",
        "https://www.wsyu.edu.cn/xxgkw/2021/0618/c944a95387/pagem.htm",
        "武昌首义学院信息公开网",
        "2021-06-18",
        "official_university_html",
        "武昌首义学院2021年停招专业名单HTML表列出11个本科专业条目，含校内专业代码、校内专业名称、标准专业名称、专业代码、所属单位、设置年份、学制、学位门类、招生状态和是否新专业。",
    ),
    SourceSeed(
        "wsyu_2020_stop_enrollment",
        "5-28-63 武昌首义学院2020年停招专业名单",
        "https://www.wsyu.edu.cn/xxgkw/2020/0608/c944a95388/pagem.htm",
        "武昌首义学院信息公开网",
        "2020-06-08",
        "official_university_html",
        "武昌首义学院2020年停招专业名单HTML表列出5个本科专业条目，含校内专业代码、校内专业名称、标准专业名称、专业代码、所属单位、设置年份、学制、学位门类、招生状态和是否新专业。",
    ),
    SourceSeed(
        "wsyu_2019_stop_enrollment",
        "5-28-63 武昌首义学院2019年停招专业名单",
        "https://www.wsyu.edu.cn/xxgkw/2019/0602/c944a95389/pagem.htm",
        "武昌首义学院信息公开网",
        "2019-06-02",
        "official_university_html",
        "武昌首义学院2019年停招专业名单HTML表列出7个本科专业条目，含校内专业代码、校内专业名称、标准专业名称、专业代码、所属单位、设置年份、学制、学位门类、招生状态和是否新专业。",
    ),
    SourceSeed(
        "wsyu_2018_stop_enrollment_empty",
        "5-28-63 武昌首义学院停招专业名单",
        "https://www.wsyu.edu.cn/xxgkw/2018/0629/c944a95422/pagem.htm",
        "武昌首义学院信息公开网",
        "2018-06-29",
        "official_university_html",
        "武昌首义学院2018年停招专业名单页面说明学校目前没有停招专业，本轮作为年度覆盖页归档。",
    ),
    SourceSeed(
        "jlu_2026_undergrad_stop_enrollment",
        "本科专业141个",
        "https://www.jlu.edu.cn/info/1078/1772.htm",
        "吉林大学",
        "2026-01-12",
        "official_university_html",
        "吉林大学统计资料页称本科专业141个，其中121个招生专业、19个专业已停招但仍有在校学生；表格备注列给出停招年份。",
    ),
    SourceSeed(
        "hnu_2025_major_cancellation",
        "关于湖南大学2025年度本科专业设置的公示",
        "https://jwc.hnu.edu.cn/info/1021/12176.htm",
        "湖南大学教务处",
        "2025-07-04",
        "official_university_html",
        "湖南大学2025年度本科专业设置公示列出3个撤销专业。",
    ),
    SourceSeed(
        "hnu_2024_major_catalog_pdf",
        "湖南大学2024年本科专业目录",
        "https://xxgk.hnu.edu.cn/__local/F/FF/52/20556644D1E8ED116ACB204BD56_DF21C489_27346.pdf",
        "湖南大学信息公开网",
        "2024",
        "official_university_pdf",
        "湖南大学2024年本科专业目录列出84个本科专业、76个招生专业、4个新增专业和8个当年停招专业，并在备注列逐项标注当年停招。",
    ),
    SourceSeed(
        "hnu_2025_major_catalog_pdf",
        "湖南大学2025年本科专业目录",
        "https://xxgk.hnu.edu.cn/__local/4/DC/92/E5B975334B616B52FE8B6C74CB3_49D6C73A_199B3.pdf",
        "湖南大学信息公开网",
        "2025",
        "official_university_pdf",
        "湖南大学2025年本科专业目录列出90个本科专业、77个招生专业、3个新增专业和13个当年停招专业，并在备注列逐项标注当年停招。",
    ),
    SourceSeed(
        "ayit_2025_stop_enrollment",
        "安阳工学院2025年停招本科专业一览表",
        "https://xxgk.ayit.edu.cn/info/1201/2555.htm",
        "安阳工学院信息公开网",
        "2025-11-09",
        "official_university_html",
        "安阳工学院2025年停招本科专业一览表列出信息管理与信息系统、城乡规划、材料成型及控制工程等9个停招本科专业及专业代码。",
    ),
    SourceSeed(
        "nwpu_2024_major_setting_pdf",
        "西北工业大学本科专业设置（2024年10月更新）",
        "https://xxgk.nwpu.edu.cn/__local/2/F8/A4/E1666E2B18C4F7CFE99E02A9ACB_9826B9A2_17578.pdf",
        "西北工业大学信息公开网",
        "2024-10",
        "official_university_pdf",
        "西北工业大学2024年本科专业设置PDF列出72个本科专业，备注说明黄色为当年停招、绿色为本年新增；表格中8个专业标注停招、2个专业标注新增。",
    ),
    SourceSeed(
        "nwpu_2025_major_setting_pdf",
        "西北工业大学本科专业设置（2025年12月更新）",
        "https://xxgk.nwpu.edu.cn/__local/7/2B/7C/B746BAEA5B42D0BA44C176B4191_85E461A5_1788D.pdf",
        "西北工业大学信息公开网",
        "2025-12",
        "official_university_pdf",
        "西北工业大学2025年本科专业设置PDF列出73个本科专业，备注说明黄色为当年停招、绿色为本年新增；表格中9个专业标注停招、2个专业标注新增。",
    ),
    SourceSeed(
        "jxufe_2025_major_adjustment",
        "江西财经大学2025年普通本科专业设置动态调整名单公示",
        "https://www.jxufe.edu.cn/news-show-274.html",
        "江西财经大学",
        "2025-07-07",
        "official_university_html",
        "江西财经大学2025年度普通本科专业设置动态调整公示列出2个撤销专业、3个新增停招专业和2个继续停招专业。",
    ),
    SourceSeed(
        "hbesxy_2024_major_cancellation",
        "湖北恩施学院关于2024年度拟撤销本科专业的公示",
        "https://jwc.hbesxy.edu.cn/info/1083/3395.htm",
        "湖北恩施学院教务处",
        "2024-08-08",
        "official_university_html",
        "湖北恩施学院2024年度拟撤销连续停招五年以上的8个本科专业。",
    ),
    SourceSeed(
        "jxau_2025_major_cancellation",
        "关于2025年度本科专业设置申报材料的公示",
        "https://www.jxau.edu.cn/info/1041/304581.htm",
        "江西农业大学",
        "2025-07-18",
        "official_university_html",
        "江西农业大学2025年度拟撤销信息与计算科学本科专业。",
    ),
    SourceSeed(
        "xupt_2025_major_cancellation",
        "关于2025年度调整本科专业情况的公示",
        "https://jyc.xupt.edu.cn/info/1195/3533.htm",
        "西安邮电大学教务处",
        "2025-07-18",
        "official_university_html",
        "西安邮电大学2025年度拟撤销广播电视工程本科专业。",
    ),
    SourceSeed(
        "lnu_2025_major_cancellation",
        "辽宁大学2025年度拟增设本科专业、拟撤销本科专业公示",
        "https://jwc.lnu.edu.cn/info/1124/2378.htm",
        "辽宁大学教务处",
        "2025-07-21",
        "official_university_html",
        "辽宁大学2025年度拟撤销信息与计算科学、摄影、广播电视学、网络与新媒体4个本科专业，其中后两项为第二学士学位。",
    ),
    SourceSeed(
        "lnu_2024_major_cancellation",
        "辽宁大学关于申报2024年度本科专业的公示",
        "https://jwc.lnu.edu.cn/info/1143/2238.htm",
        "辽宁大学教务处",
        "2024-08-23",
        "official_university_html",
        "辽宁大学2024年度拟撤销广播电视学、网络与新媒体、生物科学、生态学、材料化学、电子科学与技术、环境科学、信息管理与信息系统、旅游管理9个本科专业。",
    ),
    SourceSeed(
        "czjtu_2024_major_cancellation",
        "关于拟撤销过程装备与控制工程专业的公示",
        "https://jwc.czjtu.edu.cn/info/1083/1855.htm",
        "沧州交通学院教务处",
        "2024-08-22",
        "official_university_html",
        "沧州交通学院2024年度决定撤销过程装备与控制工程专业，并向教育部申请撤销备案。",
    ),
    SourceSeed(
        "gxust_2025_major_cancellation",
        "广西科技大学2025年度普通本科专业设置调整情况公示",
        "https://www.gxust.edu.cn/info/1037/33793.htm",
        "广西科技大学",
        "2025-07-22",
        "official_university_html",
        "广西科技大学2025年度拟撤销7个本科专业。",
    ),
    SourceSeed(
        "gxust_2016_undergrad_setting_stop_page",
        "普通本专科专业设置、当年新增专业、停招专业名单",
        "https://www.gxust.edu.cn/xxgk/info/1251/2362.htm",
        "广西科技大学信息公开网",
        "2017-01-04",
        "official_university_html",
        "广西科技大学信息公开页说明2015—2016学年停招教育技术学等14个本科专业。",
    ),
    SourceSeed(
        "bistu_2025_major_cancellation",
        "北京信息科技大学2025年度拟新增本科专业、预备案专业、拟撤销本科专业公示",
        "https://www.bistu.edu.cn/tzgg/07f55cc2576e4b29aad9462722b816c1.html",
        "北京信息科技大学",
        "2025-07-23",
        "official_university_html",
        "北京信息科技大学2025年度拟撤销英语、行政管理2个本科专业。",
    ),
    SourceSeed(
        "kmust_2025_major_cancellation",
        "关于昆明理工大学2025年度本科专业设置推荐结果的公示",
        "https://www.kmust.edu.cn/info/1013/51267.htm",
        "昆明理工大学",
        "2025-07-23",
        "official_university_html",
        "昆明理工大学2025年度本科专业设置推荐结果公示列出4个拟撤销本科专业。",
    ),
    SourceSeed(
        "haut_2025_major_cancellation",
        "关于对2025年本科新专业申报、预备案和撤销专业进行公示的通知",
        "https://jwc.haut.edu.cn/info/1114/10861.htm",
        "河南工业大学教务处",
        "2025-07-13",
        "official_university_html",
        "河南工业大学2025年度拟撤销食品营养与检验教育、交通工程、电子信息科学与技术3个本科专业。",
    ),
    SourceSeed(
        "wit_2025_major_cancellation",
        "武汉工程大学2025年度拟增设、撤销及调整本科专业的公示",
        "https://jwc.wit.edu.cn/info/1016/23181.htm",
        "武汉工程大学本科生院",
        "2025-07-23",
        "official_university_html",
        "武汉工程大学2025年度拟向教育部申请撤销5个本科专业。",
    ),
    SourceSeed(
        "ahstu_2025_major_adjustment",
        "关于2025年度预申报、撤销和调整修业年限本科专业的公示",
        "https://www.ahstu.edu.cn/jwc/info/1887/11108.htm",
        "安徽科技学院教务处",
        "2025-07-20",
        "official_university_html",
        "安徽科技学院2025年度本科专业调整公示列出预申报专业5个、调整修业年限专业1个、撤销专业2个，并提供专业代码、修业年限、学位授予门类和所在学院。",
    ),
    SourceSeed(
        "lnnu_2025_major_cancellation",
        "关于2025年度我校拟申报新增及撤销本科专业的公示",
        "https://news.lnnu.edu.cn/info/1163/8816.htm",
        "辽宁师范大学教务处",
        "2025-07-22",
        "official_university_html",
        "辽宁师范大学2025年度拟撤销绘画、环境科学、应用化学3个本科专业。",
    ),
    SourceSeed(
        "utibet_2025_major_adjustment",
        "西藏大学2025年度拟设置与调整本科专业公示",
        "https://jwc.utibet.edu.cn/info/1016/1960.htm",
        "西藏大学教务处",
        "2025-07-09",
        "official_university_html",
        "西藏大学2025年度拟新增新能源科学与工程、人工智能2个本科专业，拟撤销市场营销、财务管理、服装与服饰设计、新闻学4个本科专业，拟调整城乡规划修业年限。",
    ),
    SourceSeed(
        "sxdzkj_2025_major_cancellation",
        "山西电子科技学院关于2025年度拟增设、调整和撤销专业的公示",
        "https://www.sxdzkj.edu.cn/xxgkw/info/1052/1169.htm",
        "山西电子科技学院信息公开网",
        "2025-07-23",
        "official_university_html",
        "山西电子科技学院2025年度拟撤销思想政治教育、学前教育、汉语言文学、戏剧影视文学、广播电视编导、数字媒体艺术6个本科专业。",
    ),
    SourceSeed(
        "guit_2025_major_cancellation",
        "关于我校2025年度拟新设及撤销本科专业的公示",
        "https://www.guit.edu.cn/jwkj/info/1025/2502.htm",
        "桂林信息科技学院教务处",
        "2025-07-20",
        "official_university_html",
        "桂林信息科技学院2025年度申请撤销已停招5年以上且无在校生的材料成型及控制工程专业。",
    ),
    SourceSeed(
        "xawl_2025_major_cancellation",
        "2025年度本科新专业申报结果公示",
        "https://jwc.xawl.edu.cn/info/1116/13089.htm",
        "西安文理学院教务处",
        "2025-07-13",
        "official_university_html",
        "西安文理学院2025年度拟撤销4个本科专业。",
    ),
    SourceSeed(
        "xawl_2026_major_cancellation",
        "2026年度本科新专业申报结果公示",
        "https://www.xawl.edu.cn/jwc/info/1116/14259.htm",
        "西安文理学院教务处",
        "2026-05-14",
        "official_university_html",
        "西安文理学院2026年度拟撤销化学工程与工艺本科专业。",
    ),
    SourceSeed(
        "xync_2026_major_adjustment",
        "关于咸阳师范学院2026年度拟新增、拟撤销、预申报本科专业的公示",
        "https://jwc.xync.edu.cn/content.jsp?urltype=news.NewsContentUrl&wbnewsid=25231&wbtreeid=1032",
        "咸阳师范学院教务处",
        "2026-05-11",
        "official_university_html",
        "咸阳师范学院2026年度拟新增低空经济与管理等5个本科专业，拟撤销视觉传达设计专业，拟预申报供应链管理等25个本科专业。",
    ),
    SourceSeed(
        "cuggw_2024_major_adjustment",
        "关于2024年度新增、撤销本科专业情况的公示",
        "https://www.cuggw.com/3g/show.asp?d=19510&m=1",
        "保定理工学院教务处",
        "2024-08-22",
        "official_university_html",
        "保定理工学院2024年度新增、撤销本科专业情况公示列出拟新增软件工程、应用心理学2个本科专业，拟撤销机械电子工程、建筑学2个本科专业。",
    ),
    SourceSeed(
        "changdian_2023_major_cancellation",
        "关于我校2023年度拟撤销专业的公示",
        "https://www3.changdian2001.com/html/jwc/tzgg/2023/0813/5221.html",
        "长春电子科技学院教务处",
        "2023-07-26",
        "official_university_html",
        "长春电子科技学院2023年度拟撤销光源与照明、网络工程2个本科专业，并说明光源与照明2021年停止招生、网络工程2020年停止招生。",
    ),
    SourceSeed(
        "hynu_2025_major_cancellation",
        "衡阳师范学院关于2025年本科专业设置调整情况的公示",
        "https://jwc.hynu.edu.cn/info/1098/6503.htm",
        "衡阳师范学院教务处",
        "2025-07-16",
        "official_university_html",
        "衡阳师范学院2025年度拟撤销编辑出版学本科专业。",
    ),
    SourceSeed(
        "lyun_2025_major_cancellation",
        "龙岩学院关于2025年度本科专业设置的公示",
        "https://jwc.lyun.edu.cn/info/1191/5301.htm",
        "龙岩学院教务处",
        "2025-07-16",
        "official_university_html",
        "龙岩学院2025年度拟申请撤销日语本科专业。",
    ),
    SourceSeed(
        "henau_2025_major_cancellation",
        "关于2025年度本科专业设置及专业结构调整优化的公示",
        "https://jwc.henau.edu.cn/a/tongzhigonggao/xjk/2025/0729/7109.html",
        "河南农业大学教务处",
        "2025-07-29",
        "official_university_html",
        "河南农业大学2025年度拟撤销应用生物科学、草业科学、市场营销、植物科学与技术、食品营养与检验教育、资源循环科学与工程6个本科专业。",
    ),
    SourceSeed(
        "jift_2025_major_cancellation",
        "关于2025年拟新增税收学、智能影像艺术、时尚传播普通本科专业及拟撤销国际商务专业的公示",
        "https://jwc.jift.edu.cn/info/1068/5971.htm",
        "江西服装学院教务处",
        "2025-07-22",
        "official_university_html",
        "江西服装学院2025年度拟撤销国际商务普通本科专业。",
    ),
    SourceSeed(
        "lshhxy_2025_major_cancellation",
        "关于辽宁师范大学海华学院2025年度撤销1个本科专业的公示",
        "https://jwc.lshhxy.edu.cn/info/1086/6333.htm",
        "辽宁师范大学海华学院教务处",
        "2025-07-10",
        "official_university_html",
        "辽宁师范大学海华学院2025年度撤销电子商务及法律本科专业。",
    ),
    SourceSeed(
        "glc_2025_major_adjustment_notice",
        "【教务处】桂林学院2025年度拟申报专业材料公示",
        "https://www.glc.edu.cn/info/1374/45421.htm",
        "桂林学院",
        "2025-07-31",
        "official_university_html",
        "桂林学院2025年度拟停招专业和拟撤销专业公示页，正文链接拟撤销和拟停招专业名单Word附件。",
    ),
    SourceSeed(
        "glc_2025_major_cancellation_docx",
        "2025年桂林学院拟撤销专业名单",
        "https://www.glc.edu.cn/system/_content/download.jsp?urltype=news.DownloadAttachUrl&owner=1401022929&wbfileid=9B700D5BE299BA841F3BFD5A9D7E2CE1",
        "桂林学院",
        "2025-07-28",
        "official_university_docx",
        "桂林学院2025年度拟撤销互联网金融、资产评估、数字出版3个本科专业。",
    ),
    SourceSeed(
        "glc_2025_stop_enrollment_docx",
        "2025年桂林学院拟停招专业名单",
        "https://www.glc.edu.cn/system/_content/download.jsp?urltype=news.DownloadAttachUrl&owner=1401022929&wbfileid=7DEFF53F040CD1221A2016A156987570",
        "桂林学院",
        "2025-07-31",
        "official_university_docx",
        "桂林学院2025年度拟停招经济学、保险学、投资学、会计学、播音与主持艺术、酒店管理、会展经济与管理、工艺美术8个本科专业。",
    ),
    SourceSeed(
        "jhuni_2024_major_cancellation",
        "教学指导委员会关于2024年度专业设置调整评议结果的公示",
        "https://jwzx.jhun.edu.cn/0d/77/c3825a200055/pagem.htm",
        "江汉大学教学指导委员会",
        "2024-08-02",
        "official_university_html",
        "江汉大学2024年度拟撤销7个本科专业。",
    ),
    SourceSeed(
        "xzhmu_2025_major_cancellation",
        "关于徐州医科大学2025年拟新增本科专业、预备案本科专业、撤销专业的公示",
        "https://jwc.xzhmu.edu.cn/info/1093/5247.htm",
        "徐州医科大学教务处",
        "2025-07-14",
        "official_university_html",
        "徐州医科大学2025年度拟撤销物联网工程本科专业。",
    ),
    SourceSeed(
        "sjtu_2024_major_cancellation_pdf",
        "2024年度拟撤销专业列表",
        "https://www.sjtu.edu.cn/resource/upload/202408/20240811_102131_964.pdf",
        "上海交通大学",
        "2024-08-11",
        "official_university_pdf",
        "上海交通大学2024年度拟撤销资源环境科学、公共事业管理、软件工程（二年制）、信息安全（二年制）4个专业，撤销原因均为已停招5年及以上。",
    ),
    SourceSeed(
        "sjtu_2025_psychology_application_pdf",
        "普通高等学校本科专业设置申请表：心理学",
        "https://www.sjtu.edu.cn/resource/upload/202507/20250722_101923_495.pdf",
        "上海交通大学",
        "2025-07-22",
        "official_university_pdf",
        "上海交通大学2025年心理学本科专业设置申请表在学校近五年专业增设、停招、撤并情况栏说明2023年撤销5个专业，2024年撤销4个专业、停招3个专业；原文使用“等”，不作为完整名单。",
    ),
    SourceSeed(
        "hue_2025_major_cancellation",
        "关于2025年度拟新增本科专业、预备案本科专业、撤销专业的公示",
        "https://jwc.hue.edu.cn/2025/0713/c9502a190014/pagem.htm",
        "湖北第二师范学院教务处",
        "2025-07-13",
        "official_university_html",
        "湖北第二师范学院2025年度拟撤销信息与计算科学、物流工程、公共事业管理、汽车服务工程4个本科专业。",
    ),
    SourceSeed(
        "swjtu_2025_major_adjustment_notice",
        "西南交通大学2025年本科专业调整公示",
        "https://news.swjtu.edu.cn/info/1020/84345.htm",
        "西南交通大学新闻网",
        "2025-07-17",
        "official_university_html",
        "西南交通大学2025年本科专业调整公示列出拟撤销电气工程与智能控制、交通设备与控制工程、国际经济与贸易、视觉传达设计、商务英语5个本科专业。",
    ),
    SourceSeed(
        "jzun_2025_major_adjustment_notice",
        "荆州学院2025年本科专业设置调整情况的公示",
        "https://www.jzun.edu.cn/jwc/info/389/1614.htm",
        "荆州学院教务处",
        "2025-07-22",
        "official_university_html",
        "荆州学院2025年本科专业设置调整情况公示列出拟撤销市场营销、车辆工程2个本科专业。",
    ),
    SourceSeed(
        "njucm_2025_major_cancellation",
        "南京中医药大学关于2025年度拟撤销本科专业的公示",
        "https://jwc.njucm.edu.cn/2026/0416/c4498a167633/page.htm",
        "南京中医药大学教务处",
        "2025-07-24",
        "official_university_html",
        "南京中医药大学2025年度拟撤销生物技术、劳动与社会保障、市场营销、健康服务与管理4个本科专业。",
    ),
    SourceSeed(
        "sohu_2025_ten_university_cancellation_summary",
        "多所高校公示：撤销这些专业",
        "https://www.sohu.com/a/914669400_100262768",
        "搜狐转载高考直通车",
        "2025-07-17",
        "secondary_news_html",
        "二级汇总来源，整理10所高校2025年度撤销/停招本科专业公示；仅结构化未被本数据集官方来源覆盖的学校名单，并标为medium置信度。",
    ),
    SourceSeed(
        "gkztc_2025_multi_university_adjustment",
        "多所大学官宣撤销一批本科专业",
        "https://app.gaokaozhitongche.com/news/h/WO8KZ3Yz",
        "高考直通车/求学杂志",
        "2025-07-21",
        "secondary_news_html",
        "二级汇总来源，整理甘肃农业大学、成都信息工程大学、重庆师范大学、宁波大学、西华师范大学等高校2025年度本科专业动态调整信息。",
    ),
    SourceSeed(
        "qdfa_2025_major_cancellation",
        "关于2025年度拟撤销本科专业的公示",
        "https://www.qdfa.edu.cn/index/list/59.html",
        "青岛电影学院",
        "2025-07-17",
        "official_university_html",
        "青岛电影学院官网通知列表页直接展示2025年度拟撤销视觉传达设计、音乐表演、流行音乐3个本科专业的完整正文。",
    ),
    SourceSeed(
        "jxust_2024_major_cancellation",
        "江西理工大学2024年拟新增专业及拟撤销本科专业公示",
        "https://www.jxust.edu.cn/info/1061/32775.htm",
        "江西理工大学",
        "2024-08-21",
        "official_university_html",
        "江西理工大学2024年度拟撤销物联网工程、物流管理、电子商务、视觉传达设计4个本科专业。",
    ),
    SourceSeed(
        "qq_eol_2025_multi_university_cancellation",
        "10余所高校，拟撤销这些本科专业！",
        "https://news.qq.com/rain/a/20250715A05A1E00",
        "腾讯新闻转载中国教育在线",
        "2025-07-15",
        "secondary_news_html",
        "二级汇总来源，整理10余所高校2025年度本科专业动态调整公示；仅结构化本数据集中官方来源暂未覆盖的学校名单，并标为medium置信度。",
    ),
    SourceSeed(
        "eol_2025_multi_university_cancellation",
        "多所高校拟撤销部分本科专业！",
        "https://www.eol.cn/news/dongtai/202507/t20250713_2680665.shtml",
        "教育在线",
        "2025-07-13",
        "secondary_news_html",
        "二级汇总来源，正文列出成都信息工程大学、宁波大学、重庆师范大学、井冈山大学、马鞍山学院、合肥师范学院、四川美术学院等高校2025年度拟撤销/停招本科专业，并注明来源为各高校官网。",
    ),
    SourceSeed(
        "cqnews_eol_2025_multi_university_cancellation",
        "10余所高校，拟撤销这些本科专业！",
        "https://www.cqnews.net/1/detail/1396405287704801280/app/content_1396405287704801280.html",
        "华龙网转载中国教育在线",
        "2025-07-20",
        "secondary_news_html",
        "二级汇总来源，转载中国教育在线整理，正文列出成都信息工程大学、马鞍山学院、合肥师范学院、四川美术学院等高校2025年度拟撤销/停招本科专业，并注明来源为中国教育在线综合各高校官网。",
    ),
    SourceSeed(
        "qq_zsgk_2025_multi_university_adjustment",
        "重点关注！这些本科专业即将被撤销",
        "https://news.qq.com/rain/a/20250725A0839K00",
        "腾讯新闻转载掌上高考",
        "2025-07-25",
        "secondary_news_html",
        "二级汇总来源，转载掌上高考整理，正文列出江西师范大学、福建师范大学、成都信息工程大学、西华师范大学等高校2025年度拟撤销/停招本科专业，并标注图片来源于各校。",
    ),
    SourceSeed(
        "gz55zs_2025_university_adjustment_table",
        "近30所高校专业调整，这些本科专业被裁撤",
        "https://gz.55zs.com/content/detail/93273",
        "新期教育网",
        "2025-07-31",
        "secondary_news_html",
        "二级汇总来源，以表格形式列出2025年近30所高校拟撤销和拟新增专业，覆盖长春大学、福建师范大学、泉州信息工程学院等学校。",
    ),
    SourceSeed(
        "gkzxw_cxtc_2025_major_cancellation",
        "楚雄师范学院关于2025年拟新设本科专业和拟撤销本科专业的公示",
        "https://www.gkzxw.com/gxzs/202507/71459.html",
        "高考资讯网转载楚雄师范学院",
        "2025-07-28",
        "secondary_news_html",
        "二级镜像来源，注明文章来源为楚雄师范学院，正文完整列出2025年拟撤销中国少数民族语言文学、酒店管理2个专业，并给出专业代码、修业年限、学位和所属学院。",
    ),
    SourceSeed(
        "gkztc_2025_sjtu_major_cancellation",
        "上海交通大学拟增设4个专业",
        "https://app.gaokaozhitongche.com/newsguide/h/VGOM0paX",
        "高考直通车",
        "2025-07-28",
        "secondary_news_html",
        "二级来源，整理上海交通大学2025年度拟增设/撤销本科专业公示，正文提到拟撤销临床医学（七年制）、口腔医学（七年制）等专业。",
    ),
    SourceSeed(
        "qq_qingtah_2025_sjtu_major_cancellation",
        "C9，拟撤销多个专业",
        "https://news.qq.com/rain/a/20250727A05TL600",
        "腾讯新闻转载青塔",
        "2025-07-27",
        "secondary_news_html",
        "二级来源，正文称上海交通大学2025年度拟撤销临床医学（七年制）、口腔医学（七年制）、法学、传播学、计算机科学与技术、行政管理等6个专业，并注明来源为上海交通大学。",
    ),
    SourceSeed(
        "sohu_2025_ten_university_cancellation_aug02",
        "暨南大学等10校共36个专业拟被撤销！",
        "https://m.sohu.com/a/920159136_121124337/?pvid=000115_3w_a",
        "搜狐转载",
        "2025-08-02",
        "secondary_news_html",
        "二级汇总来源，整理暨南大学、扬州大学、成都理工大学、唐山学院、中南林业科技大学涉外学院等10所高校2025年度拟撤销本科专业信息；仅结构化正文明确列名的学校。",
    ),
    SourceSeed(
        "acabridge_2025_six_university_cancellation",
        "武汉工程大学等6校19个专业拟被撤销！",
        "https://www.acabridge.cn/zxhz/202507/t20250730_2683898.shtml",
        "学术桥",
        "2025-07-30",
        "secondary_news_html",
        "二级汇总来源，正文列出武汉工程大学、扬州大学、齐鲁工业大学、河北工业大学、信阳师范大学、北京工业大学等6校拟撤销专业名单。",
    ),
    SourceSeed(
        "qlu_2025_major_cancellation",
        "齐鲁工业大学（山东省科学院）2025年度撤销本科专业的公示",
        "https://www.sdas.org/2025/0725/c2a259690/page.htm",
        "齐鲁工业大学（山东省科学院）",
        "2025-07-25",
        "official_university_html",
        "齐鲁工业大学（山东省科学院）2025年度撤销本科专业公示列出安全工程、汽车服务工程、保险学3个拟申请撤销专业。",
    ),
    SourceSeed(
        "xynu_2025_major_cancellation",
        "关于2025年度拟申报、撤销专业的公示",
        "http://jwc.xynu.edu.cn/info/1038/10545.htm",
        "信阳师范大学教务处",
        "2025-07-27",
        "official_university_html",
        "信阳师范大学2025年度拟申请撤销人文教育、应用物理学、信息管理与信息系统3个本科专业。",
    ),
    SourceSeed(
        "bjut_2025_major_cancellation",
        "2025年度拟增设新专业和拟撤销专业情况公示",
        "https://undergrad.bjut.edu.cn/info/1243/3705.htm",
        "北京工业大学教务处",
        "2025-07-28",
        "official_university_html",
        "北京工业大学2025年度拟撤销日语、广告学、食品质量与安全、风景园林4个本科专业。",
    ),
    SourceSeed(
        "nwsuaf_2025_major_cancellation",
        "关于对 2025 年拟新增与撤销专业的公示",
        "https://zhxy.nwsuaf.edu.cn/tzgg/75cb7fd7f6174986842b7a6393eb715a.htm",
        "西北农林科技大学资源环境学院",
        "2025-07-17",
        "official_university_html",
        "西北农林科技大学2025年度拟撤销保险学、视觉传达设计2个本科专业。",
    ),
    SourceSeed(
        "ymzy_2025_multi_university_adjustment",
        "多所高校调整开设！这些本科专业撤销",
        "https://m.ymzy.cn/article/7929",
        "四川省招生考试指导中心优志愿",
        "2025-08",
        "secondary_news_html",
        "二级汇总来源，整理上海交通大学、暨南大学、北京信息科技大学、沈阳航空航天大学、成都理工大学等高校2025年本科专业撤销/调整信息。",
    ),
    SourceSeed(
        "sau_2025_major_adjustment_notice",
        "关于2025年度新增、撤销专业的公示",
        "https://jwc.sau.edu.cn/info/1052/1457.htm",
        "沈阳航空航天大学本科生院",
        "2025-08-07",
        "official_university_html",
        "沈阳航空航天大学官网公示2025年度新增、撤销专业工作；正文列出车辆工程、工业工程、功能材料、焊接技术与工程、日语、英语6个已停招拟撤销专业。",
    ),
    SourceSeed(
        "gkzxw_cdut_2025_major_adjustment",
        "成都理工大学关于2025年拟新设专业与拟撤销专业的公示",
        "https://www.gkzxw.com/gxzs/202507/71455.html",
        "高考资讯网转载成都理工大学",
        "2025-07-28",
        "secondary_news_html",
        "二级镜像来源，注明文章来源为成都理工大学，正文完整列出2025年拟撤销地球信息科学与技术等11个本科专业及专业代码。",
    ),
    SourceSeed(
        "gkzxw_nenu_2025_major_adjustment",
        "关于东北师范大学2025年度本科专业设置情况的公示",
        "https://www.gkzxw.com/gxzs/202507/71470.html",
        "高考资讯网转载东北师范大学",
        "2025-07-28",
        "secondary_news_html",
        "二级镜像来源，注明文章来源为东北师范大学，正文说明2025年度拟撤销财政学等7个本科专业；完整名单需结合公开汇总表。",
    ),
    SourceSeed(
        "nenu_2025_major_adjustment_notice",
        "关于东北师范大学2025年度本科专业设置情况的公示",
        "https://jwc.nenu.edu.cn/info/1084/11034.htm",
        "东北师范大学教务处",
        "2025-07-22",
        "official_university_html",
        "东北师范大学教务处官网保留2025年度本科专业设置情况公示标题、发布时间和“公示已结束”状态；当前页面未暴露专业名单正文。",
    ),
    SourceSeed(
        "nenu_2024_teaching_quality_report_pdf",
        "东北师范大学2023-2024学年本科教育教学质量报告",
        "https://publish.nenu.edu.cn/__local/C/BD/7C/42C6B27B0AF77D70D26029AF35B_D3B4F6AB_50525E.pdf",
        "东北师范大学信息公开网",
        "2025-05-26",
        "official_university_pdf",
        "东北师范大学2023-2024学年本科教育教学质量报告PDF，列出本科专业总数、招生专业数、一流专业建设点数量，并说明主动撤销停招5年以上专业。",
    ),
    SourceSeed(
        "cwnu_2024_teaching_quality_report_pdf",
        "西华师范大学2023-2024学年本科教学质量报告",
        "https://cfteatd.cwnu.edu.cn/__local/E/08/50/85917A7AC8CA4083F992E8A538C_A9D698A5_7854B.pdf",
        "西华师范大学教师教育学院/教师教学发展中心",
        "2024-12-18",
        "official_university_pdf",
        "西华师范大学2023-2024学年本科教学质量报告PDF，列出本科专业总数、师范/非师范专业数量、一流专业建设点数量，以及近五年新增、规划申报、撤销专业数量。",
    ),
    SourceSeed(
        "fjnu_2024_teaching_quality_report_pdf",
        "福建师范大学2023-2024学年本科教学质量报告",
        "https://jwc.fjnu.edu.cn/_upload/article/files/f4/e5/1fe740724920b0bc9fdc18543ddc/0d9bf179-e675-4794-a718-eacef684e901.pdf",
        "福建师范大学教务处",
        "2024-12-31",
        "official_university_pdf",
        "福建师范大学2023-2024学年本科教学质量报告PDF，列出本科专业总数、招生专业数，并说明本学年停招酒店管理、动画、环境科学、复合材料与工程4个专业，新增足球运动专业。",
    ),
    SourceSeed(
        "jsnu_2024_teaching_quality_report_pdf",
        "江苏师范大学2023-2024学年本科教学质量报告",
        "http://www.jsnu.edu.cn/_upload/article/files/da/78/1ba9aa7141ac9e95addcc86c6cf2/45cd2468-ff8c-4fe0-9e12-fda74351fd80.pdf",
        "江苏师范大学信息公开网",
        "2024-12-26",
        "official_university_pdf",
        "江苏师范大学2023-2024学年本科教学质量报告PDF，列出本科专业总数、招生专业数、一流本科专业建设点数量，并说明学校完善专业动态调整机制。",
    ),
    SourceSeed(
        "gkzxw_jlau_2025_major_cancellation",
        "吉林农业大学2025年度拟撤销“广告学”4个本科专业",
        "https://www.gkzxw.com/gxzs/202507/71487.html",
        "高考资讯网转载吉林农业大学",
        "2025-07-28",
        "secondary_news_html",
        "二级镜像来源，注明文章来源为吉林农业大学，正文表格列出2025年拟申请撤销广告学、应用心理学、农业机械化及其自动化、野生动物与自然保护区管理4个专业及代码。",
    ),
    SourceSeed(
        "gkzxw_jsnu_2025_major_adjustment",
        "江苏师范大学2025年拟设置3个本科专业、拟撤销3个本科专业",
        "https://www.gkzxw.com/gxzs/202507/71515.html",
        "高考资讯网转载江苏师范大学",
        "2025-07-28",
        "secondary_news_html",
        "二级镜像来源，注明文章来源为江苏师范大学，正文表格列出2025年拟撤销环境科学、土地资源管理、电子科学与技术3个本科专业及撤销理由。",
    ),
    SourceSeed(
        "gkzxw_bjut_2025_major_adjustment",
        "北京工业大学2025年度拟增设新专业和拟撤销专业情况公示",
        "https://www.gkzxw.com/gxzs/202507/71536.html",
        "高考资讯网转载北京工业大学",
        "2025-07-29",
        "secondary_news_html",
        "二级镜像来源，注明文章来源为北京工业大学，正文列出2025年拟撤销日语、广告学、食品质量与安全、风景园林4个本科专业。",
    ),
    SourceSeed(
        "gkzxw_hfnu_2025_major_adjustment",
        "合肥师范学院2025年拟增设3个本科专业、撤销1个本科专业",
        "https://www.gkzxw.com/gxzs/202507/71494.html",
        "高考资讯网转载合肥师范学院",
        "2025-07-28",
        "secondary_news_html",
        "二级镜像来源，注明文章来源为合肥师范学院，正文表格列出2025年撤销视觉传达设计（中外合作）并给出专业代码、修业年限、学位和学院。",
    ),
    SourceSeed(
        "tsc_2025_major_adjustment",
        "唐山学院关于2025年度拟申报本科专业和拟撤销本科专业的公示",
        "https://wljyzx.tsc.edu.cn/col/1414661452486/2025/07/22/1756455584212.html",
        "唐山学院现代教育技术中心",
        "2025-07-22",
        "official_university_html",
        "唐山学院官网公示2025年度拟撤销汽车服务工程、酒店管理2个本科专业。",
    ),
    SourceSeed(
        "csuft_swxy_2025_major_adjustment",
        "关于学院2025年拟新增、撤销本科专业的公示",
        "https://swxy.csuft.edu.cn/jwc/info/1019/1051.htm",
        "中南林业科技大学涉外学院教务处",
        "2025-07-21",
        "official_university_html",
        "中南林业科技大学涉外学院官网公示2025年拟撤销机械设计制造及其自动化、交通运输、环境科学、城乡规划、园林5个本科专业。",
    ),
    SourceSeed(
        "hebut_2025_major_cancellation",
        "关于2025年度拟申请撤销本科专业的公示",
        "https://www.hebut.edu.cn/tzgg/18cce0fa04184f098d3657d15114b89c.htm",
        "河北工业大学",
        "2025-07-26",
        "official_university_html",
        "河北工业大学官网公示2025年度拟申请撤销智能科学与技术、交通运输2个本科专业。",
    ),
    SourceSeed(
        "sthu_2025_major_cancellation",
        "关于我校2025年度拟申报及撤销本科专业的公示",
        "https://info.sthu.edu.cn/47/b2/c1303a83890/page.htm",
        "上海师范大学天华学院教务处",
        "2025-07-21",
        "official_university_html",
        "上海师范大学天华学院官网公示2025年度拟撤销机械设计制造及其自动化、汉语国际教育2个本科专业。",
    ),
    SourceSeed(
        "sohu_2025_six_university_cancellation_summary",
        "读不到了！6校共25个专业将被撤销！",
        "https://www.sohu.com/a/917880765_121124337",
        "搜狐转载",
        "2025-07-26",
        "secondary_news_html",
        "二级汇总来源，整理6所高校2025年度拟撤销本科专业；仅结构化本数据集中官方来源暂未覆盖的学校名单，并标为medium置信度。",
    ),
    SourceSeed(
        "zafu_2025_major_cancellation_pdf",
        "浙江农林大学拟撤销专业汇总表",
        "https://www.zafu.edu.cn/system/_content/download.jsp?owner=1261846942&urltype=news.DownloadAttachUrl&wbfileid=7166403C8A363998F50CE48BAD0EF68B",
        "浙江农林大学",
        "2025-07",
        "official_university_pdf",
        "浙江农林大学官网PDF附件列出2025年度拟撤销环境工程、森林保护、市场营销3个本科专业，并给出专业代码、修业年限、学位授予门类、开设年份和停招年份。",
    ),
    SourceSeed(
        "wsu_2025_major_cancellation",
        "关于2025年拟撤销本科专业的公示",
        "https://www.wsu.edu.cn/info/1011/37042.htm",
        "文山学院",
        "2025-07-22",
        "official_university_html",
        "文山学院2025年拟向教育部申请撤销食品质量与安全、工程管理、视觉传达设计3个本科专业；正文列出专业代码、修业年限和学位门类。",
    ),
    SourceSeed(
        "xmhbjy_2025_multi_university_cancellation",
        "官宣！福师大等多校公布撤销这些专业",
        "https://xmhbjy.com/sys-nd/88.html",
        "浩博教育",
        "2025-08-09",
        "secondary_news_html",
        "二级汇总来源，整理多所高校2025年度拟撤销本科专业；仅结构化本数据集中官方来源暂未覆盖的学校名单，并标为medium置信度。",
    ),
    SourceSeed(
        "nuaa_2025_major_adjustment",
        "关于我校2025年度本科专业设置拟调整的公示",
        "https://nuaa.edu.cn/_t55/2025/0722/c295a380353/page.htm",
        "南京航空航天大学",
        "2025-07-22",
        "official_university_html",
        "南京航空航天大学2025年度拟撤销美术学、空间科学与技术、空间信息与数字技术3个本科专业。",
    ),
    SourceSeed(
        "jou_2025_major_adjustment",
        "关于2025年学校本科专业设置情况的公示",
        "https://www.jou.edu.cn/info/1052/40229.htm",
        "江苏海洋大学",
        "2025-06-27",
        "official_university_html",
        "江苏海洋大学2025年度本科专业设置公示列出2个申报新专业、2个预备案专业和4个拟撤销专业，并提供专业代码、专业类、学科门类和学院。",
    ),
    SourceSeed(
        "zju_2025_undergrad_major_table",
        "浙江大学本科专业情况表",
        "https://www.zju.edu.cn/xxgk/2024/0912/c17970a2961786/page.htm",
        "浙江大学信息公开网",
        "2025-10-12",
        "official_university_html",
        "浙江大学本科专业情况表列出130个本科专业，并在备注列标注28个已停招专业。",
    ),
    SourceSeed(
        "zjyc_2024_cancel_stop",
        "关于学院2024年申请增设专业和撤销停招专业的公示",
        "https://jwb.zjyc.edu.cn/info/10475/92043.htm",
        "浙江农林大学暨阳学院教务部",
        "2024-08-25",
        "official_university_html",
        "浙江农林大学暨阳学院2024年申请增设专业和撤销停招专业公示列出5个拟申请增设专业和2个拟撤销停招本科专业，并提供专业代码、学位和停招起始年份。",
    ),
    SourceSeed(
        "hfit_2025_major_setting",
        "2025年专业设置、当年新增专业、停招专业名单",
        "https://xxgk.hfit.edu.cn/2025/0522/c162a354/page.htm",
        "合肥理工学院信息公开网",
        "2025-05-22",
        "official_university_html",
        "合肥理工学院2025年专业基本情况数据表列出41个本科专业，其中2个当年新增、31个当年停招，并提供专业代码和门类名称。",
    ),
    SourceSeed(
        "xjnu_2025_major_adjustment",
        "关于2025年度本科专业设置情况的公示",
        "https://jwc.xjnu.edu.cn/info/1432/9003.htm",
        "新疆师范大学教务处",
        "2025-05-30",
        "official_university_html",
        "新疆师范大学2025年度本科专业设置情况公示列出4个拟新增专业、5个预备案专业，并说明拟撤销舞蹈表演专业。",
    ),
    SourceSeed(
        "ccit_2023_major_cancellation",
        "关于我校2023年度拟撤销本科专业的公示",
        "https://jwc.ccit.edu.cn/info/1057/4367.htm",
        "长春工程学院教务处",
        "2023-08-22",
        "official_university_html",
        "长春工程学院2023年度拟撤销本科专业公示列出信息与计算科学、机械电子工程、工业设计、数字媒体技术、视觉传达设计、勘查技术与工程6个已连续停招5年以上的本科专业，并提供专业代码和学位门类。",
    ),
    SourceSeed(
        "gsau_undergrad_cancel_list",
        "甘肃农业大学撤销本科专业名单",
        "https://jiaowu.gsau.edu.cn/info/1383/11150.htm",
        "甘肃农业大学教务处",
        "2024-09-29",
        "official_university_html",
        "甘肃农业大学撤销本科专业名单，列出开设年份、停招年份、专业代码、专业名称、学位授予门类、所属学院和撤销年份。",
    ),
    SourceSeed(
        "gsau_undergrad_stop_list",
        "甘肃农业大学停招本科专业名单",
        "https://jiaowu.gsau.edu.cn/info/1383/11140.htm",
        "甘肃农业大学教务处",
        "2024-09-29",
        "official_university_html",
        "甘肃农业大学停招本科专业名单，列出开设年份、停招年份、专业代码、专业名称、学位授予门类和所属学院。",
    ),
    SourceSeed(
        "jhun_2025_major_adjustment",
        "关于2025年度专业设置调整评议结果的公示",
        "https://jwzx.jhun.edu.cn/4f/96/c3825a216982/page.htm",
        "江汉大学教务处",
        "2025-07-15",
        "official_university_html",
        "江汉大学2025年度专业设置调整评议结果公示，列出拟撤销市场营销、财务管理、通信工程、网络工程、过程装备与控制工程、工业设计6个本科专业。",
    ),
    SourceSeed(
        "sciencenet_sdu_2025_undergrad_adjustment",
        "山东大学27个本科专业暂停招生",
        "https://news.sciencenet.cn/htmlnews/2025/2/538970.shtm",
        "科学网转载山东大学本科生院",
        "2025-02-19",
        "secondary_news_html",
        "二级来源转载山东大学本科生院《2023年9月至2024年8月本科专业设置及调整情况》，列出27个暂停招生专业和10个撤销专业。",
    ),
    SourceSeed(
        "gkztc_sdu_2025_undergrad_adjustment",
        "山东大学27个专业暂停招生，10个专业被撤销",
        "https://app.gaokaozhitongche.com/newsguide/h/Rg1KYJzN",
        "高考直通车转载求学杂志",
        "2025-02-26",
        "secondary_news_html",
        "高考直通车文章称山东大学本科生院发布《2023年9月至2024年8月本科专业设置及调整情况》，逐名列出27个暂停招生专业和10个撤销专业。",
    ),
    SourceSeed(
        "sohu_2024_multi_university_cancellation_summary",
        "山大、青大……全国多所高校撤销上百个专业！今年它们缘何“大洗牌”？",
        "https://www.sohu.com/a/806210079_121311734",
        "搜狐转载",
        "2024-09-04",
        "secondary_news_html",
        "二级汇总来源，整理2024年度多所高校撤销/停招专业信息；仅结构化本数据集中官方来源暂未覆盖且正文列出完整名单的学校。",
    ),
]


def red_ug_2010_2019() -> dict[int, list[str]]:
    return {
        2010: ["动画", "法学", "生物技术", "生物科学与工程", "数学与应用数学", "体育教育", "生物工程", "计算机科学与技术", "英语", "国际经济与贸易"],
        2011: ["动画", "法学", "生物技术", "生物科学与工程", "数学与应用数学", "体育教育", "生物工程", "计算机科学与技术", "英语", "国际经济与贸易"],
        2012: ["动画", "法学", "生物技术", "生物科学与工程", "数学与应用数学", "体育教育", "生物工程", "英语", "国际经济与贸易"],
        2013: ["动画", "法学", "生物技术", "生物科学与工程", "数学与应用数学", "体育教育", "生物工程", "英语", "美术学"],
        2014: ["生物科学与工程", "法学", "生物技术", "生物工程", "动画", "美术学", "艺术设计", "体育教育"],
        2015: ["生物工程", "美术学", "生物科学", "应用物理学", "应用心理学", "法学", "音乐表演"],
        2016: ["应用心理学", "化学", "音乐表演", "生物技术", "生物科学", "美术学"],
        2017: ["历史学", "音乐表演", "生物技术", "法学", "美术学", "生物工程"],
        2018: ["绘画", "化学", "美术学", "音乐表演", "法学", "历史学"],
        2019: ["绘画", "历史学", "应用心理学", "音乐表演", "化学", "法学"],
    }


def curated_rows() -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []

    def add_many(
        year: int,
        level: str,
        risk: str,
        majors: list[str],
        source_ids: list[str],
        evidence_type: str = "explicit_list",
        confidence: str = "high",
        notes: str = "",
    ) -> None:
        evidence_text = f"{year}年{level}就业{risk_name(risk)}专业包括：" + "、".join(majors) + "。"
        for major in majors:
            rows.append(
                {
                    "report_year": year,
                    "graduate_cohort": year - 1,
                    "education_level": level,
                    "risk_level": risk,
                    "reported_major_name": major,
                    "source_ids": ";".join(source_ids),
                    "evidence_type": evidence_type,
                    "evidence_text": evidence_text,
                    "confidence": confidence,
                    "notes": notes,
                }
            )

    for year, majors in red_ug_2010_2019().items():
        source = "eol_2020_undergrad_red_2010_2019"
        if year == 2011:
            source = "edu_cn_2011_red"
        elif year == 2012:
            source = "sina_2012_red_green;edu_cn_2012_red_green;eol_2020_undergrad_red_2010_2019"
        elif year == 2013:
            source = "edu_cn_2013_undergrad_yellow_green;eol_2020_undergrad_red_2010_2019"
        elif year == 2014:
            source = "haedu_2014_red_yellow_green;eol_2020_undergrad_red_2010_2019"
        elif year == 2015:
            source = "edu_cn_2015_red_yellow_green;eol_2020_undergrad_red_2010_2019"
        elif year == 2016:
            source = "sina_2016_red_yellow_green;eol_2020_undergrad_red_2010_2019"
        add_many(year, "本科", "red", majors, source.split(";"))

    highvoc_red = {
        2010: ["临床医学", "法律文秘", "计算机科学与技术", "国际金融", "工商管理", "经济管理", "法律事务", "汉语言文学教育", "计算机应用技术", "电子商务"],
        2011: ["临床医学", "法律文秘", "计算机科学与技术", "国际金融", "工商管理", "法律事务", "汉语言文学教育", "计算机应用技术", "电子商务"],
        2012: ["临床医学", "法律文秘", "计算机科学与技术", "国际金融", "工商管理", "法律事务", "汉语言文学教育", "计算机应用技术", "电子商务"],
        2013: ["法律文秘", "计算机科学与技术", "国际金融", "工商管理", "法律事务", "汉语言文学教育", "计算机应用技术", "电子商务"],
        2014: ["法律事务", "语文教育", "电子商务", "会计电算化", "生物技术及应用", "工商企业管理", "计算机信息管理", "计算机应用技术"],
        2015: ["法律事务", "语文教育", "初等教育", "投资与理财", "应用日语", "国际金融"],
        2016: ["法律事务", "语文教育", "工程监理", "建筑工程管理", "税务"],
        2017: ["法律事务", "语文教育", "图形图像制作", "初等教育", "会计电算化"],
        2018: ["法律事务", "汉语", "食品营养与检测", "初等教育", "语文教育"],
        2019: ["语文教育", "英语教育", "法律事务", "汉语", "初等教育"],
        2020: ["法律事务", "语文教育", "烹调工艺与营养", "小学教育", "导游"],
        2021: ["法律事务", "烹调工艺与营养", "小学教育", "英语教育", "语文教育"],
        2022: ["数学教育", "小学教育", "英语教育", "语文教育", "法律事务"],
        2023: ["小学教育", "数学教育", "法律事务", "英语教育", "语文教育"],
        2024: ["法律事务", "美术教育", "小学教育", "英语教育", "语文教育"],
        2025: ["法律事务", "美术教育", "小学英语教育", "小学教育", "小学语文教育"],
    }
    for year, majors in highvoc_red.items():
        if year <= 2016:
            source_ids = ["sohu_2017_archive_highvoc_red_2010_2016"]
            if year == 2011:
                source_ids.append("edu_cn_2011_red")
            if year == 2012:
                source_ids.extend(["sina_2012_red_green", "edu_cn_2012_red_green"])
            if year == 2014:
                source_ids.append("haedu_2014_red_yellow_green")
            if year == 2015:
                source_ids.append("edu_cn_2015_red_yellow_green")
            if year == 2016:
                source_ids.append("sina_2016_red_yellow_green")
        elif year == 2017:
            source_ids = ["thepaper_2017_red_green"]
        elif year == 2018:
            source_ids = ["sina_2018_red"]
        elif year == 2019:
            source_ids = ["people_2019_red_green"]
        elif year == 2020:
            source_ids = ["sdwm_2020_red_green"]
        elif year == 2021:
            source_ids = ["sohu_2025_highvoc_red", "zjnu_2023_red_green"]
        elif year == 2022:
            source_ids = ["sohu_myc_2022_red_green"]
        elif year == 2023:
            source_ids = ["zjnu_2023_red_green"]
        elif year == 2024:
            source_ids = ["labor_2024_red_green", "sohu_2025_highvoc_red"]
        else:
            source_ids = ["sohu_2025_highvoc_red"]
        evidence_type = "derived_from_consecutive_count" if year == 2021 else "explicit_list"
        confidence = "medium" if year == 2021 else "high"
        notes = "第五个专业由2025近五年说明与2023连续三届说明交叉推定。" if year == 2021 else ""
        add_many(year, "高职高专", "red", majors, source_ids, evidence_type, confidence, notes)

    add_many(2009, "本科", "red", ["法学", "计算机科学与技术", "英语", "国际经济与贸易", "工商管理", "汉语言文学", "电子信息工程", "会计学", "中医学", "音乐表演", "运动训练"], ["sczk_2017_archive_red_yellow_green"], "secondary_archive_table", "medium", "四川招考网2017归档表格列出的2009年大学红、黄、绿牌专业。")
    add_many(2009, "本科", "yellow", ["临床医学", "旅游管理", "公共事业管理", "信息与计算科学", "数学与应用数学"], ["sczk_2017_archive_red_yellow_green"], "secondary_archive_table", "medium", "四川招考网2017归档表格列出的2009年大学红、黄、绿牌专业。")
    add_many(2009, "本科", "green", ["石油工程", "材料物理", "金属材料工程", "工程力学", "地理科学"], ["sczk_2017_archive_red_yellow_green"], "secondary_archive_table", "medium", "四川招考网2017归档表格列出的2009年大学红、黄、绿牌专业。")

    add_many(2010, "本科", "yellow", ["统计学", "艺术设计", "美术学", "电子信息科学与技术", "公共事业管理", "信息管理与信息系统", "汉语言文学", "工商管理"], ["sczk_2017_archive_red_yellow_green"], "secondary_archive_table", "medium")
    add_many(2010, "高职高专", "yellow", ["计算机网络技术", "计算机信息管理", "物流管理", "商务英语", "会计电算化"], ["sczk_2017_archive_red_yellow_green"], "secondary_archive_table", "medium")
    add_many(2010, "本科", "green", ["地质工程", "港口航道与海岸工程", "船舶与海洋工程", "石油工程", "采矿工程", "油气储运工程", "矿物加工工程", "过程装备与控制工程", "水文与水资源工程"], ["sczk_2017_archive_red_yellow_green"], "secondary_archive_table", "medium")
    add_many(2010, "高职高专", "green", ["道路桥梁工程技术", "生产过程自动化技术", "应用化工技术", "焊接技术及自动化", "楼宇智能化工程技术"], ["sczk_2017_archive_red_yellow_green"], "secondary_archive_table", "medium")

    add_many(2011, "本科", "yellow", ["工商管理", "艺术设计", "美术学", "电子信息科学与技术", "公共事业管理", "信息管理与信息系统", "汉语言文学"], ["sczk_2017_archive_red_yellow_green"], "secondary_archive_table", "medium")
    add_many(2011, "高职高专", "yellow", ["计算机网络技术", "计算机信息管理", "物流管理", "商务英语", "会计电算化", "经济管理"], ["sczk_2017_archive_red_yellow_green"], "secondary_archive_table", "medium")
    add_many(2011, "本科", "green", ["地质工程", "港口航道与海岸工程", "船舶与海洋工程", "石油工程", "采矿工程", "油气储运工程", "矿物加工工程", "过程装备与控制工程", "水文与水资源工程", "审计学"], ["edu_cn_2011_red"])
    add_many(2011, "高职高专", "green", ["道路桥梁工程技术", "生产过程自动化技术", "应用化工技术", "焊接技术及自动化", "楼宇智能化工程技术", "供热通风与空调工程技术"], ["edu_cn_2011_red"])

    add_many(2012, "本科", "yellow", ["计算机科学与技术", "艺术设计", "美术学", "电子信息科学与技术", "公共事业管理", "信息管理与信息系统", "工商管理", "汉语言文学"], ["gkzxw_2012_undergrad_yellow"], "explicit_named_in_partial_list", "medium", "原文以“等”结尾；仅保留公开文本明确列出的专业。")
    add_many(2012, "高职高专", "yellow", ["计算机网络技术", "计算机信息管理"], ["sczk_2017_archive_red_yellow_green"], "secondary_archive_table", "medium")
    add_many(2012, "高职高专", "yellow", ["高护", "机电一体化", "汽修"], ["sczk_2017_archive_red_yellow_green"], "secondary_archive_table_abbreviation", "medium", "原表使用简称；已按高职目录别名映射到当前专业名称和代码，仍保留源文本名称供复核。")
    add_many(2012, "高职高专", "yellow", ["建筑类", "水利类"], ["sczk_2017_archive_red_yellow_green"], "secondary_archive_table_category", "low", "原表使用宽泛类别，不能可靠映射到单个专业，保留源文本名称并标记人工复核。")
    add_many(2012, "本科", "green", ["地质工程", "港口航道与海岸工程", "船舶与海洋工程", "石油工程", "采矿工程", "油气储运工程", "矿物加工工程", "过程装备与控制工程", "水文与水资源工程", "审计学"], ["edu_cn_2012_red_green"])
    add_many(2012, "高职高专", "green", ["道路桥梁工程技术", "生产过程自动化技术", "应用化工技术", "焊接技术及自动化", "楼宇智能化工程技术", "供热通风与空调工程技术"], ["edu_cn_2012_red_green"])

    add_many(2013, "本科", "yellow", ["计算机科学与技术", "艺术设计", "电子信息科学与技术", "公共事业管理", "信息管理与信息系统", "工商管理", "汉语言文学", "国际经济与贸易"], ["edu_cn_2013_undergrad_yellow_green"])
    add_many(2013, "本科", "green", ["地质工程", "港口航道与海岸工程", "船舶与海洋工程", "石油工程", "采矿工程", "油气储运工程", "矿物加工工程", "过程装备与控制工程", "水文与水资源工程", "审计学"], ["edu_cn_2013_undergrad_yellow_green"])
    add_many(2013, "高职高专", "yellow", ["计算机网络技术", "计算机信息管理", "物流管理", "商务英语", "临床医学"], ["sczk_2017_archive_red_yellow_green"], "secondary_archive_table", "medium")
    add_many(2013, "高职高专", "green", ["道路桥梁工程技术", "生产过程自动化技术", "应用化工技术", "焊接技术及自动化", "供热通风与空调工程技术"], ["sczk_2017_archive_red_yellow_green"], "secondary_archive_table", "medium")

    add_many(2014, "本科", "yellow", ["数学与应用数学", "电子信息科学与技术", "公共事业管理", "汉语言文学", "英语", "工商管理", "国际经济与贸易"], ["haedu_2014_red_yellow_green"])
    add_many(2014, "高职高专", "yellow", ["人力资源管理", "国际金融", "商务英语", "计算机网络技术", "物流管理"], ["haedu_2014_red_yellow_green"])
    add_many(2014, "本科", "green", ["建筑学", "地质工程", "矿物加工工程", "采矿工程", "油气储运工程", "车辆工程", "城市规划", "船舶与海洋工程", "审计学"], ["haedu_2014_red_yellow_green"])
    add_many(2014, "高职高专", "green", ["电气化铁道技术", "供热通风与空调工程技术", "铁道工程技术", "楼宇智能化工程技术", "石油化工生产技术", "道路桥梁工程技术"], ["haedu_2014_red_yellow_green"])

    add_many(2015, "本科", "yellow", ["体育教育", "动画", "英语", "工商管理", "汉语言文学"], ["edu_cn_2015_red_yellow_green"])
    add_many(2015, "高职高专", "yellow", ["会计电算化", "工商企业管理", "计算机多媒体技术", "计算机应用技术"], ["edu_cn_2015_red_yellow_green"])
    add_many(2015, "本科", "green", ["建筑学", "软件工程", "网络工程", "通信工程", "建筑环境与设备工程", "车辆工程", "矿物加工工程"], ["edu_cn_2015_red_yellow_green"])
    add_many(2015, "高职高专", "green", ["铁道工程技术", "电气化铁道技术", "石油化工生产技术", "电力系统自动化技术", "供用电技术", "楼宇智能化工程技术"], ["edu_cn_2015_red_yellow_green"])

    add_many(2016, "本科", "yellow", ["生物工程", "动画", "艺术设计", "法学", "应用物理学"], ["sina_2016_red_yellow_green"])
    add_many(2016, "高职高专", "yellow", ["会计电算化", "图形图像制作", "影视动画", "应用日语"], ["sina_2016_red_yellow_green"])
    add_many(2016, "本科", "green", ["软件工程", "网络工程", "通信工程", "电气工程及其自动化", "审计学", "广告学", "车辆工程"], ["sina_2016_red_yellow_green"])
    add_many(2016, "高职高专", "green", ["铁道工程技术", "电力系统自动化技术", "市场营销", "房地产经营与估价", "发电厂及电力系统", "视觉传达"], ["sina_2016_red_yellow_green"])

    add_many(2017, "本科", "green", ["信息安全", "软件工程", "网络工程", "数字媒体艺术", "通信工程", "电气工程及其自动化", "广告学"], ["thepaper_2017_red_green"])
    add_many(2017, "高职高专", "green", ["市场营销", "电气化铁道技术", "电力系统自动化技术", "软件技术", "视觉传达", "发电厂及电力系统"], ["thepaper_2017_red_green"])
    add_many(2017, "本科", "yellow", ["化学", "应用心理学", "临床医学", "广播电视编导", "生物科学"], ["sina_2017_undergrad_yellow", "sczk_2017_archive_red_yellow_green"], "explicit_list_secondary", "medium", "新浪教育列出五个本科黄牌专业，四川招考网归档表格交叉列名。")
    add_many(2017, "高职高专", "yellow", ["财务管理", "建筑工程管理", "食品营养与检测", "影视动画"], ["sczk_2017_archive_red_yellow_green"], "secondary_archive_table", "medium")

    add_many(2018, "本科", "yellow", ["生物技术", "生物工程", "应用心理学", "生物科学"], ["whysw_2018_yellow"], "explicit_list_secondary", "medium", "文化艺术网转载文本明确列名；公开抓取未找到更权威的2018黄牌全文页。")
    add_many(2018, "高职高专", "yellow", ["图形图像制作", "会计电算化", "财务管理"], ["whysw_2018_yellow"], "explicit_list_secondary", "medium", "文化艺术网转载文本明确列名；公开抓取未找到更权威的2018黄牌全文页。")
    add_many(2018, "本科", "green", ["信息安全", "软件工程", "网络工程", "物联网工程", "数字媒体技术", "通信工程", "数字媒体艺术"], ["sina_2018_red"])
    add_many(2018, "高职高专", "green", ["社会体育", "市场营销", "信息安全技术", "软件技术", "电气化铁道技术", "电力系统自动化技术"], ["sina_2018_red"])

    # Red, yellow, and green records with explicit public text from 2019 onward.
    add_many(2019, "本科", "green", ["信息安全", "软件工程", "网络工程", "物联网工程", "数字媒体技术", "通信工程"], ["people_2019_red_green", "bluebook_pdf_2023_undergrad"])
    add_many(2019, "本科", "yellow", ["心理学", "美术学", "生物技术", "生物工程", "应用物理学"], ["ms315_2019_undergrad_yellow"], "explicit_list_secondary", "medium", "中国美术高考网转载人民日报文本明确列名；保留为2019本科黄牌补齐来源。")
    add_many(2019, "高职高专", "green", ["电气化铁道技术", "社会体育", "软件技术", "电力系统自动化技术", "发电厂及电力系统", "道路桥梁工程技术"], ["people_2019_red_green"])

    add_many(2020, "本科", "green", ["信息安全", "软件工程", "信息工程", "网络工程", "计算机科学与技术", "数字媒体艺术", "电气工程及其自动化"], ["sdwm_2020_red_green"])
    add_many(2020, "本科", "red", ["绘画", "音乐表演", "法学", "应用心理学", "化学"], ["sdwm_2020_red_green"])
    add_many(2020, "高职高专", "green", ["铁道机车", "铁道工程技术", "社会体育", "电力系统继电保护与自动化技术", "移动互联应用技术", "发电厂及电力系统", "物联网应用技术"], ["sdwm_2020_red_green"])

    add_many(2021, "本科", "green", ["信息安全", "软件工程", "信息工程", "网络工程", "数字媒体技术", "电气工程及其自动化", "数字媒体艺术"], ["bluebook_pdf_2023_undergrad"])
    add_many(2021, "本科", "red", ["应用心理学", "历史学", "音乐表演", "绘画", "法学"], ["bluebook_pdf_2023_undergrad"])
    add_many(2021, "高职高专", "green", ["铁道机车", "铁道供电技术", "铁道工程技术", "社会体育"], ["ycwb_2021_highvoc_green"])

    add_many(2022, "本科", "green", ["信息安全", "网络工程", "信息工程", "微电子科学与工程", "数字媒体技术", "能源与动力工程"], ["sohu_myc_2022_red_green", "bluebook_pdf_2023_undergrad"])
    add_many(2022, "本科", "red", ["汉语国际教育", "绘画", "应用心理学", "音乐表演", "法学"], ["sohu_myc_2022_red_green"])
    add_many(2022, "高职高专", "green", ["铁道机车", "铁道工程技术", "铁道供电技术", "社会体育", "发电厂及电力系统", "道路桥梁工程技术"], ["sohu_myc_2022_red_green"])

    add_many(2023, "本科", "green", ["信息工程", "微电子科学与工程", "电气工程及其自动化", "能源与动力工程", "道路桥梁与渡河工程", "机械电子工程"], ["zjnu_2023_red_green", "bluebook_pdf_2023_undergrad"])
    add_many(2023, "本科", "yellow", ["英语", "美术学", "翻译", "音乐表演"], ["bluebook_pdf_2023_undergrad"])
    add_many(2023, "本科", "red", ["汉语国际教育", "法学", "教育技术学", "绘画", "应用心理学"], ["zjnu_2023_red_green", "bluebook_pdf_2023_undergrad"])
    add_many(2023, "高职高专", "green", ["铁道机车", "铁道工程技术", "石油化工技术", "发电厂及电力系统", "应用化工技术", "道路桥梁工程技术"], ["zjnu_2023_red_green"])

    add_many(2024, "本科", "green", ["微电子科学与工程", "电气工程及其自动化", "新能源科学与工程", "能源与动力工程", "机械电子工程", "机器人工程"], ["labor_2024_red_green"])
    add_many(2024, "本科", "yellow", ["公共事业管理", "教育技术学", "生物技术", "汉语国际教育"], ["china_2025_undergrad_yellow"], "explicit_list_secondary", "medium", "中华网教育页面明确列出2024本科黄牌专业和2010-2024累计黄牌次数；用于替代Scribd受阻文本页。")
    add_many(2024, "本科", "red", ["音乐表演", "绘画", "美术学", "应用心理学", "法学"], ["labor_2024_red_green", "eol_2025_undergrad_red"])
    add_many(2024, "高职高专", "green", ["新能源汽车技术", "智能控制技术", "电气自动化技术", "应用化工技术", "石油化工技术", "铁道机车"], ["labor_2024_red_green"])

    add_many(2025, "本科", "green", ["电气工程及其自动化", "微电子科学与工程", "机械电子工程", "新能源科学与工程", "车辆工程", "机器人工程"], ["people_2025_green", "people_2025_green_income", "gkztc_2025_undergrad_yellow"])
    add_many(2025, "本科", "yellow", ["投资学", "应用心理学", "广播电视学", "学前教育"], ["gkztc_2025_undergrad_yellow"])
    add_many(2025, "本科", "red", ["公共事业管理", "音乐表演", "绘画", "法学", "美术学"], ["eol_2025_undergrad_red", "gkztc_2025_undergrad_yellow"])
    add_many(2025, "高职高专", "green", ["铁道机车运用与维护", "电气自动化技术", "应用化工技术", "工业机器人技术", "新能源汽车技术", "智能控制技术"], ["people_2025_green", "people_2025_green_income"])

    add_many(2026, "本科", "green", ["电气工程及其自动化", "微电子科学与工程", "自动化", "能源与动力工程", "车辆工程", "新能源科学与工程"], ["sina_2026_undergrad_green", "gkztc_2026_green"])
    add_many(2026, "高职高专", "green", ["铁道机车运用与维护", "电气自动化技术", "铁道工程技术", "发电厂及电力系统", "应用化工技术", "工业过程自动化技术"], ["gkztc_2026_green"])

    return rows


STANDARD_NAME_ALIASES = {
    "生物科学与工程": "生物科学与工程",
    "艺术设计": "艺术设计学",
    "社会体育": "社会体育指导与管理",
    "地理信息系统": "地理信息科学",
    "广播电视新闻学": "广播电视学",
    "临床医学（拔尖卓越培养）": "临床医学",
    "文秘教育": "秘书学",
    "信息网络": "网络工程",
    "汉语国际教育": "国际中文教育",
    "英语教育": "小学英语教育",
    "语文教育": "小学语文教育",
    "汉语言文学教育": "小学语文教育",
    "汉语言文学（师、非师）": "汉语言文学",
    "美术学（师范）": "美术学",
    "美术学（中国书画方向）（师范）": "美术学",
    "历史学（师范）": "历史学",
    "物理学（师范）": "物理学",
    "体育教育（师范）": "体育教育",
    "英语（师范）": "英语",
    "英 语（师范）": "英语",
    "数学与应用数学（师范）": "数学与应用数学",
    "教育技术学（师范）": "教育技术学",
    "小学教育（师范）": "小学教育",
    "学前教育（师范）": "学前教育",
    "视觉传达设计（中韩2+2）": "视觉传达设计",
    "酒店管理（中爱3+1）": "酒店管理",
    "初等教育": "小学教育",
    "电气化铁道技术": "铁道供电技术",
    "电力系统自动化技术": "电力系统自动化技术",
    "电力系统继电保护与自动化技术": "电力系统继电保护技术",
    "铁道机车": "铁道机车运用与维护",
    "铁道供电技术": "铁道供电技术",
    "计算机多媒体技术": "数字媒体技术",
    "会计电算化": "大数据与会计",
    "建筑工程管理": "建设工程管理",
    "工程监理": "建设工程监理",
    "烹调工艺与营养": "烹饪工艺与营养",
    "生物技术及应用": "生物技术",
    "城市规划": "城乡规划",
    "建筑环境与设备工程": "建筑环境与能源应用工程",
    "电气工程及自动化": "电气工程及其自动化",
    "生产过程自动化技术": "工业过程自动化技术",
    "焊接技术及自动化": "智能焊接技术",
    "楼宇智能化工程技术": "建筑智能化工程技术",
    "石油化工生产技术": "石油化工技术",
    "房地产经营与估价": "房地产经营与管理",
    "视觉传达": "视觉传达设计",
    "信息安全技术": "信息安全技术应用",
    "云计算技术与应用": "云计算技术应用",
    "大数据技术与应用": "大数据技术",
    "信息安全与管理": "信息安全技术应用",
    "机械制造与自动化": "机械制造及自动化",
    "焊接技术与自动化": "智能焊接技术",
    "道路桥梁工程技术": "道路与桥梁工程技术",
    "数学教育": "小学数学教育",
    "汉语": "中文",
    "税务": "财税大数据应用",
    "食品营养与检测": "食品检验检测技术",
    "投资与理财": "金融服务与管理",
    "物联网": "物联网工程",
    "视觉传达设计（中外合作）": "视觉传达设计",
    "临床医学（七年制）": "临床医学",
    "口腔医学（七年制）": "口腔医学",
    "电子商务（工学）": "电子商务",
    "数据科学与大数据技术（理学）": "数据科学与大数据技术",
    "软件工程（合作办学）": "软件工程",
    "徳语": "德语",
    "物联网工程（腾讯云精英班）": "物联网工程",
    "物联网工程（厚溥智能信息特色班）": "物联网工程",
    "通信工程（讯方5G特色班）": "通信工程",
    "通信工程（5G特色班）": "通信工程",
    "自动化（智能制造）": "自动化",
    "通信工程（讯方技术特色班）": "通信工程",
    "软件工程（中软国际特色班）": "软件工程",
    "计算机科学与技术（中软国际特色班）": "计算机科学与技术",
    "电子信息工程（东软大数据特色班）": "电子信息工程",
    "金融学（金融智能创新实验班）": "金融学",
    "计算机科学与技术(中软国际特色班)": "计算机科学与技术",
    "计算机科学与技术(华胜天成大数据方向)": "计算机科学与技术",
    "计算机科学与技术（华胜天成大数据方向）": "计算机科学与技术",
    "软件工程(中软国际特色班)": "软件工程",
    "通信工程(讯方技术特色班)": "通信工程",
    "电子信息工程(东软大数据特色班)": "电子信息工程",
    "机械电子工程(先进制造与工业机器人方向)": "机械电子工程",
    "金融学(金融智能创新实验班)": "金融学",
    "土木工程(智能建筑信息化方向)": "土木工程",
}

LEVEL_STANDARD_NAME_ALIASES = {
    ("财务管理", "高职高专"): "大数据与财务管理",
    ("机电一体化", "高职高专"): "机电一体化技术",
    ("高护", "高职高专"): "护理",
    ("汽修", "高职高专"): "汽车检测与维修技术",
}

CATALOG_LOOKUP_ALIASES = {
    "国际中文教育": "汉语国际教育",
}


def risk_name(value: str) -> str:
    return {"red": "红牌", "yellow": "黄牌", "green": "绿牌"}.get(value, value)


def safe_name(value: str) -> str:
    return re.sub(r"[^A-Za-z0-9_.-]+", "_", value).strip("._") or "source"


SOURCE_REFERERS = {
    "chu_2023_major_setting_pdf": "https://www.chu.edu.cn/xxgk/2023/1025/c5128a168196/page.htm",
    "chu_2024_major_setting_pdf": "https://www.chu.edu.cn/xxgk/2024/1016/c5128a186150/page.htm",
    "chu_2025_major_setting_pdf": "https://www.chu.edu.cn/xxgk/2025/0603/c5128a194985/page.htm",
    "cqnu_2023_major_setting_docx": "https://xxgk.cqnu.edu.cn/info/1229/4123.htm",
    "hgnu_2025_stop_enrollment_xls": "https://xxgk.hgnu.edu.cn/?Article57/Article50/Article91/Article167/677.html",
    "huat_2025_stop_enrollment_xlsx": "https://xxgk.huat.edu.cn/info/1048/2114.htm",
    "jci_2024_academic_year_major_setting_docx": "https://jwc.jci.edu.cn/info/1132/5089.htm",
    "jci_undergrad_major_setting_docx": "https://jwc.jci.edu.cn/info/1132/4429.htm",
    "jiangnan_2025_major_setting_pdf": "https://xxgk.jiangnan.edu.cn/info/1007/1108.htm",
    "szu_2025_major_setting_docx": "https://xxgk.szu.edu.cn/jxzl1/zysz_dnxzzy_tzzymd.htm",
}


def same_origin_referer(url: str) -> str:
    parsed = urlparse(url)
    if not parsed.scheme or not parsed.netloc:
        return ""
    return f"{parsed.scheme}://{parsed.netloc}/"


def request_headers_for(source: SourceSeed) -> dict[str, str]:
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/125 Safari/537.36",
        "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
    }
    referer = SOURCE_REFERERS.get(source.source_id) or same_origin_referer(source.url)
    if referer:
        headers["Referer"] = referer
    return headers


def office_zip_suffix(content: bytes) -> str:
    with tempfile.NamedTemporaryFile(suffix=".zip", delete=False) as tmp:
        tmp.write(content)
        tmp_path = Path(tmp.name)
    try:
        with zipfile.ZipFile(tmp_path) as zf:
            names = set(zf.namelist())
        if "word/document.xml" in names:
            return ".docx"
        if "xl/workbook.xml" in names:
            return ".xlsx"
        if "ppt/presentation.xml" in names:
            return ".pptx"
    except zipfile.BadZipFile:
        return ""
    finally:
        tmp_path.unlink(missing_ok=True)
    return ""


def response_suffix(url: str, content_type: str, content: bytes) -> str:
    if content.startswith(b"%PDF"):
        return ".pdf"
    if content.startswith(b"\xff\xd8\xff"):
        return ".jpg"
    if content.startswith(b"\x89PNG\r\n\x1a\n"):
        return ".png"
    if content.startswith(b"RIFF") and content[8:12] == b"WEBP":
        return ".webp"
    if content.startswith(b"PK"):
        detected = office_zip_suffix(content)
        if detected:
            return detected
    if content.startswith(b"\xd0\xcf\x11\xe0"):
        return ".doc"
    content_type = (content_type or "").lower()
    if "pdf" in content_type:
        return ".pdf"
    if "officedocument.wordprocessingml.document" in content_type:
        return ".docx"
    if "msword" in content_type:
        return ".doc"
    if "officedocument.spreadsheetml.sheet" in content_type:
        return ".xlsx"
    if "ms-excel" in content_type or "vnd.ms-excel" in content_type:
        return ".xls"
    if "image/" in content_type:
        image_type = content_type.split("image/", 1)[1].split(";", 1)[0].strip().lower()
        return ".jpg" if image_type == "jpeg" else f".{image_type}"
    suffix = Path(urlparse(url).path).suffix.lower()
    if suffix in {".pdf", ".html", ".htm", ".shtml", ".txt", ".doc", ".docx", ".xls", ".xlsx"}:
        return suffix
    return ".html" if "html" in content_type or not suffix else suffix


def curl_fetch_binary(url: str, referer: str = "") -> bytes:
    command = [
        "curl.exe",
        "-L",
        "-k",
        "--retry",
        "3",
        "--connect-timeout",
        "20",
        "-A",
        "Mozilla/5.0",
    ]
    if referer:
        command.extend(["-e", referer])
    command.append(url)
    result = subprocess.run(
        command,
        cwd=str(ROOT),
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        timeout=120,
    )
    return result.stdout


def decode_text(content: bytes, content_type: str) -> str:
    header_match = re.search(r"charset=([\w.-]+)", content_type or "", flags=re.I)
    head = content[:4096].decode("ascii", errors="ignore")
    meta_match = re.search(r"charset=[\"']?\s*([\w.-]+)", head, flags=re.I)
    encodings = [
        meta_match.group(1) if meta_match else "",
        header_match.group(1) if header_match and header_match.group(1).lower() not in {"iso-8859-1", "latin1"} else "",
        "utf-8",
        "gb18030",
    ]
    seen: set[str] = set()
    for encoding in encodings:
        if not encoding:
            continue
        normalized = encoding.lower()
        if normalized in seen:
            continue
        seen.add(normalized)
        try:
            return content.decode(encoding)
        except (LookupError, UnicodeDecodeError):
            continue
    return content.decode("utf-8", errors="replace")


def html_to_text(html: str) -> str:
    soup = BeautifulSoup(html, "html.parser")
    for node in soup(["script", "style", "noscript"]):
        node.decompose()
    return re.sub(r"\n{3,}", "\n\n", soup.get_text("\n", strip=True))


def pdf_to_text(pdf_path: Path, txt_path: Path) -> str:
    try:
        subprocess.run(
            ["pdftotext", "-layout", "-enc", "UTF-8", str(pdf_path), str(txt_path)],
            cwd=str(ROOT),
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=120,
        )
        return txt_path.read_text(encoding="utf-8", errors="replace")
    except (FileNotFoundError, subprocess.CalledProcessError, subprocess.TimeoutExpired):
        return ""


def docx_to_text(docx_path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        return ""
    try:
        document = Document(docx_path)
    except Exception:
        return ""
    parts: list[str] = []
    parts.extend(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " | ") for cell in row.cells]
            if any(cells):
                parts.append("\t".join(cells))
    return "\n".join(parts)


def fetch_sources(raw_dir: Path) -> list[dict[str, Any]]:
    raw_dir.mkdir(parents=True, exist_ok=True)
    session = requests.Session()
    session.headers.update(
        {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/125 Safari/537.36",
            "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
        }
    )
    source_rows: list[dict[str, Any]] = []
    for source in SOURCES:
        base_row = {
            "source_id": source.source_id,
            "title": source.title,
            "url": source.url,
            "publisher": source.publisher,
            "published_date": source.published_date,
            "source_type": source.source_type,
            "coverage_note": source.coverage_note,
            "fetched_at": CAPTURED_AT,
        }
        try:
            request_headers = request_headers_for(source)
            try:
                response = session.get(source.url, headers=request_headers, timeout=40)
            except requests.exceptions.SSLError:
                response = session.get(source.url, headers=request_headers, timeout=40, verify=False)
            if response.status_code in {403, 405}:
                try:
                    response = requests.get(source.url, headers=request_headers, timeout=40)
                except requests.exceptions.SSLError:
                    response = requests.get(source.url, headers=request_headers, timeout=40, verify=False)
            content = response.content
            content_type = response.headers.get("content-type", "")
            suffix = response_suffix(response.url, content_type, content)
            if source.source_type.endswith("_pdf") and suffix != ".pdf":
                try:
                    fallback_content = curl_fetch_binary(source.url, SOURCE_REFERERS.get(source.source_id, ""))
                except Exception:
                    fallback_content = b""
                if fallback_content.startswith(b"%PDF"):
                    content = fallback_content
                    content_type = "application/pdf"
                    suffix = ".pdf"
            if source.source_type.endswith("_pdf") and suffix != ".pdf":
                cached_raw_path = raw_dir / f"{source.source_id}.pdf"
                cached_text_path = raw_dir / f"{source.source_id}.txt"
                if cached_raw_path.exists() and cached_text_path.exists() and cached_text_path.stat().st_size > 500:
                    cached_content = cached_raw_path.read_bytes()
                    source_rows.append(
                        {
                            **base_row,
                            "status": "cached_after_fetch_error",
                            "http_status": response.status_code,
                            "content_type": content_type,
                            "content_length": cached_raw_path.stat().st_size,
                            "sha256": hashlib.sha256(cached_content).hexdigest(),
                            "raw_path": str(cached_raw_path.relative_to(ROOT)),
                            "text_path": str(cached_text_path.relative_to(ROOT)),
                            "error": f"Fetched non-PDF response with suffix {suffix}; reused cached PDF/text.",
                        }
                    )
                    continue
            if suffix == ".doc" and (
                source.source_type.endswith("_xls")
                or "excel" in content_type.lower()
                or urlparse(response.url).path.lower().endswith(".xls")
            ):
                suffix = ".xls"
            if source.source_type.endswith("_xlsx") and suffix not in {".xls", ".xlsx"}:
                fallback_content = curl_fetch_binary(source.url, SOURCE_REFERERS.get(source.source_id, ""))
                if fallback_content.startswith(b"PK") or fallback_content.startswith(b"\xd0\xcf\x11\xe0"):
                    content = fallback_content
                    content_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    suffix = ".xlsx"
            if (source.source_type.endswith("_docx") or source.source_type.endswith("_doc")) and suffix not in {".doc", ".docx"}:
                fallback_content = curl_fetch_binary(source.url, SOURCE_REFERERS.get(source.source_id, ""))
                if fallback_content.startswith(b"PK"):
                    detected_suffix = office_zip_suffix(fallback_content)
                    if detected_suffix in {".docx", ".xlsx", ".pptx"}:
                        content = fallback_content
                        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        suffix = ".docx"
                elif fallback_content.startswith(b"\xd0\xcf\x11\xe0"):
                    content = fallback_content
                    content_type = "application/msword"
                    suffix = ".doc"
            digest = hashlib.sha256(content).hexdigest()
            raw_path = raw_dir / f"{source.source_id}{suffix}"
            raw_path.write_bytes(content)
            text_path = raw_dir / f"{source.source_id}.txt"
            if suffix == ".pdf":
                text = pdf_to_text(raw_path, text_path)
                if not text:
                    text_path.write_text("", encoding="utf-8")
            elif suffix == ".docx":
                text = docx_to_text(raw_path)
                text_path.write_text(text, encoding="utf-8")
            elif suffix == ".doc":
                text = ""
                try:
                    with tempfile.TemporaryDirectory() as tmpdir:
                        converted_path = Path(tmpdir) / f"{source.source_id}.docx"
                        convert_doc_with_word(raw_path, converted_path)
                        text = docx_to_text(converted_path)
                except Exception:
                    text = ""
                text_path.write_text(text, encoding="utf-8")
            elif suffix in {".xls", ".xlsx"}:
                text = ""
                text_path.write_text("", encoding="utf-8")
            elif suffix in {".png", ".jpg", ".jpeg", ".gif", ".webp"}:
                text = ""
                text_path.write_text("", encoding="utf-8")
            else:
                html_or_text = decode_text(content, response.headers.get("content-type", ""))
                text = html_to_text(html_or_text) if "<html" in html_or_text[:1000].lower() else html_or_text
                text_path.write_text(text, encoding="utf-8")
            status = "fetched"
            error = ""
            if "Client Challenge" in text or "required part of this site" in text:
                status = "blocked_client_challenge"
                error = "fetched challenge page; no usable article text"
            base_row.update(
                {
                    "status": status,
                    "http_status": response.status_code,
                    "content_type": content_type,
                    "content_length": len(content),
                    "sha256": digest,
                    "raw_path": str(raw_path.relative_to(ROOT)),
                    "text_path": str(text_path.relative_to(ROOT)),
                    "error": error,
                }
            )
        except Exception as exc:  # noqa: BLE001 - source audit should not stop the dataset build.
            cached_text_path = raw_dir / f"{source.source_id}.txt"
            cached_raw_paths = sorted(
                path
                for path in raw_dir.glob(f"{source.source_id}.*")
                if path.is_file() and path.suffix.lower() != ".txt"
            )
            if cached_raw_paths and cached_text_path.exists():
                cached_raw_path = cached_raw_paths[0]
                cached_content = cached_raw_path.read_bytes()
                base_row.update(
                    {
                        "status": "cached_after_fetch_error",
                        "http_status": "",
                        "content_type": "",
                        "content_length": len(cached_content),
                        "sha256": hashlib.sha256(cached_content).hexdigest(),
                        "raw_path": str(cached_raw_path.relative_to(ROOT)),
                        "text_path": str(cached_text_path.relative_to(ROOT)),
                        "error": f"fetch_error_reused_cached_raw: {repr(exc)}",
                    }
                )
            else:
                base_row.update(
                    {
                        "status": "error",
                        "http_status": "",
                        "content_type": "",
                        "content_length": "",
                        "sha256": "",
                        "raw_path": "",
                        "text_path": "",
                        "error": repr(exc),
                    }
                )
        source_rows.append(base_row)
    return source_rows


def read_catalog() -> dict[tuple[str, str], dict[str, str]]:
    if not CATALOG_PATH.exists():
        return {}
    catalog: dict[tuple[str, str], dict[str, str]] = {}
    with CATALOG_PATH.open("r", encoding="utf-8-sig", newline="") as handle:
        for row in csv.DictReader(handle):
            catalog[(row["major_name"], row["level"])] = row
            catalog.setdefault((row["major_name"], ""), row)
    return catalog


def pdf_table_noise(line: str) -> bool:
    compact = re.sub(r"\s+", "", line)
    if not compact:
        return True
    if compact.startswith("—") or compact.endswith("—"):
        return True
    noise_terms = {
        "序号",
        "学校名称",
        "专业名称",
        "专业代码",
        "备注",
        "学位授",
        "予门类",
        "修业",
        "年限",
        "附件",
        "教育部办公厅",
        "中华人民共和国教育部",
    }
    if any(term in compact for term in noise_terms):
        return True
    province_suffixes = ("省", "市", "自治区", "壮族自治区", "回族自治区", "维吾尔自治区", "新疆生产建设兵团")
    if compact.endswith(province_suffixes) and len(compact) <= 12:
        return True
    note_fragments = ("该校此专业", "尚有", "撤销合作办", "学专业", "已调整为", "点")
    return any(fragment == compact or compact.startswith(fragment) for fragment in note_fragments)


def nearby_table_fragment(lines: list[str], index: int, step: int) -> str:
    current = index + step
    while 0 <= current < len(lines):
        text = re.sub(r"\s+", "", lines[current])
        if not text:
            current += step
            continue
        if re.match(r"^\d{1,4}", text) or pdf_table_noise(text):
            return ""
        return text
    return ""


MOE_CANCEL_PDF_CONFIGS = [
    {
        "source_id": "moe_2022_2021_undergrad_setup_results_pdf",
        "policy_year": 2021,
        "expected_count": 804,
        "title_candidates": ["撤销专业名单"],
        "warning_label": "2021年度普通高等学校撤销专业名单",
        "notice_ref": "教育部教高函〔2021〕14号附件1",
    },
    {
        "source_id": "moe_2023_2022_undergrad_setup_results_pdf",
        "policy_year": 2022,
        "expected_count": 925,
        "title_candidates": ["撤销专业名单"],
        "warning_label": "2022年度普通高等学校撤销专业名单",
        "notice_ref": "教育部教高函〔2023〕3号附件1",
    },
    {
        "source_id": "moe_2024_2023_undergrad_cancel_notice_pdf",
        "policy_year": 2023,
        "expected_count": 1670,
        "title_candidates": ["2023年度普通高等学校撤销本科专业名单"],
        "warning_label": "2023年度普通高等学校撤销本科专业名单",
        "notice_ref": "教育部教高函〔2024〕7号附件",
    },
    {
        "source_id": "moe_2025_2024_undergrad_cancel_notice_pdf",
        "policy_year": 2024,
        "expected_count": 1428,
        "title_candidates": ["2024年度普通高等学校撤销本科专业名单"],
        "warning_label": "2024年度普通高等学校撤销本科专业名单",
        "notice_ref": "教育部教高函〔2025〕5号附件",
    },
]

MOE_CANCEL_XLS_CONFIGS = [
    {
        "source_id": "yingkou_2019_undergrad_setup_results_xls_mirror",
        "source_ids": [
            "moe_2020_2019_undergrad_setup_results_pdf",
            "moe_zwfw_2020_2019_undergrad_setup_notice_json",
            "yingkou_2019_undergrad_setup_results_xls_mirror",
        ],
        "policy_year": 2019,
        "expected_count": 367,
        "title_candidates": ["撤销本科专业名单", "撤销专业名单"],
        "warning_label": "2019年度普通高等学校撤销本科专业名单",
        "notice_ref": "教育部教高函〔2020〕2号官方正文、附件1扫描PDF及机器可读XLS镜像",
        "sheet_name": "Sheet1",
        "serial_col": 0,
        "school_col": 1,
        "major_col": 2,
        "code_col": 3,
        "degree_col": None,
        "duration_col": None,
        "note_col": None,
        "confidence": "medium",
    },
    {
        "source_id": "moe_2021_2020_undergrad_setup_results_xls",
        "policy_year": 2020,
        "expected_count": 518,
        "title_candidates": ["撤销本科专业名单", "撤销专业名单"],
        "warning_label": "2020年度普通高等学校撤销本科专业名单",
        "notice_ref": "教育部2020年度普通高等学校本科专业备案和审批结果附件1",
        "serial_col": 0,
        "school_col": 1,
        "major_col": 2,
        "code_col": 3,
        "degree_col": 4,
        "duration_col": 5,
        "note_col": 6,
    },
]

MOE_CANCEL_DOCX_CONFIGS = [
    {
        "source_id": "haedu_2011_undergrad_cancel_doc_mirror",
        "policy_year": 2011,
        "expected_count": 32,
        "warning_label": "2011年度同意撤销的高等学校本科专业名单",
        "notice_ref": "教育部2011年度高等学校本科专业设置备案或审批结果附件4",
        "raw_suffix": ".doc",
        "table_index": 0,
        "serial_col": 0,
        "school_col": 2,
        "code_col": 3,
        "major_col": 4,
        "duration_col": 5,
        "degree_col": None,
        "note_col": None,
    },
    {
        "source_id": "sdipct_2013_undergrad_setup_results_doc_mirror",
        "policy_year": 2013,
        "expected_count": 26,
        "warning_label": "2013年度普通高等学校撤销专业名单",
        "notice_ref": "教育部2013年度普通高等学校本科专业备案或审批结果附件",
        "raw_suffix": ".doc",
        "table_index": 2,
        "school_col": 0,
        "code_col": 1,
        "major_col": 2,
        "duration_col": 3,
        "degree_col": 4,
        "note_col": 5,
    },
    {
        "source_id": "moe_2015_2014_undergrad_setup_results_docx_mirror",
        "policy_year": 2014,
        "expected_count": 67,
        "warning_label": "2014年度普通高等学校撤销本科专业名单",
        "notice_ref": "教育部2014年度普通高等学校本科专业备案或审批结果附件",
        "table_index": 2,
        "school_col": 1,
        "major_col": 2,
        "code_col": 3,
        "degree_col": 4,
        "duration_col": 5,
        "note_col": None,
    },
    {
        "source_id": "moe_2016_2015_undergrad_setup_results_docx",
        "policy_year": 2015,
        "expected_count": 118,
        "warning_label": "2015年度普通高等学校撤销本科专业名单",
        "notice_ref": "教育部教高函〔2016〕2号附件",
        "table_index": 3,
        "school_col": 0,
        "major_col": 1,
        "code_col": 2,
        "degree_col": 3,
        "duration_col": 4,
        "note_col": 5,
    },
    {
        "source_id": "moe_2017_2016_undergrad_setup_results_docx",
        "policy_year": 2016,
        "expected_count": 149,
        "warning_label": "2016年度普通高等学校撤销本科专业名单",
        "notice_ref": "教育部教高函〔2017〕2号附件",
        "table_index": 3,
        "school_col": 0,
        "major_col": 1,
        "code_col": 2,
        "degree_col": 3,
        "duration_col": 4,
        "note_col": 5,
    },
    {
        "source_id": "moe_2018_2017_undergrad_setup_results_docx",
        "policy_year": 2017,
        "expected_count": 241,
        "warning_label": "2017年度普通高等学校撤销本科专业名单",
        "notice_ref": "教育部教高函〔2018〕4号附件",
        "table_index": 3,
        "school_col": 0,
        "major_col": 1,
        "code_col": 2,
        "degree_col": 3,
        "duration_col": 4,
        "note_col": 5,
    },
    {
        "source_id": "moe_2019_2018_undergrad_setup_results_docx",
        "policy_year": 2018,
        "expected_count": 416,
        "warning_label": "2018年度普通高等学校撤销本科专业名单",
        "notice_ref": "教育部教高函〔2019〕7号附件",
        "table_index": 3,
        "school_col": 0,
        "major_col": 1,
        "code_col": 2,
        "degree_col": 3,
        "duration_col": 4,
        "note_col": 5,
    },
    {
        "source_id": "beijing_2019_municipal_undergrad_setup_doc",
        "policy_year": 2019,
        "expected_count": 2,
        "warning_label": "2019年度市属普通高等学校撤销本科专业名单",
        "notice_ref": "北京市教委2019年度市属普通高等学校本科专业备案和审批结果附件",
        "raw_suffix": ".doc",
        "table_index": 2,
        "serial_col": 0,
        "school_col": 1,
        "major_col": 2,
        "code_col": 3,
        "degree_col": 4,
        "duration_col": 5,
        "note_col": 6,
        "complete_cancel_index": False,
    },
]

MOE_COMPLETE_CANCEL_SOURCE_IDS = {
    config["source_id"]
    for config in [*MOE_CANCEL_DOCX_CONFIGS, *MOE_CANCEL_PDF_CONFIGS, *MOE_CANCEL_XLS_CONFIGS]
    if config.get("complete_cancel_index", True)
}


def moe_cancel_row(
    config: dict[str, Any],
    serial_no: int,
    school: str,
    major: str,
    code: str,
    degree: str,
    duration: str,
    note: str = "",
) -> dict[str, Any]:
    detail_parts = [part for part in [code, degree, duration] if part]
    detail = f"（{'，'.join(detail_parts)}）" if detail_parts else ""
    evidence = f"{config['notice_ref']}第{serial_no}条：{school}撤销{major}{detail}。"
    if note:
        evidence += f"备注：{note}。"
    source_ids = ";".join(config.get("source_ids", [config["source_id"]]))
    return {
        "policy_year": config["policy_year"],
        "region": school,
        "education_level": "本科",
        "record_type": "major_cancel",
        "warning_label": config["warning_label"],
        "reported_major_name": major,
        "major_code": code,
        "source_row_no": str(serial_no),
        "study_duration": duration,
        "policy_action": "教育部公布/备案撤销。",
        "criterion_text": "年度普通高等学校本科专业设置和调整结果中的撤销专业名单。",
        "source_ids": source_ids,
        "evidence_text": evidence,
        "confidence": config.get("confidence", "high"),
    }


def parse_moe_pdf_cancel_rows(raw_dir: Path, config: dict[str, Any]) -> list[dict[str, Any]]:
    text_path = raw_dir / f"{config['source_id']}.txt"
    if not text_path.exists():
        return []
    lines = text_path.read_text(encoding="utf-8", errors="replace").splitlines()
    starts = [
        index
        for index, line in enumerate(lines)
        for title in config["title_candidates"]
        if title in line
    ]
    if not starts:
        return []
    start = max(starts)

    parsed: list[dict[str, Any]] = []
    expected = int(config["expected_count"])
    for index, line in enumerate(lines[start + 1 :], start + 1):
        stripped = line.strip()
        if not re.match(r"^\d{1,4}\s+", stripped):
            continue
        parts = stripped.split()
        if not parts or not parts[0].isdigit():
            continue
        serial_no = int(parts[0])
        if serial_no > expected:
            break
        school = ""
        major = ""
        code = ""
        degree = ""
        duration = ""
        note = ""
        if len(parts) >= 6:
            school, major, code, degree, duration = parts[1:6]
            note = "".join(parts[6:])
        elif len(parts) == 5:
            major, code, degree, duration = parts[1:5]
            school = nearby_table_fragment(lines, index, -1) + nearby_table_fragment(lines, index, 1)
        if not school or not major or not code or not degree or not duration:
            continue
        parsed.append(moe_cancel_row(config, serial_no, school, major, code, degree, duration, note))
        if serial_no == expected:
            break
    if len(parsed) != expected:
        raise ValueError(
            f"Parsed {len(parsed)} rows from {config['source_id']}, expected {expected}."
        )
    return parsed


def convert_xls_with_excel(xls_path: Path, xlsx_path: Path) -> None:
    try:
        import win32com.client as win32  # type: ignore[import-not-found]
    except ImportError as exc:
        raise RuntimeError("Parsing legacy .xls sources requires pywin32 and Microsoft Excel.") from exc
    excel = win32.DispatchEx("Excel.Application")
    excel.Visible = False
    excel.DisplayAlerts = False
    workbook = None
    try:
        workbook = excel.Workbooks.Open(str(xls_path.resolve()))
        workbook.SaveAs(str(xlsx_path.resolve()), FileFormat=51)
    finally:
        if workbook is not None:
            workbook.Close(False)
        excel.Quit()


def convert_doc_with_word(doc_path: Path, docx_path: Path) -> None:
    try:
        import win32com.client as win32  # type: ignore[import-not-found]
    except ImportError as exc:
        raise RuntimeError("Parsing legacy .doc sources requires pywin32 and Microsoft Word.") from exc
    word = win32.DispatchEx("Word.Application")
    word.Visible = False
    document = None
    try:
        document = word.Documents.Open(str(doc_path.resolve()))
        document.SaveAs2(str(docx_path.resolve()), FileFormat=16)
    finally:
        if document is not None:
            document.Close(False)
        word.Quit()


def normalize_source_major_code(code: str) -> str:
    code = str(code or "").strip()
    if re.fullmatch(r"\d{5}", code):
        return f"0{code}"
    return code


def parse_moe_xls_cancel_rows(raw_dir: Path, config: dict[str, Any]) -> list[dict[str, Any]]:
    raw_path = raw_dir / f"{config['source_id']}.xls"
    if not raw_path.exists():
        return []
    with tempfile.TemporaryDirectory() as tmpdir:
        converted_path = Path(tmpdir) / f"{config['source_id']}.xlsx"
        convert_xls_with_excel(raw_path, converted_path)
        with pd.ExcelFile(converted_path) as workbook:
            sheet_name = config.get("sheet_name", workbook.sheet_names[0])
            frame = pd.read_excel(workbook, sheet_name=sheet_name, header=None, dtype=str).fillna("")

    starts = [
        index
        for index, row in frame.iterrows()
        for title in config["title_candidates"]
        if any(title in str(cell) for cell in row)
    ]
    if not starts:
        return []
    start = max(starts)

    parsed: list[dict[str, Any]] = []
    expected = int(config["expected_count"])
    for _, row in frame.iloc[start + 1 :].iterrows():
        serial_col = int(config.get("serial_col", 0))
        serial = str(row.iloc[serial_col]).strip()
        if not serial.isdigit():
            continue
        serial_no = int(serial)
        if serial_no > expected:
            break
        school = str(row.iloc[int(config.get("school_col", 1))]).strip()
        major = str(row.iloc[int(config.get("major_col", 2))]).strip()
        code = normalize_source_major_code(str(row.iloc[int(config.get("code_col", 3))]).strip())
        degree_col = config.get("degree_col")
        duration_col = config.get("duration_col")
        note_col = config.get("note_col")
        degree = str(row.iloc[int(degree_col)]).strip() if degree_col is not None else ""
        duration = str(row.iloc[int(duration_col)]).strip() if duration_col is not None else ""
        note = str(row.iloc[int(note_col)]).strip() if note_col is not None and int(note_col) < len(row) else ""
        if not school or not major or not code:
            continue
        parsed.append(moe_cancel_row(config, serial_no, school, major, code, degree, duration, note))
        if serial_no == expected:
            break
    if len(parsed) != expected:
        raise ValueError(
            f"Parsed {len(parsed)} rows from {config['source_id']}, expected {expected}."
        )
    return parsed


def parse_moe_docx_cancel_rows(raw_dir: Path, config: dict[str, Any]) -> list[dict[str, Any]]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("Parsing Word .docx sources requires python-docx.") from exc

    raw_suffix = config.get("raw_suffix", ".docx")
    raw_path = raw_dir / f"{config['source_id']}{raw_suffix}"
    if not raw_path.exists():
        return []
    if raw_path.suffix.lower() == ".doc":
        with tempfile.TemporaryDirectory() as tmpdir:
            converted_path = Path(tmpdir) / f"{config['source_id']}.docx"
            convert_doc_with_word(raw_path, converted_path)
            document = Document(converted_path)
    else:
        document = Document(raw_path)
    table = document.tables[int(config["table_index"])]
    parsed: list[dict[str, Any]] = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", "") for cell in row.cells]
        try:
            school = cells[int(config["school_col"])].replace("※", "").strip()
            major = cells[int(config["major_col"])].strip()
            code = normalize_source_major_code(cells[int(config["code_col"])].strip())
            degree_col = config.get("degree_col")
            duration_col = config.get("duration_col")
            degree = cells[int(degree_col)].strip() if degree_col is not None else ""
            duration = cells[int(duration_col)].strip() if duration_col is not None else ""
            note_col = config.get("note_col")
            note = cells[int(note_col)].strip() if note_col is not None and int(note_col) < len(cells) else ""
        except IndexError:
            continue
        if not school or not major or not code:
            continue
        if "专业名称" in major or "学校名称" in school:
            continue
        if len({school, major, code, degree, duration}) == 1:
            continue
        if not any(char.isdigit() for char in code):
            continue
        serial_col = config.get("serial_col")
        if serial_col is not None:
            serial = cells[int(serial_col)].strip()
            if not serial.isdigit():
                continue
            serial_no = int(serial)
        else:
            serial_no = len(parsed) + 1
        parsed.append(moe_cancel_row(config, serial_no, school, major, code, degree, duration, note))
    expected = int(config["expected_count"])
    if len(parsed) != expected:
        raise ValueError(
            f"Parsed {len(parsed)} rows from {config['source_id']}, expected {expected}."
        )
    return parsed


def parse_moe_complete_cancel_rows(raw_dir: Path) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for config in MOE_CANCEL_DOCX_CONFIGS:
        rows.extend(parse_moe_docx_cancel_rows(raw_dir, config))
    for config in MOE_CANCEL_XLS_CONFIGS:
        rows.extend(parse_moe_xls_cancel_rows(raw_dir, config))
    for config in MOE_CANCEL_PDF_CONFIGS:
        rows.extend(parse_moe_pdf_cancel_rows(raw_dir, config))
    return rows


def parse_hunan_2019_undergrad_catalog_cancel_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "hunan_2019_undergrad_catalog_xls"
    raw_path = raw_dir / f"{source_id}.xls"
    if not raw_path.exists():
        return []
    with tempfile.TemporaryDirectory() as tmpdir:
        converted_path = Path(tmpdir) / f"{source_id}.xlsx"
        convert_xls_with_excel(raw_path, converted_path)
        with pd.ExcelFile(converted_path) as workbook:
            frame = pd.read_excel(workbook, sheet_name=workbook.sheet_names[0], header=None, dtype=str).fillna("")

    parsed: list[dict[str, Any]] = []
    for _, row in frame.iterrows():
        if len(row) < 7:
            continue
        note = str(row.iloc[6]).strip()
        if note != "拟撤销":
            continue
        serial = str(row.iloc[0]).strip()
        if not serial.isdigit():
            continue
        school = str(row.iloc[1]).strip()
        code = normalize_source_major_code(str(row.iloc[2]).strip())
        major = str(row.iloc[3]).strip()
        duration = str(row.iloc[4]).strip()
        teacher_flag = str(row.iloc[5]).strip()
        if not school or not major or not code:
            continue
        evidence = (
            f"湖南省教育厅湘教发[2019]8号附件第{serial}条：{school}{major}"
            f"（{code}，学制{duration}年）备注为拟撤销。通知正文规定，目录备注栏标注"
            "“拟撤销”的专业从2019年起停止招生，相关学校须在2019年7月本科专业申报时"
            "正式提出撤销申请。"
        )
        if teacher_flag:
            evidence += f"师范标识：{teacher_flag}。"
        parsed.append(
            {
                "policy_year": 2019,
                "region": school,
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "湖南省2019年招生专业目录拟撤销本科专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": serial,
                "study_duration": duration,
                "policy_action": "湖南省招生专业目录标注拟撤销，2019年起停止招生并要求正式提出撤销申请。",
                "criterion_text": "湖南省省属普通本科高校2019年招生专业目录备注栏标注“拟撤销”。",
                "source_ids": "hunan_2019_undergrad_catalog_notice;hunan_2019_undergrad_catalog_xls",
                "evidence_text": evidence,
                "confidence": "high",
            }
        )
    expected = 14
    if len(parsed) != expected:
        raise ValueError(f"Parsed {len(parsed)} rows from {source_id}, expected {expected}.")
    return parsed


def parse_scnu_2019_teaching_quality_cancel_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "scnu_2019_teaching_quality_report_pdf"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    majors = [
        "机械电子工程",
        "工业设计",
        "统计学",
        "经济统计学",
        "服装与服饰设计",
        "摄影",
        "影视摄影与制作",
    ]
    if "申请撤销" not in text or any(major not in text for major in majors):
        raise ValueError(f"Could not verify all SCNU 2019 cancellation majors in {source_id}.")
    rows: list[dict[str, Any]] = []
    for index, major in enumerate(majors, start=1):
        rows.append(
            {
                "policy_year": 2019,
                "region": "华南师范大学",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "华南师范大学2018-2019学年本科教学质量报告申请撤销专业",
                "reported_major_name": major,
                "major_code": "",
                "source_row_no": "",
                "study_duration": "",
                "policy_action": "学校本科教学质量报告列为申请撤销专业；教育部2019年度结果公布撤销。",
                "criterion_text": "华南师范大学2018-2019学年本科教学质量报告专业建设部分列出申请撤销专业。",
                "source_ids": source_id,
                "evidence_text": (
                    "华南师范大学2018-2019学年本科教学质量报告专业建设部分称，学校申请撤销"
                    "机械电子工程、工业设计、统计学、经济统计学、服装与服饰设计、摄影、"
                    f"影视摄影与制作7个专业；本行对应其中第{index}个：{major}。"
                ),
                "confidence": "high",
            }
        )
    return rows


def parse_ynnu_2019_teaching_quality_cancel_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "ynnu_2019_teaching_quality_report_pdf"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    validation_text = compact_text.replace("学校近五年专业增设、", "").replace("停招、撤并情况", "")
    validation_text = compact_text.replace("学校近五年专业增设、", "").replace("停招、撤并情况", "")
    rows_data = [
        ("应用统计学", "071202", "数学学院"),
        ("科学教育", "040102", "化学化工学院"),
        ("旅游管理与服务教育", "120904T", "旅游与地理科学学院"),
        ("市场营销", "120202", "经济与管理学院"),
        ("应用生物科学", "090109T", "生命科学学院"),
        ("软件工程", "080902", "信息学院"),
        ("网络工程", "080903", "信息学院"),
    ]
    if "共申请撤销7个本科专业" not in compact_text or "2019年云南师范大学申请撤销专业汇总表" not in compact_text:
        raise ValueError(f"Could not find YNNU 2019 cancellation table in {source_id}.")
    for major, code, college in rows_data:
        if major not in compact_text or code not in compact_text or college not in compact_text:
            raise ValueError(f"Could not verify {major} {code} {college} in {source_id}.")
    rows: list[dict[str, Any]] = []
    for index, (major, code, college) in enumerate(rows_data, start=1):
        rows.append(
            {
                "policy_year": 2019,
                "region": "云南师范大学",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "云南师范大学2019年申请撤销本科专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(index),
                "study_duration": "",
                "policy_action": "学校本科教学质量报告列为2019年申请撤销本科专业；教育部2019年度结果公布撤销。",
                "criterion_text": "云南师范大学2018-2019学年本科教学质量报告表3“2019年云南师范大学申请撤销专业汇总表”。",
                "source_ids": source_id,
                "evidence_text": (
                    "云南师范大学2018-2019学年本科教学质量报告专业建设部分称，2019年经学院申请、专家论证、"
                    "校长办公会研究决定，共申请撤销7个本科专业；"
                    f"表3第{index}项为{major}（{code}，{college}，调整情况：撤销）。"
                ),
                "confidence": "high",
            }
        )
    return rows


def parse_dzu_2019_teaching_quality_stop_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "dzu_2019_teaching_quality_report_pdf"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    rows_data = [
        ("产品设计", "130504"),
        ("动物科学", "090301"),
        ("经济统计学", "020102"),
        ("应用心理学", "071102"),
        ("电子信息科学与技术", "080714T"),
    ]
    if "2018-2019学年，停招产品设计、动物科学、经济统计学、应用心理学、电子信息科学与技术5个专业" not in compact_text:
        raise ValueError(f"Could not find DZU 2018-2019 stop-enrollment statement in {source_id}.")
    for major, code in rows_data:
        if major not in compact_text:
            raise ValueError(f"Could not verify {major} in {source_id}.")
    rows: list[dict[str, Any]] = []
    for index, (major, code) in enumerate(rows_data, start=1):
        rows.append(
            {
                "policy_year": 2019,
                "region": "德州学院",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "德州学院2018-2019学年停招本科专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(index),
                "study_duration": "",
                "policy_action": "学校本科教学质量报告列为2018-2019学年停招本科专业。",
                "criterion_text": "德州学院2018-2019学年本科教学质量报告本科专业设置情况段落列出的停招专业。",
                "source_ids": source_id,
                "evidence_text": (
                    "德州学院2018-2019学年本科教学质量报告本科专业设置情况说明，2018-2019学年停招产品设计、"
                    "动物科学、经济统计学、应用心理学、电子信息科学与技术5个专业；"
                    f"本行对应第{index}项：{major}（{code}）。"
                ),
                "confidence": "high",
            }
        )
    return rows


def parse_glut_2019_teaching_quality_stop_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "glut_2019_teaching_quality_report_pdf"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    rows_data = [
        ("工程管理", "120103", "工程管理（应用技能班）"),
        ("市场营销", "120202", "市场营销BCU"),
        ("酒店管理", "120902", "酒店管理（专升本）"),
        ("给排水科学与工程", "081003", "给排水科学与工程（应用技能班）"),
        ("产品设计", "130504", "产品设计（珠宝首饰工艺与设计）"),
        ("机械设计制造及其自动化", "080202", "机械设计制造及其自动化"),
        ("环境工程", "082502", "环境工程"),
        ("会计学", "120203K", "会计学"),
    ]
    required_fragments = [
        "2018-2019学年本科教学质量报告桂林理工大学",
        "停招的校内专业8个",
        "工程管理（应用技能班）,市场营销BCU，酒店管理（专升本）",
        "给排水科学与工程（应用技能班），产品设计（珠宝首饰工艺与设计）",
        "机械设计制造及其自动化，环境工程，会计学",
    ]
    missing = [fragment for fragment in required_fragments if fragment not in compact_text]
    missing.extend(source_name for _major, _code, source_name in rows_data if source_name not in compact_text)
    if missing:
        raise ValueError(f"Could not verify GLUT 2018-2019 stop-enrollment fragments in {source_id}: {missing}")

    rows: list[dict[str, Any]] = []
    for index, (major, code, source_name) in enumerate(rows_data, start=1):
        rows.append(
            {
                "policy_year": 2019,
                "region": "桂林理工大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "桂林理工大学2018-2019学年停招校内专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(index),
                "study_duration": "",
                "policy_action": "学校本科教学质量报告列为2018-2019学年停招校内专业。",
                "criterion_text": "桂林理工大学2018-2019学年本科教学质量报告专业建设部分列出的当年停招校内专业；reported_major_name标准化为基础本科专业名，来源括注保留在证据文本。",
                "source_ids": source_id,
                "evidence_text": (
                    "桂林理工大学2018-2019学年本科教学质量报告专业建设部分称，当年学校招生的校内专业92个，"
                    "停招的校内专业8个，分别为工程管理（应用技能班）、市场营销BCU、酒店管理（专升本）、"
                    "给排水科学与工程（应用技能班）、产品设计（珠宝首饰工艺与设计）、机械设计制造及其自动化、"
                    f"环境工程、会计学；本行对应第{index}项：{source_name}（基础专业：{major}，{code}）。"
                ),
                "confidence": "high",
            }
        )
    return rows


def parse_hrbu_2020_teaching_quality_stop_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "hrbu_2020_teaching_quality_report_pdf"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    rows_data = [
        ("物理学", "070201"),
        ("网络工程", "080903"),
        ("服装与服饰设计", "130505"),
        ("公共艺术", "130506"),
        ("秘书学", "050107T"),
        ("酒店管理", "120902"),
        ("思想政治教育", "030503"),
    ]
    if "附表4专业设置及调整情况" not in compact_text or "当年停招专业名单" not in compact_text:
        raise ValueError(f"Could not find HRBU 2019-2020 stop-enrollment table in {source_id}.")
    for major, code in rows_data:
        if major not in compact_text:
            raise ValueError(f"Could not verify {major} in {source_id}.")
    rows: list[dict[str, Any]] = []
    for index, (major, code) in enumerate(rows_data, start=1):
        rows.append(
            {
                "policy_year": 2020,
                "region": "哈尔滨学院",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "哈尔滨学院2019-2020学年当年停招本科专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(index),
                "study_duration": "",
                "policy_action": "学校本科教学质量报告附表列为当年停招本科专业。",
                "criterion_text": "哈尔滨学院2019-2020学年本科教学质量报告附表4“专业设置及调整情况”中的当年停招专业名单。",
                "source_ids": source_id,
                "evidence_text": (
                    "哈尔滨学院2019-2020学年本科教学质量报告附表4“专业设置及调整情况”列出当年停招专业名单："
                    "物理学、网络工程、服装与服饰设计、公共艺术、秘书学、酒店管理、思想政治教育；"
                    f"本行对应第{index}项：{major}（{code}）。"
                ),
                "confidence": "high",
            }
        )
    return rows


def parse_hrbu_2020_teaching_quality_cancel_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "hrbu_2020_teaching_quality_report_pdf"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    rows_data = [
        ("体育教育", "040201"),
        ("历史学", "060101"),
        ("地理科学", "070501"),
        ("生物科学", "071001"),
        ("心理学", "071101"),
        ("美术学", "130401"),
    ]
    required_fragments = [
        "2020年，学校撤销了历史学、心理学、体育教育、地理科学、美术学、生物科学等6个专业",
        "停止秘书学、物理学、酒店管理、思想政治教育、服装与服饰、公共艺术、网络工程等7个专业的招生",
    ]
    missing = [fragment for fragment in required_fragments if fragment not in compact_text]
    if missing:
        raise ValueError(f"Could not verify HRBU 2020 cancellation fragments in {source_id}: {missing}")

    rows: list[dict[str, Any]] = []
    for index, (major, code) in enumerate(rows_data, start=1):
        rows.append(
            {
                "policy_year": 2019,
                "region": "哈尔滨学院",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "哈尔滨学院2019年度本科专业备案审批结果撤销专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": "",
                "study_duration": "",
                "policy_action": "学校本科教学质量报告列为2020年撤销专业；教育部2019年度结果公布撤销。",
                "criterion_text": "哈尔滨学院2019-2020学年本科教学质量报告专业动态调整机制段落中的撤销专业。",
                "source_ids": source_id,
                "evidence_text": (
                    "哈尔滨学院2019-2020学年本科教学质量报告说明，2020年学校撤销了历史学、心理学、"
                    "体育教育、地理科学、美术学、生物科学等6个专业；"
                    f"本行对应可与教育部2019年度撤销名单合并的第{index}项：{major}（{code}）。"
                ),
                "confidence": "high",
            }
        )
    return rows


def parse_cust_2020_application_history_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "cust_2020_applied_statistics_application_pdf"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    validation_text = compact_text.replace("停招、撤并情况（300字", "").replace("以内）", "")
    stop_rows = [
        (2016, "经济学", "020101"),
        (2016, "材料物理", "080402"),
        (2020, "网络工程", "080903"),
    ]
    cancel_rows = [
        ("工业设计", "080205"),
        ("环境科学", "082503"),
        ("劳动与社会保障", "120403"),
        ("服装与服饰设计", "130505"),
        ("轨道交通信号与控制", "080802T"),
    ]
    required_fragments = [
        "近五年停招",
        "近五年撤销",
        "经济学”专业（2016年）",
        "材料物理”专业（2016年）",
        "网络工程”专业（2020年）",
        "工业设计”专业（2019年）",
        "环境科学”专业（2019年）",
        "劳动与社会保障”专业（2019年）",
        "服装与服饰设计”专业（2019年）",
        "轨道交通信号与控制”专业（2019年）",
    ]
    missing = [fragment for fragment in required_fragments if fragment not in validation_text]
    if missing:
        raise ValueError(f"Could not verify CUST application history fragments in {source_id}: {missing}")

    rows: list[dict[str, Any]] = []
    for index, (year, major, code) in enumerate(stop_rows, start=1):
        rows.append(
            {
                "policy_year": year,
                "region": "长春理工大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": f"长春理工大学{year}年近五年专业建设情况停招本科专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(index),
                "study_duration": "",
                "policy_action": "学校专业备案申请表列为近五年停招本科专业。",
                "criterion_text": "长春理工大学应用统计学专业备案申请表“学校近五年专业增设、停招、撤并情况”中的停招专业。",
                "source_ids": source_id,
                "evidence_text": (
                    "长春理工大学应用统计学专业备案申请表在“学校近五年专业增设、停招、撤并情况”中说明，"
                    "近五年停招经济学专业（2016年）、材料物理专业（2016年）、网络工程专业（2020年）；"
                    f"本行对应第{index}项：{major}（{code}，{year}年）。"
                ),
                "confidence": "high",
            }
        )
    for index, (major, code) in enumerate(cancel_rows, start=1):
        rows.append(
            {
                "policy_year": 2019,
                "region": "长春理工大学",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "长春理工大学2019年近五年专业建设情况撤销本科专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": "",
                "study_duration": "",
                "policy_action": "学校专业备案申请表列为2019年撤销本科专业；教育部2019年度结果公布撤销。",
                "criterion_text": "长春理工大学应用统计学专业备案申请表“学校近五年专业增设、停招、撤并情况”中的撤销专业。",
                "source_ids": source_id,
                "evidence_text": (
                    "长春理工大学应用统计学专业备案申请表在“学校近五年专业增设、停招、撤并情况”中说明，"
                    "近五年撤销工业设计专业（2019年）、环境科学专业（2019年）、劳动与社会保障专业（2019年）、"
                    "服装与服饰设计专业（2019年）、轨道交通信号与控制专业（2019年）；"
                    f"本行对应其中第{index}项：{major}（{code}）。"
                ),
                "confidence": "high",
            }
        )
    return rows


def parse_suse_2019_teaching_quality_cancel_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "suse_2019_teaching_quality_report_pdf"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    majors = [
        "统计学",
        "经济统计学",
        "物理学",
        "化学生物学",
        "轨道交通信号与控制",
        "工业工程",
        "社会工作",
        "信息管理与信息系统",
    ]
    if "申请撤销" not in text or any(major not in text for major in majors):
        raise ValueError(f"Could not verify all SUSE 2019 cancellation majors in {source_id}.")
    rows: list[dict[str, Any]] = []
    for index, major in enumerate(majors, start=1):
        rows.append(
            {
                "policy_year": 2019,
                "region": "四川轻化工大学",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "四川轻化工大学2018-2019学年本科教学质量报告申请撤销专业",
                "reported_major_name": major,
                "major_code": "",
                "source_row_no": "",
                "study_duration": "",
                "policy_action": "学校本科教学质量报告列为申请撤销专业；教育部2019年度结果公布撤销。",
                "criterion_text": "四川轻化工大学2018-2019学年本科教学质量报告专业建设部分列出申请撤销专业。",
                "source_ids": source_id,
                "evidence_text": (
                    "四川轻化工大学2018-2019学年本科教学质量报告专业建设部分称，2019年学校"
                    "申请撤销统计学、经济统计学、物理学、化学生物学、轨道交通信号与控制、"
                    f"工业工程、社会工作、信息管理与信息系统等8个本科专业；本行对应其中第{index}个：{major}。"
                ),
                "confidence": "high",
            }
        )
    return rows


def parse_ourjiangsu_2020_2019_jiangsu_cancel_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "ourjiangsu_2020_2019_jiangsu_cancel_summary"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    school_majors = {
        "南京工业大学": [("生物技术", "生物技术"), ("乳品工程", "乳品工程")],
        "南通大学": [
            ("教育学", "教育学"),
            ("广播电视学", "广播电视学"),
            ("轨道交通信号与控制", "轨道交通信号与控制"),
            ("船舶与海洋工程", "船舶与海洋工程"),
        ],
        "中国传媒大学南广学院": [("徳语", "德语"), ("阿拉伯语", "阿拉伯语")],
        "江苏师范大学科文学院": [
            ("应用心理学", "应用心理学"),
            ("应用统计学", "应用统计学"),
            ("制药工程", "制药工程"),
            ("环境设计", "环境设计"),
            ("产品设计", "产品设计"),
        ],
    }
    required_fragments = [
        "（三）撤销专业13个",
        "南京工业大学：生物技术，乳品工程",
        "南通大学：教育学，广播电视学，轨道交通信号与控制，船舶与海洋工程",
        "中国传媒大学南广学院：德语，阿拉伯语",
        "江苏师范大学科文学院：应用心理学，应用统计学，制药工程，环境设计，产品设计",
        "来源：中华人民共和国教育部官方网站",
    ]
    missing = [fragment for fragment in required_fragments if fragment not in compact_text]
    if missing:
        raise ValueError(f"Could not verify Jiangsu 2019 cancellation summary fragments in {source_id}: {missing}")

    rows: list[dict[str, Any]] = []
    for school, majors in school_majors.items():
        for index, (reported_major, source_major) in enumerate(majors, start=1):
            rows.append(
                {
                    "policy_year": 2019,
                    "region": school,
                    "education_level": "本科",
                    "record_type": "major_cancel",
                    "warning_label": "2019年度江苏高校撤销专业二级汇总",
                    "reported_major_name": reported_major,
                    "major_code": "",
                    "source_row_no": "",
                    "study_duration": "",
                    "policy_action": "二级转载来源按校列入2019年度撤销专业名单；教育部2019年度结果公布撤销。",
                    "criterion_text": "我苏网转载教育部官网信息，按校列出江苏省2019年度普通高校撤销专业13个。",
                    "source_ids": source_id,
                    "evidence_text": (
                        "我苏网2020年3月3日转载教育部官网信息，在“撤销专业13个”段落按校列出："
                        "南京工业大学生物技术、乳品工程；南通大学教育学、广播电视学、轨道交通信号与控制、船舶与海洋工程；"
                        "中国传媒大学南广学院德语、阿拉伯语；江苏师范大学科文学院应用心理学、应用统计学、制药工程、环境设计、产品设计；"
                        f"本行对应{school}第{index}项：{source_major}。"
                    ),
                    "confidence": "medium",
                }
            )
    return rows


def parse_cyol_2018_hdu_major_stop_adjustment_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "cyol_2018_hdu_major_stop_adjustment"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    future_stop_majors = [
        "信息工程",
        "物联网工程",
        "应用物理学",
        "海洋工程与技术",
        "生物医学工程",
        "物流管理",
        "市场营销",
        "环境科学",
        "经济统计学",
        "编辑出版学",
        "工业设计",
    ]
    past_cancel_or_stop = [
        (2013, "教育技术学"),
        (2019, "印刷工程"),
        (2019, "包装工程"),
        (2019, "智能电网信息工程"),
        (2019, "功能材料"),
        (2019, "应用统计学"),
    ]
    required_fragments = [
        "未来两年内，该校将停招信息工程、物联网工程、应用物理学、海洋工程与技术、生物医学工程、物流管理、市场营销、环境科学、经济统计学、编辑出版学、工业设计等11个本科专业",
        "过去5年，该校已先后有教育技术学、印刷工程、包装工程、智能电网信息工程、功能材料、应用统计学等6个专业被撤销或停招",
        "印刷工程专业，在2014年被红牌警告后停招",
        "包装工程，连续3次因“低效表现”被黄牌警告，后于2016年停止招生",
    ]
    missing = [fragment for fragment in required_fragments if fragment not in compact_text]
    if missing:
        raise ValueError(f"Could not verify HDU stop-adjustment fragments in {source_id}: {missing}")

    rows: list[dict[str, Any]] = []
    for index, major in enumerate(future_stop_majors, start=1):
        rows.append(
            {
                "policy_year": 2018,
                "region": "杭州电子科技大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "杭州电子科技大学2018年未来两年拟停招本科专业",
                "reported_major_name": major,
                "major_code": "",
                "source_row_no": str(index),
                "study_duration": "",
                "policy_action": "二级媒体来源报道学校未来两年内将停招该本科专业。",
                "criterion_text": "中国青年报报道杭州电子科技大学基于专业第一志愿报考率、就业率、转专业率等指标优化低效专业。",
                "source_ids": source_id,
                "evidence_text": (
                    "中国青年报2018年4月16日报道，杭州电子科技大学未来两年内将停招信息工程、物联网工程、应用物理学、"
                    "海洋工程与技术、生物医学工程、物流管理、市场营销、环境科学、经济统计学、编辑出版学、工业设计等11个本科专业；"
                    f"本行对应第{index}项：{major}。"
                ),
                "confidence": "medium",
            }
        )
    for year, major in past_cancel_or_stop:
        rows.append(
            {
                "policy_year": year,
                "region": "杭州电子科技大学",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "杭州电子科技大学过去五年已撤销或停招专业媒体证据",
                "reported_major_name": major,
                "major_code": "",
                "source_row_no": "",
                "study_duration": "",
                "policy_action": "二级媒体来源列为过去五年已撤销或停招专业；教育部年度结果中对应撤销记录仍以教育部来源为准。",
                "criterion_text": "中国青年报报道杭州电子科技大学过去五年已有6个专业被撤销或停招。",
                "source_ids": source_id,
                "evidence_text": (
                    "中国青年报2018年4月16日报道，杭州电子科技大学过去5年先后有教育技术学、印刷工程、包装工程、"
                    f"智能电网信息工程、功能材料、应用统计学等6个专业被撤销或停招；本行对应：{major}。"
                ),
                "confidence": "medium",
            }
        )
    return rows


def parse_xpu_2019_major_dynamic_adjustment_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "xpu_2019_major_dynamic_adjustment"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    cancel_rows = [
        ("机电工程学院", "包装工程", "081702", "工学"),
        ("管理学院", "市场营销", "120202", "管理学"),
        ("理学院", "光电信息科学与工程", "080705", "理学"),
        ("服装与艺术设计学院", "摄影", "130404", "艺术学"),
        ("电子信息学院", "测控技术与仪器", "080301", "工学"),
    ]
    stop_rows = [
        ("管理学院", "信息管理与信息系统", "120102", "管理学"),
        ("理学院", "信息与计算科学", "070102", "理学"),
        ("计算机科学学院", "数字媒体技术", "080906", "工学"),
    ]
    warning_rows = [
        ("纺织科学与工程学院", "非织造材料与工程", "081603T", "工学"),
        ("机电工程学院", "工业工程", "120701", "工学"),
        ("机电工程学院", "过程装备与控制工程", "080206", "工学"),
        ("环境与化学工程学院", "环境科学", "082503", "理学"),
        ("环境与化学工程学院", "生物工程", "083001", "工学"),
        ("管理学院", "行政管理", "120402", "管理学"),
        ("人文社会科学学院", "汉语国际教育", "050103", "文学"),
        ("新媒体艺术学院", "戏剧影视美术设计", "130307", "艺术学"),
    ]
    required_fragments = [
        "确定2019年拟增设4个新专业、撤销5个专业、预警8个专业，2020年拟停招3个专业",
        "表2撤销专业名单",
        "表3拟停招专业名单",
        "表4预警专业名单",
        "公示七天（2018年7月8日至7月15日）",
    ]
    for table_rows in [cancel_rows, stop_rows, warning_rows]:
        for college, major, code, degree in table_rows:
            required_fragments.append(f"{college}{major}{code}{degree}")
    missing = [fragment for fragment in required_fragments if fragment not in compact_text]
    if missing:
        raise ValueError(f"Could not verify XPU dynamic-adjustment fragments in {source_id}: {missing}")

    rows: list[dict[str, Any]] = []
    for index, (college, major, code, degree) in enumerate(cancel_rows, start=1):
        rows.append(
            {
                "policy_year": 2019,
                "region": "西安工程大学",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "西安工程大学2019年现有专业动态调整拟撤销专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": "",
                "study_duration": "",
                "policy_action": "学校教务处公示列为2019年拟撤销专业；教育部2019年度结果公布撤销。",
                "criterion_text": "西安工程大学根据专业发展规划和专业动态调整工作要求，经专业动态调整工作小组评审、征求教学单位意见、校教学委员会审议和校长办公会讨论通过。",
                "source_ids": source_id,
                "evidence_text": (
                    "西安工程大学教务处2018年7月8日公示称，确定2019年撤销5个专业；"
                    f"表2第{index}项列出{college}{major}（{code}，{degree}）。"
                ),
                "confidence": "high",
            }
        )
    for index, (college, major, code, degree) in enumerate(stop_rows, start=1):
        rows.append(
            {
                "policy_year": 2020,
                "region": "西安工程大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "西安工程大学2020年拟停招本科专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(index),
                "study_duration": "",
                "policy_action": "学校教务处公示列为2020年拟停招专业。",
                "criterion_text": "西安工程大学现有专业动态调整公示表3“拟停招专业名单”。",
                "source_ids": source_id,
                "evidence_text": (
                    "西安工程大学教务处2018年7月8日公示表3列出2020年拟停招专业；"
                    f"第{index}项为{college}{major}（{code}，{degree}）。"
                ),
                "confidence": "high",
            }
        )
    for index, (college, major, code, degree) in enumerate(warning_rows, start=1):
        rows.append(
            {
                "policy_year": 2019,
                "region": "西安工程大学",
                "education_level": "本科",
                "record_type": "major_warning_list",
                "warning_label": "西安工程大学2019年现有专业动态调整预警专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(index),
                "study_duration": "",
                "policy_action": "学校教务处公示列为2019年拟预警专业。",
                "criterion_text": "西安工程大学现有专业动态调整公示表4“预警专业名单”。",
                "source_ids": source_id,
                "evidence_text": (
                    "西安工程大学教务处2018年7月8日公示表4列出预警专业名单；"
                    f"第{index}项为{college}{major}（{code}，{degree}）。"
                ),
                "confidence": "high",
            }
        )
    return rows


def parse_zjnu_2020_2019_undergrad_setup_news_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "zjnu_2020_2019_undergrad_setup_news"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    majors = ["教育学", "艺术教育", "环境科学"]
    required_fragments = [
        "教育部印发《关于公布2019年度普通高等学校本科专业备案和审批结果的通知》(教高函〔2020〕2号)",
        "同意浙江师范大学增设“智能制造工程专业”专业",
        "撤销教育学、艺术教育、环境科学三个本科专业",
        "浙师大本科专业总数为71个",
    ]
    missing = [fragment for fragment in required_fragments if fragment not in compact_text]
    if missing:
        raise ValueError(f"Could not verify ZJNU 2019 cancellation fragments in {source_id}: {missing}")

    rows: list[dict[str, Any]] = []
    for index, major in enumerate(majors, start=1):
        rows.append(
            {
                "policy_year": 2019,
                "region": "浙江师范大学",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "浙江师范大学2019年度本科专业备案审批结果撤销专业",
                "reported_major_name": major,
                "major_code": "",
                "source_row_no": "",
                "study_duration": "",
                "policy_action": "学校学院官网页面引用教育部2019年度结果，列为撤销本科专业。",
                "criterion_text": "浙江师范大学工学院页面引用教高函〔2020〕2号，说明学校2019年度本科专业备案和审批结果中的撤销专业。",
                "source_ids": source_id,
                "evidence_text": (
                    "浙江师范大学工学院页面称，教育部印发教高函〔2020〕2号，同意学校增设智能制造工程专业，"
                    f"并撤销教育学、艺术教育、环境科学三个本科专业；本行对应第{index}项：{major}。"
                ),
                "confidence": "high",
            }
        )
    return rows


def parse_jxau_smart_agriculture_application_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "jxau_smart_agriculture_application_pdf"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    validation_text = compact_text.replace("停招、撤并情况（300字", "").replace("以内）", "")
    stop_rows = [
        (2020, "商务英语", "050262"),
        (2021, "信息与计算科学", "070102"),
    ]
    cancel_rows = [
        (2018, "农村区域发展"),
        (2018, "财务管理"),
        (2018, "音乐表演"),
        (2018, "网络工程"),
        (2018, "信息管理与信息系统"),
        (2018, "工程管理"),
        (2018, "交通运输"),
        (2018, "管理科学"),
        (2018, "生物科学"),
        (2018, "劳动与社会保障"),
        (2019, "环境科学"),
        (2019, "环境工程"),
        (2019, "市场营销"),
        (2019, "城乡规划"),
        (2020, "视觉传达设计"),
    ]
    required_fragments = [
        "专业名称智慧农业",
        "专业代码090112T",
        "停招：2020年停招商务英语专业，2021年停招信息与计算科学专业",
        "2018年撤销农村区域发展、财务管理、音乐表演、网络工程、信息管理与信息系统、工程管理、交通运输、管理科学、生物科学、劳动与社会保障等10个专业",
        "2019年撤销环境科学、环境工程、市场营销、城乡规划（四年制）等4个专业",
        "2020年撤销视觉传达设计等1个专业",
    ]
    missing = [fragment for fragment in required_fragments if fragment not in validation_text]
    if missing:
        raise ValueError(f"Could not verify JXAU professional-history fragments in {source_id}: {missing}")

    rows: list[dict[str, Any]] = []
    for index, (year, major, code) in enumerate(stop_rows, start=1):
        rows.append(
            {
                "policy_year": year,
                "region": "江西农业大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": f"江西农业大学{year}年近五年专业建设情况停招本科专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(index),
                "study_duration": "",
                "policy_action": "学校专业备案申请表列为近五年停招本科专业。",
                "criterion_text": "江西农业大学智慧农业专业备案申请表“学校近五年专业增设、调整、停招、撤并情况”中的停招专业。",
                "source_ids": source_id,
                "evidence_text": (
                    "江西农业大学智慧农业专业备案申请表在学校近五年专业增设、调整、停招、撤并情况中说明，"
                    "2020年停招商务英语专业，2021年停招信息与计算科学专业；"
                    f"本行对应第{index}项：{major}（{code}，{year}年）。"
                ),
                "confidence": "high",
            }
        )
    for index, (year, major) in enumerate(cancel_rows, start=1):
        rows.append(
            {
                "policy_year": year,
                "region": "江西农业大学",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": f"江西农业大学{year}年近五年专业建设情况撤销本科专业",
                "reported_major_name": major,
                "major_code": "",
                "source_row_no": "",
                "study_duration": "",
                "policy_action": "学校专业备案申请表列为近五年撤销本科专业；教育部年度结果中对应撤销记录以教育部来源为准。",
                "criterion_text": "江西农业大学智慧农业专业备案申请表“学校近五年专业增设、调整、停招、撤并情况”中的撤销专业。",
                "source_ids": source_id,
                "evidence_text": (
                    "江西农业大学智慧农业专业备案申请表说明，2018年撤销农村区域发展、财务管理、音乐表演、网络工程、"
                    "信息管理与信息系统、工程管理、交通运输、管理科学、生物科学、劳动与社会保障等10个专业；"
                    "2019年撤销环境科学、环境工程、市场营销、城乡规划（四年制）等4个专业；2020年撤销视觉传达设计1个专业；"
                    f"本行对应第{index}项：{major}（{year}年）。"
                ),
                "confidence": "high",
            }
        )
    return rows


def parse_hbu_jijianjiancha_application_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "hbu_jijianjiancha_application_pdf"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    validation_text = compact_text.replace("学校近五年专业增设、", "").replace("停招、撤并情况", "")
    cancel_rows = [
        ("网络工程", "080903"),
        ("信息与计算科学", "070102"),
    ]
    required_fragments = [
        "专业代码030108TK",
        "专业名称纪检监察",
        "（二）停招专业：教育技术学、网络工程、公共事业管理、生物医学工程",
        "朝鲜语、俄语、法语、葡萄牙语、政治学与行政学、工程力学、广播电视编导、电子商务",
        "（三）撤销专业：网络工程、信息与计算科学、生物医学工程、工程力学、政治学与行政学、教育技术学、汉语言、朝鲜语、电子信息工程、生物工程",
        "市场营销、公共事业管理",
    ]
    missing = [fragment for fragment in required_fragments if fragment not in validation_text]
    if missing:
        raise ValueError(f"Could not verify HBU professional-history fragments in {source_id}: {missing}")

    rows: list[dict[str, Any]] = []
    for index, (major, code) in enumerate(cancel_rows, start=1):
        rows.append(
            {
                "policy_year": 2019,
                "region": "河北大学",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "河北大学2019年度本科专业备案审批结果撤销专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": "",
                "study_duration": "",
                "policy_action": "学校专业备案申请表列为近五年撤销专业；教育部2019年度结果公布撤销。",
                "criterion_text": "河北大学纪检监察专业备案申请表“学校近五年专业增设、停招、撤并情况”中的撤销专业；仅结构化可与教育部2019年度撤销名单精确对应的专业。",
                "source_ids": source_id,
                "evidence_text": (
                    "河北大学纪检监察专业备案申请表在“学校近五年专业增设、停招、撤并情况”中列出撤销专业包括"
                    "网络工程、信息与计算科学、生物医学工程、工程力学、政治学与行政学、教育技术学、汉语言、"
                    "朝鲜语、电子信息工程、生物工程、市场营销、公共事业管理；"
                    f"本行对应可与教育部2019年度撤销名单合并的第{index}项：{major}（{code}）。"
                ),
                "confidence": "high",
            }
        )
    return rows


def parse_kmust_2025_stomatology_application_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "kmust_2025_stomatology_application_pdf"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    cancel_rows = [
        ("思想政治教育", "030503"),
        ("教育技术学", "040104"),
        ("自然地理与资源环境", "070502"),
        ("金属材料工程", "080405"),
        ("无机非金属材料工程", "080406"),
    ]
    required_fragments = [
        "专业代码100301K",
        "专业名称口腔医学",
        "学校近五年专业增设、停招",
        "撤并：金属材料工程、无机非金属材料工程、思想政治教育、教育技术学、自然地理与资源环境",
        "电子科学与技术、信息工程、资源循环科学与工程、包装工程、农业电气化、能源化学工程",
        "编辑出版学、勘查技术与工程、动画、城市管理、软件工程、广告学、园林、智能科学与技术",
    ]
    missing = [fragment for fragment in required_fragments if fragment not in compact_text]
    if missing:
        raise ValueError(f"Could not verify KMUST professional-history fragments in {source_id}: {missing}")

    rows: list[dict[str, Any]] = []
    for index, (major, code) in enumerate(cancel_rows, start=1):
        rows.append(
            {
                "policy_year": 2019,
                "region": "昆明理工大学",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "昆明理工大学2019年度本科专业备案审批结果撤销专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": "",
                "study_duration": "",
                "policy_action": "学校专业申报材料列为近五年撤并专业；教育部2019年度结果公布撤销。",
                "criterion_text": "昆明理工大学口腔医学专业申请表“学校近五年专业增设、停招、撤并情况”中的撤并专业；仅结构化可与教育部2019年度撤销名单精确对应的专业。",
                "source_ids": source_id,
                "evidence_text": (
                    "昆明理工大学口腔医学专业申请表在“学校近五年专业增设、停招、撤并情况”中列出撤并专业包括"
                    "金属材料工程、无机非金属材料工程、思想政治教育、教育技术学、自然地理与资源环境、"
                    "电子科学与技术、信息工程、资源循环科学与工程、包装工程、农业电气化、能源化学工程、"
                    "编辑出版学、勘查技术与工程、动画、城市管理、软件工程、广告学、园林、智能科学与技术等；"
                    f"本行对应可与教育部2019年度撤销名单合并的第{index}项：{major}（{code}）。"
                ),
                "confidence": "high",
            }
        )
    return rows


def parse_nhky_2020_cybersecurity_application_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "nhky_2020_cybersecurity_application_pdf"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    stop_rows = [
        (2015, "电子科学与技术", "080702"),
        (2015, "电子信息科学与技术", "080714T"),
        (2015, "信息管理与信息系统", "120102"),
        (2015, "音乐学", "130202"),
        (2015, "应用化学", "070302"),
        (2015, "给排水科学与工程", "081003"),
        (2016, "高分子材料与工程", "080407"),
        (2016, "网络工程", "080903"),
        (2016, "通信工程", "080703"),
        (2016, "工程管理", "120103"),
    ]
    cancel_rows = [
        (2020, "信息管理与信息系统", "120102"),
        (2020, "音乐学", "130202"),
        (2020, "应用化学", "070302"),
        (2020, "电子信息科学与技术", "080714T"),
    ]
    required_fragments = [
        "专业代码080911TK专业名称网络空间安全",
        "2015-2016学年无新增专业，停招电子科学与技术、电子信息科学与技术、信息管理与信息系统、音乐学、应用化学、给排水科学与工程",
        "2016-2017学年新增专业软件工程、动画，停招信息管理与信息系统、音乐学、应用化学、给排水科学与工程、高分子材料与工程、网络工程、通信工程、工程管理、电子科学与技术、电子信息科学与技术",
        "2018-2019学年和2019-2020学年无新增专业，停招本科专业同上学年",
        "现报撤销信息管理与信息系统、音乐学、应用化学、电子信息科学与技术待批",
    ]
    missing = [fragment for fragment in required_fragments if fragment not in compact_text]
    if missing:
        raise ValueError(f"Could not verify NHKY professional-history fragments in {source_id}: {missing}")

    rows: list[dict[str, Any]] = []
    for index, (year, major, code) in enumerate(stop_rows, start=1):
        rows.append(
            {
                "policy_year": year,
                "region": "南昌航空大学科技学院",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": f"南昌航空大学科技学院{year}学年近五年专业建设情况停招本科专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(index),
                "study_duration": "",
                "policy_action": "学校专业申报材料列为近五年停招本科专业。",
                "criterion_text": "南昌航空大学科技学院网络空间安全专业申请表“学校近五年专业增设、停招、撤并情况”中的停招专业；年份取首次列入停招的学年起始年。",
                "source_ids": source_id,
                "evidence_text": (
                    "南昌航空大学科技学院网络空间安全专业申请表说明，2015-2016学年停招电子科学与技术、"
                    "电子信息科学与技术、信息管理与信息系统、音乐学、应用化学、给排水科学与工程，"
                    "2016-2017学年新增停招高分子材料与工程、网络工程、通信工程、工程管理；"
                    f"本行对应第{index}项：{major}（{code}，{year}学年起）。"
                ),
                "confidence": "high",
            }
        )
    for index, (year, major, code) in enumerate(cancel_rows, start=1):
        rows.append(
            {
                "policy_year": year,
                "region": "南昌航空大学科技学院",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "南昌航空大学科技学院2020年度本科专业备案审批结果撤销专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": "",
                "study_duration": "",
                "policy_action": "学校专业申报材料列为现报撤销待批；教育部2020年度结果公布撤销。",
                "criterion_text": "南昌航空大学科技学院网络空间安全专业申请表“学校近五年专业增设、停招、撤并情况”中的现报撤销待批专业；用于补充可与教育部2020年度撤销名单精确对应的专业证据。",
                "source_ids": source_id,
                "evidence_text": (
                    "南昌航空大学科技学院网络空间安全专业申请表说明，现报撤销信息管理与信息系统、音乐学、"
                    "应用化学、电子信息科学与技术待批；"
                    f"本行对应可与教育部2020年度撤销名单合并的第{index}项：{major}（{code}）。"
                ),
                "confidence": "high",
            }
        )
    return rows


def parse_qjnu_2025_area_studies_application_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "qjnu_2025_area_studies_application_pdf"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    validation_text = compact_text.replace("学校近五年专业增设、", "").replace("停招、撤并情况", "")
    stop_rows = [
        ("金融工程", "020302"),
        ("休闲体育", "040207T"),
        ("秘书学", "050107T"),
        ("印度尼西亚语", "050212"),
        ("商务英语", "050262"),
        ("广告学", "050303"),
        ("电子信息科学与技术", "080714T"),
        ("软件工程", "080902"),
        ("信息管理与信息系统", "120102"),
        ("工程造价", "120105"),
        ("市场营销", "120202"),
        ("国际商务", "120205"),
        ("审计学", "120207"),
        ("环境设计", "130503"),
    ]
    cancel_rows = [
        (2022, "艺术教育", "040105"),
        (2024, "房地产开发与管理", "120104"),
        (2024, "航空服务艺术与管理", "130208TK"),
    ]
    required_fragments = [
        "专业名称：区域国别学",
        "专业代码：0502104TK",
        "先后停招金融工程、休闲体育、秘书学",
        "共14个专业，撤销艺术教育、房地产开发与管理、航空服务艺术与管理3个专业",
        "撤销艺术教育、房地产开发与管理、航空服务艺术与管理3个专业",
    ]
    missing = [fragment for fragment in required_fragments if fragment not in validation_text]
    missing.extend(major for major, _code in stop_rows if major not in validation_text)
    missing.extend(major for _year, major, _code in cancel_rows if major not in validation_text)
    if missing:
        raise ValueError(f"Could not verify QJNU professional-history fragments in {source_id}: {missing}")

    rows: list[dict[str, Any]] = []
    for index, (major, code) in enumerate(stop_rows, start=1):
        rows.append(
            {
                "policy_year": 2025,
                "region": "曲靖师范学院",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "曲靖师范学院2025年专业申报材料近五年停招本科专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(index),
                "study_duration": "",
                "policy_action": "学校专业申报材料列为近五年已停招本科专业。",
                "criterion_text": "曲靖师范学院区域国别学专业申请表“学校近五年专业增设、停招、撤并情况”中的停招专业；来源未给出逐项停招年份，policy_year采用申请材料年份。",
                "source_ids": source_id,
                "evidence_text": (
                    "曲靖师范学院区域国别学专业申请表说明，近五年来学校先后停招金融工程、休闲体育、秘书学、"
                    "印度尼西亚语、商务英语、广告学、电子信息科学与技术、软件工程、信息管理与信息系统、"
                    "工程造价、市场营销、国际商务、审计学、环境设计共14个专业；"
                    f"本行对应第{index}项：{major}（{code}）。"
                ),
                "confidence": "high",
            }
        )
    for index, (year, major, code) in enumerate(cancel_rows, start=1):
        rows.append(
            {
                "policy_year": year,
                "region": "曲靖师范学院",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": f"曲靖师范学院{year}年度本科专业备案审批结果撤销专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": "",
                "study_duration": "",
                "policy_action": "学校专业申报材料列为近五年撤销本科专业；教育部年度结果公布撤销。",
                "criterion_text": "曲靖师范学院区域国别学专业申请表“学校近五年专业增设、停招、撤并情况”中的撤销专业；用于补充可与教育部年度撤销名单精确对应的专业证据。",
                "source_ids": source_id,
                "evidence_text": (
                    "曲靖师范学院区域国别学专业申请表说明，近五年来学校撤销艺术教育、房地产开发与管理、"
                    "航空服务艺术与管理3个专业；"
                    f"本行对应可与教育部年度撤销名单合并的第{index}项：{major}（{code}，{year}年度）。"
                ),
                "confidence": "high",
            }
        )
    return rows


def parse_aqnu_2025_sports_training_application_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "aqnu_2025_sports_training_application_pdf"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    stop_rows = [
        ("教育学", "040101"),
        ("人文教育", "040103"),
        ("经济统计学", "020102"),
        ("文化产业管理", "120210"),
    ]
    cancel_rows = [
        (2024, "动植物检疫", "090403T"),
        (2024, "艺术教育", "040105"),
        (2024, "统计学", "071201"),
        (2023, "国际经济与贸易", "020401H"),
        (2023, "英语", "050201H"),
        (2022, "信息管理与信息系统", "120102"),
        (2022, "农村区域发展", "120302"),
        (2022, "酒店管理", "120902"),
        (2022, "休闲体育", "040207T"),
    ]
    required_fragments = [
        "专业代码040202K专业名称运动训练",
        "校停招了以下专业：教育学、人文教育、经济统计学、文化产业管理",
        "2024年撤销了动植物检疫、艺术教育、统计学",
        "2023年撤销了国际经济与贸易（中外合作）、英语（中外合作）",
        "2022年撤销了信息管理与信息系统、农村区域发展、酒店管理、休闲体育",
    ]
    missing = [fragment for fragment in required_fragments if fragment not in compact_text]
    if missing:
        raise ValueError(f"Could not verify AQNU professional-history fragments in {source_id}: {missing}")

    rows: list[dict[str, Any]] = []
    for index, (major, code) in enumerate(stop_rows, start=1):
        rows.append(
            {
                "policy_year": 2025,
                "region": "安庆师范大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "安庆师范大学2025年专业申报材料近五年停招本科专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(index),
                "study_duration": "",
                "policy_action": "学校专业申报材料列为近五年已停招本科专业。",
                "criterion_text": "安庆师范大学运动训练专业申请表“学校近五年专业增设、停招、撤并情况”中的停招专业；来源未给出逐项停招年份，policy_year采用申请材料年份。",
                "source_ids": source_id,
                "evidence_text": (
                    "安庆师范大学运动训练专业申请表说明，近五年学校停招教育学、人文教育、经济统计学、"
                    "文化产业管理4个专业；"
                    f"本行对应第{index}项：{major}（{code}）。"
                ),
                "confidence": "high",
            }
        )
    for index, (year, major, code) in enumerate(cancel_rows, start=1):
        rows.append(
            {
                "policy_year": year,
                "region": "安庆师范大学",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": f"安庆师范大学{year}年度本科专业备案审批结果撤销专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": "",
                "study_duration": "",
                "policy_action": "学校专业申报材料逐年列为近五年撤销本科专业；教育部年度结果公布撤销。",
                "criterion_text": "安庆师范大学运动训练专业申请表“学校近五年专业增设、停招、撤并情况”中的撤销专业。",
                "source_ids": source_id,
                "evidence_text": (
                    "安庆师范大学运动训练专业申请表说明，2024年撤销动植物检疫、艺术教育、统计学，"
                    "2023年撤销国际经济与贸易（中外合作）、英语（中外合作），2022年撤销信息管理与信息系统、"
                    "农村区域发展、酒店管理、休闲体育；"
                    f"本行对应第{index}项：{major}（{code}，{year}年）。"
                ),
                "confidence": "high",
            }
        )
    return rows


def parse_yxu_2025_application_history_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_ids = [
        "yxu_2025_smart_energy_application_pdf",
        "yxu_2025_biological_breeding_application_pdf",
    ]
    text_parts: list[str] = []
    for source_id in source_ids:
        text_path = raw_dir / f"{source_id}.txt"
        if text_path.exists():
            text_parts.append(text_path.read_text(encoding="utf-8", errors="replace"))
    if not text_parts:
        return []
    compact_text = re.sub(r"\s+", "", "\n".join(text_parts))
    validation_text = compact_text.replace("学校近五年专业增设、", "").replace("停招、撤并情况", "")
    source_id_text = ";".join(source_ids)
    cancel_rows = [
        (2019, "农业资源与环境", "090201"),
        (2019, "绘画", "130402"),
        (2019, "舞蹈表演", "130204"),
        (2019, "工艺美术", "130507"),
        (2020, "书法学", "130405T"),
        (2020, "老挝语", "050215"),
        (2023, "市场营销", "120202"),
        (2023, "信息管理与信息系统", "120102"),
        (2023, "测绘工程", "081201"),
        (2023, "电子信息工程", "080701"),
        (2023, "国际经济与贸易", "020401"),
        (2023, "教育技术学", "040104"),
        (2023, "电子商务", "120801"),
    ]
    stop_rows = [
        ("信息与计算科学", "070102"),
        ("汽车服务工程", "080208"),
        ("汉语国际教育", "050103"),
        ("自然地理与资源环境", "070502"),
        ("自动化", "080801"),
        ("环境科学", "082503"),
        ("财务管理", "120204"),
    ]
    required_fragments = [
        "专业名称智慧能源工程",
        "专业名称生物育种科学",
        "近五年学校撤销专业：农业资源与环境、绘画、舞蹈表演、工艺美术、书法学、老挝语、市场营销",
        "信息管理与信息系统、测绘工程、电子信息工程、国际经济与贸易、教育技术学、电子商务",
        "停招专业：信息与计算科学、汽车服务工程、汉语国际教育、自然地理与资源环境、自动化、环境科学、财务管理",
    ]
    missing = [fragment for fragment in required_fragments if fragment not in validation_text]
    missing.extend(major for _year, major, _code in cancel_rows if major not in validation_text)
    missing.extend(major for major, _code in stop_rows if major not in validation_text)
    if missing:
        raise ValueError(f"Could not verify YXU professional-history fragments in {source_id_text}: {missing}")

    rows: list[dict[str, Any]] = []
    for index, (year, major, code) in enumerate(cancel_rows, start=1):
        rows.append(
            {
                "policy_year": year,
                "region": "玉溪师范学院",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": f"玉溪师范学院{year}年度本科专业备案审批结果撤销专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": "",
                "study_duration": "",
                "policy_action": "学校专业申报材料列为近五年撤销专业；教育部年度结果公布撤销。",
                "criterion_text": "玉溪师范学院2025年度专业申请表“学校近五年专业增设、停招、撤并情况”中的撤销专业；policy_year按可与教育部年度撤销名单精确对应的结果年份记录。",
                "source_ids": source_id_text,
                "evidence_text": (
                    "玉溪师范学院智慧能源工程和生物育种科学专业申请表均说明，近五年学校撤销农业资源与环境、"
                    "绘画、舞蹈表演、工艺美术、书法学、老挝语、市场营销、信息管理与信息系统、测绘工程、"
                    "电子信息工程、国际经济与贸易、教育技术学、电子商务；"
                    f"本行对应可与教育部年度撤销名单合并的第{index}项：{major}（{code}，{year}年度）。"
                ),
                "confidence": "high",
            }
        )
    for index, (major, code) in enumerate(stop_rows, start=1):
        rows.append(
            {
                "policy_year": 2025,
                "region": "玉溪师范学院",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "玉溪师范学院2025年专业申报材料近五年停招本科专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(index),
                "study_duration": "",
                "policy_action": "学校专业申报材料列为近五年已停招本科专业。",
                "criterion_text": "玉溪师范学院2025年度专业申请表“学校近五年专业增设、停招、撤并情况”中的停招专业；来源未给出逐项停招年份，policy_year采用申请材料年份。",
                "source_ids": source_id_text,
                "evidence_text": (
                    "玉溪师范学院智慧能源工程和生物育种科学专业申请表均说明，近五年学校停招信息与计算科学、"
                    "汽车服务工程、汉语国际教育、自然地理与资源环境、自动化、环境科学、财务管理；"
                    f"本行对应第{index}项：{major}（{code}）。"
                ),
                "confidence": "high",
            }
        )
    return rows


def parse_zcmu_2025_nursing_application_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "zcmu_2025_nursing_application_pdf"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    validation_text = compact_text.replace("学校近五年专业增设、", "").replace("停招、撤并情况", "")
    stop_rows = [
        (2023, "市场营销", "120202"),
        (2025, "英语", "050201"),
        (2025, "食品卫生与营养学", "100402"),
        (2025, "数据科学与大数据技术", "080910T"),
        (2025, "生物技术", "071002"),
        (2025, "医学影像技术", "101003"),
    ]
    cancel_rows = [
        (2024, "药物制剂", "100702"),
    ]
    required_fragments = [
        "专业代码101101KH专业名称护理学",
        "2023年增设临床药学、国际商务专业",
        "2023年停招市场营销专业",
        "2025年停招英语、食品卫生与营养学、数据科学与大数据技术、生物技术和医学影像技术专业",
        "2025年撤销药物制剂专业",
    ]
    missing = [fragment for fragment in required_fragments if fragment not in validation_text]
    missing.extend(major for _year, major, _code in stop_rows if major not in validation_text)
    missing.extend(major for _year, major, _code in cancel_rows if major not in validation_text)
    if missing:
        raise ValueError(f"Could not verify ZCMU professional-history fragments in {source_id}: {missing}")

    rows: list[dict[str, Any]] = []
    for index, (year, major, code) in enumerate(stop_rows, start=1):
        rows.append(
            {
                "policy_year": year,
                "region": "浙江中医药大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": f"浙江中医药大学{year}年专业申报材料停招本科专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(index),
                "study_duration": "",
                "policy_action": "学校专业申报材料列为近五年停招本科专业。",
                "criterion_text": "浙江中医药大学护理学专业申请表“学校近五年专业增设、停招、撤并情况”中的停招专业。",
                "source_ids": source_id,
                "evidence_text": (
                    "浙江中医药大学护理学专业申请表说明，2023年停招市场营销专业，2025年停招英语、"
                    "食品卫生与营养学、数据科学与大数据技术、生物技术和医学影像技术专业；"
                    f"本行对应第{index}项：{major}（{code}，{year}年）。"
                ),
                "confidence": "high",
            }
        )
    for index, (year, major, code) in enumerate(cancel_rows, start=1):
        rows.append(
            {
                "policy_year": year,
                "region": "浙江中医药大学",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": f"浙江中医药大学{year}年度本科专业备案审批结果撤销专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": "",
                "study_duration": "",
                "policy_action": "学校专业申报材料列为2025年撤销专业；教育部2024年度结果公布撤销。",
                "criterion_text": "浙江中医药大学护理学专业申请表“学校近五年专业增设、停招、撤并情况”中的撤销专业；policy_year按可与教育部年度撤销名单精确对应的结果年份记录。",
                "source_ids": source_id,
                "evidence_text": (
                    "浙江中医药大学护理学专业申请表说明，2025年撤销药物制剂专业；"
                    f"本行对应可与教育部2024年度撤销名单合并的第{index}项：{major}（{code}）。"
                ),
                "confidence": "high",
            }
        )
    return rows


def parse_cxtc_2025_uav_application_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "cxtc_2025_uav_smart_network_application_pdf"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    cancel_rows = [
        (2019, "应用物理学", "070202"),
        (2019, "运动康复", "040206T"),
        (2022, "政治学与行政学", "030201"),
        (2022, "统计学", "071201"),
        (2022, "雕塑", "130403"),
        (2023, "翻译", "050261"),
        (2024, "信息管理与信息系统", "120102"),
    ]
    stop_rows = [
        ("中国少数民族语言文学", "050104", ["中国少数民族语言文学"]),
        ("城乡规划", "082802", ["城乡规划"]),
        ("公共事业管理", "120401", ["公共事业管", "理、音乐表演"]),
        ("音乐表演", "130201", ["音乐表演"]),
    ]
    required_fragments = [
        "专业名称无人机与智能网络",
        "近五年增设专业情况：增设农学、食用菌科学与工程、华文教育、人工智能4个本科专业",
        "近五年停招专业情况：停招中国少数民族语言文学、城乡规划、公共事业管",
        "理、音乐表演等11个本科专业",
        "近五年撤并专业情况：撤销应用物理学、运动康复、政治学与行政学、统计学、雕塑、翻译、信息管理与信息系统7个专业",
    ]
    missing = [fragment for fragment in required_fragments if fragment not in compact_text]
    for major, _code, fragments in stop_rows:
        missing.extend(fragment for fragment in fragments if fragment not in compact_text)
    missing.extend(major for _year, major, _code in cancel_rows if major not in compact_text)
    if missing:
        raise ValueError(f"Could not verify CXTC professional-history fragments in {source_id}: {missing}")

    rows: list[dict[str, Any]] = []
    for index, (major, code, _fragments) in enumerate(stop_rows, start=1):
        rows.append(
            {
                "policy_year": 2025,
                "region": "楚雄师范学院",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "楚雄师范学院2025年专业申报材料近五年停招本科专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(index),
                "study_duration": "",
                "policy_action": "学校专业申报材料列为近五年已停招本科专业。",
                "criterion_text": "楚雄师范学院无人机与智能网络专业申请表“学校近五年专业增设、停招、撤并情况”中的停招专业；来源未给出逐项停招年份，policy_year采用申请材料年份。",
                "source_ids": source_id,
                "evidence_text": (
                    "楚雄师范学院无人机与智能网络专业申请表说明，近五年停招专业情况包括停招中国少数民族语言文学、"
                    "城乡规划、公共事业管理、音乐表演等11个本科专业；"
                    f"本行对应其中明确点名的第{index}项：{major}（{code}）。"
                ),
                "confidence": "high",
            }
        )
    for index, (year, major, code) in enumerate(cancel_rows, start=1):
        rows.append(
            {
                "policy_year": year,
                "region": "楚雄师范学院",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": f"楚雄师范学院{year}年度本科专业备案审批结果撤销专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": "",
                "study_duration": "",
                "policy_action": "学校专业申报材料列为近五年撤销专业；教育部年度结果公布撤销。",
                "criterion_text": "楚雄师范学院无人机与智能网络专业申请表“学校近五年专业增设、停招、撤并情况”中的撤销专业；policy_year按可与教育部年度撤销名单精确对应的结果年份记录。",
                "source_ids": source_id,
                "evidence_text": (
                    "楚雄师范学院无人机与智能网络专业申请表说明，近五年撤并专业情况包括撤销应用物理学、"
                    "运动康复、政治学与行政学、统计学、雕塑、翻译、信息管理与信息系统7个专业；"
                    f"本行对应可与教育部年度撤销名单合并的第{index}项：{major}（{code}，{year}年度）。"
                ),
                "confidence": "high",
            }
        )
    return rows


def parse_wust_2025_fintech_application_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "wust_2025_fintech_application_pdf"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    stop_rows = [
        ("人文地理与城乡规划", "070503", ["人文地理与城乡规划"]),
        ("汽车服务工程", "080208", ["汽车服务工程"]),
        ("电子商务", "120801", ["电子商务"]),
        ("交通运输", "081801", ["交通", "运输、马克思主义理论"]),
        ("马克思主义理论", "030504T", ["马克思主义理论"]),
        ("光电信息科学与工程", "080705", ["光电信息科学与工程"]),
    ]
    required_fragments = [
        "专业代码020310T专业名称金融科技",
        "近五年停招专业：人文地理与城乡规划、汽车服务工程、电子商务、交通",
        "运输、马克思主义理论、光电信息科学与工程专业及10个第二学士学位专业",
        "近五年撤销专业：风景园林",
    ]
    missing = [fragment for fragment in required_fragments if fragment not in compact_text]
    for _major, _code, fragments in stop_rows:
        missing.extend(fragment for fragment in fragments if fragment not in compact_text)
    if missing:
        raise ValueError(f"Could not verify WUST fintech application fragments in {source_id}: {missing}")

    rows: list[dict[str, Any]] = []
    for index, (major, code, _fragments) in enumerate(stop_rows, start=1):
        rows.append(
            {
                "policy_year": 2025,
                "region": "武汉科技大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "武汉科技大学2025年专业申报材料近五年停招本科专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(index),
                "study_duration": "",
                "policy_action": "学校专业申报材料列为近五年已停招本科专业。",
                "criterion_text": "武汉科技大学金融科技专业备案申请表“学校近五年专业增设、停招、撤并情况”中的停招专业；来源未给出逐项停招年份，policy_year采用申请材料年份。",
                "source_ids": source_id,
                "evidence_text": (
                    "武汉科技大学金融科技专业备案申请表说明，近五年停招专业包括人文地理与城乡规划、"
                    "汽车服务工程、电子商务、交通运输、马克思主义理论、光电信息科学与工程专业及10个第二学士学位专业；"
                    f"本行对应第{index}项：{major}（{code}）。"
                ),
                "confidence": "high",
            }
        )
    rows.append(
        {
            "policy_year": 2020,
            "region": "武汉科技大学",
            "education_level": "本科",
            "record_type": "major_cancel",
            "warning_label": "武汉科技大学近五年撤销本科专业",
            "reported_major_name": "风景园林",
            "major_code": "082803",
            "source_row_no": "",
            "study_duration": "",
            "policy_action": "学校专业申报材料列为近五年撤销专业；教育部年度结果公布撤销。",
            "criterion_text": "武汉科技大学金融科技专业备案申请表“学校近五年专业增设、停招、撤并情况”中的撤销专业；policy_year按可与教育部年度撤销名单精确对应的结果年份记录。",
            "source_ids": source_id,
            "evidence_text": "武汉科技大学金融科技专业备案申请表说明，近五年撤销专业为风景园林；本行与教育部2020年度撤销结果中的风景园林专业对应。",
            "confidence": "high",
        }
    )
    return rows


def parse_bigc_2025_digital_economy_application_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "bigc_2025_digital_economy_application_pdf"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    stop_rows = [
        ("摄影", "130404", ["停招摄影"]),
        ("电子信息工程", "080701", ["电子信息工程"]),
        ("物流管理", "120601", ["物流管", "理、物流工程"]),
        ("物流工程", "120602", ["物流工程"]),
        ("工业设计", "080205", ["工业设计"]),
        ("市场营销", "120202", ["市场营销"]),
    ]
    required_fragments = [
        "专业代码020109T专业名称数字经济",
        "近五年增设产品设计、智能制造工程、新媒体艺术、大数据管理",
        "停招摄影、电子信息工程、物流管",
        "理、物流工程、工业设计、市场营销等专业",
    ]
    missing = [fragment for fragment in required_fragments if fragment not in compact_text]
    for _major, _code, fragments in stop_rows:
        missing.extend(fragment for fragment in fragments if fragment not in compact_text)
    if missing:
        raise ValueError(f"Could not verify BIGC digital economy application fragments in {source_id}: {missing}")

    rows: list[dict[str, Any]] = []
    for index, (major, code, _fragments) in enumerate(stop_rows, start=1):
        rows.append(
            {
                "policy_year": 2025,
                "region": "北京印刷学院",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "北京印刷学院2025年专业申报材料近五年停招本科专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(index),
                "study_duration": "",
                "policy_action": "学校专业申报材料列为近五年已停招本科专业。",
                "criterion_text": "北京印刷学院数字经济专业备案申请表“学校近五年专业增设、停招、撤并情况”中的停招专业；来源未给出逐项停招年份，policy_year采用申请材料年份。",
                "source_ids": source_id,
                "evidence_text": (
                    "北京印刷学院数字经济专业备案申请表说明，近五年停招摄影、电子信息工程、"
                    "物流管理、物流工程、工业设计、市场营销等专业；"
                    f"本行对应第{index}项：{major}（{code}）。"
                ),
                "confidence": "high",
            }
        )
    return rows


def parse_hbkjxy_2025_cybersecurity_application_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "hbkjxy_2025_cybersecurity_application_pdf"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    stop_rows = [
        (2020, "风景园林", "082803"),
        (2020, "建筑电气与智能化", "081004"),
        (2020, "机械电子工程", "080204"),
        (2020, "工业设计", "080205"),
        (2020, "投资学", "020304"),
        (2021, "建筑电气与智能化", "081004"),
        (2021, "风景园林", "082803"),
        (2021, "自动化", "080801"),
        (2021, "投资学", "020304"),
        (2021, "机械电子工程", "080204"),
        (2022, "风景园林", "082803"),
        (2022, "自动化", "080801"),
        (2022, "投资学", "020304"),
        (2022, "机械电子工程", "080204"),
        (2022, "表演", "130301"),
        (2023, "投资学", "020304"),
        (2023, "表演", "130301"),
    ]
    cancel_rows = [
        (2021, "建筑电气与智能化", "081004"),
        (2023, "风景园林", "082803"),
        (2024, "投资学", "020304"),
    ]
    required_fragments = [
        "专业代码080911TK专业名称网络空间安全",
        "学校近五年停招：2020年暂停风景园林、建筑电气与智能化、机械电子工程",
        "工业设计、投资学；2021年暂停建筑电气与智能化、风景园林、自动化、投资学、机械电子工程",
        "2022年暂停风景园林、自动化、投资学、机械电子工程、表演",
        "2023年暂停投资学、表演；2024年无暂停专业",
        "学校近五年撤销：2021年撤销建筑电气与智能化；2023年撤销风景园林；2024年撤销投资学",
    ]
    missing = [fragment for fragment in required_fragments if fragment not in compact_text]
    missing.extend(major for _year, major, _code in stop_rows + cancel_rows if major not in compact_text)
    if missing:
        raise ValueError(f"Could not verify HBKJXY cybersecurity application fragments in {source_id}: {missing}")

    rows: list[dict[str, Any]] = []
    for index, (year, major, code) in enumerate(stop_rows, start=1):
        rows.append(
            {
                "policy_year": year,
                "region": "河北科技学院",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": f"河北科技学院{year}年专业申报材料停招本科专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(index),
                "study_duration": "",
                "policy_action": f"{year}年暂停招生。",
                "criterion_text": "河北科技学院网络空间安全专业备案申请表“学校近五年专业增设、停招、撤并情况”逐年列出暂停招生专业。",
                "source_ids": source_id,
                "evidence_text": (
                    "河北科技学院网络空间安全专业备案申请表说明，2020年暂停风景园林、建筑电气与智能化、"
                    "机械电子工程、工业设计、投资学；2021年暂停建筑电气与智能化、风景园林、自动化、"
                    "投资学、机械电子工程；2022年暂停风景园林、自动化、投资学、机械电子工程、表演；"
                    "2023年暂停投资学、表演；2024年无暂停专业；"
                    f"本行对应第{index}项：{major}（{code}，{year}年）。"
                ),
                "confidence": "high",
            }
        )
    for index, (year, major, code) in enumerate(cancel_rows, start=1):
        rows.append(
            {
                "policy_year": year,
                "region": "河北科技学院",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": f"河北科技学院{year}年专业申报材料撤销本科专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": "",
                "study_duration": "",
                "policy_action": "学校专业申报材料列为近五年撤销专业；教育部年度结果公布撤销。",
                "criterion_text": "河北科技学院网络空间安全专业备案申请表“学校近五年专业增设、停招、撤并情况”中的撤销专业；policy_year按可与教育部年度撤销名单精确对应的结果年份记录。",
                "source_ids": source_id,
                "evidence_text": (
                    "河北科技学院网络空间安全专业备案申请表说明，学校近五年撤销专业包括2021年撤销建筑电气与智能化、"
                    "2023年撤销风景园林、2024年撤销投资学；"
                    f"本行对应第{index}项：{major}（{code}，{year}年）。"
                ),
                "confidence": "high",
            }
        )
    return rows


def parse_whwl_2025_ai_education_application_rows(raw_dir: Path) -> list[dict[str, Any]]:
    stop_rows = [
        ("旅游管理", "120901"),
        ("互联网金融", "020309"),
        ("投资学", "020304"),
        ("艺术教育", "040105"),
        ("日语", "050207"),
        ("工业智能", "080807"),
    ]
    cancel_rows = [
        ("环境工程", "082502"),
        ("园艺", "090102"),
        ("会展经济与管理", "120903"),
        ("测控技术与仪器", "080301"),
        ("材料成型及控制工程", "080203"),
        ("自动化", "080801"),
        ("生物技术", "071002"),
        ("公共事业管理", "120401"),
        ("酒店管理", "120902"),
    ]
    source_checks = [
        (
            "whwl_2025_ai_education_application_pdf",
            ["专业代码：040117TK", "专业名称：人工智能教育"],
        ),
        (
            "whwl_2025_medical_device_application_pdf",
            ["专业代码101015TK", "专业名称医疗器械与装备工程"],
        ),
    ]
    common_fragments = [
        "武汉文理学院",
        "学校近五年停招专业：旅游管理、互联网金融、投资学、艺术教育、日语、工业智能",
        "学校近五年撤销专业：环境工程、园艺、会展经济与管理、测控技术与仪器、材料成型及控制工程、自动化、生物技术、公共事业管理、酒店管理",
    ]
    verified_source_ids: list[str] = []
    for source_id, source_fragments in source_checks:
        text_path = raw_dir / f"{source_id}.txt"
        if not text_path.exists():
            continue
        text = text_path.read_text(encoding="utf-8", errors="replace")
        compact_text = re.sub(r"\s+", "", text)
        required_fragments = common_fragments + source_fragments
        missing = [fragment for fragment in required_fragments if fragment not in compact_text]
        missing.extend(major for major, _code in stop_rows + cancel_rows if major not in compact_text)
        if missing:
            raise ValueError(f"Could not verify WHWL application fragments in {source_id}: {missing}")
        verified_source_ids.append(source_id)
    if not verified_source_ids:
        return []
    source_ids = ";".join(verified_source_ids)

    rows: list[dict[str, Any]] = []
    for index, (major, code) in enumerate(stop_rows, start=1):
        rows.append(
            {
                "policy_year": 2025,
                "region": "武汉文理学院",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "武汉文理学院2025年专业申报材料近五年停招本科专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(index),
                "study_duration": "",
                "policy_action": "学校专业申报材料列为近五年已停招本科专业。",
                "criterion_text": "武汉文理学院人工智能教育专业申请表“学校近五年专业增设、停招、撤并情况”中的停招专业；来源未给出逐项停招年份，policy_year采用申请材料年份。",
                "source_ids": source_ids,
                "evidence_text": (
                    "武汉文理学院专业申请表说明，学校近五年停招专业包括旅游管理、互联网金融、"
                    "投资学、艺术教育、日语、工业智能；"
                    f"本行对应第{index}项：{major}（{code}）。"
                ),
                "confidence": "high",
            }
        )
    for index, (major, code) in enumerate(cancel_rows, start=1):
        rows.append(
            {
                "policy_year": 2025,
                "region": "武汉文理学院",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "武汉文理学院2025年专业申报材料近五年撤销本科专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(index),
                "study_duration": "",
                "policy_action": "学校专业申报材料列为近五年撤销本科专业。",
                "criterion_text": "武汉文理学院人工智能教育专业申请表“学校近五年专业增设、停招、撤并情况”中的撤销专业；来源未给出逐项撤销年份，policy_year采用申请材料年份。",
                "source_ids": source_ids,
                "evidence_text": (
                    "武汉文理学院专业申请表说明，学校近五年撤销专业包括环境工程、园艺、"
                    "会展经济与管理、测控技术与仪器、材料成型及控制工程、自动化、生物技术、公共事业管理、酒店管理；"
                    f"本行对应第{index}项：{major}（{code}）。"
                ),
                "confidence": "high",
            }
        )
    return rows


def parse_muc_2025_sports_training_application_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "muc_2025_sports_training_application_pdf"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    required_fragments = [
        "学校名称中央民族大学",
        "专业代码040202K专业名称运动训练",
        "停招朝鲜语等10个专业",
        "朝鲜语",
    ]
    missing = [fragment for fragment in required_fragments if fragment not in compact_text]
    if missing:
        raise ValueError(f"Could not verify MUC sports training application fragments in {source_id}: {missing}")
    return [
        {
            "policy_year": 2025,
            "region": "中央民族大学",
            "education_level": "本科",
            "record_type": "major_stop_enrollment",
            "warning_label": "中央民族大学2025年专业申报材料近五年停招本科专业",
            "reported_major_name": "朝鲜语",
            "major_code": "050209",
            "source_row_no": "1",
            "study_duration": "",
            "policy_action": "学校专业申报材料列为停招朝鲜语等10个专业。",
            "criterion_text": "中央民族大学运动训练专业申请表“学校近五年专业增设、停招、撤并情况”中说明停招朝鲜语等10个专业；本数据集只结构化来源直接点名的朝鲜语，来源未给出逐项停招年份，policy_year采用申请材料年份。",
            "source_ids": source_id,
            "evidence_text": "中央民族大学运动训练专业申请表说明，学校近五年新增数字经济学、人工智能、音乐教育、网络新媒体、考古学等专业，停招朝鲜语等10个专业；本行对应直接点名的朝鲜语（050209）。",
            "confidence": "high",
        }
    ]


def parse_sus_2025_football_application_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "sus_2025_football_application_pdf"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    stop_rows = [
        ("特殊教育", "040108", "未招生专业"),
        ("艺术教育", "040105", "停招并撤销专业"),
        ("公共事业管理", "120401", "停招并撤销专业"),
        ("信息管理与信息系统", "120102", "停招并撤销专业"),
    ]
    cancel_rows = [
        (2023, "艺术教育", "040105"),
        (2023, "公共事业管理", "120401"),
        (2023, "信息管理与信息系统", "120102"),
    ]
    required_fragments = [
        "学校名称上海体育大学",
        "专业名称足球运动",
        "未招生专业：特殊教育",
        "停招并撤销专业：艺术教育、公共事业管理、信息管理与信息系统3个专业",
    ]
    missing = [fragment for fragment in required_fragments if fragment not in compact_text]
    missing.extend(major for major, _code, _action in stop_rows if major not in compact_text)
    if missing:
        raise ValueError(f"Could not verify SUS football application fragments in {source_id}: {missing}")

    rows: list[dict[str, Any]] = []
    for index, (major, code, action) in enumerate(stop_rows, start=1):
        rows.append(
            {
                "policy_year": 2025,
                "region": "上海体育大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "上海体育大学2025年专业申报材料停招或未招生本科专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(index),
                "study_duration": "",
                "policy_action": f"学校专业申报材料列为{action}。",
                "criterion_text": "上海体育大学足球运动专业申请表“学校近五年专业增设、停招、撤并情况”列出未招生专业和停招并撤销专业；来源未给出逐项停招年份，policy_year采用申请材料年份。",
                "source_ids": source_id,
                "evidence_text": (
                    "上海体育大学足球运动专业申请表说明，未招生专业为特殊教育；停招并撤销专业为艺术教育、"
                    "公共事业管理、信息管理与信息系统3个专业；"
                    f"本行对应第{index}项：{major}（{code}）。"
                ),
                "confidence": "high",
            }
        )
    for year, major, code in cancel_rows:
        rows.append(
            {
                "policy_year": year,
                "region": "上海体育大学",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": f"上海体育大学{year}年度本科专业备案审批结果撤销专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": "",
                "study_duration": "",
                "policy_action": "学校专业申报材料列为停招并撤销专业；教育部年度结果公布撤销。",
                "criterion_text": "上海体育大学足球运动专业申请表列为停招并撤销专业；policy_year按教育部年度撤销名单中可精确对应的结果年份记录。",
                "source_ids": source_id,
                "evidence_text": (
                    "上海体育大学足球运动专业申请表说明，停招并撤销专业为艺术教育、公共事业管理、"
                    f"信息管理与信息系统3个专业；本行对应{major}（{code}，{year}年度撤销结果）。"
                ),
                "confidence": "high",
            }
        )
    return rows


def parse_qtnu_2025_sports_training_application_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "qtnu_2025_sports_training_application_pdf"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    required_fragments = [
        "学校名称琼台师范学院",
        "专业代码040202K专业名称运动训练",
        "2023年、2024年、2025年连续三年停招艺术设计学专业",
    ]
    missing = [fragment for fragment in required_fragments if fragment not in compact_text]
    if missing:
        raise ValueError(f"Could not verify QTNU sports training application fragments in {source_id}: {missing}")

    rows: list[dict[str, Any]] = []
    for index, year in enumerate([2023, 2024, 2025], start=1):
        rows.append(
            {
                "policy_year": year,
                "region": "琼台师范学院",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": f"琼台师范学院{year}年连续停招本科专业",
                "reported_major_name": "艺术设计学",
                "major_code": "130501",
                "source_row_no": str(index),
                "study_duration": "",
                "policy_action": f"{year}年停招艺术设计学专业。",
                "criterion_text": "琼台师范学院运动训练专业申请表说明，依据地方经济社会发展需要和人才需求，2023年、2024年、2025年连续三年停招艺术设计学专业。",
                "source_ids": source_id,
                "evidence_text": f"琼台师范学院运动训练专业申请表列明2023年、2024年、2025年连续三年停招艺术设计学专业；本行对应{year}年艺术设计学（130501）。",
                "confidence": "high",
            }
        )
    return rows


def parse_gmc_2025_geriatric_medicine_application_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "gmc_2025_geriatric_medicine_application_pdf"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    required_fragments = [
        "学校名称贵州医科大学",
        "专业名称老年医学与健康",
        "2023年，暂停翻译专业招生",
        "2024年，暂停英语和翻译等2个本科专业招生",
    ]
    missing = [fragment for fragment in required_fragments if fragment not in compact_text]
    if missing:
        raise ValueError(f"Could not verify GMC geriatric medicine application fragments in {source_id}: {missing}")

    rows: list[dict[str, Any]] = []
    for index, (year, major, code) in enumerate([(2023, "翻译", "050261"), (2024, "英语", "050201"), (2024, "翻译", "050261")], start=1):
        rows.append(
            {
                "policy_year": year,
                "region": "贵州医科大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": f"贵州医科大学{year}年专业申报材料停招本科专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(index),
                "study_duration": "",
                "policy_action": f"{year}年暂停{major}专业招生。",
                "criterion_text": "贵州医科大学老年医学与健康专业申请表“学校近五年专业增设、停招、撤并情况”逐年说明暂停招生专业。",
                "source_ids": source_id,
                "evidence_text": "贵州医科大学老年医学与健康专业申请表说明，2023年暂停翻译专业招生；2024年暂停英语和翻译等2个本科专业招生。",
                "confidence": "high",
            }
        )
    return rows


def parse_gnnu_2025_ai_education_application_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "gnnu_2025_ai_education_application_pdf"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    required_fragments = [
        "学校名称赣南师范大学",
        "专业代码040117TK专业名称人工智能教育",
        "2022年增设区块链工程和数字媒体艺术2个专业，停招环境设计等23个专业",
        "撤销广告学等10个专业",
        "2023年停招表演等10个专业，撤销软件工程1个专业",
        "2025年增设数字人文等5个专业，撤销秘书学等2个专业，停招区块链工程等8个专业",
    ]
    missing = [fragment for fragment in required_fragments if fragment not in compact_text]
    if missing:
        raise ValueError(f"Could not verify GNNU AI education application fragments in {source_id}: {missing}")

    specs = [
        (2022, "major_stop_enrollment", "环境设计", "130503", "停招环境设计等23个专业"),
        (2022, "major_cancel", "广告学", "050303", "撤销广告学等10个专业"),
        (2023, "major_stop_enrollment", "表演", "130301", "停招表演等10个专业"),
        (2023, "major_cancel", "软件工程", "080902", "撤销软件工程1个专业"),
        (2025, "major_cancel", "秘书学", "050107", "撤销秘书学等2个专业"),
        (2025, "major_stop_enrollment", "区块链工程", "080917", "停招区块链工程等8个专业"),
    ]
    rows: list[dict[str, Any]] = []
    for index, (year, record_type, major, code, action) in enumerate(specs, start=1):
        rows.append(
            {
                "policy_year": year,
                "region": "赣南师范大学",
                "education_level": "本科",
                "record_type": record_type,
                "warning_label": f"赣南师范大学{year}年专业申报材料专业调整本科专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(index),
                "study_duration": "",
                "policy_action": action + "。",
                "criterion_text": "赣南师范大学人工智能教育专业申请表“学校近五年专业增设、停招、撤并情况”逐年说明停招和撤销专业；本数据集只结构化原文直接点名的专业。",
                "source_ids": source_id,
                "evidence_text": (
                    "赣南师范大学人工智能教育专业申请表说明，2022年停招环境设计等23个专业并撤销广告学等10个专业；"
                    "2023年停招表演等10个专业并撤销软件工程1个专业；2025年撤销秘书学等2个专业并停招区块链工程等8个专业。"
                ),
                "confidence": "high",
            }
        )
    return rows


def parse_hbesxy_2025_stomatology_application_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "hbesxy_2025_stomatology_application_pdf"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    required_fragments = [
        "学校名称湖北恩施学院",
        "专业代码100301K专业名称口腔医学",
        "2018年40个本科专业有通信工程等22个专业停招",
        "2019年45个本科专业中有社会学等21个专业停招",
        "2020年45个本科专业中有社会学等26个专业停招",
        "2021年48个本科专业中有社会学等25个专业停招",
        "2022年49个本科专业有社会学等24个专业停招",
        "无撤并情况",
    ]
    missing = [fragment for fragment in required_fragments if fragment not in compact_text]
    if missing:
        raise ValueError(f"Could not verify HBESXY stomatology application fragments in {source_id}: {missing}")

    specs = [(2018, "通信工程", "080703")] + [(year, "社会学", "030301") for year in [2019, 2020, 2021, 2022]]
    rows: list[dict[str, Any]] = []
    for index, (year, major, code) in enumerate(specs, start=1):
        rows.append(
            {
                "policy_year": year,
                "region": "湖北恩施学院",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": f"湖北恩施学院{year}年专业申报材料停招本科专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(index),
                "study_duration": "",
                "policy_action": f"{year}年列入停招专业示例。",
                "criterion_text": "湖北恩施学院口腔医学专业申请表“学校近五年专业增设、停招、撤并情况”逐年说明停招专业数量；本数据集只结构化原文直接点名的专业。",
                "source_ids": source_id,
                "evidence_text": "湖北恩施学院口腔医学专业申请表说明，2018年通信工程等22个专业停招，2019-2022年社会学等专业连续被列为停招示例，且无撤并情况。",
                "confidence": "high",
            }
        )
    return rows


def parse_nwupl_2025_forensic_science_application_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "nwupl_2025_forensic_science_application_pdf"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    required_fragments = [
        "学校名称西北政法大学",
        "专业代码030110TK专业名称司法鉴定学",
        "2021年社会工作、治安学、商务英语、广播电视学、公共事业管理、朝鲜语、经济统计学、思想政治教育、信息管理与信息系统等8个专业未招生",
        "2022年经济统计学、社会学、治安学、商务英语、广播电视学、朝鲜语",
        "劳动与社会保障等7个专业未招生",
        "撤销思想政治教育、信息管理与信息系",
        "2023年社会工作、广播电视学、朝鲜语、商务英语、公共事业管理等5个专业未招生",
        "撤销经济统计学",
        "2024年治安学、市场营销、公共事业管理、商务英语专业不招生",
        "撤销广播电视学、朝鲜语",
        "2025年设置监狱学，撤销市场营销，治安学、公共事业管理、商务英语、监狱学未招生",
    ]
    missing = [fragment for fragment in required_fragments if fragment not in compact_text]
    if missing:
        raise ValueError(f"Could not verify NWUPL forensic science application fragments in {source_id}: {missing}")

    stop_specs = [
        (2021, "社会工作", "030302"),
        (2021, "治安学", "030601"),
        (2021, "商务英语", "050262"),
        (2021, "广播电视学", "050302"),
        (2021, "公共事业管理", "120401"),
        (2021, "朝鲜语", "050209"),
        (2021, "经济统计学", "020102"),
        (2021, "思想政治教育", "030503"),
        (2021, "信息管理与信息系统", "120102"),
        (2022, "经济统计学", "020102"),
        (2022, "社会学", "030301"),
        (2022, "治安学", "030601"),
        (2022, "商务英语", "050262"),
        (2022, "广播电视学", "050302"),
        (2022, "朝鲜语", "050209"),
        (2022, "劳动与社会保障", "120403"),
        (2023, "社会工作", "030302"),
        (2023, "广播电视学", "050302"),
        (2023, "朝鲜语", "050209"),
        (2023, "商务英语", "050262"),
        (2023, "公共事业管理", "120401"),
        (2024, "治安学", "030601"),
        (2024, "市场营销", "120202"),
        (2024, "公共事业管理", "120401"),
        (2024, "商务英语", "050262"),
        (2025, "治安学", "030601"),
        (2025, "公共事业管理", "120401"),
        (2025, "商务英语", "050262"),
        (2025, "监狱学", "030103"),
    ]
    cancel_specs = [
        (2022, "思想政治教育", "030503"),
        (2022, "信息管理与信息系统", "120102"),
        (2023, "经济统计学", "020102"),
        (2024, "广播电视学", "050302"),
        (2024, "朝鲜语", "050209"),
        (2025, "市场营销", "120202"),
    ]
    rows: list[dict[str, Any]] = []
    for index, (year, major, code) in enumerate(stop_specs, start=1):
        rows.append(
            {
                "policy_year": year,
                "region": "西北政法大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": f"西北政法大学{year}年专业申报材料未招生本科专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(index),
                "study_duration": "",
                "policy_action": f"{year}年列为未招生专业。",
                "criterion_text": "西北政法大学司法鉴定学专业申请表“学校近五年专业增设、停招、撤并情况”逐年列出未招生专业。",
                "source_ids": source_id,
                "evidence_text": "西北政法大学司法鉴定学专业申请表逐年列出2022-2025年未招生专业；本行对应原文直接点名的未招生专业。",
                "confidence": "high",
            }
        )
    for index, (year, major, code) in enumerate(cancel_specs, start=1):
        rows.append(
            {
                "policy_year": year,
                "region": "西北政法大学",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": f"西北政法大学{year}年专业申报材料撤销本科专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": "",
                "study_duration": "",
                "policy_action": f"{year}年列为撤销专业。",
                "criterion_text": "西北政法大学司法鉴定学专业申请表“学校近五年专业增设、停招、撤并情况”逐年列出撤销专业。",
                "source_ids": source_id,
                "evidence_text": "西北政法大学司法鉴定学专业申请表逐年列出2022-2025年撤销专业；本行对应原文直接点名的撤销专业。",
                "confidence": "high",
            }
        )
    return rows


def parse_sicnu_2025_sports_tourism_application_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "sicnu_2025_sports_tourism_application_pdf"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    required_fragments = [
        "学校名称四川师范大学",
        "专业代码040212TK专业名称体育旅游",
        "2020年，停招“酒店管理”“服装设计与工程”2个专业",
        "2021年，停招“国际经济与贸易”“工业工程”“工业设计”3个专业",
    ]
    missing = [fragment for fragment in required_fragments if fragment not in compact_text]
    if missing:
        raise ValueError(f"Could not verify SICNU sports tourism application fragments in {source_id}: {missing}")

    specs = [
        (2020, "酒店管理", "120902"),
        (2020, "服装设计与工程", "081602"),
        (2021, "国际经济与贸易", "020401"),
        (2021, "工业工程", "120701"),
        (2021, "工业设计", "080205"),
    ]
    rows: list[dict[str, Any]] = []
    for index, (year, major, code) in enumerate(specs, start=1):
        rows.append(
            {
                "policy_year": year,
                "region": "四川师范大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": f"四川师范大学{year}年专业申报材料停招本科专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(index),
                "study_duration": "",
                "policy_action": f"{year}年停招{major}专业。",
                "criterion_text": "四川师范大学体育旅游专业申请表“学校近五年专业增设、停招、撤并情况”逐年列出停招专业。",
                "source_ids": source_id,
                "evidence_text": "四川师范大学体育旅游专业申请表说明，2020年停招酒店管理、服装设计与工程，2021年停招国际经济与贸易、工业工程、工业设计。",
                "confidence": "high",
            }
        )
    return rows


def parse_tjufezj_2025_fiscal_application_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "tjufezj_2025_fiscal_application_pdf"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    required_fragments = [
        "天津财经大学珠江",
        "专业代码020201K专业名称财政学",
        "专业撤销情况：2024年，撤销信息管理与信息系统、网络工程、俄语三个专业",
        "专业停招情况：2020年，停招产品设计、服装与服饰设计两个专业",
    ]
    missing = [fragment for fragment in required_fragments if fragment not in compact_text]
    if missing:
        raise ValueError(f"Could not verify TJUFEZJ fiscal application fragments in {source_id}: {missing}")

    specs = [
        (2020, "major_stop_enrollment", "产品设计", "130504", "2020年停招产品设计专业"),
        (2020, "major_stop_enrollment", "服装与服饰设计", "130505", "2020年停招服装与服饰设计专业"),
        (2024, "major_cancel", "信息管理与信息系统", "120102", "2024年撤销信息管理与信息系统专业"),
        (2024, "major_cancel", "网络工程", "080903", "2024年撤销网络工程专业"),
        (2024, "major_cancel", "俄语", "050202", "2024年撤销俄语专业"),
    ]
    rows: list[dict[str, Any]] = []
    for index, (year, record_type, major, code, action) in enumerate(specs, start=1):
        rows.append(
            {
                "policy_year": year,
                "region": "天津财经大学珠江学院",
                "education_level": "本科",
                "record_type": record_type,
                "warning_label": f"天津财经大学珠江学院{year}年专业申报材料专业调整本科专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(index),
                "study_duration": "",
                "policy_action": action + "。",
                "criterion_text": "天津财经大学珠江学院财政学专业申请表“学校近五年专业增设、停招、撤并情况”分别列出专业停招和专业撤销情况。",
                "source_ids": source_id,
                "evidence_text": "天津财经大学珠江学院财政学专业申请表说明，2024年撤销信息管理与信息系统、网络工程、俄语，2020年停招产品设计、服装与服饰设计。",
                "confidence": "high",
            }
        )
    return rows


def parse_whtyxykj_2025_smart_sports_engineering_application_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "whtyxykj_2025_smart_sports_engineering_application_pdf"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    required_fragments = [
        "武汉体育学院体育科",
        "专业代码040211TK专业名称智能体育工程",
        "近五年学校停招专业：2019年停招文化产业管理专业",
        "近五年学校撤并情况：2024年撤销表演专业",
    ]
    missing = [fragment for fragment in required_fragments if fragment not in compact_text]
    if missing:
        raise ValueError(f"Could not verify WHTYXYKJ smart sports engineering application fragments in {source_id}: {missing}")
    return [
        {
            "policy_year": 2019,
            "region": "武汉体育学院体育科技学院",
            "education_level": "本科",
            "record_type": "major_stop_enrollment",
            "warning_label": "武汉体育学院体育科技学院2019年专业申报材料停招本科专业",
            "reported_major_name": "文化产业管理",
            "major_code": "120210",
            "source_row_no": "1",
            "study_duration": "",
            "policy_action": "2019年停招文化产业管理专业。",
            "criterion_text": "武汉体育学院体育科技学院智能体育工程专业申请表“学校近五年专业增设、停招、撤并情况”列出近五年学校停招专业。",
            "source_ids": source_id,
            "evidence_text": "武汉体育学院体育科技学院智能体育工程专业申请表说明，近五年学校停招专业为2019年停招文化产业管理专业。",
            "confidence": "high",
        },
        {
            "policy_year": 2024,
            "region": "武汉体育学院体育科技学院",
            "education_level": "本科",
            "record_type": "major_cancel",
            "warning_label": "武汉体育学院体育科技学院2024年专业申报材料撤销本科专业",
            "reported_major_name": "表演",
            "major_code": "130301",
            "source_row_no": "",
            "study_duration": "",
            "policy_action": "2024年撤销表演专业。",
            "criterion_text": "武汉体育学院体育科技学院智能体育工程专业申请表“学校近五年专业增设、停招、撤并情况”列出近五年学校撤并情况。",
            "source_ids": source_id,
            "evidence_text": "武汉体育学院体育科技学院智能体育工程专业申请表说明，近五年学校撤并情况为2024年撤销表演专业。",
            "confidence": "high",
        },
    ]


def parse_ntu_2025_smart_energy_application_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "ntu_2025_smart_energy_application_pdf"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    validation_text = compact_text.replace("学校近五年专业增设、", "").replace("停招、撤并情况", "")
    stop_rows = [
        ("应用物理学", "070202"),
        ("社会工作", "030302"),
        ("公共事业管理", "120401"),
        ("人力资源管理", "120206"),
        ("电子商务", "120801"),
        ("环境工程", "082502"),
        ("生物工程", "083001"),
        ("海洋技术", "070702"),
        ("网络工程", "080903"),
        ("动画", "130310"),
        ("产品设计", "130504"),
        ("人文地理与城乡规划", "070503"),
        ("旅游管理", "120901K"),
        ("交通运输", "081801"),
        ("建筑电气与智能化", "081004"),
        ("表演", "130301"),
        ("生物信息学", "071003"),
        ("物流管理", "120601"),
        ("机械设计制造及其自动化", "080202"),
        ("秘书学", "050107T"),
    ]
    required_fragments = [
        "专业代码080608TK专业名称智慧能源工程",
        "新增专业5个：新能源材料与器件、智能制造工程、智能医学工程、人工智能、基础医学",
        "停招专业20个：应用物理学、社会工作、公共事业管理、人力资源管理",
        "撤销专业7个：电气工程与智能控制、电子信息科学与技术、自然地理与资源环境",
    ]
    missing = [fragment for fragment in required_fragments if fragment not in validation_text]
    missing.extend(major for major, _code in stop_rows if major not in validation_text)
    if missing:
        raise ValueError(f"Could not verify NTU professional-history fragments in {source_id}: {missing}")

    rows: list[dict[str, Any]] = []
    for index, (major, code) in enumerate(stop_rows, start=1):
        rows.append(
            {
                "policy_year": 2025,
                "region": "南通大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "南通大学2025年专业申报材料近五年停招本科专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(index),
                "study_duration": "",
                "policy_action": "学校专业申报材料列为近五年已停招本科专业。",
                "criterion_text": "南通大学智慧能源工程专业申请表“学校近五年专业增设、停招、撤并情况”中的停招专业；来源未给出逐项停招年份，policy_year采用申请材料年份。",
                "source_ids": source_id,
                "evidence_text": (
                    "南通大学智慧能源工程专业申请表说明，近五年学校停招应用物理学、社会工作、公共事业管理、"
                    "人力资源管理、电子商务、环境工程、生物工程、海洋技术、网络工程、动画、产品设计、"
                    "人文地理与城乡规划、旅游管理、交通运输、建筑电气与智能化、表演、生物信息学、物流管理、"
                    "机械设计制造及其自动化、秘书学20个专业；"
                    f"本行对应第{index}项：{major}（{code}）。"
                ),
                "confidence": "high",
            }
        )
    return rows


def parse_tjcu_2025_low_altitude_application_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "tjcu_2025_low_altitude_application_pdf"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    validation_text = compact_text.replace("学校近五年专业增设、", "").replace("停招、撤并情况", "")
    stop_rows = [
        ("药事管理", "100704T"),
        ("信息与计算科学", "070102"),
        ("应用物理学", "070202"),
        ("产品设计", "130504"),
        ("信息管理与信息系统", "120102"),
    ]
    required_fragments = [
        "专业代码083203TK专业名称低空技术与工程",
        "专业增设：经济统计学、翻译",
        "停招专业：药事管理专业、信息与计算科学专业、应用物理学专业、产品设计专业、信息管理与信息系统",
        "撤销专业：药事管理专业、信息与计算科学专业、应用物理学专业、产品设计专业",
    ]
    missing = [fragment for fragment in required_fragments if fragment not in validation_text]
    missing.extend(major for major, _code in stop_rows if major not in validation_text)
    if missing:
        raise ValueError(f"Could not verify TJCU professional-history fragments in {source_id}: {missing}")

    rows: list[dict[str, Any]] = []
    for index, (major, code) in enumerate(stop_rows, start=1):
        rows.append(
            {
                "policy_year": 2025,
                "region": "天津商业大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "天津商业大学2025年专业申报材料近五年停招本科专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(index),
                "study_duration": "",
                "policy_action": "学校专业申报材料列为近五年已停招本科专业。",
                "criterion_text": "天津商业大学低空技术与工程专业申请表“学校近五年专业增设、停招、撤并情况”中的停招专业；来源未给出逐项停招年份，policy_year采用申请材料年份。",
                "source_ids": source_id,
                "evidence_text": (
                    "天津商业大学低空技术与工程专业申请表说明，近五年学校停招药事管理、信息与计算科学、"
                    "应用物理学、产品设计、信息管理与信息系统5个专业；"
                    f"本行对应第{index}项：{major}（{code}）。"
                ),
                "confidence": "high",
            }
        )
    return rows


def parse_csuft_2025_low_altitude_application_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "csuft_2025_low_altitude_application_pdf"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    validation_text = compact_text.replace("学校近五年专业增设、", "").replace("停招、撤并情况", "")
    stop_rows = [
        (2021, "电子科学与技术", "080702"),
        (2021, "交通运输", "081801"),
        (2023, "材料成型及控制工程", "080203"),
        (2024, "国际商务", "120205"),
        (2024, "保险学", "020303"),
        (2024, "法语", "050204"),
        (2024, "行政管理", "120402"),
        (2025, "朝鲜语", "050209"),
        (2025, "地理信息科学", "070504"),
        (2025, "工程力学", "080102"),
        (2025, "材料化学", "080403"),
        (2025, "高分子材料与工程", "080407"),
        (2025, "能源与动力工程", "080501"),
        (2025, "建筑环境与能源应用工程", "081002"),
        (2025, "森林工程", "082401"),
        (2025, "环境生态工程", "082504"),
        (2025, "酒店管理", "120902"),
        (2025, "会展经济与管理", "120903"),
        (2025, "舞蹈学", "130205"),
    ]
    required_fragments = [
        "专业代码083203TK专业名称低空技术与工程",
        "停招：电子科学与技术、交通运输（2021）",
        "材料成型及控制工程（2023）",
        "国际商务、保险学、法语、行政管理（2024年）",
        "朝鲜语、地理信息科学、工程力学、材料化学、高分子材料与工程",
    ]
    missing = [fragment for fragment in required_fragments if fragment not in validation_text]
    missing.extend(major for _year, major, _code in stop_rows if major not in validation_text)
    if missing:
        raise ValueError(f"Could not verify CSUFT professional-history fragments in {source_id}: {missing}")

    rows: list[dict[str, Any]] = []
    for index, (year, major, code) in enumerate(stop_rows, start=1):
        rows.append(
            {
                "policy_year": year,
                "region": "中南林业科技大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": f"中南林业科技大学{year}年专业申报材料停招本科专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(index),
                "study_duration": "",
                "policy_action": "学校专业申报材料列为近五年停招本科专业。",
                "criterion_text": "中南林业科技大学低空技术与工程专业申请表“学校近五年专业增设、停招、撤并情况”中的停招专业。",
                "source_ids": source_id,
                "evidence_text": (
                    "中南林业科技大学低空技术与工程专业申请表说明，2021年停招电子科学与技术、交通运输，"
                    "2023年停招材料成型及控制工程，2024年停招国际商务、保险学、法语、行政管理，"
                    "2025年停招朝鲜语、地理信息科学、工程力学、材料化学、高分子材料与工程、能源与动力工程、"
                    "建筑环境与能源应用工程、森林工程、环境生态工程、酒店管理、会展经济与管理、舞蹈学；"
                    f"本行对应第{index}项：{major}（{code}，{year}年）。"
                ),
                "confidence": "high",
            }
        )
    return rows


def parse_zju_2025_undergrad_stop_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "zju_2025_undergrad_major_table"
    html_path = raw_dir / f"{source_id}.html"
    if not html_path.exists():
        html_path = raw_dir / f"{source_id}.htm"
    if not html_path.exists():
        return []
    html = html_path.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(html, "html.parser")
    target_table = None
    for table in soup.find_all("table"):
        header_text = re.sub(r"\s+", "", table.get_text("", strip=True))
        if "序号专业代码专业名称学位学科门类备注" in header_text:
            target_table = table
            break
    if target_table is None:
        raise ValueError(f"Could not find Zhejiang University undergraduate major table in {source_id}.")

    rows: list[dict[str, Any]] = []
    for tr in target_table.find_all("tr")[1:]:
        cells = [cell.get_text(" ", strip=True) for cell in tr.find_all(["td", "th"])]
        if len(cells) < 6:
            continue
        source_row_no, major_code, major, degree, discipline, note = cells[:6]
        if "已停招" not in note:
            continue
        compact_code = re.sub(r"\s+", "", major_code)
        rows.append(
            {
                "policy_year": 2025,
                "region": "浙江大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "浙江大学本科专业情况表已停招专业",
                "reported_major_name": major,
                "major_code": compact_code,
                "study_duration": "",
                "policy_action": "已停招。",
                "criterion_text": "浙江大学本科专业情况表备注列标注“已停招”。",
                "source_row_no": source_row_no,
                "source_ids": source_id,
                "evidence_text": f"浙江大学本科专业情况表第{source_row_no}项列出{major}（{compact_code}，{degree}，{discipline}），备注为已停招。",
                "confidence": "high",
            }
        )
    if len(rows) != 28:
        raise ValueError(f"Expected 28 Zhejiang University stopped majors in {source_id}, parsed {len(rows)}.")
    return rows


def parse_hfit_2025_major_setting_stop_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "hfit_2025_major_setting"
    html_path = raw_dir / f"{source_id}.html"
    if not html_path.exists():
        html_path = raw_dir / f"{source_id}.htm"
    if not html_path.exists():
        return []
    html = html_path.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(html, "html.parser")
    target_table = None
    for table in soup.find_all("table"):
        header_text = re.sub(r"\s+", "", table.get_text("", strip=True))
        if "合肥理工学院2025年专业基本情况数据表" in header_text and "专业名称门类名称专业代码备注" in header_text:
            target_table = table
            break
    if target_table is None:
        raise ValueError(f"Could not find HFIT 2025 major-setting table in {source_id}.")

    rows: list[dict[str, Any]] = []
    for tr in target_table.find_all("tr"):
        cells = [cell.get_text(" ", strip=True) for cell in tr.find_all(["td", "th"])]
        if len(cells) != 5 or cells[0] in {"序号", "合肥理工学院2025年专业基本情况数据表"}:
            continue
        source_row_no, major, discipline, major_code, note = cells
        if "当年停招" not in note:
            continue
        rows.append(
            {
                "policy_year": 2025,
                "region": "合肥理工学院",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2025年专业基本情况数据表当年停招专业",
                "reported_major_name": major,
                "major_code": major_code,
                "study_duration": "",
                "policy_action": "当年停招。",
                "criterion_text": "合肥理工学院2025年专业基本情况数据表备注列标注“当年停招”。",
                "source_row_no": source_row_no,
                "source_ids": source_id,
                "evidence_text": f"合肥理工学院2025年专业基本情况数据表第{source_row_no}项列出{major}（{major_code}，{discipline}），备注为当年停招。",
                "confidence": "high",
            }
        )
    if len(rows) != 31:
        raise ValueError(f"Expected 31 HFIT stopped majors in {source_id}, parsed {len(rows)}.")
    return rows


def parse_shiep_stop_enrollment_rows(raw_dir: Path) -> list[dict[str, Any]]:
    def read_cached_text(source_id: str) -> str:
        text_path = raw_dir / f"{source_id}.txt"
        if text_path.exists():
            return text_path.read_text(encoding="utf-8", errors="replace")
        return ""

    text_2019 = read_cached_text("shiep_2019_stop_enrollment")
    text_2021 = read_cached_text("shiep_2021_stop_enrollment")
    text_2023 = read_cached_text("shiep_2023_stop_enrollment")
    text_pdf = read_cached_text("shiep_2024_teaching_quality_report_pdf")
    compact_2019 = re.sub(r"\s+", "", text_2019)
    compact_2021 = re.sub(r"\s+", "", text_2021)
    compact_2023 = re.sub(r"\s+", "", text_2023)
    compact_pdf = re.sub(r"\s+", "", text_pdf)
    if not any([compact_2019, compact_2021, compact_2023, compact_pdf]):
        return []

    required_by_source = {
        "shiep_2019_stop_enrollment": (
            compact_2019,
            ["2019年停招专业：电力工程与管理"],
        ),
        "shiep_2021_stop_enrollment": (
            compact_2021,
            ["2021年停招专业：公共事业管理"],
        ),
        "shiep_2023_stop_enrollment": (
            compact_2023,
            [
                "2023年停招本科专业",
                "120401公共事业管理2021-2023",
                "080401材料科学与工程2022-2023",
                "080903网络工程2023",
            ],
        ),
        "shiep_2024_teaching_quality_report_pdf": (
            compact_pdf,
            [
                "当年停招专业4个：机械电子工程、材料化学、物流管理、日语",
                "说明：机械电子工程、材料化学、材料科学与工程、网络工程、公共事业管理、物流管理、日语7个专业已停招",
            ],
        ),
    }
    for source_id, (text, fragments) in required_by_source.items():
        if not text:
            raise ValueError(f"Missing cached text for {source_id}.")
        missing = [fragment for fragment in fragments if fragment not in text]
        if missing:
            raise ValueError(f"Could not verify SHIEP fragments in {source_id}: {missing}")

    rows: list[dict[str, Any]] = []

    def add_row(
        policy_year: int,
        major: str,
        code: str,
        source_row_no: str,
        source_ids: str,
        warning_label: str,
        criterion_text: str,
        evidence_text: str,
    ) -> None:
        rows.append(
            {
                "policy_year": policy_year,
                "region": "上海电力大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": warning_label,
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": source_row_no,
                "study_duration": "",
                "policy_action": "停招。",
                "criterion_text": criterion_text,
                "source_ids": source_ids,
                "evidence_text": evidence_text,
                "confidence": "high",
            }
        )

    add_row(
        2019,
        "电力工程与管理",
        "",
        "1",
        "shiep_2019_stop_enrollment",
        "2019年上海电力大学本科停招专业",
        "上海电力大学信息公开页“2019年上海电力大学本科新增或停招专业名单”正文列出停招专业。",
        "上海电力大学信息公开页说明2019年停招专业为电力工程与管理。",
    )

    shiep_2023_rows = [
        (2021, "公共事业管理", "120401", "1", "2021-2023"),
        (2022, "公共事业管理", "120401", "1", "2021-2023"),
        (2023, "公共事业管理", "120401", "1", "2021-2023"),
        (2022, "材料科学与工程", "080401", "2", "2022-2023"),
        (2023, "材料科学与工程", "080401", "2", "2022-2023"),
        (2023, "网络工程", "080903", "3", "2023"),
    ]
    for year, major, code, row_no, stop_years in shiep_2023_rows:
        source_ids = "shiep_2023_stop_enrollment"
        if year == 2021 and major == "公共事业管理":
            source_ids = "shiep_2021_stop_enrollment;shiep_2023_stop_enrollment"
        add_row(
            year,
            major,
            code,
            row_no,
            source_ids,
            f"{year}年上海电力大学本科停招专业",
            "上海电力大学信息公开页“2023年新增、停招本科专业情况”表格列出专业代码、专业名称、停招年份、修业年限和学位授予门类。",
            f"上海电力大学2023年停招本科专业表第{row_no}行列出{major}（{code}），停招年份为{stop_years}。",
        )

    shiep_stopped_majors = [
        ("8", "机械电子工程", "080204"),
        ("12", "材料化学", "080403"),
        ("13", "材料科学与工程", "080401"),
        ("23", "网络工程", "080903"),
        ("33", "物流管理", "120601"),
        ("36", "公共事业管理", "120401"),
        ("40", "日语", "050207"),
    ]
    for row_no, major, code in shiep_stopped_majors:
        add_row(
            2024,
            major,
            code,
            row_no,
            "shiep_2024_teaching_quality_report_pdf",
            "2023-2024学年本科教学质量报告已停招专业",
            "上海电力大学2023-2024学年本科教学质量报告附表2备注标记“停招”，报告正文还说明当年停招机械电子工程、材料化学、物流管理、日语4个专业。",
            f"上海电力大学2023-2024学年本科教学质量报告附表2第{row_no}项列出{major}（{code}）并标记停招；报告说明7个专业已停招。",
        )
        add_row(
            2025,
            major,
            code,
            row_no,
            "shiep_2025_stop_enrollment;shiep_2024_teaching_quality_report_pdf",
            "2025年停招本科专业",
            "上海电力大学本科生院保留2025年新增、停招本科专业情况页；学校2023-2024学年本科教学质量报告附表2逐名列出同一批已停招专业。",
            f"上海电力大学2025年新增、停招本科专业情况页确认有停招信息；学校2023-2024学年本科教学质量报告附表2第{row_no}项列出{major}（{code}）并标记停招。",
        )
    return rows


def parse_shiep_2026_professional_adjustment_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "shiep_2026_professional_adjustment_meeting"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    required_fragments = [
        "上海电力大学教学指导委员会召开2026年第一次会议",
        "会议第三项议程为审议2026年专业优化调整方案",
        "拟停招环境工程、光电信息科学与工程、信息与计算科学三个专业",
        "拟撤销公共事业管理专业",
    ]
    missing = [fragment for fragment in required_fragments if fragment not in compact_text]
    if missing:
        raise ValueError(f"Could not verify SHIEP 2026 professional-adjustment fragments in {source_id}: {missing}")

    rows: list[dict[str, Any]] = []
    stop_rows = [
        ("环境工程", "082502"),
        ("光电信息科学与工程", "080705"),
        ("信息与计算科学", "070102"),
    ]
    for index, (major, code) in enumerate(stop_rows, start=1):
        rows.append(
            {
                "policy_year": 2026,
                "region": "上海电力大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "上海电力大学2026年专业优化调整方案拟停招本科专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(index),
                "study_duration": "",
                "policy_action": "2026年专业优化调整方案拟停招。",
                "criterion_text": "上海电力大学官网新闻“上海电力大学教学指导委员会召开2026年第一次会议”正文列出2026年专业优化调整方案拟停招专业。",
                "source_ids": source_id,
                "evidence_text": f"上海电力大学官网新闻说明，2026年专业优化调整方案拟停招环境工程、光电信息科学与工程、信息与计算科学三个专业；本行对应第{index}项：{major}（{code}）。",
                "confidence": "high",
            }
        )
    rows.append(
        {
            "policy_year": 2026,
            "region": "上海电力大学",
            "education_level": "本科",
            "record_type": "major_cancel",
            "warning_label": "上海电力大学2026年专业优化调整方案拟撤销本科专业",
            "reported_major_name": "公共事业管理",
            "major_code": "120401",
            "source_row_no": "",
            "study_duration": "",
            "policy_action": "2026年专业优化调整方案拟撤销。",
            "criterion_text": "上海电力大学官网新闻“上海电力大学教学指导委员会召开2026年第一次会议”正文列出2026年专业优化调整方案拟撤销专业。",
            "source_ids": source_id,
            "evidence_text": "上海电力大学官网新闻说明，2026年专业优化调整方案拟撤销公共事业管理专业。",
            "confidence": "high",
        }
    )
    return rows


def parse_chu_2019_2024_major_setting_rows(raw_dir: Path) -> list[dict[str, Any]]:
    def read_cached_text(source_id: str) -> str:
        text_path = raw_dir / f"{source_id}.txt"
        if text_path.exists():
            return text_path.read_text(encoding="utf-8", errors="replace")
        return ""

    configs = [
        {
            "year": 2019,
            "source_ids": "chu_2019_major_setting",
            "validation_source_id": "chu_2019_major_setting",
            "warning_label": "2019年本科专业设置及停招专业情况一览表",
            "rows": [
                ("3", "历史学（师范）", "060101", "四年", "历史学", "旅游管理学院"),
                ("4", "物理学（师范）", "070201", "四年", "理学", "电子工程学院"),
                ("11", "教育技术学（师范）", "040104", "四年", "理学", "文学传媒与教育科学学院"),
                ("18", "小学教育（师范）", "040107", "四年", "教育学", "文学传媒与教育科学学院"),
                ("21", "应用心理学", "071102", "四年", "理学", "文学传媒与教育科学学院"),
                ("23", "公共事业管理", "120401", "四年", "管理学", "工商管理学院"),
                ("27", "信息与计算科学", "070102", "四年", "理学", "数学与统计学院"),
                ("28", "微电子科学与工程", "080704", "四年", "工学", "电子工程学院"),
                ("32", "视觉传达设计（中韩2+2）", "130502", "四年", "艺术学", "艺术学院"),
                ("35", "信息管理与信息系统", "120102", "四年", "工学", "信息工程学院"),
                ("36", "统计学", "071201", "四年", "理学", "数学与统计学院"),
                ("42", "文化产业管理", "120201", "四年", "管理学", "文化传媒与教育科学学院"),
            ],
        },
        {
            "year": 2020,
            "source_ids": "chu_2020_major_setting",
            "validation_source_id": "chu_2020_major_setting",
            "warning_label": "2020年本科专业设置及新增专业、停招专业情况一览表",
            "rows": [
                ("3", "历史学（师范）", "060101", "四年", "历史学", "旅游管理学院"),
                ("7", "视觉传达设计（中韩2+2）", "130502", "四年", "艺术学", "艺术学院"),
                ("11", "教育技术学（师范）", "040104", "四年", "理学", "文学传媒与教育科学学院"),
                ("18", "小学教育（师范）", "040107", "四年", "教育学", "文学传媒与教育科学学院"),
                ("23", "公共事业管理", "120401", "四年", "管理学", "工商管理学院"),
                ("27", "信息与计算科学", "070102", "四年", "理学", "数学与统计学院"),
                ("28", "微电子科学与工程", "080704", "四年", "工学", "电子工程学院"),
                ("33", "信息管理与信息系统", "120102", "四年", "工学", "信息工程学院"),
                ("34", "统计学", "071201", "四年", "理学", "数学与统计学院"),
                ("40", "文化产业管理", "120201", "四年", "管理学", "文学传媒与教育科学学院"),
            ],
        },
        {
            "year": 2021,
            "source_ids": "chu_2021_major_setting",
            "validation_source_id": "chu_2021_major_setting",
            "warning_label": "2021年本科专业设置及新增专业、停招专业情况一览表",
            "rows": [
                ("3", "历史学（师范）", "060101", "四年", "历史学", "旅游管理学院"),
                ("7", "视觉传达设计（中韩2+2）", "130502", "四年", "艺术学", "艺术学院"),
                ("11", "教育技术学（师范）", "040104", "四年", "理学", "文学传媒与教育科学学院"),
                ("18", "小学教育（师范）", "040107", "四年", "教育学", "文学传媒与教育科学学院"),
                ("23", "公共事业管理", "120401", "四年", "管理学", "工商管理学院"),
                ("27", "信息与计算科学", "070102", "四年", "理学", "数学与统计学院"),
                ("28", "微电子科学与工程", "080704", "四年", "工学", "电子工程学院"),
                ("33", "信息管理与信息系统", "120102", "四年", "工学", "信息工程学院"),
                ("34", "统计学", "071201", "四年", "理学", "数学与统计学院"),
                ("40", "文化产业管理", "120201", "四年", "管理学", "文学传媒与教育科学学院"),
                ("44", "酒店管理", "120902", "四年", "管理学", "旅游管理学院"),
                ("54", "人工智能", "080717T", "四年", "工学", "信息工程学院"),
                ("56", "食品科学与工程", "082701", "四年", "工学", "化学与材料工程学院"),
            ],
        },
        {
            "year": 2022,
            "source_ids": "chu_2022_major_setting",
            "validation_source_id": "chu_2022_major_setting",
            "warning_label": "2022年本科专业设置及新增专业、停招专业情况一览表",
            "rows": [
                ("3", "历史学（师范）", "060101", "四年", "历史学", "旅游管理学院"),
                ("7", "视觉传达设计（中韩2+2）", "130502", "四年", "艺术学", "美术与设计学院"),
                ("11", "教育技术学（师范）", "040104", "四年", "理学", "教师教育学院"),
                ("18", "小学教育（师范）", "040107", "四年", "教育学", "教师教育学院"),
                ("23", "公共事业管理", "120401", "四年", "管理学", "工商管理学院"),
                ("24", "广播电视学", "050302", "四年", "文学", "文学与传媒学院"),
                ("27", "信息与计算科学", "070102", "四年", "理学", "数学与大数据学院"),
                ("28", "微电子科学与工程", "080704", "四年", "工学", "电子工程学院"),
                ("33", "信息管理与信息系统", "120102", "四年", "工学", "计算机与人工智能学院"),
                ("34", "统计学", "071201", "四年", "理学", "数学与大数据学院"),
                ("40", "文化产业管理", "120201", "四年", "管理学", "文学与传媒学院"),
                ("44", "酒店管理", "120902", "四年", "管理学", "旅游管理学院"),
            ],
        },
        {
            "year": 2023,
            "source_ids": "chu_2023_major_setting_notice;chu_2023_major_setting_pdf",
            "validation_source_id": "chu_2023_major_setting_pdf",
            "warning_label": "2023年本科专业设置及新增专业、停招专业情况一览表",
            "rows": [
                ("3", "历史学（师范）", "060101", "四年", "历史学", "旅游管理学院"),
                ("11", "教育技术学（师范）", "040104", "四年", "理学", "教师教育学院"),
                ("18", "小学教育（师范）", "040107", "四年", "教育学", "教师教育学院"),
                ("23", "广播电视学", "050302", "四年", "文学", "文学与传媒学院"),
                ("31", "统计学", "071201", "四年", "理学", "数学与大数据学院"),
                ("37", "文化产业管理", "120201", "四年", "管理学", "文学与传媒学院"),
                ("41", "酒店管理", "120902", "四年", "管理学", "旅游管理学院"),
            ],
        },
        {
            "year": 2024,
            "source_ids": "chu_2024_major_setting_notice;chu_2024_major_setting_pdf",
            "validation_source_id": "chu_2024_major_setting_pdf",
            "warning_label": "2024年本科专业设置及新增专业、停招专业情况一览表",
            "rows": [
                ("3", "历史学（师范）", "060101", "四年", "历史学", "旅游管理学院"),
                ("4", "物理学（师范）", "070201", "四年", "理学", "电子工程学院"),
                ("21", "广播电视学", "050302", "四年", "文学", "文学与传媒学院"),
                ("29", "统计学", "071201", "四年", "理学", "数学与大数据学院"),
                ("38", "酒店管理", "120902", "四年", "管理学", "旅游管理学院"),
                ("41", "会展经济与管理", "120903", "四年", "管理学", "旅游管理学院"),
                ("43", "审计学", "120207", "四年", "管理学", "工商管理学院"),
                ("48", "互联网金融", "020309T", "四年", "经济学", "经济与法学学院"),
            ],
        },
    ]

    validation_texts = {
        str(config["validation_source_id"]): re.sub(
            r"\s+",
            "",
            read_cached_text(str(config["validation_source_id"])),
        )
        for config in configs
    }
    if not any(validation_texts.values()):
        return []

    rows: list[dict[str, Any]] = []
    for config in configs:
        year = int(config["year"])
        validation_source_id = str(config["validation_source_id"])
        compact_text = validation_texts[validation_source_id]
        if not compact_text:
            raise ValueError(f"Missing cached text for {validation_source_id}.")
        if "停招" not in compact_text:
            raise ValueError(f"Could not verify CHU stop-enrollment marker in {validation_source_id}.")
        missing_fragments = [
            fragment
            for row_no, major, code, _duration, _degree, _school_unit in config["rows"]
            for fragment in (major, code)
            if fragment not in compact_text
        ]
        if missing_fragments:
            raise ValueError(f"Could not verify CHU fragments in {validation_source_id}: {missing_fragments}")

        for source_row_no, major, code, duration, degree, school_unit in config["rows"]:
            rows.append(
                {
                    "policy_year": year,
                    "region": "巢湖学院",
                    "education_level": "本科",
                    "record_type": "major_stop_enrollment",
                    "warning_label": str(config["warning_label"]),
                    "reported_major_name": major,
                    "major_code": code,
                    "source_row_no": source_row_no,
                    "study_duration": duration,
                    "policy_action": f"{year}年停招。",
                    "criterion_text": f"巢湖学院{year}年本科专业设置及新增专业、停招专业情况一览表在“{year}年停招”列标记为停招。",
                    "source_ids": str(config["source_ids"]),
                    "evidence_text": (
                        f"巢湖学院{year}年本科专业设置及新增专业、停招专业情况一览表第{source_row_no}项"
                        f"列出{major}（{code}，{duration}，{degree}，{school_unit}），"
                        f"{year}年停招列标记为停招。"
                    ),
                    "confidence": "high",
                    "notes": school_unit,
                }
            )
    return rows


def parse_2025_university_adjustment_notice_rows(raw_dir: Path) -> list[dict[str, Any]]:
    configs = [
        {
            "source_id": "swjtu_2025_major_adjustment_notice",
            "region": "西南交通大学",
            "warning_label": "2025年本科专业调整公示拟撤销专业",
            "criterion_text": "学校本科专业调整公示列出2025年拟撤销专业名单。",
            "evidence_prefix": "西南交通大学2025年本科专业调整公示说明，经校本科教学工作委员会审议、校长办公会批准，2025年学校拟撤销专业5个，拟撤销专业为",
            "required_fragments": [
                "2025年学校拟申报新专业9个，拟预申报专业2个，拟撤销专业5个",
                "拟撤销专业：电气工程与智能控制、交通设备与控制工程、国际经济与贸易、视觉传达设计、商务英语",
            ],
            "majors": ["电气工程与智能控制", "交通设备与控制工程", "国际经济与贸易", "视觉传达设计", "商务英语"],
        },
        {
            "source_id": "jzun_2025_major_adjustment_notice",
            "region": "荆州学院",
            "warning_label": "2025年本科专业设置调整情况公示拟撤销专业",
            "criterion_text": "学校本科专业设置调整情况公示列出2025年拟撤销专业名单。",
            "evidence_prefix": "荆州学院2025年本科专业设置调整情况公示说明，2025年学校拟申报新专业4个、撤销专业2个、预备案专业1个，拟撤销专业为",
            "required_fragments": [
                "2025年学校拟申报新专业4个，撤销专业2个，预备案专业",
                "拟撤销专业：市场营销、车辆工程",
            ],
            "majors": ["市场营销", "车辆工程"],
        },
        {
            "source_id": "jzmu_2025_major_cancellation_notice",
            "region": "锦州医科大学",
            "warning_label": "2025年撤销本科专业公示",
            "criterion_text": "学校撤销本科专业公示说明保险学专业停招已超过五年，目前已无在校生，决定予以撤销。",
            "evidence_prefix": "锦州医科大学2025年撤销本科专业公示说明，保险学专业停招已超过五年，目前已无在校生，决定予以撤销；本来源列出的撤销专业为",
            "required_fragments": [
                "保险学专业停招已超过五年，目前已无在校生，决定予以撤销",
                "现对撤销专业进行公示",
            ],
            "majors": ["保险学"],
            "policy_action": "决定予以撤销。",
        },
    ]

    rows: list[dict[str, Any]] = []
    for config in configs:
        source_id = str(config["source_id"])
        text_path = raw_dir / f"{source_id}.txt"
        if not text_path.exists():
            continue
        compact_text = re.sub(r"\s+", "", text_path.read_text(encoding="utf-8", errors="replace"))
        missing = [fragment for fragment in config["required_fragments"] if fragment not in compact_text]
        missing.extend(major for major in config["majors"] if str(major) not in compact_text)
        if missing:
            raise ValueError(f"Could not verify 2025 adjustment fragments in {source_id}: {missing}")

        for index, major in enumerate(config["majors"], 1):
            rows.append(
                {
                    "policy_year": 2025,
                    "region": str(config["region"]),
                    "education_level": "本科",
                    "record_type": "major_cancel",
                    "warning_label": str(config["warning_label"]),
                    "reported_major_name": str(major),
                    "source_row_no": str(index),
                    "policy_action": str(config.get("policy_action", "拟撤销。")),
                    "criterion_text": str(config["criterion_text"]),
                    "source_ids": source_id,
                    "evidence_text": f"{config['evidence_prefix']}{'、'.join(config['majors'])}；本记录对应第{index}项：{major}。",
                    "confidence": "high",
                }
            )
    return rows


def parse_pxc_2025_smart_construction_application_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "pxc_2025_smart_construction_application_pdf"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    required_fragments = [
        "专业代码081008T专业名称智能建造",
        "思想政治教育专业停招",
        "数字媒体技术专业停招",
        "工程造价专业停招",
        "2025年，拟撤销设施",
        "农业科学与工程",
    ]
    missing = [fragment for fragment in required_fragments if fragment not in compact_text]
    if missing:
        raise ValueError(f"Could not verify PXC smart construction application fragments in {source_id}: {missing}")

    rows: list[dict[str, Any]] = []
    stop_rows = [
        (2021, "思想政治教育", "040503", "2021年，新增知识产权、旅游管理与服务教育2个本科专业，思想政治教育专业停招。"),
        (2022, "数字媒体技术", "080906", "2022年，撤销特殊教育专业，数字媒体技术专业停招。"),
        (2023, "工程造价", "120105", "2023年，学校近五年工程造价专业停招。"),
    ]
    for index, (year, major, code, evidence) in enumerate(stop_rows, 1):
        rows.append(
            {
                "policy_year": year,
                "region": "萍乡学院",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "智能建造专业备案申请表近五年专业停招情况",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(index),
                "policy_action": "停招。",
                "criterion_text": "萍乡学院智能建造专业备案申请表“学校近五年专业增设、停招、撤并情况”列明对应专业停招。",
                "source_ids": source_id,
                "evidence_text": f"萍乡学院智能建造专业备案申请表在学校近五年专业增设、停招、撤并情况中说明：{evidence}",
                "confidence": "high",
            }
        )
    rows.append(
        {
            "policy_year": 2025,
            "region": "萍乡学院",
            "education_level": "本科",
            "record_type": "major_cancel",
            "warning_label": "智能建造专业备案申请表2025年拟撤销专业",
            "reported_major_name": "设施农业科学与工程",
            "major_code": "090106",
            "source_row_no": "4",
            "policy_action": "拟撤销。",
            "criterion_text": "萍乡学院智能建造专业备案申请表“学校近五年专业增设、停招、撤并情况”列明2025年拟撤销设施农业科学与工程。",
            "source_ids": source_id,
            "evidence_text": "萍乡学院智能建造专业备案申请表说明2025年拟撤销设施农业科学与工程，并新增数字经济、机器人工程、社会体育指导与管理（中外合作办学）3个本科专业。",
            "confidence": "high",
        }
    )
    return rows


def parse_eol_2020_2019_moe_direct_cancel_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "eol_2020_2019_moe_direct_cancel_table"
    html_path = raw_dir / f"{source_id}.shtml"
    if not html_path.exists():
        html_path = raw_dir / f"{source_id}.html"
    if not html_path.exists():
        return []
    html = html_path.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(html, "html.parser")
    table = soup.find("table")
    if table is None:
        raise ValueError(f"Could not find EOL 2019 direct-admin cancellation table in {source_id}.")
    rows: list[dict[str, Any]] = []
    parsed: list[tuple[str, str, str, str, str]] = []
    current_department = ""
    for tr in table.find_all("tr"):
        cells = [cell.get_text(" ", strip=True).replace("\xa0", " ") for cell in tr.find_all(["td", "th"])]
        cells = [re.sub(r"\s+", " ", cell).strip() for cell in cells]
        if not cells:
            continue
        if cells[0].startswith("主管部门、学校名称"):
            continue
        if len(cells) == 1:
            current_department = cells[0]
            continue
        if len(cells) < 5:
            continue
        school, major, code, degree, duration = cells[:5]
        note = cells[5] if len(cells) > 5 else ""
        code = code.replace(" ", "")
        parsed.append((school, major, code, degree, duration))
        rows.append(
            {
                "policy_year": 2019,
                "region": school,
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "中国教育在线2019年度教育部直属高校撤销本科专业表",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(len(rows) + 1),
                "study_duration": duration,
                "policy_action": "中国教育在线转列表格列为2019年度撤销本科专业；教育部2019年度结果以官方来源为准。",
                "criterion_text": "中国教育在线页面“教育部2019年度撤销本科专业”表格，覆盖教育部直属高校撤销专业。",
                "source_ids": source_id,
                "evidence_text": (
                    f"中国教育在线2020年7月1日表格第{len(rows) + 1}项列出{current_department}{school}"
                    f"撤销{major}（{code}，{degree}，{duration}）。"
                    + (f"备注：{note}。" if note else "")
                ),
                "confidence": "medium",
            }
        )
    expected = {
        ("华中师范大学", "公共事业管理", "120401", "管理学", "四年"),
        ("华中师范大学", "戏剧影视文学", "130304", "艺术学", "四年"),
        ("华中师范大学", "动画", "130310", "艺术学", "四年"),
        ("同济大学", "临床医学", "100201K", "医学", "七年"),
        ("长安大学", "工业工程", "120701", "管理学", "四年"),
    }
    observed = set(parsed)
    missing = expected - observed
    if len(rows) != 11 or missing:
        raise ValueError(f"Unexpected EOL 2019 direct-admin cancellation rows in {source_id}: count={len(rows)}, missing={missing}")
    return rows


def parse_gxust_2016_undergrad_setting_stop_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "gxust_2016_undergrad_setting_stop_page"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    rows_data = [
        ("教育技术学", "040104"),
        ("材料成型及控制工程", "080203"),
        ("管理科学", "120101"),
        ("统计学", "071201"),
        ("电子信息科学与技术", "080714T"),
        ("计算机科学与技术", "080901"),
        ("服装设计与工程", "081602"),
        ("机械电子工程", "080204"),
        ("公共事业管理", "120401"),
        ("经济统计学", "020102"),
        ("食品质量与安全", "082702"),
        ("产品设计", "130504"),
        ("机电技术教育", "080211T"),
        ("数学与应用数学", "070101"),
    ]
    required_fragments = [
        "2015—2016学年",
        "设置本科专业67个",
        "新增临床医学1个专业",
        "停招教育技术学、材料成型及控制工程、管理科学、统计学、电子信息科学与技术、计算机科学与技术、服装设计与工程、机械电子工程、公共事业管理、经济统计学、食品质量与安全、产品设计、机电技术教育、数学与应用数学14个专业",
    ]
    missing = [fragment for fragment in required_fragments if fragment not in compact_text]
    if missing:
        raise ValueError(f"Could not verify GXUST 2015-2016 stop-enrollment fragments in {source_id}: {missing}")

    rows: list[dict[str, Any]] = []
    for index, (major, code) in enumerate(rows_data, start=1):
        rows.append(
            {
                "policy_year": 2016,
                "region": "广西科技大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "广西科技大学2015-2016学年停招本科专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(index),
                "study_duration": "",
                "policy_action": "学校信息公开页列为2015—2016学年停招本科专业。",
                "criterion_text": "广西科技大学信息公开页“普通本专科专业设置、当年新增专业、停招专业名单”正文列出的停招专业。",
                "source_ids": source_id,
                "evidence_text": (
                    "广西科技大学信息公开页说明，2015—2016学年学校设置本科专业67个，其中新增临床医学1个专业，"
                    "停招教育技术学、材料成型及控制工程、管理科学、统计学、电子信息科学与技术、计算机科学与技术、"
                    "服装设计与工程、机械电子工程、公共事业管理、经济统计学、食品质量与安全、产品设计、机电技术教育、数学与应用数学14个专业；"
                    f"本行对应第{index}项：{major}（{code}）。"
                ),
                "confidence": "high",
            }
        )
    return rows


def parse_guizhou_2019_undergrad_setup_cancel_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "guizhou_2019_undergrad_setup_results_pdf_mirror"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    rows_data = [
        ("贵州民族大学", "经济统计学", "020102", "经济学", "四年"),
        ("贵州民族大学", "应用统计学", "071202", "理学", "四年"),
        ("贵州民族大学", "采矿工程", "081501", "工学", "四年"),
        ("贵州民族大学", "环境科学与工程", "082501", "工学", "四年"),
        ("贵州民族大学", "环境科学", "082503", "理学", "四年"),
        ("贵州民族大学", "信息管理与信息系统", "120102", "工学", "四年"),
        ("贵阳学院", "秘书学", "050107T", "文学", "四年"),
        ("贵阳学院", "城乡规划", "082802", "工学", "四年"),
        ("贵阳学院", "电子商务", "120801", "管理学", "四年"),
        ("贵阳学院", "服装与服饰设计", "130505", "艺术学", "四年"),
        ("贵州民族大学人文科技学院", "信息与计算科学", "070102", "理学", "四年"),
    ]
    if "撤销本科专业名单" not in text:
        raise ValueError(f"Could not find cancellation section in {source_id}.")
    for school, major, code, _, _ in rows_data:
        if school not in text or major not in text or code not in text:
            raise ValueError(f"Could not verify {school} {major} {code} in {source_id}.")
    rows: list[dict[str, Any]] = []
    for index, (school, major, code, degree, duration) in enumerate(rows_data, start=1):
        rows.append(
            {
                "policy_year": 2019,
                "region": school,
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "2019年度贵州省普通高等学校撤销本科专业名单",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": "",
                "study_duration": duration,
                "policy_action": "贵州省2019年度普通高等学校本科专业备案和审批结果列为撤销本科专业；教育部2019年度结果公布撤销。",
                "criterion_text": "2019年度贵州省普通高等学校本科专业备案和审批结果中的撤销本科专业名单。",
                "source_ids": source_id,
                "evidence_text": (
                    f"2019年度贵州省普通高等学校本科专业备案和审批结果撤销本科专业名单第{index}条："
                    f"{school}撤销{major}（{code}，{degree}，{duration}）。"
                ),
                "confidence": "high",
            }
        )
    return rows


def parse_hebtu_2019_performance_target_cancel_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "hebtu_2019_performance_target_pdf"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    compact_text = re.sub(r"\s+", "", text)
    majors = [
        "秘书学",
        "广播电视学",
        "戏剧影视文学",
        "经济学",
        "经济统计学",
        "应用物理学",
        "房地产开发与管理",
        "社会体育指导与管理",
        "汽车服务工程",
        "表演",
    ]
    if "撤销秘书学" not in compact_text or "上报教育部" not in compact_text or "审批" not in compact_text:
        raise ValueError(f"Could not find HEBTU 2019 cancellation statement in {source_id}.")
    verification_terms = [
        "秘书学",
        "广播电视学",
        "戏剧影视文学",
        "经济学",
        "经济统",
        "计学",
        "应用物理",
        "房地产开发",
        "与管理",
        "社会体",
        "育指导与管理",
        "汽车服务工程",
        "表演",
    ]
    for term in verification_terms:
        if term not in compact_text:
            raise ValueError(f"Could not verify {term} in {source_id}.")
    rows: list[dict[str, Any]] = []
    for index, major in enumerate(majors, start=1):
        rows.append(
            {
                "policy_year": 2019,
                "region": "河北师范大学",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "河北师范大学2019年度职能绩效管理目标撤销本科专业",
                "reported_major_name": major,
                "major_code": "",
                "source_row_no": "",
                "study_duration": "",
                "policy_action": "学校2019年度职能绩效管理目标完成情况列为撤销并已上报教育部备案审批；教育部2019年度结果公布撤销。",
                "criterion_text": "河北师范大学2019年度职能绩效管理目标一览表列出撤销本科专业并说明已上报教育部备案审批。",
                "source_ids": source_id,
                "evidence_text": (
                    "河北师范大学2019年度职能绩效管理目标一览表称学校撤销秘书学、广播电视学、"
                    "戏剧影视文学、经济学、经济统计学、应用物理学、房地产开发与管理、"
                    f"社会体育指导与管理、汽车服务工程、表演等10个本科专业并已上报教育部备案审批；本行对应第{index}个：{major}。"
                ),
                "confidence": "high",
            }
        )
    return rows


def parse_aufe_2019_undergrad_stop_rows(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "aufe_2019_undergrad_stop_list_html"
    raw_candidates = sorted(
        path
        for path in raw_dir.glob(f"{source_id}.*")
        if path.is_file() and path.suffix.lower() != ".txt"
    )
    if not raw_candidates:
        return []
    html = decode_text(raw_candidates[0].read_bytes(), "")
    soup = BeautifulSoup(html, "html.parser")
    parsed: list[dict[str, Any]] = []
    for table in soup.find_all("table"):
        rows = table.find_all("tr")
        if not rows:
            continue
        header = [cell.get_text(" ", strip=True) for cell in rows[0].find_all(["td", "th"])]
        if "招生状态" not in header or "专业代码" not in header:
            continue
        for row_index, tr in enumerate(rows[1:], start=1):
            cells = [cell.get_text(" ", strip=True) for cell in tr.find_all(["td", "th"])]
            if len(cells) < len(header):
                continue
            values = dict(zip(header, cells, strict=False))
            if values.get("招生状态") != "当年停招":
                continue
            school_major = values.get("校内专业名称", "").strip()
            major = values.get("专业名称", "").strip()
            code = normalize_source_major_code(values.get("专业代码", "").strip())
            school_unit = values.get("所属单位名称", "").strip()
            opened_year = values.get("专业设置年份", "").strip()
            degree = values.get("授予学位门类", "").strip()
            new_major = values.get("是否新专业", "").strip()
            if not school_major or not major or not code:
                continue
            parsed.append(
                {
                    "policy_year": 2019,
                    "region": "安徽财经大学",
                    "education_level": "本科",
                    "record_type": "major_stop_enrollment",
                    "warning_label": "安徽财经大学2019年当年停招本科专业名单",
                    "reported_major_name": school_major,
                    "major_code": code,
                    "source_row_no": str(row_index),
                    "study_duration": "",
                    "policy_action": "学校2019年本科专业设置表列为当年停招。",
                    "criterion_text": "安徽财经大学本科专业设置、当年新增专业、停招专业名单信息一览表中的招生状态为“当年停招”。",
                    "source_ids": source_id,
                    "evidence_text": (
                        f"安徽财经大学2019年本科专业设置表第{row_index}行列出{school_major}（标准专业名称{major}，"
                        f"专业代码{code}，{degree}，所属单位{school_unit}，设置年份{opened_year}），招生状态为当年停招。"
                        f"是否新专业：{new_major}。"
                    ),
                    "confidence": "high",
                }
            )
    expected = 14
    if len(parsed) != expected:
        raise ValueError(f"Parsed {len(parsed)} rows from {source_id}, expected {expected}.")
    return parsed


def parse_aufe_2019_cancel_notice_rows(raw_dir: Path) -> list[dict[str, Any]]:
    page_id = "aufe_2019_cancel_notice_page"
    pdf_id = "aufe_2019_cancel_notice_pdf"
    page_text_path = raw_dir / f"{page_id}.txt"
    pdf_text_path = raw_dir / f"{pdf_id}.txt"
    pdf_raw_path = raw_dir / f"{pdf_id}.pdf"
    if not page_text_path.exists() and not pdf_raw_path.exists():
        return []
    page_text = page_text_path.read_text(encoding="utf-8", errors="replace") if page_text_path.exists() else ""
    pdf_text = pdf_text_path.read_text(encoding="utf-8", errors="replace") if pdf_text_path.exists() else ""
    combined_text = page_text + "\n" + pdf_text
    if "关于撤销相关本科专业" not in combined_text:
        raise ValueError(f"Could not verify AUFE cancellation notice title in {page_id}/{pdf_id}.")
    if pdf_raw_path.exists() and not pdf_raw_path.read_bytes().startswith(b"%PDF"):
        raise ValueError(f"{pdf_id} is not a PDF file.")
    rows_data = [
        ("公共事业管理", "120401"),
        ("文化产业管理", "120210"),
        ("体育经济与管理", "120212T"),
        ("管理科学", "120101"),
        ("房地产开发与管理", "120104"),
        ("国际政治", "030202"),
        ("社会工作", "030302"),
        ("服装与服饰设计", "130505"),
        ("广播电视编导", "130305"),
        ("英语", "050201"),
        ("汉语言文学", "050101"),
        ("传播学", "050304"),
        ("金融学（国际金融）", "020301K"),
    ]
    rows: list[dict[str, Any]] = []
    for index, (major, code) in enumerate(rows_data, start=1):
        rows.append(
            {
                "policy_year": 2019,
                "region": "安徽财经大学",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "安徽财经大学2019年撤销本科专业（方向）",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(index),
                "study_duration": "",
                "policy_action": "学校校政字〔2019〕39号通知列为撤销本科专业（方向）；教育部2019年度结果对其中本科专业公布撤销。",
                "criterion_text": "安徽财经大学关于撤销相关本科专业（方向）的通知列出撤销专业（方向）名单。",
                "source_ids": f"{page_id};{pdf_id}",
                "evidence_text": (
                    "安徽财经大学校政字〔2019〕39号《关于撤销相关本科专业（方向）的通知》称，经2019年3月20日"
                    "校长办公会研究决定，对公共事业管理等13个专业（方向）实施撤销；"
                    f"本行对应第{index}项：{major}。"
                ),
                "confidence": "high",
            }
        )
    return rows


SHANGHAI_UNDERGRAD_WARNING_CONFIGS = [
    {
        "policy_year": 2012,
        "source_id": "shanghai_2012_undergrad_warning_doc",
        "notice_ref": "沪教委高〔2012〕4号",
        "majors": [
            "社会工作",
            "社会体育",
            "广告学",
            "艺术设计",
            "表演",
            "动画",
            "播音与主持艺术",
            "广播电视编导",
            "信息与计算科学",
            "材料化学",
            "电子信息工程",
            "网络工程",
            "信息显示与光电技术",
            "食品质量与安全",
            "国际商务",
            "公共事业管理",
            "劳动与社会保障",
            "会展经济与管理",
        ],
    },
    {
        "policy_year": 2013,
        "source_id": "shanghai_2013_undergrad_warning_doc",
        "notice_ref": "沪教委高〔2013〕9号",
        "majors": [
            "日语",
            "信息管理与信息系统",
            "工商管理",
            "艺术设计",
            "市场营销",
            "物流管理",
            "行政管理",
            "公共事业管理",
            "电子信息工程",
            "环境工程",
            "信息与计算科学",
            "广告学",
            "电子商务",
            "社会工作",
            "交通运输",
        ],
    },
    {
        "policy_year": 2014,
        "source_id": "shanghai_2014_undergrad_warning_doc",
        "source_ids": ["shanghai_2014_undergrad_warning_doc", "moe_2014_shanghai_warning_news"],
        "notice_ref": "沪教委高〔2014〕11号",
        "majors": [
            "信息管理与信息系统",
            "市场营销",
            "电子信息工程",
            "信息与计算科学",
            "应用物理学",
            "社会工作",
            "社会学",
        ],
    },
    {
        "policy_year": 2016,
        "source_id": "shanghai_2016_undergrad_warning_doc",
        "notice_ref": "沪教委高〔2016〕40号",
        "majors": [
            "英语",
            "国际经济与贸易",
            "法学",
            "工商管理",
            "物流管理",
            "新闻学",
            "旅游管理",
            "信息管理与信息系统",
            "市场营销",
            "行政管理",
        ],
    },
]


def source_doc_text(raw_dir: Path, source_id: str) -> str:
    text_path = raw_dir / f"{source_id}.txt"
    if text_path.exists():
        text = text_path.read_text(encoding="utf-8", errors="replace")
        if text.strip():
            return text
    raw_path = raw_dir / f"{source_id}.doc"
    if not raw_path.exists():
        return ""
    with tempfile.TemporaryDirectory() as tmpdir:
        converted_path = Path(tmpdir) / f"{source_id}.docx"
        convert_doc_with_word(raw_path, converted_path)
        return docx_to_text(converted_path)


def parse_shanghai_undergrad_warning_rows(raw_dir: Path) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for config in SHANGHAI_UNDERGRAD_WARNING_CONFIGS:
        text = source_doc_text(raw_dir, config["source_id"])
        missing = [major for major in config["majors"] if major not in text]
        if missing:
            raise ValueError(f"Missing expected majors in {config['source_id']}: {missing}")
        source_ids = config.get("source_ids", [config["source_id"]])
        for source_row_no, major in enumerate(config["majors"], start=1):
            rows.append(
                {
                    "policy_year": config["policy_year"],
                    "region": "上海市",
                    "education_level": "本科",
                    "record_type": "major_warning_list",
                    "warning_label": f"{config['policy_year']}年度上海市本科预警专业名单",
                    "reported_major_name": major,
                    "policy_action": "减少预警专业招生总量；对办学条件严重不足、教学质量低下的专业严格控制招生计划，甚至暂停招生；拟增设已列入预警范围专业原则上不予受理备案申请。",
                    "criterion_text": "本市高校中重复设置较多，连续多年招生第一志愿录取率偏低、调剂和征求志愿录取率偏高，且毕业生签约率偏低。",
                    "source_row_no": str(source_row_no),
                    "source_ids": ";".join(source_ids),
                    "evidence_text": f"上海市教委{config['notice_ref']}正文列出{major}为{config['policy_year']}年度本科预警专业。",
                    "confidence": "high",
                }
            )
    return rows


HUST_HISTORICAL_STOP_CONFIGS = [
    (2025, "hust_2025_stop_enrollment", 17),
    (2024, "hust_2024_stop_enrollment_empty", 0),
    (2023, "hust_2023_stop_enrollment_empty", 0),
    (2022, "hust_2022_stop_enrollment_empty", 0),
    (2021, "hust_2021_stop_enrollment", 13),
    (2020, "hust_2020_stop_enrollment", 13),
    (2019, "hust_2019_stop_enrollment", 9),
    (2018, "hust_2018_stop_enrollment", 5),
    (2016, "hust_2016_stop_enrollment_empty", 0),
]


HUST_STOP_MARKERS = ("暂停招生", "停止招生", "已停止招生", "暂未招生", "未招生")


def parse_hust_historical_stop_rows(raw_dir: Path) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    code_pattern = re.compile(r"^\d{6}[A-Z]{0,2}$")
    for policy_year, source_id, expected_count in HUST_HISTORICAL_STOP_CONFIGS:
        raw_candidates = sorted(
            path for path in raw_dir.glob(f"{source_id}.*") if path.is_file() and path.suffix.lower() != ".txt"
        )
        if not raw_candidates:
            raise FileNotFoundError(f"Missing raw HTML for {source_id}")
        html = decode_text(raw_candidates[0].read_bytes(), "")
        soup = BeautifulSoup(html, "html.parser")
        parsed: list[dict[str, Any]] = []
        for table in soup.find_all("table"):
            last_school_unit = ""
            for table_row_no, table_row in enumerate(table.find_all("tr"), start=1):
                cells = [
                    " ".join(cell.get_text(" ", strip=True).replace("\xa0", " ").split())
                    for cell in table_row.find_all(["th", "td"])
                ]
                if len(cells) >= 6:
                    last_school_unit = cells[0]
                    school_unit, code, major, duration, degree, note = cells[:6]
                elif len(cells) == 5 and last_school_unit:
                    school_unit = last_school_unit
                    code, major, duration, degree, note = cells[:5]
                else:
                    continue
                if not code_pattern.match(code) or not any(marker in note for marker in HUST_STOP_MARKERS):
                    continue
                parsed.append(
                    {
                        "policy_year": policy_year,
                        "region": "华中科技大学",
                        "education_level": "本科",
                        "record_type": "major_stop_enrollment",
                        "warning_label": f"{policy_year}年本科专业设置停招/未招生专业",
                        "reported_major_name": major,
                        "major_code": code,
                        "study_duration": duration,
                        "policy_action": f"{note}。",
                        "criterion_text": f"华中科技大学{policy_year}年本科专业设置、新增专业、停招专业名单表备注列标记停招、暂停招生或未招生。",
                        "source_row_no": str(table_row_no),
                        "source_ids": source_id,
                        "evidence_text": f"华中科技大学{policy_year}年本科专业设置表第{table_row_no}行列出{school_unit}{major}（{code}，{duration}，{degree}），备注为{note}。",
                        "confidence": "high",
                    }
                )
        if len(parsed) != expected_count:
            raise ValueError(f"Parsed {len(parsed)} rows from {source_id}, expected {expected_count}.")
        rows.extend(parsed)
    return rows


SUFE_HISTORICAL_STOP_CONFIGS = [
    {
        "policy_year": 2024,
        "source_id": "sufe_2024_stop_enrollment",
        "mode": "symbol",
        "symbol": "◆",
        "expected_count": 11,
        "warning_label": "2024本科专业设置情况已停招专业",
        "policy_action": "已停招。",
        "criterion_text": "上海财经大学2024本科专业设置HTML表注释说明标“◆”为已停招专业。",
    },
    {
        "policy_year": 2023,
        "source_id": "sufe_2023_stop_enrollment",
        "mode": "symbol",
        "symbol": "◆",
        "expected_count": 11,
        "warning_label": "2023本科专业设置情况已停招专业",
        "policy_action": "已停招。",
        "criterion_text": "上海财经大学2023本科专业设置HTML表注释说明标“◆”为已停招专业。",
    },
    {
        "policy_year": 2022,
        "source_id": "sufe_2022_stop_enrollment",
        "mode": "symbol",
        "symbol": "◆",
        "expected_count": 12,
        "warning_label": "2022本科专业设置情况已停招专业",
        "policy_action": "已停招。",
        "criterion_text": "上海财经大学2022本科专业设置HTML表注释说明标“◆”为已停招专业。",
    },
    {
        "policy_year": 2021,
        "source_id": "sufe_2021_stop_enrollment",
        "mode": "symbol",
        "symbol": "◆",
        "expected_count": 12,
        "warning_label": "2020-2021学年本科专业设置已停招专业",
        "policy_action": "已停招。",
        "criterion_text": "上海财经大学2020-2021学年学科专业设置HTML表注释说明标“◆”为已停招专业。",
    },
    {
        "policy_year": 2020,
        "source_id": "sufe_2020_stop_enrollment",
        "mode": "symbol",
        "symbol": "◆",
        "expected_count": 9,
        "warning_label": "2020本科专业设置情况已停招专业",
        "policy_action": "已停招。",
        "criterion_text": "上海财经大学2020本科专业设置HTML表注释说明标“◆”为已停招专业。",
    },
    {
        "policy_year": 2019,
        "source_id": "sufe_2019_stop_enrollment",
        "mode": "symbol",
        "symbol": "■",
        "expected_count": 11,
        "warning_label": "2019本科专业设置情况停招专业",
        "policy_action": "停招。",
        "criterion_text": "上海财经大学2019本科专业设置HTML表注释说明标“■”为停招专业。",
    },
    {
        "policy_year": 2018,
        "source_id": "sufe_2018_stop_enrollment",
        "mode": "stars_2018",
        "expected_count": 12,
        "warning_label": "2018本科专业设置情况停招/停止招生专业",
        "criterion_text": "上海财经大学2018本科专业设置HTML表注释说明标“***”为已停招专业，标“****”为2018年停止招生专业。",
    },
    {
        "policy_year": 2017,
        "source_id": "sufe_2017_stop_enrollment",
        "mode": "stars_2017",
        "expected_count": 10,
        "warning_label": "2017本科专业设置情况停止招生专业",
        "policy_action": "2017年停止招生。",
        "criterion_text": "上海财经大学2017本科专业设置HTML表注释说明标“***”为2017年停止招生专业。",
    },
    {
        "policy_year": 2016,
        "source_id": "sufe_2014_2016_stop_enrollment",
        "mode": "remark",
        "expected_count": 5,
        "warning_label": "2014、2015、2016级本科专业设置停招/未招生专业",
        "criterion_text": "上海财经大学2014、2015、2016级本科专业设置HTML表备注列标注停招或未招生。",
    },
]


def normalize_sufe_cell(value: Any) -> str:
    text = str(value if value is not None else "").replace("\xa0", " ")
    text = "" if text.lower() == "nan" else text
    return " ".join(text.split())


def parse_sufe_major_cell(value: str) -> tuple[str, str]:
    text = normalize_sufe_cell(value)
    match = re.search(r"[（(]\s*(\d{6}[A-Z]{0,2})\s*[）)]", text)
    if not match:
        return "", ""
    major = text[: match.start()]
    major = re.sub(r"[◆■☆*]+", "", major)
    major = re.sub(r"\s+", "", major).strip()
    return major, match.group(1)


def sufe_historical_row(
    *,
    policy_year: int,
    source_id: str,
    source_row_no: str,
    major: str,
    code: str,
    degree: str,
    warning_label: str,
    policy_action: str,
    criterion_text: str,
    evidence_detail: str,
) -> dict[str, Any]:
    degree_phrase = f"，{degree}" if degree else ""
    return {
        "policy_year": policy_year,
        "region": "上海财经大学",
        "education_level": "本科",
        "record_type": "major_stop_enrollment",
        "warning_label": warning_label,
        "reported_major_name": major,
        "major_code": code,
        "policy_action": policy_action,
        "criterion_text": criterion_text,
        "source_row_no": source_row_no,
        "source_ids": source_id,
        "evidence_text": f"上海财经大学{policy_year}年本科专业设置表第{source_row_no}项列出{major}（{code}{degree_phrase}），{evidence_detail}",
        "confidence": "high",
    }


def parse_sufe_historical_stop_rows(raw_dir: Path) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for config in SUFE_HISTORICAL_STOP_CONFIGS:
        source_id = str(config["source_id"])
        raw_candidates = sorted(
            path for path in raw_dir.glob(f"{source_id}.*") if path.is_file() and path.suffix.lower() != ".txt"
        )
        if not raw_candidates:
            raise FileNotFoundError(f"Missing raw HTML for {source_id}")
        html = decode_text(raw_candidates[0].read_bytes(), "")
        parsed: list[dict[str, Any]] = []
        for table in pd.read_html(StringIO(html)):
            table = table.fillna("")
            for _, table_row in table.iterrows():
                cells = [normalize_sufe_cell(value) for value in table_row.tolist()]
                if str(config["mode"]) == "remark":
                    if len(cells) < 6:
                        continue
                    source_row_no, code, major, degree, _, note = cells[:6]
                    if not source_row_no.isdigit() or note not in {"停招", "未招生"}:
                        continue
                    parsed.append(
                        sufe_historical_row(
                            policy_year=int(config["policy_year"]),
                            source_id=source_id,
                            source_row_no=source_row_no,
                            major=major,
                            code=code,
                            degree=degree,
                            warning_label=str(config["warning_label"]),
                            policy_action=f"{note}。",
                            criterion_text=str(config["criterion_text"]),
                            evidence_detail=f"备注为{note}。",
                        )
                    )
                    continue

                if len(cells) < 3 or not cells[0].isdigit():
                    continue
                source_row_no, major_cell, degree = cells[:3]
                mode = str(config["mode"])
                policy_action = str(config.get("policy_action", ""))
                evidence_detail = ""
                if mode == "symbol":
                    symbol = str(config["symbol"])
                    if symbol not in major_cell:
                        continue
                    evidence_detail = f"并以{symbol}标注为{policy_action.rstrip('。')}。"
                elif mode == "stars_2017":
                    star_match = re.search(r"\*+", major_cell)
                    if not star_match or len(star_match.group(0)) != 3:
                        continue
                    evidence_detail = "并以***标注为2017年停止招生专业。"
                elif mode == "stars_2018":
                    star_match = re.search(r"\*+", major_cell)
                    if not star_match:
                        continue
                    star_count = len(star_match.group(0))
                    if star_count == 3:
                        policy_action = "已停招。"
                        evidence_detail = "并以***标注为已停招专业。"
                    elif star_count == 4:
                        policy_action = "2018年停止招生。"
                        evidence_detail = "并以****标注为2018年停止招生专业。"
                    else:
                        continue
                else:
                    continue
                major, code = parse_sufe_major_cell(major_cell)
                if not major or not code:
                    continue
                parsed.append(
                    sufe_historical_row(
                        policy_year=int(config["policy_year"]),
                        source_id=source_id,
                        source_row_no=source_row_no,
                        major=major,
                        code=code,
                        degree=degree,
                        warning_label=str(config["warning_label"]),
                        policy_action=policy_action,
                        criterion_text=str(config["criterion_text"]),
                        evidence_detail=evidence_detail,
                    )
                )
        expected_count = int(config["expected_count"])
        if len(parsed) != expected_count:
            raise ValueError(f"Parsed {len(parsed)} rows from {source_id}, expected {expected_count}.")
        rows.extend(parsed)
    return rows


JNU_HISTORICAL_STOP_CONFIGS = [
    (2024, "jnu_2024_stop_enrollment", 20),
    (2023, "jnu_2023_stop_enrollment", 20),
    (2022, "jnu_2022_stop_enrollment", 19),
    (2021, "jnu_2021_stop_enrollment", 14),
    (2020, "jnu_2020_stop_enrollment", 12),
    (2019, "jnu_2019_stop_enrollment", 2),
    (2018, "jnu_2018_stop_enrollment", 1),
]


NJTECH_HISTORICAL_STOP_CONFIGS = [
    (2024, "njtech_2024_stop_enrollment", 11),
    (2023, "njtech_2023_stop_enrollment", 5),
    (2022, "njtech_2022_stop_enrollment", 4),
    (2021, "njtech_2021_stop_enrollment", 4),
    (2020, "njtech_2020_stop_enrollment", 7),
    (2019, "njtech_2019_stop_enrollment", 9),
]


UJS_HISTORICAL_STOP_CONFIGS = [
    (2024, "ujs_2024_stop_enrollment", 2),
    (2023, "ujs_2023_stop_enrollment", 6),
    (2022, "ujs_2022_stop_enrollment", 11),
    (2021, "ujs_2021_stop_enrollment", 13),
    (2020, "ujs_2020_stop_enrollment", 13),
]


def parse_ujs_historical_stop_rows(raw_dir: Path) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    code_pattern = re.compile(r"^\d{6}[A-Z]{0,2}$")
    for page_year, source_id, expected_count in UJS_HISTORICAL_STOP_CONFIGS:
        raw_candidates = sorted(
            path for path in raw_dir.glob(f"{source_id}.*") if path.is_file() and path.suffix.lower() != ".txt"
        )
        if not raw_candidates:
            raise FileNotFoundError(f"Missing raw HTML for {source_id}")
        html = decode_text(raw_candidates[0].read_bytes(), "")
        parsed: list[dict[str, Any]] = []
        seen_keys: set[tuple[str, str, str, str]] = set()
        for table in pd.read_html(StringIO(html)):
            table = table.fillna("")
            for _, table_row in table.iterrows():
                cells = [str(value).strip() for value in table_row.tolist()]
                if len(cells) == 10:
                    source_row_no, school_unit, code, major, discipline, major_class, teacher_flag, duration, degree, note = cells[:10]
                elif len(cells) == 9:
                    source_row_no, school_unit, code, major, discipline, teacher_flag, duration, degree, note = cells[:9]
                    major_class = ""
                else:
                    continue
                if not source_row_no.isdigit() or not code_pattern.match(code) or "停招" not in note:
                    continue
                key = (source_row_no, code, major, note)
                if key in seen_keys:
                    continue
                seen_keys.add(key)
                match = re.search(r"(\d{4})年停招", note)
                policy_year = int(match.group(1)) if match else page_year
                class_phrase = f"{major_class}，" if major_class else ""
                teacher_phrase = f"{teacher_flag}，" if teacher_flag else ""
                parsed.append(
                    {
                        "policy_year": policy_year,
                        "region": "江苏大学",
                        "education_level": "本科",
                        "record_type": "major_stop_enrollment",
                        "warning_label": f"{policy_year}年本科专业停招情况",
                        "reported_major_name": major,
                        "major_code": code,
                        "study_duration": duration,
                        "policy_action": f"{note}。",
                        "criterion_text": f"江苏大学{page_year}年本科专业设置表备注/是否停招列标记停招专业。",
                        "source_row_no": source_row_no,
                        "source_ids": source_id,
                        "evidence_text": f"江苏大学{page_year}年本科专业设置表第{source_row_no}项列出{school_unit}{major}（{code}，{discipline}，{class_phrase}{teacher_phrase}{duration}，{degree}），备注为{note}。",
                        "confidence": "high",
                    }
                )
        if len(parsed) != expected_count:
            raise ValueError(f"Parsed {len(parsed)} rows from {source_id}, expected {expected_count}.")
        rows.extend(parsed)
    return rows


def parse_njtech_historical_stop_rows(raw_dir: Path) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for page_year, source_id, expected_count in NJTECH_HISTORICAL_STOP_CONFIGS:
        raw_candidates = sorted(
            path for path in raw_dir.glob(f"{source_id}.*") if path.is_file() and path.suffix.lower() != ".txt"
        )
        if not raw_candidates:
            raise FileNotFoundError(f"Missing raw HTML for {source_id}")
        html = decode_text(raw_candidates[0].read_bytes(), "")
        parsed: list[dict[str, Any]] = []
        for table in pd.read_html(StringIO(html)):
            table = table.fillna("")
            for _, table_row in table.iterrows():
                cells = [str(value).strip() for value in table_row.tolist()]
                if len(cells) < 7 or "停招" not in cells[6]:
                    continue
                source_row_no, school_unit, major, code, major_class, degree, note = cells[:7]
                match = re.search(r"(\d{4})年停招", note)
                policy_year = int(match.group(1)) if match else page_year
                parsed.append(
                    {
                        "policy_year": policy_year,
                        "region": "南京工业大学",
                        "education_level": "本科",
                        "record_type": "major_stop_enrollment",
                        "warning_label": f"{policy_year}年本科专业停招情况",
                        "reported_major_name": major,
                        "major_code": code,
                        "study_duration": "四年",
                        "policy_action": f"{note}。",
                        "criterion_text": f"南京工业大学{page_year}年本科专业设置及新增、停招情况表备注列标记停招专业。",
                        "source_row_no": source_row_no,
                        "source_ids": source_id,
                        "evidence_text": f"南京工业大学{page_year}年本科专业设置及新增、停招情况表第{source_row_no}项列出{school_unit}{major}（{code}，{major_class}，四年，{degree}），备注为{note}。",
                        "confidence": "high",
                    }
                )
        if len(parsed) != expected_count:
            raise ValueError(f"Parsed {len(parsed)} rows from {source_id}, expected {expected_count}.")
        rows.extend(parsed)
    return rows


def parse_jnu_historical_stop_rows(raw_dir: Path) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for policy_year, source_id, expected_count in JNU_HISTORICAL_STOP_CONFIGS:
        raw_candidates = sorted(
            path for path in raw_dir.glob(f"{source_id}.*") if path.is_file() and path.suffix.lower() != ".txt"
        )
        if not raw_candidates:
            raise FileNotFoundError(f"Missing raw HTML for {source_id}")
        html = decode_text(raw_candidates[0].read_bytes(), "")
        parsed: list[dict[str, Any]] = []
        for table in pd.read_html(StringIO(html)):
            table = table.fillna("")
            for _, table_row in table.iterrows():
                cells = [str(value).strip() for value in table_row.tolist()]
                if len(cells) < 8 or "停招专业" not in cells[7]:
                    continue
                source_row_no, code, major, major_class, _, degree, duration, note = cells[:8]
                if code.isdigit() and len(code) < 6:
                    code = code.zfill(6)
                parsed.append(
                    {
                        "policy_year": policy_year,
                        "region": "暨南大学",
                        "education_level": "本科",
                        "record_type": "major_stop_enrollment",
                        "warning_label": f"{policy_year}年专业设置一览表停招专业",
                        "reported_major_name": major,
                        "major_code": code,
                        "study_duration": duration,
                        "policy_action": "停招。",
                        "criterion_text": f"暨南大学{policy_year}年专业设置一览表备注列标记为停招专业。",
                        "source_row_no": source_row_no,
                        "source_ids": source_id,
                        "evidence_text": f"暨南大学{policy_year}年专业设置一览表第{source_row_no}项列出{major}（{code}，{major_class}，{duration}，{degree}），备注为{note}。",
                        "confidence": "high",
                    }
                )
        if len(parsed) != expected_count:
            raise ValueError(f"Parsed {len(parsed)} rows from {source_id}, expected {expected_count}.")
        rows.extend(parsed)
    return rows


def merge_semicolon_values(*values: str) -> str:
    seen: set[str] = set()
    merged: list[str] = []
    for value in values:
        for part in str(value or "").split(";"):
            part = part.strip()
            if part and part not in seen:
                seen.add(part)
                merged.append(part)
    return ";".join(merged)


def merge_evidence(primary: str, secondary: str) -> str:
    primary = primary or ""
    secondary = secondary or ""
    if not primary:
        return secondary
    if not secondary or secondary == primary:
        return primary
    return f"{primary} 补充来源：{secondary}"


def merge_official_warning_rows(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    confidence_rank = {"low": 0, "medium": 1, "high": 2}
    merged: dict[tuple[Any, ...], dict[str, Any]] = {}
    official_broad_index: dict[tuple[Any, ...], list[tuple[Any, ...]]] = defaultdict(list)

    def broad_key(row: dict[str, Any]) -> tuple[Any, ...]:
        return (
            row.get("policy_year", ""),
            row.get("region", ""),
            row.get("education_level", ""),
            row.get("record_type", ""),
            row.get("reported_major_name", ""),
        )

    def precise_key(row: dict[str, Any]) -> tuple[Any, ...]:
        return broad_key(row) + (row.get("major_code", ""), row.get("study_duration", ""), row.get("source_row_no", ""))

    ordered_rows = sorted(
        rows,
        key=lambda item: 0
        if any(source_id in item.get("source_ids", "") for source_id in MOE_COMPLETE_CANCEL_SOURCE_IDS)
        else 1,
    )
    for row in ordered_rows:
        broad = broad_key(row)
        is_official_complete = any(source_id in row.get("source_ids", "") for source_id in MOE_COMPLETE_CANCEL_SOURCE_IDS)
        if is_official_complete:
            key = precise_key(row)
            official_broad_index[broad].append(key)
        elif not row.get("major_code") and official_broad_index.get(broad):
            key = official_broad_index[broad][0]
        else:
            key = precise_key(row)
            if not is_official_complete:
                for existing_key in merged:
                    if existing_key[:5] != broad:
                        continue
                    key_has_empty_detail = not all(key[5:])
                    existing_has_empty_detail = not all(existing_key[5:])
                    details_match = key[5:] == existing_key[5:]
                    if details_match or key_has_empty_detail or existing_has_empty_detail:
                        key = existing_key
                        break
        existing = merged.get(key)
        if existing is None:
            merged[key] = dict(row)
            continue

        row_rank = confidence_rank.get(row.get("confidence", ""), -1)
        existing_rank = confidence_rank.get(existing.get("confidence", ""), -1)
        prefer_new = row_rank > existing_rank or (
            row_rank == existing_rank and bool(row.get("major_code")) and not existing.get("major_code")
        )
        primary = dict(row if prefer_new else existing)
        secondary = existing if prefer_new else row
        primary["source_ids"] = merge_semicolon_values(primary.get("source_ids", ""), secondary.get("source_ids", ""))
        primary["evidence_text"] = merge_evidence(primary.get("evidence_text", ""), secondary.get("evidence_text", ""))
        for field in ["major_code", "study_duration", "source_row_no"]:
            if not primary.get(field) and secondary.get(field):
                primary[field] = secondary[field]
        primary["confidence"] = "high" if max(row_rank, existing_rank) >= confidence_rank["high"] else primary.get("confidence", "")
        merged[key] = primary
    return list(merged.values())


def enrich_records(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    catalog = read_catalog()
    enriched: list[dict[str, Any]] = []
    seen: set[str] = set()
    for row in rows:
        standard_name = LEVEL_STANDARD_NAME_ALIASES.get(
            (row["reported_major_name"], row["education_level"]),
            STANDARD_NAME_ALIASES.get(row["reported_major_name"], row["reported_major_name"]),
        )
        catalog_level = "本科" if row["education_level"] == "本科" else "专科"
        catalog_name = CATALOG_LOOKUP_ALIASES.get(standard_name, standard_name)
        catalog_row = catalog.get((catalog_name, catalog_level), catalog.get((catalog_name, ""), {}))
        record_id_base = "|".join(
            [
                str(row["report_year"]),
                row["education_level"],
                row["risk_level"],
                row["reported_major_name"],
            ]
        )
        record_id = hashlib.sha1(record_id_base.encode("utf-8")).hexdigest()[:16]
        if record_id in seen:
            continue
        seen.add(record_id)
        updated = {
            "record_id": record_id,
            "schema_version": SCHEMA_VERSION,
            "standard_major_name": standard_name,
            "major_code": catalog_row.get("major_code", ""),
            "discipline": catalog_row.get("category", ""),
            "major_category": catalog_row.get("subject", ""),
            "captured_at": CAPTURED_AT,
            **row,
        }
        enriched.append(updated)
    return sorted(enriched, key=lambda item: (item["report_year"], item["education_level"], item["risk_level"], item["reported_major_name"]))


def parse_hebei_2018_initial_employment_metrics(raw_dir: Path) -> list[dict[str, Any]]:
    source_id = "hebei_2019_undergrad_application_policy_pdf_mirror"
    text_path = raw_dir / f"{source_id}.txt"
    if not text_path.exists():
        return []
    text = text_path.read_text(encoding="utf-8", errors="replace")
    start = text.find("2018届毕业生本科毕业生分专业初次就业率统计表")
    end = text.find("附件3", start)
    if start < 0 or end < 0:
        return []
    section = text[start:end]
    pattern = re.compile(
        r"^\s*(?P<major>[^\d\s][^\d\n]*?)\s+(?P<graduates>\d+)\s+(?P<employed>\d+)\s+(?P<rate>\d+(?:\.\d+)?)\s*$",
        re.M,
    )
    parsed_rows: list[tuple[str, int, int, float]] = []
    for match in pattern.finditer(section):
        major = re.sub(r"\s+", "", match.group("major"))
        if major in {"学科门类"}:
            continue
        parsed_rows.append(
            (
                major,
                int(match.group("graduates")),
                int(match.group("employed")),
                float(match.group("rate")),
            )
        )
    expected_count = 308
    if len(parsed_rows) != expected_count:
        raise ValueError(f"Parsed {len(parsed_rows)} Hebei employment-rate rows, expected {expected_count}.")

    rows: list[dict[str, Any]] = []
    for major, graduates, employed, rate in parsed_rows:
        evidence = f"河北省2018届本科毕业生分专业初次就业率统计表列出{major}：毕业生{graduates}人，就业{employed}人，初次就业率{rate:g}%。"
        rows.extend(
            [
                {
                    "report_year": 2019,
                    "graduate_cohort": 2018,
                    "education_level": "本科",
                    "reported_major_name": major,
                    "metric_name": "initial_employment_rate",
                    "metric_value": rate,
                    "metric_unit": "percent",
                    "metric_rank": "",
                    "rank_scope": "河北省2018届本科毕业生分专业初次就业率统计表",
                    "source_ids": source_id,
                    "evidence_text": evidence,
                    "confidence": "high",
                },
                {
                    "report_year": 2019,
                    "graduate_cohort": 2018,
                    "education_level": "本科",
                    "reported_major_name": major,
                    "metric_name": "graduate_count",
                    "metric_value": graduates,
                    "metric_unit": "persons",
                    "metric_rank": "",
                    "rank_scope": "河北省2018届本科毕业生分专业初次就业率统计表",
                    "source_ids": source_id,
                    "evidence_text": evidence,
                    "confidence": "high",
                },
                {
                    "report_year": 2019,
                    "graduate_cohort": 2018,
                    "education_level": "本科",
                    "reported_major_name": major,
                    "metric_name": "employed_count",
                    "metric_value": employed,
                    "metric_unit": "persons",
                    "metric_rank": "",
                    "rank_scope": "河北省2018届本科毕业生分专业初次就业率统计表",
                    "source_ids": source_id,
                    "evidence_text": evidence,
                    "confidence": "high",
                },
            ]
        )
    return rows


def curated_metric_rows(raw_dir: Path | None = None) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    if raw_dir is not None:
        rows.extend(parse_hebei_2018_initial_employment_metrics(raw_dir))

    def add_metric(
        report_year: int,
        graduate_cohort: int,
        level: str,
        major: str,
        metric_name: str,
        metric_value: float | int,
        metric_unit: str,
        source_ids: list[str],
        evidence_text: str,
        metric_rank: int | str = "",
        rank_scope: str = "",
        confidence: str = "high",
    ) -> None:
        rows.append(
            {
                "report_year": report_year,
                "graduate_cohort": graduate_cohort,
                "education_level": level,
                "reported_major_name": major,
                "metric_name": metric_name,
                "metric_value": metric_value,
                "metric_unit": metric_unit,
                "metric_rank": metric_rank,
                "rank_scope": rank_scope,
                "source_ids": ";".join(source_ids),
                "evidence_text": evidence_text,
                "confidence": confidence,
            }
        )

    for major, count in [
        ("公共事业管理", 6),
        ("生物技术", 3),
        ("教育技术学", 1),
        ("汉语国际教育", 1),
    ]:
        add_metric(
            2024,
            2023,
            "本科",
            major,
            "cumulative_yellow_warning_count_2010_2024",
            count,
            "times",
            ["china_2025_undergrad_yellow"],
            f"中华网教育页面列出{major}专业2024年黄牌，2010年至2024年黄牌累计次数{count}个。",
            "",
            "2010-2024本科黄牌累计次数",
            "medium",
        )

    highvoc_2020_income = [
        ("空中乘务", 5533),
        ("铁道工程技术", 5503),
        ("铁道机车", 5385),
        ("铁道供电技术", 5225),
        ("社会体育", 5225),
        ("软件技术", 5166),
        ("云计算技术与应用", 5058),
        ("铁道交通运营管理", 5027),
        ("移动应用开发", 4952),
        ("航海技术", 4949),
        ("大数据技术与应用", 4945),
        ("数控设备应用与维护", 4861),
        ("石油化工技术", 4858),
        ("数控技术", 4839),
        ("机械制造与自动化", 4832),
        ("信息安全与管理", 4806),
        ("焊接技术与自动化", 4762),
        ("道路桥梁工程技术", 4754),
        ("医学美容技术", 4752),
        ("电力系统继电保护与自动化技术", 4741),
    ]
    for rank, (major, income) in enumerate(highvoc_2020_income, start=1):
        add_metric(
            2021,
            2020,
            "高职高专",
            major,
            "monthly_income",
            income,
            "CNY/month",
            ["ycwb_2021_highvoc_green"],
            "2020届高职专业毕业生月收入前20位。",
            rank,
            "高职高专专业毕业半年后月收入前20",
        )

    add_metric(2021, 2020, "高职高专", "铁道机车", "employment_satisfaction", 79, "percent", ["ycwb_2021_highvoc_green"], "铁道机车专业就业满意度为79%。")
    add_metric(2021, 2020, "高职高专", "铁道机车", "major_job_relevance", 88, "percent", ["ycwb_2021_highvoc_green"], "铁道机车专业工作与专业相关度为88%。")
    add_metric(2021, 2020, "高职高专", "铁道供电技术", "employment_satisfaction", 80, "percent", ["ycwb_2021_highvoc_green"], "铁道供电技术专业就业满意度为80%。")
    add_metric(2021, 2020, "高职高专", "铁道工程技术", "major_job_relevance", 83, "percent", ["ycwb_2021_highvoc_green"], "铁道工程技术专业工作与专业相关度为83%。")

    undergrad_2024_income = [
        ("信息安全", 7599),
        ("微电子科学与工程", 7282),
        ("电子科学与技术", 7215),
        ("自动化", 7108),
        ("软件工程", 7092),
        ("材料成型及控制工程", 7077),
        ("光电信息科学与工程", 7076),
        ("电子信息科学与技术", 7058),
        ("机械设计制造及其自动化", 7051),
        ("机械电子工程", 7018),
    ]
    for rank, (major, income) in enumerate(undergrad_2024_income, start=1):
        add_metric(
            2025,
            2024,
            "本科",
            major,
            "monthly_income",
            income,
            "CNY/month",
            ["people_2025_green_income"],
            "2024届本科专业毕业生月收入较高的10个专业。",
            rank,
            "本科专业毕业半年后月收入前10",
        )

    add_metric(2025, 2024, "本科", "电气工程及其自动化", "major_job_relevance", 86, "percent", ["people_2025_green"], "电气工程及其自动化工作与专业相关度为86%，排名第15。", 15, "本科专业工作与专业相关度")
    add_metric(2025, 2024, "本科", "电气工程及其自动化", "employment_satisfaction", 85, "percent", ["people_2025_green"], "电气工程及其自动化就业满意度为85%，排名第8。", 8, "本科专业就业满意度")
    add_metric(2025, 2024, "高职高专", "铁道机车运用与维护", "monthly_income", 5926, "CNY/month", ["people_2025_green_income"], "铁道机车运用与维护月收入5926元，排名第1。", 1, "高职专业毕业半年后月收入")
    add_metric(2025, 2024, "高职高专", "铁道机车运用与维护", "employment_satisfaction", 85, "percent", ["people_2025_green_income"], "铁道机车运用与维护就业满意度85%，排名第7。", 7, "高职专业就业满意度")
    add_metric(2025, 2024, "高职高专", "铁道机车运用与维护", "major_job_relevance", 84, "percent", ["people_2025_green_income"], "铁道机车运用与维护工作与专业相关度84%，排名第7。", 7, "高职专业工作与专业相关度")
    add_metric(2025, 2024, "高职高专", "应用化工技术", "monthly_income", 5481, "CNY/month", ["people_2025_green"], "应用化工技术月收入5481元，排名第13。", 13, "高职专业毕业半年后月收入")
    add_metric(2025, 2024, "高职高专", "应用化工技术", "major_job_relevance", 81, "percent", ["people_2025_green"], "应用化工技术工作与专业相关度81%，排名第9。", 9, "高职专业工作与专业相关度")
    add_metric(2026, 2025, "本科", "电气工程及其自动化", "monthly_income", 7160, "CNY/month", ["sina_2026_undergrad_green", "gkztc_2026_green"], "2025届电气工程及其自动化本科毕业生月收入7160元，排名第16。", 16, "本科专业毕业半年后月收入")
    add_metric(2026, 2025, "本科", "电气工程及其自动化", "major_job_relevance", 87, "percent", ["sina_2026_undergrad_green", "gkztc_2026_green"], "2025届电气工程及其自动化本科毕业生工作与专业相关度87%，排名第15。", 15, "本科专业工作与专业相关度")
    add_metric(2026, 2025, "本科", "电气工程及其自动化", "employment_satisfaction", 87, "percent", ["sina_2026_undergrad_green", "gkztc_2026_green"], "2025届电气工程及其自动化本科毕业生就业满意度87%，排名第3。", 3, "本科专业就业满意度")
    add_metric(2026, 2025, "本科", "全国本科就业质量统计", "average_monthly_income", 6435, "CNY/month", ["sina_2026_undergrad_green"], "麦可思-中国2025届本科生培养质量跟踪评价显示，2025届本科毕业生平均月收入为6435元。", "", "2025届本科毕业生就业质量总体指标")
    add_metric(2026, 2025, "本科", "全国本科就业质量统计", "survey_sample_graduate_count", 195000, "persons-approx", ["sina_2026_undergrad_green"], "麦可思-中国2025届本科生培养质量跟踪评价覆盖全国本科生样本约19.5万人。", "", "2025届本科毕业生就业质量总体指标")
    add_metric(2026, 2025, "本科", "全国本科就业质量统计", "survey_major_count", 522, "majors", ["sina_2026_undergrad_green"], "麦可思-中国2025届本科生培养质量跟踪评价涉及522个本科专业。", "", "2025届本科毕业生就业质量总体指标")
    add_metric(2026, 2025, "本科", "全国本科就业质量统计", "survey_occupation_count", 578, "occupations", ["sina_2026_undergrad_green"], "麦可思-中国2025届本科生培养质量跟踪评价涵盖毕业生从事的职业578个。", "", "2025届本科毕业生就业质量总体指标")
    add_metric(2026, 2025, "本科", "全国本科就业质量统计", "survey_industry_count", 322, "industries", ["sina_2026_undergrad_green"], "麦可思-中国2025届本科生培养质量跟踪评价涵盖毕业生从事的行业322个。", "", "2025届本科毕业生就业质量总体指标")
    add_metric(2026, 2025, "本科", "教育业", "employment_industry_share", 14.2, "percent", ["sina_2026_undergrad_green"], "2025届本科毕业生就业量较大的行业中，教育业占14.2%。", 1, "2025届本科毕业生就业量较大行业")
    add_metric(2026, 2025, "本科", "信息传输、软件和信息技术服务业", "employment_industry_share", 9.3, "percent", ["sina_2026_undergrad_green"], "2025届本科毕业生就业量较大的行业中，信息传输、软件和信息技术服务业占9.3%。", 2, "2025届本科毕业生就业量较大行业")
    add_metric(2026, 2025, "本科", "电子电气设备制造业", "employment_industry_share", 7.0, "percent", ["sina_2026_undergrad_green"], "2025届本科毕业生就业量较大的行业中，电子电气设备制造业占7.0%。", 3, "2025届本科毕业生就业量较大行业")
    add_metric(2026, 2025, "本科", "政府及公共管理", "employment_industry_share", 6.9, "percent", ["sina_2026_undergrad_green"], "2025届本科毕业生就业量较大的行业中，政府及公共管理占6.9%。", 4, "2025届本科毕业生就业量较大行业")
    add_metric(2026, 2025, "本科", "金融业", "employment_industry_share", 6.6, "percent", ["sina_2026_undergrad_green"], "2025届本科毕业生就业量较大的行业中，金融业占6.6%。", 5, "2025届本科毕业生就业量较大行业")
    add_metric(2026, 2025, "本科", "电子电气设备制造业", "employment_industry_share_change_since_2021", 0.8, "percentage-points", ["sina_2026_undergrad_green"], "2025届本科毕业生在电子电气设备制造业就业比例为7.0%，较2021届提高0.8个百分点。", "", "2021-2025届本科毕业生制造业就业比例变化")
    add_metric(2026, 2025, "本科", "机械设备制造业", "employment_industry_share", 4.2, "percent", ["sina_2026_undergrad_green"], "2025届本科毕业生在机械设备制造业就业比例为4.2%。", "", "2021-2025届本科毕业生制造业就业比例变化")
    add_metric(2026, 2025, "本科", "机械设备制造业", "employment_industry_share_change_since_2021", 1.6, "percentage-points", ["sina_2026_undergrad_green"], "2025届本科毕业生在机械设备制造业就业比例为4.2%，较2021届提高1.6个百分点。", "", "2021-2025届本科毕业生制造业就业比例变化")
    add_metric(2026, 2025, "本科", "交通运输设备制造业", "employment_industry_share", 2.3, "percent", ["sina_2026_undergrad_green"], "2025届本科毕业生在交通运输设备制造业就业比例为2.3%。", "", "2021-2025届本科毕业生制造业就业比例变化")
    add_metric(2026, 2025, "本科", "交通运输设备制造业", "employment_industry_share_change_since_2021", 1.0, "percentage-points", ["sina_2026_undergrad_green"], "2025届本科毕业生在交通运输设备制造业就业比例为2.3%，较2021届提高1.0个百分点。", "", "2021-2025届本科毕业生制造业就业比例变化")
    add_metric(2025, 2024, "本科", "全国本科专业设置调整统计", "national_new_major_point_count", 1839, "major-points", ["moe_2025_2024_undergrad_setting_result_news", "moe_2025_2024_undergrad_setting_qna"], "教育部政务服务平台称2024年度全国高校新增专业点1839个。", "", "2024年度普通高等学校本科专业备案和审批结果", "high")
    add_metric(2025, 2024, "本科", "全国本科专业设置调整统计", "national_degree_or_duration_adjusted_major_point_count", 157, "major-points", ["moe_2025_2024_undergrad_setting_result_news", "moe_2025_2024_undergrad_setting_qna"], "教育部政务服务平台称2024年度调整学位授予门类或修业年限专业点157个。", "", "2024年度普通高等学校本科专业备案和审批结果", "high")
    add_metric(2025, 2024, "本科", "全国本科专业设置调整统计", "national_stop_enrollment_major_point_count", 2220, "major-points", ["moe_2025_2024_undergrad_setting_result_news", "moe_2025_2024_undergrad_setting_qna"], "教育部政务服务平台称2024年度全国高校停招专业点2220个。", "", "2024年度普通高等学校本科专业备案和审批结果", "high")
    add_metric(2025, 2024, "本科", "全国本科专业设置调整统计", "national_cancel_major_point_count", 1428, "major-points", ["moe_2025_2024_undergrad_setting_result_news", "moe_2025_2024_undergrad_setting_qna", "moe_2025_2024_undergrad_cancel_notice_pdf"], "教育部政务服务平台称2024年度全国高校撤销专业点1428个；教育部教高函〔2025〕5号附件逐项列出1428个撤销本科专业点。", "", "2024年度普通高等学校本科专业备案和审批结果", "high")
    add_metric(2025, 2025, "本科", "全国本科专业目录统计", "national_undergrad_major_point_count", 62800, "major-points", ["moe_2025_2024_undergrad_setting_result_news"], "教育部政务服务平台称目前全国高校本科专业布点共有6.28万个。", "", "2025年普通高等学校本科专业目录发布新闻", "high")
    add_metric(2025, 2025, "本科", "全国本科专业目录统计", "catalog_major_category_count", 93, "major-categories", ["moe_2025_2024_undergrad_setting_result_news"], "教育部政务服务平台称2025版本科专业目录包含93个专业类。", "", "2025年普通高等学校本科专业目录发布新闻", "high")
    add_metric(2025, 2025, "本科", "全国本科专业目录统计", "catalog_major_kind_count", 845, "majors", ["moe_2025_2024_undergrad_setting_result_news", "moe_2025_2024_undergrad_setting_qna"], "教育部政务服务平台称《普通高等学校本科专业目录（2025年）》有845种专业。", "", "2025年普通高等学校本科专业目录发布新闻", "high")
    add_metric(2025, 2025, "本科", "全国本科专业目录统计", "catalog_new_major_kind_count", 29, "majors", ["moe_2025_2024_undergrad_setting_result_news", "moe_2025_2024_undergrad_setting_qna"], "教育部政务服务平台称2025版本科专业目录增列29种新专业。", "", "2025年普通高等学校本科专业目录发布新闻", "high")
    add_metric(2025, 2024, "本科", "专业设置与区域发展匹配度试点统计", "regional_matching_pilot_province_count", 5, "regions", ["moe_2025_2024_undergrad_setting_qna"], "教育部高等教育司负责人答记者问称2024年重点指导黑龙江、浙江、河南、重庆、陕西5省市率先开展高校专业设置与区域发展匹配度提升工作试点。", "", "2024年度普通高等学校本科专业设置工作答记者问", "high")
    add_metric(2025, 2024, "本科", "专业设置与区域发展匹配度试点统计", "cross_school_featured_major_cluster_count", 172, "clusters", ["moe_2025_2024_undergrad_setting_qna"], "教育部高等教育司负责人答记者问称5省市围绕区域千亿、万亿级产业集群打造172个跨校特色专业集群。", "", "2024年度普通高等学校本科专业设置工作答记者问", "high")
    add_metric(2025, 2024, "本科", "黑龙江省本科专业设置调整统计", "province_major_adjustment_ratio", 29.3, "percent", ["moe_2025_2024_undergrad_setting_qna"], "教育部高等教育司负责人答记者问称黑龙江全省高校2024年专业增撤调整比例达29.3%。", "", "2024年度普通高等学校本科专业设置工作答记者问", "high")
    add_metric(2025, 2024, "本科", "黑龙江省本科专业设置调整统计", "key_industry_supporting_major_share", 74, "percent", ["moe_2025_2024_undergrad_setting_qna"], "教育部高等教育司负责人答记者问称黑龙江支撑重点产业的专业比例达到74%。", "", "2024年度普通高等学校本科专业设置工作答记者问", "high")
    add_metric(2025, 2024, "本科", "河南省本科专业设置调整统计", "new_emerging_interdisciplinary_major_count", 254, "major-points", ["moe_2025_2024_undergrad_setting_qna"], "教育部高等教育司负责人答记者问称河南超常布局新兴交叉专业254个。", "", "2024年度普通高等学校本科专业设置工作答记者问", "high")
    add_metric(2025, 2024, "本科", "浙江省本科专业设置调整统计", "province_major_adjustment_ratio", 31.4, "percent", ["moe_2025_2024_undergrad_setting_qna"], "教育部高等教育司负责人答记者问称浙江推动省域高校专业调整比例达31.4%。", "", "2024年度普通高等学校本科专业设置工作答记者问", "high")
    add_metric(2025, 2022, "本科", "陕西省本科专业设置调整统计", "direct_key_industry_supporting_major_count", 1348, "major-points", ["moe_2025_2024_undergrad_setting_qna"], "教育部高等教育司负责人答记者问称陕西直接支撑重点产业发展的专业数量2022年为1348个。", "", "2024年度普通高等学校本科专业设置工作答记者问", "high")
    add_metric(2025, 2024, "本科", "陕西省本科专业设置调整统计", "direct_key_industry_supporting_major_count", 1602, "major-points", ["moe_2025_2024_undergrad_setting_qna"], "教育部高等教育司负责人答记者问称陕西直接支撑重点产业发展的专业数量2024年增至1602个。", "", "2024年度普通高等学校本科专业设置工作答记者问", "high")
    add_metric(2025, 2024, "本科", "陕西省本科专业设置调整统计", "direct_key_industry_supporting_major_growth", 19, "percent", ["moe_2025_2024_undergrad_setting_qna"], "教育部高等教育司负责人答记者问称陕西直接支撑重点产业发展的专业数量由2022年的1348个增至2024年的1602个、增长19%。", "", "2024年度普通高等学校本科专业设置工作答记者问", "high")
    add_metric(2025, 2024, "本科", "低空技术与工程", "extraordinary_setup_school_count", 6, "schools", ["moe_2025_2024_undergrad_setting_result_news", "moe_2025_2024_undergrad_setting_qna"], "教育部政务服务平台称教育部指导北京航空航天大学等6所高校超常增设低空技术与工程专业。", "", "2024年度普通高等学校本科专业设置工作答记者问", "high")
    add_metric(2026, 2026, "本科", "全国本科专业目录统计", "catalog_discipline_count", 13, "disciplines", ["moe_2026_undergrad_catalog_release_news"], "教育部官网新闻称2026年本科专业目录共涵盖13个门类、92个专业类、883种专业。", "", "普通高等学校本科专业目录（2026年）发布新闻", "high")
    add_metric(2026, 2026, "本科", "全国本科专业目录统计", "catalog_major_category_count", 92, "major-categories", ["moe_2026_undergrad_catalog_release_news"], "教育部官网新闻称2026年本科专业目录共涵盖13个门类、92个专业类、883种专业。", "", "普通高等学校本科专业目录（2026年）发布新闻", "high")
    add_metric(2026, 2026, "本科", "全国本科专业目录统计", "catalog_major_kind_count", 883, "majors", ["moe_2026_undergrad_catalog_release_news", "moe_2026_undergrad_catalog_pdf"], "教育部官网新闻称2026年本科专业目录共涵盖883种专业；目录PDF逐项列出专业。", "", "普通高等学校本科专业目录（2026年）发布新闻及目录PDF", "high")
    add_metric(2026, 2026, "本科", "交叉学科专业目录统计", "cross_discipline_existing_catalog_major_count", 11, "majors", ["moe_2026_undergrad_catalog_release_news", "moe_2026_undergrad_catalog_pdf"], "教育部官网新闻称2026年本科专业目录在交叉学科门类中首批列入未来机器人、交叉工程等11种目录内已有专业。", "", "普通高等学校本科专业目录（2026年）发布新闻及目录PDF", "high")
    add_metric(2026, 2026, "本科", "交叉学科专业目录统计", "cross_discipline_new_catalog_major_count", 4, "majors", ["moe_2026_undergrad_catalog_release_news", "moe_2026_undergrad_catalog_pdf"], "教育部官网新闻称2026年本科专业目录在交叉学科门类中列入具身智能、脑机科学与技术等4种本次列入目录的新专业；PDF列出具身智能、脑机科学与技术、工程互联网、深地科学与工程。", "", "普通高等学校本科专业目录（2026年）发布新闻及目录PDF", "high")
    add_metric(2026, 2026, "本科", "交叉学科专业目录统计", "cross_discipline_catalog_major_count", 15, "majors", ["moe_2026_undergrad_catalog_release_news", "moe_2026_undergrad_catalog_pdf"], "教育部官网新闻称交叉学科门类首批列入11种目录内已有专业和4种本次列入目录的新专业，合计15种；PDF交叉学科门类列出140001TK至140015T。", "", "普通高等学校本科专业目录（2026年）发布新闻及目录PDF", "high")
    add_metric(2026, 2021, "本科", "全国本科专业设置调整统计", "fourteenth_five_year_new_major_point_count", 10200, "major-points", ["moe_2026_undergrad_catalog_release_news"], "教育部官网新闻称“十四五”期间全国高校新增本科专业布点1.02万个。", "", "普通高等学校本科专业目录（2026年）发布新闻", "high")
    add_metric(2026, 2021, "本科", "全国本科专业设置调整统计", "fourteenth_five_year_cancel_or_stop_major_point_count", 12200, "major-points", ["moe_2026_undergrad_catalog_release_news"], "教育部官网新闻称“十四五”期间全国高校撤销或停招1.22万个本科专业布点。", "", "普通高等学校本科专业目录（2026年）发布新闻", "high")
    add_metric(2026, 2026, "本科", "全国本科专业设置调整统计", "national_major_adjustment_ratio_lower_bound", 10, "percent-lower-bound", ["moe_2026_undergrad_catalog_release_news"], "教育部官网新闻称今年全国高校专业调整比例首次突破10%。", "", "普通高等学校本科专业目录（2026年）发布新闻", "high")
    add_metric(2026, 2026, "本科", "具身智能", "school_new_major_count", 9, "schools", ["moe_2026_undergrad_catalog_release_news"], "教育部官网新闻称支持哈尔滨工业大学、北京航空航天大学等9所高校增设具身智能新专业。", "", "普通高等学校本科专业目录（2026年）发布新闻", "high")
    add_metric(2026, 2026, "高职高专", "全国高职专科专业设置统计", "national_enrolling_specialty_point_count", 69414, "specialty-points", ["moe_2026_vocational_specialty_setup_results_notice"], "教育部官网通知称经省级教育行政部门备案并在教育部汇总的2026年拟招生高职专科专业点共69414个。", "", "2026年高等职业教育专科专业设置备案和审批结果", "high")
    add_metric(2026, 2026, "高职高专", "高职专科国家控制布点专业审批统计", "controlled_specialty_application_count", 12, "specialty-points", ["moe_2026_vocational_specialty_setup_results_notice", "moe_2026_vocational_controlled_specialty_approval_pdf"], "教育部官网通知称2026年共受理拟新设高职专科国家控制布点专业申请12个。", "", "2026年新设高职专科国家控制布点专业审批结果", "high")
    add_metric(2026, 2026, "高职高专", "高职专科国家控制布点专业审批统计", "controlled_specialty_approved_count", 8, "specialty-points", ["moe_2026_vocational_specialty_setup_results_notice", "moe_2026_vocational_controlled_specialty_approval_pdf"], "教育部官网通知和附件列出2026年同意设置高职专科国家控制布点专业点8个。", "", "2026年新设高职专科国家控制布点专业审批结果", "high")
    add_metric(2026, 2026, "高职高专", "高职专科国家控制布点专业审批统计", "controlled_specialty_rejected_count", 4, "specialty-points", ["moe_2026_vocational_specialty_setup_results_notice", "moe_2026_vocational_controlled_specialty_approval_pdf"], "教育部官网通知和附件列出2026年不同意设置高职专科国家控制布点专业点4个。", "", "2026年新设高职专科国家控制布点专业审批结果", "high")
    add_metric(2026, 2026, "高职高专", "司法类国家控制布点专业审批统计", "controlled_specialty_approved_count", 4, "specialty-points", ["moe_2026_vocational_controlled_specialty_approval_pdf"], "教育部官网PDF附件司法类部分列出同意设置国家控制专业点4个。", "", "2026年新设高职专科国家控制布点专业审批结果", "high")
    add_metric(2026, 2026, "高职高专", "教育类国家控制布点专业审批统计", "controlled_specialty_approved_count", 4, "specialty-points", ["moe_2026_vocational_controlled_specialty_approval_pdf"], "教育部官网PDF附件教育类部分列出同意设置国家控制专业点4个。", "", "2026年新设高职专科国家控制布点专业审批结果", "high")
    add_metric(2026, 2026, "高职高专", "司法类国家控制布点专业审批统计", "controlled_specialty_rejected_count", 1, "specialty-points", ["moe_2026_vocational_controlled_specialty_approval_pdf"], "教育部官网PDF附件司法类部分列出不同意设置国家控制专业点1个。", "", "2026年新设高职专科国家控制布点专业审批结果", "high")
    add_metric(2026, 2026, "高职高专", "教育类国家控制布点专业审批统计", "controlled_specialty_rejected_count", 3, "specialty-points", ["moe_2026_vocational_controlled_specialty_approval_pdf"], "教育部官网PDF附件教育类部分列出不同意设置国家控制专业点3个。", "", "2026年新设高职专科国家控制布点专业审批结果", "high")
    add_metric(2025, 2025, "本科", "上海交通大学专业调整统计", "school_2023_cancel_major_count", 5, "majors", ["sjtu_2025_psychology_application_pdf"], "上海交通大学2025年心理学专业申请表称，2023年撤销“园林”等5个本科专业；原文使用“等”，不作为完整名单。", "", "学校近五年专业增设、停招、撤并情况", "high")
    add_metric(2025, 2025, "本科", "上海交通大学专业调整统计", "school_2024_cancel_major_count", 4, "majors", ["sjtu_2025_psychology_application_pdf"], "上海交通大学2025年心理学专业申请表称，2024年撤销“资源环境科学”等4个本科专业；原文使用“等”，不作为完整名单。", "", "学校近五年专业增设、停招、撤并情况", "high")
    add_metric(2025, 2025, "本科", "上海交通大学专业调整统计", "school_2024_stop_enrollment_major_count", 3, "majors", ["sjtu_2025_psychology_application_pdf"], "上海交通大学2025年心理学专业申请表称，2024年停招“建筑学”等3个本科专业；原文使用“等”，不作为完整名单。", "", "学校近五年专业增设、停招、撤并情况", "high")
    sjtu_yearly_metrics = [
        (2020, 71, 71, 10, 2, "sjtu_2020_stop_enrollment"),
        (2021, 77, 74, 9, 0, "sjtu_2021_stop_enrollment_empty"),
        (2022, 77, 74, 11, 1, "sjtu_2022_stop_enrollment"),
        (2023, 82, 79, 17, 1, "sjtu_2023_stop_enrollment"),
        (2024, 85, 80, 16, 2, "sjtu_2024_stop_enrollment"),
    ]
    for policy_year, total_count, enrolling_count, new_count, stop_count, source_id in sjtu_yearly_metrics:
        scope = f"上海交通大学{policy_year}年专业设置及调整情况（新增、停招）"
        add_metric(policy_year, policy_year, "本科", "上海交通大学专业设置统计", "school_undergrad_major_count", total_count, "majors", [source_id], f"{scope}表格列出本科专业总数为{total_count}个。", "", scope, "high")
        add_metric(policy_year, policy_year, "本科", "上海交通大学专业设置统计", "school_enrolling_major_count", enrolling_count, "majors", [source_id], f"{scope}表格列出在招/招生专业数为{enrolling_count}个。", "", scope, "high")
        add_metric(policy_year, policy_year, "本科", "上海交通大学专业设置统计", "school_new_major_count", new_count, "majors", [source_id], f"{scope}表格列出新专业名单共{new_count}个。", "", scope, "high")
        add_metric(policy_year, policy_year, "本科", "上海交通大学专业设置统计", "school_current_year_stop_enrollment_major_count", stop_count, "majors", [source_id], f"{scope}表格列出当年停招专业{stop_count}个。", "", scope, "high")
    add_metric(2024, 2024, "本科", "东北师范大学专业设置统计", "school_undergrad_major_count", 72, "majors", ["nenu_2024_teaching_quality_report_pdf"], "东北师范大学2023-2024学年本科教育教学质量报告称，截至2024年7月，学校设有本科专业72个。", "", "东北师范大学2023-2024学年本科教育教学质量报告本科专业设置情况", "high")
    add_metric(2024, 2024, "本科", "东北师范大学专业设置统计", "school_enrolling_major_count", 60, "majors", ["nenu_2024_teaching_quality_report_pdf"], "东北师范大学2023-2024学年本科教育教学质量报告称，学校招生专业60个（含2个只招收留学生的专业）。", "", "东北师范大学2023-2024学年本科教育教学质量报告本科专业设置情况", "high")
    add_metric(2024, 2024, "本科", "东北师范大学专业设置统计", "school_enrolling_chinese_student_major_count", 58, "majors", ["nenu_2024_teaching_quality_report_pdf"], "东北师范大学2023-2024学年本科教育教学质量报告称，招收中国学生的专业为58个。", "", "东北师范大学2023-2024学年本科教育教学质量报告本科专业设置情况", "high")
    add_metric(2024, 2024, "本科", "东北师范大学专业设置统计", "school_national_first_class_major_count", 46, "majors", ["nenu_2024_teaching_quality_report_pdf"], "东北师范大学2023-2024学年本科教育教学质量报告称，在招收中国学生的58个专业中，46个专业被认定为国家级一流本科专业建设点。", "", "东北师范大学2023-2024学年本科教育教学质量报告本科专业设置情况", "high")
    add_metric(2024, 2024, "本科", "东北师范大学专业设置统计", "school_provincial_first_class_major_count", 54, "majors", ["nenu_2024_teaching_quality_report_pdf"], "东北师范大学2023-2024学年本科教育教学质量报告称，54个专业被认定为省级一流本科专业建设点。", "", "东北师范大学2023-2024学年本科教育教学质量报告本科专业设置情况", "high")
    add_metric(2024, 2024, "本科", "东北师范大学专业调整统计", "school_cancel_major_count", 10, "majors", ["nenu_2024_teaching_quality_report_pdf"], "东北师范大学2023-2024学年本科教育教学质量报告称，学校主动撤销10个不适应经济社会发展需求和学校办学定位、且停招5年以上的本科专业；正文可见9个专业名称，明细行仅结构化可见名称。", "", "东北师范大学2023-2024学年本科教育教学质量报告专业建设", "high")
    add_metric(2024, 2024, "本科", "西华师范大学专业设置统计", "school_undergrad_major_count", 79, "majors", ["cwnu_2024_teaching_quality_report_pdf"], "西华师范大学2023-2024学年本科教学质量报告称学校现有本科专业79个。", "", "西华师范大学2023-2024学年本科教学质量报告本科专业设置情况", "high")
    add_metric(2024, 2024, "本科", "西华师范大学专业设置统计", "school_teacher_education_major_count", 21, "majors", ["cwnu_2024_teaching_quality_report_pdf"], "西华师范大学2023-2024学年本科教学质量报告称学校师范类专业21个。", "", "西华师范大学2023-2024学年本科教学质量报告本科专业设置情况", "high")
    add_metric(2024, 2024, "本科", "西华师范大学专业设置统计", "school_non_teacher_applied_major_count", 58, "majors", ["cwnu_2024_teaching_quality_report_pdf"], "西华师范大学2023-2024学年本科教学质量报告称学校非师范应用型专业58个。", "", "西华师范大学2023-2024学年本科教学质量报告本科专业设置情况", "high")
    add_metric(2024, 2024, "本科", "西华师范大学专业设置统计", "school_national_first_class_major_count", 11, "majors", ["cwnu_2024_teaching_quality_report_pdf"], "西华师范大学2023-2024学年本科教学质量报告称获批国家级一流本科专业建设点11个。", "", "西华师范大学2023-2024学年本科教学质量报告本科专业设置情况", "high")
    add_metric(2024, 2024, "本科", "西华师范大学专业设置统计", "school_provincial_first_class_major_count", 30, "majors", ["cwnu_2024_teaching_quality_report_pdf"], "西华师范大学2023-2024学年本科教学质量报告称获批省级一流本科专业建设点30个。", "", "西华师范大学2023-2024学年本科教学质量报告本科专业设置情况", "high")
    add_metric(2024, 2024, "本科", "西华师范大学专业调整统计", "school_recent_five_year_new_major_count", 9, "majors", ["cwnu_2024_teaching_quality_report_pdf"], "西华师范大学2023-2024学年本科教学质量报告称近五年新增人工智能、文物与博物馆学、飞行器控制与信息工程等9个专业。", "", "西华师范大学2023-2024学年本科教学质量报告专业布局", "high")
    add_metric(2024, 2024, "本科", "西华师范大学专业调整统计", "school_planned_application_major_count", 2, "majors", ["cwnu_2024_teaching_quality_report_pdf"], "西华师范大学2023-2024学年本科教学质量报告称规划申报国家公园建设与管理、特殊教育等专业，明确列名2个。", "", "西华师范大学2023-2024学年本科教学质量报告专业布局", "high")
    add_metric(2024, 2024, "本科", "西华师范大学专业调整统计", "school_recent_five_year_cancel_major_count", 5, "majors", ["cwnu_2024_teaching_quality_report_pdf"], "西华师范大学2023-2024学年本科教学质量报告称近五年撤销“双低”专业5个；报告未逐名列出这5个专业。", "", "西华师范大学2023-2024学年本科教学质量报告专业布局", "high")
    add_metric(2024, 2024, "本科", "福建师范大学专业设置统计", "school_undergrad_major_count", 85, "majors", ["fjnu_2024_teaching_quality_report_pdf"], "福建师范大学2023-2024学年本科教学质量报告称学校共有本科专业85个。", "", "福建师范大学2023-2024学年本科教学质量报告本科专业设置与分布", "high")
    add_metric(2024, 2024, "本科", "福建师范大学专业设置统计", "school_enrolling_major_count", 75, "majors", ["fjnu_2024_teaching_quality_report_pdf"], "福建师范大学2023-2024学年本科教学质量报告支撑数据列出本科专业85个，其中招生专业75个。", "", "福建师范大学2023-2024学年本科教学质量报告支撑数据", "high")
    add_metric(2024, 2024, "本科", "福建师范大学专业调整统计", "school_current_year_stop_enrollment_major_count", 4, "majors", ["fjnu_2024_teaching_quality_report_pdf"], "福建师范大学2023-2024学年本科教学质量报告称本学年停招酒店管理、动画、环境科学、复合材料与工程4个本科专业。", "", "福建师范大学2023-2024学年本科教学质量报告推进专业动态调整", "high")
    add_metric(2024, 2024, "本科", "福建师范大学专业调整统计", "school_new_major_count", 1, "majors", ["fjnu_2024_teaching_quality_report_pdf"], "福建师范大学2023-2024学年本科教学质量报告称新增足球运动本科专业。", "", "福建师范大学2023-2024学年本科教学质量报告推进专业动态调整", "high")
    add_metric(2024, 2024, "本科", "江苏师范大学专业设置统计", "school_undergrad_major_count", 78, "majors", ["jsnu_2024_teaching_quality_report_pdf"], "江苏师范大学2023-2024学年本科教学质量报告称学校现有本科专业78个。", "", "江苏师范大学2023-2024学年本科教学质量报告本科专业设置情况", "high")
    add_metric(2024, 2024, "本科", "江苏师范大学专业设置统计", "school_enrolling_major_count", 60, "majors", ["jsnu_2024_teaching_quality_report_pdf"], "江苏师范大学2023-2024学年本科教学质量报告称2023-2024学年招生专业60个。", "", "江苏师范大学2023-2024学年本科教学质量报告本科专业设置情况", "high")
    add_metric(2024, 2024, "本科", "江苏师范大学专业设置统计", "school_national_first_class_major_count", 24, "majors", ["jsnu_2024_teaching_quality_report_pdf"], "江苏师范大学2023-2024学年本科教学质量报告称有国家级一流本科专业建设点24个。", "", "江苏师范大学2023-2024学年本科教学质量报告本科专业设置情况", "high")
    add_metric(2024, 2024, "本科", "江苏师范大学专业设置统计", "school_provincial_first_class_major_count", 16, "majors", ["jsnu_2024_teaching_quality_report_pdf"], "江苏师范大学2023-2024学年本科教学质量报告称有省级一流本科专业建设点16个。", "", "江苏师范大学2023-2024学年本科教学质量报告本科专业设置情况", "high")
    add_metric(2026, 2026, "本科", "浙江财经大学专业调整统计", "school_recent_cancel_major_count", 2, "majors", ["zufe_2026_training_system_news"], "浙江财经大学官网新闻称，学校持续优化学科专业布局，撤销2个本科专业。", "", "浙江财经大学专业结构优化近年统计", "high")
    add_metric(2026, 2026, "本科", "浙江财经大学专业调整统计", "school_recent_stop_enrollment_major_count", 8, "majors", ["zufe_2026_training_system_news"], "浙江财经大学官网新闻称，学校持续优化学科专业布局，停招8个本科专业。", "", "浙江财经大学专业结构优化近年统计", "high")
    add_metric(2026, 2026, "本科", "浙江财经大学专业调整统计", "school_new_major_count", 5, "majors", ["zufe_2026_training_system_news"], "浙江财经大学官网新闻称，学校持续优化学科专业布局，新增5个本科专业。", "", "浙江财经大学专业结构优化近年统计", "high")
    add_metric(2026, 2026, "本科", "浙江财经大学专业调整统计", "school_applied_or_preapplied_major_count", 3, "majors", ["zufe_2026_training_system_news"], "浙江财经大学官网新闻称，学校申报和预申报3个本科专业。", "", "浙江财经大学专业结构优化近年统计", "high")
    add_metric(2024, 2024, "本科", "浙江财经大学专业设置统计", "school_enrolling_major_count", 45, "majors", ["zufe_2024_teaching_quality_report_pdf"], "浙江财经大学2023-2024学年本科教学质量报告称，2024年学校共有在招本科专业45个。", "", "浙江财经大学2023-2024学年本科教学质量报告专业设置", "high")
    add_metric(2024, 2024, "本科", "浙江财经大学专业设置统计", "school_economic_management_major_count", 26, "majors", ["zufe_2024_teaching_quality_report_pdf"], "浙江财经大学2023-2024学年本科教学质量报告称，经管类专业26个。", "", "浙江财经大学2023-2024学年本科教学质量报告专业设置", "high")
    add_metric(2024, 2024, "本科", "浙江财经大学专业设置统计", "school_economic_management_major_share", 57.8, "percent", ["zufe_2024_teaching_quality_report_pdf"], "浙江财经大学2023-2024学年本科教学质量报告称，经管类专业占57.8%。", "", "浙江财经大学2023-2024学年本科教学质量报告专业设置", "high")
    add_metric(2024, 2024, "本科", "浙江财经大学招生统计", "school_undergrad_admitted_count", 5198, "persons", ["zufe_2024_teaching_quality_report_pdf"], "浙江财经大学2023-2024学年本科教学质量报告称，2024年共计招收5198人。", "", "浙江财经大学2023-2024学年本科教学质量报告生源报到情况", "high")
    add_metric(2024, 2024, "本科", "浙江财经大学招生统计", "school_undergrad_no_show_count", 168, "persons", ["zufe_2024_teaching_quality_report_pdf"], "浙江财经大学2023-2024学年本科教学质量报告表2-2合计列出未报到人数168人，注释说明未报到学生中未包含保留入学资格10人。", "", "浙江财经大学2023-2024学年本科教学质量报告生源报到情况", "high")
    add_metric(2024, 2024, "本科", "浙江财经大学专业预警统计", "school_warning_major_count", 1, "majors", ["zufe_2024_teaching_quality_report_pdf"], "浙江财经大学2023-2024学年本科教学质量报告称，根据《浙江财经大学本科专业动态调整管理办法》，资产评估专业列入预警名单。", "", "浙江财经大学2023-2024学年本科教学质量报告专业动态调整", "high")
    add_metric(2024, 2024, "本科", "湖北经济学院专业调整统计", "school_paused_major_count", 15, "majors", ["hbue_2024_teaching_quality_report_pdf"], "湖北经济学院2023-2024学年本科教学质量报告称，2024年招生专业数为48个，15个专业暂停招生。", "", "湖北经济学院2023-2024学年本科教学质量报告专业设置及调整情况", "high")
    add_metric(2024, 2024, "本科", "湖北经济学院专业调整统计", "school_current_year_stop_enrollment_major_count", 3, "majors", ["hbue_2024_teaching_quality_report_pdf"], "湖北经济学院2023-2024学年本科教学质量报告称，本年度停招行政管理、劳动与社会保障、会展经济与管理3个专业。", "", "湖北经济学院2023-2024学年本科教学质量报告专业设置及调整情况", "high")
    add_metric(2024, 2024, "本科", "湖北经济学院专业调整统计", "school_applied_cancel_major_count", 7, "majors", ["hbue_2024_teaching_quality_report_pdf"], "湖北经济学院2023-2024学年本科教学质量报告称，申请撤销信用管理、物流工程、资产评估、社会工作、艺术设计学、广告学、应用统计学7个专业。", "", "湖北经济学院2023-2024学年本科教学质量报告专业设置及调整情况", "high")
    add_metric(2024, 2024, "本科", "武汉科技大学专业调整统计", "school_undergrad_major_count", 81, "majors", ["wust_2024_teaching_quality_report_pdf"], "武汉科技大学2023-2024学年本科教学质量报告称学校现有81个本科专业。", "", "武汉科技大学2023-2024学年本科教学质量报告专业设置情况", "high")
    add_metric(2024, 2024, "本科", "武汉科技大学专业调整统计", "school_enrolling_major_count", 72, "majors", ["wust_2024_teaching_quality_report_pdf"], "武汉科技大学2023-2024学年本科教学质量报告称2024年招生专业72个。", "", "武汉科技大学2023-2024学年本科教学质量报告专业设置情况", "high")
    add_metric(2024, 2024, "本科", "武汉科技大学专业调整统计", "school_current_year_stop_enrollment_major_count", 4, "majors", ["wust_2024_teaching_quality_report_pdf"], "武汉科技大学2023-2024学年本科教学质量报告称2024年停招绘画、国际经济与贸易、人力资源管理、财务管理4个专业。", "", "武汉科技大学2023-2024学年本科教学质量报告专业设置情况", "high")
    add_metric(2024, 2024, "本科", "武汉科技大学专业调整统计", "school_applied_cancel_major_count", 2, "majors", ["wust_2024_teaching_quality_report_pdf"], "武汉科技大学2023-2024学年本科教学质量报告称申请撤销人文地理与城乡规划和交通运输2个专业。", "", "武汉科技大学2023-2024学年本科教学质量报告专业设置情况", "high")
    add_metric(2024, 2024, "本科", "成都信息工程大学专业调整统计", "school_undergrad_major_count", 60, "majors", ["cuit_2024_teaching_quality_report_pdf"], "成都信息工程大学2023-2024学年本科教学质量报告称学校设置本科专业60个。", "", "成都信息工程大学2023-2024学年本科教学质量报告本科专业设置情况", "high")
    add_metric(2024, 2024, "本科", "成都信息工程大学专业调整统计", "school_enrolling_major_count", 51, "majors", ["cuit_2024_teaching_quality_report_pdf"], "成都信息工程大学2023-2024学年本科教学质量报告称目前本科招生专业51个。", "", "成都信息工程大学2023-2024学年本科教学质量报告本科专业设置情况", "high")
    add_metric(2024, 2024, "本科", "成都信息工程大学专业调整统计", "school_first_class_major_count", 37, "majors", ["cuit_2024_teaching_quality_report_pdf"], "成都信息工程大学2023-2024学年本科教学质量报告称37个专业获批国家级、省级一流本科专业建设点。", "", "成都信息工程大学2023-2024学年本科教学质量报告专业建设", "high")
    add_metric(2024, 2024, "本科", "成都信息工程大学专业调整统计", "school_engineering_education_certified_major_count", 9, "majors", ["cuit_2024_teaching_quality_report_pdf"], "成都信息工程大学2023-2024学年本科教学质量报告称9个专业已经通过工程教育专业认证。", "", "成都信息工程大学2023-2024学年本科教学质量报告专业建设", "high")
    add_metric(2024, 2024, "本科", "成都信息工程大学专业调整统计", "school_cancel_major_count", 2, "majors", ["cuit_2024_teaching_quality_report_pdf"], "成都信息工程大学2023-2024学年本科教学质量报告称学校撤销材料物理和信息对抗技术两个专业并上报备案。", "", "成都信息工程大学2023-2024学年本科教学质量报告专业布局优化", "high")
    add_metric(2024, 2024, "本科", "成都信息工程大学专业调整统计", "school_minor_program_count", 5, "programs", ["cuit_2024_teaching_quality_report_pdf"], "成都信息工程大学2023-2024学年本科教学质量报告称首批设立5个辅修专业试点。", "", "成都信息工程大学2023-2024学年本科教学质量报告专业建设", "high")
    add_metric(2024, 2024, "本科", "成都信息工程大学专业调整统计", "school_micro_major_count", 18, "programs", ["cuit_2024_teaching_quality_report_pdf"], "成都信息工程大学2023-2024学年本科教学质量报告称首批设立18个微专业试点。", "", "成都信息工程大学2023-2024学年本科教学质量报告专业建设", "high")
    add_metric(2022, 2022, "本科", "陇东学院专业调整统计", "school_not_enrolled_major_count", 15, "majors", ["ldxy_2024_teaching_quality_report_pdf"], "陇东学院2023-2024学年本科教学质量报告称2022年15个专业未招生。", "", "陇东学院2023-2024学年本科教学质量报告专业设置情况", "high")
    add_metric(2023, 2023, "本科", "陇东学院专业调整统计", "school_not_enrolled_major_count", 24, "majors", ["ldxy_2024_teaching_quality_report_pdf"], "陇东学院2023-2024学年本科教学质量报告称2023年24个专业未招生。", "", "陇东学院2023-2024学年本科教学质量报告专业设置情况", "high")
    add_metric(2024, 2024, "本科", "陇东学院专业调整统计", "school_paused_major_count", 22, "majors", ["ldxy_2024_teaching_quality_report_pdf"], "陇东学院2023-2024学年本科教学质量报告称2024年实际招生专业36个、暂停招生专业22个。", "", "陇东学院2023-2024学年本科教学质量报告专业设置情况", "high")
    add_metric(2024, 2024, "本科", "陇东学院专业调整统计", "school_applied_cancel_major_count", 8, "majors", ["ldxy_2024_teaching_quality_report_pdf"], "陇东学院2023-2024学年本科教学质量报告称学校主动申请撤销8个停招五年以上且没有在校生的本科专业。", "", "陇东学院2023-2024学年本科教学质量报告专业设置情况", "high")
    add_metric(2024, 2024, "本科", "大连理工大学专业调整统计", "school_undergrad_major_count", 95, "majors", ["dlut_2024_teaching_quality_report_pdf"], "大连理工大学2023-2024学年本科教学质量报告称截至2024年8月学校共有95个本科专业。", "", "大连理工大学2023-2024学年本科教学质量报告专业设置脚注", "high")
    add_metric(2024, 2024, "本科", "大连理工大学专业调整统计", "school_second_bachelor_major_count", 8, "majors", ["dlut_2024_teaching_quality_report_pdf"], "大连理工大学2023-2024学年本科教学质量报告称学校另有8个第二学士学位专业。", "", "大连理工大学2023-2024学年本科教学质量报告专业设置脚注", "high")
    add_metric(2024, 2024, "本科", "大连理工大学专业调整统计", "school_current_year_stop_enrollment_major_count", 4, "majors", ["dlut_2024_teaching_quality_report_pdf"], "大连理工大学2023-2024学年本科教学质量报告脚注称资源循环科学与工程、食品科学与工程、运动康复、商务英语专业于2024年停止招生。", "", "大连理工大学2023-2024学年本科教学质量报告专业设置脚注", "high")
    add_metric(2024, 2024, "本科", "大连理工大学专业调整统计", "school_applied_cancel_major_count", 6, "majors", ["dlut_2024_teaching_quality_report_pdf"], "大连理工大学2023-2024学年本科教学质量报告脚注称学校已于2024年申请撤销物流工程、纳米材料与技术、材料物理、无机非金属材料工程、法学、管理科学等6个专业。", "", "大连理工大学2023-2024学年本科教学质量报告专业设置脚注", "high")
    add_metric(2024, 2024, "本科", "沈阳工程学院专业调整统计", "school_undergrad_major_count", 37, "majors", ["sie_2024_teaching_quality_report_pdf"], "沈阳工程学院2023-2024学年本科教学质量报告称学校现有普通本科专业37个。", "", "沈阳工程学院2023-2024学年本科教学质量报告专业设置情况", "high")
    add_metric(2023, 2023, "本科", "沈阳工程学院专业调整统计", "school_enrolling_major_count", 34, "majors", ["sie_2024_teaching_quality_report_pdf"], "沈阳工程学院2023-2024学年本科教学质量报告称2023年本科招生专业34个。", "", "沈阳工程学院2023-2024学年本科教学质量报告专业设置情况", "high")
    add_metric(2024, 2024, "本科", "沈阳工程学院专业调整统计", "school_paused_major_count", 3, "majors", ["sie_2024_teaching_quality_report_pdf"], "沈阳工程学院2023-2024学年本科教学质量报告专业设置表标注机械电子工程、机械工艺技术和商务英语停招。", "", "沈阳工程学院2023-2024学年本科教学质量报告专业设置情况", "high")
    add_metric(2023, 2023, "本科", "沈阳工程学院专业调整统计", "school_cancel_major_count", 1, "majors", ["sie_2024_teaching_quality_report_pdf"], "沈阳工程学院2023-2024学年本科教学质量报告称2023年撤销测控技术与仪器专业。", "", "沈阳工程学院2023-2024学年本科教学质量报告专业建设情况", "high")
    add_metric(2024, 2024, "本科", "沈阳工程学院专业调整统计", "school_applied_cancel_major_count", 2, "majors", ["sie_2024_teaching_quality_report_pdf"], "沈阳工程学院2023-2024学年本科教学质量报告称2024年申请撤销机械电子工程、商务英语专业。", "", "沈阳工程学院2023-2024学年本科教学质量报告专业建设情况", "high")
    add_metric(2024, 2024, "本科", "湖州学院专业调整统计", "school_undergrad_major_count", 36, "majors", ["zjhzu_2024_teaching_quality_report_pdf"], "湖州学院2023-2024学年本科教学质量报告称学校现有36个本科专业。", "", "湖州学院2023-2024学年本科教学质量报告专业设置情况", "high")
    add_metric(2024, 2024, "本科", "湖州学院专业调整统计", "school_enrolling_major_count", 27, "majors", ["zjhzu_2024_teaching_quality_report_pdf"], "湖州学院2023-2024学年本科教学质量报告称招生专业27个。", "", "湖州学院2023-2024学年本科教学质量报告专业设置情况", "high")
    add_metric(2024, 2024, "本科", "湖州学院专业调整统计", "school_paused_major_count", 9, "majors", ["zjhzu_2024_teaching_quality_report_pdf"], "湖州学院2023-2024学年本科教学质量报告称停招专业9个，并在专业设置表逐项标注。", "", "湖州学院2023-2024学年本科教学质量报告专业设置情况", "high")
    add_metric(2023, 2023, "本科", "湖州学院专业调整统计", "school_cancel_major_count", 2, "majors", ["zjhzu_2024_teaching_quality_report_pdf"], "湖州学院2023-2024学年本科教学质量报告称2023年撤销历史学、美术学2个专业。", "", "湖州学院2023-2024学年本科教学质量报告专业设置情况", "high")
    add_metric(2024, 2024, "本科", "湖州学院专业调整统计", "school_applied_cancel_major_count", 1, "majors", ["zjhzu_2024_teaching_quality_report_pdf"], "湖州学院2023-2024学年本科教学质量报告称2024年申请撤销物联网工程1个专业。", "", "湖州学院2023-2024学年本科教学质量报告专业设置情况", "high")
    add_metric(2024, 2024, "本科", "沈阳师范大学专业调整统计", "school_undergrad_major_count", 68, "majors", ["synu_2024_teaching_quality_report_pdf"], "沈阳师范大学2023-2024学年本科教学质量报告称学校现有68个本科专业。", "", "沈阳师范大学2023-2024学年本科教学质量报告专业调整情况", "high")
    add_metric(2024, 2024, "本科", "沈阳师范大学专业调整统计", "school_enrolling_major_count", 56, "majors", ["synu_2024_teaching_quality_report_pdf"], "沈阳师范大学2023-2024学年本科教学质量报告称当年本科招生专业56个。", "", "沈阳师范大学2023-2024学年本科教学质量报告专业调整情况", "high")
    add_metric(2024, 2024, "本科", "沈阳师范大学专业调整统计", "school_paused_major_count", 12, "majors", ["synu_2024_teaching_quality_report_pdf"], "沈阳师范大学2023-2024学年本科教学质量报告称停招专业12个。", "", "沈阳师范大学2023-2024学年本科教学质量报告专业调整情况", "high")
    add_metric(2024, 2024, "本科", "沈阳师范大学专业调整统计", "school_cancel_major_count", 9, "majors", ["synu_2024_teaching_quality_report_pdf"], "沈阳师范大学2023-2024学年本科教学质量报告称今年撤销9个专业。", "", "沈阳师范大学2023-2024学年本科教学质量报告专业调整情况", "high")
    add_metric(2024, 2024, "本科", "南京师范大学泰州学院专业调整统计", "school_applied_cancel_major_count", 3, "majors", ["nnutc_2024_teaching_quality_report_pdf"], "南京师范大学泰州学院2023-2024学年本科教学质量报告称申请撤销广告学、戏剧影视文学、园艺等3个长期未招生专业。", "", "南京师范大学泰州学院2023-2024学年本科教学质量报告专业调整情况", "high")
    add_metric(2024, 2024, "本科", "四川音乐学院专业调整统计", "school_applied_cancel_major_count", 5, "majors", ["sccm_2024_teaching_quality_report_pdf"], "四川音乐学院2023-2024学年本科教学质量报告称2023年、2024年申请撤销工业设计、公共事业管理，一并撤销戏剧学、戏剧影视导演、服装与服饰设计3个未招生专业。", "", "四川音乐学院2023-2024学年本科教学质量报告专业调整情况", "high")
    add_metric(2024, 2024, "本科", "四川音乐学院专业调整统计", "school_undergrad_major_count", 29, "majors", ["sccm_2024_teaching_quality_report_pdf"], "四川音乐学院2023-2024学年本科教学质量报告支撑数据称全校本科专业总数29个。", "", "四川音乐学院2023-2024学年本科教学质量报告支撑数据", "high")
    add_metric(2024, 2024, "本科", "四川音乐学院专业调整统计", "school_enrolling_major_count", 27, "majors", ["sccm_2024_teaching_quality_report_pdf"], "四川音乐学院2023-2024学年本科教学质量报告支撑数据称当年本科招生专业总数27个。", "", "四川音乐学院2023-2024学年本科教学质量报告支撑数据", "high")
    add_metric(2024, 2024, "本科", "四川音乐学院专业调整统计", "school_current_year_stop_enrollment_major_count", 1, "majors", ["sccm_2024_teaching_quality_report_pdf"], "四川音乐学院2023-2024学年本科教学质量报告支撑数据列出当年停招专业1个：新媒体艺术。", "", "四川音乐学院2023-2024学年本科教学质量报告支撑数据", "high")
    add_metric(2024, 2024, "本科", "湖南农业大学专业调整统计", "school_undergrad_major_count", 79, "majors", ["hunau_2024_teaching_quality_report_pdf"], "湖南农业大学2023-2024学年本科教学质量报告称学校现有本科专业79个。", "", "湖南农业大学2023-2024学年本科教学质量报告专业设置情况", "high")
    add_metric(2024, 2024, "本科", "湖南农业大学专业调整统计", "school_enrolling_major_count", 74, "majors", ["hunau_2024_teaching_quality_report_pdf"], "湖南农业大学2023-2024学年本科教学质量报告称当年在招本科专业74个。", "", "湖南农业大学2023-2024学年本科教学质量报告专业设置情况", "high")
    add_metric(2024, 2024, "本科", "湖南农业大学专业调整统计", "school_paused_major_count", 5, "majors", ["hunau_2024_teaching_quality_report_pdf"], "湖南农业大学2023-2024学年本科教学质量报告称近5年停招社会工作、信息工程、汽车服务工程、机械电子工程、水族科学与技术等5个本科专业。", "", "湖南农业大学2023-2024学年本科教学质量报告专业设置情况", "high")
    add_metric(2024, 2024, "本科", "湖南农业大学专业调整统计", "school_recent_cancel_major_count", 2, "majors", ["hunau_2024_teaching_quality_report_pdf"], "湖南农业大学2023-2024学年本科教学质量报告称近5年撤销植物科学与技术、表演等2个本科专业。", "", "湖南农业大学2023-2024学年本科教学质量报告专业设置情况", "high")
    add_metric(2024, 2024, "本科", "湖南农业大学专业调整统计", "school_applied_cancel_major_count", 3, "majors", ["hunau_2024_teaching_quality_report_pdf"], "湖南农业大学2023-2024学年本科教学质量报告称2024年申请撤销信息工程、汽车服务工程、社会工作3个本科专业。", "", "湖南农业大学2023-2024学年本科教学质量报告专业建设情况", "high")
    add_metric(2024, 2024, "本科", "浙江科技大学专业调整统计", "school_paused_major_count", 5, "majors", ["zust_2024_teaching_quality_report_pdf"], "浙江科技大学2023-2024学年本科教学质量报告脚注称2023-2024学年停招专业为测控技术与仪器、物联网工程、包装工程、电子商务、汽车服务工程。", "", "浙江科技大学2023-2024学年本科教学质量报告专业设置脚注", "high")
    add_metric(2024, 2024, "本科", "浙江科技大学专业调整统计", "school_cancel_major_count", 2, "majors", ["zust_2024_teaching_quality_report_pdf"], "浙江科技大学2023-2024学年本科教学质量报告称拟增设1个新专业、1个第二学士学位专业，撤销2个专业。", "", "浙江科技大学2023-2024学年本科教学质量报告专业建设情况", "high")
    add_metric(2023, 2023, "本科", "淮北师范大学专业调整统计", "school_cancel_major_count", 7, "majors", ["chnu_2024_teaching_quality_report_pdf"], "淮北师范大学2023-2024学年本科教学质量报告称2024年完成金融数学、机械电子工程、计算机科学与技术（理学）、数字媒体技术、财务管理、公共事业管理、劳动与社会保障等7个专业的撤销工作。", "", "淮北师范大学2023-2024学年本科教学质量报告专业建设情况", "high")
    add_metric(2024, 2024, "本科", "淮北师范大学专业调整统计", "school_applied_cancel_major_count", 1, "majors", ["chnu_2024_teaching_quality_report_pdf"], "淮北师范大学2023-2024学年本科教学质量报告称申请撤销信息管理与信息系统专业。", "", "淮北师范大学2023-2024学年本科教学质量报告专业建设情况", "high")
    add_metric(2024, 2024, "本科", "淮北师范大学专业调整统计", "school_paused_major_count", 5, "majors", ["chnu_2024_teaching_quality_report_pdf"], "淮北师范大学2023-2024学年本科教学质量报告称停招广告学、戏剧影视文学、社会学、国际经济与贸易、审计学5个专业。", "", "淮北师范大学2023-2024学年本科教学质量报告专业建设情况", "high")
    add_metric(2024, 2024, "本科", "淮北师范大学专业调整统计", "school_undergrad_major_count", 71, "majors", ["chnu_2024_teaching_quality_report_pdf"], "淮北师范大学2023-2024学年本科教学质量报告支撑数据列出本科专业总数71个。", "", "淮北师范大学2023-2024学年本科教学质量报告支撑数据", "high")
    add_metric(2024, 2024, "本科", "淮北师范大学专业调整统计", "school_enrolling_major_count", 60, "majors", ["chnu_2024_teaching_quality_report_pdf"], "淮北师范大学2023-2024学年本科教学质量报告支撑数据列出在招专业数60个。", "", "淮北师范大学2023-2024学年本科教学质量报告支撑数据", "high")
    add_metric(2024, 2024, "本科", "温州理工学院专业调整统计", "school_undergrad_major_count", 34, "majors", ["wzut_2024_teaching_quality_report_pdf"], "温州理工学院2023-2024学年本科教学质量报告称学校现有本科专业34个。", "", "温州理工学院2023-2024学年本科教学质量报告专业设置情况", "high")
    add_metric(2024, 2024, "本科", "温州理工学院专业调整统计", "school_enrolling_major_count", 28, "majors", ["wzut_2024_teaching_quality_report_pdf"], "温州理工学院2023-2024学年本科教学质量报告称在招本科专业28个。", "", "温州理工学院2023-2024学年本科教学质量报告专业设置情况", "high")
    add_metric(2024, 2024, "本科", "温州理工学院专业调整统计", "school_paused_major_count", 6, "majors", ["wzut_2024_teaching_quality_report_pdf"], "温州理工学院2023-2024学年本科教学质量报告称停招专业6个。", "", "温州理工学院2023-2024学年本科教学质量报告专业设置情况", "high")
    add_metric(2024, 2024, "本科", "池州学院专业调整统计", "school_undergrad_major_count", 50, "majors", ["czu_2024_teaching_quality_report_page"], "池州学院2023-2024学年本科教学质量报告支撑数据列出本科专业总数50个。", "", "池州学院2023-2024学年本科教学质量报告支撑数据", "high")
    add_metric(2024, 2024, "本科", "池州学院专业调整统计", "school_enrolling_major_count", 49, "majors", ["czu_2024_teaching_quality_report_page"], "池州学院2023-2024学年本科教学质量报告支撑数据列出在招专业数49个。", "", "池州学院2023-2024学年本科教学质量报告支撑数据", "high")
    add_metric(2024, 2024, "本科", "池州学院专业调整统计", "school_current_year_stop_enrollment_major_count", 1, "majors", ["czu_2024_teaching_quality_report_page"], "池州学院2023-2024学年本科教学质量报告支撑数据列出当年停招专业名单：材料化学。", "", "池州学院2023-2024学年本科教学质量报告支撑数据", "high")
    add_metric(2024, 2024, "本科", "池州学院专业调整统计", "school_internal_paused_major_count", 3, "majors", ["czu_2024_teaching_quality_report_page"], "池州学院2023-2024学年本科教学质量报告正文称当年学校招生的校内专业69个，停招的校内专业3个，分别是材料化学、市场营销（专升本）、知识产权（安警院联培）。", "", "池州学院2023-2024学年本科教学质量报告专业建设情况", "high")
    add_metric(2024, 2024, "本科", "北京物资学院专业调整统计", "school_undergrad_major_count", 26, "majors", ["bwu_2024_teaching_quality_report_pdf"], "北京物资学院2023-2024学年本科教学质量报告支撑数据列出本科专业总数26个。", "", "北京物资学院2023-2024学年本科教学质量报告支撑数据", "high")
    add_metric(2024, 2024, "本科", "北京物资学院专业调整统计", "school_enrolling_major_count", 24, "majors", ["bwu_2024_teaching_quality_report_pdf"], "北京物资学院2023-2024学年本科教学质量报告支撑数据列出在招专业数24个。", "", "北京物资学院2023-2024学年本科教学质量报告支撑数据", "high")
    add_metric(2023, 2023, "本科", "北京物资学院专业调整统计", "school_cancel_major_count", 3, "majors", ["bwu_2024_teaching_quality_report_pdf"], "北京物资学院2023-2024学年本科教学质量报告称2023年撤销经济统计学、英语、劳动关系三个专业。", "", "北京物资学院2023-2024学年本科教学质量报告专业建设情况", "high")
    add_metric(2024, 2024, "本科", "许昌学院专业调整统计", "school_enrolling_major_count", 57, "majors", ["xcu_2024_teaching_quality_report_pdf"], "许昌学院2023-2024学年本科教学质量报告称学校本科招生专业数量稳定在57个。", "", "许昌学院2023-2024学年本科教学质量报告专业建设情况", "high")
    add_metric(2024, 2024, "本科", "许昌学院专业调整统计", "school_current_year_stop_enrollment_major_count", 9, "majors", ["xcu_2024_teaching_quality_report_pdf"], "许昌学院2023-2024学年本科教学质量报告称停招通信工程、网络工程、交通设备与控制工程、酒店管理、财务管理、绘画、汉语国际教育、商务英语、工商管理等9个本科专业。", "", "许昌学院2023-2024学年本科教学质量报告专业建设情况", "high")
    add_metric(2024, 2024, "本科", "许昌学院专业调整统计", "school_applied_cancel_major_count", 4, "majors", ["xcu_2024_teaching_quality_report_pdf"], "许昌学院2023-2024学年本科教学质量报告称城乡规划、市场营销、音乐表演、社会体育指导与管理等4个专业申请撤销。", "", "许昌学院2023-2024学年本科教学质量报告专业建设情况", "high")
    add_metric(2024, 2024, "本科", "上海海洋大学专业调整统计", "school_new_major_count", 2, "majors", ["shou_2024_teaching_quality_report_pdf"], "上海海洋大学2023-2024学年本科教学质量报告称2023学年新增人工智能、环境科学与工程2个专业。", "", "上海海洋大学2023-2024学年本科教学质量报告专业建设", "high")
    add_metric(2024, 2024, "本科", "上海海洋大学专业调整统计", "school_current_year_stop_enrollment_major_count", 3, "majors", ["shou_2024_teaching_quality_report_pdf"], "上海海洋大学2023-2024学年本科教学质量报告称停招环境科学、环境工程、文化产业管理3个专业。", "", "上海海洋大学2023-2024学年本科教学质量报告专业建设", "high")
    add_metric(2024, 2024, "本科", "上海大学专业调整统计", "school_undergrad_major_count", 101, "majors", ["shu_2024_teaching_quality_report_pdf"], "上海大学2023-2024学年本科教育教学质量报告附录列出本科专业总数101个。", "", "上海大学2023-2024学年本科教育教学质量报告专业设置及调整情况", "high")
    add_metric(2024, 2024, "本科", "上海大学专业调整统计", "school_enrolling_major_count", 98, "majors", ["shu_2024_teaching_quality_report_pdf"], "上海大学2023-2024学年本科教育教学质量报告附录列出在招专业数98个。", "", "上海大学2023-2024学年本科教育教学质量报告专业设置及调整情况", "high")
    add_metric(2024, 2024, "本科", "上海大学专业调整统计", "school_current_year_stop_enrollment_major_count", 3, "majors", ["shu_2024_teaching_quality_report_pdf"], "上海大学2023-2024学年本科教育教学质量报告列出停招材料物理、工业设计、包装工程3个本科专业。", "", "上海大学2023-2024学年本科教育教学质量报告专业设置及调整情况", "high")
    add_metric(2024, 2024, "本科", "上海大学专业调整统计", "school_applied_cancel_major_count", 1, "majors", ["shu_2024_teaching_quality_report_pdf"], "上海大学2023-2024学年本科教育教学质量报告称向教育主管部门申请备案撤销包装工程专业。", "", "上海大学2023-2024学年本科教育教学质量报告基本情况", "high")
    add_metric(2024, 2024, "本科", "内蒙古财经大学专业调整统计", "school_undergrad_major_count", 55, "majors", ["imufe_2024_teaching_quality_report_pdf"], "内蒙古财经大学2023-2024学年本科教学质量报告称学校现设本科专业55个。", "", "内蒙古财经大学2023-2024学年本科教学质量报告本科专业设置", "high")
    add_metric(2024, 2024, "本科", "内蒙古财经大学专业调整统计", "school_enrolling_major_count", 48, "majors", ["imufe_2024_teaching_quality_report_pdf"], "内蒙古财经大学2023-2024学年本科教学质量报告称实际招生本科专业48个。", "", "内蒙古财经大学2023-2024学年本科教学质量报告本科专业设置", "high")
    add_metric(2024, 2024, "本科", "内蒙古财经大学专业调整统计", "school_paused_major_count", 7, "majors", ["imufe_2024_teaching_quality_report_pdf"], "内蒙古财经大学2023-2024学年本科教学质量报告称暂停招生专业7个。", "", "内蒙古财经大学2023-2024学年本科教学质量报告本科专业设置", "high")
    add_metric(2023, 2023, "本科", "内蒙古财经大学专业调整统计", "school_cancel_major_count", 5, "majors", ["imufe_2024_teaching_quality_report_pdf"], "内蒙古财经大学2023-2024学年本科教学质量报告称2023年撤销信息与计算科学等5个本科专业。", "", "内蒙古财经大学2023-2024学年本科教学质量报告专业动态调整机制", "high")
    add_metric(2023, 2023, "本科", "内蒙古财经大学专业调整统计", "school_paused_major_count", 6, "majors", ["imufe_2024_teaching_quality_report_pdf"], "内蒙古财经大学2023-2024学年本科教学质量报告称2023年对物业管理等6个本科专业暂缓招生。", "", "内蒙古财经大学2023-2024学年本科教学质量报告专业动态调整机制", "high")
    add_metric(2024, 2024, "本科", "内蒙古财经大学专业调整统计", "school_current_year_paused_major_count", 2, "majors", ["imufe_2024_teaching_quality_report_pdf"], "内蒙古财经大学2023-2024学年本科教学质量报告称2024年对保险学和公共事业管理共2个本科专业暂缓招生。", "", "内蒙古财经大学2023-2024学年本科教学质量报告专业动态调整机制", "high")
    add_metric(2024, 2024, "本科", "重庆文理学院专业调整统计", "school_undergrad_major_count", 63, "majors", ["cqwu_2024_teaching_quality_report_pdf"], "重庆文理学院2023-2024学年本科教学质量报告称学校共有本科专业63个。", "", "重庆文理学院2023-2024学年本科教学质量报告学科专业设置", "high")
    add_metric(2024, 2024, "本科", "重庆文理学院专业调整统计", "school_enrolling_major_count", 59, "majors", ["cqwu_2024_teaching_quality_report_pdf"], "重庆文理学院2023-2024学年本科教学质量报告称2024年招生专业59个。", "", "重庆文理学院2023-2024学年本科教学质量报告生源质量", "high")
    add_metric(2024, 2024, "本科", "重庆文理学院专业调整统计", "school_current_year_stop_enrollment_major_count", 4, "majors", ["cqwu_2024_teaching_quality_report_pdf"], "重庆文理学院2023-2024学年本科教学质量报告称2024年停招广播电视学、美术学、旅游管理与服务教育、金融数学等4个专业。", "", "重庆文理学院2023-2024学年本科教学质量报告学科专业设置", "high")
    add_metric(2024, 2024, "本科", "重庆文理学院专业调整统计", "school_cancel_major_count", 1, "majors", ["cqwu_2024_teaching_quality_report_pdf"], "重庆文理学院2023-2024学年本科教学质量报告称2024年主动撤销经济统计学专业。", "", "重庆文理学院2023-2024学年本科教学质量报告学科专业设置", "high")
    add_metric(2024, 2024, "本科", "重庆文理学院专业调整统计", "school_new_major_count", 2, "majors", ["cqwu_2024_teaching_quality_report_pdf"], "重庆文理学院2023-2024学年本科教学质量报告称2024年成功增设智慧农业和智能车辆工程2个新专业。", "", "重庆文理学院2023-2024学年本科教学质量报告学科专业设置", "high")
    add_metric(2024, 2024, "本科", "东北石油大学专业调整统计", "school_undergrad_major_count", 68, "majors", ["nepu_2024_teaching_quality_report_pdf"], "东北石油大学2023-2024学年本科教学质量报告支撑数据列出全校本科专业总数68个。", "", "东北石油大学2023-2024学年本科教学质量报告专业设置情况", "high")
    add_metric(2024, 2024, "本科", "东北石油大学专业调整统计", "school_enrolling_major_count", 53, "majors", ["nepu_2024_teaching_quality_report_pdf"], "东北石油大学2023-2024学年本科教学质量报告支撑数据列出当年本科招生专业总数53个。", "", "东北石油大学2023-2024学年本科教学质量报告专业设置情况", "high")
    add_metric(2024, 2024, "本科", "东北石油大学专业调整统计", "school_new_major_count", 2, "majors", ["nepu_2024_teaching_quality_report_pdf"], "东北石油大学2023-2024学年本科教学质量报告称2024年增设地球信息科学与技术、储能科学与工程2个本科专业。", "", "东北石油大学2023-2024学年本科教学质量报告专业设置情况", "high")
    add_metric(2024, 2024, "本科", "东北石油大学专业调整统计", "school_current_year_stop_enrollment_major_count", 3, "majors", ["nepu_2024_teaching_quality_report_pdf"], "东北石油大学2023-2024学年本科教学质量报告称2024年停招光电信息科学与工程、智能电网信息工程、财务管理3个专业。", "", "东北石油大学2023-2024学年本科教学质量报告专业设置情况", "high")
    add_metric(2024, 2024, "本科", "东北石油大学专业调整统计", "school_cancel_major_count", 1, "majors", ["nepu_2024_teaching_quality_report_pdf"], "东北石油大学2023-2024学年本科教学质量报告称撤销公共事业管理专业。", "", "东北石油大学2023-2024学年本科教学质量报告专业设置情况", "high")
    add_metric(2024, 2024, "本科", "重庆科技大学专业调整统计", "school_undergrad_major_count", 70, "majors", ["cqust_2024_teaching_quality_report_pdf"], "重庆科技大学2023-2024学年本科教学质量报告支撑数据列出全校本科专业总数70个。", "", "重庆科技大学2023-2024学年本科教学质量报告教学基本状态数据", "high")
    add_metric(2024, 2024, "本科", "重庆科技大学专业调整统计", "school_enrolling_major_count", 61, "majors", ["cqust_2024_teaching_quality_report_pdf"], "重庆科技大学2023-2024学年本科教学质量报告称其他61个本科专业以及物流管理1个中外合作办学项目实际招生；支撑数据列出当年本科招生专业61个。", "", "重庆科技大学2023-2024学年本科教学质量报告本科专业设置", "high")
    add_metric(2024, 2024, "本科", "重庆科技大学专业调整统计", "school_paused_major_count", 9, "majors", ["cqust_2024_teaching_quality_report_pdf"], "重庆科技大学2023-2024学年本科教学质量报告称截至2024年9月矿物加工工程、过程装备与控制工程等9个专业暂停招生。", "", "重庆科技大学2023-2024学年本科教学质量报告本科专业设置", "high")
    add_metric(2024, 2024, "本科", "重庆科技大学专业调整统计", "school_current_year_stop_enrollment_major_count", 7, "majors", ["cqust_2024_teaching_quality_report_pdf"], "重庆科技大学2023-2024学年本科教学质量报告称停招地球物理学、材料物理等9个专业，其中当年停招7个。", "", "重庆科技大学2023-2024学年本科教学质量报告专业建设", "high")
    add_metric(2024, 2024, "本科", "宜春学院专业调整统计", "school_enrolling_major_count", 53, "majors", ["jxycu_2024_teaching_quality_report_pdf"], "宜春学院2023-2024学年本科教学质量报告称学校2024年招生专业总数减少至53个。", "", "宜春学院2023-2024学年本科教学质量报告专业建设", "high")
    add_metric(2024, 2024, "本科", "宜春学院专业调整统计", "school_current_year_stop_enrollment_major_count", 14, "majors", ["jxycu_2024_teaching_quality_report_pdf"], "宜春学院2023-2024学年本科教学质量报告称停招广播电视编导等14个专业；本指标记录总数，专业名单只结构化报告明示名称。", "", "宜春学院2023-2024学年本科教学质量报告专业建设", "high")
    add_metric(2024, 2024, "本科", "宜春学院专业调整统计", "school_integrated_major_count", 2, "majors", ["jxycu_2024_teaching_quality_report_pdf"], "宜春学院2023-2024学年本科教学质量报告称交叉融合生物工程和制药工程2个专业。", "", "宜春学院2023-2024学年本科教学质量报告专业建设", "high")
    add_metric(2024, 2024, "本科", "山东航空学院专业调整统计", "school_undergrad_major_count", 59, "majors", ["sdua_2024_teaching_quality_report_pdf"], "山东航空学院2023-2024学年本科教学质量报告核心支撑数据列出全校本科专业总数59个。", "", "山东航空学院2023-2024学年本科教学质量报告核心支撑数据", "high")
    add_metric(2024, 2024, "本科", "山东航空学院专业调整统计", "school_enrolling_major_count", 52, "majors", ["sdua_2024_teaching_quality_report_pdf"], "山东航空学院2023-2024学年本科教学质量报告核心支撑数据列出当年本科招生专业总数52个。", "", "山东航空学院2023-2024学年本科教学质量报告核心支撑数据", "high")
    add_metric(2024, 2024, "本科", "山东航空学院专业调整统计", "school_new_major_count", 1, "majors", ["sdua_2024_teaching_quality_report_pdf"], "山东航空学院2023-2024学年本科教学质量报告核心支撑数据列出当年新增专业为空间信息与数字技术。", "", "山东航空学院2023-2024学年本科教学质量报告核心支撑数据", "high")
    add_metric(2024, 2024, "本科", "山东航空学院专业调整统计", "school_current_year_stop_enrollment_major_count", 2, "majors", ["sdua_2024_teaching_quality_report_pdf"], "山东航空学院2023-2024学年本科教学质量报告核心支撑数据列出当年停招专业为法语、飞行器适航技术。", "", "山东航空学院2023-2024学年本科教学质量报告核心支撑数据", "high")
    add_metric(2024, 2024, "本科", "山东航空学院专业调整统计", "school_cancel_major_count", 1, "majors", ["sdua_2024_teaching_quality_report_pdf"], "山东航空学院2023-2024学年本科教学质量报告称撤销能源化学工程专业。", "", "山东航空学院2023-2024学年本科教学质量报告专业建设", "high")
    add_metric(2024, 2024, "本科", "北京科技大学专业调整统计", "school_undergrad_major_count", 61, "majors", ["ustb_2024_teaching_quality_report_pdf"], "北京科技大学2023-2024学年本科教学质量报告称学校设有61个本科专业。", "", "北京科技大学2023-2024学年本科教学质量报告本科专业设置", "high")
    add_metric(2024, 2024, "本科", "北京科技大学专业调整统计", "school_enrolling_major_group_count", 25, "major-groups", ["ustb_2024_teaching_quality_report_pdf"], "北京科技大学2023-2024学年本科教学质量报告称2023-2024学年共有25个招生专业（类），其中14个为大类招生专业。", "", "北京科技大学2023-2024学年本科教学质量报告本科专业设置", "high")
    add_metric(2024, 2024, "本科", "北京科技大学专业调整统计", "school_paused_major_count", 6, "majors", ["ustb_2024_teaching_quality_report_pdf"], "北京科技大学2023-2024学年本科教学质量报告称本科专业中共有6个专业停招。", "", "北京科技大学2023-2024学年本科教学质量报告本科专业设置", "high")
    add_metric(2024, 2024, "本科", "北京科技大学专业调整统计", "school_new_major_count", 2, "majors", ["ustb_2024_teaching_quality_report_pdf"], "北京科技大学2023-2024学年本科教学质量报告称2024年新增材料智能技术、工程力学2个专业。", "", "北京科技大学2023-2024学年本科教学质量报告本科专业设置", "high")
    add_metric(2024, 2024, "本科", "贵州财经大学专业调整统计", "school_undergrad_major_count", 64, "majors", ["gufe_2024_teaching_quality_report_pdf"], "贵州财经大学2023-2024学年本科教学质量报告称学校现有本科专业64个。", "", "贵州财经大学2023-2024学年本科教学质量报告专业建设", "high")
    add_metric(2024, 2024, "本科", "贵州财经大学专业调整统计", "school_enrolling_major_count", 54, "majors", ["gufe_2024_teaching_quality_report_pdf"], "贵州财经大学2023-2024学年本科教学质量报告支撑数据列出当年本科招生专业总数54个。", "", "贵州财经大学2023-2024学年本科教学质量报告支撑数据", "high")
    add_metric(2024, 2024, "本科", "贵州财经大学专业调整统计", "school_current_year_stop_enrollment_major_count", 10, "majors", ["gufe_2024_teaching_quality_report_pdf"], "贵州财经大学2023-2024学年本科教学质量报告称停招投资学、保险学、应用心理学、网络工程、教育技术学、英语、葡萄牙语、汉语国际教育、公共事业管理、艺术管理等10个专业。", "", "贵州财经大学2023-2024学年本科教学质量报告专业建设", "high")
    add_metric(2024, 2024, "本科", "贵州财经大学专业调整统计", "school_applied_cancel_major_count", 3, "majors", ["gufe_2024_teaching_quality_report_pdf"], "贵州财经大学2023-2024学年本科教学质量报告称申请撤销英语、会展经济与管理、金融科技3个专业。", "", "贵州财经大学2023-2024学年本科教学质量报告专业建设", "high")
    add_metric(2023, 2023, "本科", "上海工程技术大学专业调整统计", "school_warning_major_count", 5, "majors", ["sues_2024_teaching_quality_report_pdf"], "上海工程技术大学2023-2024学年本科教学质量报告称2023年发布本科专业综合得分，对校内5个专业进行预警。", "", "上海工程技术大学2023-2024学年本科教学质量报告教学质量保障", "high")
    add_metric(2024, 2024, "本科", "上海工程技术大学专业调整统计", "school_stop_or_cancel_major_count", 4, "majors", ["sues_2024_teaching_quality_report_pdf"], "上海工程技术大学2023-2024学年本科教学质量报告称持续推进专业动态调整，撤销及停招专业4个；报告未逐名列出这4个专业。", "", "上海工程技术大学2023-2024学年本科教学质量报告质量保障", "high")
    add_metric(2024, 2024, "本科", "湖南大学专业调整统计", "school_undergrad_major_count", 84, "majors", ["hnu_2024_major_catalog_pdf"], "湖南大学2024年本科专业目录备注称2024年全校本科专业84个。", "", "湖南大学2024年本科专业目录", "high")
    add_metric(2024, 2024, "本科", "湖南大学专业调整统计", "school_enrolling_major_count", 76, "majors", ["hnu_2024_major_catalog_pdf"], "湖南大学2024年本科专业目录备注称2024年招生专业76个。", "", "湖南大学2024年本科专业目录", "high")
    add_metric(2024, 2024, "本科", "湖南大学专业调整统计", "school_new_major_count", 4, "majors", ["hnu_2024_major_catalog_pdf"], "湖南大学2024年本科专业目录备注称2024年新增专业4个。", "", "湖南大学2024年本科专业目录", "high")
    add_metric(2024, 2024, "本科", "湖南大学专业调整统计", "school_current_year_stop_enrollment_major_count", 8, "majors", ["hnu_2024_major_catalog_pdf"], "湖南大学2024年本科专业目录备注称停招专业8个，表格备注列逐项标注当年停招。", "", "湖南大学2024年本科专业目录", "high")
    add_metric(2025, 2025, "本科", "湖南大学专业调整统计", "school_undergrad_major_count", 90, "majors", ["hnu_2025_major_catalog_pdf"], "湖南大学2025年本科专业目录备注称2025年全校本科专业90个。", "", "湖南大学2025年本科专业目录", "high")
    add_metric(2025, 2025, "本科", "湖南大学专业调整统计", "school_enrolling_major_count", 77, "majors", ["hnu_2025_major_catalog_pdf"], "湖南大学2025年本科专业目录备注称2025年招生专业77个。", "", "湖南大学2025年本科专业目录", "high")
    add_metric(2025, 2025, "本科", "湖南大学专业调整统计", "school_new_major_count", 3, "majors", ["hnu_2025_major_catalog_pdf"], "湖南大学2025年本科专业目录备注称2025年新增专业3个。", "", "湖南大学2025年本科专业目录", "high")
    add_metric(2025, 2025, "本科", "湖南大学专业调整统计", "school_current_year_stop_enrollment_major_count", 13, "majors", ["hnu_2025_major_catalog_pdf"], "湖南大学2025年本科专业目录备注称停招专业13个，表格备注列逐项标注当年停招。", "", "湖南大学2025年本科专业目录", "high")
    add_metric(2025, 2025, "本科", "北京交通大学专业调整统计", "school_applied_cancel_major_count", 6, "majors", ["bjtu_2025_major_adjustment_notice", "bjtu_2025_major_adjustment_materials_zip"], "北京交通大学2025年度专业设置与调整申报材料zip附件含申请撤销专业材料目录，列出材料化学、生物信息学、给排水科学与工程、电子信息工程、汉语言、思想政治教育6个专业。", "", "北京交通大学2025年度专业设置与调整申报材料", "high")
    add_metric(2024, 2024, "本科", "安徽中医药大学专业调整统计", "school_catalog_listed_major_count", 32, "majors", ["ahtcm_2024_major_setting"], "安徽中医药大学2024年专业设置情况表序号列出1-32个本科专业。", "", "安徽中医药大学2024年专业设置情况", "high")
    add_metric(2024, 2024, "本科", "安徽中医药大学专业调整统计", "school_current_year_stop_enrollment_major_count", 4, "majors", ["ahtcm_2024_stop_enrollment", "ahtcm_2024_major_setting"], "安徽中医药大学2024年停招专业情况说明列出药物分析、中药资源与开发、中医儿科学、保险学4个停招专业；2024年专业设置情况表在招生状态列逐项标注停招。", "", "安徽中医药大学2024年停招专业情况说明及专业设置情况", "high")
    add_metric(2025, 2025, "本科", "安徽中医药大学专业调整统计", "school_catalog_listed_major_count", 32, "majors", ["ahtcm_2025_major_setting"], "安徽中医药大学2025年本科专业设置情况表序号列出1-32个本科专业。", "", "安徽中医药大学2025年本科专业设置情况", "high")
    add_metric(2025, 2025, "本科", "安徽中医药大学专业调整统计", "school_current_year_stop_enrollment_major_count", 3, "majors", ["ahtcm_2025_stop_enrollment", "ahtcm_2025_major_setting"], "安徽中医药大学2025年停招专业情况说明列出中医儿科学、保险学、人力资源管理3个停招专业；2025年本科专业设置情况表在招生状态列逐项标注停招。", "", "安徽中医药大学2025年停招专业情况说明及本科专业设置情况", "high")
    add_metric(2025, 2025, "本科", "安阳工学院专业调整统计", "school_current_year_stop_enrollment_major_count", 9, "majors", ["ayit_2025_stop_enrollment"], "安阳工学院2025年停招本科专业一览表列出9个停招本科专业及专业代码。", "", "安阳工学院2025年停招本科专业一览表", "high")
    add_metric(2024, 2024, "本科", "西北工业大学专业调整统计", "school_catalog_listed_major_count", 72, "majors", ["nwpu_2024_major_setting_pdf"], "西北工业大学2024年本科专业设置PDF表格序号列出1-72个本科专业。", "", "西北工业大学本科专业设置（2024年10月更新）", "high")
    add_metric(2024, 2024, "本科", "西北工业大学专业调整统计", "school_current_year_stop_enrollment_major_count", 8, "majors", ["nwpu_2024_major_setting_pdf"], "西北工业大学2024年本科专业设置PDF表格备注列标注8个停招专业。", "", "西北工业大学本科专业设置（2024年10月更新）", "high")
    add_metric(2024, 2024, "本科", "西北工业大学专业调整统计", "school_new_major_count", 2, "majors", ["nwpu_2024_major_setting_pdf"], "西北工业大学2024年本科专业设置PDF表格备注列标注2个新增专业。", "", "西北工业大学本科专业设置（2024年10月更新）", "high")
    add_metric(2025, 2025, "本科", "西北工业大学专业调整统计", "school_catalog_listed_major_count", 73, "majors", ["nwpu_2025_major_setting_pdf"], "西北工业大学2025年本科专业设置PDF表格序号列出1-73个本科专业。", "", "西北工业大学本科专业设置（2025年12月更新）", "high")
    add_metric(2025, 2025, "本科", "西北工业大学专业调整统计", "school_current_year_stop_enrollment_major_count", 9, "majors", ["nwpu_2025_major_setting_pdf"], "西北工业大学2025年本科专业设置PDF表格备注列标注9个停招专业。", "", "西北工业大学本科专业设置（2025年12月更新）", "high")
    add_metric(2025, 2025, "本科", "西北工业大学专业调整统计", "school_new_major_count", 2, "majors", ["nwpu_2025_major_setting_pdf"], "西北工业大学2025年本科专业设置PDF表格备注列标注2个新增专业。", "", "西北工业大学本科专业设置（2025年12月更新）", "high")
    add_metric(2024, 2024, "本科", "吉林工程技术师范学院专业调整统计", "school_undergrad_major_count", 54, "majors", ["jlenu_2024_teaching_quality_report_pdf"], "吉林工程技术师范学院2023-2024学年本科教学质量报告称学校设有本科专业54个。", "", "吉林工程技术师范学院2023-2024学年本科教学质量报告本科专业基本情况", "high")
    add_metric(2024, 2024, "本科", "吉林工程技术师范学院专业调整统计", "school_with_students_major_count", 53, "majors", ["jlenu_2024_teaching_quality_report_pdf"], "吉林工程技术师范学院2023-2024学年本科教学质量报告称本学年有在校生专业数53个。", "", "吉林工程技术师范学院2023-2024学年本科教学质量报告本科专业基本情况", "high")
    add_metric(2024, 2024, "本科", "吉林工程技术师范学院专业调整统计", "school_enrolling_major_count", 41, "majors", ["jlenu_2024_teaching_quality_report_pdf"], "吉林工程技术师范学院2023-2024学年本科教学质量报告称在招专业41个。", "", "吉林工程技术师范学院2023-2024学年本科教学质量报告专业建设情况", "high")
    add_metric(2024, 2024, "本科", "吉林工程技术师范学院专业调整统计", "school_paused_major_count", 10, "majors", ["jlenu_2024_teaching_quality_report_pdf"], "吉林工程技术师范学院2023-2024学年本科教学质量报告表2-1注释说明44-53项为停招专业。", "", "吉林工程技术师范学院2023-2024学年本科专业设置一览表", "high")
    add_metric(2024, 2024, "本科", "吉林工程技术师范学院专业调整统计", "school_recent_three_year_stop_enrollment_major_count", 14, "majors", ["jlenu_2024_teaching_quality_report_pdf"], "吉林工程技术师范学院2023-2024学年本科教学质量报告称近三年累计停招与吉林省经济社会发展不相适应的专业14个。", "", "吉林工程技术师范学院2023-2024学年本科教学质量报告专业建设情况", "high")
    add_metric(2023, 2023, "本科", "吉林工程技术师范学院专业调整统计", "school_cancel_major_count", 1, "majors", ["jlenu_2023_major_cancellation"], "吉林工程技术师范学院2023年拟撤销专业公示称经学校专业建设委员会审议通过，决定拟撤销行政管理专业。", "", "吉林工程技术师范学院2023年拟撤销专业公示", "high")
    add_metric(2025, 2025, "本科", "吉林工程技术师范学院专业调整统计", "school_enrolling_major_count", 42, "majors", ["jlenu_2025_teaching_quality_report_pdf"], "吉林工程技术师范学院2024-2025学年本科教学质量报告称现有42个本科专业招生。", "", "吉林工程技术师范学院2024-2025学年本科专业设置情况", "high")
    add_metric(2025, 2025, "本科", "吉林工程技术师范学院专业调整统计", "school_cancel_major_count", 4, "majors", ["jlenu_2025_teaching_quality_report_pdf"], "吉林工程技术师范学院2024-2025学年本科教学质量报告称2025年撤销汽车服务工程、人工智能、环境设计、表演4个专业。", "", "吉林工程技术师范学院2024-2025学年本科教学质量报告专业建设成效", "high")
    add_metric(2025, 2025, "本科", "吉林工程技术师范学院专业调整统计", "school_new_major_count", 3, "majors", ["jlenu_2025_teaching_quality_report_pdf"], "吉林工程技术师范学院2024-2025学年本科教学质量报告称2025年新增新能源汽车工程、新能源汽车工程技术、网络空间安全3个新工科专业。", "", "吉林工程技术师范学院2024-2025学年本科教学质量报告专业建设成效", "high")
    add_metric(2025, 2025, "本科", "吉林工程技术师范学院专业调整统计", "school_recent_three_year_stop_enrollment_major_count", 10, "majors", ["jlenu_2025_teaching_quality_report_pdf"], "吉林工程技术师范学院2024-2025学年本科教学质量报告称近三年累计停招与吉林省经济社会发展不相适应的专业10个。", "", "吉林工程技术师范学院2024-2025学年本科教学质量报告专业建设成效", "high")
    add_metric(2025, 2025, "本科", "江苏海洋大学专业调整统计", "school_new_major_count", 2, "majors", ["jou_2025_major_adjustment"], "江苏海洋大学2025年度本科专业设置公示的申报新专业汇总表列出土木、水利与海洋工程、储能科学与工程2个申报新专业。", "", "江苏海洋大学2025年度本科专业设置公示", "high")
    add_metric(2025, 2025, "本科", "江苏海洋大学专业调整统计", "school_pre_filing_major_count", 2, "majors", ["jou_2025_major_adjustment"], "江苏海洋大学2025年度本科专业设置公示的申报预备案专业汇总表列出智能海洋装备、人工智能2个预备案专业。", "", "江苏海洋大学2025年度本科专业设置公示", "high")
    add_metric(2025, 2025, "本科", "江苏海洋大学专业调整统计", "school_cancel_major_count", 4, "majors", ["jou_2025_major_adjustment"], "江苏海洋大学2025年度拟撤销专业汇总表列出测控技术与仪器、动植物检疫、广告学、水质科学与技术4个拟撤销专业。", "", "江苏海洋大学2025年度本科专业设置公示", "high")
    add_metric(2025, 2025, "本科", "浙江大学专业调整统计", "school_catalog_listed_major_count", 130, "majors", ["zju_2025_undergrad_major_table"], "浙江大学本科专业情况表序号列出1-130个本科专业。", "", "浙江大学本科专业情况表", "high")
    add_metric(2025, 2025, "本科", "浙江大学专业调整统计", "school_paused_major_count", 28, "majors", ["zju_2025_undergrad_major_table"], "浙江大学本科专业情况表备注列标注28个已停招专业。", "", "浙江大学本科专业情况表", "high")
    add_metric(2024, 2024, "本科", "浙江农林大学暨阳学院专业调整统计", "school_new_major_count", 5, "majors", ["zjyc_2024_cancel_stop"], "浙江农林大学暨阳学院2024年申请增设专业和撤销停招专业公示列出拟申请增设智能科学与技术、智能建造、数字经济、网络与新媒体、应急管理5个本科专业。", "", "浙江农林大学暨阳学院2024年申请增设专业和撤销停招专业公示", "high")
    add_metric(2024, 2024, "本科", "浙江农林大学暨阳学院专业调整统计", "school_cancel_major_count", 2, "majors", ["zjyc_2024_cancel_stop"], "浙江农林大学暨阳学院2024年申请增设专业和撤销停招专业公示列出拟撤销生物技术、服装与服饰设计2个本科专业。", "", "浙江农林大学暨阳学院2024年申请增设专业和撤销停招专业公示", "high")
    add_metric(2025, 2025, "本科", "合肥理工学院专业调整统计", "school_catalog_listed_major_count", 41, "majors", ["hfit_2025_major_setting"], "合肥理工学院2025年专业基本情况数据表序号列出1-41个本科专业。", "", "合肥理工学院2025年专业设置、当年新增专业、停招专业名单", "high")
    add_metric(2025, 2025, "本科", "合肥理工学院专业调整统计", "school_enrolling_major_count", 10, "majors", ["hfit_2025_major_setting"], "合肥理工学院2025年专业基本情况数据表备注列标注10个专业当年招生，其中2个同时为当年新增。", "", "合肥理工学院2025年专业设置、当年新增专业、停招专业名单", "high")
    add_metric(2025, 2025, "本科", "合肥理工学院专业调整统计", "school_new_major_count", 2, "majors", ["hfit_2025_major_setting"], "合肥理工学院2025年专业基本情况数据表备注列标注储能科学与工程、光电信息科学与工程2个专业当年新增。", "", "合肥理工学院2025年专业设置、当年新增专业、停招专业名单", "high")
    add_metric(2025, 2025, "本科", "合肥理工学院专业调整统计", "school_current_year_stop_enrollment_major_count", 31, "majors", ["hfit_2025_major_setting"], "合肥理工学院2025年专业基本情况数据表备注列标注31个专业当年停招。", "", "合肥理工学院2025年专业设置、当年新增专业、停招专业名单", "high")
    add_metric(2025, 2025, "本科", "新疆师范大学专业调整统计", "school_new_major_count", 4, "majors", ["xjnu_2025_major_adjustment"], "新疆师范大学2025年度本科专业设置情况公示列出食品安全与检测、中国画、国际法、足球运动4个拟新增专业。", "", "新疆师范大学2025年度本科专业设置情况公示", "high")
    add_metric(2025, 2025, "本科", "新疆师范大学专业调整统计", "school_pre_filing_major_count", 5, "majors", ["xjnu_2025_major_adjustment"], "新疆师范大学2025年度本科专业设置情况公示列出合成生物学、人工智能教育、地理国情监测、旅游地学与规划工程、城市管理5个预备案专业。", "", "新疆师范大学2025年度本科专业设置情况公示", "high")
    add_metric(2025, 2025, "本科", "新疆师范大学专业调整统计", "school_cancel_major_count", 1, "majors", ["xjnu_2025_major_adjustment"], "新疆师范大学2025年度本科专业设置情况公示说明拟撤销舞蹈表演专业1个。", "", "新疆师范大学2025年度本科专业设置情况公示", "high")
    add_metric(2023, 2023, "本科", "长春工程学院专业调整统计", "school_applied_cancel_major_count", 6, "majors", ["ccit_2023_major_cancellation"], "长春工程学院2023年度拟撤销本科专业公示列出信息与计算科学等6个已连续停招5年以上的本科专业。", "", "长春工程学院2023年度拟撤销本科专业公示", "high")
    add_metric(2025, 2025, "本科", "华东师范大学专业调整统计", "school_new_major_count", 1, "majors", ["ecnu_2025_stop_enrollment"], "华东师范大学2025年本科专业设置及新增专业、停招专业情况称学校新增数据科学1个专业。", "", "华东师范大学2025年本科专业设置及新增专业、停招专业情况", "high")
    add_metric(2025, 2025, "本科", "华东师范大学专业调整统计", "school_undergrad_major_count", 86, "majors", ["ecnu_2025_stop_enrollment"], "华东师范大学2025年本科专业设置及新增专业、停招专业情况称专业总数为86个。", "", "华东师范大学2025年本科专业设置及新增专业、停招专业情况", "high")
    add_metric(2025, 2025, "本科", "华东师范大学专业调整统计", "school_current_year_stop_enrollment_major_count", 24, "majors", ["ecnu_2025_stop_enrollment"], "华东师范大学2025年本科专业设置及新增专业、停招专业情况称停招专业24个；正文实际明示22个专业名称。", "", "华东师范大学2025年本科专业设置及新增专业、停招专业情况", "high")
    add_metric(2025, 2025, "本科", "信阳师范大学专业调整统计", "school_new_major_count", 3, "majors", ["xynu_2025_major_cancellation"], "信阳师范大学2025年度拟申报、撤销专业公示称拟向教育部申报密码科学与技术、康复治疗学、旅游管理与服务教育3个本科专业。", "", "信阳师范大学2025年度拟申报、撤销专业公示", "high")
    add_metric(2025, 2025, "本科", "信阳师范大学专业调整统计", "school_applied_cancel_major_count", 3, "majors", ["xynu_2025_major_cancellation"], "信阳师范大学2025年度拟申报、撤销专业公示称拟申请撤销人文教育、应用物理学、信息管理与信息系统3个本科专业。", "", "信阳师范大学2025年度拟申报、撤销专业公示", "high")
    add_metric(2025, 2025, "本科", "南阳理工学院专业调整统计", "school_applied_cancel_major_count", 5, "majors", ["nyist_2025_major_cancellation"], "南阳理工学院2025年度本科专业设置与调整公示称对停招五年及以上且无在籍学生的法学、音乐表演、网络工程、电子科学与技术、汽车服务工程5个专业拟申请撤销。", "", "南阳理工学院2025年度本科专业设置与调整公示", "high")
    add_metric(2025, 2025, "本科", "南阳理工学院专业调整统计", "school_new_major_count", 8, "majors", ["nyist_2025_major_cancellation"], "南阳理工学院2025年度本科专业设置与调整公示列出工业软件、智慧水利、智慧城市与空间规划、化学工程与工业生物工程、智能车辆工程、储能科学与工程、内部审计、智能建造8个拟申请新增专业。", "", "南阳理工学院2025年度本科专业设置与调整公示", "high")
    add_metric(2025, 2025, "本科", "南阳理工学院专业调整统计", "school_pre_filing_major_count", 2, "majors", ["nyist_2025_major_cancellation"], "南阳理工学院2025年度本科专业设置与调整公示列出智慧能源工程、音乐科技2个拟申请预申报专业。", "", "南阳理工学院2025年度本科专业设置与调整公示", "high")
    add_metric(2026, 2025, "本科", "四川省本科专业调整统计", "province_adjusted_major_count", 296, "major-instances", ["scol_2026_sichuan_undergrad_adjustment_stats"], "四川在线报道称，据悉四川全省2025年共调整优化专业296个。", "", "四川省2025年本科专业调整优化汇总统计", "medium")
    add_metric(2026, 2025, "本科", "四川省本科专业调整统计", "province_new_major_count", 120, "major-instances", ["scol_2026_sichuan_undergrad_adjustment_stats"], "四川在线报道称，四川高校2025年共新增本科专业120个，其中普通本科专业106个、第二学士学位专业14个。", "", "四川省2025年本科专业调整优化汇总统计", "medium")
    add_metric(2026, 2025, "本科", "四川省本科专业调整统计", "province_standard_new_major_count", 106, "major-instances", ["scol_2026_sichuan_undergrad_adjustment_stats"], "四川在线报道称，四川高校2025年新增普通本科专业106个。", "", "四川省2025年本科专业调整优化汇总统计", "medium")
    add_metric(2026, 2025, "本科", "四川省本科专业调整统计", "province_second_bachelor_major_count", 14, "major-instances", ["scol_2026_sichuan_undergrad_adjustment_stats"], "四川在线报道称，四川高校2025年新增第二学士学位专业14个。", "", "四川省2025年本科专业调整优化汇总统计", "medium")
    add_metric(2026, 2025, "本科", "四川省本科专业调整统计", "province_duration_degree_adjusted_major_count", 11, "major-instances", ["scol_2026_sichuan_undergrad_adjustment_stats"], "四川在线报道称，四川高校2025年调整修业年限及学位授予门类专业点11个。", "", "四川省2025年本科专业调整优化汇总统计", "medium")
    add_metric(2026, 2025, "本科", "四川省本科专业调整统计", "province_stop_enrollment_major_count", 132, "major-instances", ["scol_2026_sichuan_undergrad_adjustment_stats"], "四川在线报道称，四川全省2025年停招132个本科专业点。", "", "四川省2025年本科专业调整优化汇总统计", "medium")
    add_metric(2026, 2025, "本科", "四川省本科专业调整统计", "province_cancel_major_count", 33, "major-instances", ["scol_2026_sichuan_undergrad_adjustment_stats"], "四川在线报道称，四川高校2025年主动撤销专业33个，且12所高校共撤销33个专业。", "", "四川省2025年本科专业调整优化汇总统计", "medium")
    add_metric(2026, 2025, "本科", "四川省本科专业调整统计", "province_three_year_adjusted_major_count", 669, "major-instances", ["scol_2026_sichuan_undergrad_adjustment_stats"], "四川在线报道称，近三年来四川已累计调整优化专业669个。", "", "四川省近三年本科专业调整优化汇总统计", "medium")
    add_metric(2026, 2025, "本科", "四川省本科专业调整统计", "province_new_major_industry_match_count", 103, "major-instances", ["scol_2026_sichuan_undergrad_adjustment_stats"], "四川在线报道称，四川2025年新设专业中共有103个专业深度匹配六大优势产业与重点产业链。", "", "四川省2025年新设本科专业产业匹配统计", "medium")
    add_metric(2026, 2025, "本科", "四川省本科专业调整统计", "province_new_major_industry_match_share", 85.8, "percent", ["scol_2026_sichuan_undergrad_adjustment_stats"], "四川在线报道称，四川2025年新设专业中深度匹配六大优势产业与重点产业链的专业占比85.8%。", "", "四川省2025年新设本科专业产业匹配统计", "medium")
    add_metric(2026, 2025, "本科", "机器人工程", "province_new_major_school_count", 8, "schools", ["scol_2026_sichuan_undergrad_adjustment_stats"], "四川在线报道称，四川2025年机器人工程有8所高校同时增设。", "", "四川省2025年热门新设本科专业学校数", "medium")
    add_metric(2026, 2025, "本科", "数字经济", "province_new_major_school_count", 5, "schools", ["scol_2026_sichuan_undergrad_adjustment_stats"], "四川在线报道称，四川2025年数字经济有5所高校增设。", "", "四川省2025年热门新设本科专业学校数", "medium")
    add_metric(2026, 2025, "本科", "人工智能", "province_new_major_school_count", 4, "schools", ["scol_2026_sichuan_undergrad_adjustment_stats"], "四川在线报道称，四川2025年人工智能有4所高校增设。", "", "四川省2025年热门新设本科专业学校数", "medium")
    add_metric(2025, 2025, "本科", "武汉理工大学专业调整统计", "school_stop_enrollment_major_count", 24, "majors", ["whut_2025_undergrad_major_catalog"], "武汉理工大学2025年本科专业目录图片表格备注列共标注24个2020、2021、2024或2025年停招专业。", "", "武汉理工大学2025年本科专业目录备注列停招标记", "high")
    add_metric(2025, 2025, "本科", "武汉理工大学专业调整统计", "school_current_year_stop_enrollment_major_count", 14, "majors", ["whut_2025_undergrad_major_catalog"], "武汉理工大学2025年本科专业目录图片表格备注列标注14个2025年停招专业。", "", "武汉理工大学2025年本科专业目录备注列停招标记", "high")
    add_metric(2025, 2025, "本科", "绍兴理工学院专业调整统计", "school_new_major_count", 3, "majors", ["zsit_2025_major_setting_status"], "绍兴理工学院2025年专业设置表中设置年份为2025的专业有智能建造、跨境电子商务、电子封装技术3个。", "", "绍兴理工学院2025年专业设置及历年招生情况表", "high")
    add_metric(2025, 2025, "本科", "绍兴理工学院专业调整统计", "school_stop_enrollment_major_count", 3, "majors", ["zsit_2025_major_setting_status"], "绍兴理工学院2025年专业设置表的2022-2025历年招生情况列中，机械电子工程、跨境电子商务、电子封装技术出现停招标记。", "", "绍兴理工学院2025年专业设置及历年招生情况表", "high")
    add_metric(2025, 2025, "本科", "绍兴理工学院专业调整统计", "school_current_year_stop_enrollment_major_count", 2, "majors", ["zsit_2025_major_setting_status"], "绍兴理工学院2025年专业设置表在2025年招生情况列标记跨境电子商务、电子封装技术为停招。", "", "绍兴理工学院2025年专业设置及历年招生情况表", "high")
    add_metric(2025, 2025, "本科", "安徽科技学院专业调整统计", "school_pre_filing_major_count", 5, "majors", ["ahstu_2025_major_adjustment"], "安徽科技学院2025年度预申报、撤销和调整修业年限本科专业公示称学校预申报专业5个。", "", "安徽科技学院2025年度本科专业调整公示", "high")
    add_metric(2025, 2025, "本科", "安徽科技学院专业调整统计", "school_cancel_major_count", 2, "majors", ["ahstu_2025_major_adjustment"], "安徽科技学院2025年度预申报、撤销和调整修业年限本科专业公示称学校撤销专业2个。", "", "安徽科技学院2025年度本科专业调整公示", "high")
    add_metric(2025, 2025, "本科", "安徽科技学院专业调整统计", "school_duration_adjusted_major_count", 1, "majors", ["ahstu_2025_major_adjustment"], "安徽科技学院2025年度预申报、撤销和调整修业年限本科专业公示称学校调整修业年限专业1个。", "", "安徽科技学院2025年度本科专业调整公示", "high")
    add_metric(2024, 2024, "本科", "保定理工学院专业调整统计", "school_new_major_count", 2, "majors", ["cuggw_2024_major_adjustment"], "保定理工学院2024年度新增、撤销本科专业情况公示称拟申报增设软件工程、应用心理学2个本科专业。", "", "保定理工学院2024年度新增、撤销本科专业情况公示", "high")
    add_metric(2024, 2024, "本科", "保定理工学院专业调整统计", "school_cancel_major_count", 2, "majors", ["cuggw_2024_major_adjustment"], "保定理工学院2024年度新增、撤销本科专业情况公示称拟撤销机械电子工程、建筑学2个本科专业。", "", "保定理工学院2024年度新增、撤销本科专业情况公示", "high")
    add_metric(2026, 2026, "本科", "咸阳师范学院专业调整统计", "school_new_major_count", 5, "majors", ["xync_2026_major_adjustment"], "咸阳师范学院2026年度拟新增、拟撤销、预申报本科专业公示称拟新增低空经济与管理、数字文旅、虚拟空间艺术、新媒体艺术、储能科学与工程等5个本科专业。", "", "咸阳师范学院2026年度拟新增、拟撤销、预申报本科专业公示", "high")
    add_metric(2026, 2026, "本科", "咸阳师范学院专业调整统计", "school_cancel_major_count", 1, "majors", ["xync_2026_major_adjustment"], "咸阳师范学院2026年度拟新增、拟撤销、预申报本科专业公示称拟撤销视觉传达设计专业。", "", "咸阳师范学院2026年度拟新增、拟撤销、预申报本科专业公示", "high")
    add_metric(2026, 2026, "本科", "咸阳师范学院专业调整统计", "school_pre_filing_major_count", 25, "majors", ["xync_2026_major_adjustment"], "咸阳师范学院2026年度拟新增、拟撤销、预申报本科专业公示称拟预申报供应链管理等25个本科专业。", "", "咸阳师范学院2026年度拟新增、拟撤销、预申报本科专业公示", "high")
    add_metric(2025, 2025, "本科", "西藏大学专业调整统计", "school_new_major_count", 2, "majors", ["utibet_2025_major_adjustment"], "西藏大学2025年度拟设置与调整本科专业公示称拟新增新能源科学与工程、人工智能2个本科专业。", "", "西藏大学2025年度拟设置与调整本科专业公示", "high")
    add_metric(2025, 2025, "本科", "西藏大学专业调整统计", "school_cancel_major_count", 4, "majors", ["utibet_2025_major_adjustment"], "西藏大学2025年度拟设置与调整本科专业公示称拟撤销市场营销、财务管理、服装与服饰设计、新闻学4个本科专业。", "", "西藏大学2025年度拟设置与调整本科专业公示", "high")
    add_metric(2025, 2025, "本科", "西藏大学专业调整统计", "school_duration_adjusted_major_count", 1, "majors", ["utibet_2025_major_adjustment"], "西藏大学2025年度拟设置与调整本科专业公示称拟将城乡规划学制从五年调整为四年。", "", "西藏大学2025年度拟设置与调整本科专业公示", "high")
    add_metric(2023, 2023, "本科", "长春电子科技学院专业调整统计", "school_cancel_major_count", 2, "majors", ["changdian_2023_major_cancellation"], "长春电子科技学院2023年度拟撤销专业公示称拟撤销光源与照明、网络工程2个专业。", "", "长春电子科技学院2023年度拟撤销专业公示", "high")
    add_metric(2023, 2023, "本科", "长春电子科技学院专业调整统计", "school_stop_enrollment_major_count", 2, "majors", ["changdian_2023_major_cancellation"], "长春电子科技学院2023年度拟撤销专业公示说明光源与照明2021年停止招生、网络工程2020年停止招生。", "", "长春电子科技学院2023年度拟撤销专业公示", "high")

    add_metric(
        2026,
        2026,
        "本科",
        "本科专业停招统计",
        "stop_enrollment_university_count",
        70,
        "schools",
        ["peopleapp_2026_undergrad_stop_summary", "qq_2026_undergrad_stop_summary"],
        "人民日报客户端人民号/教育在线、掌上高考腾讯新闻页统计70所本科高校最新公布的停招专业名单。",
        "",
        "70所本科高校最新公布停招专业名单不完全统计",
        "medium",
    )
    add_metric(
        2026,
        2026,
        "本科",
        "本科专业停招统计",
        "stop_enrollment_major_instance_count",
        525,
        "major-instances",
        ["peopleapp_2026_undergrad_stop_summary", "qq_2026_undergrad_stop_summary"],
        "人民日报客户端人民号/教育在线、掌上高考腾讯新闻页统计70所本科高校最新公布的停招专业名单，共涉及525个本科专业。",
        "",
        "70所本科高校最新公布停招专业名单不完全统计",
        "medium",
    )
    stop_frequency_metrics = [
        ("市场营销", 16, 1),
        ("公共事业管理", 11, 2),
        ("物流管理", 10, 3),
        ("物联网工程", 9, 4),
        ("旅游管理", 8, 5),
        ("日语", 8, 5),
        ("网络工程", 8, 5),
        ("财务管理", 7, 8),
        ("电子商务", 7, 8),
        ("广告学", 7, 8),
        ("汉语国际教育", 7, 8),
        ("会展经济与管理", 7, 8),
        ("汽车服务工程", 7, 8),
        ("保险学", 6, 14),
        ("材料成型及控制工程", 6, 14),
        ("材料化学", 6, 14),
        ("环境设计", 6, 14),
        ("酒店管理", 6, 14),
        ("劳动与社会保障", 6, 14),
        ("人力资源管理", 6, 14),
        ("应用心理学", 6, 14),
        ("产品设计", 5, 22),
        ("德语", 5, 22),
        ("动画", 5, 22),
        ("翻译", 5, 22),
        ("风景园林", 5, 22),
        ("工业工程", 5, 22),
        ("光电信息科学与工程", 5, 22),
        ("国际经济与贸易", 5, 22),
        ("建筑环境与能源应用工程", 5, 22),
        ("经济统计学", 5, 22),
        ("食品质量与安全", 5, 22),
        ("视觉传达设计", 5, 22),
        ("数字媒体技术", 5, 22),
        ("信息管理与信息系统", 5, 22),
        ("应用化学", 5, 22),
        ("房地产开发与管理", 4, 37),
        ("工程管理", 4, 37),
        ("工业设计", 4, 37),
        ("广播电视学", 4, 37),
    ]
    for major, school_count, rank in stop_frequency_metrics:
        add_metric(
            2026,
            2026,
            "本科",
            major,
            "stop_enrollment_school_count",
            school_count,
            "schools",
            [
                "peopleapp_2026_undergrad_stop_summary",
                "peopleapp_2026_undergrad_stop_frequency_image",
                "qq_2026_undergrad_stop_summary",
                "qq_2026_undergrad_stop_frequency_image_part1",
                "qq_2026_undergrad_stop_frequency_image_part2",
            ],
            f"人民日报客户端人民号/教育在线内嵌长图及掌上高考腾讯新闻页榜单统计，{major}在70所本科高校最新公布停招专业名单中有{school_count}所高校暂停招生。",
            rank,
            "70所本科高校最新公布停招专业名单专业频次（停招数量较多的本科专业部分）",
            "medium",
        )

    return rows


def enrich_metric_records(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    catalog = read_catalog()
    enriched: list[dict[str, Any]] = []
    for row in rows:
        standard_name = LEVEL_STANDARD_NAME_ALIASES.get(
            (row["reported_major_name"], row["education_level"]),
            STANDARD_NAME_ALIASES.get(row["reported_major_name"], row["reported_major_name"]),
        )
        catalog_level = "本科" if row["education_level"] == "本科" else "专科"
        catalog_name = CATALOG_LOOKUP_ALIASES.get(standard_name, standard_name)
        catalog_row = catalog.get((catalog_name, catalog_level), catalog.get((catalog_name, ""), {}))
        metric_id_base = "|".join(
            [
                str(row["report_year"]),
                row["education_level"],
                row["reported_major_name"],
                row["metric_name"],
                str(row["metric_value"]),
                row["source_ids"],
            ]
        )
        enriched.append(
            {
                "metric_id": hashlib.sha1(metric_id_base.encode("utf-8")).hexdigest()[:16],
                "schema_version": SCHEMA_VERSION,
                "standard_major_name": standard_name,
                "major_code": catalog_row.get("major_code", ""),
                "captured_at": CAPTURED_AT,
                **row,
            }
        )
    return sorted(
        enriched,
        key=lambda item: (
            item["report_year"],
            item["education_level"],
            item["reported_major_name"],
            item["metric_name"],
        ),
    )


def curated_official_warning_rows(raw_dir: Path | None = None) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []

    def add_warning(
        policy_year: int,
        region: str,
        level: str,
        record_type: str,
        warning_label: str,
        major: str,
        policy_action: str,
        criterion_text: str,
        source_ids: list[str],
        evidence_text: str,
        confidence: str = "high",
    ) -> None:
        rows.append(
            {
                "policy_year": policy_year,
                "region": region,
                "education_level": level,
                "record_type": record_type,
                "warning_label": warning_label,
                "reported_major_name": major,
                "policy_action": policy_action,
                "criterion_text": criterion_text,
                "source_ids": ";".join(source_ids),
                "evidence_text": evidence_text,
                "confidence": confidence,
            }
        )

    def add_warning_list(
        policy_year: int,
        region: str,
        level: str,
        record_type: str,
        warning_label: str,
        majors: list[str],
        policy_action: str,
        criterion_text: str,
        source_ids: list[str],
        evidence_text: str,
        confidence: str = "medium",
    ) -> None:
        for major in majors:
            add_warning(
                policy_year,
                region,
                level,
                record_type,
                warning_label,
                major,
                policy_action,
                criterion_text,
                source_ids,
                evidence_text,
                confidence,
            )

    def add_school_table_warning(
        policy_year: int,
        record_type: str,
        warning_label: str,
        major: str,
        major_code: str,
        opened_year: str,
        stopped_year: str,
        degree: str,
        school_unit: str,
        source_id: str,
        source_row_no: int,
        action: str,
        note: str = "",
    ) -> None:
        evidence_parts = [
            f"甘肃农业大学表格第{source_row_no}行列出{major}（{major_code}，{degree}），",
            f"开设年份{opened_year}，停招年份{stopped_year}，所属学院为{school_unit}。",
        ]
        if note:
            evidence_parts.append(f"备注：{note}。")
        rows.append(
            {
                "policy_year": policy_year,
                "region": "甘肃农业大学",
                "education_level": "本科",
                "record_type": record_type,
                "warning_label": warning_label,
                "reported_major_name": major,
                "major_code": major_code,
                "policy_action": action,
                "criterion_text": f"甘肃农业大学本科专业动态调整表；开设年份{opened_year}，停招年份{stopped_year}。",
                "source_row_no": str(source_row_no),
                "source_ids": source_id,
                "evidence_text": "".join(evidence_parts),
                "confidence": "high",
            }
        )

    guizhou_majors = ["计算机科学与技术", "艺术设计", "汉语言文学", "数学与应用数学", "表演", "绘画", "美术学"]
    for major in guizhou_majors:
        add_warning(
            2012,
            "贵州省",
            "本科",
            "major_warning_list",
            "省级本科专业预警名单",
            major,
            "调低预警专业招生计划；连续3次被预警的，除个别特殊专业外实行退出机制、停止招生。",
            "社会认同度不高，社会需求量明显下降，毕业生就业率较低且布点较多。",
            ["edu_guizhou_2012_undergrad_warning"],
            "贵州省教育厅将该省本科院校中的计算机科学与技术、艺术设计、汉语言文学、数学与应用数学、表演、绘画、美术学等专业列入预警名单。",
        )

    add_warning(
        2012,
        "贵州省",
        "本科",
        "policy_rule",
        "省级本科专业预警及退出机制",
        "",
        "调低预警专业招生计划；连续3次预警后退出并停止招生。",
        "社会认同度不高，社会需求量明显下降，毕业生就业率较低且布点较多。",
        ["edu_guizhou_2012_undergrad_warning"],
        "社会认同度不高、社会需求量明显下降、毕业生就业率较低且布点较多的专业，将被列入预警名单；连续3次被预警的，除个别特殊专业外，将实行退出机制，停止招生。",
    )
    guizhou_2021_majors = ["法学", "数字媒体艺术", "劳动与社会保障", "公共事业管理", "汉语言文学"]
    for major in guizhou_2021_majors:
        add_warning(
            2021,
            "贵州省",
            "本科",
            "major_warning_list",
            "2021年贵州省普通本科高校专业预警名单",
            major,
            "设置预警专业的高校需论证社会需求和人才培养环节、加强教学质量管理、优化课程设置；连续3次列入预警名单的专业除个别特殊专业外实行退出机制、停止招生。",
            "贵州省高校毕业生就业率较低（就业率排名倒数前十名）且布点较多。",
            ["eol_guizhou_2021_undergrad_warning"],
            "贵州省教育厅发布2021年普通本科高校专业预警通知，对法学、数字媒体艺术、劳动与社会保障、公共事业管理、汉语言文学五个专业进行预警。",
        )

    cctv_source = ["cctv_2016_shanghai_low_employment_warning"]
    cctv_edu_cn_2014_low_source = ["cctv_2016_shanghai_low_employment_warning", "edu_cn_2014_low_employment_warning"]
    cctv_edu_cn_2014_low_page2_source = ["cctv_2016_shanghai_low_employment_warning", "edu_cn_2014_low_employment_warning_page2"]
    shanghai_warning_lists = {
        2012: ["社会工作", "社会体育", "广告学", "艺术设计", "表演", "动画", "播音与主持艺术", "广播电视编导", "信息与计算科学", "材料化学", "电子信息工程", "网络工程", "信息显示与光电技术", "食品质量与安全", "国际商务", "公共事业管理", "劳动与社会保障", "会展经济与管理"],
        2013: ["信息管理与信息系统", "公共事业管理", "汉语言文学", "行政管理", "工商管理", "物流管理", "电子商务", "市场营销", "社会体育", "旅游管理", "社会工作", "法学", "园林", "食品卫生与营养学", "体育教育"],
        2014: ["信息管理与信息系统", "电子商务", "市场营销", "公共事业管理", "汉语言文学", "社会工作", "交通运输"],
        2016: ["英语", "国际经济与贸易", "法学", "工商管理", "物流管理", "新闻学", "旅游管理", "信息管理与信息系统", "市场营销", "行政管理"],
    }
    for year, majors in shanghai_warning_lists.items():
        add_warning_list(
            year,
            "上海市",
            "本科",
            "major_warning_list",
            f"{year}年上海市高校本科预警专业名单",
            majors,
            "减少预警专业招生总量；对办学条件严重不足、培养质量低、特色不明显的专业严格控制招生计划，甚至暂停招生；拟增设已列入预警范围专业原则上不予受理备案申请。",
            "重复设置相对较多，连续多年招生录取率和毕业生签约情况不太理想。",
            cctv_source,
            f"央视网转载央视新闻/新华社报道列出上海市{year}年本科预警专业名单。",
            "medium",
        )

    add_warning_list(
        2014,
        "全国",
        "本科",
        "national_low_employment_list",
        "教育部近两年全国就业率较低本科专业名单",
        ["食品卫生与营养学", "生物科学", "旅游管理", "社会体育指导与管理", "市场营销", "动画", "知识产权", "广播电视编导", "表演", "艺术设计学", "播音与主持艺术", "音乐表演", "电子商务", "贸易经济", "公共事业管理"],
        "列入就业率较低本科专业名单，供专业设置、招生计划和专业结构调整参考。",
        "教育部发布的近两年全国就业率较低本科专业。",
        cctv_edu_cn_2014_low_source,
        "央视网报道和中国教育和科研计算机网转载新华网-大河报报道均列出教育部2014年发布的15个全国就业率较低本科专业名单。",
        "medium",
    )

    provincial_low_employment_text = """
北京：音乐学、社会学、法学、公共事业管理、应用物理学、新闻学、国际经济与贸易、表演、工商管理、经济学
天津：药物制剂、作曲与作曲技术理论、历史学、针灸推拿学、文物与博物馆学、化学生物学、测绘工程、教育学
河北：通信工程、电气工程及其自动化、汉语言文学、土木工程、英语、电子信息工程、工程管理、法学、计算机科学与技术、会计学
山西：音乐学、英语、旅游管理、经济学、学前教育、计算机科学与技术、行政管理、社会体育指导与管理、市场营销、美术学
内蒙古：播音与主持艺术、社会学、农村区域发展、生态学、蒙古语、艺术设计学、法学、英语、计算机科学与技术、会计学
辽宁：数学与应用数学、艺术设计学、表演、广告学、音乐表演、护理学、广播电视编导
吉林：英语、市场营销、计算机科学与技术、艺术设计学、国际经济与贸易、日语、工商管理、汉语言文学、美术学、动画
黑龙江：播音与主持艺术、武术与民族传统体育、摄影、运动训练、表演、体育教育、音乐表演
上海：汉语言文学、工业设计、法学、软件工程、行政管理、物流管理、公共事业管理、电子商务、市场营销、数学与应用数学
江苏：应用心理学、园艺、旅游管理、汉语国际教育、社会工作、纺织工程、工业设计、社会体育指导与管理
浙江：英语、古典文献学、信息管理与信息系统、电子信息工程、汉语言文学、哲学、市场营销、财务管理、法学、文秘教育
安徽：应用心理学、工程管理、法学、信息与计算科学、国际经济与贸易、动画、公共事业管理、计算机科学与技术、金融工程、信息管理与信息系统
福建：社会学、法学、汉语言文学、行政管理、体育教育、政治学与行政学、数学与应用数学、教育学、应用心理学、信息工程
江西：秘书学、翻译、动画、资产评估
山东：音乐表演、应用心理学、音乐学、公共事业管理、艺术设计学、汉语言文学
河南：法学、应用心理学、汉语国际教育、应用物理学、人力资源管理
湖北：中西医临床医学、法学、口腔医学、动画、中医学、金融工程、土地资源管理、音乐表演、社会体育指导与管理、音乐学
湖南：计算机科学与技术、英语、市场营销、国际经济与贸易、法学、旅游管理、信息与计算科学、汉语言文学、会计学、工商管理
广东：表演、应用心理学、新闻学、美术学、公共事业管理、汉语言文学、治安学、考古学、音乐表演、资源环境科学
广西：市场营销、英语、计算机科学与技术、国际经济与贸易、社会体育指导与管理、法学、化学、应用心理学
海南：化学、计算机科学与技术、数学与应用数学、物理学、网络工程、法学
重庆：戏剧影视导演、法医学、戏剧影视美术设计、地理信息系统、农林经济管理
四川：材料物理、地理信息科学、交通工程、教育学、物业管理、文物与博物馆学、西班牙语、辐射防护与核安全、文化产业管理、心理学
贵州：播音与主持艺术、行政管理、土地资源管理、运动训练、工业设计
云南：动画、体育教育、生物科学、教育技术学、物理学、美术学、英语、汉语言文学、思想政治教育、公共事业管理
西藏：档案学、历史学
陕西：音乐表演
甘肃：英语、汉语言文学、经济学、艺术设计学
青海：音乐表演、经济学、旅游管理
宁夏：农业水利工程、数学与应用数学、日语、工商管理、信息管理与信息系统、公共事业管理、广告学、信息与计算科学、美术学、信息工程
新疆：应用物理学、法学、社会体育、小学教育、物理学、应用化学、美术学、生物技术、计算机科学与技术、新闻学
新疆生产建设兵团：广播电视新闻学、汉语言文学、化学、中国少数民族语言文学
""".strip()
    province_region_aliases = {
        "北京": "北京市",
        "天津": "天津市",
        "河北": "河北省",
        "山西": "山西省",
        "内蒙古": "内蒙古自治区",
        "辽宁": "辽宁省",
        "吉林": "吉林省",
        "黑龙江": "黑龙江省",
        "上海": "上海市",
        "江苏": "江苏省",
        "浙江": "浙江省",
        "安徽": "安徽省",
        "福建": "福建省",
        "江西": "江西省",
        "山东": "山东省",
        "河南": "河南省",
        "湖北": "湖北省",
        "湖南": "湖南省",
        "广东": "广东省",
        "广西": "广西壮族自治区",
        "海南": "海南省",
        "重庆": "重庆市",
        "四川": "四川省",
        "贵州": "贵州省",
        "云南": "云南省",
        "西藏": "西藏自治区",
        "陕西": "陕西省",
        "甘肃": "甘肃省",
        "青海": "青海省",
        "宁夏": "宁夏回族自治区",
        "新疆": "新疆维吾尔自治区",
    }
    for line in provincial_low_employment_text.splitlines():
        region, majors_text = line.split("：", 1)
        normalized_region = province_region_aliases.get(region, region)
        majors = [major.strip() for major in majors_text.split("、") if major.strip()]
        has_edu_cn_page1_evidence = region in {"北京", "天津", "河北", "山西", "内蒙古", "辽宁", "吉林", "黑龙江", "上海", "江苏"}
        has_edu_cn_page2_evidence = region in {"浙江", "安徽", "福建", "江西", "山东", "河南", "湖北", "湖南", "广东", "广西", "海南", "重庆", "四川", "贵州", "云南", "西藏", "陕西", "甘肃", "青海", "宁夏", "新疆", "新疆生产建设兵团"}
        source_ids = (
            cctv_edu_cn_2014_low_source
            if has_edu_cn_page1_evidence
            else cctv_edu_cn_2014_low_page2_source
            if has_edu_cn_page2_evidence
            else cctv_source
        )
        evidence_source_text = "央视网报道和中国教育和科研计算机网转载报道均列出" if (has_edu_cn_page1_evidence or has_edu_cn_page2_evidence) else "央视网报道列出"
        add_warning_list(
            2014,
            normalized_region,
            "本科",
            "provincial_low_employment_list",
            "全国各省市低就业率本科专业名单",
            majors,
            "列入低就业率本科专业名单，供地方和高校专业结构调整、招生计划优化参考。",
            "各省市低就业率本科专业。",
            source_ids,
            f"{evidence_source_text}{region}低就业率本科专业名单。",
            "medium",
        )

    caijing_source = ["caijing_2016_provincial_warning"]
    liaoning_2016_negative_majors = [
        "经济学",
        "财政学",
        "金融学",
        "国际经济与贸易",
        "法学",
        "思想政治教育",
        "体育教育",
        "运动训练",
        "社会体育指导与管理",
        "汉语言文学",
        "汉语言",
        "英语",
        "日语",
        "新闻学",
        "广播电视学",
        "广告学",
        "历史学",
        "数学与应用数学",
        "信息与计算科学",
        "物理学",
        "化学",
        "应用化学",
        "生物科学",
        "生物技术",
        "应用心理学",
        "通信工程",
        "自动化",
        "计算机科学与技术",
        "服装设计与工程",
        "动物医学",
        "园林",
        "临床医学",
        "麻醉学",
        "口腔医学",
        "预防医学",
        "医学影像技术",
        "康复治疗学",
        "护理学",
        "医学影像学",
        "信息管理与信息系统",
        "工程管理",
        "工商管理",
        "市场营销",
        "会计学",
        "财务管理",
        "公共事业管理",
        "行政管理",
        "劳动与社会保障",
        "物流管理",
        "电子商务",
        "旅游管理",
        "音乐表演",
        "音乐学",
        "舞蹈表演",
        "表演",
        "广播电视编导",
        "播音与主持艺术",
        "动画",
        "美术学",
        "绘画",
        "雕塑",
        "视觉传达设计",
        "环境设计",
        "产品设计",
        "服装与服饰设计",
        "数字媒体艺术",
    ]
    if raw_dir is not None:
        liaoning_text_path = raw_dir / "liaoning_kjt_2016_undergrad_negative_list.txt"
        if liaoning_text_path.exists():
            liaoning_text = liaoning_text_path.read_text(encoding="utf-8", errors="replace")
            missing = [major for major in liaoning_2016_negative_majors if major not in liaoning_text]
            if missing:
                raise ValueError(f"Missing expected Liaoning 2016 negative-list majors: {missing}")
    for source_row_no, major in enumerate(liaoning_2016_negative_majors, start=1):
        rows.append(
            {
                "policy_year": 2016,
                "region": "辽宁省",
                "education_level": "本科",
                "record_type": "major_warning_list",
                "warning_label": "2016年度辽宁省建议高校暂缓申请增设本科专业名单",
                "reported_major_name": major,
                "policy_action": "建议高校2016年度暂缓申请增设。",
                "criterion_text": "省内布点数较多、就业率持续较低、不宜再重复设置的本科专业。",
                "source_row_no": str(source_row_no),
                "source_ids": "liaoning_kjt_2016_undergrad_negative_list",
                "evidence_text": f"辽宁省科技厅官网转载沈阳日报/沈阳网报道称辽宁省教育厅印发通知，2016年度建议高校暂缓申请增设本科专业名单第{source_row_no}项为{major}。",
                "confidence": "medium",
            }
        )
    add_warning_list(
        2016,
        "广东省",
        "本科",
        "major_warning_list",
        "广东省本科预警专业名单",
        ["历史学", "思想政治教育", "物理学", "表演", "中医学", "国际政治", "公共关系学", "安全工程", "生态学", "资源环境科学", "社会学", "应用电子技术教育", "法学", "汉语言文学"],
        "2016年广东省高校暂缓增设这些专业。",
        "广东省教育厅根据全省本科专业布点率、就业率、第一志愿录取率等情况制定。",
        caijing_source,
        "21世纪经济报道列出广东14个本科预警专业，并说明名单由广东省教育厅根据本科专业布点率、就业率、第一志愿录取率等情况制定。",
        "medium",
    )
    add_warning_list(
        2016,
        "河北省",
        "本科",
        "major_warning_list",
        "河北省本科预警专业名单",
        ["通信工程", "电气工程及其自动化", "汉语言文学", "土木工程", "英语", "电子信息工程", "工程管理", "法学", "计算机科学与技术", "会计学"],
        "列入省级预警名单，供高校和考生参考。",
        "主要依据毕业生就业率制定。",
        caijing_source,
        "21世纪经济报道列出河北10个本科预警专业，并说明名单制定主要依据毕业生就业率。",
        "medium",
    )
    add_warning_list(
        2016,
        "山东省",
        "本科",
        "provincial_low_employment_list",
        "山东省低就业率前六名本科专业",
        ["音乐表演", "应用心理学", "音乐学", "公共事业管理", "艺术设计学", "汉语言文学"],
        "位居山东低就业率前六名，供专业结构调整和招生计划优化参考。",
        "低就业率前六名。",
        caijing_source,
        "21世纪经济报道列出山东低就业率前六名专业：音乐表演、应用心理学、音乐学、公共事业管理、艺术设计学、汉语言文学。",
        "medium",
    )
    add_warning_list(
        2016,
        "辽宁省",
        "本科",
        "major_warning_list",
        "辽宁省本科预警专业部分列名",
        ["法学", "金融学", "计算机科学与技术", "通信工程", "新闻学", "国际经济与贸易", "工商管理", "电子商务", "园林", "表演", "美术学"],
        "列入辽宁省本科预警专业范围；原文称共66个专业，此处只结构化正文明确列出的专业。",
        "辽宁省布点数较多、就业率持续较低、不宜再重复设置的本科专业。",
        caijing_source,
        "21世纪经济报道称辽宁有66个预警专业，并明确列出法学、金融学、计算机科学与技术、通信工程、新闻学、国际经济与贸易、工商管理、电子商务、园林、表演、美术学等。",
        "low",
    )

    add_warning_list(
        2026,
        "武汉商学院",
        "本科",
        "major_warning_list",
        "2026年拟预警专业",
        ["经济与金融", "经济统计学", "会展经济与管理"],
        "拟列入学校本科专业预警名单。",
        "学院自查论证、学校专项督查与综合研判、校本科专业设置评议委员会评议。",
        ["acabridge_2026_wuhan_business_warning"],
        "学术桥转载武汉商学院公示，列出2026年拟预警专业为经济与金融、经济统计学、会展经济与管理。",
        "medium",
    )
    add_warning(
        2026,
        "武汉商学院",
        "本科",
        "major_stop_enrollment",
        "2026年拟停招专业",
        "国际商务",
        "拟停招。",
        "学院自查论证、学校专项督查与综合研判、校本科专业设置评议委员会评议。",
        ["acabridge_2026_wuhan_business_warning"],
        "学术桥转载武汉商学院公示，列出2026年拟停招专业为国际商务。",
        "medium",
    )

    add_warning_list(
        2018,
        "西安工程大学",
        "本科",
        "major_cancel",
        "2019年度拟撤销专业名单",
        ["包装工程", "市场营销", "光电信息科学与工程", "摄影", "测控技术与仪器"],
        "拟撤销。",
        "学校2019年度本科专业动态调整。",
        ["xpu_2019_dynamic_adjustment"],
        "西安工程大学2019年拟增设新专业和专业动态调整公示列出拟撤销专业：包装工程、市场营销、光电信息科学与工程、摄影、测控技术与仪器。",
        "high",
    )
    add_warning_list(
        2018,
        "西安工程大学",
        "本科",
        "major_stop_enrollment",
        "2019年度拟停招专业名单",
        ["信息管理与信息系统", "信息与计算科学", "数字媒体技术"],
        "拟停招。",
        "学校2019年度本科专业动态调整。",
        ["xpu_2019_dynamic_adjustment"],
        "西安工程大学2019年拟增设新专业和专业动态调整公示列出拟停招专业：信息管理与信息系统、信息与计算科学、数字媒体技术。",
        "high",
    )
    add_warning_list(
        2018,
        "西安工程大学",
        "本科",
        "major_warning_list",
        "2019年度预警专业名单",
        ["非织造材料与工程", "工业工程", "过程装备与控制工程", "环境科学", "生物工程", "行政管理", "汉语国际教育", "戏剧影视美术设计"],
        "拟列入校内预警专业名单。",
        "学校2019年度本科专业动态调整。",
        ["xpu_2019_dynamic_adjustment"],
        "西安工程大学2019年拟增设新专业和专业动态调整公示列出预警专业：非织造材料与工程、工业工程、过程装备与控制工程、环境科学、生物工程、行政管理、汉语国际教育、戏剧影视美术设计。",
        "high",
    )

    add_warning_list(
        2025,
        "武汉商学院",
        "本科",
        "major_warning_list",
        "2025年度拟预警专业",
        ["经济与金融", "会展经济与管理", "食品科学与工程", "软件工程"],
        "拟列入学校本科专业预警名单。",
        "武汉商学院2025年度本科专业调整公示。",
        ["wbu_2025_warning_stop"],
        "武汉商学院2025年度拟申报本科专业、预备案专业、预警和停招专业公示列出拟预警专业：经济与金融、会展经济与管理、食品科学与工程、软件工程。",
        "high",
    )
    add_warning(
        2025,
        "武汉商学院",
        "本科",
        "major_stop_enrollment",
        "2025年度拟停招专业",
        "经济与金融",
        "拟于2026年停招。",
        "武汉商学院2025年度本科专业调整公示。",
        ["wbu_2025_warning_stop"],
        "武汉商学院2025年度拟申报本科专业、预备案专业、预警和停招专业公示列出拟停招专业：经济与金融。",
        "high",
    )

    add_warning_list(
        2025,
        "沈阳化工大学",
        "本科",
        "major_cancel",
        "2025年拟撤销专业名单",
        ["信息与计算科学", "市场营销", "信息管理与信息系统", "公共事业管理", "通信工程"],
        "拟申请撤销。",
        "停招已满五年。",
        ["syuct_2025_major_cancellation"],
        "沈阳化工大学2025年拟撤销专业公示列出5个停招已满五年、拟申请撤销的专业：信息与计算科学、市场营销、信息管理与信息系统、公共事业管理、通信工程。",
        "high",
    )

    add_warning_list(
        2025,
        "南昌大学",
        "本科",
        "major_cancel",
        "2025年度拟撤销专业清单",
        ["戏剧影视文学", "广播电视编导", "动画", "艺术设计学", "劳动与社会保障", "经济统计学", "信息安全", "医学影像技术"],
        "拟撤销。",
        "连续五年停止招生且无在校学生的专业，原则上应予撤销；结合学校专业设置实际与专业结构持续优化需求。",
        ["ncu_2025_major_cancellation"],
        "南昌大学2025年度本科专业调整方案公示列出拟撤销专业清单，拟撤销戏剧影视文学、广播电视编导、动画、艺术设计学、劳动与社会保障、经济统计学、信息安全、医学影像技术8个本科专业。",
        "high",
    )

    add_warning_list(
        2025,
        "华侨大学",
        "本科",
        "major_cancel",
        "2025年度拟撤销专业清单",
        ["自动化", "会计学"],
        "拟撤销。",
        "校内专业设置评议专家组审议；自动化备注停招满5年且无在校生。",
        ["hqu_2025_major_cancellation"],
        "华侨大学2025年本科专业设置调整公示列出拟撤销自动化、会计学2个本科专业。",
        "high",
    )

    gznu_cancel_rows = [
        ("冶金工程", "080404"),
        ("工程管理", "120103"),
        ("电子信息科学与技术", "080714T"),
        ("数字媒体艺术", "130508"),
    ]
    for row_no, (major, code) in enumerate(gznu_cancel_rows, start=1):
        rows.append(
            {
                "policy_year": 2025,
                "region": "贵州师范大学",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "2025年度拟撤销专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(row_no),
                "policy_action": "拟撤销。",
                "criterion_text": "经学院申报、学校审议，结合学校专业调整优化方案。",
                "source_ids": "gznu_2025_major_cancellation",
                "evidence_text": f"贵州师范大学教务处2025年度本科专业设置情况公示列出拟撤销{major}（{code}）。",
                "confidence": "high",
            }
        )

    synu_cancel_rows = [
        ("食品质量与安全", "082702"),
        ("人力资源管理", "120206"),
        ("会展经济与管理", "120903"),
        ("动画", "130310"),
    ]
    for row_no, (major, code) in enumerate(synu_cancel_rows, start=1):
        rows.append(
            {
                "policy_year": 2025,
                "region": "沈阳师范大学",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "2025年度拟撤销本科专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(row_no),
                "policy_action": "拟撤销。",
                "criterion_text": "经学院申报、专家论证、校长办公会审议通过。",
                "source_ids": "synu_2025_major_cancellation_docx",
                "evidence_text": f"沈阳师范大学2025年度本科专业设置公示附件列出拟撤销{major}（{code}）。",
                "confidence": "high",
            }
        )

    ncpu_cancel_rows = [
        ("采购管理", "120603T", "经济与管理学院"),
        ("交通运输", "081801", "机械与车辆工程学院"),
        ("工业工程", "120701", "机械与车辆工程学院"),
        ("测绘工程", "081201", "建筑与环境工程学院"),
    ]
    for row_no, (major, code, school_unit) in enumerate(ncpu_cancel_rows, start=1):
        rows.append(
            {
                "policy_year": 2025,
                "region": "南昌工学院",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "2025年度拟撤销本科专业",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(row_no),
                "policy_action": "拟撤销。",
                "criterion_text": f"公示表列明停招年份为2020、2021、2022、2023、2024、2025；所属学院为{school_unit}。",
                "source_ids": "ncpu_2025_major_cancellation_notice;ncpu_2025_major_cancellation_table_image",
                "evidence_text": f"南昌工学院2025年度拟新增设、撤销本科专业公示页内嵌表格列出拟撤销{major}（{code}），停招年份为2020-2025。",
                "confidence": "high",
            }
        )

    rows.append(
        {
            "policy_year": 2025,
            "region": "武汉学院",
            "education_level": "本科",
            "record_type": "major_cancel",
            "warning_label": "2025年度拟撤销本科专业",
            "reported_major_name": "投资学",
            "major_code": "020304",
            "source_row_no": "1",
            "policy_action": "拟撤销。",
            "criterion_text": "学校研究决定；学科门类为经济学，所属学院为金融与经济学院。",
            "source_ids": "whxy_2025_major_cancellation",
            "evidence_text": "武汉学院2025年拟撤销本科专业公示列出拟撤销投资学（020304，经济学），所属学院为金融与经济学院。",
            "confidence": "high",
        }
    )

    gufe_cancel_rows = [
        ("投资学", "020304", "应用经济学院", "经济学"),
        ("公共事业管理", "120401", "公共管理学院", "管理学"),
        ("应用心理学", "071102", "公共管理学院", "理学"),
        ("工程造价", "120105", "管理科学与工程学院", "管理学"),
        ("房地产开发与管理", "120104", "管理科学与工程学院", "管理学"),
        ("教育技术学", "040104", "信息学院", "理学"),
        ("日语", "050207", "外语学院", "文学"),
    ]
    for row_no, (major, code, school_unit, degree) in enumerate(gufe_cancel_rows, start=1):
        rows.append(
            {
                "policy_year": 2025,
                "region": "贵州财经大学",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "2025年度拟撤销专业名单",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(row_no),
                "policy_action": "拟向教育部申请撤销。",
                "criterion_text": f"经学院论证分析、教务处审核、学校教学指导委员会审议通过；学院为{school_unit}，授予学位门类为{degree}。",
                "source_ids": "gufe_2025_major_cancellation_pdf",
                "evidence_text": f"贵州财经大学教务处PDF公示表第{row_no}行列出拟撤销{major}（{code}，{degree}），学院为{school_unit}。",
                "confidence": "high",
            }
        )

    hcnu_cancel_rows = [
        ("微电子科学与工程", "080704", "四年", "工学"),
        ("网络工程", "080903", "四年", "工学"),
    ]
    for row_no, (major, code, duration, degree) in enumerate(hcnu_cancel_rows, start=1):
        rows.append(
            {
                "policy_year": 2025,
                "region": "河池学院",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "2025年度拟撤销本科专业",
                "reported_major_name": major,
                "major_code": code,
                "study_duration": duration,
                "source_row_no": str(row_no),
                "policy_action": "拟撤销。",
                "criterion_text": "根据教育部普通高等学校本科专业设置工作通知要求，经学校研究。",
                "source_ids": "hcnu_2025_major_cancellation",
                "evidence_text": f"河池学院2025年拟撤销本科专业公示表第{row_no}行列出拟撤销{major}（{code}，{duration}，{degree}）。",
                "confidence": "high",
            }
        )

    wjut_cancel_rows = [
        ("公共艺术", "130506", "艺术学", "2015年"),
        ("应用化学", "070302", "理学", "2018年"),
        ("材料成型及控制工程", "080203", "工学", "2021年"),
    ]
    for row_no, (major, code, degree, stop_year) in enumerate(wjut_cancel_rows, start=1):
        rows.append(
            {
                "policy_year": 2025,
                "region": "皖江工学院",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "2025年撤销本科专业名单",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(row_no),
                "policy_action": "申请撤销。",
                "criterion_text": f"经学院申请、教学工作委员会审议；授予学位为{degree}，停招起始年份为{stop_year}。",
                "source_ids": "wjut_2025_cancel_stop",
                "evidence_text": f"皖江工学院2025年撤销本科专业名单第{row_no}行列出{major}（{code}，{degree}），停招起始年份为{stop_year}。",
                "confidence": "high",
            }
        )
    rows.append(
        {
            "policy_year": 2025,
            "region": "皖江工学院",
            "education_level": "本科",
            "record_type": "major_stop_enrollment",
            "warning_label": "2025年停招本科专业名单",
            "reported_major_name": "汽车服务工程",
            "major_code": "080208",
            "source_row_no": "1",
            "policy_action": "2025年起停止招生。",
            "criterion_text": "经学院申请、教学工作委员会审议；授予学位为工学，停招起始年份为2025年。",
            "source_ids": "wjut_2025_cancel_stop",
            "evidence_text": "皖江工学院2025年停招本科专业名单列出汽车服务工程（080208，工学），停招起始年份为2025年。",
            "confidence": "high",
        }
    )

    add_warning_list(
        2022,
        "西安建筑科技大学",
        "本科",
        "major_cancel",
        "2022年度拟撤销专业清单",
        ["产品设计", "摄影"],
        "拟撤销。",
        "征求学院意见、校内专业设置评议专家组和学校教学工作委员会审议、校长办公会审批通过。",
        ["xauat_2022_cancel_stop"],
        "西安建筑科技大学2022年度专业调整公示明确拟撤销产品设计、摄影2个本科专业。",
        "high",
    )
    add_warning(
        2022,
        "西安建筑科技大学",
        "本科",
        "major_stop_enrollment",
        "2022年度继续停招专业",
        "软件工程",
        "继续停招。",
        "征求学院意见、校内专业设置评议专家组和学校教学工作委员会审议、校长办公会审批通过。",
        ["xauat_2022_cancel_stop"],
        "西安建筑科技大学2022年度专业调整公示明确继续停招软件工程1个本科专业。",
        "high",
    )

    add_warning_list(
        2025,
        "南阳理工学院",
        "本科",
        "major_cancel",
        "2025年度拟申请撤销专业",
        ["法学", "音乐表演", "网络工程", "电子科学与技术", "汽车服务工程"],
        "拟申请撤销。",
        "停招五年及以上且无在籍学生。",
        ["nyist_2025_major_cancellation"],
        "南阳理工学院2025年度本科专业设置与调整公示列出拟申请撤销法学、音乐表演、网络工程、电子科学与技术、汽车服务工程5个本科专业。",
        "high",
    )

    add_warning_list(
        2025,
        "江西师范大学",
        "本科",
        "major_cancel",
        "2025年度本科专业撤销情况",
        ["网络工程", "物联网工程", "经济统计学", "信息与计算科学", "汉语言", "翻译", "劳动与社会保障"],
        "拟实施撤销处理。",
        "学院申报、学校审查。",
        ["jxnu_2025_major_cancellation"],
        "江西师范大学教务处《关于我校2025年度本科专业撤销情况的公示》列出拟对网络工程、物联网工程、经济统计学、信息与计算科学、汉语言、翻译、劳动与社会保障七个专业实施撤销处理。",
        "high",
    )

    add_warning_list(
        2025,
        "井冈山大学",
        "本科",
        "major_cancel",
        "2025年拟新增和撤销本科专业",
        ["商务英语", "应用物理学", "动画"],
        "拟撤销。",
        "学院申报、专家论证、学校学术委员会评议、校长办公会审议。",
        ["jgsu_2025_major_cancellation"],
        "井冈山大学教务处《关于2025年拟新增和撤销本科专业的公示》列出拟撤销商务英语、应用物理学、动画3个本科专业。",
        "high",
    )

    add_warning(
        2025,
        "青岛大学",
        "本科",
        "major_cancel",
        "2025年度拟撤销专业",
        "数字媒体技术",
        "予以撤销。",
        "近五年专业停招情况。",
        ["qdu_2025_major_cancellation"],
        "青岛大学2025年本科专业设置调整情况公示列出拟撤销近五年停招专业：数字媒体技术。",
        "high",
    )

    add_warning_list(
        2025,
        "湖北大学",
        "本科",
        "major_cancel",
        "2025年度拟申请撤销专业",
        ["工程管理", "公共事业管理"],
        "拟申请撤销。",
        "已连续停招五年及以上的专业，原则上应当予以撤销。",
        ["hubu_2025_major_cancellation"],
        "湖北大学2025年度本科专业设置公示列出拟申请撤销工程管理、公共事业管理2个本科专业。",
        "high",
    )

    add_warning(
        2024,
        "武汉商学院",
        "本科",
        "major_warning_list",
        "2024年度拟预警专业",
        "经济与金融",
        "拟预警。",
        "学校本科专业设置评议委员会评议。",
        ["wbu_2024_warning_stop"],
        "武汉商学院2024年度拟申报本科专业、预备案专业、预警和停招专业公示列出拟预警专业：经济与金融。",
        "high",
    )
    add_warning_list(
        2024,
        "武汉商学院",
        "本科",
        "major_stop_enrollment",
        "2024年度拟停招专业",
        ["汽车服务工程", "马业科学"],
        "拟停招。",
        "学校本科专业设置评议委员会评议。",
        ["wbu_2024_warning_stop"],
        "武汉商学院2024年度拟申报本科专业、预备案专业、预警和停招专业公示列出拟停招专业：汽车服务工程、马业科学。",
        "high",
    )

    add_warning(
        2025,
        "河西学院",
        "本科",
        "major_cancel",
        "2025年度拟申请撤销专业",
        "绘画",
        "拟申请撤销。",
        "学院论证申报、学校专业设置评议专家组评议、校长办公会审议通过。",
        ["hxu_2025_major_cancellation"],
        "河西学院2025年度拟申请增设本科专业、预备案本科专业和撤销本科专业公示列出拟申请撤销绘画本科专业。",
        "high",
    )

    add_warning_list(
        2024,
        "清华大学",
        "本科",
        "major_stop_enrollment",
        "2023—2024学年度不招生专业",
        ["微机电系统工程", "广告学"],
        "2023—2024学年度不招生。",
        "清华大学2023—2024学年度专业设置页在专业说明中明确标注“不招生”。",
        ["tsinghua_2024_not_enrolling"],
        "清华大学2023—2024学年度专业设置、当年新增专业、停招专业名单页说明：微机电系统工程、广告学，不招生。",
        "high",
    )

    fudan_stop_majors = ["应用化学", "电气工程及其自动化", "保密管理"]
    add_warning_list(
        2023,
        "复旦大学",
        "本科",
        "major_stop_enrollment",
        "2023年招生专业调整停招专业",
        fudan_stop_majors,
        "2023年停招。",
        "复旦大学教务处招生专业调整情况页正文列出2023年停招的专业。",
        ["fudan_2023_stop_enrollment"],
        "复旦大学教务处招生专业调整情况页写明：2023年停招应用化学、电气工程及其自动化、保密管理3个专业。",
        "high",
    )
    add_warning_list(
        2024,
        "复旦大学",
        "本科",
        "major_stop_enrollment",
        "2024年学科专业设置停招本科专业",
        ["应用化学", "保密管理", "电气工程及其自动化"],
        "2024年停招。",
        "复旦大学2024年学科专业设置、当年新增或停招专业名单页正文列出2024年停招的本科专业。",
        ["fudan_2024_stop_enrollment"],
        "复旦大学2024年学科专业设置、当年新增或停招专业名单页写明：2024年停招应用化学、保密管理、电气工程及其自动化3个本科专业。",
        "high",
    )
    add_warning(
        2024,
        "复旦大学",
        "本科第二学士学位",
        "major_stop_enrollment",
        "2024年第二学士学位项目停招",
        "电子信息科学与技术（智能科学与技术方向）",
        "2024年停招第二学士学位项目。",
        "复旦大学2024年学科专业设置、当年新增或停招专业名单页正文列出停招第二学士学位项目。",
        ["fudan_2024_stop_enrollment"],
        "复旦大学2024年学科专业设置、当年新增或停招专业名单页写明：停招第二学士学位项目电子信息科学与技术（智能科学与技术方向）。",
        "high",
    )

    sjtu_stop_rows = [
        (2024, "2024年专业设置及调整情况当年停招专业", ["建筑学", "风景园林"], "sjtu_2024_stop_enrollment"),
        (2023, "2023年专业设置及调整情况当年停招专业", ["交通运输"], "sjtu_2023_stop_enrollment"),
        (2022, "2022年专业设置及调整情况当年停招专业", ["测控技术与仪器"], "sjtu_2022_stop_enrollment"),
        (2020, "2020年度专业设置及调整情况当年停招专业", ["工程力学(海洋科学与技术)", "生物科学(海洋科学)"], "sjtu_2020_stop_enrollment"),
    ]
    for policy_year, label, majors, source_id in sjtu_stop_rows:
        add_warning_list(
            policy_year,
            "上海交通大学",
            "本科",
            "major_stop_enrollment",
            label,
            majors,
            f"{policy_year}年停招。",
            "上海交通大学信息公开网年度表格“当年停招专业名单”直接列出。",
            [source_id],
            f"上海交通大学信息公开网《{label}》列出当年停招专业：{'、'.join(majors)}。",
            "high",
        )

    bnu_2023_stop_rows = [
        ("66", "081102", "水文与水资源工程", "工学"),
        ("67", "081202", "遥感科学与技术", "工学"),
        ("68", "082902T", "应急技术与管理", "工学"),
        ("69", "070302", "应用化学", "理学"),
        ("70", "071102", "应用心理学", "理学"),
        ("71", "080402", "材料物理", "理学"),
        ("72", "120212T", "体育经济与管理", "管理学"),
        ("73", "130303", "电影学", "艺术学"),
        ("74", "130502", "视觉传达设计", "艺术学"),
        ("75", "100701", "药学", "理学"),
        ("76", "080714T", "电子信息科学与技术", "理学"),
        ("77", "120101", "管理科学", "管理学"),
    ]
    for source_row_no, code, major, degree in bnu_2023_stop_rows:
        rows.append(
            {
                "policy_year": 2023,
                "region": "北京师范大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2022-2023学年专业设置、新增专业、停招专业名单",
                "reported_major_name": major,
                "major_code": code,
                "study_duration": "四年",
                "policy_action": "暂停招生。",
                "criterion_text": "北京师范大学2022-2023学年专业设置、新增专业、停招专业名单PDF在备注列标记“暂停招生”。",
                "source_row_no": source_row_no,
                "source_ids": "bnu_2023_stop_enrollment_pdf",
                "evidence_text": f"北京师范大学2022-2023学年专业设置、新增专业、停招专业名单PDF第{source_row_no}项列出{major}（{code}，四年，{degree}），备注为暂停招生。",
                "confidence": "high",
            }
        )

    bnu_2025_stop_rows = [
        ("69", "020109T", "数字经济", "经济学"),
        ("70", "130509T", "艺术与科技", "艺术学"),
        ("71", "081102", "水文与水资源工程", "工学"),
        ("72", "081202", "遥感科学与技术", "工学"),
        ("73", "082902T", "应急技术与管理", "工学"),
        ("74", "070302", "应用化学", "理学"),
        ("75", "071102", "应用心理学", "理学"),
        ("76", "080714T", "电子信息科学与技术", "理学"),
        ("77", "120101", "管理科学", "管理学"),
    ]
    for source_row_no, code, major, degree in bnu_2025_stop_rows:
        rows.append(
            {
                "policy_year": 2025,
                "region": "北京师范大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2024-2025学年专业设置、新增专业、停招专业名单",
                "reported_major_name": major,
                "major_code": code,
                "study_duration": "四年",
                "policy_action": "暂停招生。",
                "criterion_text": "北京师范大学2024-2025学年专业设置、新增专业、停招专业名单PDF在备注列标记“暂停招生”。",
                "source_row_no": source_row_no,
                "source_ids": "bnu_2025_stop_enrollment_index;bnu_2025_stop_enrollment_pdf",
                "evidence_text": f"北京师范大学2024-2025学年专业设置、新增专业、停招专业名单PDF第{source_row_no}项列出{major}（{code}，四年，{degree}），备注为暂停招生。",
                "confidence": "high",
            }
        )

    add_warning(
        2017,
        "电子科技大学",
        "本科",
        "major_stop_enrollment",
        "2017年本科专业停招情况",
        "环境工程",
        "2017年停招。",
        "电子科技大学信息公开页正文明确说明2017年停招环境工程一个专业。",
        ["uestc_2017_stop_enrollment"],
        "电子科技大学专业设置、当年新增专业及停招专业情况页说明：截至2017年8月，学校共有本科专业66个，2017年新增5个专业，停招环境工程1个专业。",
        "high",
    )

    ecnu_2023_stop_rows = [
        ("3", "自然地理与资源环境", "070502"),
        ("8", "环境工程", "082502"),
        ("9", "教育学", "040101"),
        ("11", "艺术教育", "040105"),
        ("23", "会展经济与管理", "120903"),
        ("24", "房地产开发与管理", "120104"),
        ("25", "国际经济与贸易", "020401"),
        ("26", "经济统计学", "020102"),
        ("29", "金融工程", "020302"),
        ("43", "公共关系学", "120409T"),
        ("64", "广告学", "050303"),
        ("76", "信息与计算科学", "070102"),
        ("78", "电子科学与技术", "080702"),
        ("79", "材料科学与工程", "080401"),
        ("81", "光电信息科学与工程", "080705"),
        ("83", "应用化学", "070302"),
    ]
    for source_row_no, major, code in ecnu_2023_stop_rows:
        rows.append(
            {
                "policy_year": 2023,
                "region": "华东师范大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2023年度暂停招生专业",
                "reported_major_name": major,
                "major_code": code,
                "policy_action": "暂停招生。",
                "criterion_text": "华东师范大学2023年本科专业一览表备注列标记“暂停招生”。",
                "source_row_no": source_row_no,
                "source_ids": "ecnu_2023_stop_enrollment",
                "evidence_text": f"华东师范大学2023年本科专业设置及新增专业、停招专业情况表第{source_row_no}项列出{major}（{code}），备注为暂停招生。",
                "confidence": "high",
            }
        )

    ecnu_2024_stop_majors = [
        "自然地理与资源环境",
        "环境工程",
        "教育学",
        "艺术教育",
        "信息管理与信息系统",
        "旅游管理",
        "会展经济与管理",
        "房地产开发与管理",
        "国际经济与贸易",
        "经济统计学",
        "金融工程",
        "公共关系学",
        "翻译",
        "德语",
        "广告学",
        "绘画",
        "雕塑",
        "信息与计算科学",
        "电子科学与技术",
        "材料科学与工程",
        "光电信息科学与工程",
        "应用化学",
    ]
    add_warning_list(
        2024,
        "华东师范大学",
        "本科",
        "major_stop_enrollment",
        "2024年度暂停招生专业",
        ecnu_2024_stop_majors,
        "暂停招生。",
        "华东师范大学本科专业一览表备注列标记“暂停招生”。",
        ["ecnu_2024_stop_enrollment"],
        "华东师范大学2024年本科专业设置及新增专业、停招专业情况页在本科专业一览表备注列标记22个暂停招生专业。",
        "high",
    )

    add_warning_list(
        2025,
        "华东师范大学",
        "本科",
        "major_stop_enrollment",
        "2025年度停招专业可见名单",
        ecnu_2024_stop_majors,
        "停招。",
        "信息公开页称停招专业24个；正文和meta描述实际列出22个专业名称，本数据集仅结构化这些明示名称。",
        ["ecnu_2025_stop_enrollment"],
        "华东师范大学2025年本科专业设置及新增专业、停招专业情况页称停招专业24个，正文可见自然地理与资源环境、环境工程、教育学、艺术教育、信息管理与信息系统、旅游管理、会展经济与管理、房地产开发与管理、国际经济与贸易、经济统计学、金融工程、公共关系学、翻译、德语、广告学、绘画、雕塑、信息与计算科学、电子科学与技术、材料科学与工程、光电信息科学与工程、应用化学22个专业名称。",
        "high",
    )

    add_warning_list(
        2025,
        "上海商学院",
        "本科",
        "major_stop_enrollment",
        "2025年本科专业停招情况",
        ["园林", "广告学"],
        "停招。",
        "上海商学院信息公开页列出2025年停招本科专业。",
        ["sbs_2025_stop_enrollment"],
        "上海商学院2025年本科专业新增、停招情况信息公开页列出2025年停招的本科专业为园林、广告学。",
        "high",
    )

    add_warning_list(
        2025,
        "上海海洋大学",
        "本科",
        "major_stop_enrollment",
        "2025年停招专业",
        ["行政管理", "物流管理", "工业工程", "软件工程", "朝鲜语"],
        "停招。",
        "上海海洋大学信息公开页列出2025年停招专业；页面正文明确列名。",
        ["shou_2025_stop_enrollment"],
        "上海海洋大学2025年新增和停招专业信息公开页列出2025年停招专业为行政管理、物流管理、工业工程、软件工程、朝鲜语。",
        "high",
    )
    add_warning_list(
        2024,
        "上海海洋大学",
        "本科",
        "major_stop_enrollment",
        "2023-2024学年本科教学质量报告停招专业",
        ["环境科学", "环境工程", "文化产业管理"],
        "停招。",
        "上海海洋大学2023-2024学年本科教学质量报告说明2023学年制定拟停招专业工作方案并停招3个专业。",
        ["shou_2024_teaching_quality_report_pdf"],
        "上海海洋大学2023-2024学年本科教学质量报告专业建设部分称，因专业师资不足、专业质量指标水平偏低等因素，停招环境科学、环境工程、文化产业管理3个专业。",
        "high",
    )
    add_warning_list(
        2024,
        "上海大学",
        "本科",
        "major_stop_enrollment",
        "2023-2024学年本科教育教学质量报告已停招专业",
        ["材料物理", "工业设计", "包装工程"],
        "停招。",
        "上海大学2023-2024学年本科教育教学质量报告附录专业设置及调整情况列出已停招专业名单。",
        ["shu_2024_teaching_quality_report_pdf"],
        "上海大学2023-2024学年本科教育教学质量报告附录列明已停招专业名单为材料物理、工业设计、包装工程。",
        "high",
    )
    add_warning_list(
        2024,
        "上海大学",
        "本科",
        "major_cancel",
        "2023-2024学年本科教育教学质量报告申请备案撤销专业",
        ["包装工程"],
        "申请备案撤销。",
        "上海大学2023-2024学年本科教育教学质量报告称向教育主管部门申请备案撤销包装工程专业。",
        ["shu_2024_teaching_quality_report_pdf"],
        "上海大学2023-2024学年本科教育教学质量报告基本情况部分称，学校向教育主管部门申请备案撤销包装工程专业；该专业同时见教育部2024年度普通高等学校本科专业备案和审批结果撤销名单。",
        "high",
    )
    add_warning_list(
        2023,
        "内蒙古财经大学",
        "本科",
        "major_cancel",
        "2023-2024学年本科教学质量报告撤销专业",
        ["信息与计算科学"],
        "撤销。",
        "内蒙古财经大学2023-2024学年本科教学质量报告称2023年撤销信息与计算科学等5个本科专业；本数据只结构化报告明示名称。",
        ["imufe_2024_teaching_quality_report_pdf"],
        "内蒙古财经大学2023-2024学年本科教学质量报告专业动态调整机制部分称，2023年撤销信息与计算科学等5个本科专业；本行对应报告明示的信息与计算科学专业。",
        "high",
    )
    add_warning_list(
        2023,
        "内蒙古财经大学",
        "本科",
        "major_stop_enrollment",
        "2023-2024学年本科教学质量报告暂缓招生专业",
        ["物业管理"],
        "暂缓招生。",
        "内蒙古财经大学2023-2024学年本科教学质量报告称2023年对物业管理等6个本科专业暂缓招生；本数据只结构化报告明示名称。",
        ["imufe_2024_teaching_quality_report_pdf"],
        "内蒙古财经大学2023-2024学年本科教学质量报告专业动态调整机制部分称，2023年对物业管理等6个本科专业暂缓招生；本行对应报告明示的物业管理专业。",
        "high",
    )
    add_warning_list(
        2024,
        "内蒙古财经大学",
        "本科",
        "major_stop_enrollment",
        "2023-2024学年本科教学质量报告暂缓招生专业",
        ["保险学", "公共事业管理"],
        "暂缓招生。",
        "内蒙古财经大学2023-2024学年本科教学质量报告称2024年对保险学和公共事业管理共2个本科专业暂缓招生。",
        ["imufe_2024_teaching_quality_report_pdf"],
        "内蒙古财经大学2023-2024学年本科教学质量报告专业动态调整机制部分称，2024年对保险学和公共事业管理共2个本科专业暂缓招生。",
        "high",
    )
    add_warning_list(
        2024,
        "重庆文理学院",
        "本科",
        "major_stop_enrollment",
        "2023-2024学年本科教学质量报告停招专业",
        ["广播电视学", "美术学", "旅游管理与服务教育", "金融数学"],
        "停招。",
        "重庆文理学院2023-2024学年本科教学质量报告说明2024年停招4个专业。",
        ["cqwu_2024_teaching_quality_report_pdf"],
        "重庆文理学院2023-2024学年本科教学质量报告学科专业设置部分称，2024年停招广播电视学、美术学、旅游管理与服务教育、金融数学等4个专业。",
        "high",
    )
    add_warning_list(
        2024,
        "重庆文理学院",
        "本科",
        "major_cancel",
        "2023-2024学年本科教学质量报告主动撤销专业",
        ["经济统计学"],
        "主动撤销。",
        "重庆文理学院2023-2024学年本科教学质量报告说明2024年主动撤销经济统计学专业。",
        ["cqwu_2024_teaching_quality_report_pdf"],
        "重庆文理学院2023-2024学年本科教学质量报告学科专业设置部分称，2024年主动撤销经济统计学专业。",
        "high",
    )
    add_warning_list(
        2024,
        "东北石油大学",
        "本科",
        "major_stop_enrollment",
        "2023-2024学年本科教学质量报告停招专业",
        ["光电信息科学与工程", "智能电网信息工程", "财务管理"],
        "停招。",
        "东北石油大学2023-2024学年本科教学质量报告列明2024年停招3个本科专业。",
        ["nepu_2024_teaching_quality_report_pdf"],
        "东北石油大学2023-2024学年本科教学质量报告专业设置部分和支撑数据均列出当年停招专业名单：光电信息科学与工程、智能电网信息工程、财务管理。",
        "high",
    )
    add_warning_list(
        2024,
        "东北石油大学",
        "本科",
        "major_cancel",
        "2023-2024学年本科教学质量报告撤销专业",
        ["公共事业管理"],
        "撤销。",
        "东北石油大学2023-2024学年本科教学质量报告说明撤销公共事业管理专业。",
        ["nepu_2024_teaching_quality_report_pdf"],
        "东北石油大学2023-2024学年本科教学质量报告专业设置部分称撤销公共事业管理专业，支撑数据注释说明该专业已无在校生并于2024年撤销。",
        "high",
    )
    add_warning_list(
        2024,
        "重庆科技大学",
        "本科",
        "major_stop_enrollment",
        "2023-2024学年本科教学质量报告暂停招生/停招专业",
        ["矿物加工工程", "过程装备与控制工程", "地球物理学", "材料物理"],
        "暂停招生或停招。",
        "重庆科技大学2023-2024学年本科教学质量报告正文逐名提及部分暂停招生/停招专业；未明示的“等9个”不作结构化名单行。",
        ["cqust_2024_teaching_quality_report_pdf"],
        "重庆科技大学2023-2024学年本科教学质量报告本科专业设置部分称矿物加工工程、过程装备与控制工程等9个专业暂停招生，专业建设部分称停招地球物理学、材料物理等9个专业（当年停招7个）；本行仅对应报告明示专业名。",
        "high",
    )
    add_warning_list(
        2024,
        "宜春学院",
        "本科",
        "major_stop_enrollment",
        "2023-2024学年本科教学质量报告停招专业",
        ["广播电视编导"],
        "停招。",
        "宜春学院2023-2024学年本科教学质量报告称停招广播电视编导等14个专业；本数据只结构化报告明示名称。",
        ["jxycu_2024_teaching_quality_report_pdf"],
        "宜春学院2023-2024学年本科教学质量报告专业建设部分称停招“广播电视编导”等14个专业；本行对应报告明示的广播电视编导专业。",
        "high",
    )
    add_warning_list(
        2024,
        "山东航空学院",
        "本科",
        "major_stop_enrollment",
        "2023-2024学年本科教学质量报告当年停招专业",
        ["法语", "飞行器适航技术"],
        "停招。",
        "山东航空学院2023-2024学年本科教学质量报告核心支撑数据列出当年停招专业。",
        ["sdua_2024_teaching_quality_report_pdf"],
        "山东航空学院2023-2024学年本科教学质量报告核心支撑数据列出当年停招专业为法语、飞行器适航技术。",
        "high",
    )
    add_warning_list(
        2024,
        "山东航空学院",
        "本科",
        "major_cancel",
        "2023-2024学年本科教学质量报告撤销专业",
        ["能源化学工程"],
        "撤销。",
        "山东航空学院2023-2024学年本科教学质量报告说明撤销能源化学工程专业。",
        ["sdua_2024_teaching_quality_report_pdf"],
        "山东航空学院2023-2024学年本科教学质量报告专业建设部分称撤销能源化学工程专业；该专业同时见教育部2024年度普通高等学校本科专业备案和审批结果撤销名单。",
        "high",
    )
    add_warning_list(
        2024,
        "北京科技大学",
        "本科",
        "major_stop_enrollment",
        "2023-2024学年本科教学质量报告停招专业",
        ["思想政治教育", "生态学", "电子信息工程", "矿物资源工程", "工业工程", "智能科学与技术"],
        "停招。",
        "北京科技大学2023-2024学年本科教学质量报告列出6个停招本科专业。",
        ["ustb_2024_teaching_quality_report_pdf"],
        "北京科技大学2023-2024学年本科教学质量报告本科专业设置部分称，本科专业中共有6个专业停招，包括思想政治教育、生态学、电子信息工程、矿物资源工程、工业工程、智能科学与技术。",
        "high",
    )
    add_warning_list(
        2024,
        "贵州财经大学",
        "本科",
        "major_stop_enrollment",
        "2023-2024学年本科教学质量报告停招专业",
        ["投资学", "保险学", "应用心理学", "网络工程", "教育技术学", "英语", "葡萄牙语", "汉语国际教育", "公共事业管理", "艺术管理"],
        "停招。",
        "贵州财经大学2023-2024学年本科教学质量报告列出10个停招本科专业。",
        ["gufe_2024_teaching_quality_report_pdf"],
        "贵州财经大学2023-2024学年本科教学质量报告专业建设部分称，学校停招了投资学、保险学、应用心理学、网络工程、教育技术学、英语、葡萄牙语、汉语国际教育、公共事业管理、艺术管理等10个专业。",
        "high",
    )
    add_warning_list(
        2024,
        "贵州财经大学",
        "本科",
        "major_cancel",
        "2023-2024学年本科教学质量报告申请撤销专业",
        ["英语", "会展经济与管理", "金融科技"],
        "申请撤销。",
        "贵州财经大学2023-2024学年本科教学质量报告称申请撤销英语、会展经济与管理、金融科技3个专业。",
        ["gufe_2024_teaching_quality_report_pdf"],
        "贵州财经大学2023-2024学年本科教学质量报告专业建设部分称，学校申请撤销英语、会展经济与管理、金融科技专业；这些专业同时见教育部2024年度普通高等学校本科专业备案和审批结果撤销名单。",
        "high",
    )

    add_warning_list(
        2025,
        "上海戏剧学院",
        "本科",
        "major_stop_enrollment",
        "2025-2026学年当年停招专业",
        ["动画", "数字演艺设计"],
        "当年停招。",
        "上海戏剧学院信息公开页列出2025-2026学年当年停招专业。",
        ["sta_2025_stop_enrollment"],
        "上海戏剧学院2025-2026学年学科专业建设、当年新增或停招专业名单页列出当年停招专业为动画、数字演艺设计。",
        "high",
    )

    lixin_stop_rows = [
        ("1", "金融工程（中美合作）", "020302H", "2023、2024、2025年停招"),
        ("2", "信用管理", "020306T", "2025年停招"),
        ("3", "房地产开发与管理", "120104", "2024、2025年停招"),
        ("4", "劳动与社会保障", "120403", "2024、2025年停招"),
        ("5", "日语", "050207", "2025年停招"),
    ]
    for source_row_no, major, code, note in lixin_stop_rows:
        rows.append(
            {
                "policy_year": 2025,
                "region": "上海立信会计金融学院",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2025年当年新增或停招专业名单/本科专业设置情况",
                "reported_major_name": major,
                "major_code": code,
                "policy_action": "停招。",
                "criterion_text": f"上海立信会计金融学院本科专业设置情况表在专业名称旁标注“{note}”。",
                "source_row_no": source_row_no,
                "source_ids": "lixin_2025_stop_enrollment;lixin_2025_undergrad_setting",
                "evidence_text": f"上海立信会计金融学院2025年本科专业设置情况表列出{major}（{code}），备注为{note}；当年新增或停招专业名单页同步说明学校2025年停招专业。",
                "confidence": "high",
            }
        )

    tongji_historical_stop_configs = [
        (
            2020,
            "tongji_2020_stop_enrollment",
            "2020级停招专业名单",
            [
                ("1", "交通运输工程学院", "15034", "物流工程"),
                ("2", "电子与信息工程学院", "10074", "电子科学与技术"),
                ("3", "海洋与地球科学学院", "31014", "地质学"),
            ],
        ),
        (
            2021,
            "tongji_2021_stop_enrollment",
            "2021级停招专业名单",
            [("1", "医学院", "14054", "康复治疗学")],
        ),
        (
            2022,
            "tongji_2022_stop_enrollment",
            "2022级停招专业名单",
            [
                ("1", "医学院", "", "临床医学（拔尖卓越培养）"),
                ("2", "政治与国际关系学院", "", "社会学"),
            ],
        ),
        (
            2023,
            "tongji_2023_stop_enrollment",
            "2023级停招专业名单",
            [
                ("1", "经济与管理学院", "", "市场营销"),
                ("2", "经济与管理学院", "", "行政管理"),
                ("3", "经济与管理学院", "", "国际经济与贸易"),
                ("4", "化学科学与工程学院", "", "化学工程与工艺"),
                ("5", "电子与信息工程学院", "", "电子信息工程"),
                ("6", "海洋与地球科学学院", "", "海洋技术"),
                ("7", "土木工程学院", "", "港口航道与海岸工程"),
                ("8", "艺术与传媒学院", "", "广告学"),
                ("9", "机械与能源工程学院", "", "工业工程"),
            ],
        ),
        (
            2024,
            "tongji_2024_stop_enrollment",
            "2024级停招专业",
            [("1", "艺术与传媒学院", "", "广播电视学")],
        ),
    ]
    for policy_year, source_id, warning_label, stop_rows in tongji_historical_stop_configs:
        for source_row_no, school_unit, internal_code, major in stop_rows:
            internal_phrase = f"，校内专业编号{internal_code}" if internal_code else ""
            rows.append(
                {
                    "policy_year": policy_year,
                    "region": "同济大学",
                    "education_level": "本科",
                    "record_type": "major_stop_enrollment",
                    "warning_label": warning_label,
                    "reported_major_name": major,
                    "policy_action": "停招。",
                    "criterion_text": f"同济大学信息公开页“{warning_label}”表列出停招专业；页面未提供教育部专业代码，若有编号则为校内编号。",
                    "source_row_no": source_row_no,
                    "source_ids": source_id,
                    "evidence_text": f"同济大学{warning_label}表第{source_row_no}项列出{school_unit}{major}{internal_phrase}，为停招专业。",
                    "confidence": "high",
                }
            )

    tongji_stop_rows = [
        ("1", "海洋科学", "070701", "海洋与地球科学学院", ""),
        ("2", "汽车服务工程", "080208", "中德工程学院", ""),
        ("3", "物流管理", "120601", "经济与管理学院", ""),
        ("4", "视觉传达设计", "130502", "设计创意学院", "第二学士学位同步停招"),
        ("5", "环境设计", "130503", "设计创意学院", ""),
        ("6", "产品设计", "130504", "设计创意学院", ""),
    ]
    for source_row_no, major, code, school_unit, note in tongji_stop_rows:
        rows.append(
            {
                "policy_year": 2025,
                "region": "同济大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2025级停招专业",
                "reported_major_name": major,
                "major_code": code,
                "policy_action": "停招。",
                "criterion_text": "同济大学信息公开页“2025级停招专业”表列出专业代码、专业名称、所属学院和备注。",
                "source_row_no": source_row_no,
                "source_ids": "tongji_2025_stop_enrollment",
                "evidence_text": f"同济大学2025级停招专业表第{source_row_no}项列出{major}（{code}，{school_unit}）{('，备注为' + note) if note else ''}。",
                "confidence": "high",
            }
        )

    ahtcm_stop_rows = [
        (2024, "ahtcm_2024_stop_enrollment;ahtcm_2024_major_setting", "2024年停招专业情况说明", "14", "药物分析", "100705T", "四年", "理学"),
        (2024, "ahtcm_2024_stop_enrollment;ahtcm_2024_major_setting", "2024年停招专业情况说明", "19", "保险学", "020303", "四年", "经济学"),
        (2024, "ahtcm_2024_stop_enrollment;ahtcm_2024_major_setting", "2024年停招专业情况说明", "21", "中药资源与开发", "100802", "四年", "理学"),
        (2024, "ahtcm_2024_stop_enrollment;ahtcm_2024_major_setting", "2024年停招专业情况说明", "24", "中医儿科学", "100512TK", "五年", "医学"),
        (2025, "ahtcm_2025_stop_enrollment;ahtcm_2025_major_setting", "2025年停招专业情况说明", "15", "人力资源管理", "120206", "四年", "管理学"),
        (2025, "ahtcm_2025_stop_enrollment;ahtcm_2025_major_setting", "2025年停招专业情况说明", "19", "保险学", "020303", "四年", "经济学"),
        (2025, "ahtcm_2025_stop_enrollment;ahtcm_2025_major_setting", "2025年停招专业情况说明", "24", "中医儿科学", "100512TK", "五年", "医学"),
    ]
    for policy_year, source_ids, warning_label, source_row_no, major, major_code, duration, degree in ahtcm_stop_rows:
        rows.append(
            {
                "policy_year": policy_year,
                "region": "安徽中医药大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": warning_label,
                "reported_major_name": major,
                "major_code": major_code,
                "study_duration": duration,
                "policy_action": "停招。",
                "criterion_text": f"安徽中医药大学教务处{policy_year}年停招专业情况说明明确列名；同年专业设置情况表招生状态列逐项标注停招。",
                "source_row_no": source_row_no,
                "source_ids": source_ids,
                "evidence_text": f"安徽中医药大学{policy_year}年专业设置情况表第{source_row_no}项列出{major}（{major_code}，{duration}，{degree}），招生状态为停招；停招专业情况说明页同步列出该专业。",
                "confidence": "high",
            }
        )

    add_warning_list(
        2025,
        "上海电力大学",
        "本科",
        "major_stop_enrollment",
        "2025年停招本科专业",
        ["公共事业管理", "材料科学与工程", "网络工程", "机械电子工程", "材料化学", "物流管理", "日语"],
        "停招。",
        "上海电力大学本科生院保留2025年新增、停招本科专业情况页；公开HTML正文未暴露名单，结构化名单来自可抓取二级汇总。",
        ["shiep_2025_stop_enrollment", "sohu_2025_shanghai_university_stop_summary"],
        "搜狐转载汇总列出上海电力大学2025年停招公共事业管理、材料科学与工程、网络工程、机械电子工程、材料化学、物流管理、日语7个专业；上海电力大学本科生院同期公开2025年新增、停招本科专业情况页。",
        "medium",
    )

    add_warning_list(
        2025,
        "浙江财经大学",
        "本科",
        "major_stop_enrollment",
        "2025年停招专业",
        ["城市管理", "资产评估", "保险学", "信用管理", "物流管理", "国民经济管理", "日语", "汉语国际教育"],
        "停招。",
        "可抓取二级汇总列出浙江财经大学2025年停招专业；学校原信息公开链接本地抓取返回404，需后续复核官方原页。",
        ["sohu_2025_shanghai_university_stop_summary"],
        "搜狐转载汇总列出浙江财经大学2025年停招城市管理、资产评估、保险学、信用管理、物流管理、国民经济管理、日语、汉语国际教育8个专业。",
        "medium",
    )
    add_warning_list(
        2025,
        "浙江财经大学",
        "本科",
        "major_cancel",
        "2025年撤销专业",
        ["公共事业管理", "数字媒体艺术"],
        "撤销。",
        "可抓取二级汇总列出浙江财经大学2025年撤销专业；学校原信息公开链接本地抓取返回404，需后续复核官方原页。",
        ["sohu_2025_shanghai_university_stop_summary"],
        "搜狐转载汇总列出浙江财经大学2025年撤销公共事业管理、数字媒体艺术2个专业。",
        "medium",
    )
    add_warning(
        2024,
        "浙江财经大学",
        "本科",
        "major_warning_list",
        "2023-2024学年本科教学质量报告专业动态调整预警名单",
        "资产评估",
        "列入预警名单。",
        "根据《浙江财经大学本科专业动态调整管理办法》，结合专业准入、准出和专业持续改进压力形成动态调整预警。",
        ["zufe_2024_teaching_quality_report_pdf"],
        "浙江财经大学2023-2024学年本科教学质量报告称，根据《浙江财经大学本科专业动态调整管理办法》，资产评估专业列入预警名单。",
        "high",
    )

    usst_adjustment_rows = [
        (2025, "major_cancel", "2025年本科招生专业设置变化", "网络工程", "080903", "撤销。", "光电信息与计算机工程学院", "上海理工大学官网转载中国科技网文章称2025年撤销网络工程专业；教务处本科专业汇总表列出该专业代码为080903，2019年停招。", "usst_2025_adjustment_article;usst_undergrad_major_setting"),
        (2025, "major_cancel", "2025年本科招生专业设置变化", "假肢矫形工程", "082602T", "撤销。", "健康科学与工程学院", "上海理工大学官网转载中国科技网文章称2025年撤销假肢矫形工程专业；教务处本科专业汇总表列出该专业代码为082602T，2020年停招。", "usst_2025_adjustment_article;usst_undergrad_major_setting"),
        (2025, "major_stop_enrollment", "2025年本科招生专业设置变化", "制药工程", "081302", "停招。", "健康科学与工程学院", "上海理工大学官网转载中国科技网文章称2025年停招制药工程等4个专业；仅结构化正文明确列名的制药工程，其他未列名专业不猜测。", "usst_2025_adjustment_article"),
        (2019, "major_stop_enrollment", "上海理工大学本科专业汇总表停招备注", "网络工程", "080903", "停招。", "光电信息与计算机工程学院", "上海理工大学教务处本科专业汇总表列出网络工程（080903）备注为2019年停招。", "usst_undergrad_major_setting"),
        (2020, "major_stop_enrollment", "上海理工大学本科专业汇总表停招备注", "假肢矫形工程", "082602T", "停招。", "健康科学与工程学院", "上海理工大学教务处本科专业汇总表列出假肢矫形工程（082602T）备注为2020年停招。", "usst_undergrad_major_setting"),
    ]
    for source_row_no, (policy_year, record_type, warning_label, major, code, action, school_unit, evidence, source_ids) in enumerate(usst_adjustment_rows, start=1):
        rows.append(
            {
                "policy_year": policy_year,
                "region": "上海理工大学",
                "education_level": "本科",
                "record_type": record_type,
                "warning_label": warning_label,
                "reported_major_name": major,
                "major_code": code,
                "study_duration": "四年",
                "policy_action": action,
                "criterion_text": "学校官网转载综合教改报道和教务处本科专业汇总表显示专业设置动态调整、停招或撤销信息。",
                "source_row_no": str(source_row_no),
                "source_ids": source_ids,
                "evidence_text": evidence,
                "confidence": "high",
                "notes": school_unit,
            }
        )

    chu_stop_rows = [
        ("3", "物理学（师范）", "070201", "理学", "电子工程学院"),
        ("17", "广告学", "050303", "文学", "文学与传媒学院"),
        ("20", "广播电视学", "050302", "文学", "文学与传媒学院"),
        ("30", "学前教育（师范）", "040106", "教育学", "教师教育学院"),
        ("35", "环境设计", "130503", "艺术学", "美术与设计学院"),
        ("36", "酒店管理", "120902", "管理学", "旅游管理学院"),
        ("39", "会展经济与管理", "120903", "管理学", "旅游管理学院"),
        ("41", "审计学", "120207", "管理学", "工商管理学院"),
        ("46", "互联网金融", "020309T", "经济学", "经济与法学学院"),
    ]
    for source_row_no, major, code, degree, school_unit in chu_stop_rows:
        rows.append(
            {
                "policy_year": 2025,
                "region": "巢湖学院",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2025年本科专业设置及新增专业、停招专业情况一览表",
                "reported_major_name": major,
                "major_code": code,
                "study_duration": "四年",
                "policy_action": "2025年停招。",
                "criterion_text": "巢湖学院2025年本科专业设置及新增专业、停招专业情况一览表在“2025年停招”列标记为停招。",
                "source_row_no": source_row_no,
                "source_ids": "chu_2025_major_setting_notice;chu_2025_major_setting_pdf",
                "evidence_text": f"巢湖学院2025年本科专业设置及新增专业、停招专业情况一览表第{source_row_no}项列出{major}（{code}，四年，{degree}，{school_unit}），2025年停招列标记为停招。",
                "confidence": "high",
            }
        )

    whut_stop_rows = [
        ("3", "2025", "材料物理", "080402"),
        ("4", "2025", "材料化学", "080403"),
        ("14", "2025", "油气储运工程", "081504"),
        ("16", "2025", "物流管理", "120601"),
        ("19", "2025", "港口航道与海岸工程", "081103"),
        ("26", "2025", "过程装备与控制工程", "080206"),
        ("30", "2020", "包装工程", "081702"),
        ("36", "2025", "城乡规划", "082802"),
        ("37", "2021", "人文地理与城乡规划", "070503"),
        ("39", "2025", "环境科学", "082503"),
        ("48", "2024", "信息工程", "080706"),
        ("52", "2024", "物联网工程", "080905"),
        ("56", "2021", "机器人工程", "080803T"),
        ("69", "2024", "生物制药", "083002T"),
        ("75", "2025", "市场营销", "120202"),
        ("76", "2024", "财务管理", "120204"),
        ("80", "2025", "国际经济与贸易", "020401"),
        ("81", "2024", "电子商务", "120801"),
        ("83", "2025", "艺术设计学", "130501"),
        ("90", "2025", "日语", "050207"),
        ("95", "2025", "汉语国际教育", "050103"),
        ("96", "2025", "广告学", "050303"),
        ("97", "2021", "编辑出版学", "050305"),
        ("98", "2020", "教育技术学", "040104"),
    ]
    for source_row_no, stop_year, major, code in whut_stop_rows:
        rows.append(
            {
                "policy_year": int(stop_year),
                "region": "武汉理工大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2025年本科专业目录备注停招专业",
                "reported_major_name": major,
                "major_code": code,
                "study_duration": "",
                "policy_action": f"{stop_year}年停招。",
                "criterion_text": "武汉理工大学2025年本科专业目录图片表格在备注列标注停招年份。",
                "source_row_no": source_row_no,
                "source_ids": "whut_2025_undergrad_major_catalog",
                "evidence_text": f"武汉理工大学2025年本科专业目录图片表格第{source_row_no}项列出{major}（{code}），备注列标注{stop_year}停招。",
                "confidence": "high",
            }
        )

    zsit_stop_rows = [
        ("17", 2025, "跨境电子商务", "120803T", "管理学", "电子商务类", "管理学", "2025年招生情况列标记为停。"),
        ("20", 2025, "电子封装技术", "080709T", "工学", "电子信息类", "工学", "2025年招生情况列标记为停。"),
        ("21", 2022, "机械电子工程", "080204", "工学", "机械类", "工学", "2022年招生情况列标记为停。"),
        ("21", 2023, "机械电子工程", "080204", "工学", "机械类", "工学", "2023年招生情况列标记为停。"),
        ("21", 2024, "机械电子工程", "080204", "工学", "机械类", "工学", "2024年招生情况列标记为停。"),
    ]
    for source_row_no, stop_year, major, code, discipline, major_category, degree, note in zsit_stop_rows:
        rows.append(
            {
                "policy_year": stop_year,
                "region": "绍兴理工学院",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2025年专业设置表历年招生情况停招标记",
                "reported_major_name": major,
                "major_code": code,
                "study_duration": "4",
                "policy_action": f"{stop_year}年停招。",
                "criterion_text": "绍兴理工学院2025年专业设置表在历年招生情况列用“停”标注停招年份。",
                "source_row_no": source_row_no,
                "source_ids": "zsit_2025_major_setting_status",
                "evidence_text": f"绍兴理工学院2025年专业设置表第{source_row_no}项列出{major}（{code}，{discipline}，{major_category}，{degree}，学制4年），{note}",
                "confidence": "high",
            }
        )

    szu_stop_rows = [
        ("20", "050103", "汉语国际教育", "当年普通本科停招，仍招收来华留学生", "当年普通本科停招，仍招收来华留学生。"),
        ("46", "080301", "测控技术与仪器", "当年停招", "当年停招。"),
        ("77", "082803", "风景园林", "当年停招", "当年停招。"),
        ("11", "030302", "社会工作", "往年停招", "往年停招。"),
        ("12", "040101", "教育学", "往年停招", "往年停招。"),
        ("14", "040105", "艺术教育", "往年停招", "往年停招。"),
        ("28", "050304", "传播学", "往年停招", "往年停招。"),
        ("34", "070202", "应用物理学", "往年停招", "往年停招。"),
        ("36", "070302", "应用化学", "往年停招", "往年停招。"),
        ("41", "071102", "应用心理学", "往年停招", "往年停招。"),
        ("44", "080205", "工业设计", "往年停招", "往年停招。"),
        ("45", "080208", "汽车服务工程", "往年停招", "往年停招。"),
        ("50", "080601", "电气工程及其自动化", "往年停招", "往年停招。"),
        ("51", "080603T", "光源与照明", "往年停招", "往年停招。"),
        ("57", "080710T", "集成电路设计与集成系统", "往年停招", "往年停招。"),
        ("62", "080903", "网络工程", "往年停招", "往年停招。"),
        ("63", "080905", "物联网工程", "往年停招", "往年停招。"),
        ("67", "081602", "服装设计与工程", "往年停招", "往年停招。"),
        ("68", "081801", "交通运输", "往年停招", "往年停招。"),
        ("69", "081802", "交通工程", "往年停招", "往年停招。"),
        ("78", "083001", "生物工程", "往年停招", "往年停招。"),
        ("83", "101101", "护理学", "往年停招", "往年停招。"),
        ("91", "120601", "物流管理", "往年停招", "往年停招。"),
        ("98", "130310", "动画", "往年停招", "往年停招。"),
        ("100", "130501", "艺术设计学", "往年停招", "往年停招。"),
        ("102", "130503", "环境设计", "往年停招", "往年停招。"),
    ]
    for source_row_no, code, major, note, action in szu_stop_rows:
        rows.append(
            {
                "policy_year": 2025,
                "region": "深圳大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2025年专业设置、当年新增专业、停招专业名单",
                "reported_major_name": major,
                "major_code": code,
                "policy_action": action,
                "criterion_text": f"深圳大学2025年专业设置Word附件在备注列标注“{note}”。",
                "source_row_no": source_row_no,
                "source_ids": "szu_2025_major_setting_notice;szu_2025_major_setting_docx",
                "evidence_text": f"深圳大学2025年专业设置、当年新增专业、停招专业名单附件第{source_row_no}行列出{major}（{code}），备注为{note}。",
                "confidence": "high",
            }
        )

    jiangnan_stop_rows = [
        ("32", "国际经济与贸易", "020401", "四年", "经济学", "商学院", "2025年停招专业", "2025年停招。"),
        ("40", "汉语言文学", "050101", "四年", "文学", "人文学院", "2025年停招师范班级，保留非师范招生", "2025年停招师范班级，保留非师范招生。"),
    ]
    for source_row_no, major, code, duration, degree, school_unit, note, action in jiangnan_stop_rows:
        rows.append(
            {
                "policy_year": 2025,
                "region": "江南大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2025年本科专业设置、新增专业、停招专业名单",
                "reported_major_name": major,
                "major_code": code,
                "study_duration": duration,
                "policy_action": action,
                "criterion_text": f"江南大学2025年本科专业设置、新增专业、停招专业名单PDF在备注列标注“{note}”。",
                "source_row_no": source_row_no,
                "source_ids": "jiangnan_2025_major_setting_notice;jiangnan_2025_major_setting_pdf",
                "evidence_text": f"江南大学2025年本科专业设置、新增专业、停招专业名单第{source_row_no}项列出{major}（{code}，{duration}，{degree}，{school_unit}），备注为{note}。",
                "confidence": "high",
            }
        )

    upc_stop_rows = [
        ("16", "机械工程", "080201", "工学", "机电工程学院"),
        ("21", "土木工程", "081001", "工学", "储运与建筑工程学院"),
        ("28", "材料物理", "080402", "理学", "材料科学与工程学院"),
        ("29", "材料化学", "080403", "理学", "材料科学与工程学院"),
        ("38", "地理信息科学", "070504", "理学", "海洋与空间信息学院"),
        ("46", "物联网工程", "080905", "工学", "青岛软件学院、计算机科学与技术学院"),
        ("57", "市场营销", "120202", "管理学", "经济管理学院"),
    ]
    for source_row_no, major, code, degree, school_unit in upc_stop_rows:
        rows.append(
            {
                "policy_year": 2025,
                "region": "中国石油大学（华东）",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2025年本科专业设置情况（含新增专业、停招专业）",
                "reported_major_name": major,
                "major_code": code,
                "study_duration": "四年",
                "policy_action": "停招，有在校生。",
                "criterion_text": "中国石油大学（华东）2025年本科专业设置情况PDF在备注列标注“停招，有在校生”。",
                "source_row_no": source_row_no,
                "source_ids": "upc_2025_major_setting_pdf",
                "evidence_text": f"中国石油大学（华东）2025年本科专业设置情况（含新增专业、停招专业）第{source_row_no}项列出{major}（{code}，四年，{degree}，{school_unit}），备注为停招，有在校生。",
                "confidence": "high",
            }
        )

    jci_stop_rows = [
        ("39", 2021, "物流管理", "120601", "管理学", "21停招"),
        ("40", 2021, "公共事业管理", "120401", "管理学", "21停招"),
        ("42", 2022, "金融工程", "020302", "经济学", "22停招"),
        ("52", 2022, "翻译", "050261", "文学", "22停招"),
    ]
    for source_row_no, policy_year, major, code, degree, note in jci_stop_rows:
        rows.append(
            {
                "policy_year": policy_year,
                "region": "景德镇陶瓷大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "本科专业设置一览表停招备注",
                "reported_major_name": major,
                "major_code": code,
                "study_duration": "四年",
                "policy_action": f"{policy_year}年停招。",
                "criterion_text": f"景德镇陶瓷大学本科专业设置一览表Word附件在备注列标注“{note}”。",
                "source_row_no": source_row_no,
                "source_ids": "jci_2024_major_setting_notice;jci_undergrad_major_setting_docx",
                "evidence_text": f"景德镇陶瓷大学本科专业设置一览表第{source_row_no}项列出{major}（{code}，四年，{degree}），备注为{note}。",
                "confidence": "high",
            }
        )

    jci_current_stop_rows = [
        ("28", "30", "自动化", "080801", "工学"),
        ("31", "33", "电子信息工程", "080701", "工学"),
        ("32", "34", "机械电子工程", "080204", "工学"),
        ("40", "42", "物流管理", "120601", "管理学"),
        ("41", "43", "公共事业管理", "120401", "管理学"),
        ("43", "45", "金融工程", "020302", "经济学"),
        ("53", "55", "翻译", "050261", "文学"),
    ]
    for source_seq_no, table_row_no, major, code, degree in jci_current_stop_rows:
        rows.append(
            {
                "policy_year": 2024,
                "region": "景德镇陶瓷大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2023-2024学年本科专业设置信息停招专业",
                "reported_major_name": major,
                "major_code": code,
                "study_duration": "四年",
                "policy_action": "2023-2024学年停招。",
                "criterion_text": "景德镇陶瓷大学2023-2024学年本科专业设置信息正文说明现有招生专业53个、7个专业停招，并逐名列出停招专业。",
                "source_row_no": source_seq_no,
                "source_ids": "jci_2024_academic_year_major_status_notice;jci_2024_academic_year_major_setting_docx",
                "evidence_text": f"景德镇陶瓷大学2023-2024学年本科专业设置信息正文将{major}列入7个停招专业；本科专业设置一览表第{source_seq_no}项（DOCX表格第{table_row_no}行）列出{major}（{code}，四年，{degree}）。",
                "confidence": "high",
            }
        )

    cqnu_2022_stop_rows = [
        ("3", "金融数学", "020305T", "数学科学学院", "经济学"),
        ("9", "艺术教育", "040105", "地理与旅游学院", "艺术学"),
        ("39", "新能源材料与器件", "080414T", "物理与电子工程学院", "工学"),
        ("46", "服装设计与工程", "081602", "美术学院", "工学"),
        ("48", "城乡规划", "082802", "地理与旅游学院", "工学"),
        ("50", "信息管理与信息系统", "120102", "计算机与信息科学学院", "工学"),
        ("68", "摄影", "130404", "新闻与传媒学院", "艺术学"),
    ]
    for source_row_no, major, code, school_unit, degree in cqnu_2022_stop_rows:
        rows.append(
            {
                "policy_year": 2022,
                "region": "重庆师范大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2022-2023学年度专业设置、当年新增专业、停招专业名单",
                "reported_major_name": major,
                "major_code": code,
                "policy_action": "2022年停招。",
                "criterion_text": "重庆师范大学2022-2023学年度专业设置Word附件在“2022年招生情况”列标注“停招”。",
                "source_row_no": source_row_no,
                "source_ids": "cqnu_2023_major_setting_notice;cqnu_2023_major_setting_docx",
                "evidence_text": f"重庆师范大学2022-2023学年度专业设置、当年新增专业、停招专业名单附件第{source_row_no}项列出{major}（{code}，{school_unit}，授予{degree}学位），2022年招生情况为停招。",
                "confidence": "high",
            }
        )

    cqnu_2024_stop_rows = [
        ("3", "金融数学", "020305T", "数学科学学院", ""),
        ("8", "教育技术学", "040104", "教育科学学院", "理学"),
        ("9", "艺术教育", "040105", "地理与旅游学院", "艺术学"),
        ("21", "僧伽罗语", "050219", "外国语学院", ""),
        ("39", "新能源材料与器件", "080414T", "物理与电子工程学院", "工学"),
        ("46", "服装设计与工程", "081602", "美术学院", "工学"),
        ("48", "城乡规划", "082802", "地理与旅游学院", ""),
        ("50", "信息管理与信息系统", "120102", "计算机与信息科学学院", "工学"),
        ("60", "舞蹈表演", "130204", "音乐学院", ""),
        ("68", "摄影", "130404", "新闻与传媒学院", ""),
    ]
    for source_row_no, major, code, school_unit, degree in cqnu_2024_stop_rows:
        degree_phrase = f"，授予{degree}学位" if degree else ""
        rows.append(
            {
                "policy_year": 2024,
                "region": "重庆师范大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2023-2024学年度专业设置、当年新增专业、停招专业名单",
                "reported_major_name": major,
                "major_code": code,
                "policy_action": "2023-2024学年度标注停招。",
                "criterion_text": "重庆师范大学2023-2024学年度专业设置HTML表在招生情况列标注“停招”。",
                "source_row_no": source_row_no,
                "source_ids": "cqnu_2024_major_setting_html",
                "evidence_text": f"重庆师范大学2023-2024学年度专业设置、当年新增专业、停招专业名单HTML表第{source_row_no}项列出{major}（{code}，{school_unit}{degree_phrase}），招生情况标注为停招。",
                "confidence": "high",
            }
        )

    scnu_stop_rows = [
        ("", "小学教育"),
        ("", "特殊教育"),
        ("11", "翻译"),
        ("14", "视觉传达设计"),
        ("21", "金融数学"),
        ("38", "会展经济与管理"),
        ("39", "酒店管理"),
        ("41", "管理科学"),
        ("47", "电子商务"),
        ("51", "国际经济与贸易"),
        ("53", "金融工程"),
        ("59", "编辑出版学"),
        ("60", "汉语言"),
        ("63", "音乐表演"),
        ("65", "信息工程"),
        ("", "财务管理"),
        ("77", "文化产业管理"),
        ("78", "网络与新媒体"),
        ("", "法语"),
        ("82", "通信工程"),
        ("83", "物联网工程"),
        ("", "科学教育"),
    ]
    for source_row_no, major in scnu_stop_rows:
        rows.append(
            {
                "policy_year": 2025,
                "region": "华南师范大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2025本科专业设置、新增专业、停招专业一览表",
                "reported_major_name": major,
                "policy_action": "停招。",
                "criterion_text": "华南师范大学2025本科专业设置、新增专业、停招专业一览表的招生情况列标记为停招；未招专业未纳入本组记录。",
                "source_row_no": source_row_no,
                "source_ids": "scnu_2025_stop_enrollment_pdf",
                "evidence_text": f"华南师范大学2025本科专业设置、新增专业、停招专业一览表列出{major}，招生情况为停招。",
                "confidence": "high",
            }
        )

    sufe_stop_rows = [
        ("6", "金融工程", "020302"),
        ("9", "信用管理", "020306T"),
        ("13", "工程管理", "120103"),
        ("19", "国际商务", "120205"),
        ("20", "人力资源管理", "120206"),
        ("21", "公共事业管理", "120401"),
        ("25", "物流管理", "120601"),
        ("29", "社会学", "030301"),
        ("32", "英语", "050201"),
        ("35", "新闻学", "050301"),
        ("39", "应用统计学", "071202"),
        ("40", "数据科学与大数据技术（理学）", "080910T"),
    ]
    for source_row_no, major, code in sufe_stop_rows:
        rows.append(
            {
                "policy_year": 2025,
                "region": "上海财经大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2025本科专业设置情况已停招专业",
                "reported_major_name": major,
                "major_code": code,
                "policy_action": "已停招。",
                "criterion_text": "上海财经大学2025本科专业设置PDF注释说明标“◆”为已停招专业。",
                "source_row_no": source_row_no,
                "source_ids": "sufe_2025_stop_enrollment_pdf",
                "evidence_text": f"上海财经大学2025本科专业设置PDF第{source_row_no}项列出{major}（{code}）并以◆标注为已停招专业。",
                "confidence": "high",
            }
        )

    bjtu_stop_rows = [
        ("66", "思想政治教育", "030503", "四年", "法学"),
        ("67", "汉语言", "050102", "四年", "文学"),
        ("68", "给排水科学与工程", "081003", "四年", "工学"),
        ("69", "材料化学", "080403", "四年", "理学"),
        ("70", "电子信息工程", "080701", "四年", "工学"),
        ("71", "生物信息学", "071003", "四年", "理学"),
    ]
    for source_row_no, major, code, duration, degree in bjtu_stop_rows:
        rows.append(
            {
                "policy_year": 2025,
                "region": "北京交通大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2025专业设置停招专业名单",
                "reported_major_name": major,
                "major_code": code,
                "study_duration": duration,
                "policy_action": "停招。",
                "criterion_text": "北京交通大学信息公开页本科专业表备注列标记为停招。",
                "source_row_no": source_row_no,
                "source_ids": "bjtu_2025_stop_enrollment",
                "evidence_text": f"北京交通大学专业设置、当年新增专业、停招专业名单第{source_row_no}项列出{major}（{code}，{duration}，{degree}），备注为停招。",
                "confidence": "high",
            }
        )

    for source_row_no, major, code, duration, degree in bjtu_stop_rows:
        rows.append(
            {
                "policy_year": 2025,
                "region": "北京交通大学",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "2025年度专业设置与调整申请撤销专业材料",
                "reported_major_name": major,
                "major_code": code,
                "study_duration": duration,
                "policy_action": "申请撤销。",
                "criterion_text": "北京交通大学2025年度专业设置与调整申报材料zip附件设有申请撤销专业材料目录，逐项提供撤销材料文件。",
                "source_row_no": str(source_row_no),
                "source_ids": "bjtu_2025_major_adjustment_notice;bjtu_2025_major_adjustment_materials_zip",
                "evidence_text": f"北京交通大学2025年度专业设置与调整申报材料zip附件“申请撤销专业材料”目录列出{major}撤销材料；该专业同时见信息公开本科专业表第{source_row_no}项（{code}，{duration}，{degree}）并标注停招。",
                "confidence": "high",
            }
        )

    jnu_stop_rows = [
        ("86", "音乐学", "130202", "四年", "艺术学"),
        ("87", "信息工程", "080706", "四年", "工学"),
        ("88", "电子信息科学与技术", "080714T", "四年", "工学"),
        ("89", "智能科学与技术", "080907T", "四年", "工学"),
        ("90", "通信工程", "080703", "四年", "工学"),
        ("91", "物流管理", "120601", "四年", "管理学"),
        ("92", "食品质量与安全", "082702", "四年", "工学"),
        ("93", "化学工程与工艺", "081301", "四年", "工学"),
        ("94", "高分子材料与工程", "080407", "四年", "工学"),
        ("95", "材料物理", "080402", "四年", "工学"),
        ("96", "给排水科学与工程", "081003", "四年", "工学"),
        ("97", "美术学", "130401", "四年", "艺术学"),
        ("98", "人力资源管理", "120206", "四年", "管理学"),
        ("99", "环境科学", "082503", "四年", "工学"),
        ("100", "环境工程", "082502", "四年", "工学"),
        ("101", "网络工程", "080903", "四年", "工学"),
        ("102", "公共事业管理", "120401", "四年", "管理学"),
        ("103", "会展经济与管理", "120903", "四年", "管理学"),
        ("104", "信息安全", "080904K", "四年", "工学"),
        ("105", "临床药学", "100703TK", "五年", "理学"),
        ("106", "财务管理", "120204", "四年", "管理学"),
        ("107", "风景园林", "082803", "四年", "工学"),
        ("108", "建筑学", "082801", "五年", "工学"),
    ]
    for source_row_no, major, code, duration, degree in jnu_stop_rows:
        rows.append(
            {
                "policy_year": 2025,
                "region": "暨南大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2025年专业设置一览表停招专业",
                "reported_major_name": major,
                "major_code": code,
                "study_duration": duration,
                "policy_action": "停招。",
                "criterion_text": "暨南大学2025年专业设置一览表备注列标记为停招专业。",
                "source_row_no": source_row_no,
                "source_ids": "jnu_2025_stop_enrollment",
                "evidence_text": f"暨南大学2025年专业设置一览表第{source_row_no}项列出{major}（{code}，{duration}，{degree}），备注为停招专业。",
                "confidence": "high",
            }
        )

    njtech_stop_rows = [
        ("7", "环境科学", "082503", "四年", "工学（授理学）"),
        ("8", "资源环境科学", "082506T", "四年", "工学"),
        ("11", "冶金工程", "080404", "四年", "工学"),
        ("32", "焊接技术与工程", "080411T", "四年", "工学"),
        ("36", "能源与环境系统工程", "080502T", "四年", "工学"),
        ("42", "风景园林", "082803", "四年", "工学"),
        ("44", "工业设计", "080205", "四年", "工学"),
        ("48", "数字媒体艺术", "130508", "四年", "艺术学"),
        ("56", "市场营销", "120202", "四年", "管理学"),
        ("58", "人力资源管理", "120206", "四年", "管理学"),
        ("62", "社会工作", "030302", "四年", "法学"),
        ("63", "公共事业管理", "120401", "四年", "管理学"),
        ("67", "德语", "050203", "四年", "文学"),
        ("74", "食品质量与安全", "082702", "四年", "工学"),
        ("80", "信息与计算科学", "070102", "四年", "理学"),
        ("85", "地理信息科学", "070504", "四年", "理学"),
        ("92", "铁道工程", "081007T", "四年", "工学"),
    ]
    for source_row_no, major, code, duration, degree in njtech_stop_rows:
        rows.append(
            {
                "policy_year": 2025,
                "region": "南京工业大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2025年本科专业设置及新增、停招情况",
                "reported_major_name": major,
                "major_code": code,
                "study_duration": duration,
                "policy_action": "停招。",
                "criterion_text": "南京工业大学2025年本科专业设置及新增、停招情况表备注列标记为2025年停招。",
                "source_row_no": source_row_no,
                "source_ids": "njtech_2025_stop_enrollment",
                "evidence_text": f"南京工业大学2025年本科专业设置及新增、停招情况表第{source_row_no}项列出{major}（{code}，{duration}，{degree}），备注为2025年停招。",
                "confidence": "high",
            }
        )

    nuist_stop_rows = [
        (2024, "nuist_2024_stop_enrollment_image", "2024年停招本科专业列表", "停招。", "南京信息工程大学2024年停招本科专业列表图片列出停招专业。", [
            ("1", "水利科学与工程", "081105T"),
            ("2", "市场营销", "120202"),
            ("3", "公共事业管理", "120401"),
        ]),
        (2025, "nuist_2025_stop_enrollment_image", "2025年停招本科专业列表", "停招。", "南京信息工程大学2025年停招本科专业列表图片列出停招专业。", [
            ("1", "水利科学与工程", "081105T"),
            ("2", "市场营销", "120202"),
            ("3", "公共事业管理", "120401"),
            ("4", "保险学", "020303"),
            ("5", "翻译", "050261"),
        ]),
    ]
    for policy_year, source_id, warning_label, action, criterion, major_rows in nuist_stop_rows:
        for source_row_no, major, code in major_rows:
            rows.append(
                {
                    "policy_year": policy_year,
                    "region": "南京信息工程大学",
                    "education_level": "本科",
                    "record_type": "major_stop_enrollment",
                    "warning_label": warning_label,
                    "reported_major_name": major,
                    "major_code": code,
                    "policy_action": action,
                    "criterion_text": criterion,
                    "source_row_no": source_row_no,
                    "source_ids": source_id,
                    "evidence_text": f"南京信息工程大学{warning_label}图片第{source_row_no}行列出{major}（{code}）。",
                    "confidence": "high",
                }
            )

    nuist_cancel_rows = [
        ("1", "统计学", "071201", "2001", "2013年停招，2024年撤销"),
        ("2", "轨道交通信号与控制", "080802T", "2003", "未招生，2024年撤销"),
    ]
    for source_row_no, major, code, setup_year, note in nuist_cancel_rows:
        rows.append(
            {
                "policy_year": 2024,
                "region": "南京信息工程大学",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "2024年撤销本科专业列表",
                "reported_major_name": major,
                "major_code": code,
                "policy_action": "撤销。",
                "criterion_text": f"南京信息工程大学2024年撤销本科专业列表图片备注为“{note}”。",
                "source_row_no": source_row_no,
                "source_ids": "nuist_2024_major_cancellation_image",
                "evidence_text": f"南京信息工程大学2024年撤销本科专业列表图片第{source_row_no}行列出{major}（{code}），设置年份{setup_year}，备注为{note}。",
                "confidence": "high",
            }
        )

    hbue_paused_rows = [
        ("1", "国际商务", "120205", "已停招。"),
        ("2", "信用管理", "020306T", "已停招。"),
        ("3", "物流工程", "120602", "已停招。"),
        ("4", "工程造价", "120105", "已停招。"),
        ("5", "资产评估", "120208", "已停招。"),
        ("6", "应用统计学", "071202", "已停招。"),
        ("7", "社会工作", "030302", "已停招。"),
        ("8", "艺术设计学", "130501", "已停招。"),
        ("9", "产品设计", "130504", "已停招。"),
        ("10", "广告学", "050303", "已停招。"),
        ("11", "数字媒体技术", "080906", "已停招。"),
        ("12", "体育经济与管理", "120212T", "已停招。"),
        ("13", "行政管理", "120402", "当年停招。"),
        ("14", "劳动与社会保障", "120403", "当年停招。"),
        ("15", "会展经济与管理", "120903", "当年停招。"),
    ]
    for source_row_no, major, code, action in hbue_paused_rows:
        rows.append(
            {
                "policy_year": 2024,
                "region": "湖北经济学院",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2023-2024学年本科教学质量报告暂停招生专业",
                "reported_major_name": major,
                "major_code": code,
                "policy_action": action,
                "criterion_text": "湖北经济学院2023-2024学年本科教学质量报告说明2024年15个专业暂停招生，其中行政管理、劳动与社会保障、会展经济与管理为本年度停招。",
                "source_row_no": source_row_no,
                "source_ids": "hbue_2024_teaching_quality_report_pdf",
                "evidence_text": f"湖北经济学院2023-2024学年本科教学质量报告列出{major}（{code}）为暂停招生专业；状态为{action}",
                "confidence": "high",
            }
        )

    hbue_cancel_rows = [
        ("1", "信用管理", "020306T"),
        ("2", "物流工程", "120602"),
        ("3", "资产评估", "120208"),
        ("4", "社会工作", "030302"),
        ("5", "艺术设计学", "130501"),
        ("6", "广告学", "050303"),
        ("7", "应用统计学", "071202"),
    ]
    for source_row_no, major, code in hbue_cancel_rows:
        rows.append(
            {
                "policy_year": 2024,
                "region": "湖北经济学院",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "2023-2024学年本科教学质量报告申请撤销专业",
                "reported_major_name": major,
                "major_code": code,
                "policy_action": "申请撤销。",
                "criterion_text": "湖北经济学院2023-2024学年本科教学质量报告说明申请撤销7个专业。",
                "source_row_no": source_row_no,
                "source_ids": "hbue_2024_teaching_quality_report_pdf",
                "evidence_text": f"湖北经济学院2023-2024学年本科教学质量报告列出申请撤销{major}（{code}）。",
                "confidence": "high",
            }
        )

    wit_stop_rows = [
        ("5", "工程力学", "080102", "四年", "工学", "停招"),
        ("7", "材料成型及控制工程", "080203", "四年", "工学", "当年停招"),
        ("19", "道路桥梁与渡河工程", "081006T", "四年", "工学", "停招"),
        ("23", "城市地下空间工程", "081005T", "四年", "工学", "停招"),
        ("32", "市场营销", "120202", "四年", "管理学", "停招"),
        ("34", "财务管理", "120204", "四年", "管理学", "停招"),
        ("35", "公共事业管理", "120401", "四年", "管理学", "停招"),
        ("47", "国际经济与贸易", "020401", "四年", "经济学", "当年停招"),
        ("59", "广告学", "050303", "四年", "文学", "停招"),
    ]
    for source_row_no, major, code, duration, degree, note in wit_stop_rows:
        rows.append(
            {
                "policy_year": 2025,
                "region": "武汉工程大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "武汉工程大学本科专业目录（2025年）停招专业",
                "reported_major_name": major,
                "major_code": code,
                "study_duration": duration,
                "policy_action": "停招。",
                "criterion_text": f"武汉工程大学本科专业目录（2025年）在专业名称中标注“{note}”。",
                "source_row_no": source_row_no,
                "source_ids": "wit_2025_stop_enrollment",
                "evidence_text": f"武汉工程大学本科专业目录（2025年）第{source_row_no}项列出{major}（{code}，{duration}，{degree}），专业名称标注{note}。",
                "confidence": "high",
            }
        )

    ujs_status_rows = [
        ("4", 2025, "高分子材料与工程", "080407", "四年", "工学", "2025年停招", "停招。"),
        ("8", 2025, "财务管理", "120204", "四年", "管理学", "2025年停招", "停招。"),
        ("10", 2024, "保险学", "020303", "四年", "经济学", "2024年停招", "停招。"),
        ("17", 2024, "生物医学工程", "082601", "四年", "工学", "2024年停招", "停招。"),
        ("22", 2025, "市场营销", "120202", "四年", "管理学", "2025年停招", "停招。"),
        ("27", 2024, "电子商务", "120801", "四年", "管理学", "2024年停招", "停招。"),
        ("28", 2025, "海洋资源开发技术", "081903T", "四年", "工学", "暂缓招生", "暂缓招生。"),
        ("36", 2025, "化工安全工程", "081306T", "四年", "工学", "暂缓招生", "暂缓招生。"),
        ("48", 2025, "教育技术学", "040104", "四年", "教育学", "2025年停招", "停招。"),
        ("52", 2024, "建筑环境与能源应用工程", "081002", "四年", "工学", "2024年停招", "停招。"),
        ("55", 2024, "设施农业科学与工程", "090106", "四年", "农学", "2024年停招", "停招。"),
        ("61", 2025, "智慧交通", "081811T", "四年", "工学", "新增、暂缓招生", "新增后暂缓招生。"),
        ("68", 2024, "金融数学", "020305T", "四年", "经济学", "2024年停招", "停招。"),
        ("71", 2025, "工程管理", "120103", "四年", "工学", "2025年停招", "停招。"),
        ("76", 2025, "语言学", "0502100T", "四年", "文学", "暂缓招生", "暂缓招生。"),
        ("88", 2024, "护理学", "101101K", "四年", "理学", "2024年停招", "停招。"),
        ("91", 2024, "动画", "130310", "四年", "艺术学", "2024年停招", "停招。"),
    ]
    for source_row_no, policy_year, major, code, duration, degree, note, action in ujs_status_rows:
        rows.append(
            {
                "policy_year": policy_year,
                "region": "江苏大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2025本科专业一览表停招/暂缓招生备注",
                "reported_major_name": major,
                "major_code": code,
                "study_duration": duration,
                "policy_action": action,
                "criterion_text": f"江苏大学现有本科专业一览表（2025.11）备注列标记“{note}”。",
                "source_row_no": source_row_no,
                "source_ids": "ujs_2025_stop_enrollment",
                "evidence_text": f"江苏大学现有本科专业一览表（2025.11）第{source_row_no}项列出{major}（{code}，{duration}，{degree}），备注为{note}。",
                "confidence": "high",
            }
        )

    add_warning_list(
        2025,
        "广东财经大学",
        "本科",
        "major_stop_enrollment",
        "2025年暂停招生本科专业",
        ["房地产开发与管理", "公共事业管理", "信息管理与信息系统", "戏剧影视文学"],
        "暂停招生。",
        "广东财经大学2025年专业设置、当年新增专业、停招专业名单列出暂停招生本科专业。",
        ["gdufe_2025_stop_cancel"],
        "广东财经大学2025年专业设置、当年新增专业、停招专业名单列出暂停招生本科专业4个：房地产开发与管理、公共事业管理、信息管理与信息系统、戏剧影视文学。",
        "high",
    )
    add_warning_list(
        2025,
        "广东财经大学",
        "本科",
        "major_cancel",
        "2025年撤销本科专业",
        ["应用心理学", "编辑出版学"],
        "撤销。",
        "广东财经大学2025年专业设置、当年新增专业、停招专业名单列出撤销本科专业。",
        ["gdufe_2025_stop_cancel"],
        "广东财经大学2025年专业设置、当年新增专业、停招专业名单列出撤销本科专业2个：应用心理学、编辑出版学。",
        "high",
    )
    add_warning_list(
        2025,
        "东北林业大学",
        "本科",
        "major_stop_enrollment",
        "2025年停招专业名单",
        ["汽车服务工程", "工业工程", "森林工程", "包装工程", "物流工程", "公共事业管理", "旅游管理"],
        "停招。",
        "东北林业大学2025年专业设置，当年新增专业、停招专业名单正文列出停招专业。",
        ["nefu_2025_stop_enrollment"],
        "东北林业大学2025年专业设置，当年新增专业、停招专业名单正文列出停招专业：汽车服务工程、工业工程、森林工程、包装工程、物流工程、公共事业管理、旅游管理。",
        "high",
    )
    add_warning(
        2025,
        "武汉轻工大学",
        "本科",
        "major_stop_enrollment",
        "2025年学校停招专业",
        "给排水科学与工程",
        "停招。",
        "武汉轻工大学2025年学校停招专业页列出学校停招专业。",
        ["whpu_2025_stop_enrollment"],
        "武汉轻工大学2025年学校停招专业页列出2025年学校停招专业1个：给排水科学与工程。",
        "high",
    )

    hgnu_2025_rows = [
        ("1", "文学院（苏东坡书院）", "秘书学", "050107T", "文学", "停招", ""),
        ("2", "商学院", "电子商务", "120801", "管理学", "停招", ""),
        ("3", "物理与电信学院", "电子信息科学与技术", "080714T", "工学", "停招", ""),
        ("4", "化学化工学院", "应用化学", "070302", "理学", "停招", ""),
        ("5", "计算机学院", "信息管理与信息系统", "120102", "工学", "停招", ""),
        ("6", "机电与智能制造学院", "车辆工程", "080207", "工学", "停招", ""),
        ("7", "商学院", "市场营销", "120202", "管理学", "撤销", "2020年停招"),
        ("8", "计算机学院", "物联网工程", "080905", "工学", "撤销", "2020年停招"),
        ("9", "机电与智能制造学院", "汽车服务工程", "080208", "工学", "撤销", "2020年停招"),
    ]
    for source_row_no, school_unit, major, code, degree, adjust_type, note in hgnu_2025_rows:
        is_cancel = adjust_type == "撤销"
        action = "撤销，备注2020年停招。" if is_cancel else "停招。"
        criterion = "黄冈师范学院2025年停招专业XLS在调整类型列标记“撤销”，备注为“2020年停招”。" if is_cancel else "黄冈师范学院2025年停招专业XLS在调整类型列标记“停招”。"
        rows.append(
            {
                "policy_year": 2025,
                "region": "黄冈师范学院",
                "education_level": "本科",
                "record_type": "major_cancel" if is_cancel else "major_stop_enrollment",
                "warning_label": "黄冈师范学院停招专业（2025）",
                "reported_major_name": major,
                "major_code": code,
                "policy_action": action,
                "criterion_text": criterion,
                "source_row_no": source_row_no,
                "source_ids": "hgnu_2025_stop_enrollment_notice;hgnu_2025_stop_enrollment_xls",
                "evidence_text": f"黄冈师范学院停招专业（2025）XLS第{source_row_no}项列出{major}（{code}，{school_unit}，授予{degree}学位），调整类型为{adjust_type}{('，备注为' + note) if note else ''}。",
                "confidence": "high",
            }
        )

    hgnu_2022_rows = [
        ("1", "化学化工学院", "工业工程", "120701", "工学", "2012", "教高[2012]2号"),
        ("2", "教育学院", "数字媒体技术", "080906", "工学", "2009", "鄂教高[2009]3号"),
        ("3", "地理与旅游学院", "酒店管理", "120902", "管理学", "2014", "教高[2014]1号"),
        ("4", "商学院", "市场营销", "120202", "管理学", "2008", "鄂教高[2008]11号"),
        ("5", "机电与汽车工程学院", "汽车服务工程", "080208", "工学", "2015", "教高函[2015]2号"),
        ("6", "计算机学院", "物联网工程", "080905", "工学", "2017", "教高[2017]2号"),
    ]
    for source_row_no, school_unit, major, code, degree, start_year, approval_no in hgnu_2022_rows:
        rows.append(
            {
                "policy_year": 2022,
                "region": "黄冈师范学院",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "黄冈师范学院停招专业（2022）",
                "reported_major_name": major,
                "major_code": code,
                "policy_action": "停招。",
                "criterion_text": "黄冈师范学院2022年停招专业HTML表在备注列标记“停招”。",
                "source_row_no": source_row_no,
                "source_ids": "hgnu_2022_stop_enrollment",
                "evidence_text": f"黄冈师范学院停招专业（2022）HTML表第{source_row_no}项列出{major}（{code}，{school_unit}，授予{degree}学位，开始招生时间{start_year}，批准文号{approval_no}），备注为停招。",
                "confidence": "high",
            }
        )

    huat_stop_rows = [
        ("6", "汽车服务工程", "080208"),
        ("7", "焊接技术与工程", "080411T"),
    ]
    for source_row_no, major, code in huat_stop_rows:
        rows.append(
            {
                "policy_year": 2025,
                "region": "湖北汽车工业学院",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2025年新增专业、停招专业名单",
                "reported_major_name": major,
                "major_code": code,
                "policy_action": "停招。",
                "criterion_text": "湖北汽车工业学院2025年新增专业、停招专业名单附件在备注列标记“2025年停招专业”，该段包含汽车服务工程、焊接技术与工程。",
                "source_row_no": source_row_no,
                "source_ids": "huat_2025_stop_enrollment_xlsx",
                "evidence_text": f"湖北汽车工业学院2025年新增专业、停招专业名单附件第{source_row_no}行列出{major}（{code}），归入2025年停招专业。",
                "confidence": "high",
            }
        )

    add_warning(
        2025,
        "西安欧亚学院",
        "本科",
        "major_stop_enrollment",
        "2025年专业设置、当年新增专业、停招专业名单",
        "软件工程（合作办学）",
        "停招。",
        "西安欧亚学院信息公开页正文明确“软件工程（合作办学）停招”。",
        ["eurasia_2025_stop_enrollment"],
        "西安欧亚学院2025年专业设置、当年新增专业、停招专业名单页写明：截至2025年9月，软件工程（合作办学）停招。",
        "high",
    )

    add_warning_list(
        2023,
        "中南大学",
        "本科",
        "major_stop_enrollment",
        "2023年停招专业名单",
        ["视觉传达设计", "环境设计"],
        "停止招生。",
        "中南大学信息公开页正文列出2023年停止招生的本科专业。",
        ["csu_2023_stop_enrollment"],
        "中南大学2023年停招专业名单页列出2023年停止招生的本科专业为视觉传达设计、环境设计。",
        "high",
    )
    add_warning_list(
        2024,
        "中南大学",
        "本科",
        "major_stop_enrollment",
        "2024年停招专业名单",
        ["产品设计", "生物技术"],
        "停止招生。",
        "中南大学信息公开页正文列出2024年停止招生的本科专业。",
        ["csu_2024_stop_enrollment"],
        "中南大学2024年停招专业名单页列出2024年停止招生的本科专业为产品设计、生物技术。",
        "high",
    )

    csu_2025_stop_rows = [
        ("90", "资源加工与生物工程学院", "071002", "生物技术", "理学"),
        ("91", "自动化学院", "080907T", "智能科学与技术", "工学"),
        ("92", "计算机学院", "080905", "物联网工程", "工学"),
        ("93", "粉末冶金研究院", "080403", "材料化学", "工学"),
        ("94", "粉末冶金研究院", "080407", "高分子材料与工程", "工学"),
        ("95", "自动化学院", "082103", "探测制导与控制技术", "工学"),
        ("96", "交通运输工程学院", "120602", "物流工程", "工学"),
        ("97", "人文学院", "050302", "广播电视学", "文学"),
        ("98", "商学院", "120202", "市场营销", "管理学"),
        ("99", "商学院", "120204", "财务管理", "管理学"),
        ("100", "商学院", "120801", "电子商务", "管理学"),
        ("101", "建筑与艺术学院", "130502", "视觉传达设计", "艺术学"),
        ("102", "建筑与艺术学院", "130503", "环境设计", "艺术学"),
        ("103", "建筑与艺术学院", "130504", "产品设计", "艺术学"),
        ("104", "公共管理学院", "120403", "劳动与社会保障", "管理学"),
    ]
    for source_row_no, school_unit, code, major, degree in csu_2025_stop_rows:
        rows.append(
            {
                "policy_year": 2025,
                "region": "中南大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2025年本科专业设置表停止招生专业",
                "reported_major_name": major,
                "major_code": code,
                "study_duration": "四年",
                "policy_action": "停止招生。",
                "criterion_text": "中南大学2025年本科专业设置表PDF在“停止招生”列标记为“是”。",
                "source_row_no": source_row_no,
                "source_ids": "csu_2025_undergrad_setting;csu_2025_undergrad_setting_pdf",
                "evidence_text": f"中南大学2025年本科专业设置表PDF第{source_row_no}项列出{major}（{code}，{school_unit}，{degree}，四年），停止招生列为是。",
                "confidence": "high",
            }
        )

    imnc_2025_stop_rows = [
        ("信息管理与信息系统", "120102"),
        ("视觉传达设计", "130502"),
        ("行政管理", "120402"),
        ("市场营销", "120202"),
    ]
    for major, code in imnc_2025_stop_rows:
        rows.append(
            {
                "policy_year": 2025,
                "region": "呼和浩特民族学院",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2024-2025学年本科教学质量报告停招专业",
                "reported_major_name": major,
                "major_code": code,
                "policy_action": "2025年停招，并申请撤销。",
                "criterion_text": "呼和浩特民族学院2024-2025学年本科教学质量报告在本科专业设置情况中明确说明2025年停招4个专业，后文说明申请撤销这4个本科专业。",
                "source_row_no": "2",
                "source_ids": "imnc_2025_teaching_quality_report_pdf",
                "evidence_text": f"呼和浩特民族学院2024-2025学年本科教学质量报告第2页说明2025年停招信息管理与信息系统、视觉传达设计、行政管理、市场营销4个专业；后文说明申请撤销{major}等4个本科专业。",
                "confidence": "high",
            }
        )

    tjfsu_bhws_2025_stop_rows = [
        ("国际事务与国际关系", "030204T"),
        ("新闻学", "050301"),
        ("电子商务", "120801"),
        ("财务管理", "120204"),
        ("广告学", "050303"),
    ]
    for major, code in tjfsu_bhws_2025_stop_rows:
        rows.append(
            {
                "policy_year": 2025,
                "region": "天津外国语大学滨海外事学院",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2024-2025学年本科教学质量报告停招专业",
                "reported_major_name": major,
                "major_code": code,
                "policy_action": "本学年停招。",
                "criterion_text": "天津外国语大学滨海外事学院2024-2025学年本科教学质量报告“2.专业设置”明确列出本学年停招5个本科专业。",
                "source_row_no": "2",
                "source_ids": "tjfsu_bhws_2025_teaching_quality_report",
                "evidence_text": f"天津外国语大学滨海外事学院2024-2025学年本科教学质量报告“2.专业设置”说明本学年停招专业为国际事务与国际关系、新闻学、电子商务、财务管理和广告学；本行对应{major}（{code}）。",
                "confidence": "high",
            }
        )

    sxufe_2025_stop_rows = [
        ("1", "050103", "汉语国际教育", "中国语言文学类", "四年", "文学", "2018"),
        ("2", "080902", "软件工程", "计算机类", "四年", "工学", "2022"),
        ("3", "120104", "房地产开发与管理", "管理科学与工程类", "四年", "管理学", "2011"),
    ]
    for source_row_no, code, major, major_class, duration, degree, opened_year in sxufe_2025_stop_rows:
        rows.append(
            {
                "policy_year": 2025,
                "region": "山西财经大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2025年停招本科专业名单",
                "reported_major_name": major,
                "major_code": code,
                "study_duration": duration,
                "policy_action": "停招。",
                "criterion_text": "山西财经大学教务部2025年停招本科专业名单表格逐项列出停招本科专业。",
                "source_row_no": source_row_no,
                "source_ids": "sxufe_2025_stop_enrollment",
                "evidence_text": f"山西财经大学2025年停招本科专业名单第{source_row_no}行列出{major}（{code}，{major_class}，{duration}，{degree}，开设年份{opened_year}）。",
                "confidence": "high",
            }
        )

    hbmzu_stop_rows = [
        (2025, "hbmzu_2025_stop_enrollment", "停招专业（2025年）", "1", "医学影像技术", "101003"),
        (2025, "hbmzu_2025_stop_enrollment", "停招专业（2025年）", "2", "广播电视编导", "130305"),
        (2024, "hbmzu_2024_stop_enrollment", "停招专业（2024年）", "1", "编辑出版学", "050305"),
        (2022, "hbmzu_2022_stop_enrollment", "停招专业（2022年）", "1", "应用化学", "070302"),
        (2022, "hbmzu_2022_stop_enrollment", "停招专业（2022年）", "2", "翻译", "050261"),
        (2022, "hbmzu_2022_stop_enrollment", "停招专业（2022年）", "3", "人文地理与城乡规划", "070503"),
    ]
    for policy_year, source_id, page_title, source_row_no, major, code in hbmzu_stop_rows:
        rows.append(
            {
                "policy_year": policy_year,
                "region": "湖北民族大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": page_title,
                "reported_major_name": major,
                "major_code": code,
                "policy_action": "停招。",
                "criterion_text": f"湖北民族大学信息公开网“{page_title}”页面列出停止招生专业。",
                "source_row_no": source_row_no,
                "source_ids": source_id,
                "evidence_text": f"湖北民族大学信息公开网“{page_title}”页面第{source_row_no}项列出{major}。",
                "confidence": "high",
            }
        )

    hbxytc_vocational_stop_rows = [
        (2025, "hbxytc_2025_vocational_stop_pdf", "2025年停招专业", "1", "530810", "供应链运营", "三年"),
        (2025, "hbxytc_2025_vocational_stop_pdf", "2025年停招专业", "2", "520202", "助产", "三年"),
        (2025, "hbxytc_2025_vocational_stop_pdf", "2025年停招专业", "3", "510201", "计算机应用技术", "二年"),
        (2025, "hbxytc_2025_vocational_stop_pdf", "2025年停招专业", "4", "510102", "物联网应用技术", "二年"),
        (2025, "hbxytc_2025_vocational_stop_pdf", "2025年停招专业", "5", "540106", "酒店管理与数字化运营", "二年"),
        (2025, "hbxytc_2025_vocational_stop_pdf", "2025年停招专业", "6", "530302", "大数据与会计", "二年"),
        (2025, "hbxytc_2025_vocational_stop_pdf", "2025年停招专业", "7", "530701", "电子商务", "二年"),
        (2025, "hbxytc_2025_vocational_stop_pdf", "2025年停招专业", "8", "550105", "服装与服饰设计", "二年"),
        (2024, "hbxytc_2024_vocational_stop_pdf", "2024年停招专业名单", "1", "510206", "云计算技术应用", "三年"),
        (2024, "hbxytc_2024_vocational_stop_pdf", "2024年停招专业名单", "2", "530301", "大数据与财务管理", "三年"),
        (2024, "hbxytc_2024_vocational_stop_pdf", "2024年停招专业名单", "3", "460103", "数控技术", "二年"),
        (2024, "hbxytc_2024_vocational_stop_pdf", "2024年停招专业名单", "4", "430705", "装配式建筑构件智能制造技术", "三年"),
        (2024, "hbxytc_2024_vocational_stop_pdf", "2024年停招专业名单", "5", "520202", "助产", "三年"),
        (2024, "hbxytc_2024_vocational_stop_pdf", "2024年停招专业名单", "6", "530802", "现代物流管理", "三年"),
        (2023, "hbxytc_2023_vocational_stop_pdf", "2023年停招专业", "1", "520301", "药学", "三年"),
        (2023, "hbxytc_2023_vocational_stop_pdf", "2023年停招专业", "2", "460103", "数控技术", "三年"),
        (2022, "hbxytc_2022_vocational_stop_pdf", "2022年停招专业", "1", "410301", "动物医学", "三年"),
        (2022, "hbxytc_2022_vocational_stop_pdf", "2022年停招专业", "2", "430705", "装配式建筑构件智能制造技术", "三年"),
        (2022, "hbxytc_2022_vocational_stop_pdf", "2022年停招专业", "3", "460405", "轨道交通工程机械制造与维护", "三年"),
        (2022, "hbxytc_2022_vocational_stop_pdf", "2022年停招专业", "4", "510206", "云计算技术应用", "三年"),
        (2022, "hbxytc_2022_vocational_stop_pdf", "2022年停招专业", "5", "520202", "助产", "三年"),
        (2022, "hbxytc_2022_vocational_stop_pdf", "2022年停招专业", "6", "530301", "大数据与财务管理", "三年"),
        (2021, "hbxytc_2021_vocational_stop_pdf", "2021年停招专业", "1", "520504", "口腔医学技术", "三年"),
        (2021, "hbxytc_2021_vocational_stop_pdf", "2021年停招专业", "2", "520901", "眼视光技术", "三年"),
        (2021, "hbxytc_2021_vocational_stop_pdf", "2021年停招专业", "3", "510103", "应用电子技术", "三年"),
        (2021, "hbxytc_2021_vocational_stop_pdf", "2021年停招专业", "4", "510202", "计算机网络技术", "三年"),
        (2021, "hbxytc_2021_vocational_stop_pdf", "2021年停招专业", "5", "460103", "数控技术", "三年"),
        (2021, "hbxytc_2021_vocational_stop_pdf", "2021年停招专业", "6", "460113", "模具设计与制造", "三年"),
        (2021, "hbxytc_2021_vocational_stop_pdf", "2021年停招专业", "7", "460703", "汽车电子技术", "三年"),
        (2021, "hbxytc_2021_vocational_stop_pdf", "2021年停招专业", "8", "530205", "投资与理财", "三年"),
        (2021, "hbxytc_2021_vocational_stop_pdf", "2021年停招专业", "9", "410401", "水产养殖技术", "三年"),
        (2021, "hbxytc_2021_vocational_stop_pdf", "2021年停招专业", "10", "510213", "移动应用开发", "三年"),
        (2021, "hbxytc_2021_vocational_stop_pdf", "2021年停招专业", "11", "490202", "生物制药技术", "三年"),
        (2021, "hbxytc_2021_vocational_stop_pdf", "2021年停招专业", "12", "570202", "应用英语", "三年"),
        (2021, "hbxytc_2021_vocational_stop_pdf", "2021年停招专业", "13", "570201", "商务英语", "三年"),
        (2021, "hbxytc_2021_vocational_stop_pdf", "2021年停招专业", "14", "510215", "动漫制作技术", "三年"),
        (2021, "hbxytc_2021_vocational_stop_pdf", "2021年停招专业", "15", "410105", "园艺技术", "三年"),
        (2021, "hbxytc_2021_vocational_stop_pdf", "2021年停招专业", "16", "590302", "智慧健康养老服务与管理", "三年"),
        (2021, "hbxytc_2021_vocational_stop_pdf", "2021年停招专业", "17", "530605", "市场营销", "三年"),
        (2021, "hbxytc_2021_vocational_stop_pdf", "2021年停招专业", "18", "440102", "建筑装饰工程技术", "三年"),
        (2021, "hbxytc_2021_vocational_stop_pdf", "2021年停招专业", "19", "520703K", "预防医学", "三年"),
        (2021, "hbxytc_2021_vocational_stop_pdf", "2021年停招专业", "20", "420802", "环境工程技术", "三年"),
    ]
    for policy_year, source_id, label, source_row_no, code, major, duration in hbxytc_vocational_stop_rows:
        rows.append(
            {
                "policy_year": policy_year,
                "region": "襄阳职业技术学院",
                "education_level": "高职高专",
                "record_type": "major_stop_enrollment",
                "warning_label": label,
                "reported_major_name": major,
                "major_code": code,
                "study_duration": duration,
                "policy_action": "停招。",
                "criterion_text": f"襄阳职业技术学院信息公开网“{label}”PDF表格逐项列出停招专业。",
                "source_row_no": source_row_no,
                "source_ids": source_id,
                "evidence_text": f"襄阳职业技术学院“{label}”PDF第{source_row_no}行列出{major}（{code}，学制{duration}）。",
                "confidence": "high",
            }
        )

    vocational_controlled_approval_rows = [
        ("approved", "司法类", "1.1-1", "山西省", "山西警官职业学院", "580605K", "罪犯心理测量与矫正技术", "3"),
        ("approved", "司法类", "1.1-2", "山东省", "山东司法警官职业学院", "580605K", "罪犯心理测量与矫正技术", "3"),
        ("approved", "司法类", "1.1-3", "湖南省", "湖南司法警官职业学院", "580506K", "专门矫治教育", "3"),
        ("approved", "司法类", "1.1-4", "山东省", "山东司法警官职业学院", "580506K", "专门矫治教育", "3"),
        ("approved", "教育类", "1.2-1", "北京市", "北京青年政治学院", "570101K", "早期教育", "2"),
        ("approved", "教育类", "1.2-2", "云南省", "云南文化艺术职业学院", "570101K", "早期教育", "2"),
        ("approved", "教育类", "1.2-3", "浙江省", "绍兴职业技术学院", "570101K", "早期教育", "2"),
        ("approved", "教育类", "1.2-4", "吉林省", "四平现代职业学院", "570101K", "早期教育", "3"),
        ("rejected", "司法类", "2.1-1", "四川省", "四川司法警官职业学院", "580506K", "专门矫治教育", "3"),
        ("rejected", "教育类", "2.2-1", "黑龙江省", "黑龙江民族职业学院", "570102K", "学前教育", "2"),
        ("rejected", "教育类", "2.2-2", "湖北省", "咸宁职业技术学院", "570102K", "学前教育", "2"),
        ("rejected", "教育类", "2.2-3", "上海市", "上海震旦职业学院", "570102K", "学前教育", "5"),
    ]
    for decision, category, source_row_no, province, school, code, major, duration in vocational_controlled_approval_rows:
        approved = decision == "approved"
        rows.append(
            {
                "policy_year": 2026,
                "region": school,
                "education_level": "高职高专",
                "record_type": "major_controlled_approved" if approved else "major_controlled_rejected",
                "warning_label": "2026年新设高职专科国家控制布点专业审批结果",
                "reported_major_name": major,
                "major_code": code,
                "study_duration": f"{duration}年",
                "policy_action": "同意设置国家控制专业点。" if approved else "不同意设置国家控制专业点。",
                "criterion_text": f"教育部根据《普通高等学校高等职业教育（专科）专业设置管理办法》组织2026年国家控制布点专业审批；类别为{category}，省份为{province}。",
                "source_row_no": source_row_no,
                "source_ids": "moe_2026_vocational_specialty_setup_results_notice;moe_2026_vocational_controlled_specialty_approval_pdf",
                "evidence_text": f"教育部官网附件《2026年新设高职专科国家控制布点专业审批结果》{category}第{source_row_no.split('-')[-1]}行列出{school}申报{major}（{code}，{duration}年），审批结果为{'同意设置' if approved else '不同意设置'}。",
                "confidence": "high",
            }
        )

    hzec_stop_rows = [
        (2024, "33", "教育体育学院", "商务日语", "570205", "三年"),
        (2024, "34", "教育体育学院", "婴幼儿托育服务与管理", "520802", "二、三年"),
        (2024, "35", "经济管理学院", "财富管理", "530205", "三年"),
        (2024, "36", "经济管理学院", "金融科技应用", "530202", "三年"),
        (2024, "37", "经济管理学院", "中小企业创业与经营", "530604", "三年"),
        (2024, "38", "经济管理学院", "机场运行服务与管理", "500408", "三年"),
        (2024, "39", "智慧城市设计学院", "书画艺术", "550107", "三年"),
        (2024, "27", "智慧城市设计学院", "建筑动画技术", "440107", "三年"),
        (2024, "40", "智能制造与汽车工程学院", "汽车造型与改装技术", "460705", "三年"),
        (2025, "43", "经济管理学院", "工商企业管理", "530601", "三年"),
        (2025, "44", "信息工程学院", "大数据技术", "510205", "三年"),
        (2025, "45", "智慧城市设计学院", "建筑工程技术", "440301", "三年"),
        (2025, "46", "智慧城市设计学院", "广告艺术设计", "550113", "二、三年"),
        (2025, "47", "智能制造与汽车工程学院", "机械设计与制造", "460101", "三年"),
    ]
    for policy_year, source_row_no, school_unit, major, code, duration in hzec_stop_rows:
        rows.append(
            {
                "policy_year": policy_year,
                "region": "惠州经济职业技术学院",
                "education_level": "高职高专",
                "record_type": "major_stop_enrollment",
                "warning_label": f"{policy_year}年停招专业",
                "reported_major_name": major,
                "major_code": code,
                "study_duration": duration,
                "policy_action": f"{policy_year}年停招。",
                "criterion_text": "惠州经济职业技术学院2024-2025学年专业设置、当年新增专业、停招专业名单PDF在“新招、停招专业情况”列标记停招专业。",
                "source_row_no": source_row_no,
                "source_ids": "hzec_2025_major_setting_page;hzec_2025_major_setting_pdf",
                "evidence_text": f"惠州经济职业技术学院2024-2025学年专业设置PDF第{source_row_no}行列出{school_unit}{major}（{code}，学制{duration}），新招、停招专业情况为{policy_year}年停招专业。",
                "confidence": "high",
            }
        )

    gdjmcmc_2024_stop_rows = [
        ("30", "老年保健与管理"),
        ("31", "生殖健康管理"),
        ("32", "呼吸治疗技术"),
    ]
    for source_row_no, major in gdjmcmc_2024_stop_rows:
        rows.append(
            {
                "policy_year": 2024,
                "region": "广东江门中医药职业学院",
                "education_level": "高职高专",
                "record_type": "major_stop_enrollment",
                "warning_label": "2024年停招专业",
                "reported_major_name": major,
                "policy_action": "2024年停招。",
                "criterion_text": "广东江门中医药职业学院2024-2025学年专业设置、当年新增专业、停招专业名单PDF在备注列标记2024年停招专业。",
                "source_row_no": source_row_no,
                "source_ids": "gdjmcmc_2025_major_setting_pdf",
                "evidence_text": f"广东江门中医药职业学院2024-2025学年专业设置PDF第{source_row_no}行列出{major}，备注为2024年停招。",
                "confidence": "high",
            }
        )

    xzcit_stop_rows = [
        (2025, "xzcit_2025_stop_page;xzcit_2025_stop_pdf", "徐州工业职业技术学院2025年停招专业", "1", "化学工程学院", "470102", "药品生物技术"),
        (2025, "xzcit_2025_stop_page;xzcit_2025_stop_pdf", "徐州工业职业技术学院2025年停招专业", "2", "化学工程学院", "490104", "食品检验检测技术"),
        (2025, "xzcit_2025_stop_page;xzcit_2025_stop_pdf", "徐州工业职业技术学院2025年停招专业", "3", "建筑工程学院", "440502", "建设工程管理"),
        (2025, "xzcit_2025_stop_page;xzcit_2025_stop_pdf", "徐州工业职业技术学院2025年停招专业", "4", "信息工程学院", "510205", "大数据技术"),
        (2024, "xzcit_2024_stop_page;xzcit_2024_stop_pdf", "徐州工业职业技术学院2024年停招专业", "1", "", "540101", "旅游管理"),
        (2024, "xzcit_2024_stop_page;xzcit_2024_stop_pdf", "徐州工业职业技术学院2024年停招专业", "2", "", "540106", "酒店管理与数字化运营"),
        (2024, "xzcit_2024_stop_page;xzcit_2024_stop_pdf", "徐州工业职业技术学院2024年停招专业", "3", "", "570201", "商务英语"),
        (2024, "xzcit_2024_stop_page;xzcit_2024_stop_pdf", "徐州工业职业技术学院2024年停招专业", "4", "", "500603", "城市轨道交通机电技术"),
        (2022, "xzcit_2022_stop_html", "徐州工业职业技术学院2022年停招专业", "1", "", "480108", "皮具制作与工艺"),
        (2022, "xzcit_2022_stop_html", "徐州工业职业技术学院2022年停招专业", "2", "", "470205", "煤化工技术"),
        (2022, "xzcit_2022_stop_html", "徐州工业职业技术学院2022年停招专业", "3", "", "510206", "云计算技术应用"),
        (2021, "xzcit_2021_stop_page;xzcit_2021_stop_pdf", "徐州工业职业技术学院2021年停招专业", "1", "", "530304", "光伏发电技术与应用"),
        (2021, "xzcit_2021_stop_page;xzcit_2021_stop_pdf", "徐州工业职业技术学院2021年停招专业", "2", "", "650104", "数字媒体艺术设计"),
    ]
    for policy_year, source_ids, label, source_row_no, school_unit, code, major in xzcit_stop_rows:
        owner_text = f"{school_unit}" if school_unit else ""
        rows.append(
            {
                "policy_year": policy_year,
                "region": "徐州工业职业技术学院",
                "education_level": "高职高专",
                "record_type": "major_stop_enrollment",
                "warning_label": label,
                "reported_major_name": major,
                "major_code": code,
                "policy_action": "停招。",
                "criterion_text": f"{label}来源表格逐项列出停招专业。",
                "source_row_no": source_row_no,
                "source_ids": source_ids,
                "evidence_text": f"{label}第{source_row_no}行列出{owner_text}{major}（{code}）。",
                "confidence": "high",
            }
        )

    wust_2023_stop_rows = [
        ("1", "管理学院（恒大管理学院）", "120801", "电子商务", "四年", "管理学"),
        ("2", "管理学院（恒大管理学院）", "120209", "物业管理", "四年", "管理学"),
        ("3", "汽车与交通工程学院", "080208", "汽车服务工程", "四年", "工学"),
        ("4", "资源与环境工程学院", "070503", "人文地理与城乡规划", "四年", "理学"),
        ("5", "汽车与交通工程学院", "081801", "交通运输", "四年", "工学"),
        ("6", "马克思主义学院", "030504T", "马克思主义理论", "四年", "法学"),
        ("7", "法学与经济学院", "120403", "劳动与社会保障", "四年", "管理学"),
    ]
    for source_row_no, school_unit, code, major, duration, degree in wust_2023_stop_rows:
        rows.append(
            {
                "policy_year": 2023,
                "region": "武汉科技大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2023年停招专业",
                "reported_major_name": major,
                "major_code": code,
                "study_duration": duration,
                "policy_action": "停招。",
                "criterion_text": "武汉科技大学信息公开网2023年停招专业HTML表格逐项列出停招本科专业。",
                "source_row_no": source_row_no,
                "source_ids": "wust_2023_stop_enrollment",
                "evidence_text": f"武汉科技大学2023年停招专业表第{source_row_no}行列出{school_unit}{major}（{code}，修业年限{duration}，学科{degree}）。",
                "confidence": "high",
            }
        )

    add_warning_list(
        2024,
        "武汉科技大学",
        "本科",
        "major_stop_enrollment",
        "2023-2024学年本科教学质量报告2024年停招专业",
        ["绘画", "国际经济与贸易", "人力资源管理", "财务管理"],
        "停招。",
        "武汉科技大学2023-2024学年本科教学质量报告列明2024年停招专业。",
        ["wust_2024_teaching_quality_report_pdf"],
        "武汉科技大学2023-2024学年本科教学质量报告专业设置情况说明2024年停招绘画、国际经济与贸易、人力资源管理、财务管理4个专业。",
        "high",
    )
    add_warning_list(
        2024,
        "武汉科技大学",
        "本科",
        "major_cancel",
        "2023-2024学年本科教学质量报告申请撤销专业",
        ["人文地理与城乡规划", "交通运输"],
        "申请撤销。",
        "武汉科技大学2023-2024学年本科教学质量报告说明学校申请撤销2个专业。",
        ["wust_2024_teaching_quality_report_pdf"],
        "武汉科技大学2023-2024学年本科教学质量报告说明申请撤销人文地理与城乡规划和交通运输专业。",
        "high",
    )
    cuit_cancel_rows = [
        ("1", "材料物理", "080402"),
        ("2", "信息对抗技术", "080706T"),
    ]
    for source_row_no, major, code in cuit_cancel_rows:
        rows.append(
            {
                "policy_year": 2024,
                "region": "成都信息工程大学",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "2023-2024学年本科教学质量报告撤销备案专业",
                "reported_major_name": major,
                "major_code": code,
                "policy_action": "撤销并上报备案。",
                "criterion_text": "成都信息工程大学2023-2024学年本科教学质量报告说明，学校常态开展专业评估监测，实施专业“红橙黄牌”管理，全面开展专业预警与动态调整。",
                "source_row_no": source_row_no,
                "source_ids": "cuit_2024_teaching_quality_report_pdf",
                "evidence_text": f"成都信息工程大学2023-2024学年本科教学质量报告专业建设部分称，学校撤销“材料物理”和“信息对抗技术”两个专业并上报备案；本行对应第{source_row_no}项：{major}（{code}）。",
                "confidence": "high",
            }
        )

    ldxy_2022_not_enrolled = [
        "人文教育",
        "秘书学",
        "汉语国际教育",
        "文化产业管理",
        "地理信息科学",
        "公共事业管理",
        "产品设计",
        "舞蹈表演",
        "测控技术与仪器",
        "应用电子技术教育",
        "科学教育",
        "通信工程",
        "教育技术学",
        "自然地理与资源环境",
        "材料成型及控制工程",
    ]
    add_warning_list(
        2022,
        "陇东学院",
        "本科",
        "major_stop_enrollment",
        "2023-2024学年本科教学质量报告2022年未招生专业",
        ldxy_2022_not_enrolled,
        "未招生。",
        "陇东学院2023-2024学年本科教学质量报告逐名列出2022年未招生专业。",
        ["ldxy_2024_teaching_quality_report_pdf"],
        "陇东学院2023-2024学年本科教学质量报告称2022年人文教育、秘书学、汉语国际教育、文化产业管理、地理信息科学、公共事业管理、产品设计、舞蹈表演、测控技术与仪器、应用电子技术教育、科学教育、通信工程、教育技术学、自然地理与资源环境、材料成型及控制工程等15个专业未招生。",
        "high",
    )
    ldxy_2023_not_enrolled = [
        "地理信息科学",
        "自然地理与资源环境",
        "教育技术学",
        "公共事业管理",
        "法学",
        "园艺",
        "生物技术",
        "应用化学",
        "科学教育",
        "油气储运工程",
        "信息与计算科学",
        "网络工程",
        "工程造价",
        "秘书学",
        "汉语国际教育",
        "人文教育",
        "文化产业管理",
        "绘画",
        "音乐表演",
        "通信工程",
        "测控技术与仪器",
        "应用电子技术教育",
        "材料成型及控制工程",
        "机械电子工程",
    ]
    add_warning_list(
        2023,
        "陇东学院",
        "本科",
        "major_stop_enrollment",
        "2023-2024学年本科教学质量报告2023年未招生专业",
        ldxy_2023_not_enrolled,
        "未招生。",
        "陇东学院2023-2024学年本科教学质量报告逐名列出2023年未招生专业。",
        ["ldxy_2024_teaching_quality_report_pdf"],
        "陇东学院2023-2024学年本科教学质量报告称2023年地理信息科学、自然地理与资源环境、教育技术学、公共事业管理、法学、园艺、生物技术、应用化学、科学教育、油气储运工程、信息与计算科学、网络工程、工程造价、秘书学、汉语国际教育、人文教育、文化产业管理、绘画、音乐表演、通信工程、测控技术与仪器、应用电子技术教育、材料成型及控制工程、机械电子工程等24个专业未招生。",
        "high",
    )
    add_warning_list(
        2023,
        "陇东学院",
        "本科",
        "major_cancel",
        "2023-2024学年本科教学质量报告申请撤销专业",
        ["人文教育", "测控技术与仪器", "应用电子技术教育", "材料成型及控制工程", "通信工程", "教育技术学", "自然地理与资源环境", "科学教育"],
        "主动申请撤销。",
        "陇东学院2023-2024学年本科教学质量报告说明学校主动申请撤销8个停招五年以上且没有在校生的本科专业。",
        ["ldxy_2024_teaching_quality_report_pdf"],
        "陇东学院2023-2024学年本科教学质量报告称，学校主动申请撤销人文教育、测控技术与仪器、应用电子技术教育、材料成型及控制工程、通信工程、教育技术学、自然地理与资源环境和科学教育等8个停招五年以上且没有在校生的本科专业。",
        "high",
    )

    dlut_stop_rows = [
        (2013, "note-3-1", "物流工程", "120602", "2013年停止招生"),
        (2014, "note-3-2", "纳米材料与技术", "080413T", "2014年停止招生"),
        (2014, "note-3-3", "物联网工程", "080905", "2014年停止招生"),
        (2016, "note-3-4", "材料物理", "080402", "2016年停止招生"),
        (2016, "note-3-5", "通信工程", "080703", "2016年停止招生"),
        (2016, "note-3-6", "生物技术", "071002", "2016年停止招生"),
        (2016, "note-3-7", "无机非金属材料工程", "080406", "2016年停止招生"),
        (2016, "note-3-8", "能源化学工程", "081304T", "2016年停止招生并于2018年在盘锦校区恢复招生"),
        (2018, "note-3-9", "能源与环境系统工程", "080502T", "2018年停止招生"),
        (2018, "note-3-10", "法学", "030101K", "2018年停止招生"),
        (2021, "note-3-11", "人力资源管理", "120206", "2011年停止招生并于2013年在盘锦校区恢复招生，于2021年再次停止招生"),
        (2022, "note-3-12", "俄语", "050202", "2022年停止招生"),
        (2024, "note-3-13", "资源循环科学与工程", "081303T", "2024年停止招生"),
        (2024, "note-3-14", "食品科学与工程", "082701", "2024年停止招生"),
        (2024, "note-3-15", "运动康复", "040206T", "2024年停止招生"),
        (2024, "note-3-16", "商务英语", "050262", "2024年停止招生"),
        (2024, "note-4-1", "管理科学", "120101", "管理科学暂未招生"),
    ]
    for policy_year, source_row_no, major, code, note in dlut_stop_rows:
        rows.append(
            {
                "policy_year": policy_year,
                "region": "大连理工大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2023-2024学年本科教学质量报告历年停招/未招生专业",
                "reported_major_name": major,
                "major_code": code,
                "policy_action": "停止招生或暂未招生。",
                "criterion_text": "大连理工大学2023-2024学年本科教学质量报告专业设置脚注列明历年停止招生专业和暂未招生专业。",
                "source_row_no": source_row_no,
                "source_ids": "dlut_2024_teaching_quality_report_pdf",
                "evidence_text": f"大连理工大学2023-2024学年本科教学质量报告专业设置脚注列明：{note}；本行对应{major}（{code}）。",
                "confidence": "high",
            }
        )
    add_warning_list(
        2024,
        "大连理工大学",
        "本科",
        "major_cancel",
        "2023-2024学年本科教学质量报告申请撤销专业",
        ["物流工程", "纳米材料与技术", "材料物理", "无机非金属材料工程", "法学", "管理科学"],
        "申请撤销。",
        "大连理工大学2023-2024学年本科教学质量报告脚注说明学校已于2024年申请撤销6个专业。",
        ["dlut_2024_teaching_quality_report_pdf"],
        "大连理工大学2023-2024学年本科教学质量报告脚注称，学校已于2024年申请撤销物流工程、纳米材料与技术、材料物理、无机非金属材料工程、法学、管理科学等6个专业。",
        "high",
    )

    sie_stop_rows = [
        ("1", "机械电子工程", "080204", "专业设置表在机械类中标注机械电子工程（停招）。"),
        ("2", "机械工艺技术", "080209T", "专业设置表在机械类中标注机械工艺技术（停招）。"),
        ("3", "商务英语", "050262", "专业设置表在文学类中标注商务英语（停招）。"),
    ]
    for source_row_no, major, code, note in sie_stop_rows:
        rows.append(
            {
                "policy_year": 2024,
                "region": "沈阳工程学院",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2023-2024学年本科教学质量报告停招专业",
                "reported_major_name": major,
                "major_code": code,
                "policy_action": "停招。",
                "criterion_text": "沈阳工程学院2023-2024学年本科教学质量报告本科专业设置情况一览表标注停招专业。",
                "source_row_no": source_row_no,
                "source_ids": "sie_2024_teaching_quality_report_pdf",
                "evidence_text": f"沈阳工程学院2023-2024学年本科教学质量报告本科专业设置情况一览表{note}",
                "confidence": "high",
            }
        )
    add_warning_list(
        2023,
        "沈阳工程学院",
        "本科",
        "major_cancel",
        "2023-2024学年本科教学质量报告撤销专业",
        ["测控技术与仪器"],
        "撤销。",
        "沈阳工程学院2023-2024学年本科教学质量报告说明2023年撤销测控技术与仪器专业。",
        ["sie_2024_teaching_quality_report_pdf"],
        "沈阳工程学院2023-2024学年本科教学质量报告称2023年撤销测控技术与仪器专业。",
        "high",
    )
    add_warning_list(
        2024,
        "沈阳工程学院",
        "本科",
        "major_cancel",
        "2023-2024学年本科教学质量报告申请撤销专业",
        ["机械电子工程", "商务英语"],
        "申请撤销。",
        "沈阳工程学院2023-2024学年本科教学质量报告说明2024年申请撤销机械电子工程、商务英语专业。",
        ["sie_2024_teaching_quality_report_pdf"],
        "沈阳工程学院2023-2024学年本科教学质量报告称2024年申请撤销机械电子工程、商务英语专业。",
        "high",
    )

    zjhzu_stop_rows = [
        ("2", "机械电子工程", "080204"),
        ("12", "物联网工程", "080905"),
        ("19", "电子商务", "120801"),
        ("20", "市场营销", "120202"),
        ("21", "物流管理", "120601"),
        ("22", "行政管理", "120402"),
        ("28", "秘书学", "050107T"),
        ("30", "日语", "050207"),
        ("31", "商务英语", "050262"),
    ]
    for source_row_no, major, code in zjhzu_stop_rows:
        rows.append(
            {
                "policy_year": 2024,
                "region": "湖州学院",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2023-2024学年本科教学质量报告停招专业",
                "reported_major_name": major,
                "major_code": code,
                "policy_action": "停招。",
                "criterion_text": "湖州学院2023-2024学年本科教学质量报告专业设置与学科支撑情况表在招生情况列标注停招。",
                "source_row_no": source_row_no,
                "source_ids": "zjhzu_2024_teaching_quality_report_pdf",
                "evidence_text": f"湖州学院2023-2024学年本科教学质量报告专业设置表第{source_row_no}项列出{major}（{code}）招生情况为停招。",
                "confidence": "high",
            }
        )
    add_warning_list(
        2023,
        "湖州学院",
        "本科",
        "major_cancel",
        "2023-2024学年本科教学质量报告撤销专业",
        ["历史学", "美术学"],
        "撤销。",
        "湖州学院2023-2024学年本科教学质量报告说明2023年撤销历史学、美术学2个专业。",
        ["zjhzu_2024_teaching_quality_report_pdf"],
        "湖州学院2023-2024学年本科教学质量报告称2023年获批高分子材料与工程、供应链管理2个专业，撤销历史学、美术学2个专业。",
        "high",
    )
    add_warning_list(
        2024,
        "湖州学院",
        "本科",
        "major_cancel",
        "2023-2024学年本科教学质量报告申请撤销专业",
        ["物联网工程"],
        "申请撤销。",
        "湖州学院2023-2024学年本科教学质量报告说明2024年申请撤销物联网工程1个专业。",
        ["zjhzu_2024_teaching_quality_report_pdf"],
        "湖州学院2023-2024学年本科教学质量报告称2024年申请增设人工智能1个专业，申请撤销物联网工程1个专业。",
        "high",
    )

    add_warning_list(
        2024,
        "南京师范大学泰州学院",
        "本科",
        "major_cancel",
        "2023-2024学年本科教学质量报告申请撤销长期未招生专业",
        ["广告学", "戏剧影视文学", "园艺"],
        "申请撤销。",
        "南京师范大学泰州学院2023-2024学年本科教学质量报告说明申请撤销3个长期未招生专业。",
        ["nnutc_2024_teaching_quality_report_pdf"],
        "南京师范大学泰州学院2023-2024学年本科教学质量报告称申请撤销广告学、戏剧影视文学、园艺等3个长期未招生专业。",
        "high",
    )
    add_warning_list(
        2023,
        "四川音乐学院",
        "本科",
        "major_cancel",
        "2023-2024学年本科教学质量报告申请撤销/撤销未招生专业",
        ["工业设计", "公共事业管理", "戏剧学", "戏剧影视导演", "服装与服饰设计"],
        "申请撤销并一并撤销未招生专业。",
        "四川音乐学院2023-2024学年本科教学质量报告说明2023年、2024年申请撤销或一并撤销5个专业。",
        ["sccm_2024_teaching_quality_report_pdf"],
        "四川音乐学院2023-2024学年本科教学质量报告称2023年、2024年申请撤销工业设计、公共事业管理，一并撤销戏剧学、戏剧影视导演、服装与服饰设计3个未招生专业。",
        "high",
    )
    rows.append(
        {
            "policy_year": 2024,
            "region": "四川音乐学院",
            "education_level": "本科",
            "record_type": "major_stop_enrollment",
            "warning_label": "2023-2024学年本科教学质量报告当年停招专业",
            "reported_major_name": "新媒体艺术",
            "major_code": "130511T",
            "policy_action": "当年停招。",
            "criterion_text": "四川音乐学院2023-2024学年本科教学质量报告支撑数据列出当年停招专业名单。",
            "source_row_no": "1",
            "source_ids": "sccm_2024_teaching_quality_report_pdf",
            "evidence_text": "四川音乐学院2023-2024学年本科教学质量报告支撑数据列出当年停招专业1个：新媒体艺术（130511T，艺术学）。",
            "confidence": "high",
        }
    )

    hunau_stop_rows = [
        ("社会工作", "030302"),
        ("信息工程", "080706"),
        ("汽车服务工程", "080208"),
        ("机械电子工程", "080204"),
        ("水族科学与技术", "090603T"),
    ]
    for index, (major, code) in enumerate(hunau_stop_rows, start=1):
        rows.append(
            {
                "policy_year": 2024,
                "region": "湖南农业大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2023-2024学年本科教学质量报告近5年停招专业",
                "reported_major_name": major,
                "major_code": code,
                "policy_action": "停招。",
                "criterion_text": "湖南农业大学2023-2024学年本科教学质量报告列明近5年停招本科专业。",
                "source_row_no": str(index),
                "source_ids": "hunau_2024_teaching_quality_report_pdf",
                "evidence_text": f"湖南农业大学2023-2024学年本科教学质量报告称近5年停招社会工作、信息工程、汽车服务工程、机械电子工程、水族科学与技术等5个本科专业；本行对应{major}（{code}）。",
                "confidence": "high",
            }
        )
    add_warning_list(
        2022,
        "湖南农业大学",
        "本科",
        "major_cancel",
        "2023-2024学年本科教学质量报告近5年撤销专业",
        ["植物科学与技术", "表演"],
        "撤销。",
        "湖南农业大学2023-2024学年本科教学质量报告说明近5年撤销2个本科专业。",
        ["hunau_2024_teaching_quality_report_pdf"],
        "湖南农业大学2023-2024学年本科教学质量报告称近5年撤销植物科学与技术、表演等2个本科专业。",
        "high",
    )
    add_warning_list(
        2024,
        "湖南农业大学",
        "本科",
        "major_cancel",
        "2023-2024学年本科教学质量报告申请撤销专业",
        ["信息工程", "汽车服务工程", "社会工作"],
        "申请撤销。",
        "湖南农业大学2023-2024学年本科教学质量报告说明2024年申请撤销3个本科专业。",
        ["hunau_2024_teaching_quality_report_pdf"],
        "湖南农业大学2023-2024学年本科教学质量报告称2024年申请撤销信息工程、汽车服务工程、社会工作3个本科专业。",
        "high",
    )

    zust_stop_rows = [
        ("测控技术与仪器", "080301"),
        ("物联网工程", "080905"),
        ("包装工程", "081702"),
        ("电子商务", "120801"),
        ("汽车服务工程", "080208"),
    ]
    for index, (major, code) in enumerate(zust_stop_rows, start=1):
        rows.append(
            {
                "policy_year": 2024,
                "region": "浙江科技大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2023-2024学年本科教学质量报告停招专业",
                "reported_major_name": major,
                "major_code": code,
                "policy_action": "停招。",
                "criterion_text": "浙江科技大学2023-2024学年本科教学质量报告专业设置脚注列出本学年停招专业。",
                "source_row_no": str(index),
                "source_ids": "zust_2024_teaching_quality_report_pdf",
                "evidence_text": f"浙江科技大学2023-2024学年本科教学质量报告脚注称2023-2024学年停招专业为测控技术与仪器、物联网工程、包装工程、电子商务、汽车服务工程；本行对应{major}（{code}）。",
                "confidence": "high",
            }
        )
    add_warning_list(
        2024,
        "浙江科技大学",
        "本科",
        "major_cancel",
        "2023-2024学年本科教学质量报告撤销专业",
        ["包装工程", "电子商务"],
        "撤销。",
        "浙江科技大学2023-2024学年本科教学质量报告说明撤销2个专业。",
        ["zust_2024_teaching_quality_report_pdf"],
        "浙江科技大学2023-2024学年本科教学质量报告称拟增设1个新专业、1个第二学士学位专业，撤销2个专业；本行对应可与教育部2024年度撤销名单合并的包装工程和电子商务。",
        "high",
    )

    chnu_stop_rows = [
        ("广告学", "050303"),
        ("戏剧影视文学", "130304"),
        ("社会学", "030301"),
        ("国际经济与贸易", "020401"),
        ("审计学", "120207"),
    ]
    for index, (major, code) in enumerate(chnu_stop_rows, start=1):
        rows.append(
            {
                "policy_year": 2024,
                "region": "淮北师范大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2023-2024学年本科教学质量报告停招专业",
                "reported_major_name": major,
                "major_code": code,
                "policy_action": "停招。",
                "criterion_text": "淮北师范大学2023-2024学年本科教学质量报告说明停招5个本科专业。",
                "source_row_no": str(index),
                "source_ids": "chnu_2024_teaching_quality_report_pdf",
                "evidence_text": f"淮北师范大学2023-2024学年本科教学质量报告称停招广告学、戏剧影视文学、社会学、国际经济与贸易、审计学专业；本行对应{major}（{code}）。",
                "confidence": "high",
            }
        )
    add_warning_list(
        2023,
        "淮北师范大学",
        "本科",
        "major_cancel",
        "2023-2024学年本科教学质量报告完成撤销专业",
        ["金融数学", "机械电子工程", "计算机科学与技术", "数字媒体技术", "财务管理", "公共事业管理", "劳动与社会保障"],
        "完成撤销。",
        "淮北师范大学2023-2024学年本科教学质量报告说明2024年完成7个专业的撤销工作。",
        ["chnu_2024_teaching_quality_report_pdf"],
        "淮北师范大学2023-2024学年本科教学质量报告称2024年完成金融数学、机械电子工程、计算机科学与技术（理学）、数字媒体技术、财务管理、公共事业管理、劳动与社会保障等7个专业的撤销工作。",
        "high",
    )
    add_warning_list(
        2024,
        "淮北师范大学",
        "本科",
        "major_cancel",
        "2023-2024学年本科教学质量报告申请撤销专业",
        ["信息管理与信息系统"],
        "申请撤销。",
        "淮北师范大学2023-2024学年本科教学质量报告说明申请撤销信息管理与信息系统专业。",
        ["chnu_2024_teaching_quality_report_pdf"],
        "淮北师范大学2023-2024学年本科教学质量报告称申请撤销信息管理与信息系统专业。",
        "high",
    )

    wzut_stop_rows = [
        (2020, "29", "市场营销", "120202", "管理学"),
        (2020, "30", "音乐表演", "130201", "艺术学"),
        (2021, "31", "人力资源管理", "120206", "管理学"),
        (2022, "32", "电子商务", "120801", "工学"),
        (2023, "33", "车辆工程", "080207", "工学"),
        (2023, "34", "广告学", "050303", "文学"),
    ]
    for policy_year, source_row_no, major, code, degree in wzut_stop_rows:
        rows.append(
            {
                "policy_year": policy_year,
                "region": "温州理工学院",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2023-2024学年本科教学质量报告历年停招专业",
                "reported_major_name": major,
                "major_code": code,
                "study_duration": "四年",
                "policy_action": f"{policy_year}年停招。",
                "criterion_text": "温州理工学院2023-2024学年本科教学质量报告本科专业设置一览表逐项列出停招年份。",
                "source_row_no": source_row_no,
                "source_ids": "wzut_2024_teaching_quality_report_pdf",
                "evidence_text": f"温州理工学院2023-2024学年本科教学质量报告本科专业设置表第{source_row_no}项列出{major}（{code}，{degree}），当年招生情况为{policy_year}年停招。",
                "confidence": "high",
            }
        )
    czu_stop_rows = [
        ("1", "材料化学", "080403", "当年停招专业名单列出材料化学"),
        ("2", "市场营销（专升本）", "120202", "正文列出校内停招专业市场营销（专升本）"),
        ("3", "知识产权（安警院联培）", "030102T", "正文列出校内停招专业知识产权（安警院联培）"),
    ]
    for source_row_no, major, code, note in czu_stop_rows:
        rows.append(
            {
                "policy_year": 2024,
                "region": "池州学院",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2023-2024学年本科教学质量报告停招专业",
                "reported_major_name": major,
                "major_code": code,
                "policy_action": "停招。",
                "criterion_text": "池州学院2023-2024学年本科教学质量报告正文和支撑数据列出停招专业。",
                "source_row_no": source_row_no,
                "source_ids": "czu_2024_teaching_quality_report_page",
                "evidence_text": f"池州学院2023-2024学年本科教学质量报告称停招的校内专业为材料化学、市场营销（专升本）、知识产权（安警院联培）；支撑数据当年停招专业名单列出材料化学。本行对应{major}（{code}），证据位置：{note}。",
                "confidence": "high",
            }
        )
    add_warning_list(
        2023,
        "北京物资学院",
        "本科",
        "major_cancel",
        "2023-2024学年本科教学质量报告撤销专业",
        ["经济统计学", "英语", "劳动关系"],
        "撤销。",
        "北京物资学院2023-2024学年本科教学质量报告说明2023年撤销3个本科专业。",
        ["bwu_2024_teaching_quality_report_pdf"],
        "北京物资学院2023-2024学年本科教学质量报告称2023年学校增设金融科技专业，开设大数据管理与应用（中外合作办学）专业方向，撤销经济统计学、英语、劳动关系三个专业。",
        "high",
    )
    xcu_stop_rows = [
        ("1", "通信工程", "080703"),
        ("2", "网络工程", "080903"),
        ("3", "交通设备与控制工程", "081806T"),
        ("4", "酒店管理", "120902"),
        ("5", "财务管理", "120204"),
        ("6", "绘画", "130402"),
        ("7", "汉语国际教育", "050103"),
        ("8", "商务英语", "050262"),
        ("9", "工商管理", "120201K"),
    ]
    for source_row_no, major, code in xcu_stop_rows:
        rows.append(
            {
                "policy_year": 2024,
                "region": "许昌学院",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2023-2024学年本科教学质量报告停招专业",
                "reported_major_name": major,
                "major_code": code,
                "policy_action": "停招。",
                "criterion_text": "许昌学院2023-2024学年本科教学质量报告专业建设部分列出本学年停招本科专业。",
                "source_row_no": source_row_no,
                "source_ids": "xcu_2024_teaching_quality_report_pdf",
                "evidence_text": f"许昌学院2023-2024学年本科教学质量报告称停招通信工程、网络工程、交通设备与控制工程、酒店管理、财务管理、绘画、汉语国际教育、商务英语、工商管理等9个本科专业；本行对应{major}（{code}）。",
                "confidence": "high",
            }
        )
    add_warning_list(
        2024,
        "许昌学院",
        "本科",
        "major_cancel",
        "2023-2024学年本科教学质量报告申请撤销专业",
        ["城乡规划", "市场营销", "音乐表演", "社会体育指导与管理"],
        "申请撤销。",
        "许昌学院2023-2024学年本科教学质量报告说明4个专业申请撤销。",
        ["xcu_2024_teaching_quality_report_pdf"],
        "许昌学院2023-2024学年本科教学质量报告称城乡规划、市场营销、音乐表演、社会体育指导与管理等4个专业申请撤销。",
        "high",
    )

    wsyu_stop_rows = [
        (2025, "wsyu_2025_stop_enrollment", "1", "01124", "物联网工程（腾讯云精英班）", "物联网工程", "080905", "信息科学与工程学院", "2022", "4", "工学", "当年停招"),
        (2025, "wsyu_2025_stop_enrollment", "2", "01117", "物联网工程（厚溥智能信息特色班）", "物联网工程", "080905", "信息科学与工程学院", "2017", "4", "工学", "已停招"),
        (2025, "wsyu_2025_stop_enrollment", "3", "01110", "物联网工程", "物联网工程", "080905", "信息科学与工程学院", "2014", "4", "工学", "当年停招"),
        (2025, "wsyu_2025_stop_enrollment", "4", "01121", "通信工程（讯方5G特色班）", "通信工程", "080703", "信息科学与工程学院", "2021", "4", "工学", "已停招"),
        (2025, "wsyu_2025_stop_enrollment", "5", "01125", "通信工程（5G特色班）", "通信工程", "080703", "信息科学与工程学院", "2022", "4", "工学", "当年停招"),
        (2025, "wsyu_2025_stop_enrollment", "6", "01105", "通信工程", "通信工程", "080703", "信息科学与工程学院", "2001", "4", "工学", "已停招"),
        (2025, "wsyu_2025_stop_enrollment", "7", "02109", "自动化（智能制造）", "自动化", "080801", "机电与自动化学院", "2001", "4", "工学", "已停招"),
        (2025, "wsyu_2025_stop_enrollment", "8", "03103", "环境工程", "环境工程", "082502", "城市建设学院", "2002", "4", "工学", "当年停招"),
        (2025, "wsyu_2025_stop_enrollment", "9", "03108", "给排水科学与工程", "给排水科学与工程", "081003", "城市建设学院", "2003", "4", "工学", "当年停招"),
        (2025, "wsyu_2025_stop_enrollment", "10", "05101", "国际经济与贸易", "国际经济与贸易", "020401", "经济管理学院", "2002", "4", "经济学", "当年停招"),
        (2025, "wsyu_2025_stop_enrollment", "11", "06107", "广播电视编导", "广播电视编导", "130305", "新闻与文法学院", "2016", "4", "艺术学", "当年停招"),
        (2024, "wsyu_2024_stop_enrollment", "1", "01117", "物联网工程（厚溥智能信息特色班）", "物联网工程", "080905", "信息科学与工程学院", "2017", "4", "工学", "已停招"),
        (2024, "wsyu_2024_stop_enrollment", "2", "01121", "通信工程（讯方5G特色班）", "通信工程", "080703", "信息科学与工程学院", "2021", "4", "工学", "已停招"),
        (2024, "wsyu_2024_stop_enrollment", "3", "01105", "通信工程", "通信工程", "080703", "信息科学与工程学院", "2001", "4", "工学", "已停招"),
        (2024, "wsyu_2024_stop_enrollment", "4", "02109", "自动化（智能制造）", "自动化", "080801", "机电与自动化学院", "2001", "4", "工学", "当年停招"),
        (2024, "wsyu_2024_stop_enrollment", "5", "01116", "通信工程（讯方技术特色班）", "通信工程", "080703", "信息科学与工程学院", "2017", "4", "工学", "已停招"),
        (2024, "wsyu_2024_stop_enrollment", "6", "01112", "软件工程（中软国际特色班）", "软件工程", "080902", "信息科学与工程学院", "2016", "4", "工学", "已停招"),
        (2024, "wsyu_2024_stop_enrollment", "7", "01111", "计算机科学与技术（中软国际特色班）", "计算机科学与技术", "080901", "信息科学与工程学院", "2016", "4", "工学", "已停招"),
        (2024, "wsyu_2024_stop_enrollment", "8", "01109", "光电信息科学与工程", "光电信息科学与工程", "080705", "信息科学与工程学院", "2001", "4", "工学", "已停招"),
        (2024, "wsyu_2024_stop_enrollment", "9", "01118", "电子信息工程（东软大数据特色班）", "电子信息工程", "080701", "信息科学与工程学院", "2017", "4", "工学", "已停招"),
        (2024, "wsyu_2024_stop_enrollment", "10", "03104", "工程管理", "工程管理", "120103", "城市建设学院", "2004", "4", "工学", "已停招"),
        (2024, "wsyu_2024_stop_enrollment", "11", "05111", "金融学（金融智能创新实验班）", "金融学", "020301K", "经济管理学院", "2019", "4", "经济学", "已停招"),
        (2023, "wsyu_2023_stop_enrollment", "1", "", "计算机科学与技术（中软国际特色班）", "计算机科学与技术", "080901", "信息科学与工程学院", "2016", "4", "工学", "已停招"),
        (2023, "wsyu_2023_stop_enrollment", "2", "", "软件工程（中软国际特色班）", "软件工程", "080902", "信息科学与工程学院", "2016", "4", "工学", "已停招"),
        (2023, "wsyu_2023_stop_enrollment", "3", "", "通信工程", "通信工程", "080703", "信息科学与工程学院", "2001", "4", "工学", "已停招"),
        (2023, "wsyu_2023_stop_enrollment", "4", "", "通信工程（讯方技术特色班）", "通信工程", "080703", "信息科学与工程学院", "2017", "4", "工学", "已停招"),
        (2023, "wsyu_2023_stop_enrollment", "5", "", "光电信息科学与工程", "光电信息科学与工程", "080705", "信息科学与工程学院", "2001", "4", "工学", "已停招"),
        (2023, "wsyu_2023_stop_enrollment", "6", "", "电子信息工程（东软大数据特色班）", "电子信息工程", "080701", "信息科学与工程学院", "2017", "4", "工学", "已停招"),
        (2023, "wsyu_2023_stop_enrollment", "7", "", "物联网工程（厚溥智能信息特色班）", "物联网工程", "080905", "信息科学与工程学院", "2017", "4", "工学", "已停招"),
        (2023, "wsyu_2023_stop_enrollment", "8", "", "通信工程（讯方5G特色班）", "通信工程", "080703", "信息科学与工程学院", "2021", "4", "工学", "已停招"),
        (2023, "wsyu_2023_stop_enrollment", "9", "", "工程管理", "工程管理", "120103", "城市建设学院", "2004", "4", "工学", "已停招"),
        (2023, "wsyu_2023_stop_enrollment", "10", "", "道路桥梁与渡河工程", "道路桥梁与渡河工程", "081006T", "城市建设学院", "2015", "4", "工学", "已停招"),
        (2023, "wsyu_2023_stop_enrollment", "11", "", "金融学（金融智能创新实验班）", "金融学", "020301K", "经济管理学院", "2019", "4", "经济学", "已停招"),
        (2021, "wsyu_2021_stop_enrollment", "1", "01111", "计算机科学与技术(中软国际特色班)", "计算机科学与技术", "080901", "信息科学与工程学院", "2016", "4", "工学", "当年停招"),
        (2021, "wsyu_2021_stop_enrollment", "2", "01115", "计算机科学与技术(华胜天成大数据方向)", "计算机科学与技术", "080901", "信息科学与工程学院", "2017", "4", "工学", "已停招"),
        (2021, "wsyu_2021_stop_enrollment", "3", "01112", "软件工程(中软国际特色班)", "软件工程", "080902", "信息科学与工程学院", "2016", "4", "工学", "当年停招"),
        (2021, "wsyu_2021_stop_enrollment", "4", "01116", "通信工程(讯方技术特色班)", "通信工程", "080703", "信息科学与工程学院", "2017", "4", "工学", "当年停招"),
        (2021, "wsyu_2021_stop_enrollment", "5", "01109", "光电信息科学与工程", "光电信息科学与工程", "080705", "信息科学与工程学院", "2001", "4", "工学", "当年停招"),
        (2021, "wsyu_2021_stop_enrollment", "6", "01118", "电子信息工程(东软大数据特色班)", "电子信息工程", "080701", "信息科学与工程学院", "2017", "4", "工学", "当年停招"),
        (2021, "wsyu_2021_stop_enrollment", "7", "02105", "机械电子工程(先进制造与工业机器人方向)", "机械电子工程", "080204", "机电与自动化学院", "2016", "4", "工学", "已停招"),
        (2021, "wsyu_2021_stop_enrollment", "8", "03103", "环境工程", "环境工程", "082502", "城市建设学院", "2002", "4", "工学", "当年停招"),
        (2021, "wsyu_2021_stop_enrollment", "9", "03111", "道路桥梁与渡河工程", "道路桥梁与渡河工程", "081006T", "城市建设学院", "2015", "4", "工学", "已停招"),
        (2021, "wsyu_2021_stop_enrollment", "10", "05111", "金融学(金融智能创新实验班)", "金融学", "020301K", "经济管理学院", "2019", "4", "经济学", "当年停招"),
        (2021, "wsyu_2021_stop_enrollment", "11", "07106", "风景园林", "风景园林", "082803", "艺术设计学院", "2013", "4", "艺术学", "已停招"),
        (2020, "wsyu_2020_stop_enrollment", "1", "01115", "计算机科学与技术(华胜天成大数据方向)", "计算机科学与技术", "080901", "信息科学与工程学院", "2017", "4", "工学", "已停招"),
        (2020, "wsyu_2020_stop_enrollment", "2", "02105", "机械电子工程(先进制造与工业机器人方向)", "机械电子工程", "080204", "机电与自动化学院", "2016", "4", "工学", "已停招"),
        (2020, "wsyu_2020_stop_enrollment", "3", "03113", "土木工程(智能建筑信息化方向)", "土木工程", "081001", "城市建设学院", "2017", "4", "工学", "已停招"),
        (2020, "wsyu_2020_stop_enrollment", "4", "03111", "道路桥梁与渡河工程", "道路桥梁与渡河工程", "081006T", "城市建设学院", "2015", "4", "工学", "当年停招"),
        (2020, "wsyu_2020_stop_enrollment", "5", "07106", "风景园林", "风景园林", "082803", "艺术设计学院", "2013", "4", "艺术学", "已停招"),
        (2019, "wsyu_2019_stop_enrollment", "1", "01115", "计算机科学与技术（华胜天成大数据方向）", "计算机科学与技术", "080901", "信息科学与工程学院", "2016", "4", "工学", "当年停招"),
        (2019, "wsyu_2019_stop_enrollment", "2", "02101", "自动化", "自动化", "080801", "机电与自动化学院", "2001", "4", "工学", "当年停招"),
        (2019, "wsyu_2019_stop_enrollment", "3", "03104", "工程管理", "工程管理", "120103", "城市建设学院", "2004", "4", "工学", "当年停招"),
        (2019, "wsyu_2019_stop_enrollment", "4", "05102", "市场营销", "市场营销", "120202", "经济管理学院", "2002", "4", "管理学", "当年停招"),
        (2019, "wsyu_2019_stop_enrollment", "5", "06102", "新闻学", "新闻学", "050301", "新闻与法学学院", "2001", "4", "文学", "当年停招"),
        (2019, "wsyu_2019_stop_enrollment", "6", "06105", "广播电视学", "广播电视学", "050302", "新闻与法学学院", "2005", "4", "文学", "当年停招"),
        (2019, "wsyu_2019_stop_enrollment", "7", "07106", "风景园林", "风景园林", "082803", "艺术设计学院", "2013", "4", "艺术学", "当年停招"),
    ]
    for policy_year, source_id, source_row_no, internal_code, campus_major, base_major, code, school_unit, setup_year, duration, degree, status in wsyu_stop_rows:
        internal_phrase = f"校内专业代码{internal_code}，" if internal_code else ""
        rows.append(
            {
                "policy_year": policy_year,
                "region": "武昌首义学院",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": f"{policy_year}年停招专业名单",
                "reported_major_name": campus_major,
                "major_code": code,
                "study_duration": duration,
                "policy_action": f"{status}。",
                "criterion_text": f"武昌首义学院{policy_year}年停招专业名单HTML表“招生状态”列标记当年停招或已停招。",
                "source_row_no": source_row_no,
                "source_ids": source_id,
                "evidence_text": f"武昌首义学院{policy_year}年停招专业名单第{source_row_no}行列出{internal_phrase}校内专业名称{campus_major}（专业名称{base_major}，专业代码{code}，{school_unit}，设置年份{setup_year}，学制{duration}，授予{degree}学位），招生状态为{status}。",
                "confidence": "high",
            }
        )

    jlu_stop_rows = [
        (4, 2024, "劳动与社会保障"),
        (5, 2024, "应用心理学"),
        (11, 2024, "广播电视编导"),
        (15, 2024, "音乐表演"),
        (16, 2024, "作曲与作曲技术理论"),
        (18, 2024, "产品设计"),
        (20, 2024, "环境设计"),
        (21, 2024, "视觉传达设计"),
        (39, 2025, "信用管理"),
        (58, 2024, "药物制剂"),
        (60, 2023, "生物工程"),
        (66, 2022, "金融数学"),
        (69, 2024, "核物理"),
        (70, 2024, "光电信息科学与工程"),
        (75, 2024, "物流工程"),
        (76, 2024, "汽车服务工程"),
        (93, 2024, "电子信息工程"),
        (94, 2023, "生物医学工程"),
        (134, 2024, "园艺"),
    ]
    for source_row_no, stop_year, major in jlu_stop_rows:
        rows.append(
            {
                "policy_year": stop_year,
                "region": "吉林大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "吉林大学本科专业停招备注",
                "reported_major_name": major,
                "policy_action": "停招。",
                "criterion_text": f"吉林大学统计资料表格备注列标记“{stop_year}年停招”；页面说明19个专业已停招但目前仍有在校学生。",
                "source_row_no": str(source_row_no),
                "source_ids": "jlu_2026_undergrad_stop_enrollment",
                "evidence_text": f"吉林大学《本科专业141个》表格第{source_row_no}行列出{major}，备注为{stop_year}年停招。",
                "confidence": "high",
            }
        )

    add_warning_list(
        2025,
        "湖南大学",
        "本科",
        "major_cancel",
        "2025年度撤销专业",
        ["政治学与行政学", "视觉传达设计", "保密管理"],
        "撤销。",
        "学院申请、学校初审、学校教学委员会审议、校长办公会议审定。",
        ["hnu_2025_major_cancellation"],
        "湖南大学2025年度本科专业设置公示列出撤销政治学与行政学、视觉传达设计、保密管理3个本科专业。",
        "high",
    )
    hnu_stop_rows = [
        (2024, "hnu_2024_major_catalog_pdf", "湖南大学2024年本科专业目录当年停招专业", "12", "政治学与行政学", "030201", "四年"),
        (2024, "hnu_2024_major_catalog_pdf", "湖南大学2024年本科专业目录当年停招专业", "14", "运动训练", "040202K", "四年"),
        (2024, "hnu_2024_major_catalog_pdf", "湖南大学2024年本科专业目录当年停招专业", "31", "材料成型及控制工程", "080203", "四年"),
        (2024, "hnu_2024_major_catalog_pdf", "湖南大学2024年本科专业目录当年停招专业", "66", "风景园林", "082803", "五年"),
        (2024, "hnu_2024_major_catalog_pdf", "湖南大学2024年本科专业目录当年停招专业", "70", "保密管理", "120106TK", "四年"),
        (2024, "hnu_2024_major_catalog_pdf", "湖南大学2024年本科专业目录当年停招专业", "82", "视觉传达设计", "130502", "四年"),
        (2024, "hnu_2024_major_catalog_pdf", "湖南大学2024年本科专业目录当年停招专业", "83", "环境设计", "130503", "四年"),
        (2024, "hnu_2024_major_catalog_pdf", "湖南大学2024年本科专业目录当年停招专业", "84", "产品设计", "130504", "四年"),
        (2025, "hnu_2025_major_catalog_pdf", "湖南大学2025年本科专业目录当年停招专业", "7", "金融学", "020301K", "二年"),
        (2025, "hnu_2025_major_catalog_pdf", "湖南大学2025年本科专业目录当年停招专业", "8", "保险学", "020303", "四年"),
        (2025, "hnu_2025_major_catalog_pdf", "湖南大学2025年本科专业目录当年停招专业", "13", "法学", "030101K", "二年"),
        (2025, "hnu_2025_major_catalog_pdf", "湖南大学2025年本科专业目录当年停招专业", "14", "政治学与行政学", "030201", "四年"),
        (2025, "hnu_2025_major_catalog_pdf", "湖南大学2025年本科专业目录当年停招专业", "16", "运动训练", "040202K", "四年"),
        (2025, "hnu_2025_major_catalog_pdf", "湖南大学2025年本科专业目录当年停招专业", "34", "材料成型及控制工程", "080203", "四年"),
        (2025, "hnu_2025_major_catalog_pdf", "湖南大学2025年本科专业目录当年停招专业", "54", "计算机科学与技术", "080901", "二年"),
        (2025, "hnu_2025_major_catalog_pdf", "湖南大学2025年本科专业目录当年停招专业", "72", "风景园林", "082803", "五年"),
        (2025, "hnu_2025_major_catalog_pdf", "湖南大学2025年本科专业目录当年停招专业", "74", "信息管理与信息系统", "120102", "四年"),
        (2025, "hnu_2025_major_catalog_pdf", "湖南大学2025年本科专业目录当年停招专业", "76", "保密管理", "120106TK", "四年"),
        (2025, "hnu_2025_major_catalog_pdf", "湖南大学2025年本科专业目录当年停招专业", "85", "电子商务", "120801", "四年"),
        (2025, "hnu_2025_major_catalog_pdf", "湖南大学2025年本科专业目录当年停招专业", "88", "视觉传达设计", "130502", "四年"),
        (2025, "hnu_2025_major_catalog_pdf", "湖南大学2025年本科专业目录当年停招专业", "89", "环境设计", "130503", "四年"),
    ]
    for policy_year, source_id, warning_label, source_row_no, major, major_code, duration in hnu_stop_rows:
        rows.append(
            {
                "policy_year": policy_year,
                "region": "湖南大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": warning_label,
                "reported_major_name": major,
                "major_code": major_code,
                "study_duration": duration,
                "policy_action": "当年停招。",
                "criterion_text": f"湖南大学{policy_year}年本科专业目录表格备注列标注“当年停招”。",
                "source_row_no": source_row_no,
                "source_ids": source_id,
                "evidence_text": f"湖南大学{policy_year}年本科专业目录第{source_row_no}行列出{major}（{major_code}，{duration}），备注为当年停招。",
                "confidence": "high",
            }
        )

    ayit_stop_rows = [
        ("1", "信息管理与信息系统", "120102"),
        ("2", "城乡规划", "082802"),
        ("3", "材料成型及控制工程", "080203"),
        ("4", "工程管理", "120103"),
        ("5", "社会工作", "030302"),
        ("6", "商务英语", "050262"),
        ("7", "物联网工程", "080905"),
        ("8", "给排水科学与工程", "081003"),
        ("9", "高分子材料与工程", "080407"),
    ]
    for source_row_no, major, major_code in ayit_stop_rows:
        rows.append(
            {
                "policy_year": 2025,
                "region": "安阳工学院",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2025年停招本科专业一览表",
                "reported_major_name": major,
                "major_code": major_code,
                "policy_action": "停招。",
                "criterion_text": "安阳工学院信息公开网2025年停招本科专业一览表逐项列出停招本科专业。",
                "source_row_no": source_row_no,
                "source_ids": "ayit_2025_stop_enrollment",
                "evidence_text": f"安阳工学院2025年停招本科专业一览表第{source_row_no}行列出{major}（{major_code}）。",
                "confidence": "high",
            }
        )

    nwpu_stop_rows = [
        (2024, "nwpu_2024_major_setting_pdf", "西北工业大学2024年本科专业设置停招专业", "1", "金融科技", "020310T", "四年", "经济学"),
        (2024, "nwpu_2024_major_setting_pdf", "西北工业大学2024年本科专业设置停招专业", "22", "测控技术与仪器", "080301", "四年", "工学"),
        (2024, "nwpu_2024_major_setting_pdf", "西北工业大学2024年本科专业设置停招专业", "37", "电磁场与无线技术", "080712T", "四年", "工学"),
        (2024, "nwpu_2024_major_setting_pdf", "西北工业大学2024年本科专业设置停招专业", "51", "交通工程", "081802", "四年", "工学"),
        (2024, "nwpu_2024_major_setting_pdf", "西北工业大学2024年本科专业设置停招专业", "52", "交通设备与控制工程", "081806T", "四年", "工学"),
        (2024, "nwpu_2024_major_setting_pdf", "西北工业大学2024年本科专业设置停招专业", "59", "飞行器环境与生命保障工程", "082005", "四年", "工学"),
        (2024, "nwpu_2024_major_setting_pdf", "西北工业大学2024年本科专业设置停招专业", "60", "飞行器适航技术", "082007T", "四年", "工学"),
        (2024, "nwpu_2024_major_setting_pdf", "西北工业大学2024年本科专业设置停招专业", "71", "电子商务", "120801", "四年", "工学"),
        (2025, "nwpu_2025_major_setting_pdf", "西北工业大学2025年本科专业设置停招专业", "65", "金融科技", "020310T", "四年", "经济学"),
        (2025, "nwpu_2025_major_setting_pdf", "西北工业大学2025年本科专业设置停招专业", "66", "测控技术与仪器", "080301", "四年", "工学"),
        (2025, "nwpu_2025_major_setting_pdf", "西北工业大学2025年本科专业设置停招专业", "67", "电磁场与无线技术", "080712T", "四年", "工学"),
        (2025, "nwpu_2025_major_setting_pdf", "西北工业大学2025年本科专业设置停招专业", "68", "土木工程", "081001", "四年", "工学"),
        (2025, "nwpu_2025_major_setting_pdf", "西北工业大学2025年本科专业设置停招专业", "69", "交通工程", "081802", "四年", "工学"),
        (2025, "nwpu_2025_major_setting_pdf", "西北工业大学2025年本科专业设置停招专业", "70", "飞行器环境与生命保障工程", "082005", "四年", "工学"),
        (2025, "nwpu_2025_major_setting_pdf", "西北工业大学2025年本科专业设置停招专业", "71", "飞行器适航技术", "082007T", "四年", "工学"),
        (2025, "nwpu_2025_major_setting_pdf", "西北工业大学2025年本科专业设置停招专业", "72", "市场营销", "120202", "四年", "管理学"),
        (2025, "nwpu_2025_major_setting_pdf", "西北工业大学2025年本科专业设置停招专业", "73", "电子商务", "120801", "四年", "工学"),
    ]
    for policy_year, source_id, warning_label, source_row_no, major, major_code, duration, discipline in nwpu_stop_rows:
        rows.append(
            {
                "policy_year": policy_year,
                "region": "西北工业大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": warning_label,
                "reported_major_name": major,
                "major_code": major_code,
                "study_duration": duration,
                "policy_action": "停招。",
                "criterion_text": f"西北工业大学{policy_year}年本科专业设置PDF备注列标注“停招”；文件注释说明黄色为当年停招专业。",
                "source_row_no": source_row_no,
                "source_ids": source_id,
                "evidence_text": f"西北工业大学{policy_year}年本科专业设置PDF第{source_row_no}行列出{major}（{major_code}，{duration}，{discipline}），备注为停招。",
                "confidence": "high",
            }
        )

    add_warning(
        2023,
        "吉林工程技术师范学院",
        "本科",
        "major_cancel",
        "2023年拟撤销专业公示",
        "行政管理",
        "拟撤销。",
        "根据教育部和吉林省教育厅本科专业设置工作要求，结合学校专业建设实际，经学校专业建设委员会审议通过。",
        ["jlenu_2023_major_cancellation"],
        "吉林工程技术师范学院2023年拟撤销专业公示称，结合学校专业建设实际，经学校专业建设委员会审议通过，决定拟撤销行政管理专业。",
        "high",
    )

    jlenu_2024_stop_rows = [
        ("44", "电子商务", "120801"),
        ("45", "旅游管理", "120901K"),
        ("46", "航空服务艺术与管理", "130208TK"),
        ("47", "表演", "130301"),
        ("48", "动画", "130310"),
        ("49", "美术学", "130401"),
        ("50", "视觉传达设计", "130502"),
        ("51", "环境设计", "130503"),
        ("52", "服装与服饰设计", "130505"),
        ("53", "工艺美术", "130507"),
    ]
    for source_row_no, major, major_code in jlenu_2024_stop_rows:
        rows.append(
            {
                "policy_year": 2024,
                "region": "吉林工程技术师范学院",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2023-2024学年本科专业设置表停招专业",
                "reported_major_name": major,
                "major_code": major_code,
                "policy_action": "停招。",
                "criterion_text": "吉林工程技术师范学院2023-2024学年本科教学质量报告表2-1注释说明44-53项为停招专业。",
                "source_row_no": source_row_no,
                "source_ids": "jlenu_2024_teaching_quality_report_pdf",
                "evidence_text": f"吉林工程技术师范学院2023-2024学年本科教学质量报告表2-1第{source_row_no}项列出{major}（{major_code}）；表下注释说明44-53项为停招专业。",
                "confidence": "high",
            }
        )

    jlenu_2025_cancel_rows = [
        ("p24", "汽车服务工程", "080208"),
        ("p24", "人工智能", "080717T"),
        ("p24", "环境设计", "130503"),
        ("p24", "表演", "130301"),
    ]
    for source_row_no, major, major_code in jlenu_2025_cancel_rows:
        rows.append(
            {
                "policy_year": 2025,
                "region": "吉林工程技术师范学院",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "2024-2025学年本科教学质量报告撤销专业",
                "reported_major_name": major,
                "major_code": major_code,
                "policy_action": "撤销。",
                "criterion_text": "吉林工程技术师范学院2024-2025学年本科教学质量报告称以社会需求为导向实施专业增、减、并、撤动态调整。",
                "source_row_no": source_row_no,
                "source_ids": "jlenu_2025_teaching_quality_report_pdf",
                "evidence_text": f"吉林工程技术师范学院2024-2025学年本科教学质量报告专业建设成效部分称，2025年撤销汽车服务工程、人工智能、环境设计、表演4个专业；本行对应{major}（{major_code}）。",
                "confidence": "high",
            }
        )

    add_warning_list(
        2025,
        "江西财经大学",
        "本科",
        "major_cancel",
        "2025年度撤销专业",
        ["文化产业管理", "广告学"],
        "撤销。",
        "连续五年停止招生且无在校生。",
        ["jxufe_2025_major_adjustment"],
        "江西财经大学2025年普通本科专业设置动态调整名单公示列出撤销连续五年停止招生且无在校生的文化产业管理、广告学2个专业。",
        "high",
    )
    add_warning_list(
        2025,
        "江西财经大学",
        "本科",
        "major_stop_enrollment",
        "2025年度新增停招专业",
        ["国民经济管理", "土地资源管理", "汉语国际教育"],
        "新增停招。",
        "学校本科教学指导委员会评议专家组评议、校长办公会审定。",
        ["jxufe_2025_major_adjustment"],
        "江西财经大学2025年普通本科专业设置动态调整名单公示列出新增停招国民经济管理、土地资源管理、汉语国际教育3个专业。",
        "high",
    )
    add_warning_list(
        2025,
        "江西财经大学",
        "本科",
        "major_stop_enrollment",
        "2025年度继续停招专业",
        ["房地产开发与管理", "日语"],
        "继续停招。",
        "学校本科教学指导委员会评议专家组评议、校长办公会审定。",
        ["jxufe_2025_major_adjustment"],
        "江西财经大学2025年普通本科专业设置动态调整名单公示列出继续停招房地产开发与管理、日语2个专业。",
        "high",
    )

    add_warning_list(
        2024,
        "湖北恩施学院",
        "本科",
        "major_cancel",
        "2024年度拟撤销专业",
        ["市场营销", "社会体育指导与管理", "电子信息科学与技术", "机械电子工程", "广播电视学", "卫生检验与检疫", "茶学", "广播电视编导"],
        "拟撤销。",
        "连续停招五年以上；学校专业建设委员会评议专家审议通过。",
        ["hbesxy_2024_major_cancellation"],
        "湖北恩施学院2024年度拟撤销本科专业公示列出拟撤销市场营销、社会体育指导与管理、电子信息科学与技术、机械电子工程、广播电视学、卫生检验与检疫、茶学、广播电视编导8个连续停招五年以上的本科专业。",
        "high",
    )

    add_warning(
        2025,
        "江西农业大学",
        "本科",
        "major_cancel",
        "2025年度拟撤销专业",
        "信息与计算科学",
        "拟撤销。",
        "学院申请、教务处审核、校长办公会审议通过。",
        ["jxau_2025_major_cancellation"],
        "江西农业大学2025年度本科专业设置申报材料公示列出拟撤销信息与计算科学本科专业。",
        "high",
    )

    add_warning(
        2025,
        "西安邮电大学",
        "本科",
        "major_cancel",
        "2025年度撤销专业",
        "广播电视工程",
        "撤销。",
        "学院申报、评议专家组评审、校教学委员会审议通过。",
        ["xupt_2025_major_cancellation"],
        "西安邮电大学2025年度调整本科专业情况公示列出广播电视工程调整情况为撤销。",
        "high",
    )

    add_warning_list(
        2025,
        "辽宁大学",
        "本科",
        "major_cancel",
        "2025年度拟撤销专业",
        ["信息与计算科学", "摄影", "广播电视学", "网络与新媒体"],
        "拟撤销。",
        "学院申报、学校组织专家论证、校学术委员会审议通过、校长办公会审定；广播电视学和网络与新媒体为第二学士学位。",
        ["lnu_2025_major_cancellation"],
        "辽宁大学2025年度拟增设本科专业、拟撤销本科专业公示列出拟撤销信息与计算科学、摄影、广播电视学（第二学士学位）、网络与新媒体（第二学士学位）4个本科专业。",
        "high",
    )

    add_warning_list(
        2024,
        "辽宁大学",
        "本科",
        "major_cancel",
        "2024年度拟撤销本科专业",
        ["广播电视学", "网络与新媒体", "生物科学", "生态学", "材料化学", "电子科学与技术", "环境科学", "信息管理与信息系统", "旅游管理"],
        "拟撤销。",
        "经学院申报、学校组织专家论证、校学术委员会审议通过。",
        ["lnu_2024_major_cancellation"],
        "辽宁大学关于申报2024年度本科专业的公示列出拟撤销广播电视学、网络与新媒体、生物科学、生态学、材料化学、电子科学与技术、环境科学、信息管理与信息系统、旅游管理9个本科专业。",
        "high",
    )

    add_warning(
        2024,
        "沧州交通学院",
        "本科",
        "major_cancel",
        "2024年度拟撤销本科专业",
        "过程装备与控制工程",
        "拟撤销。",
        "经化学与制药工程学院申请、学术委员会审定，学校决定撤销并向教育部申请撤销备案。",
        ["czjtu_2024_major_cancellation"],
        "沧州交通学院关于拟撤销过程装备与控制工程专业的公示说明，学校决定撤销过程装备与控制工程专业，并向教育部申请撤销备案。",
        "high",
    )

    add_warning_list(
        2025,
        "广西科技大学",
        "本科",
        "major_cancel",
        "2025年度撤销专业",
        ["材料科学与工程", "微电子科学与工程", "数字媒体技术", "市场营销", "人力资源管理", "物流工程", "信息与计算科学"],
        "撤销。",
        "学校组织专家评议论证，并报学校教学指导委员会审议、校长办公会审定。",
        ["gxust_2025_major_cancellation"],
        "广西科技大学2025年度普通本科专业设置调整情况公示列出撤销材料科学与工程、微电子科学与工程、数字媒体技术、市场营销、人力资源管理、物流工程、信息与计算科学7个专业。",
        "high",
    )

    add_warning_list(
        2025,
        "北京信息科技大学",
        "本科",
        "major_cancel",
        "2025年度拟撤销专业",
        ["英语", "行政管理"],
        "拟撤销。",
        "学院申报、校外专家论证、学校本科教学工作委员会审议、校长办公会审定。",
        ["bistu_2025_major_cancellation"],
        "北京信息科技大学2025年度拟新增本科专业、预备案专业、拟撤销本科专业公示列出拟撤销英语、行政管理2个本科专业。",
        "high",
    )

    add_warning_list(
        2025,
        "河南工业大学",
        "本科",
        "major_cancel",
        "2025年度拟撤销专业",
        ["食品营养与检验教育", "交通工程", "电子信息科学与技术"],
        "拟撤销。",
        "学院申报和论证，学校专业设置评议专家组评议。",
        ["haut_2025_major_cancellation"],
        "河南工业大学2025年本科新专业申报、预备案和撤销专业公示列出拟撤销食品营养与检验教育、交通工程、电子信息科学与技术3个专业。",
        "high",
    )

    add_warning_list(
        2025,
        "武汉工程大学",
        "本科",
        "major_cancel",
        "2025年度拟申请撤销专业",
        ["工程力学", "城市地下空间工程", "道路桥梁与渡河工程", "公共事业管理", "广告学"],
        "拟向教育部申请撤销。",
        "学校组织2025年本科专业设置评审工作，经评审决定。",
        ["wit_2025_major_cancellation"],
        "武汉工程大学2025年度拟增设、撤销及调整本科专业公示列出拟向教育部申请撤销工程力学、城市地下空间工程、道路桥梁与渡河工程、公共事业管理、广告学5个本科专业。",
        "high",
    )

    add_warning_list(
        2025,
        "辽宁师范大学",
        "本科",
        "major_cancel",
        "2025年度拟撤销专业",
        ["绘画", "环境科学", "应用化学"],
        "拟撤销。",
        "根据2025年度普通高等学校本科专业设置工作要求，学校拟申报新增及撤销本科专业。",
        ["lnnu_2025_major_cancellation"],
        "辽宁师范大学2025年度拟申报新增及撤销本科专业公示列出拟撤销绘画、环境科学、应用化学3个专业。",
        "high",
    )

    add_warning_list(
        2025,
        "山西电子科技学院",
        "本科",
        "major_cancel",
        "2025年度拟撤销专业",
        ["思想政治教育", "学前教育", "汉语言文学", "戏剧影视文学", "广播电视编导", "数字媒体艺术"],
        "拟撤销。",
        "学校组织开展2025年度专业设置工作，并公示拟申请撤销专业。",
        ["sxdzkj_2025_major_cancellation"],
        "山西电子科技学院2025年度拟增设、调整和撤销专业公示列出拟撤销思想政治教育、学前教育、汉语言文学、戏剧影视文学、广播电视编导、数字媒体艺术6个专业。",
        "high",
    )

    ahstu_cancel_rows = [
        ("7", "120401", "公共事业管理", "4", "工学", "管理学院"),
        ("8", "090106", "设施农业科学与工程", "4", "工学", "农学院"),
    ]
    for source_row_no, code, major, study_duration, degree, school_unit in ahstu_cancel_rows:
        rows.append(
            {
                "policy_year": 2025,
                "region": "安徽科技学院",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "2025年度拟撤销专业",
                "reported_major_name": major,
                "major_code": code,
                "study_duration": study_duration,
                "policy_action": "撤销。",
                "criterion_text": "经学院申报、专家论证、学校教学指导委员会审定通过，学校公示2025年度撤销专业。",
                "source_row_no": source_row_no,
                "source_ids": "ahstu_2025_major_adjustment",
                "evidence_text": f"安徽科技学院2025年度预申报、撤销和调整修业年限本科专业公示表第{source_row_no}项列出{major}（{code}，修业年限{study_duration}，学位授予门类{degree}，{school_unit}），申报类型为撤销。",
                "confidence": "high",
            }
        )

    add_warning(
        2025,
        "桂林信息科技学院",
        "本科",
        "major_cancel",
        "2025年度申请撤销专业",
        "材料成型及控制工程",
        "申请撤销。",
        "材料成型及控制工程专业已停招5年以上，且无在校生。",
        ["guit_2025_major_cancellation"],
        "桂林信息科技学院2025年度拟新设及撤销本科专业公示称材料成型及控制工程已停招5年以上且无在校生，本年度将申请撤销。",
        "high",
    )

    add_warning_list(
        2025,
        "西安文理学院",
        "本科",
        "major_cancel",
        "2025年度拟撤销专业",
        ["电子商务", "材料科学与工程", "音乐表演", "公共艺术"],
        "拟撤销。",
        "结合学校办学定位、发展规划以及专业建设实际情况，经学校校长办公会审议。",
        ["xawl_2025_major_cancellation"],
        "西安文理学院2025年度本科新专业申报结果公示列出拟撤销电子商务、材料科学与工程、音乐表演、公共艺术4个专业。",
        "high",
    )
    add_warning_list(
        2025,
        "西藏大学",
        "本科",
        "major_cancel",
        "2025年度拟撤销本科专业",
        ["市场营销", "财务管理", "服装与服饰设计", "新闻学"],
        "拟撤销。",
        "经学院申报、教务处初审、学校本科专业调整优化工作专题会讨论、学校本科教学指导委员会审议、校长办公会议审定。",
        ["utibet_2025_major_adjustment"],
        "西藏大学2025年度拟设置与调整本科专业公示列出拟撤销本科专业：市场营销、财务管理、服装与服饰设计、新闻学。",
        "high",
    )
    add_warning(
        2026,
        "西安文理学院",
        "本科",
        "major_cancel",
        "2026年度拟撤销专业",
        "化学工程与工艺",
        "拟撤销。",
        "结合学校办学定位、发展规划以及专业建设实际情况，经学校校长办公会审议。",
        ["xawl_2026_major_cancellation"],
        "西安文理学院2026年度本科新专业申报结果公示列出拟撤销化学工程与工艺专业。",
        "high",
    )

    add_warning(
        2026,
        "咸阳师范学院",
        "本科",
        "major_cancel",
        "2026年度拟撤销本科专业",
        "视觉传达设计",
        "拟撤销。",
        "经二级学院申报、校教学指导委员会评议，学校公示2026年度拟撤销专业。",
        ["xync_2026_major_adjustment"],
        "咸阳师范学院2026年度拟新增、拟撤销、预申报本科专业公示称拟撤销视觉传达设计专业。",
        "high",
    )

    cuggw_cancel_rows = [
        ("3", "机械电子工程", "080204", "工学", "资源与工程技术学院"),
        ("4", "建筑学", "082801", "工学", "资源与工程技术学院"),
    ]
    for source_row_no, major, code, degree, school_unit in cuggw_cancel_rows:
        rows.append(
            {
                "policy_year": 2024,
                "region": "保定理工学院",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "2024年度拟撤销本科专业",
                "reported_major_name": major,
                "major_code": code,
                "study_duration": "",
                "policy_action": "拟撤销。",
                "criterion_text": "从学校建设规划、学科专业结构、学生生源及就业前景等方面综合考虑，经学校教学指导委员会审议。",
                "source_row_no": source_row_no,
                "source_ids": "cuggw_2024_major_adjustment",
                "evidence_text": f"保定理工学院2024年度新增、撤销本科专业情况公示表第{source_row_no}项列出{major}（{code}，{degree}，{school_unit}），调整类型为撤销。",
                "confidence": "high",
            }
        )

    changdian_2023_cancel_rows = [
        ("光源与照明", "080603T", "2021年停止招生"),
        ("网络工程", "080903", "2020年停止招生"),
    ]
    for source_row_no, (major, code, stop_note) in enumerate(changdian_2023_cancel_rows, start=1):
        rows.append(
            {
                "policy_year": 2023,
                "region": "长春电子科技学院",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "2023年度拟撤销专业",
                "reported_major_name": major,
                "major_code": code,
                "study_duration": "",
                "policy_action": "拟撤销。",
                "criterion_text": "结合吉林省经济与社会发展方向及学校办学实际，拟撤销已停止招生且目前无在校生的专业。",
                "source_row_no": str(source_row_no),
                "source_ids": "changdian_2023_major_cancellation",
                "evidence_text": f"长春电子科技学院2023年度拟撤销专业公示列出拟撤销{major}（{code}），并说明该专业{stop_note}，目前已无在校生。",
                "confidence": "high",
            }
        )

    add_warning(
        2025,
        "衡阳师范学院",
        "本科",
        "major_cancel",
        "2025年度拟撤销专业",
        "编辑出版学",
        "撤销。",
        "学院申请、教务处组织相关部门论证、教学工作委员会审议。",
        ["hynu_2025_major_cancellation"],
        "衡阳师范学院2025年本科专业设置调整情况公示列出撤销编辑出版学专业。",
        "high",
    )

    add_warning(
        2025,
        "龙岩学院",
        "本科",
        "major_cancel",
        "2025年度拟申请撤销专业",
        "日语",
        "拟申请撤销。",
        "根据2025年度普通高等学校本科专业设置工作要求，学校拟申请撤销日语专业。",
        ["lyun_2025_major_cancellation"],
        "龙岩学院2025年度本科专业设置公示列出拟申请撤销日语专业。",
        "high",
    )

    add_warning_list(
        2025,
        "河南农业大学",
        "本科",
        "major_cancel",
        "2025年度拟撤销专业",
        ["应用生物科学", "草业科学", "市场营销", "植物科学与技术", "食品营养与检验教育", "资源循环科学与工程"],
        "拟撤销。",
        "根据学校发展规划，经学院申报、校本科专业设置与优化调整管理领导小组专题会议审核、校长办公会议题意见征求等程序。",
        ["henau_2025_major_cancellation"],
        "河南农业大学2025年度本科专业设置及专业结构调整优化公示列出拟撤销应用生物科学、草业科学、市场营销、植物科学与技术、食品营养与检验教育、资源循环科学与工程6个专业。",
        "high",
    )

    add_warning(
        2025,
        "江西服装学院",
        "本科",
        "major_cancel",
        "2025年度拟撤销专业",
        "国际商务",
        "拟撤销。",
        "学校按要求针对撤销国际商务专业征求校内意见无异议后拟撤销。",
        ["jift_2025_major_cancellation"],
        "江西服装学院2025年拟新增普通本科专业及拟撤销国际商务专业公示列出拟撤销国际商务普通本科专业。",
        "high",
    )

    add_warning(
        2025,
        "辽宁师范大学海华学院",
        "本科",
        "major_cancel",
        "2025年度撤销本科专业",
        "电子商务及法律",
        "撤销。",
        "按照2025年度普通高等学校本科专业设置工作要求，根据学校发展需要，经研究决定撤销。",
        ["lshhxy_2025_major_cancellation"],
        "辽宁师范大学海华学院2025年度撤销1个本科专业公示列出撤销电子商务及法律（120802T）专业。",
        "high",
    )

    add_warning_list(
        2025,
        "桂林学院",
        "本科",
        "major_cancel",
        "2025年度拟撤销专业名单",
        ["互联网金融", "资产评估", "数字出版"],
        "拟撤销。",
        "桂林学院拟撤销专业名单Word附件列入普通高等学校本科专业和专业方向撤销汇总表。",
        ["glc_2025_major_adjustment_notice", "glc_2025_major_cancellation_docx"],
        "桂林学院2025年度拟申报专业材料公示页链接的《2025年桂林学院拟撤销专业名单》列出互联网金融、资产评估、数字出版3个本科专业。",
        "high",
    )

    add_warning_list(
        2025,
        "桂林学院",
        "本科",
        "major_stop_enrollment",
        "2025年度拟停招专业名单",
        ["经济学", "保险学", "投资学", "会计学", "播音与主持艺术", "酒店管理", "会展经济与管理", "工艺美术"],
        "拟停招。",
        "桂林学院拟停招专业名单Word附件列出停招原因，涉及师资力量、专业资源、招生质量和就业困难等。",
        ["glc_2025_major_adjustment_notice", "glc_2025_stop_enrollment_docx"],
        "桂林学院2025年度拟申报专业材料公示页链接的《2025年拟停招专业名单》列出经济学、保险学、投资学、会计学、播音与主持艺术、酒店管理、会展经济与管理、工艺美术8个本科专业。",
        "high",
    )

    add_warning_list(
        2024,
        "江汉大学",
        "本科",
        "major_cancel",
        "2024年度拟撤销专业",
        ["思想政治教育", "翻译", "汽车服务工程", "光电信息科学与工程", "文化产业管理", "播音与主持艺术", "测控技术与仪器"],
        "拟撤销。",
        "学校教学指导委员会对2024年度申请撤销专业组织评审。",
        ["jhuni_2024_major_cancellation"],
        "江汉大学教学指导委员会2024年度专业设置调整评议结果公示列出拟撤销思想政治教育、翻译、汽车服务工程、光电信息科学与工程、文化产业管理、播音与主持艺术、测控技术与仪器7个专业。",
        "high",
    )

    add_warning(
        2025,
        "徐州医科大学",
        "本科",
        "major_cancel",
        "2025年度拟撤销专业",
        "物联网工程",
        "拟撤销。",
        "相关学院申报、专家论证、校长办公会审议通过。",
        ["xzhmu_2025_major_cancellation"],
        "徐州医科大学2025年拟新增本科专业、预备案本科专业、撤销专业公示列出拟撤销物联网工程本科专业。",
        "high",
    )

    add_warning_list(
        2024,
        "上海交通大学",
        "本科",
        "major_cancel",
        "2024年度拟撤销专业列表",
        ["资源环境科学", "公共事业管理", "软件工程", "信息安全"],
        "拟撤销。",
        "已停招5年及以上；原表中软件工程、信息安全注明为二年制。",
        ["sjtu_2024_major_cancellation_pdf"],
        "上海交通大学2024年度拟撤销专业列表PDF列出资源环境科学、公共事业管理、软件工程（二年制）、信息安全（二年制）4个专业，撤销原因均为已停招5年及以上。",
        "high",
    )

    add_warning_list(
        2025,
        "湖北第二师范学院",
        "本科",
        "major_cancel",
        "2025年度拟撤销专业",
        ["信息与计算科学", "物流工程", "公共事业管理", "汽车服务工程"],
        "拟撤销。",
        "学院申报、专家论证、校长办公会审议通过。",
        ["hue_2025_major_cancellation"],
        "湖北第二师范学院2025年度拟新增本科专业、预备案本科专业、撤销专业公示列出拟撤销信息与计算科学、物流工程、公共事业管理、汽车服务工程4个专业。",
        "high",
    )

    add_warning_list(
        2025,
        "南京中医药大学",
        "本科",
        "major_cancel",
        "2025年度拟撤销本科专业",
        ["生物技术", "劳动与社会保障", "市场营销", "健康服务与管理"],
        "拟撤销。",
        "学院申请、校教学委员会审议、校长办公会通过。",
        ["njucm_2025_major_cancellation"],
        "南京中医药大学教务处公示列出2025年学校拟撤销生物技术、劳动与社会保障、市场营销、健康服务与管理4个本科专业。",
        "high",
    )

    add_warning_list(
        2025,
        "长春大学",
        "本科",
        "major_cancel",
        "2025年度拟撤销专业",
        ["应用统计学", "汽车服务工程", "文化产业管理"],
        "拟撤销。",
        "二级汇总来源称学校已决定拟撤销对应本科专业。",
        ["sohu_2025_ten_university_cancellation_summary", "gz55zs_2025_university_adjustment_table"],
        "搜狐转载高考直通车和新期教育网表格汇总均称长春大学2025年拟撤销应用统计学、汽车服务工程、文化产业管理3个本科专业。",
        "medium",
    )
    add_warning_list(
        2025,
        "成都信息工程大学",
        "本科",
        "major_cancel",
        "2025年度拟撤销专业",
        ["物流工程", "信息管理与信息系统"],
        "拟撤销。",
        "二级汇总来源称学校本科专业拟设置情况公示列出拟撤销专业。",
        [
            "sohu_2025_ten_university_cancellation_summary",
            "gkztc_2025_multi_university_adjustment",
            "eol_2025_multi_university_cancellation",
            "cqnews_eol_2025_multi_university_cancellation",
            "qq_zsgk_2025_multi_university_adjustment",
            "gz55zs_2025_university_adjustment_table",
        ],
        "搜狐转载高考直通车、教育在线、华龙网转载中国教育在线、腾讯新闻转载掌上高考、高考直通车和新期教育网表格均汇总称成都信息工程大学2025年本科专业拟设置情况公示列出拟撤销物流工程、信息管理与信息系统2个专业。",
        "medium",
    )
    add_warning(
        2025,
        "成都信息工程大学",
        "本科",
        "major_stop_enrollment",
        "2025年度拟停招专业",
        "翻译",
        "拟停招。",
        "二级汇总来源称学校本科专业拟设置情况公示列出拟停招专业。",
        [
            "sohu_2025_ten_university_cancellation_summary",
            "gkztc_2025_multi_university_adjustment",
            "eol_2025_multi_university_cancellation",
            "cqnews_eol_2025_multi_university_cancellation",
            "qq_zsgk_2025_multi_university_adjustment",
        ],
        "搜狐转载高考直通车、教育在线、华龙网转载中国教育在线、腾讯新闻转载掌上高考和高考直通车原文均汇总称成都信息工程大学2025年本科专业拟设置情况公示列出拟停招翻译专业。",
        "medium",
    )
    add_warning(
        2025,
        "西华师范大学",
        "本科",
        "major_cancel",
        "2025年拟调整本科专业",
        "健康服务与管理",
        "拟撤销。",
        "二级汇总来源称学校拟调整本科专业公示列出撤销本科专业1个。",
        [
            "eol_2025_multi_university_cancellation",
            "gkztc_2025_multi_university_adjustment",
            "qq_zsgk_2025_multi_university_adjustment",
            "gz55zs_2025_university_adjustment_table",
        ],
        "中国教育在线、高考直通车、腾讯新闻转载掌上高考和新期教育网表格均汇总称西华师范大学2025年拟撤销本科专业1个：健康服务与管理。",
        "medium",
    )
    add_warning_list(
        2025,
        "西华师范大学",
        "本科",
        "major_stop_enrollment",
        "2025年停招专业",
        ["测绘工程", "公共事业管理", "社会体育指导与管理", "园林"],
        "停招。",
        "二级汇总来源称学校2025年拟调整本科专业公示列出停招专业。",
        [
            "eol_2025_multi_university_cancellation",
            "gkztc_2025_multi_university_adjustment",
            "qq_zsgk_2025_multi_university_adjustment",
        ],
        "中国教育在线、高考直通车和腾讯新闻转载掌上高考均汇总称西华师范大学2025年停招测绘工程、公共事业管理、社会体育指导与管理、园林4个专业。",
        "medium",
    )
    add_warning_list(
        2025,
        "江汉大学",
        "本科",
        "major_cancel",
        "2025年度专业设置调整评议结果",
        ["市场营销", "财务管理", "通信工程", "网络工程", "过程装备与控制工程", "工业设计"],
        "申请撤销。",
        "学院申请、学校初审、学校教学指导委员会评议通过、校长办公会议审定。",
        ["jhun_2025_major_adjustment"],
        "江汉大学教务处公示列出2025年度申请撤销专业：市场营销、财务管理、通信工程、网络工程、过程装备与控制工程、工业设计。",
        "high",
    )
    add_warning_list(
        2025,
        "江汉大学",
        "本科",
        "major_cancel",
        "2025年度拟撤销专业",
        ["市场营销", "财务管理", "通信工程", "网络工程", "过程装备与控制工程", "工业设计"],
        "拟撤销。",
        "二级汇总来源称学校2025年度专业设置调整评议结果公示列出拟撤销专业。",
        ["sohu_2025_ten_university_cancellation_summary"],
        "搜狐转载高考直通车汇总称江汉大学2025年度专业设置调整评议结果公示列出拟撤销市场营销、财务管理、通信工程、网络工程、过程装备与控制工程、工业设计6个专业。",
        "medium",
    )
    add_warning_list(
        2025,
        "泉州信息工程学院",
        "本科",
        "major_cancel",
        "2025年度拟撤销专业",
        ["建筑电气与智能化", "材料成型及控制工程", "光电信息科学与工程", "投资学", "汽车服务工程"],
        "拟撤销。",
        "官网通知页确认学校公布2025年度撤销部分停招本科专业事项；当前公开页面未暴露正文名单，专业名来自二级汇总来源。",
        ["qzuie_2025_major_cancellation_notice", "sohu_2025_ten_university_cancellation_summary", "gz55zs_2025_university_adjustment_table"],
        "泉州信息工程学院教务处官网保留《关于公布2025年度撤销部分停招本科专业的通知》标题、来源和时间；搜狐转载高考直通车和新期教育网表格均称该通知列出拟撤销建筑电气与智能化、材料成型及控制工程、光电信息科学与工程、投资学、汽车服务工程5个专业。",
        "medium",
    )

    gsau_cancel_rows = [
        (2023, "农艺教育", "090110T", "1994", "1998", "农学", "农学院", "2023年撤销"),
        (2023, "园艺教育", "090111T", "1994", "1999", "农学", "园艺学院", "2023年撤销"),
        (2023, "交通运输", "081801", "2002", "2017", "工学", "机电工程学院", "2023年撤销"),
        (2025, "法学", "030101K", "2004", "2020", "法学", "人文学院", "2025年撤销"),
        (2025, "农业电气化", "082303", "2002", "2015", "工学", "机电工程学院", "2025年撤销"),
        (2025, "农村区域发展", "120302", "2001", "2017", "管理学", "管理学院", "2025年撤销"),
    ]
    for row_no, (policy_year, major, code, opened, stopped, degree, school_unit, note) in enumerate(gsau_cancel_rows, start=1):
        add_school_table_warning(
            policy_year,
            "major_cancel",
            "甘肃农业大学撤销本科专业名单",
            major,
            code,
            opened,
            stopped,
            degree,
            school_unit,
            "gsau_undergrad_cancel_list",
            row_no,
            "已撤销。",
            note,
        )

    gsau_stop_rows = [
        (2005, "农业建筑环境与能源工程", "082304", "2000", "2005", "工学", "水利水电工程学院", ""),
        (2017, "农村区域发展", "120302", "2001", "2017", "管理学", "管理学院", ""),
        (2018, "农业电气化", "082303", "2002", "2018", "工学", "机电工程学院", ""),
        (2020, "动物医学（四年制）", "090401", "1946", "2020", "农学", "动物医学院", "2020年开始招收动物医学（五年制）"),
        (2020, "法学", "030101K", "2004", "2020", "法学", "人文学院", ""),
        (2024, "植物科学与技术", "090104", "2016", "2024", "农学", "农学院", ""),
        (2024, "土地整治工程", "082306T", "2017", "2024", "工学", "资源与环境学院", ""),
        (2024, "公共事业管理", "120401", "2003", "2024", "管理学", "管理学院", ""),
        (2024, "信息管理与信息系统", "080701", "2004", "2024", "管理学", "信息科学技术学院", ""),
    ]
    for row_no, (policy_year, major, code, opened, stopped, degree, school_unit, note) in enumerate(gsau_stop_rows, start=1):
        add_school_table_warning(
            policy_year,
            "major_stop_enrollment",
            "甘肃农业大学停招本科专业名单",
            major,
            code,
            opened,
            stopped,
            degree,
            school_unit,
            "gsau_undergrad_stop_list",
            row_no,
            "已停招。",
            note,
        )

    add_warning_list(
        2025,
        "甘肃农业大学",
        "本科",
        "major_cancel",
        "2025年度拟撤销专业",
        ["农业电气化", "农村区域发展", "法学"],
        "拟撤销。",
        "二级汇总来源称学校申请新增专业、撤销专业评审结果公示列出拟撤销专业。",
        ["sohu_2025_ten_university_cancellation_summary"],
        "搜狐转载高考直通车汇总称甘肃农业大学2025年申请新增专业、撤销专业评审结果公示列出拟撤销农业电气化、农村区域发展、法学3个专业。",
        "medium",
    )
    add_warning_list(
        2025,
        "青岛电影学院",
        "本科",
        "major_cancel",
        "2025年度拟撤销专业",
        ["视觉传达设计", "音乐表演", "流行音乐"],
        "拟向教育部申请撤销。",
        "二级学院申报、本科教学指导委员会审议、校长办公会审定。",
        ["qdfa_2025_major_cancellation"],
        "青岛电影学院官网通知列表页展示的2025年度拟撤销本科专业公示正文列出拟向教育部申请撤销视觉传达设计、音乐表演、流行音乐3个本科专业。",
        "high",
    )
    add_warning_list(
        2025,
        "文山学院",
        "本科",
        "major_cancel",
        "2025年拟撤销本科专业",
        ["食品质量与安全", "工程管理", "视觉传达设计"],
        "拟向教育部申请撤销。",
        "连续五年停招且无在校学生，经充分论证、征求意见、学校教学指导委员会和校长办公会审议通过。",
        ["wsu_2025_major_cancellation"],
        "文山学院官网公示列出2025年拟向教育部申请撤销食品质量与安全、工程管理、视觉传达设计3个本科专业，并列明专业代码、修业年限和学位门类。",
        "high",
    )
    add_warning_list(
        2024,
        "江西理工大学",
        "本科",
        "major_cancel",
        "2024年度拟撤销本科专业",
        ["物联网工程", "物流管理", "电子商务", "视觉传达设计"],
        "拟撤销。",
        "学院申报、学校组织专家论证、学校教学指导委员会评审、校长办公会通过。",
        ["jxust_2024_major_cancellation"],
        "江西理工大学2024年拟新增专业及拟撤销本科专业公示列出拟撤销物联网工程、物流管理、电子商务、视觉传达设计4个本科专业。",
        "high",
    )

    add_warning_list(
        2025,
        "宁波大学",
        "本科",
        "major_cancel",
        "2025年度申请撤销备案专业",
        ["财务会计教育", "新闻学", "广告学", "旅游管理与服务教育", "产品设计", "物联网"],
        "申请撤销备案。",
        "宁波大学教务处公示称对停招五年及以上且无在籍学生的六个专业申请撤销备案。",
        [
            "nbu_2025_major_cancellation",
            "qq_eol_2025_multi_university_cancellation",
            "eol_2025_multi_university_cancellation",
            "gkztc_2025_multi_university_adjustment",
        ],
        "宁波大学教务处《关于撤销与新增设若干专业的公示》列出对财务会计教育、新闻学、广告学、旅游管理与服务教育、产品设计、物联网六个停招五年及以上且无在籍学生的专业申请撤销备案；中国教育在线和高考直通车汇总与官网名单一致。",
        "high",
    )
    add_warning(
        2025,
        "重庆师范大学",
        "本科",
        "major_cancel",
        "2025年度拟撤销专业",
        "信息管理与信息系统",
        "拟撤销。",
        "重庆师范大学官方新闻称教务处向教学指导委员会汇报拟撤销信息管理与信息系统专业。",
        [
            "cqnu_2025_teaching_committee_major_adjustment",
            "qq_eol_2025_multi_university_cancellation",
            "eol_2025_multi_university_cancellation",
            "gkztc_2025_multi_university_adjustment",
        ],
        "重庆师范大学官方新闻《学校召开第四届教学指导委员会第二次工作会》披露会议审议2025年拟新设专业、第二学士学位专业、撤销专业，并称教务处负责人就拟撤销信息管理与信息系统专业相关情况进行了汇报；中国教育在线和高考直通车汇总与官方新闻一致。",
        "high",
    )
    add_warning_list(
        2025,
        "井冈山大学",
        "本科",
        "major_cancel",
        "2025年度拟撤销专业",
        ["商务英语", "应用物理学", "动画"],
        "拟撤销。",
        "二级汇总来源称学院申报、专家论证、学校学术委员会评议、校长办公会审议。",
        [
            "qq_eol_2025_multi_university_cancellation",
            "eol_2025_multi_university_cancellation",
        ],
        "中国教育在线和腾讯新闻转载中国教育在线均汇总称井冈山大学2025年拟撤销商务英语、应用物理学、动画3个专业。",
        "medium",
    )
    add_warning(
        2025,
        "马鞍山学院",
        "本科",
        "major_cancel",
        "2025年度拟申请撤销专业",
        "体育经济与管理",
        "拟申请撤销。",
        "二级汇总来源称学院申报、教务处形式审查、校教学工作委员会审议通过。",
        [
            "qq_eol_2025_multi_university_cancellation",
            "eol_2025_multi_university_cancellation",
            "cqnews_eol_2025_multi_university_cancellation",
            "gz55zs_2025_university_adjustment_table",
        ],
        "中国教育在线、腾讯新闻转载中国教育在线、华龙网转载中国教育在线和新期教育网表格均汇总称马鞍山学院2025年拟申请撤销体育经济与管理专业。",
        "medium",
    )
    add_warning(
        2025,
        "合肥师范学院",
        "本科",
        "major_cancel",
        "2025年度拟撤销专业",
        "视觉传达设计（中外合作）",
        "拟撤销。",
        "二级汇总来源称二级学院论证、申报，学校审议。",
        [
            "qq_eol_2025_multi_university_cancellation",
            "eol_2025_multi_university_cancellation",
            "cqnews_eol_2025_multi_university_cancellation",
            "gz55zs_2025_university_adjustment_table",
        ],
        "中国教育在线、腾讯新闻转载中国教育在线、华龙网转载中国教育在线和新期教育网表格均汇总称合肥师范学院2025年拟撤销视觉传达设计（中外合作）专业。",
        "medium",
    )
    add_warning_list(
        2025,
        "四川美术学院",
        "本科",
        "major_cancel",
        "2025年度拟撤销专业",
        ["广告学", "教育技术学", "广播电视编导"],
        "拟撤销。",
        "二级汇总来源称学校学位评定委员会、校长办公会审议通过。",
        [
            "qq_eol_2025_multi_university_cancellation",
            "eol_2025_multi_university_cancellation",
            "cqnews_eol_2025_multi_university_cancellation",
            "gz55zs_2025_university_adjustment_table",
        ],
        "中国教育在线、腾讯新闻转载中国教育在线、华龙网转载中国教育在线和新期教育网表格均汇总称四川美术学院2025年拟撤销广告学、教育技术学、广播电视编导3个本科专业。",
        "medium",
    )

    add_warning_list(
        2025,
        "楚雄师范学院",
        "本科",
        "major_cancel",
        "2025年度拟撤销专业",
        ["中国少数民族语言文学", "酒店管理"],
        "拟同意撤销。",
        "二级汇总来源称需求调研、专家论证、省教育厅高教处初审同意、学校教学指导委员会审议、校长办公会议审定。",
        ["sohu_2025_six_university_cancellation_summary"],
        "搜狐汇总称楚雄师范学院2025年拟同意撤销中国少数民族语言文学、酒店管理2个专业。",
        "medium",
    )
    cxtc_cancel_rows = [
        ("中国少数民族语言文学", "050104", "4", "文学", "人文学院"),
        ("酒店管理", "120902", "4", "管理学", "管理与经济学院"),
    ]
    for row_no, (major, code, duration, degree, school_unit) in enumerate(cxtc_cancel_rows, start=1):
        rows.append(
            {
                "policy_year": 2025,
                "region": "楚雄师范学院",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "2025年拟撤销专业信息表",
                "reported_major_name": major,
                "major_code": code,
                "study_duration": duration,
                "source_row_no": str(row_no),
                "policy_action": "拟同意撤销。",
                "criterion_text": f"学校教学指导委员会审议并报请2025年第18次校长办公会议审定；所属学院为{school_unit}。",
                "source_ids": "gkzxw_cxtc_2025_major_cancellation",
                "evidence_text": f"高考资讯网转载楚雄师范学院公示表2第{row_no}行列出拟撤销{major}（{code}，{duration}年，{degree}），所属学院为{school_unit}。",
                "confidence": "medium",
            }
        )
    add_warning_list(
        2025,
        "九江学院",
        "本科",
        "major_cancel",
        "2025年度拟申请撤销本科专业",
        ["表演", "服装与服饰设计", "信息管理与信息系统", "历史学", "测控技术与仪器"],
        "拟申请撤销。",
        "二级汇总来源称已连续停招5年以上。",
        ["sohu_2025_six_university_cancellation_summary"],
        "搜狐汇总称九江学院2025年度拟申请撤销表演、服装与服饰设计、信息管理与信息系统、历史学、测控技术与仪器5个已连续停招5年以上的本科专业。",
        "medium",
    )
    add_warning_list(
        2025,
        "江西师范大学",
        "本科",
        "major_cancel",
        "2025年度拟撤销专业",
        ["网络工程", "物联网工程", "经济统计学", "信息与计算科学", "汉语言", "翻译", "劳动与社会保障"],
        "拟实施撤销处理。",
        "二级汇总来源称学院申报、学校审查。",
        ["sohu_2025_six_university_cancellation_summary"],
        "搜狐汇总称江西师范大学2025年度拟对网络工程、物联网工程、经济统计学、信息与计算科学、汉语言、翻译、劳动与社会保障7个专业实施撤销处理。",
        "medium",
    )
    add_warning_list(
        2025,
        "浙江农林大学",
        "本科",
        "major_cancel",
        "2025年度拟撤销专业",
        ["环境工程", "森林保护", "市场营销"],
        "拟撤销。",
        "二级汇总来源称学院论证、申报，校教育教学专门委员会审议。",
        ["sohu_2025_six_university_cancellation_summary"],
        "搜狐汇总称浙江农林大学2025年度拟撤销环境工程、森林保护、市场营销3个专业。",
        "medium",
    )
    zafu_cancel_rows = [
        ("环境工程", "082502", "四年", "工学", "2001", "2019"),
        ("森林保护", "090503", "四年", "农学", "1990", "2019"),
        ("市场营销", "120202", "四年", "管理学", "2006", "2019"),
    ]
    for row_no, (major, code, duration, degree, opened_year, stopped_year) in enumerate(zafu_cancel_rows, start=1):
        rows.append(
            {
                "policy_year": 2025,
                "region": "浙江农林大学",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "浙江农林大学拟撤销专业汇总表",
                "reported_major_name": major,
                "major_code": code,
                "study_duration": duration,
                "source_row_no": str(row_no),
                "policy_action": "拟撤销。",
                "criterion_text": f"浙江农林大学拟撤销专业汇总表；开设年份{opened_year}，停招年份{stopped_year}，调整类型为撤销。",
                "source_ids": "zafu_2025_major_cancellation_pdf;sohu_2025_six_university_cancellation_summary",
                "evidence_text": f"浙江农林大学官网PDF附件第{row_no}行列出拟撤销{major}（{code}，{duration}，{degree}），开设年份{opened_year}，停招年份{stopped_year}。",
                "confidence": "high",
            }
        )
    add_warning_list(
        2025,
        "文山学院",
        "本科",
        "major_cancel",
        "2025年度拟撤销本科专业",
        ["食品质量与安全", "工程管理", "视觉传达设计"],
        "拟向教育部申请撤销。",
        "二级汇总来源称连续五年停招且无在校学生，经论证、征求意见、学校教学指导委员会和校长办公会审议通过。",
        ["sohu_2025_six_university_cancellation_summary"],
        "搜狐汇总称文山学院2025年拟向教育部申请撤销食品质量与安全、工程管理、视觉传达设计3个本科专业。",
        "medium",
    )

    add_warning_list(
        2025,
        "福建师范大学",
        "本科",
        "major_cancel",
        "2025年度拟申请撤销专业",
        ["音乐表演", "信息管理与信息系统", "物联网工程"],
        "拟申请撤销。",
        "二级汇总来源称学院申报、学校审议。",
        ["xmhbjy_2025_multi_university_cancellation", "qq_zsgk_2025_multi_university_adjustment", "gz55zs_2025_university_adjustment_table", "ymzy_2025_multi_university_adjustment"],
        "浩博教育、腾讯新闻转载掌上高考、新期教育网表格和优志愿均汇总称福建师范大学2025年拟申请撤销音乐表演、信息管理与信息系统、物联网工程3个原有专业。",
        "medium",
    )
    add_warning_list(
        2025,
        "上海交通大学",
        "本科",
        "major_cancel",
        "2025年度拟撤销本科专业",
        ["临床医学（七年制）", "口腔医学（七年制）", "法学", "传播学", "计算机科学与技术", "行政管理"],
        "拟撤销。",
        "二级汇总来源称上海交通大学官网发布2025年度拟增设/撤销本科专业公示。",
        [
            "xmhbjy_2025_multi_university_cancellation",
            "gkztc_2025_sjtu_major_cancellation",
            "qq_qingtah_2025_sjtu_major_cancellation",
            "ymzy_2025_multi_university_adjustment",
        ],
        "浩博教育、腾讯新闻转载青塔和优志愿均汇总称上海交通大学2025年度拟撤销临床医学（七年制）、口腔医学（七年制）、法学、传播学、计算机科学与技术、行政管理6个专业；高考直通车单篇确认该校拟撤销临床医学（七年制）、口腔医学（七年制）等专业。",
        "medium",
    )
    add_warning_list(
        2025,
        "南京航空航天大学",
        "本科",
        "major_cancel",
        "2025年度本科专业设置拟调整",
        ["美术学", "空间科学与技术", "空间信息与数字技术"],
        "拟撤销。",
        "学院申报、初评、学校专家组评议。",
        ["nuaa_2025_major_adjustment"],
        "南京航空航天大学官网公示列出2025年度拟撤销专业3个：美术学、空间科学与技术、空间信息与数字技术。",
        "high",
    )
    add_warning_list(
        2025,
        "南京航空航天大学",
        "本科",
        "major_cancel",
        "2025年度拟撤销专业",
        ["美术学", "空间科学与技术", "空间信息与数字技术"],
        "拟撤销。",
        "二级汇总来源称学校发布2025年度本科专业设置拟调整公示。",
        ["xmhbjy_2025_multi_university_cancellation"],
        "浩博教育汇总称南京航空航天大学2025年度拟撤销美术学、空间科学与技术、空间信息与数字技术3个专业。",
        "medium",
    )

    jou_cancel_rows = [
        ("1", "测控技术与仪器", "080301", "仪器类", "工学", "电子工程学院"),
        ("2", "动植物检疫", "090403T", "动物医学类", "农学", "海洋食品与生物工程学院"),
        ("3", "广告学", "050303", "新闻传播学类", "文学", "文法学院"),
        ("4", "水质科学与技术", "082507T", "环境科学与工程类", "工学", "环境与化学工程学院"),
    ]
    for source_row_no, major, code, major_class, discipline, college in jou_cancel_rows:
        rows.append(
            {
                "policy_year": 2025,
                "region": "江苏海洋大学",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "2025年度拟撤销专业汇总表",
                "reported_major_name": major,
                "major_code": code,
                "study_duration": "",
                "policy_action": "拟撤销。",
                "criterion_text": "江苏海洋大学2025年度本科专业设置公示中的拟撤销专业汇总表逐项列出拟撤销专业。",
                "source_row_no": source_row_no,
                "source_ids": "jou_2025_major_adjustment",
                "evidence_text": f"江苏海洋大学2025年度拟撤销专业汇总表第{source_row_no}行列出{major}（{code}，{major_class}，{discipline}，{college}）。",
                "confidence": "high",
            }
        )

    zjyc_cancel_rows = [
        ("1", "生物技术", "071002", "理学", "2006年"),
        ("2", "服装与服饰设计", "130505", "艺术学", "2000年"),
    ]
    for source_row_no, major, code, degree, stop_start_year in zjyc_cancel_rows:
        rows.append(
            {
                "policy_year": 2024,
                "region": "浙江农林大学暨阳学院",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "2024年拟撤销停招本科专业名单",
                "reported_major_name": major,
                "major_code": code,
                "study_duration": "",
                "policy_action": "拟撤销。",
                "criterion_text": f"浙江农林大学暨阳学院2024年拟撤销停招本科专业名单列出该专业，停招起始年份为{stop_start_year}。",
                "source_row_no": source_row_no,
                "source_ids": "zjyc_2024_cancel_stop",
                "evidence_text": f"浙江农林大学暨阳学院2024年拟撤销停招本科专业名单第{source_row_no}行列出{major}（{code}，{degree}），停招起始年份为{stop_start_year}。",
                "confidence": "high",
            }
        )

    rows.append(
        {
            "policy_year": 2025,
            "region": "新疆师范大学",
            "education_level": "本科",
            "record_type": "major_cancel",
            "warning_label": "2025年度本科专业设置情况拟撤销专业",
            "reported_major_name": "舞蹈表演",
            "major_code": "130204",
            "study_duration": "",
            "policy_action": "拟撤销。",
            "criterion_text": "新疆师范大学2025年度本科专业设置情况公示说明，经学院申请，结合学校专业调整优化方案，学校审议通过。",
            "source_row_no": "正文",
            "source_ids": "xjnu_2025_major_adjustment",
            "evidence_text": "新疆师范大学2025年度本科专业设置情况公示第三部分“拟撤销专业”说明拟撤销舞蹈表演专业（代码：130204）。",
            "confidence": "high",
        }
    )

    ccit_cancel_rows = [
        ("1", "信息与计算科学", "070102", "理学"),
        ("2", "机械电子工程", "080204", "工学"),
        ("3", "工业设计", "080205", "工学"),
        ("4", "数字媒体技术", "080906", "工学"),
        ("5", "视觉传达设计", "130502", "艺术学"),
        ("6", "勘查技术与工程", "081402", "工学"),
    ]
    for source_row_no, major, code, degree in ccit_cancel_rows:
        rows.append(
            {
                "policy_year": 2023,
                "region": "长春工程学院",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "2023年度拟撤销本科专业",
                "reported_major_name": major,
                "major_code": code,
                "study_duration": "",
                "policy_action": "拟申请撤销。",
                "criterion_text": "长春工程学院2023年度拟撤销本科专业公示称这些专业已连续停招5年以上，并经相关学院党政联席会议决议、校内专业设置评议专家组评议、校长办公会审议。",
                "source_row_no": source_row_no,
                "source_ids": "ccit_2023_major_cancellation",
                "evidence_text": f"长春工程学院2023年度拟撤销本科专业公示表第{source_row_no}行列出{major}（{code}，{degree}），正文说明拟申请撤销6个已连续停招5年以上的本科专业。",
                "confidence": "high",
            }
        )

    xmhbjy_additional_cancel_rows = {
        "辽宁大学": ["信息与计算科学", "摄影", "广播电视学", "网络与新媒体"],
        "广西科技大学": ["材料科学与工程", "微电子科学与工程", "数字媒体技术", "市场营销", "人力资源管理", "物流工程", "信息与计算科学"],
        "北京信息科技大学": ["英语", "行政管理"],
        "江西农业大学": ["信息与计算科学"],
        "西安邮电大学": ["广播电视工程"],
        "宁波大学": ["财务会计教育", "新闻学", "广告学", "旅游管理与服务教育", "产品设计", "物联网"],
        "四川美术学院": ["广告学", "教育技术学", "广播电视编导"],
        "江西师范大学": ["网络工程", "物联网工程", "经济统计学", "信息与计算科学", "汉语言", "翻译", "劳动与社会保障"],
        "南阳理工学院": ["法学", "音乐表演", "网络工程", "电子科学与技术", "汽车服务工程"],
        "长春大学": ["应用统计学", "汽车服务工程", "文化产业管理"],
        "湖北第二师范学院": ["信息与计算科学", "物流工程", "公共事业管理", "汽车服务工程"],
        "河南工业大学": ["食品营养与检验教育", "交通工程", "电子信息科学与技术"],
        "西安文理学院": ["电子商务", "材料科学与工程", "音乐表演", "公共艺术"],
        "江汉大学": ["市场营销", "财务管理", "通信工程", "网络工程", "过程装备与控制工程", "工业设计"],
        "泉州信息工程学院": ["建筑电气与智能化", "材料成型及控制工程", "光电信息科学与工程", "投资学", "汽车服务工程"],
        "甘肃农业大学": ["农业电气化", "农村区域发展", "法学"],
    }
    for school, majors in xmhbjy_additional_cancel_rows.items():
        add_warning_list(
            2025,
            school,
            "本科",
            "major_cancel",
            "2025年度拟撤销本科专业",
            majors,
            "拟撤销。",
            "二级汇总来源列出学校2025年度拟撤销本科专业名单；用于补充交叉来源。",
            ["xmhbjy_2025_multi_university_cancellation"],
            f"新期教育网汇总称{school}2025年度拟撤销本科专业包括：{'、'.join(majors)}。",
            "medium",
        )
    add_warning_list(
        2025,
        "成都信息工程大学",
        "本科",
        "major_cancel",
        "2025年度拟撤销本科专业",
        ["物流工程", "信息管理与信息系统"],
        "拟撤销。",
        "二级汇总来源称学校撤销已连续五年不招生的本科专业。",
        ["xmhbjy_2025_multi_university_cancellation"],
        "新期教育网汇总称成都信息工程大学2025年度拟撤销物流工程、信息管理与信息系统2个本科专业。",
        "medium",
    )
    add_warning(
        2025,
        "成都信息工程大学",
        "本科",
        "major_stop_enrollment",
        "2025年度拟停招本科专业",
        "翻译",
        "拟停招。",
        "二级汇总来源称学校2025年度拟停招翻译本科专业。",
        ["xmhbjy_2025_multi_university_cancellation"],
        "新期教育网汇总称成都信息工程大学2025年度停招翻译本科专业。",
        "medium",
    )

    add_warning_list(
        2025,
        "昆明理工大学",
        "本科",
        "major_cancel",
        "2025年度本科专业设置推荐结果",
        ["运动康复", "电子信息科学与技术", "产品设计", "汽车维修工程教育"],
        "推荐撤销。",
        "学校本科专业设置新增、调整、撤销专业推荐结果，经学院申报、校内专业设置评议专家组评议形成。",
        ["kmust_2025_major_cancellation", "gz55zs_2025_university_adjustment_table"],
        "昆明理工大学官网公示称推荐撤销运动康复、电子信息科学与技术、产品设计、汽车维修工程教育共4个本科专业；二级表格来源列出的名单一致。",
        "high",
    )

    gz55zs_additional_2025 = {
        "昆明理工大学": ["运动康复", "电子信息科学与技术", "产品设计", "汽车维修工程教育"],
        "东北师范大学": ["财政学", "国际经济与贸易", "财务管理", "应用统计学", "旅游管理", "社会工作", "信息资源管理"],
    }
    for school, majors in gz55zs_additional_2025.items():
        add_warning_list(
            2025,
            school,
            "本科",
            "major_cancel",
            "2025年高校公示拟撤销专业",
            majors,
            "拟撤销。",
            "二级汇总来源列入2025年高校公示拟撤销专业表；需继续优先补学校官网原始公示。",
            ["gz55zs_2025_university_adjustment_table"],
            f"新期教育网《2025年高校公示拟撤销专业》表格列出{school}拟撤销专业：{'、'.join(majors)}。",
            "medium",
        )

    nenu_2024_report_cancel_rows = [
        "经济统计学",
        "广播电视学",
        "应用化学",
        "自然地理与资源环境",
        "电气工程及其自动化",
        "产品设计",
        "考古学",
        "表演",
        "数字媒体艺术",
    ]
    for row_no, major in enumerate(nenu_2024_report_cancel_rows, start=1):
        rows.append(
            {
                "policy_year": 2024,
                "region": "东北师范大学",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "2023-2024学年本科教育教学质量报告主动撤销专业",
                "reported_major_name": major,
                "policy_action": "主动撤销。",
                "criterion_text": "学校主动撤销不适应经济社会发展需求和学校办学定位、且停招5年以上的本科专业。",
                "source_row_no": str(row_no),
                "source_ids": "nenu_2024_teaching_quality_report_pdf",
                "evidence_text": f"东北师范大学2023-2024学年本科教育教学质量报告称，学校主动撤销经济统计学、广播电视学、应用化学、自然地理与资源环境、电气工程及其自动化、产品设计、考古学、表演、数字媒体艺术等本科专业；本行对应第{row_no}个可见名称：{major}。报告原文称主动撤销10个专业，但正文抽取和页面截图可见9个专业名称；本数据集仅结构化可见名称。",
                "confidence": "high",
            }
        )

    add_warning_list(
        2025,
        "西北农林科技大学",
        "本科",
        "major_cancel",
        "2025年拟新增与撤销专业",
        ["保险学", "视觉传达设计"],
        "拟撤销。",
        "经学院申请、学院教授委员会研究、学校教学指导专门委员会审议、校长办公会审定。",
        ["nwsuaf_2025_major_cancellation", "gz55zs_2025_university_adjustment_table"],
        "西北农林科技大学资源环境学院2025年拟新增与撤销专业公示列出撤销保险学、视觉传达设计2个本科专业；新期教育网表格与官网名单一致。",
        "high",
    )

    add_warning_list(
        2025,
        "东北师范大学",
        "本科",
        "major_cancel",
        "2025年度本科专业设置情况公示",
        ["财政学", "国际经济与贸易", "财务管理", "应用统计学", "旅游管理", "社会工作", "信息资源管理"],
        "拟撤销。",
        "东北师范大学教务处官网保留同名公示标题、发布时间和“公示已结束”状态；专业名单来自高考资讯网转载正文与公开汇总表。",
        ["nenu_2025_major_adjustment_notice", "gkzxw_nenu_2025_major_adjustment", "gz55zs_2025_university_adjustment_table"],
        "东北师范大学教务处官网保留2025年度本科专业设置情况公示标题和发布时间；高考资讯网转载称该公示拟撤销财政学等7个本科专业，新期教育网表格补充列出7个专业名称。",
        "medium",
    )

    gkzxw_jlau_cancel_rows = [
        ("广告学", "050303", "四年"),
        ("应用心理学", "071102", "四年"),
        ("农业机械化及其自动化", "082304", "四年"),
        ("野生动物与自然保护区管理", "090202", "四年"),
    ]
    for row_no, (major, code, duration) in enumerate(gkzxw_jlau_cancel_rows, start=1):
        rows.append(
            {
                "policy_year": 2025,
                "region": "吉林农业大学",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "2025年度拟申请撤销本科专业清单",
                "reported_major_name": major,
                "major_code": code,
                "study_duration": duration,
                "source_row_no": str(row_no),
                "policy_action": "拟申请撤销。",
                "criterion_text": "经学院申请、专家组论证、学校研究决定申请撤销。",
                "source_ids": "gkzxw_jlau_2025_major_cancellation",
                "evidence_text": f"高考资讯网转载吉林农业大学公示表第{row_no}行列出拟申请撤销{major}（{code}，{duration}）。",
                "confidence": "medium",
            }
        )

    add_warning_list(
        2025,
        "江苏师范大学",
        "本科",
        "major_cancel",
        "2025拟撤销本科专业",
        ["环境科学", "土地资源管理", "电子科学与技术"],
        "拟撤销。",
        "连续5年未招生；电子科学与技术保留中俄学院中外联合办学专业，普通本科多年未招生。",
        ["gkzxw_jsnu_2025_major_adjustment"],
        "高考资讯网转载江苏师范大学公示表2列出拟撤销环境科学、土地资源管理、电子科学与技术3个本科专业及撤销理由。",
        "medium",
    )

    fjnu_2024_stop_rows = [
        "酒店管理",
        "动画",
        "环境科学",
        "复合材料与工程",
    ]
    for row_no, major in enumerate(fjnu_2024_stop_rows, start=1):
        rows.append(
            {
                "policy_year": 2024,
                "region": "福建师范大学",
                "education_level": "本科",
                "record_type": "major_stop_enrollment",
                "warning_label": "2023-2024学年本科教学质量报告本学年停招专业",
                "reported_major_name": major,
                "source_row_no": str(row_no),
                "policy_action": "本学年停招。",
                "criterion_text": "落实“就业-招生-培养”联动的专业动态调整机制，积极推进本科专业调整。",
                "source_ids": "fjnu_2024_teaching_quality_report_pdf",
                "evidence_text": f"福建师范大学2023-2024学年本科教学质量报告称，本学年停招酒店管理、动画、环境科学、复合材料与工程4个本科专业；本行对应第{row_no}个专业：{major}。",
                "confidence": "high",
            }
        )

    hfnu_cancel_rows = [
        ("视觉传达设计（中外合作）", "130502H", "4", "艺术学", "美术与设计学院"),
    ]
    for row_no, (major, code, duration, degree, school_unit) in enumerate(hfnu_cancel_rows, start=1):
        rows.append(
            {
                "policy_year": 2025,
                "region": "合肥师范学院",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "2025年本科专业设置与调整工作公示",
                "reported_major_name": major,
                "major_code": code,
                "study_duration": duration,
                "source_row_no": str(row_no),
                "policy_action": "撤销专业。",
                "criterion_text": f"经二级学院论证、申报，学校审议；所在学院为{school_unit}。",
                "source_ids": "gkzxw_hfnu_2025_major_adjustment",
                "evidence_text": f"高考资讯网转载合肥师范学院公示表列出撤销{major}（{code}，{duration}年，{degree}），所在学院为{school_unit}。",
                "confidence": "medium",
            }
        )

    cdut_cancel_rows = [
        ("地球信息科学与技术", "070903T"),
        ("材料成型及控制工程", "080203"),
        ("电子信息科学与技术", "080714T"),
        ("制药工程", "081302"),
        ("海洋油气工程", "081506T"),
        ("城乡规划", "082802"),
        ("园林", "090502"),
        ("工程造价", "120105"),
        ("劳动与社会保障", "120403"),
        ("戏剧影视文学", "130304"),
        ("环境设计", "130503"),
    ]
    for row_no, (major, code) in enumerate(cdut_cancel_rows, start=1):
        rows.append(
            {
                "policy_year": 2025,
                "region": "成都理工大学",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "2025年拟新设专业与拟撤销专业公示",
                "reported_major_name": major,
                "major_code": code,
                "source_row_no": str(row_no),
                "policy_action": "拟撤销。",
                "criterion_text": "经学院申报及组织专家评议，教务处专题研讨，校长办公会审议、党委常委会审定。",
                "source_ids": "gkzxw_cdut_2025_major_adjustment;sohu_2025_ten_university_cancellation_aug02;ymzy_2025_multi_university_adjustment",
                "evidence_text": f"高考资讯网转载成都理工大学公示列出拟撤销专业{major}（{code}）；搜狐和优志愿汇总亦称成都理工大学2025年拟撤销专业11个。",
                "confidence": "medium",
            }
        )

    add_warning(
        2025,
        "暨南大学",
        "本科",
        "major_cancel",
        "2025年度拟新增及撤销本科专业",
        "给排水科学与工程",
        "申报撤销。",
        "该专业已停招5年以上且无在校生。",
        ["sohu_2025_ten_university_cancellation_aug02", "ymzy_2025_multi_university_adjustment"],
        "搜狐汇总称暨南大学2025年度将申报撤销给排水科学与工程专业；优志愿汇总与该表述一致。",
        "medium",
    )
    add_warning_list(
        2025,
        "扬州大学",
        "本科",
        "major_cancel",
        "2025年拟新增专业和拟撤销专业",
        ["通信工程", "风景园林"],
        "拟撤销。",
        "已停招五年及以上，无在校生。",
        ["sohu_2025_ten_university_cancellation_aug02", "acabridge_2025_six_university_cancellation"],
        "搜狐和学术桥均汇总称扬州大学2025年拟撤销通信工程、风景园林2个专业，原因是已停招五年及以上且无在校生。",
        "medium",
    )
    add_warning_list(
        2025,
        "唐山学院",
        "本科",
        "major_cancel",
        "2025年度拟撤销本科专业",
        ["汽车服务工程", "酒店管理"],
        "拟撤销。",
        "经校长办公会议议定。",
        ["tsc_2025_major_adjustment", "sohu_2025_ten_university_cancellation_aug02"],
        "唐山学院官网公示称拟撤销汽车服务工程、酒店管理2个本科专业；搜狐汇总与该名单一致。",
        "high",
    )
    csuft_swxy_cancel_rows = [
        ("机械设计制造及其自动化", "080202", "工学", "四年"),
        ("交通运输", "081801", "工学", "四年"),
        ("环境科学", "082503", "工学", "四年"),
        ("城乡规划", "082802", "工学", "五年"),
        ("园林", "090502", "农学", "四年"),
    ]
    for row_no, (major, code, degree, duration) in enumerate(csuft_swxy_cancel_rows, start=1):
        rows.append(
            {
                "policy_year": 2025,
                "region": "中南林业科技大学涉外学院",
                "education_level": "本科",
                "record_type": "major_cancel",
                "warning_label": "2025年拟撤销本科专业一览表",
                "reported_major_name": major,
                "major_code": code,
                "study_duration": duration,
                "source_row_no": str(row_no),
                "policy_action": "拟撤销。",
                "criterion_text": "经二级学院申报，教务处组织专家评审，教学指导委员会审议通过。",
                "source_ids": "csuft_swxy_2025_major_adjustment;sohu_2025_ten_university_cancellation_aug02",
                "evidence_text": f"中南林业科技大学涉外学院官网公示表第{row_no}行列出拟撤销{major}（{code}，{degree}，{duration}）。",
                "confidence": "high",
            }
        )

    add_warning_list(
        2025,
        "沈阳航空航天大学",
        "本科",
        "major_cancel",
        "2025年度新增、撤销专业",
        ["车辆工程", "工业工程", "功能材料", "焊接技术与工程", "日语", "英语"],
        "拟撤销。",
        "经二级学院充分论证、学院教学工作委员会和党政联席会议审议、本科生院审查、学术委员会评议、校长办公会审议；6个专业均为已停招专业。",
        ["sau_2025_major_adjustment_notice", "ymzy_2025_multi_university_adjustment"],
        "沈阳航空航天大学本科生院官网公示称2025年度申请撤销车辆工程、工业工程、功能材料、焊接技术与工程、日语、英语等6个已停招专业；优志愿汇总名单一致。",
        "high",
    )
    add_warning_list(
        2025,
        "齐鲁工业大学（山东省科学院）",
        "本科",
        "major_cancel",
        "2025年拟申请撤销专业",
        ["安全工程", "汽车服务工程", "保险学"],
        "拟申请撤销。",
        "已停招5年及以上且无在校生。",
        ["qlu_2025_major_cancellation", "acabridge_2025_six_university_cancellation"],
        "齐鲁工业大学（山东省科学院）官网公示称对已停招5年及以上且无在校生的安全工程、汽车服务工程、保险学3个本科专业向教育部申请撤销；学术桥汇总名单一致。",
        "high",
    )
    add_warning_list(
        2025,
        "河北工业大学",
        "本科",
        "major_cancel",
        "2025年度拟申请撤销本科专业",
        ["智能科学与技术", "交通运输"],
        "拟申请撤销。",
        "结合学校发展规划和专业建设实际，经学院申报、学校审查、学术委员会审议。",
        ["hebut_2025_major_cancellation", "acabridge_2025_six_university_cancellation"],
        "河北工业大学官网公示称拟申请撤销智能科学与技术、交通运输2个专业；学术桥汇总与该名单一致。",
        "high",
    )
    add_warning_list(
        2025,
        "信阳师范大学",
        "本科",
        "major_cancel",
        "2025年度申请撤销专业",
        ["人文教育", "应用物理学", "信息管理与信息系统"],
        "申请撤销。",
        "经学院论证申报、学校评议推荐和省教育厅评审，学校拟申请撤销3个本科专业。",
        ["xynu_2025_major_cancellation", "acabridge_2025_six_university_cancellation"],
        "信阳师范大学教务处2025年度拟申报、撤销专业公示列出拟申请撤销人文教育、应用物理学、信息管理与信息系统3个本科专业；学术桥汇总与官网名单一致。",
        "high",
    )
    add_warning_list(
        2025,
        "北京工业大学",
        "本科",
        "major_cancel",
        "2025年拟撤销本科专业",
        ["日语", "广告学", "食品质量与安全", "风景园林"],
        "拟撤销。",
        "根据学校本科专业优化调整方案（2023-2025年），2025年拟撤销4个本科专业。",
        ["bjut_2025_major_cancellation", "acabridge_2025_six_university_cancellation", "gkzxw_bjut_2025_major_adjustment"],
        "北京工业大学教务处2025年度拟增设新专业和拟撤销专业情况公示列出拟撤销日语、广告学、食品质量与安全、风景园林4个专业；学术桥和高考资讯网转载名单一致。",
        "high",
    )
    add_warning_list(
        2025,
        "上海师范大学天华学院",
        "本科",
        "major_cancel",
        "2025年度拟申报及撤销本科专业",
        ["机械设计制造及其自动化", "汉语国际教育"],
        "拟撤销。",
        "学校2025年度拟申报及撤销本科专业公示。",
        ["sthu_2025_major_cancellation"],
        "上海师范大学天华学院官网公示列出拟撤销机械设计制造及其自动化、汉语国际教育2个本科专业。",
        "high",
    )

    add_warning_list(
        2024,
        "山东大学",
        "本科",
        "major_stop_enrollment",
        "2023-2024学年暂停招生专业",
        [
            "金融工程",
            "保险学",
            "公共事业管理",
            "世界史",
            "文物与博物馆学",
            "文化产业管理",
            "电子信息工程",
            "物联网工程",
            "电子信息科学与技术",
            "无机非金属材料工程",
            "金属材料工程",
            "过程装备与控制工程",
            "车辆工程",
            "测控技术与仪器",
            "物流工程",
            "土木工程",
            "水利水电工程",
            "信息管理与信息系统",
            "市场营销",
            "国际商务",
            "人力资源管理",
            "物流管理",
            "工业工程",
            "旅游管理",
            "信息安全",
            "电子商务",
            "海洋资源与环境",
        ],
        "暂停招生。",
        "山东大学本科生院调整情况中序号标为“-”的为停招专业；统计时点为2023-2024学年。",
        ["sciencenet_sdu_2025_undergrad_adjustment", "gkztc_sdu_2025_undergrad_adjustment"],
        "科学网和高考直通车均转载或转述山东大学本科生院《2023年9月至2024年8月本科专业设置及调整情况》，列出金融工程、世界史、土木工程等27个暂停招生专业。",
        "medium",
    )
    add_warning_list(
        2024,
        "山东大学",
        "本科",
        "major_cancel",
        "2023-2024学年撤销专业",
        ["书法学", "高分子材料与工程", "包装工程", "材料物理", "材料化学", "工业设计", "交通运输", "资源循环科学与工程", "制药工程", "电子商务（工学）"],
        "撤销。",
        "山东大学本科生院2023-2024学年本科专业设置及调整情况。",
        ["sciencenet_sdu_2025_undergrad_adjustment", "gkztc_sdu_2025_undergrad_adjustment"],
        "科学网和高考直通车均称山东大学本科生院调整情况列出撤销书法学、高分子材料与工程、包装工程、材料物理、材料化学、工业设计、交通运输、资源循环科学与工程、制药工程、电子商务（工学）10个专业。",
        "medium",
    )

    add_warning_list(
        2024,
        "青岛大学",
        "本科",
        "major_cancel",
        "2024年度撤销近五年停招专业",
        ["经济统计学", "物联网工程", "市场营销", "财务管理", "国际商务", "人力资源管理"],
        "予以撤销。",
        "二级汇总来源称结合学校学科专业布局、优化调整需求和近五年专业停招情况，经学校审议。",
        ["sohu_2024_multi_university_cancellation_summary"],
        "搜狐汇总称青岛大学2024年本科专业设置调整公示列出撤销近五年停招专业：经济统计学、物联网工程、市场营销、财务管理、国际商务和人力资源管理。",
        "medium",
    )
    add_warning_list(
        2024,
        "齐鲁工业大学（山东省科学院）",
        "本科",
        "major_cancel",
        "2024年度申请撤销专业",
        ["朝鲜语", "投资学", "摄影"],
        "向教育部申请撤销。",
        "二级汇总来源称已停招5年及以上且无在校生。",
        ["sohu_2024_multi_university_cancellation_summary"],
        "搜狐汇总称齐鲁工业大学（山东省科学院）2024年对已停招5年及以上且无在校生的朝鲜语、投资学、摄影3个本科专业向教育部申请撤销。",
        "medium",
    )
    add_warning_list(
        2024,
        "兰州大学",
        "本科",
        "major_cancel",
        "2024年度拟撤销专业",
        ["药物制剂", "中药学", "教育学"],
        "拟撤销。",
        "二级汇总来源称学校发布拟预备案和拟撤销本科专业公示。",
        ["sohu_2024_multi_university_cancellation_summary"],
        "搜狐汇总称兰州大学2024年拟撤销药物制剂、中药学、教育学3个专业。",
        "medium",
    )
    add_warning_list(
        2024,
        "西北大学",
        "本科",
        "major_cancel",
        "2024年度拟申请撤销专业",
        ["汉语言", "金融工程", "财政学", "管理科学", "图书馆学", "广告学", "视觉传达设计"],
        "拟申请撤销。",
        "二级汇总来源称学校发布2024年度专业调整计划。",
        ["sohu_2024_multi_university_cancellation_summary"],
        "搜狐汇总称西北大学2024年度拟申请撤销汉语言、金融工程、财政学、管理科学、图书馆学、广告学、视觉传达设计7个本科专业。",
        "medium",
    )
    add_warning_list(
        2024,
        "西北农林科技大学",
        "本科",
        "major_cancel",
        "2024年度拟撤销专业",
        ["人文地理与城乡规划", "电子商务", "信息管理与信息系统"],
        "拟撤销。",
        "二级汇总来源称学校发布拟预备案和拟撤销本科专业公示。",
        ["sohu_2024_multi_university_cancellation_summary"],
        "搜狐汇总称西北农林科技大学2024年拟撤销人文地理与城乡规划、电子商务、信息管理与信息系统3个专业。",
        "medium",
    )
    add_warning_list(
        2024,
        "湘潭大学",
        "本科",
        "major_cancel",
        "2024年度撤销专业",
        ["电子信息科学与技术", "数字出版", "药学", "建筑环境与能源应用工程", "酒店管理", "广播电视学", "翻译"],
        "撤销。",
        "二级汇总来源称学院申报、学校组织专业设置评议专家审议、学校校长办公会议审定。",
        ["sohu_2024_multi_university_cancellation_summary"],
        "搜狐汇总称湘潭大学2024年本科专业设置调整公示列出撤销电子信息科学与技术、数字出版、药学、建筑环境与能源应用工程、酒店管理、广播电视学、翻译7个专业。",
        "medium",
    )
    add_warning_list(
        2024,
        "西南交通大学",
        "本科",
        "major_cancel",
        "2024年度拟撤销专业",
        ["网络工程", "物联网工程", "森林保护", "信息与计算科学", "电子商务", "旅游管理", "中药学"],
        "拟撤销。",
        "二级汇总来源称学校发布2024年拟调整本科专业公示。",
        ["sohu_2024_multi_university_cancellation_summary"],
        "搜狐汇总称西南交通大学2024年拟撤销网络工程、物联网工程、森林保护、信息与计算科学、电子商务、旅游管理、中药学7个本科专业。",
        "medium",
    )
    add_warning_list(
        2024,
        "四川大学",
        "本科",
        "major_cancel",
        "2024年度拟撤销本科专业",
        [
            "音乐学",
            "表演",
            "动画",
            "保险学",
            "广播电视学",
            "信息管理与信息系统",
            "公共事业管理",
            "电子商务",
            "应用物理学",
            "核物理",
            "生物技术",
            "材料物理",
            "材料化学",
            "金属材料工程",
            "无机非金属材料工程",
            "电子科学与技术",
            "电子信息科学与技术",
            "保密管理",
            "工业设计",
            "网络工程",
            "建筑环境与能源应用工程",
            "环境科学",
            "城乡规划",
            "工程造价",
            "风景园林",
            "水利水电工程",
            "水文与水资源工程",
            "纺织工程",
            "冶金工程",
            "安全工程",
            "信息安全",
        ],
        "拟撤销。",
        "二级汇总来源称四川大学官网发布2024年度拟新增本科专业、预备案专业、拟撤销本科专业公示。",
        ["sohu_2024_multi_university_cancellation_summary"],
        "搜狐汇总称四川大学2024年度拟撤销音乐学、表演、动画、保险学、广播电视学、信息管理与信息系统、公共事业管理、电子商务、应用物理学、核物理、生物技术、材料物理、材料化学、金属材料工程、无机非金属材料工程、电子科学与技术、电子信息科学与技术、保密管理、工业设计、网络工程、建筑环境与能源应用工程、环境科学、城乡规划、工程造价、风景园林、水利水电工程、水文与水资源工程、纺织工程、冶金工程、安全工程、信息安全31个本科专业。",
        "medium",
    )
    add_warning(
        2020,
        "河南省",
        "本科",
        "policy_rule",
        "本科专业设置负面清单/黄牌红牌预警",
        "",
        "限制招生规模、暂停招生或调整撤销；每年发布建议暂缓增设本科专业目录。",
        "全省布点较多、规模较大、近三年社会就业率较低、社会需求饱和。",
        ["henan_2020_undergrad_warning_policy"],
        "河南提出定期发布需求严重饱和本科专业、紧缺本科专业和就业率较低本科专业名单，建立本科专业设置负面清单制度，并对布点较多、规模较大、近三年社会就业率较低、社会需求饱和的本科专业发布黄牌、红牌预警。",
    )
    add_warning(
        2025,
        "全国/省级教育行政部门",
        "本科",
        "policy_rule",
        "省级急需本科专业清单和过剩专业预警清单",
        "",
        "省级教育行政部门发布清单；对布点量大、就业率过低的专业及相近专业原则上不再支持增设。",
        "人才供需关系、区域发展匹配度、布点量、就业率。",
        ["edu_moe_2025_undergrad_setup_notice"],
        "教育部高教司要求各省级教育行政部门于7月31日前发布本年度省级急需本科专业清单和过剩专业预警清单；对本地区布点量大、就业率过低的专业及相近专业，原则上不再支持增设。",
    )
    add_warning(
        2022,
        "江西省",
        "本科",
        "policy_rule",
        "本科专业黄牌红牌/停止招生机制",
        "",
        "毕业去向落实率低于50%的专业列入黄牌提示；连续两年低于50%的专业列入红牌提示并停止招生。",
        "毕业去向落实率低于50%；连续两年毕业去向落实率低于50%。",
        ["sina_2022_jiangxi_undergrad_structure_policy"],
        "江西本科专业结构优化调整办法明确，毕业去向落实率低于50%的专业列入黄牌提示，连续两年低于50%的专业列入红牌提示并停止招生。",
    )
    add_warning(
        2023,
        "四川省",
        "高等教育",
        "policy_rule",
        "就业-招生-培养联动红黄牌机制",
        "",
        "连续两年初次毕业去向落实率低于50%列入黄牌专业；连续三年低于50%列入红牌专业。",
        "初次毕业去向落实率连续两年或连续三年低于50%。",
        ["sichuan_2023_employment_enrollment_policy"],
        "四川提出深化高等学校就业、招生、培养联动机制改革，对初次毕业去向落实率连续两年低于50%的专业给予黄牌，连续三年低于50%的专业给予红牌。",
    )
    add_warning(
        2025,
        "福建省",
        "高等教育",
        "policy_rule",
        "就业红黄牌提示制度",
        "",
        "对就业质量不高的专业调减或停止招生。",
        "高校毕业生就业质量、专业供需适配情况。",
        ["fujian_2025_full_employment_policy"],
        "福建关于全方位促进高质量充分就业的实施意见提出健全就业红黄牌提示制度，对就业质量不高的专业调减或停止招生。",
    )
    add_warning(
        2024,
        "山东省",
        "师范类专业",
        "policy_rule",
        "师范类专业红黄牌预警和退出机制",
        "",
        "对就业率较低或专业认证一级监测不达标的师范类专业实施专业预警和退出机制。",
        "师范类专业就业率较低，或专业认证一级监测不达标。",
        ["e23_2024_shandong_teacher_warning_policy"],
        "山东推进师范教育高质量发展的措施提出建立师范类专业红黄牌制度，对就业率较低或专业认证一级监测不达标的实施专业预警和退出机制。",
    )
    add_warning(
        2022,
        "安徽省",
        "本科",
        "policy_rule",
        "连续3年就业去向落实率低于60%专业暂停招生",
        "",
        "连续3年就业去向落实率低于60%的专业暂停招生。",
        "连续3年就业去向落实率低于60%。",
        ["anhui_2022_structure_policy"],
        "安徽深化高校学科专业结构改革公开报道明确，对连续3年就业去向落实率低于60%的专业暂停招生。",
    )
    if raw_dir is not None:
        rows.extend(parse_shanghai_undergrad_warning_rows(raw_dir))
        rows.extend(parse_hust_historical_stop_rows(raw_dir))
        rows.extend(parse_sufe_historical_stop_rows(raw_dir))
        rows.extend(parse_ujs_historical_stop_rows(raw_dir))
        rows.extend(parse_njtech_historical_stop_rows(raw_dir))
        rows.extend(parse_jnu_historical_stop_rows(raw_dir))
        rows.extend(parse_moe_complete_cancel_rows(raw_dir))
        rows.extend(parse_hunan_2019_undergrad_catalog_cancel_rows(raw_dir))
        rows.extend(parse_scnu_2019_teaching_quality_cancel_rows(raw_dir))
        rows.extend(parse_ynnu_2019_teaching_quality_cancel_rows(raw_dir))
        rows.extend(parse_dzu_2019_teaching_quality_stop_rows(raw_dir))
        rows.extend(parse_glut_2019_teaching_quality_stop_rows(raw_dir))
        rows.extend(parse_hrbu_2020_teaching_quality_stop_rows(raw_dir))
        rows.extend(parse_hrbu_2020_teaching_quality_cancel_rows(raw_dir))
        rows.extend(parse_cust_2020_application_history_rows(raw_dir))
        rows.extend(parse_suse_2019_teaching_quality_cancel_rows(raw_dir))
        rows.extend(parse_ourjiangsu_2020_2019_jiangsu_cancel_rows(raw_dir))
        rows.extend(parse_cyol_2018_hdu_major_stop_adjustment_rows(raw_dir))
        rows.extend(parse_xpu_2019_major_dynamic_adjustment_rows(raw_dir))
        rows.extend(parse_zjnu_2020_2019_undergrad_setup_news_rows(raw_dir))
        rows.extend(parse_jxau_smart_agriculture_application_rows(raw_dir))
        rows.extend(parse_hbu_jijianjiancha_application_rows(raw_dir))
        rows.extend(parse_kmust_2025_stomatology_application_rows(raw_dir))
        rows.extend(parse_nhky_2020_cybersecurity_application_rows(raw_dir))
        rows.extend(parse_qjnu_2025_area_studies_application_rows(raw_dir))
        rows.extend(parse_aqnu_2025_sports_training_application_rows(raw_dir))
        rows.extend(parse_yxu_2025_application_history_rows(raw_dir))
        rows.extend(parse_zcmu_2025_nursing_application_rows(raw_dir))
        rows.extend(parse_cxtc_2025_uav_application_rows(raw_dir))
        rows.extend(parse_wust_2025_fintech_application_rows(raw_dir))
        rows.extend(parse_bigc_2025_digital_economy_application_rows(raw_dir))
        rows.extend(parse_hbkjxy_2025_cybersecurity_application_rows(raw_dir))
        rows.extend(parse_whwl_2025_ai_education_application_rows(raw_dir))
        rows.extend(parse_muc_2025_sports_training_application_rows(raw_dir))
        rows.extend(parse_sus_2025_football_application_rows(raw_dir))
        rows.extend(parse_qtnu_2025_sports_training_application_rows(raw_dir))
        rows.extend(parse_gmc_2025_geriatric_medicine_application_rows(raw_dir))
        rows.extend(parse_gnnu_2025_ai_education_application_rows(raw_dir))
        rows.extend(parse_hbesxy_2025_stomatology_application_rows(raw_dir))
        rows.extend(parse_nwupl_2025_forensic_science_application_rows(raw_dir))
        rows.extend(parse_sicnu_2025_sports_tourism_application_rows(raw_dir))
        rows.extend(parse_tjufezj_2025_fiscal_application_rows(raw_dir))
        rows.extend(parse_whtyxykj_2025_smart_sports_engineering_application_rows(raw_dir))
        rows.extend(parse_ntu_2025_smart_energy_application_rows(raw_dir))
        rows.extend(parse_tjcu_2025_low_altitude_application_rows(raw_dir))
        rows.extend(parse_csuft_2025_low_altitude_application_rows(raw_dir))
        rows.extend(parse_zju_2025_undergrad_stop_rows(raw_dir))
        rows.extend(parse_hfit_2025_major_setting_stop_rows(raw_dir))
        rows.extend(parse_shiep_stop_enrollment_rows(raw_dir))
        rows.extend(parse_shiep_2026_professional_adjustment_rows(raw_dir))
        rows.extend(parse_chu_2019_2024_major_setting_rows(raw_dir))
        rows.extend(parse_2025_university_adjustment_notice_rows(raw_dir))
        rows.extend(parse_pxc_2025_smart_construction_application_rows(raw_dir))
        rows.extend(parse_eol_2020_2019_moe_direct_cancel_rows(raw_dir))
        rows.extend(parse_gxust_2016_undergrad_setting_stop_rows(raw_dir))
        rows.extend(parse_guizhou_2019_undergrad_setup_cancel_rows(raw_dir))
        rows.extend(parse_hebtu_2019_performance_target_cancel_rows(raw_dir))
        rows.extend(parse_aufe_2019_undergrad_stop_rows(raw_dir))
        rows.extend(parse_aufe_2019_cancel_notice_rows(raw_dir))
    return rows


def enrich_official_warning_records(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    catalog = read_catalog()
    enriched: list[dict[str, Any]] = []
    for row in merge_official_warning_rows(rows):
        reported = row.get("reported_major_name", "")
        standard_name = STANDARD_NAME_ALIASES.get(reported, reported)
        catalog_level = "本科" if row["education_level"] == "本科" else "专科"
        catalog_name = CATALOG_LOOKUP_ALIASES.get(standard_name, standard_name)
        catalog_row = catalog.get((catalog_name, catalog_level), catalog.get((catalog_name, ""), {}))
        warning_id_base = "|".join(
            [
                str(row["policy_year"]),
                row["region"],
                row["record_type"],
                row["warning_label"],
                reported,
                row.get("major_code", ""),
                row.get("study_duration", ""),
                row.get("source_row_no", ""),
                row["source_ids"],
            ]
        )
        enriched.append(
            {
                **row,
                "warning_id": hashlib.sha1(warning_id_base.encode("utf-8")).hexdigest()[:16],
                "schema_version": SCHEMA_VERSION,
                "standard_major_name": standard_name,
                "major_code": row.get("major_code") or (catalog_row.get("major_code", "") if reported else ""),
                "discipline": catalog_row.get("category", "") if reported else "",
                "major_category": catalog_row.get("subject", "") if reported else "",
                "captured_at": CAPTURED_AT,
            }
        )
    return sorted(
        enriched,
        key=lambda item: (
            item["policy_year"],
            item["region"],
            item["record_type"],
            item["reported_major_name"],
        ),
    )


def write_csv(path: Path, fieldnames: list[str], rows: list[dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames, extrasaction="ignore")
        writer.writeheader()
        for row in rows:
            writer.writerow({field: row.get(field, "") for field in fieldnames})


def write_jsonl(path: Path, rows: list[dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8") as handle:
        for row in rows:
            handle.write(json.dumps(row, ensure_ascii=False) + "\n")


def read_existing_sources(output_dir: Path) -> list[dict[str, Any]]:
    path = output_dir / "major_risk_warning_sources.csv"
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def build_major_summary(records: list[dict[str, Any]]) -> list[dict[str, Any]]:
    grouped: dict[tuple[str, str], list[dict[str, Any]]] = defaultdict(list)
    for row in records:
        grouped[(row["education_level"], row["standard_major_name"])].append(row)
    summary = []
    for (level, major), group in sorted(grouped.items()):
        risk_counter = Counter(row["risk_level"] for row in group)
        years_by_risk = defaultdict(list)
        for row in group:
            years_by_risk[row["risk_level"]].append(str(row["report_year"]))
        summary.append(
            {
                "education_level": level,
                "standard_major_name": major,
                "reported_names": ";".join(sorted({row["reported_major_name"] for row in group})),
                "major_code": next((row.get("major_code") for row in group if row.get("major_code")), ""),
                "discipline": next((row.get("discipline") for row in group if row.get("discipline")), ""),
                "major_category": next((row.get("major_category") for row in group if row.get("major_category")), ""),
                "red_count": risk_counter["red"],
                "yellow_count": risk_counter["yellow"],
                "green_count": risk_counter["green"],
                "red_years": ";".join(sorted(years_by_risk["red"])),
                "yellow_years": ";".join(sorted(years_by_risk["yellow"])),
                "green_years": ";".join(sorted(years_by_risk["green"])),
                "latest_risk_level": max(group, key=lambda row: row["report_year"])["risk_level"],
                "latest_report_year": max(row["report_year"] for row in group),
            }
        )
    return summary


def build_year_summary(records: list[dict[str, Any]]) -> list[dict[str, Any]]:
    grouped: dict[tuple[int, str, str], list[dict[str, Any]]] = defaultdict(list)
    for row in records:
        grouped[(row["report_year"], row["education_level"], row["risk_level"])].append(row)
    rows = []
    for (year, level, risk), group in sorted(grouped.items()):
        rows.append(
            {
                "report_year": year,
                "graduate_cohort": year - 1,
                "education_level": level,
                "risk_level": risk,
                "record_count": len(group),
                "majors": "、".join(row["reported_major_name"] for row in group),
                "source_ids": ";".join(sorted({source_id for row in group for source_id in row["source_ids"].split(";") if source_id})),
            }
        )
    return rows


def build_coverage(records: list[dict[str, Any]]) -> list[dict[str, Any]]:
    years = range(2009, 2027)
    levels = ["本科", "高职高专"]
    risks = ["red", "yellow", "green"]
    available = {(row["report_year"], row["education_level"], row["risk_level"]) for row in records}
    rows = []
    for year in years:
        for level in levels:
            for risk in risks:
                rows.append(
                    {
                        "report_year": year,
                        "education_level": level,
                        "risk_level": risk,
                        "has_records": "true" if (year, level, risk) in available else "false",
                    }
                )
    return rows


def build_report(
    records: list[dict[str, Any]],
    sources: list[dict[str, Any]],
    metrics: list[dict[str, Any]],
    official_warnings: list[dict[str, Any]],
    report_dir: Path,
) -> Path:
    report_dir.mkdir(parents=True, exist_ok=True)
    report_path = report_dir / "major_risk_warning_dataset_report_2026-06-12.md"
    red_records = [row for row in records if row["risk_level"] == "red"]
    yellow_records = [row for row in records if row["risk_level"] == "yellow"]
    green_records = [row for row in records if row["risk_level"] == "green"]
    failed_sources = [row for row in sources if row["status"] != "fetched"]
    latest_red = [row for row in red_records if row["report_year"] == max(r["report_year"] for r in red_records)]
    source_status = Counter(row["status"] for row in sources)
    summary = build_major_summary(records)
    top_red = sorted(summary, key=lambda row: (int(row["red_count"]), row["latest_report_year"]), reverse=True)[:20]
    lines = [
        "# Major Employment Warning Dataset",
        "",
        f"- Built at: {CAPTURED_AT}",
        f"- Records: {len(records)}",
        f"- Sources: {len(sources)} ({dict(source_status)})",
        f"- Red records: {len(red_records)}",
        f"- Yellow records: {len(yellow_records)}",
        f"- Green records: {len(green_records)}",
        f"- Metric records: {len(metrics)}",
        f"- Official policy warning records: {len(official_warnings)}",
        "",
        "## Files",
        "",
        "- `data/processed/major_risk_warnings/major_risk_warning_records.csv`",
        "- `data/processed/major_risk_warnings/major_risk_warning_records.jsonl`",
        "- `data/processed/major_risk_warnings/major_risk_warning_metrics.csv`",
        "- `data/processed/major_risk_warnings/major_risk_warning_official_policy_warnings.csv`",
        "- `data/processed/major_risk_warnings/major_risk_warning_major_summary.csv`",
        "- `data/processed/major_risk_warnings/major_risk_warning_year_summary.csv`",
        "- `data/processed/major_risk_warnings/major_risk_warning_sources.csv`",
        "- `data/processed/major_risk_warnings/major_risk_warning_coverage.csv`",
        "- `data/raw/major_risk_warnings/`",
        "",
        "## Latest Explicit Red Lists",
        "",
    ]
    for level in sorted({row["education_level"] for row in latest_red}):
        majors = "、".join(row["reported_major_name"] for row in latest_red if row["education_level"] == level)
        lines.append(f"- {level} {max(r['report_year'] for r in red_records)}: {majors}")
    lines.extend(
        [
            "",
            "## Highest Red-Flag Frequency",
            "",
            "| level | major | red_count | red_years | latest_risk |",
            "|---|---:|---:|---|---|",
        ]
    )
    for row in top_red:
        if int(row["red_count"]) <= 0:
            continue
        lines.append(
            f"| {row['education_level']} | {row['standard_major_name']} | {row['red_count']} | {row['red_years']} | {row['latest_risk_level']} |"
        )
    lines.extend(
        [
            "",
            "## Caveats",
            "",
            "- This is a public-source dataset, not an official Ministry of Education prohibition list.",
            "- `confidence=medium` rows are derived from public consecutive-year counts or secondary text excerpts and should be prioritized for manual review.",
            "- 2026 public data found in this pass contains undergraduate and high-vocational green-list records; no high-confidence 2026 red/yellow list was found during this crawl.",
            "- Historical high-vocational specialty names changed across catalog revisions; `reported_major_name` keeps source wording and `standard_major_name` gives a best-effort current-name alias.",
        ]
    )
    if failed_sources:
        lines.extend(["", "## Failed Sources", ""])
        for row in failed_sources:
            lines.append(f"- `{row['source_id']}`: {row['error']}")
    report_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return report_path


def build_dataset(
    raw_dir: Path,
    output_dir: Path,
    report_dir: Path,
    *,
    skip_fetch: bool = False,
) -> dict[str, Any]:
    source_rows = read_existing_sources(output_dir) if skip_fetch else fetch_sources(raw_dir)
    records = enrich_records(curated_rows())
    metrics = enrich_metric_records(curated_metric_rows(raw_dir))
    official_warnings = enrich_official_warning_records(curated_official_warning_rows(raw_dir))

    write_csv(output_dir / "major_risk_warning_sources.csv", SOURCE_FIELDS, source_rows)
    write_csv(output_dir / "major_risk_warning_records.csv", RECORD_FIELDS, records)
    write_jsonl(output_dir / "major_risk_warning_records.jsonl", records)
    write_csv(output_dir / "major_risk_warning_metrics.csv", METRIC_FIELDS, metrics)
    write_csv(
        output_dir / "major_risk_warning_official_policy_warnings.csv",
        OFFICIAL_WARNING_FIELDS,
        official_warnings,
    )
    write_csv(output_dir / "major_risk_warning_major_summary.csv", [
        "education_level",
        "standard_major_name",
        "reported_names",
        "major_code",
        "discipline",
        "major_category",
        "red_count",
        "yellow_count",
        "green_count",
        "red_years",
        "yellow_years",
        "green_years",
        "latest_risk_level",
        "latest_report_year",
    ], build_major_summary(records))
    write_csv(output_dir / "major_risk_warning_year_summary.csv", [
        "report_year",
        "graduate_cohort",
        "education_level",
        "risk_level",
        "record_count",
        "majors",
        "source_ids",
    ], build_year_summary(records))
    write_csv(output_dir / "major_risk_warning_coverage.csv", [
        "report_year",
        "education_level",
        "risk_level",
        "has_records",
    ], build_coverage(records))

    xlsx_path = output_dir / "major_risk_warning_dataset_latest.xlsx"
    with pd.ExcelWriter(xlsx_path, engine="openpyxl") as writer:
        pd.DataFrame(records, columns=RECORD_FIELDS).to_excel(writer, index=False, sheet_name="records")
        pd.DataFrame(metrics, columns=METRIC_FIELDS).to_excel(writer, index=False, sheet_name="metrics")
        pd.DataFrame(official_warnings, columns=OFFICIAL_WARNING_FIELDS).to_excel(writer, index=False, sheet_name="official_warnings")
        pd.DataFrame(build_major_summary(records)).to_excel(writer, index=False, sheet_name="major_summary")
        pd.DataFrame(build_year_summary(records)).to_excel(writer, index=False, sheet_name="year_summary")
        pd.DataFrame(source_rows, columns=SOURCE_FIELDS).to_excel(writer, index=False, sheet_name="sources")
        pd.DataFrame(build_coverage(records)).to_excel(writer, index=False, sheet_name="coverage")

    legacy_xlsx_path = output_dir / "major_risk_warning_dataset.xlsx"
    legacy_xlsx_error = ""
    try:
        shutil.copyfile(xlsx_path, legacy_xlsx_path)
    except PermissionError as exc:
        legacy_xlsx_error = str(exc)

    report_path = build_report(records, source_rows, metrics, official_warnings, report_dir)
    return {
        "records": len(records),
        "metrics": len(metrics),
        "official_policy_warnings": len(official_warnings),
        "sources": len(source_rows),
        "fetched_sources": sum(1 for row in source_rows if row["status"] == "fetched"),
        "output_dir": str(output_dir),
        "raw_dir": str(raw_dir),
        "xlsx_path": str(xlsx_path),
        "legacy_xlsx_path": str(legacy_xlsx_path),
        "legacy_xlsx_error": legacy_xlsx_error,
        "report_path": str(report_path),
    }


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--raw-dir", type=Path, default=DEFAULT_RAW_DIR)
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT_DIR)
    parser.add_argument("--report-dir", type=Path, default=DEFAULT_REPORT_DIR)
    parser.add_argument(
        "--skip-fetch",
        action="store_true",
        help="Reuse the existing processed source index and rebuild derived tables from local raw/text files.",
    )
    args = parser.parse_args(argv)
    result = build_dataset(args.raw_dir, args.output_dir, args.report_dir, skip_fetch=args.skip_fetch)
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
